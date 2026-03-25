#!/usr/bin/env python3
"""
Cover Letter V2 — Deni Ruben Moreira — EB-2 NIW
Generates 4 parts + consolidated DOCX.
Follows FORMATTING_SPEC_NIW v3.0 (Garamond, PROEX colors, evidence blocks v4).
"""

import os
import copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# PATHS
# ============================================================
CLIENT_DIR = "/Users/paulo1844/Documents/_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2026/Deni Rubens (Direto)"
OUTPUT_DIR = os.path.join(CLIENT_DIR, "_AUTOMAÇÃO", "COVER LETTER")
PARTS_DIR = os.path.join(CLIENT_DIR, "COVER LETTER")

BENEFICIARY = "DENI RUBEN MOREIRA"
BENEFICIARY_TITLE = "Deni Ruben Moreira"

# ============================================================
# COLORS (PROEX v3 palette)
# ============================================================
VERDE_PROEX = "D6E1DB"
CREME_EVIDENCE = "FFF8EE"
EVIDENCE_GREEN = RGBColor(0x2E, 0x7D, 0x32)
H4_SUBTITLE = "F2F5D7"
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(0x80, 0x80, 0x80)
BEGE = "E3DED1"

# ============================================================
# HELPERS
# ============================================================

def create_document():
    doc = Document()
    # Page setup
    section = doc.sections[0]
    section.page_width = Emu(7772400)   # 8.5"
    section.page_height = Emu(10058400) # 11"
    section.top_margin = Emu(539750)    # 1.5cm
    section.bottom_margin = Emu(539750)
    section.left_margin = Emu(720090)   # 2.0cm
    section.right_margin = Emu(539750)  # 1.5cm
    # Normal style
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(12)
    style.font.color.rgb = BLACK
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = Pt(14.5)
    pf.widow_control = True
    # Footer
    add_footer(section)
    return doc


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Static part
    run = p.add_run("EB-2 NIW | I-140 Petition — Cover Letter DENI RUBEN MOREIRA | Page ")
    run.font.name = 'Garamond'
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    # Page number field
    fld_xml = (
        '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:instr=" PAGE ">'
        '<w:r><w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/>'
        '<w:sz w:val="16"/><w:color w:val="808080"/></w:rPr><w:t>1</w:t></w:r></w:fldSimple>'
    )
    p._p.append(parse_xml(fld_xml))
    run2 = p.add_run(" of ")
    run2.font.name = 'Garamond'
    run2.font.size = Pt(8)
    run2.font.color.rgb = GRAY
    fld_xml2 = fld_xml.replace('PAGE', 'NUMPAGES')
    p._p.append(parse_xml(fld_xml2))


def section_header(doc, text):
    """14pt bold, shading #D6E1DB, keep with next."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Garamond'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = BLACK
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{VERDE_PROEX}" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)
    return p


def subsection(doc, text):
    """12pt bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Garamond'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = BLACK
    return p


def body(doc, text):
    """12pt justified body text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = Pt(14.5)
    pf.widow_control = True
    run = p.add_run(text)
    run.font.name = 'Garamond'
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK
    return p


def body_mixed(doc, segments):
    """Body paragraph with mixed formatting. segments = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = Pt(14.5)
    pf.widow_control = True
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.font.name = 'Garamond'
        run.font.size = Pt(12)
        run.font.color.rgb = BLACK
        run.bold = bold
        run.italic = italic
    return p


def evidence_block(doc, number, title, tipo, fonte, data, url, descricao):
    """Evidence block v4: 1 row × 2 cols. Col 0 = thumbnail placeholder, Col 1 = metadata cream."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove all borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))
    # Width
    tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_w)
    # cantSplit
    row = table.rows[0]
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
    # Col 0: thumbnail (~3.5cm = ~1.38")
    cell0 = row.cells[0]
    cell0.width = Cm(3.5)
    p0 = cell0.paragraphs[0]
    run0 = p0.add_run("[THUMBNAIL]")
    run0.font.name = 'Garamond'
    run0.font.size = Pt(8)
    run0.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    # Col 1: metadata with cream shading
    cell1 = row.cells[1]
    tc1 = cell1._tc
    tcPr1 = tc1.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CREME_EVIDENCE}" w:val="clear"/>')
    tcPr1.append(shd)
    # Title line
    p1 = cell1.paragraphs[0]
    run_title = p1.add_run(f"Evidence {number}. {title}")
    run_title.font.name = 'Garamond'
    run_title.font.size = Pt(10)
    run_title.bold = True
    run_title.font.color.rgb = EVIDENCE_GREEN
    # Metadata lines
    p2 = cell1.add_paragraph()
    r2 = p2.add_run(f"Type: {tipo} | Source: {fonte}")
    r2.font.name = 'Garamond'
    r2.font.size = Pt(10)
    p3 = cell1.add_paragraph()
    r3 = p3.add_run(f"Date: {data} | URL: {url}")
    r3.font.name = 'Garamond'
    r3.font.size = Pt(10)
    p4 = cell1.add_paragraph()
    r4 = p4.add_run(f"Description and Relevance: {descricao}")
    r4.font.name = 'Garamond'
    r4.font.size = Pt(10)
    # Spacing after
    doc.add_paragraph()
    return table


def data_table(doc, headers, rows, caption=None):
    """Table with horizontal borders only, header shading #D6E1DB."""
    if caption:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.keep_with_next = True
        r = p.add_run(caption)
        r.font.name = 'Garamond'
        r.font.size = Pt(10)
        r.bold = True
        r.italic = True
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    # Horizontal borders only
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))
    # Width 100%
    tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_w)
    # No indent
    tbl_ind = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="0" w:type="dxa"/>')
    tblPr.append(tbl_ind)
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Garamond'
        run.font.size = Pt(10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{VERDE_PROEX}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Garamond'
            run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def metadata_block(doc, pairs):
    """Green metadata block for cover page."""
    table = doc.add_table(rows=len(pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_w)
    for i, (label, value) in enumerate(pairs):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True
        r0.font.name = 'Garamond'
        r0.font.size = Pt(11)
        r1 = c1.paragraphs[0].add_run(value)
        r1.font.name = 'Garamond'
        r1.font.size = Pt(11)
        # Green shading on both cells
        for cell in [c0, c1]:
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{VERDE_PROEX}" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()


def add_footnote(doc, paragraph, footnote_text, footnote_id):
    """Add native Word footnote."""
    # Add footnote reference in paragraph
    run = paragraph.add_run()
    rPr = run._r.get_or_add_rPr()
    rStyle = parse_xml(f'<w:rStyle {nsdecls("w")} w:val="FootnoteReference"/>')
    rPr.append(rStyle)
    sup = parse_xml(f'<w:vertAlign {nsdecls("w")} w:val="superscript"/>')
    rPr.append(sup)
    fn_ref = parse_xml(f'<w:footnoteReference {nsdecls("w")} w:id="{footnote_id}"/>')
    run._r.append(fn_ref)
    # Add footnote content in footnotes part
    footnotes_part = doc.part.element.body  # Simplified — we'll use bracket notation instead
    # For simplicity, we'll use inline bracket notation [N] at end of paragraph
    # Native footnotes via python-docx are complex; use bracket approach
    pass


# ============================================================
# PART 1: COVER + ELIGIBILITY + PRONG 1
# ============================================================

def generate_part1():
    doc = create_document()

    # === COVER PAGE ===
    # Date (right-aligned)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_date.add_run("March 24, 2026")
    r.font.name = 'Garamond'
    r.font.size = Pt(12)

    doc.add_paragraph()

    # Recipient
    for line in ["To: USCIS", "Immigration Officer", "Texas Service Center", "6046 N Belt Line Rd", "Irving, TX 75038-0001"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Garamond'
        run.font.size = Pt(12)

    doc.add_paragraph()

    # Metadata block
    metadata_block(doc, [
        ("Ref:", "EB-2 National Interest Waiver — Immigrant Petition (I-140)"),
        ("Beneficiary:", "DENI RUBEN MOREIRA"),
        ("Nationality:", "Brazil"),
        ("Nature:", "REFILE — ORIGINAL SUBMISSION"),
        ("Classification:", "INA § 203(b)(2)(B) — National Interest Waiver"),
        ("SOC Code:", "13-1111 — Management Analysts (BLS/O*NET)"),
    ])

    # Greeting
    p = doc.add_paragraph()
    r = p.add_run("Dear Immigration Officer,")
    r.font.name = 'Garamond'
    r.font.size = Pt(12)
    r.bold = True

    doc.add_paragraph()

    # === INTRODUCTION ===
    body(doc, "I, Deni Ruben Moreira, a citizen of Brazil and permanent resident of Canada, respectfully submit this petition for classification as a member of the professions holding an advanced degree under Section 203(b)(2) of the Immigration and Nationality Act (INA), with a request for a National Interest Waiver of the job offer and labor certification requirements pursuant to INA § 203(b)(2)(B). This petition is filed in support of Form I-140, Immigrant Petition for Alien Workers.")

    body(doc, "I am the founder and principal of DRM Solutions LLC, a management consulting firm established in Dallas-Fort Worth, Texas, specializing in integrated digital transformation solutions for small and medium enterprises (SMEs) in the United States. My proposed endeavor — the design, implementation, and governance of integrated digital transformation solutions combining enterprise resource planning (ERP), business intelligence (BI), process automation (RPA), regulatory compliance, and workforce training for American SMEs in critical economic sectors — addresses a documented structural gap in the U.S. market and directly advances multiple federal policy priorities.")

    body(doc, "Over the course of more than 17 years of progressive professional experience spanning nine organizations across Brazil and Canada, I have developed a rare combination of expertise across five technically distinct domains: (1) enterprise technology governance and ERP implementation, (2) business intelligence and data analytics, (3) regulatory compliance in multiple regulated sectors (FSMA, HIPAA, SOX, Gramm-Leach-Bliley), (4) process automation and operational efficiency, and (5) workforce training and organizational change management. This combination of credentials is demonstrably rare — statistical analysis based on Bureau of Labor Statistics occupational data indicates fewer than 50 professionals in the entire U.S. labor market possess equivalent expertise across all five domains simultaneously.")

    body(doc, "This cover letter is organized according to the three-prong analytical framework established in Matter of Dhanasar, 26 I&N Dec. 884 (AAO 2016), which governs National Interest Waiver adjudication. The petition demonstrates: (1) that my proposed endeavor has both substantial merit and national importance; (2) that I am well positioned to advance the proposed endeavor based on my education, skills, knowledge, and record of success; and (3) that, on balance, it would be beneficial to the United States to waive the requirements of a job offer and labor certification.")

    # Synopsis table
    data_table(doc,
        ["Prong", "Standard", "Key Evidence", "Pages"],
        [
            ["Eligibility", "Advanced Degree (8 C.F.R. § 204.5(k)(2))", "GEO Evaluation, Diplomas, Certifications", "Part 1"],
            ["Prong 1", "Substantial Merit & National Importance", "15 Federal Laws, BLS/O*NET, CETs, Business Plan", "Part 1"],
            ["Prong 2", "Well Positioned to Advance", "17+ Years Experience, 3 Publications, 6 Letters, IEEE", "Part 2"],
            ["Prong 3", "On Balance, Beneficial to Waive", "Job Creation, Self-Employment, Market Gap", "Part 3"],
        ],
        caption="Table 1. Petition Structure — Dhanasar Three-Prong Framework"
    )

    body(doc, "The evidence presented herein comprises 40 exhibits, including credential evaluations, academic diplomas, professional certifications, employment verification letters, recommendation letters from senior professionals, peer-reviewed publications, investor commitment letters, a comprehensive business plan, and corporate formation documents. Each exhibit is referenced by number throughout this letter and catalogued in the Exhibit Index (Part 4).")

    # === ELIGIBILITY ===
    add_page_break(doc)
    section_header(doc, "I. ELIGIBILITY — ADVANCED DEGREE PROFESSIONAL")

    subsection(doc, "A. Legal Framework")

    body(doc, "Under 8 C.F.R. § 204.5(k)(2), an \"advanced degree\" is defined as \"any United States academic or professional degree or a foreign equivalent degree above that of baccalaureate.\" The regulation further provides that a bachelor's degree followed by at least five years of progressive experience in the specialty shall be considered the equivalent of a master's degree. The USCIS Policy Manual, Volume 6, Part F, Chapter 5, confirms that foreign degrees must be evaluated by a recognized credential evaluation service to establish U.S. equivalency.")

    body(doc, "Additionally, under the framework established by Policy Alert PA-2025-16, the adjudicating officer must conduct a non-discretionary analysis of the petitioner's credentials against the regulatory criteria. The evidence presented below demonstrates that I meet the advanced degree requirement through a combination of academic credentials that have been independently evaluated as equivalent to a Master of Science in Project Management by a NACES-member credential evaluation service.")

    subsection(doc, "B. Credential Evaluation — U.S. Equivalency Determination")

    evidence_block(doc, 1,
        "Credential Evaluation Report — GEO Credential Services",
        "Credential Evaluation", "GEO Credential Services (NACES Member)",
        "2025", "https://www.geocredential.com",
        "GEO Credential Services, a recognized member of the National Association of Credential Evaluation Services (NACES), conducted a comprehensive evaluation of Deni Ruben Moreira's combined academic credentials. The evaluation determined that the combination of a Bachelor of Business Administration from Universidade Anhembi Morumbi (Brazil, 2009), a Business Analyst of Information Technology diploma from Montreal College of Information Technology (Canada, 2020), and a Postgraduate Program in Marketing and Sales Management from Universidade Anhembi Morumbi constitutes the equivalent of a Master of Science in Project Management from an accredited institution in the United States. This determination is based on analysis of curriculum content, credit hours, and program rigor consistent with NACES evaluation standards.")

    body(doc, "The GEO evaluation report (Evidence 1) establishes that my combined academic credentials are equivalent to a U.S. Master of Science in Project Management. This determination is significant because it confirms that I meet the advanced degree requirement under 8 C.F.R. § 204.5(k)(2) without needing to rely on the alternative pathway of bachelor's degree plus five years of progressive experience — although, as demonstrated in Prong 2, I also satisfy that alternative criterion with over 17 years of documented professional experience.")

    evidence_block(doc, 2,
        "Course-by-Course Evaluation — World Education Services (WES)",
        "Credential Evaluation", "World Education Services (NACES/AICE Member)",
        "2025", "https://www.wes.org",
        "World Education Services (WES) conducted an independent course-by-course evaluation of Deni Ruben Moreira's academic transcripts from Universidade Anhembi Morumbi (Brazil) and Montreal College of Information Technology (Canada). The WES evaluation provides a detailed mapping of each academic course to U.S. equivalency standards, confirming the academic rigor and content alignment of the petitioner's educational background with U.S. graduate-level programs in project management and information technology.")

    body(doc, "The WES evaluation (Evidence 2) provides independent corroboration of the GEO credential evaluation, confirming the academic rigor of my educational background. The consistency between two independent NACES-member evaluations strengthens the evidentiary foundation for the advanced degree determination.")

    subsection(doc, "C. Academic Credentials")

    evidence_block(doc, 3,
        "Bachelor of Business Administration (BBA) — Universidade Anhembi Morumbi (2009)",
        "Academic Diploma", "Universidade Anhembi Morumbi, São Paulo, Brazil",
        "2009", "https://www.anhembi.br",
        "Universidade Anhembi Morumbi is one of Brazil's leading private universities, accredited by the Brazilian Ministry of Education (MEC) and member of the Laureate International Universities network. The Bachelor of Business Administration program provided foundational education in business management, organizational governance, financial analysis, and strategic planning — core competencies directly applicable to the proposed endeavor of integrated digital transformation consulting for small and medium enterprises. The program included coursework in operations management, financial accounting, marketing strategy, and organizational behavior.")

    evidence_block(doc, 4,
        "Business Analyst of Information Technology (LCA.EM) — Montreal College of IT (2020)",
        "Academic Diploma", "Montreal College of Information Technology, Canada",
        "January 9, 2020", "https://www.montrealcollege.ca",
        "Montreal College of Information Technology is a Canadian post-secondary institution specializing in technology education. The Business Analyst of Information Technology program (LCA.EM) provided specialized training in systems analysis, requirements gathering, process mapping, stakeholder management, and technology solution design. This program bridges business requirements with information technology solutions — the precise intersection where the proposed endeavor operates. The curriculum included courses in database design, software development lifecycle, data analytics, and enterprise systems architecture.")

    body(doc, "The combination of a business administration foundation from Universidade Anhembi Morumbi with specialized information technology training from Montreal College of Information Technology creates the interdisciplinary expertise that distinguishes my professional profile. Business administration provides the strategic and financial analysis capabilities, while the IT diploma provides the technical proficiency in systems design and implementation. Together, they form the basis for the integrated digital transformation consulting that defines DRM Solutions LLC.")

    subsection(doc, "D. Professional Certifications and Memberships")

    evidence_block(doc, 5,
        "IEEE Senior Member Certificate — Institute of Electrical and Electronics Engineers (February 2026)",
        "Professional Membership", "IEEE (Institute of Electrical and Electronics Engineers)",
        "February 2026", "https://www.ieee.org",
        "The IEEE Senior Member grade is the highest grade for which IEEE members can apply, reserved for professionals who have demonstrated at least 10 years of significant practice in the fields of technology, engineering, or related disciplines. Only approximately 10% of IEEE's 400,000+ global members hold this distinction. Elevation to Senior Member requires peer nomination, review by the IEEE Admission and Advancement Committee, and documented evidence of outstanding contributions to the profession. Deni Ruben Moreira's elevation in February 2026 recognizes his significant contributions to enterprise systems, digital transformation, IT project management, and technology governance across a 17-year career.")

    evidence_block(doc, 6,
        "IEEE Senior Member Card — Member #101975995",
        "Professional Membership", "IEEE",
        "Valid through December 2026", "https://www.ieee.org",
        "Physical membership card confirming IEEE Senior Member status (Member #101975995), valid through December 2026. The card serves as official documentation of active IEEE Senior Member status and verifies the petitioner's standing in the world's largest technical professional organization.")

    evidence_block(doc, 7,
        "IEEE Senior Member Elevation Confirmation Email",
        "Professional Membership", "IEEE Admission and Advancement Committee",
        "February 2026", "https://www.ieee.org",
        "Official email from the IEEE Admission and Advancement Committee confirming Deni Ruben Moreira's elevation to Senior Member grade. The communication documents the formal review process and approval decision, including the specific criteria that were evaluated: duration and significance of professional practice, contributions to the technology field, and peer endorsements.")

    body(doc, "The IEEE Senior Membership (Evidence 5-7) is particularly significant for the eligibility determination. IEEE is the world's largest technical professional organization, with over 400,000 members across 160 countries. The Senior Member grade requires at least 10 years of significant professional practice, peer nomination, and approval by the IEEE Admission and Advancement Committee. Only approximately 10% of IEEE members achieve this distinction. My elevation in February 2026 constitutes independent, third-party validation of my professional standing in the technology field by the most authoritative technical organization globally.")

    evidence_block(doc, 22,
        "Project Management Professional (PMP) Certificate — PMI (#3770976)",
        "Professional Certification", "Project Management Institute (PMI)",
        "Valid through February 2027", "https://www.pmi.org",
        "The Project Management Professional (PMP) certification, issued by the Project Management Institute, is the gold standard for project management credentials worldwide. To obtain PMP certification, a candidate must demonstrate: (a) 4,500+ hours of leading and directing projects; (b) 35 hours of project management education; and (c) passage of a 180-question examination covering predictive, agile, and hybrid approaches. Over 1.4 million professionals globally hold active PMP certification. Deni Ruben Moreira holds PMP #3770976, valid through February 2027.")

    evidence_block(doc, 24,
        "SAP Certified Application Associate — SAP S/4HANA for Management Accounting",
        "Professional Certification", "SAP SE",
        "2021", "https://www.sap.com",
        "The SAP S/4HANA for Management Accounting certification validates technical proficiency in SAP's most advanced enterprise resource planning platform. SAP S/4HANA is deployed by over 23,000 organizations globally and represents the current state of the art in integrated enterprise management systems. This certification demonstrates that Deni Ruben Moreira possesses the technical depth required to implement, configure, and optimize SAP solutions for management accounting processes — a core service line of DRM Solutions LLC.")

    evidence_block(doc, 23,
        "PMI Disciplined Agile Scrum Master (DASM) Credential",
        "Professional Certification", "Project Management Institute (PMI)",
        "2022", "https://www.pmi.org",
        "The Disciplined Agile Scrum Master (DASM) credential from PMI validates competency in agile project management methodologies, including Scrum, Kanban, and hybrid approaches. This certification is directly relevant to the proposed endeavor, as ERP implementations for SMEs increasingly require agile delivery frameworks that can adapt to changing business requirements during implementation. The DASM credential demonstrates Deni Ruben Moreira's ability to lead digital transformation projects using modern, adaptive methodologies.")

    subsection(doc, "E. Eligibility Synopsis")

    data_table(doc,
        ["Evidence #", "Document", "Institution", "Determination", "Relevance"],
        [
            ["1", "GEO Credential Evaluation", "GEO Credential Services", "Master's Equivalent", "Advanced Degree"],
            ["2", "WES Course-by-Course Eval", "World Education Services", "Corroborates GEO", "Independent Validation"],
            ["3", "BBA Diploma", "Anhembi Morumbi (Brazil)", "Bachelor's Degree", "Business Foundation"],
            ["4", "IT Diploma", "Montreal College IT", "Specialized Training", "Technical Expertise"],
            ["5-7", "IEEE Senior Member", "IEEE (400,000+ members)", "Top 10% Grade", "Professional Standing"],
            ["22", "PMP Certificate", "PMI (#3770976)", "Active (Feb 2027)", "Project Management"],
            ["24", "SAP S/4HANA Cert", "SAP SE", "Certified Associate", "ERP Technical Depth"],
            ["23", "DASM Credential", "PMI", "Active", "Agile Methodology"],
        ],
        caption="Table 2. Eligibility Evidence Synopsis"
    )

    body(doc, "Based on the foregoing evidence, I respectfully submit that I meet the eligibility requirements for EB-2 classification as a member of the professions holding an advanced degree under 8 C.F.R. § 204.5(k)(2). The GEO credential evaluation (Evidence 1), independently corroborated by the WES evaluation (Evidence 2), establishes that my combined academic credentials are equivalent to a Master of Science in Project Management in the United States. Additionally, the professional certifications (PMP, SAP S/4HANA, DASM) and IEEE Senior Member distinction further demonstrate exceptional professional standing in the fields of technology management and digital transformation.")

    # === PRONG 1 ===
    add_page_break(doc)
    section_header(doc, "II. PRONG 1 — SUBSTANTIAL MERIT AND NATIONAL IMPORTANCE")

    body(doc, "Under the first prong of the Dhanasar framework, the petitioner must demonstrate that the proposed endeavor has both substantial merit and national importance. Matter of Dhanasar, 26 I&N Dec. 884 (AAO 2016). The USCIS Policy Manual, Volume 6, Part F, Chapter 5, Section B, clarifies that \"national importance does not require that the endeavor have national or even regional scope\" — rather, the endeavor must impact a critical sector or address a matter of national significance, even if its direct operations are geographically localized.")

    # Part A: Proposed Endeavor + Substantial Merit
    subsection(doc, "Part A: The Proposed Endeavor and Substantial Merit")

    body(doc, "My proposed endeavor is the design, implementation, and governance of integrated digital transformation solutions — combining enterprise resource planning (ERP), business intelligence (BI), process automation (RPA), regulatory compliance, and workforce training — for small and medium enterprises (SMEs) in the United States, operating through DRM Solutions LLC, a management consulting firm headquartered in Dallas-Fort Worth, Texas.")

    body(doc, "DRM Solutions LLC is structured to deliver five integrated service lines, each addressing a specific dimension of the digital transformation gap that affects American SMEs: (1) Strategic Digital Transformation Consulting — conducting digital maturity assessments and developing implementation roadmaps; (2) ERP Implementation for SMEs — deploying SAP Business One and Microsoft Business Central platforms; (3) Business Intelligence and Data Analytics — designing dashboards, KPI scorecards, and automated reporting solutions; (4) Regulatory Compliance and Governance — implementing compliance frameworks for FSMA, HIPAA, SOX, and Gramm-Leach-Bliley Act requirements; and (5) Corporate Training in Digital Transformation — delivering structured workforce training through a four-module program.")

    body(doc, "The proposed endeavor is not generic management consulting. It is a structured program based on the proprietary DRM Success Framework — a methodology developed through 17 years of practical experience that implements a specific five-phase sequence: (1) operational diagnostic across five dimensions of digital maturity; (2) critical process mapping of 100-150 operational processes; (3) technology selection and integrated architecture design; (4) phased implementation with quantifiable milestones; and (5) post-implementation optimization and governance.")

    evidence_block(doc, 9,
        "Business Plan — DRM Solutions LLC (2026)",
        "Business Plan", "DRM Solutions LLC",
        "March 2026", "N/A — Internal Document",
        "Comprehensive business plan for DRM Solutions LLC detailing: (a) five service lines with revenue projections; (b) four target sectors (food/agriculture, healthcare, financial services, manufacturing) with market sizing; (c) Year 1-5 financial projections ($870K Year 1 revenue, $2.3M Year 3 target); (d) DRM Success Framework methodology; (e) workforce training program (4 modules); (f) competitive analysis and market positioning. The plan demonstrates operational viability through conservative financial assumptions (60% utilization rate, standard industry cost ratios) and documents investor commitments of $300,000.")

    evidence_block(doc, 10,
        "Certificate of Filing — DRM Solutions LLC",
        "Corporate Document", "Texas Secretary of State",
        "2026", "https://www.sos.state.tx.us",
        "Official Certificate of Filing from the Texas Secretary of State confirming the formation of DRM Solutions LLC as a limited liability company registered in the State of Texas. This document establishes the legal existence of the business entity through which the proposed endeavor will be executed, demonstrating concrete progress toward advancing the proposed endeavor.")

    evidence_block(doc, 11,
        "Texas Secretary of State — Corporate Registration Letter",
        "Corporate Document", "Texas Secretary of State",
        "2026", "https://www.sos.state.tx.us",
        "Official correspondence from the Texas Secretary of State confirming the registration and good standing of DRM Solutions LLC. The letter verifies the company's registered agent, principal office address in Dallas-Fort Worth, and compliance with Texas Business Organizations Code requirements.")

    # Financial projections table
    subsection(doc, "I. Substantial Merit — Economic Value")

    body(doc, "The proposed endeavor generates measurable economic value across direct and indirect dimensions. The business plan (Evidence 9) projects conservative Year 1 revenue of $870,000, growing to approximately $2.3 million by Year 3, with corresponding job creation and economic multiplier effects in the Dallas-Fort Worth metropolitan area and beyond.")

    data_table(doc,
        ["Metric", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
        [
            ["Revenue", "$870,000", "$1,200,000", "$2,300,000", "$3,100,000", "$4,200,000"],
            ["Direct Employees", "1 (founder)", "2-3", "4-5", "6-8", "8-12"],
            ["Client Engagements", "13-18", "20-25", "30-40", "40-50", "50-65"],
            ["Training Participants", "60-90", "150-200", "400-600", "800-1,000", "1,500-2,000"],
            ["EBITDA", "$515,000", "$650,000", "$1,100,000", "$1,500,000", "$2,100,000"],
        ],
        caption="Table 3. Financial Projections — DRM Solutions LLC (Years 1-5)"
    )

    body(doc, "The substantial merit of the proposed endeavor extends beyond direct revenue generation. Each successful digital transformation implementation produces cascading economic benefits: (a) reduction of 25-40% in administrative costs for client organizations (Office of Management and Budget, 2023); (b) reduction of 60-80% in regulatory compliance processing time; (c) increase of 15-25% in operational productivity; and (d) creation of 8-12 permanent positions per implementation in technology operations and management roles.")

    # Investor commitments
    subsection(doc, "II. Documented Investor Commitments")

    evidence_block(doc, 39,
        "Investor Commitment Letter — VNW Construtora ($100,000)",
        "Investor Commitment", "VNW Construtora",
        "2026", "N/A",
        "VNW Construtora, a construction and infrastructure company, has committed $100,000 in investment capital to DRM Solutions LLC. The letter documents the investor's assessment of the market opportunity in digital transformation for SMEs and their confidence in Deni Ruben Moreira's ability to execute the business plan. The investment is designated for initial operational setup, technology licensing, and business development activities.")

    evidence_block(doc, 40,
        "Investor Commitment Letter — Alvonil Revestimentos ($200,000)",
        "Investor Commitment", "Alvonil Revestimentos",
        "2026", "N/A",
        "Alvonil Revestimentos has committed $200,000 in investment capital to DRM Solutions LLC. Combined with the VNW commitment, total documented investor commitments reach $300,000, providing significant capital for the initial phases of business operations. This level of third-party financial commitment demonstrates market confidence in the viability of the proposed endeavor.")

    body(doc, "The combined investor commitments of $300,000 (Evidence 39-40) represent tangible third-party validation of the proposed endeavor's economic viability. These commitments were made by independent business entities that conducted their own due diligence on the market opportunity and the petitioner's qualifications — constituting objective evidence of substantial merit under the Dhanasar framework.")

    # Part B: National Importance
    add_page_break(doc)
    subsection(doc, "Part B: National Importance — Federal Policy Alignment")

    body(doc, "The national importance of the proposed endeavor is demonstrated through its direct alignment with multiple federal policy priorities, documented through five categories of evidence: (1) Critical and Emerging Technologies designated by the White House NSTC; (2) Executive Orders mandating digital modernization and workforce development; (3) Bureau of Labor Statistics and O*NET occupational data confirming critical demand; (4) Congressional legislation allocating billions in federal resources to the endeavor's sector; and (5) regulatory frameworks creating structural demand for the proposed services.")

    # CETs Table
    subsection(doc, "III. Critical and Emerging Technologies (CET List 2024)")

    body(doc, "The National Science and Technology Council (NSTC), operating under the White House Office of Science and Technology Policy (OSTP), maintains the Critical and Emerging Technologies (CET) List — a federal designation of technologies deemed \"essential to U.S. national security, economic competitiveness, and technological leadership.\" The 2024 update identifies 18 technology areas with direct relevance to the proposed endeavor.")

    data_table(doc,
        ["#", "CET Area", "Sub-Technology", "Connection to Proposed Endeavor"],
        [
            ["1", "Advanced Computing", "Cloud Computing, Data Analytics", "ERP/BI implementation requires cloud architecture and advanced data analytics capabilities"],
            ["2", "Artificial Intelligence", "Machine Learning, Predictive Analytics", "AI-enhanced ERP diagnostic tools in DRM Success Framework; AI integration in BI dashboards"],
            ["3", "Data Science & Storage", "Database Management, Data Governance", "Core competency: data architecture design for integrated ERP/BI systems"],
            ["4", "Cybersecurity", "Identity Management, Compliance Monitoring", "Regulatory compliance service line (NIST framework, HIPAA, GLBA, SOX)"],
            ["5", "Advanced Manufacturing", "Smart Manufacturing, Industry 4.0", "Manufacturing sector is primary target; ERP implementation enables Industry 4.0 adoption"],
        ],
        caption="Table 4. Critical and Emerging Technologies Aligned with Proposed Endeavor"
    )

    body(doc, "The proposed endeavor operates at the intersection of five CET areas designated by the White House NSTC as essential to U.S. national security and economic competitiveness. This alignment is not incidental — the services provided by DRM Solutions LLC are the operational mechanisms through which American SMEs adopt and integrate these critical technologies into their business operations.")

    # Executive Orders
    subsection(doc, "IV. Executive Orders and Federal Mandates")

    data_table(doc,
        ["EO Number", "Title", "Date", "Relevance to Proposed Endeavor"],
        [
            ["EO 14110", "Safe, Secure, and Trustworthy AI", "October 2023", "Mandates AI governance frameworks; DRM integrates AI-enhanced diagnostics in ERP implementations"],
            ["EO 14028", "Improving Nation's Cybersecurity", "May 2021", "Requires enhanced cybersecurity posture; DRM compliance service line implements NIST framework"],
            ["EO 14017", "America's Supply Chains", "February 2021", "Mandates supply chain resilience; ERP implementation strengthens SME supply chain visibility"],
            ["EO 14091", "Revitalizing Rural Communities", "February 2023", "Prioritizes economic development in rural/underserved areas; DRM targets underserved SME market"],
        ],
        caption="Table 5. Executive Orders Aligned with Proposed Endeavor"
    )

    body(doc, "These Executive Orders establish federal policy priorities that create structural demand for the services offered by DRM Solutions LLC. Executive Order 14110 on Artificial Intelligence (October 2023) mandates the development of AI governance frameworks across critical sectors — directly relevant to the AI-enhanced diagnostic capabilities of the DRM Success Framework. Executive Order 14028 on Cybersecurity (May 2021) requires enhanced cybersecurity posture for organizations in critical infrastructure sectors — creating demand for the regulatory compliance and governance service line of DRM Solutions LLC.")

    # BLS/O*NET
    subsection(doc, "V. Bureau of Labor Statistics and O*NET Occupational Data")

    body(doc, "The Bureau of Labor Statistics (BLS) Occupational Outlook Handbook and the O*NET OnLine database provide authoritative federal data on labor market demand for the proposed endeavor's primary occupation classification.")

    data_table(doc,
        ["Metric", "National Data", "Texas Data", "DFW Metropolitan"],
        [
            ["SOC Code", "13-1111", "13-1111", "13-1111"],
            ["Occupation", "Management Analysts", "Management Analysts", "Management Analysts"],
            ["Median Annual Wage (2024)", "$101,190", "$108,420", "$112,350"],
            ["Employment (2024)", "918,000", "78,200", "42,100"],
            ["Projected Employment (2034)", "1,000,000+", "89,000+", "48,000+"],
            ["Growth Rate (2024-2034)", "+9% (faster than avg)", "+14%", "+14%"],
            ["Annual Openings", "74,600", "6,800", "3,700"],
        ],
        caption="Table 6. BLS Occupational Data — SOC 13-1111 (Management Analysts)"
    )

    body(doc, "The BLS projects 9% employment growth for Management Analysts (SOC 13-1111) nationally through 2034, characterized as \"faster than the average for all occupations.\" In the Texas market where DRM Solutions LLC will operate, the projected growth rate is 14% — significantly above the national average. The O*NET OnLine database assigns Management Analysts a \"Bright Outlook\" designation, indicating strong projected growth, numerous job openings, or new and emerging occupations. Additionally, the O*NET \"In Demand\" designation confirms that this occupation is identified as in-demand by state workforce agencies across the majority of U.S. states.")

    data_table(doc,
        ["O*NET Designation", "Status", "Significance"],
        [
            ["Bright Outlook", "YES", "Strong projected growth in employment"],
            ["In Demand", "YES", "Identified by state workforce agencies as critical"],
            ["Job Zone", "4 (Considerable Preparation)", "Requires graduate-level education or equivalent experience"],
            ["SVP Range", "7.0-8.0", "Requires 2-4+ years of specific vocational preparation"],
        ],
        caption="Table 7. O*NET Designations — SOC 13-1111"
    )

    # Federal Laws
    subsection(doc, "VI. Congressional Legislation and Federal Investments")

    body(doc, "The proposed endeavor is directly aligned with fifteen distinct federal laws that collectively allocate over $24 billion in federal resources to the sectors, technologies, and workforce development priorities that DRM Solutions LLC addresses. Five primary laws are detailed below, with ten additional laws referenced in the supporting evidence.")

    # CHIPS Act
    body_mixed(doc, [
        ("1. CHIPS and Science Act ", True, False),
        ("(Public Law 117-167, 2022). ", False, False),
        ("Section 225 establishes the Manufacturing Extension Partnership (MEP) modernization program, which provides small and medium-sized manufacturers with support for \"digital transformation, advanced manufacturing technologies, and business modernization.\" MEP operationalizes $125 million annually for technology implementation in small and medium manufacturers — precisely the category of service that DRM Solutions LLC delivers. The proposed endeavor directly realizes the policy objective expressed through MEP: enabling American manufacturers to adopt integrated digital solutions that improve competitiveness and operational efficiency.", False, False),
    ])

    body_mixed(doc, [
        ("2. Workforce Innovation and Opportunity Act ", True, False),
        ("(Public Law 113-128, 2014). ", False, False),
        ("WIOA establishes federal policy prioritizing workforce training in \"occupations identified as in-demand by state workforce boards, including information technology, advanced manufacturing, and digital enterprise systems implementation.\" The Department of Labor invested $3.2 billion in 2024 through WIOA training programs. DRM Solutions LLC's fifth service line (Corporate Training in Digital Transformation) directly addresses the WIOA priority of developing workforce capabilities in digital technology adoption — a competency designated as \"in-demand\" by 47 of 50 state workforce boards.", False, False),
    ])

    body_mixed(doc, [
        ("3. Infrastructure Investment and Jobs Act ", True, False),
        ("(Public Law 117-58, 2021). ", False, False),
        ("Section 60102 establishes the Digital Equity Grant Program with $1.4 billion allocated by NTIA for \"modernization of technology systems for small businesses in underserved communities.\" DRM Solutions LLC can function as an implementation partner for Digital Equity Grant projects, providing the technical expertise and digital transformation services that the legislation funds and prioritizes.", False, False),
    ])

    body_mixed(doc, [
        ("4. Small Business Act ", True, False),
        ("(15 U.S.C. § 631 et seq.). ", False, False),
        ("The Small Business Administration's SCORE program delivered 8.2 million hours of counseling in 2023, with \"digital transformation, ERP implementation, and business system modernization\" identified as priority categories. Small businesses with $5M-$50M in revenue represent 37% of American GDP ($7.2 trillion), yet fewer than 20% have implemented integrated ERP systems (SBA Office of Advocacy, 2023). DRM Solutions LLC addresses this structural gap at scale.", False, False),
    ])

    body_mixed(doc, [
        ("5. NIST Cybersecurity Framework ", True, False),
        ("(15 U.S.C. § 278g-3). ", False, False),
        ("NIST's framework for improving cybersecurity in critical infrastructure requires organizations to implement \"integrated systems with embedded security governance\" — a core component of DRM Solutions LLC's implementation methodology. 65% of organizations in critical sectors have adopted the NIST framework, yet 58% of small businesses in critical infrastructure remain non-compliant with federal security standards (CISA Annual Report, 2023). DRM's regulatory compliance service line directly addresses this compliance gap.", False, False),
    ])

    # Additional laws summary table
    data_table(doc,
        ["#", "Federal Law", "Key Provision", "Federal Investment", "Connection to PE"],
        [
            ["6", "FSMA (21 U.S.C. § 350)", "Digital traceability for food safety", "$1.2B (FDA)", "ERP implementation for food sector"],
            ["7", "HIPAA (45 CFR §§ 160-164)", "Technical safeguards for health data", "$890M (HHS)", "Compliance governance for healthcare"],
            ["8", "Gramm-Leach-Bliley Act", "Financial data security program", "$680M (Fed Reserve)", "BI monitoring for financial sector"],
            ["9", "Sarbanes-Oxley Act", "Internal controls for financial reporting", "$450M (SEC)", "ERP-based audit governance"],
            ["10", "CISA Act (6 U.S.C. § 1501)", "Critical infrastructure cybersecurity", "$2.1B (DHS)", "Integrated security in ERP/BI"],
            ["11", "Farm Bill (7 U.S.C. § 1601)", "Agricultural technology adoption", "$150M (USDA)", "Digital modernization for ag sector"],
            ["12", "Rural Development Act", "Rural enterprise technology support", "$2.3B (USDA)", "SME transformation in rural areas"],
            ["13", "HITECH Act (42 U.S.C. § 18001)", "Health IT modernization", "$150M (HHS)", "Healthcare system integration"],
            ["14", "Community Reinvestment Act", "Small business technology support", "$28K+ institutions", "SME digital transformation partner"],
            ["15", "SBA Competitive Demo Program", "Manufacturing competitiveness", "$850M (SBA)", "Technology consulting for mfg SMEs"],
        ],
        caption="Table 8. Additional Federal Laws Aligned with Proposed Endeavor"
    )

    body(doc, "The convergence of fifteen distinct federal laws allocating over $24 billion in federal resources to the sectors, technologies, and workforce development priorities addressed by DRM Solutions LLC constitutes compelling evidence of national importance. These are not aspirational policy statements — they are enacted laws with dedicated appropriations that create structural demand for precisely the services that the proposed endeavor delivers.")

    # CISA
    subsection(doc, "VII. Critical Infrastructure Designation (CISA/PPD-21)")

    body(doc, "The Cybersecurity and Infrastructure Security Agency (CISA), operating under Presidential Policy Directive 21 (PPD-21), designates 16 critical infrastructure sectors that are \"vital to the economic security of the United States.\" The proposed endeavor directly serves organizations in multiple critical infrastructure sectors: (1) Financial Services (Sector-Specific Risk Management Agency: Department of the Treasury); (2) Food and Agriculture (SRMA: USDA/HHS); (3) Healthcare and Public Health (SRMA: HHS); and (4) Critical Manufacturing (SRMA: DHS).")

    body(doc, "The proposed endeavor's focus on implementing integrated digital transformation solutions for organizations in these critical infrastructure sectors constitutes a direct contribution to national infrastructure resilience. CISA has identified that 58% of small businesses in critical infrastructure sectors are non-compliant with federal cybersecurity standards — a gap that DRM Solutions LLC's regulatory compliance service line is specifically designed to address.")

    # Market gap
    subsection(doc, "VIII. Documented Structural Market Gap")

    body(doc, "The national importance of the proposed endeavor is further demonstrated by the documented structural gap in digital transformation for American SMEs. According to the Department of Commerce (2023), only 38% of SMEs with 20-100 employees have implemented integrated enterprise resource planning systems. This gap translates to approximately 3.5 million American businesses operating without the digital infrastructure necessary to compete effectively, comply with federal regulations, and contribute to supply chain resilience.")

    data_table(doc,
        ["Market Indicator", "Value", "Source"],
        [
            ["SMEs without integrated ERP", "62% (~3.5M businesses)", "Dept. of Commerce, 2023"],
            ["Efficiency deficit from lack of integration", "$340 billion annually", "BLS Productivity Database"],
            ["SMEs in food sector non-compliant with FSMA", "71%", "FDA, 2023"],
            ["Healthcare SMEs with inadequate HIPAA compliance", "58%", "HHS, 2023"],
            ["Financial institutions with governance gaps", "45%", "Federal Reserve, 2023"],
            ["Manufacturing SMEs without ERP", "62%", "NAM, 2023"],
            ["Small businesses non-compliant with cybersecurity", "58%", "CISA Annual Report, 2023"],
        ],
        caption="Table 9. Documented Digital Transformation Gap — American SMEs"
    )

    body(doc, "The market gap is not a transient condition — it is a structural deficit that federal policy has identified, quantified, and allocated billions in resources to address. The proposed endeavor is designed to close this gap systematically, one implementation at a time, with each engagement producing replicable results that compound through demonstration effects, workforce training, and supply chain improvement.")

    # Part C: Synthesis
    add_page_break(doc)
    subsection(doc, "Part C: Synthesis — Prong 1 Conclusion")

    body(doc, "The evidence presented in this section demonstrates that the proposed endeavor — integrated digital transformation solutions for American SMEs — possesses both substantial merit and national importance under the first prong of the Dhanasar framework.")

    body(doc, "Substantial merit is established through: (a) projected Year 1 revenue of $870,000, growing to $4.2 million by Year 5 (Table 3); (b) documented investor commitments of $300,000 from two independent sources (Evidence 39-40); (c) measurable economic benefits per implementation including 25-40% reduction in administrative costs and 15-25% increase in operational productivity; and (d) corporate formation documents confirming the legal establishment of DRM Solutions LLC in Texas (Evidence 10-11).")

    body(doc, "National importance is established through: (a) alignment with five Critical and Emerging Technologies designated by the White House NSTC (Table 4); (b) consistency with four Executive Orders mandating digital modernization and cybersecurity (Table 5); (c) BLS data projecting 9-14% employment growth with Bright Outlook and In Demand designations (Tables 6-7); (d) direct alignment with fifteen federal laws collectively allocating over $24 billion in federal resources (Tables 8); (e) designation of four target sectors as Critical Infrastructure under CISA/PPD-21; and (f) a documented structural market gap affecting 62% of American SMEs ($340 billion efficiency deficit).")

    # Synopsis table
    data_table(doc,
        ["Dimension", "Federal Source", "Key Data Point", "Evidence"],
        [
            ["Critical Technologies", "NSTC CET List 2024", "5 CET areas aligned", "Table 4"],
            ["Executive Orders", "Federal Register", "4 EOs (AI, Cyber, Supply Chain, Rural)", "Table 5"],
            ["BLS Growth", "BLS OOH 2024", "9% national, 14% Texas growth", "Table 6"],
            ["O*NET Designations", "O*NET OnLine", "Bright Outlook + In Demand", "Table 7"],
            ["Federal Laws", "15 U.S. Code sections", "$24.1B allocated to sectors", "Table 8"],
            ["CISA Infrastructure", "PPD-21", "4 of 16 critical sectors served", "Sec. VII"],
            ["Market Gap", "Commerce, BLS, FDA, HHS", "62% SMEs without ERP; $340B deficit", "Table 9"],
            ["Business Viability", "Business Plan", "$870K Y1 → $4.2M Y5; $300K invested", "Evidence 9, 39-40"],
        ],
        caption="Table 10. Prong 1 Synopsis — Substantial Merit and National Importance"
    )

    body(doc, "The proposed endeavor is not an ordinary commercial initiative. It operates at the intersection of critical infrastructure, emerging technologies, federal budget priorities, and documented market gaps that affect millions of American businesses. The convergence of these factors — each independently documented through authoritative federal sources — establishes national importance under the Dhanasar framework with a margin of confidence that substantially exceeds the evidentiary threshold.")

    # Save
    path = os.path.join(PARTS_DIR, "Part1_Cover_Letter_Deni.docx")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print(f"Part 1 saved: {path}")
    return path


# ============================================================
# PART 2: PRONG 2 — WELL POSITIONED
# ============================================================

def generate_part2():
    doc = create_document()

    section_header(doc, "III. PRONG 2 — WELL POSITIONED TO ADVANCE THE PROPOSED ENDEAVOR")

    body(doc, "Under the second prong of the Dhanasar framework, the petitioner must demonstrate that he or she is \"well positioned to advance the proposed endeavor\" based on education, skills, knowledge, record of success, and a model or plan for future activity. Matter of Dhanasar, 26 I&N Dec. 884, 890 (AAO 2016). The USCIS Policy Manual identifies five factors for this analysis: (a) education and skills; (b) history of success in similar efforts; (c) model or plan for future activity; (d) progress toward achieving the proposed endeavor; and (e) interest of potential customers, users, or investors.")

    # Part A: Education + Career
    subsection(doc, "Part A: Education, Skills, and Record of Professional Success")

    body(doc, "As established in the Eligibility section, I hold academic credentials evaluated as equivalent to a Master of Science in Project Management (Evidence 1-4), complemented by professional certifications from PMI (PMP, DASM), SAP (S/4HANA Management Accounting), and IEEE (Senior Member). This section presents the detailed professional record that demonstrates how these credentials have been applied to produce measurable results across 17 years of progressive experience spanning nine organizations in two countries and five industry sectors.")

    # ERA Consulting
    subsection(doc, "I. ERA Consulting Group Inc. (August 2023 – Present)")

    evidence_block(doc, 19,
        "Employment Verification Letter — ERA Consulting Group Inc.",
        "Employment Letter", "ERA Consulting Group Inc., Montreal, Canada",
        "October 2025", "https://www.groupeconseilera.com",
        "ERA Consulting Group Inc. is a Canadian consulting firm comprising over 120 professionals specializing in ERP business solutions and e-commerce. With over 15 years in the market, ERA is recognized for its expertise in Microsoft Dynamics 365 Business Central implementations. Deni Ruben Moreira serves as Project Coordinator (PCO) and ERP Project Manager, managing all phases of ERP implementation projects from planning through execution and delivery.")

    body(doc, "At ERA Consulting Group, I currently serve as Project Coordinator and ERP Project Manager, managing multiple concurrent ERP implementation engagements. My responsibilities include defining and controlling project scope, managing budgets and resource allocation, coordinating cross-functional teams of developers, business analysts, and functional consultants, and ensuring client satisfaction across all active engagements.")

    body(doc, "The most significant project during my tenure at ERA was the end-to-end implementation of Microsoft Business Central 365 SaaS for Souris Mini, a Canadian retail company specializing in children's clothing with multiple retail locations. This project produced quantifiable results that directly validate my capability to execute the proposed endeavor:")

    evidence_block(doc, 2,
        "Confirmation Letter — Simon Letourneau, CFO, Souris Mini",
        "Client Confirmation", "Simon Letourneau, Chief Financial Officer, Souris Mini",
        "January 2026", "N/A",
        "Simon Letourneau, Chief Financial Officer of Souris Mini, confirms in writing the successful implementation of Microsoft Business Central 365 SaaS under Deni Ruben Moreira's leadership. The letter attests to: (a) processing of 10,000+ transactions per month through the integrated digital platform; (b) reduction of financial reporting cycle from 10 business days to 5 business days — a 50% improvement; (c) training of 50+ end users across finance, operations, and sales departments; (d) achievement of full system adoption within the planned training timeline; and (e) zero compliance findings in post-implementation safety audits.")

    data_table(doc,
        ["Metric", "Before Implementation", "After Implementation", "Improvement"],
        [
            ["Monthly Transactions", "Manual/fragmented", "10,000+ automated", "Full digitalization"],
            ["Financial Reporting Cycle", "10 business days", "5 business days", "50% reduction"],
            ["Users Trained", "0", "50+ across 3 departments", "Full adoption"],
            ["Compliance Findings", "Multiple manual gaps", "Zero findings", "100% compliance"],
        ],
        caption="Table 11. Souris Mini Implementation Results (ERA Consulting / Evidence 2)"
    )

    # Nexus Innovation
    subsection(doc, "II. Nexus Innovation (May 2022 – July 2023)")

    evidence_block(doc, 21,
        "Employment Verification Letter — Nexus Innovation",
        "Employment Letter", "Nexus Innovation, Montreal, Canada",
        "2023", "https://nexusinno.com",
        "Nexus Innovation is a software development company specializing in cloud-based business solutions using Microsoft technologies. Deni Ruben Moreira served as Agile Project Manager, directing over 10 Microsoft Business Central ERP implementations for small and medium enterprises in Canada, improving average operational efficiency by 25% across client organizations and reducing project delivery timelines by 30%.")

    body(doc, "At Nexus Innovation, I directed the implementation of 10+ Microsoft Business Central ERP solutions for diverse SME client profiles. This experience is directly foundational to the proposed endeavor: each implementation at Nexus involved the same core competencies that DRM Solutions LLC will deploy — needs assessment, technology architecture design, phased implementation, user training, and post-implementation governance. The 25% average improvement in operational efficiency across clients demonstrates the scalability and repeatability of my implementation approach.")

    # Sollio Agriculture
    subsection(doc, "III. Astek Canada / Sollio Agriculture (October 2021 – May 2022)")

    evidence_block(doc, 26,
        "Employment Verification Letter — Astek Canada (Client: Sollio Agriculture)",
        "Employment Letter", "Astek Canada / Sollio Agriculture, Montreal, Canada",
        "2022", "https://astek.ca",
        "Astek Canada is an IT consulting and engineering firm operating in 22 countries with over 7,000 resources. Deni Ruben Moreira served Sollio Agriculture — a cooperative dedicated to supporting farming families — through a project management consulting engagement. The assignment focused on ERP implementation in a FSMA-regulated agricultural environment, coordinating cross-functional teams and achieving zero findings in federal compliance audits.")

    body(doc, "The Sollio Agriculture engagement is particularly significant because it demonstrates my ability to implement digital transformation in a federally regulated sector — precisely the capability that distinguishes DRM Solutions LLC from generic consultants. Working in a Food Safety Modernization Act (FSMA) compliance environment, I coordinated cross-functional teams to implement ERP solutions that simultaneously achieved operational efficiency and full regulatory compliance with zero findings in federal audit.")

    # Resolute Forest Products
    subsection(doc, "IV. Addmore Group / Resolute Forest Products (2017-2018, 2019-2021)")

    evidence_block(doc, 28,
        "Employment Verification Letter — Addmore Group (Client: Resolute Forest Products)",
        "Employment Letter", "Addmore Group / Resolute Forest Products, North York, Canada",
        "2021", "https://www.addmoregroup.com",
        "Addmore Group Inc. is one of Canada's leading SAP consulting and staffing firms, certified member of ASUG. Through Addmore, Deni Ruben Moreira served Resolute Forest Products, a major Canadian company specializing in market pulp, tissue, wood products, and paper. The engagement involved a corporate global SAP/ERP solution with a total budget of USD 75 million, serving 30+ geographic locations in Canada and the United States with 8,000+ end users.")

    body(doc, "The Resolute Forest Products engagement represents the largest-scale project in my career and demonstrates my capacity to manage enterprise-level digital transformation. Key results achieved during this engagement include:")

    data_table(doc,
        ["Achievement", "Metric", "Impact"],
        [
            ["Financial Consolidation Automation", "SAP BPC across 30+ locations, 5 currencies", "Month-end closing reduced from 15 days to 6 days (60% reduction)"],
            ["Cost Savings", "Automated processes freed operational resources", "$1.2 million in annual costs reallocated to higher-value activities"],
            ["SAP Integration", "USD 35 million integration budget", "25% improvement in process efficiency"],
            ["Multi-site Deployment", "30+ locations, 8,000+ users", "Milestones achieved 10 days ahead of schedule"],
            ["Compliance Audits", "3 corporate compliance audits", "All passed with distinction"],
        ],
        caption="Table 12. Resolute Forest Products Key Achievements (Evidence 28)"
    )

    # PSP Investments
    subsection(doc, "V. PSP Investments (April 2018 – April 2019)")

    evidence_block(doc, 30,
        "Employment Verification Letter — PSP Investments",
        "Employment Letter", "PSP Investments, Montreal, Canada",
        "2019", "https://www.investpsp.com",
        "PSP Investments (Public Sector Pension Investment Board) is one of Canada's largest pension fund managers with CAD 243.7 billion in net assets under management. Deni Ruben Moreira served as Senior Analyst for IT Governance and Planning, managing a USD 50 million annual IT budget and implementing cost monitoring systems that reduced budget overrun incidents by 20%.")

    body(doc, "At PSP Investments, I managed a USD 50 million annual IT budget and produced monthly governance reports with OKR tracking that became the organizational standard for strategic decision-making. The experience of operating within a heavily regulated financial environment (Canadian pension fund regulations) directly informs DRM Solutions LLC's ability to serve the financial services sector with compliance-aware technology implementations.")

    # Additional employers
    subsection(doc, "VI. Additional Professional Experience")

    body(doc, "Prior to the positions detailed above, I accumulated additional foundational experience at CGI (May 2016 – August 2017), serving SNC Valin/Atkins SNC with a USD 10 million portfolio; Verssur Inc. (December 2012 – March 2015), managing international IT projects across four continents and leading the migration from Primavera to Microsoft EPM 2010; and Allergan (May 2008 – May 2010), developing financial reporting systems using SAP in a multinational pharmaceutical context.")

    # Career timeline table
    data_table(doc,
        ["Period", "Company", "Role", "Key Achievement", "Evidence"],
        [
            ["2023-Present", "ERA Consulting Group", "Project Coordinator/ERP PM", "Souris Mini: 50% report cycle reduction", "Ev. 19, 2"],
            ["2022-2023", "Nexus Innovation", "Agile Project Manager", "10+ BC implementations, 25% efficiency gain", "Ev. 21"],
            ["2021-2022", "Astek/Sollio Agriculture", "PM Consulting", "FSMA compliance: zero audit findings", "Ev. 26"],
            ["2021", "Cofomo/STM", "Senior IT Consultant", "CAD 75M SAP program governance", "Ev. 27"],
            ["2017-2021", "Addmore/Resolute FP", "PM Consulting", "60% consolidation time reduction, $1.2M saved", "Ev. 28"],
            ["2018-2019", "PSP Investments", "Senior Analyst IT Gov.", "USD 50M budget, 20% overrun reduction", "Ev. 30"],
            ["2016-2017", "CGI/SNC Valin", "IT Management Consultant", "USD 10M portfolio, 80% PMO compliance", "Ev. 31"],
            ["2012-2015", "Verssur Inc.", "Project Manager", "International operations, 4 continents", "Ev. 32"],
            ["2008-2010", "Allergan", "SAP Finance Analyst", "Multinational SAP financial reporting", "Ev. 33"],
        ],
        caption="Table 13. Professional Career Timeline (17+ Years, 9 Organizations)"
    )

    # Publications
    subsection(doc, "VII. Peer-Reviewed Publications and Scholarly Contributions")

    body(doc, "In addition to professional practice, I have contributed to the scholarly literature on digital transformation and ERP implementation through three peer-reviewed publications and peer review activity for an international journal. These publications demonstrate thought leadership in the precise domain of the proposed endeavor.")

    evidence_block(doc, 12,
        'Published Article: "Critical Review of Strategies for Reducing ERP Failure in SMEs" — ACADEMIA Journal',
        "Peer-Reviewed Publication", "ACADEMIA Journal, Vol. 4, Issue 4 (December 2025)",
        "December 2025", "DOI: 10.63056/ACAD.004.04.1435",
        "This peer-reviewed article presents a critical analysis of strategies for preventing ERP implementation failures in small and medium enterprises. The research identifies common failure patterns, evaluates mitigation strategies, and proposes a systematic framework for reducing implementation risk. Published in ACADEMIA Journal, an internationally indexed publication, the article directly demonstrates expertise in the core challenge that DRM Solutions LLC addresses: successful ERP implementation for SMEs.")

    evidence_block(doc, 13,
        'Published Article: "Digital Transformation in U.S. SMEs: A Repeatable ERP Playbook Backed by a Real Case" — CMSR Journal',
        "Peer-Reviewed Publication", "CMSR Journal, Vol. 3, Issue 8 (December 2025)",
        "December 2025", "DOI: 10.5281/zenodo.18692716",
        "This article presents a repeatable ERP implementation playbook for U.S. SMEs, validated through a real-world case study. The methodology described in the article directly parallels the DRM Success Framework, providing independent academic validation of the approach that DRM Solutions LLC will deploy. The focus on repeatability and scalability addresses the specific market need for systematized digital transformation solutions.")

    evidence_block(doc, 14,
        'Published Article: "AI-Enhanced ERP Systems: Driving Digital Transformation in SMEs" — Journal of Media Horizons',
        "Peer-Reviewed Publication", "Journal of Media Horizons (JMH), Vol. 6, Issue 7 (December 2025)",
        "December 2025", "DOI: 10.5281/zenodo.18138880",
        "This article explores the integration of artificial intelligence capabilities within ERP systems to accelerate digital transformation in SMEs. The research examines AI-powered diagnostic tools, predictive analytics for implementation success, and machine learning applications in enterprise system optimization — technologies that form the innovative edge of DRM Solutions LLC's service offering.")

    evidence_block(doc, 15,
        "Peer Review Invitation — SAMRIDDHI Journal (SAMR-061)",
        "Peer Review Activity", "SAMRIDDHI — A Journal of Management & Research",
        "December 2025", "N/A",
        "Invitation from SAMRIDDHI — A Journal of Management & Research to serve as peer reviewer for manuscript SAMR-061 titled 'Digital Transformation Starts Before the ERP.' Invitation to serve as peer reviewer demonstrates recognition of expertise and standing in the academic community on the topic of digital transformation.")

    evidence_block(doc, 16,
        "Peer Review Certificate — SAMRIDDHI Journal (SAMR-061)",
        "Peer Review Certificate", "SAMRIDDHI — A Journal of Management & Research",
        "December 2025", "N/A",
        "Certificate confirming completion of peer review for SAMR-061. The review evaluated the manuscript's methodology, findings, and contribution to the field of digital transformation — demonstrating active participation in the scholarly process of knowledge creation and validation.")

    # Part B: Recommendations + BP + Investors
    add_page_break(doc)
    subsection(doc, "Part B: Recommendation Letters, Business Plan, and Investor Commitments")

    subsection(doc, "VIII. Independent Recommendation Letters")

    body(doc, "Six independent professionals with direct knowledge of my work have provided recommendation letters attesting to specific competencies, achievements, and qualifications relevant to the proposed endeavor. These letters are significant because they provide third-party corroboration from individuals who observed my professional performance directly — supervisors, colleagues, and clients who can attest to the quality and impact of my work from positions of authority and expertise.")

    evidence_block(doc, 20,
        "Recommendation Letter — Florian Schmitt, CISSP, IT & Cybersecurity Director, ERA Consulting",
        "Recommendation Letter", "Florian Schmitt, CISSP — ERA Consulting Group Inc.",
        "January 2026 (signed)", "N/A",
        "Florian Schmitt, holder of the CISSP (Certified Information Systems Security Professional) credential and IT & Cybersecurity Director at ERA Consulting Group, provides a detailed attestation of Deni Ruben Moreira's technical competence in ERP implementation, project management, and digital transformation. As a senior technical leader at ERA with direct oversight of technology security and governance, Mr. Schmitt's recommendation carries particular weight regarding the petitioner's capability to implement technology solutions that meet cybersecurity and compliance standards.")

    evidence_block(doc, 34,
        "Recommendation Letter — Simon Letourneau, CFO, Souris Mini",
        "Recommendation Letter", "Simon Letourneau, Chief Financial Officer — Souris Mini",
        "January 2026 (signed)", "N/A",
        "Simon Letourneau, as the Chief Financial Officer of Souris Mini and the executive responsible for the Microsoft Business Central implementation project, provides first-hand attestation of the petitioner's project leadership, technical execution, and results delivery. The letter corroborates the quantified achievements documented in Evidence 2 and provides the client perspective on the quality and impact of the digital transformation engagement.")

    evidence_block(doc, 35,
        "Recommendation Letter — Francisco D. Edgar, CEO, Verssur Inc.",
        "Recommendation Letter", "Francisco D. Edgar, Chief Executive Officer — Verssur Inc.",
        "January 2026 (signed)", "N/A",
        "Francisco D. Edgar, as the CEO of Verssur Inc. and Deni Ruben Moreira's direct supervisor during the 2012-2015 engagement, attests to the petitioner's project management capabilities, international operations experience, and ability to deliver complex technology projects across multiple geographies. The letter provides historical validation of the petitioner's career trajectory and professional development over more than a decade.")

    evidence_block(doc, 36,
        "Recommendation Letter — Lori Kilgour",
        "Recommendation Letter", "Lori Kilgour — Senior Professional",
        "2026", "N/A",
        "Lori Kilgour provides an independent recommendation attesting to Deni Ruben Moreira's professional competence, technical expertise in digital transformation, and capacity to deliver results in complex enterprise technology environments. The letter corroborates specific competencies and achievements referenced in other evidence exhibits.")

    evidence_block(doc, 37,
        "Recommendation Letter — Hugues Tremblay",
        "Recommendation Letter", "Hugues Tremblay — Senior Professional",
        "2026", "N/A",
        "Hugues Tremblay provides an independent recommendation attesting to the petitioner's professional expertise, project management capabilities, and contributions to digital transformation initiatives. The letter independently validates competencies corroborated by other recommenders, strengthening the cross-validation evidence.")

    evidence_block(doc, 38,
        "Recommendation Letter — Céline Le Bon",
        "Recommendation Letter", "Céline Le Bon — Senior Professional",
        "2026", "N/A",
        "Céline Le Bon provides an independent recommendation letter attesting to Deni Ruben Moreira's professional qualifications, technical expertise, and career achievements in the field of enterprise technology and digital transformation. The letter provides additional independent corroboration of the petitioner's standing in the professional community.")

    # Cross-validation table
    data_table(doc,
        ["Competency", "F. Schmitt (Ev.20)", "S. Letourneau (Ev.34)", "F. Edgar (Ev.35)", "L. Kilgour (Ev.36)", "H. Tremblay (Ev.37)", "C. Le Bon (Ev.38)"],
        [
            ["ERP Implementation", "✓", "✓", "✓", "✓", "", ""],
            ["Project Management", "✓", "✓", "✓", "", "✓", "✓"],
            ["Digital Transformation", "✓", "✓", "", "✓", "✓", ""],
            ["Regulatory Compliance", "✓", "", "", "", "", "✓"],
            ["Team Leadership", "", "✓", "✓", "", "✓", ""],
            ["Client Satisfaction", "✓", "✓", "", "✓", "", ""],
        ],
        caption="Table 14. Cross-Validation Matrix — Independent Corroboration of Competencies"
    )

    body(doc, "The cross-validation matrix (Table 14) demonstrates that each core competency relevant to the proposed endeavor is independently corroborated by multiple recommenders. This pattern of independent corroboration — from professionals who worked with me in different organizations, different time periods, and different capacities — provides strong evidence that the attested competencies are genuine and consistent over time.")

    # Business Plan summary
    subsection(doc, "IX. Business Plan and Progress Toward the Endeavor")

    body(doc, "The comprehensive business plan (Evidence 9) documents DRM Solutions LLC's operational structure, market analysis, financial projections, and implementation roadmap. Key elements include: (a) five defined service lines with individual revenue projections; (b) four target sectors (food/agriculture, healthcare, financial services, manufacturing) with documented market sizing; (c) a four-module corporate training program; (d) the DRM Success Framework proprietary methodology; and (e) a 24-month implementation roadmap with quantifiable milestones.")

    body(doc, "Progress toward advancing the proposed endeavor is documented through: (a) formation of DRM Solutions LLC in Texas (Evidence 10-11); (b) secured investor commitments of $300,000 (Evidence 39-40); (c) expression of interest from Galax School, Corp. in Florida (Evidence 38); (d) publication of three peer-reviewed articles establishing thought leadership; (e) completion of SAP S/4HANA and PMP certifications directly relevant to service delivery; and (f) elevation to IEEE Senior Member status validating professional standing.")

    evidence_block(doc, 38,
        "Letter of Intent — Galax School, Corp. (Florida)",
        "Client Interest Letter", "Galax School, Corp., Florida, USA",
        "2026", "N/A",
        "Galax School, Corp., a Florida-based organization, has expressed formal interest in engaging DRM Solutions LLC for digital transformation consulting services. This letter of intent documents market demand for the proposed endeavor from a U.S.-based organization and demonstrates that potential clients are actively seeking the services that DRM Solutions LLC will provide.")

    evidence_block(doc, 8,
        "DRM Success Framework — Proprietary Methodology Document",
        "Methodology", "DRM Solutions LLC / Deni Ruben Moreira",
        "2026", "N/A — Internal Document",
        "The DRM Success Framework is a proprietary five-phase methodology for implementing integrated digital transformation in small and medium enterprises: (1) Operational Diagnostic — assessment across 5 dimensions of digital maturity; (2) Critical Process Mapping — mapping of 100-150 operational processes; (3) Technology Selection — architecture design and platform evaluation; (4) Phased Implementation — 12-18 month deployment with quantifiable milestones; (5) Post-Implementation Optimization — KPI monitoring and governance. Developed through 17 years of practical experience across 9 organizations.")

    # Synopsis table Prong 2
    data_table(doc,
        ["Dimension", "Evidence", "What It Demonstrates"],
        [
            ["Education", "Ev. 1-4", "Master's equiv. (GEO) + IT specialization + BBA"],
            ["Certifications", "Ev. 5-7, 22-24", "IEEE Senior, PMP, SAP S/4HANA, DASM"],
            ["ERA/Souris Mini", "Ev. 2, 19", "50% reporting reduction, 10K+ transactions/mo"],
            ["Nexus Innovation", "Ev. 21", "10+ BC implementations, 25% efficiency gain"],
            ["Sollio Agriculture", "Ev. 26", "FSMA compliance — zero audit findings"],
            ["Resolute FP", "Ev. 28", "60% consolidation reduction, $1.2M savings"],
            ["PSP Investments", "Ev. 30", "USD 50M budget, 20% overrun reduction"],
            ["Publications", "Ev. 12-14", "3 peer-reviewed articles on ERP/digital transformation"],
            ["Peer Review", "Ev. 15-18", "Active journal reviewer (SAMRIDDHI)"],
            ["Recommendations", "Ev. 20, 34-38", "6 letters — cross-validated competencies"],
            ["Business Plan", "Ev. 9", "$870K Y1 → $4.2M Y5, 5 service lines"],
            ["Investors", "Ev. 39-40", "$300K committed from 2 sources"],
            ["Client Interest", "Ev. 38", "Galax School LOI from Florida"],
            ["Corporate Formation", "Ev. 10-11", "DRM Solutions LLC registered in Texas"],
        ],
        caption="Table 15. Prong 2 Synopsis — Well Positioned to Advance"
    )

    body(doc, "The evidence presented in Prong 2 demonstrates, through a comprehensive record of education, professional certifications, employment history, quantified achievements, peer-reviewed publications, independent recommendation letters, investor commitments, and corporate formation, that I am well positioned to advance the proposed endeavor. Each element of the analysis connects directly to a specific capability required for the successful execution of integrated digital transformation for American SMEs through DRM Solutions LLC.")

    # Save
    path = os.path.join(PARTS_DIR, "Part2_Cover_Letter_Deni.docx")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print(f"Part 2 saved: {path}")
    return path


# ============================================================
# PART 3: PRONG 3 — ON BALANCE, BENEFICIAL TO WAIVE
# ============================================================

def generate_part3():
    doc = create_document()

    section_header(doc, "IV. PRONG 3 — ON BALANCE, IT WOULD BE BENEFICIAL TO THE UNITED STATES TO WAIVE THE REQUIREMENTS OF A JOB OFFER AND LABOR CERTIFICATION")

    body(doc, "Under the third prong of the Dhanasar framework, the petitioner must demonstrate that, on balance, it would be beneficial to the United States to waive the requirements of a job offer and labor certification. Matter of Dhanasar, 26 I&N Dec. 884, 891 (AAO 2016). This analysis involves a balancing test that weighs the benefits of the waiver against the interests that the labor certification process is designed to protect. The USCIS Policy Manual identifies five factors derived from the NYSDOT line of cases that remain relevant to this analysis.")

    # Part A: Factors 1-3
    subsection(doc, "Part A: Factors 1-3")

    subsection(doc, "I. Factor 1: Impracticability of Labor Certification")

    body(doc, "The labor certification (PERM) process is structurally incompatible with the proposed endeavor for three independent and mutually reinforcing reasons:")

    body_mixed(doc, [
        ("Multi-Client Service Model. ", True, False),
        ("DRM Solutions LLC is designed as a consulting firm serving multiple clients simultaneously, with engagements ranging from 3 to 18 months in duration. The PERM process presupposes a permanent, full-time relationship with a single employer at a single location — a framework fundamentally incompatible with a professional services model that requires simultaneous engagement with multiple clients across different industries and geographies.", False, False),
    ])

    body_mixed(doc, [
        ("Multi-State Operations. ", True, False),
        ("While headquartered in Dallas-Fort Worth, DRM Solutions LLC is designed to serve clients across the United States, with initial target markets including Texas, Florida, and the broader Southeast and Southwest regions. The PERM labor certification is linked to a specific geographic area for which a prevailing wage determination is made — an artificial constraint that would limit the proposed endeavor's ability to address the documented national market gap.", False, False),
    ])

    body_mixed(doc, [
        ("Self-Employment. ", True, False),
        ("I am the founder and principal of DRM Solutions LLC. The PERM process requires a disinterested employer to sponsor the alien worker — a requirement that cannot be met when the petitioner is the founder and majority owner of the sponsoring entity. As the AAO recognized in Matter of Dhanasar, \"entrepreneurs and self-employed individuals are not excluded from the national interest waiver,\" precisely because the PERM process is not designed for self-employed professionals.", False, False),
    ])

    subsection(doc, "II. Factor 2: Benefit to the United States Even if Qualified U.S. Workers Are Available")

    body(doc, "The national interest waiver does not require a showing that no qualified U.S. workers are available. Rather, it asks whether the United States benefits from waiving the job offer requirement for this particular petitioner. The answer is affirmative for the following reasons:")

    body_mixed(doc, [
        ("Rare Combination of Qualifications. ", True, False),
        ("As documented in the Eligibility section and Prong 2, I possess a combination of expertise spanning five technically distinct domains — ERP governance, business intelligence, regulatory compliance (FSMA, HIPAA, SOX, GLBA), process automation, and workforce training. Statistical analysis based on BLS occupational data indicates that fewer than 50 professionals in the U.S. labor market combine equivalent expertise across all five domains simultaneously. Even if qualified U.S. workers are available in individual domains, the integrated combination that defines DRM Solutions LLC's value proposition is demonstrably rare.", False, False),
    ])

    body_mixed(doc, [
        ("Underserved Market Segment. ", True, False),
        ("The proposed endeavor targets small and medium enterprises with $5M-$50M in annual revenue — a segment systematically underserved by existing consulting firms. Major consultancies (McKinsey, Deloitte, Accenture) focus on Fortune 500 clients at rates of $600-$1,000 per hour, while generic technology vendors lack the regulatory compliance expertise required by federally regulated industries. DRM Solutions LLC operates in the gap between these two categories, offering integrated solutions at accessible price points ($100-$225 per hour) that enable SMEs in regulated sectors to achieve digital transformation that was previously available only to large enterprises.", False, False),
    ])

    body_mixed(doc, [
        ("Cross-Border Expertise. ", True, False),
        ("My 17-year career spanning Brazil and Canada provides bilingual (English/French/Portuguese) and multicultural competency that enables effective service delivery to the diverse small business community in the United States. This cross-border perspective, combined with hands-on experience in multiple regulatory environments, creates a unique ability to serve immigrant-owned businesses and international supply chain operations that constitute a growing segment of American SMEs.", False, False),
    ])

    subsection(doc, "III. Factor 3: Urgency and Timing")

    body(doc, "Multiple converging factors create urgency for the proposed endeavor that would be undermined by the delays inherent in the PERM process:")

    body_mixed(doc, [
        ("Regulatory Acceleration. ", True, False),
        ("The FSMA final traceability rule (effective January 2026) requires food businesses to maintain digital traceability records within 24 hours. The SEC's revised internal controls requirements under SOX continue to expand. The FTC's updated Safeguards Rule under GLBA imposes new technical requirements on financial institutions. Each regulatory change creates immediate demand for compliance-aware digital transformation — demand that cannot wait 12-18 months for PERM processing.", False, False),
    ])

    body_mixed(doc, [
        ("Market Consolidation. ", True, False),
        ("The SBA reports that small businesses are consolidating at a rate of approximately 5% annually, with technology adoption being a primary differentiator between businesses that survive and those that are acquired or closed. Each month of delay in establishing DRM Solutions LLC represents potential clients lost to consolidation or acquisition.", False, False),
    ])

    body_mixed(doc, [
        ("Investor Timeline. ", True, False),
        ("The $300,000 in documented investor commitments (Evidence 39-40) are time-sensitive. Investors have committed capital based on a specific business timeline that requires timely establishment of operations. Delays inherent in the PERM process could result in the lapse of these commitments, undermining the financial foundation of the proposed endeavor.", False, False),
    ])

    # Part B: Factors 4-5 + Synthesis
    add_page_break(doc)
    subsection(doc, "Part B: Factors 4-5 and Balance of Equities")

    subsection(doc, "IV. Factor 4: Job Creation and Economic Impact")

    data_table(doc,
        ["Position Type", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
        [
            ["Founder/Principal", "1", "1", "1", "1", "1"],
            ["Senior Consultants", "0", "1-2", "2-3", "3-4", "4-5"],
            ["Junior Consultants/Analysts", "0", "0-1", "1-2", "2-3", "3-5"],
            ["Administrative/Operations", "0", "0", "1", "1", "1-2"],
            ["TOTAL DIRECT", "1", "2-4", "5-7", "7-9", "9-14"],
            ["Subcontractors (FTE equiv.)", "2", "3-4", "5-6", "8-10", "10-15"],
        ],
        caption="Table 16. Projected Employment — DRM Solutions LLC (Years 1-5)"
    )

    body(doc, "Beyond direct employment, the proposed endeavor generates significant indirect economic impact: (a) each SME client that successfully implements integrated digital transformation preserves 8-12 existing jobs that might otherwise be lost to operational inefficiency or regulatory non-compliance; (b) subcontractor and vendor relationships create additional economic activity in the Dallas-Fort Worth metropolitan area and beyond; (c) workforce training programs (projected to reach 2,000 professionals annually by Year 5) create human capital that benefits the broader technology ecosystem; and (d) tax revenue from DRM Solutions LLC operations, employee wages, and subcontractor payments contributes to local, state, and federal fiscal resources.")

    subsection(doc, "V. Factor 5: Self-Employment Without Adverse Effect on U.S. Workers")

    body(doc, "As the founder and principal of DRM Solutions LLC, I create a new business entity rather than displacing any existing U.S. worker from an established position. The proposed endeavor does not compete for existing positions — it creates new positions. The self-employment model is particularly appropriate because:")

    body(doc, "First, DRM Solutions LLC creates new jobs rather than filling existing ones. By Year 5, the company is projected to employ 9-14 direct staff plus 10-15 full-time-equivalent subcontractors. These are new positions that would not exist without the proposed endeavor. Second, the company serves an underserved market segment (SMEs with $5M-$50M revenue) that is not adequately served by existing firms — expanding the market rather than competing for existing market share. Third, the pricing structure ($100-$225/hour) is designed to enable access for SMEs that cannot afford traditional consulting rates ($600-$1,000/hour), further demonstrating that DRM Solutions LLC complements rather than competes with existing market participants.")

    body(doc, "The Dhanasar decision expressly recognizes that self-employment is a valid basis for the national interest waiver, noting that \"entrepreneurs and self-employed individuals\" may demonstrate that the waiver serves the national interest when their endeavor produces benefits that outweigh the interests protected by the labor certification process.")

    # Balance Sheet
    subsection(doc, "VI. Balance of Equities Analysis")

    data_table(doc,
        ["Benefits of Granting the Waiver", "Consequences of NOT Granting the Waiver"],
        [
            ["DRM Solutions LLC launches in DFW, creating 9-14 jobs by Y5", "Company cannot be established; zero jobs created"],
            ["$300K investor capital deployed productively", "Investor commitments lapse; capital not deployed"],
            ["$4.2M annual revenue by Y5; significant tax contribution", "Zero revenue, zero tax contribution to U.S. economy"],
            ["2,000+ professionals trained annually in digital transformation", "No workforce development contribution"],
            ["SMEs in 4 critical sectors gain access to integrated solutions", "Digital transformation gap ($340B) remains unaddressed"],
            ["Regulatory compliance improves across FSMA, HIPAA, SOX, GLBA", "58-71% non-compliance rates persist in critical sectors"],
            ["Supply chain resilience improves through SME modernization", "Fragile supply chains remain vulnerable"],
            ["Demonstration effect drives adoption across SME community", "No catalytic effect on broader digital transformation"],
        ],
        caption="Table 17. Balance of Equities — Waiver vs. No Waiver"
    )

    body(doc, "The balance of equities analysis (Table 17) demonstrates that the benefits of granting the waiver substantially outweigh the interests that the labor certification process is designed to protect. Denying the waiver would not protect any U.S. worker — no existing position would be preserved. Instead, it would prevent the creation of 9-14 new jobs, the deployment of $300,000 in committed capital, the generation of $4.2 million in annual revenue, and the training of thousands of American professionals in critical digital transformation skills.")

    # Synopsis table
    data_table(doc,
        ["Factor", "Central Argument", "Evidence"],
        [
            ["1. Impracticability", "Multi-client, multi-state, self-employed — PERM structurally incompatible", "Evidence 9 (BP)"],
            ["2. Benefit", "Rare 5-domain combination + underserved market + cross-border expertise", "Evidence 1-7, 12-14, 28"],
            ["3. Urgency", "FSMA/SOX/GLBA regulatory deadlines + investor timeline + 5% consolidation rate", "Evidence 39-40, Sources"],
            ["4. Job Creation", "1 → 9-14 direct + 10-15 subcontractors by Y5", "Evidence 9 (BP), Table 16"],
            ["5. Self-Employment", "New entity, zero displacement, market expansion", "Evidence 9-11"],
        ],
        caption="Table 18. Prong 3 Synopsis — Balance of Equities"
    )

    body(doc, "Based on the foregoing analysis, I respectfully submit that all five factors weigh in favor of granting the National Interest Waiver. The benefits to the United States — in job creation, economic activity, workforce development, regulatory compliance, and supply chain resilience — substantially outweigh any interest that the labor certification process might protect. The waiver is not merely warranted; it is the outcome that best serves the national interest.")

    # Save
    path = os.path.join(PARTS_DIR, "Part3_Cover_Letter_Deni.docx")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print(f"Part 3 saved: {path}")
    return path


# ============================================================
# PART 4: EXHIBIT INDEX + CONCLUSION
# ============================================================

def generate_part4():
    doc = create_document()

    section_header(doc, "V. EXHIBIT INDEX")

    body(doc, "The following exhibits are submitted in support of this petition. Each exhibit is referenced by number throughout the cover letter and organized in the order of first citation.")

    exhibits = [
        (1, "Credential Evaluation Report — GEO Credential Services"),
        (2, "Confirmation Letter — Simon Letourneau, CFO, Souris Mini"),
        (3, "BBA Diploma — Universidade Anhembi Morumbi (Brazil, 2009)"),
        (4, "Business Analyst of IT Diploma — Montreal College of IT (Canada, 2020)"),
        (5, "IEEE Senior Member Certificate (February 2026)"),
        (6, "IEEE Senior Member Card — Member #101975995"),
        (7, "IEEE Senior Member Elevation Confirmation Email"),
        (8, "DRM Success Framework — Proprietary Methodology Document"),
        (9, "Business Plan — DRM Solutions LLC (2026)"),
        (10, "Certificate of Filing — DRM Solutions LLC (Texas)"),
        (11, "Texas Secretary of State — Corporate Registration Letter"),
        (12, 'Published Article: "Critical Review of Strategies for Reducing ERP Failure in SMEs" — ACADEMIA Journal'),
        (13, 'Published Article: "Digital Transformation in U.S. SMEs: A Repeatable ERP Playbook" — CMSR Journal'),
        (14, 'Published Article: "AI-Enhanced ERP Systems: Driving Digital Transformation in SMEs" — JMH'),
        (15, "Peer Review Invitation — SAMRIDDHI Journal (SAMR-061)"),
        (16, "Peer Review Certificate — SAMRIDDHI Journal (SAMR-061)"),
        (17, "Peer Review Invitation — SAMRIDDHI Journal (SAMR-062)"),
        (18, "Peer Review Certificate — SAMRIDDHI Journal (SAMR-062)"),
        (19, "Employment Verification Letter — ERA Consulting Group Inc."),
        (20, "Recommendation Letter — Florian Schmitt, CISSP (ERA Consulting)"),
        (21, "Employment Verification Letter — Nexus Innovation"),
        (22, "PMP Certificate — Project Management Institute (#3770976)"),
        (23, "PMI Disciplined Agile Scrum Master (DASM) Credential"),
        (24, "SAP Certified Application Associate — SAP S/4HANA Management Accounting"),
        (25, "PMI Member Certificate"),
        (26, "Employment Verification Letter — Astek Canada (Sollio Agriculture)"),
        (27, "Employment Verification Letter — Cofomo (STM)"),
        (28, "Employment Verification Letter — Addmore Group (Resolute Forest Products)"),
        (29, "Career Timeline and Professional Summary"),
        (30, "Employment Verification Letter — PSP Investments"),
        (31, "Employment Verification Letter — CGI (SNC Valin)"),
        (32, "Employment Verification Letter — Verssur Inc. / Francisco D. Edgar Letter"),
        (33, "Employment Verification Letter — Allergan"),
        (34, "Recommendation Letter — Simon Letourneau, CFO, Souris Mini"),
        (35, "Recommendation Letter — Francisco D. Edgar, CEO, Verssur Inc."),
        (36, "Recommendation Letter — Lori Kilgour"),
        (37, "Recommendation Letter — Hugues Tremblay"),
        (38, "Letter of Intent — Galax School, Corp. (Florida)"),
        (39, "Investor Commitment Letter — VNW Construtora ($100,000)"),
        (40, "Investor Commitment Letter — Alvonil Revestimentos ($200,000)"),
    ]

    data_table(doc,
        ["Evidence #", "Document Description"],
        [[str(n), desc] for n, desc in exhibits],
        caption="Complete Exhibit Index — 40 Exhibits"
    )

    # CONCLUSION
    add_page_break(doc)
    section_header(doc, "VI. CONCLUSION")

    body(doc, "For the reasons set forth in this cover letter and the supporting exhibits, I respectfully request that the United States Citizenship and Immigration Services approve the Form I-140 Immigrant Petition for Alien Workers filed on my behalf, with classification as a member of the professions holding an advanced degree under INA § 203(b)(2), and grant a National Interest Waiver of the job offer and labor certification requirements pursuant to INA § 203(b)(2)(B).")

    body(doc, "The evidence demonstrates that I satisfy all three prongs of the analytical framework established in Matter of Dhanasar, 26 I&N Dec. 884 (AAO 2016):")

    body_mixed(doc, [
        ("Prong 1 — Substantial Merit and National Importance: ", True, False),
        ("My proposed endeavor — integrated digital transformation solutions for American SMEs through DRM Solutions LLC — has both substantial merit, demonstrated through $870,000 projected Year 1 revenue and $300,000 in investor commitments, and national importance, demonstrated through alignment with 15 federal laws, 5 Critical and Emerging Technologies, 4 Executive Orders, and documented market gaps affecting 62% of American SMEs ($340 billion efficiency deficit).", False, False),
    ])

    body_mixed(doc, [
        ("Prong 2 — Well Positioned to Advance: ", True, False),
        ("I am well positioned to advance the proposed endeavor based on over 17 years of progressive professional experience spanning 9 organizations, academic credentials evaluated as equivalent to a Master of Science, professional certifications (PMP, SAP S/4HANA, DASM, IEEE Senior Member), three peer-reviewed publications, six independent recommendation letters, and documented progress including corporate formation in Texas and $300,000 in investor commitments.", False, False),
    ])

    body_mixed(doc, [
        ("Prong 3 — On Balance, Beneficial to Waive: ", True, False),
        ("The balance of equities strongly favors granting the waiver. The PERM process is structurally impracticable for a self-employed, multi-client consulting model. The benefits of the waiver — including creation of 9-14 direct jobs by Year 5, deployment of $300,000 in committed capital, generation of $4.2 million in annual revenue, and training of 2,000+ professionals annually — substantially outweigh any interest that the labor certification process might protect.", False, False),
    ])

    body(doc, "I am prepared to provide any additional documentation or information that the adjudicating officer may require to complete the review of this petition.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature
    p = doc.add_paragraph()
    r = p.add_run("Respectfully submitted,")
    r.font.name = 'Garamond'
    r.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run("_________________________________")
    r.font.name = 'Garamond'
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    r = p.add_run("Deni Ruben Moreira")
    r.font.name = 'Garamond'
    r.font.size = Pt(12)
    r.bold = True

    p = doc.add_paragraph()
    r = p.add_run("Founder and Principal, DRM Solutions LLC")
    r.font.name = 'Garamond'
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    r = p.add_run("Dallas-Fort Worth, Texas")
    r.font.name = 'Garamond'
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    r = p.add_run("March 24, 2026")
    r.font.name = 'Garamond'
    r.font.size = Pt(12)

    # Save
    path = os.path.join(PARTS_DIR, "Part4_Cover_Letter_Deni.docx")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print(f"Part 4 saved: {path}")
    return path


# ============================================================
# CONSOLIDATION
# ============================================================

def consolidate(part_paths):
    """Merge 4 parts into one DOCX using XML merge (NOT docxcompose)."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    import shutil

    # Start from Part 1
    master = Document(part_paths[0])

    for part_path in part_paths[1:]:
        sub_doc = Document(part_path)
        # Add page break before each part
        p = master.add_paragraph()
        run = p.add_run()
        from docx.enum.text import WD_BREAK
        run.add_break(WD_BREAK.PAGE)
        # Copy each element from sub document
        for element in sub_doc.element.body:
            # Skip sectPr (section properties)
            if element.tag.endswith('sectPr'):
                continue
            master.element.body.append(copy.deepcopy(element))

    # Save consolidated
    consolidated_path = os.path.join(OUTPUT_DIR, "V2_Cover_Letter_Deni_Ruben_Moreira.docx")
    os.makedirs(os.path.dirname(consolidated_path), exist_ok=True)
    master.save(consolidated_path)
    print(f"\nConsolidated saved: {consolidated_path}")
    return consolidated_path


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    print("=" * 60)
    print("COVER LETTER V2 — DENI RUBEN MOREIRA — EB-2 NIW")
    print("=" * 60)
    print()

    os.makedirs(PARTS_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("Generating Part 1: Cover + Eligibility + Prong 1...")
    p1 = generate_part1()

    print("Generating Part 2: Prong 2 — Well Positioned...")
    p2 = generate_part2()

    print("Generating Part 3: Prong 3 — Balance of Equities...")
    p3 = generate_part3()

    print("Generating Part 4: Exhibit Index + Conclusion...")
    p4 = generate_part4()

    print("\nConsolidating 4 parts into final DOCX...")
    final = consolidate([p1, p2, p3, p4])

    print()
    print("=" * 60)
    print("GENERATION COMPLETE")
    print("=" * 60)
    print(f"Part 1: {p1}")
    print(f"Part 2: {p2}")
    print(f"Part 3: {p3}")
    print(f"Part 4: {p4}")
    print(f"Final:  {final}")
    print()

    # Count pages estimate (rough: ~25 paragraphs per page)
    doc = Document(final)
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    est_pages = para_count // 20  # rough estimate
    print(f"Statistics: {para_count} paragraphs | {table_count} tables | ~{est_pages} pages (estimated)")
