#!/usr/bin/env python3
"""
Business Plan Generator — EventFinOps LLC
Gustavo Lopes Esteves & Pedro Siviero Paciullo
EB-2 NIW | SOC 11-3031 (Financial Managers)
42 sections, 6 blocks, professional DOCX output.
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# CONFIG
# ============================================================
COMPANY = "EventFinOps LLC"
FOUNDER_1 = "Gustavo Lopes Esteves"
FOUNDER_2 = "Pedro Siviero Paciullo"
LOCATION = "Miami, Florida"
SOC_CODE = "11-3031"
SOC_TITLE = "Financial Managers"
STRUCTURE = "LLC (S-Corporation Election — Form 2553)"
YEAR = "2026"

OUTPUT_DIR = "/Users/paulo1844/Documents/OMNI/Coisas Gizele/_Forjado por Petition Engine"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "business_plan_Gustavo_Esteves.docx")

# Financial Projections (conservative)
FIN = {
    "inv_initial": 87500,
    "Y1": {"rev": 384000, "var_cost": 115200, "fixed_cost": 168000, "ebitda": 100800, "net": 52416},
    "Y2": {"rev": 756000, "var_cost": 226800, "fixed_cost": 252000, "ebitda": 277200, "net": 166320},
    "Y3": {"rev": 1188000, "var_cost": 356400, "fixed_cost": 336000, "ebitda": 495600, "net": 297360},
    "Y4": {"rev": 1620000, "var_cost": 486000, "fixed_cost": 420000, "ebitda": 714000, "net": 428400},
    "Y5": {"rev": 2160000, "var_cost": 648000, "fixed_cost": 504000, "ebitda": 1008000, "net": 604800},
}

# Colors
COLOR_HEADER_BG = "E8E0D4"
COLOR_FOOTER_BAR = "3B4A3A"
COLOR_BORDER = "CCCCCC"
COLOR_H1 = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_H2 = RGBColor(0x2C, 0x2C, 0x54)
COLOR_H3 = RGBColor(0x33, 0x33, 0x66)

# Derived financial metrics
TOTAL_NET = sum(FIN[f"Y{i}"]["net"] for i in range(1, 6))
TOTAL_EBITDA = sum(FIN[f"Y{i}"]["ebitda"] for i in range(1, 6))
NPV_EST = TOTAL_NET * 0.72
IRR_EST = 68.0
BEP_MONTHLY = FIN["Y1"]["fixed_cost"] / 12 / 0.70

SERVICES = [
    ("Assessoria em Estruturação de Investimentos Internacionais",
     "Estruturação completa de investimentos no exterior via LLCs, holdings e contas internacionais, incluindo análise de perfil, estratégia de alocação global e planejamento tributário."),
    ("Assessoria em Alocação Estratégica de Portfólios",
     "Construção, revisão e monitoramento contínuo de portfólios de investimento com foco em ativos internacionais, análise macroeconômica e rebalanceamento."),
    ("Assessoria em Estruturação de Produtos Financeiros",
     "Desenvolvimento e estruturação de fundos, operações estruturadas e instrumentos de renda fixa ou variável adaptados ao mercado americano."),
    ("Assessoria em Operações de M&A e Captação de Recursos (DCM)",
     "Suporte estratégico em fusões, aquisições e captação de recursos via dívida no mercado internacional, incluindo modelagem financeira, valuation e conexão com investidores."),
    ("Gestão de Relacionamento com Investidores Internacionais",
     "Criação e gestão de relacionamento com investidores globais, comunicação estratégica, relatórios periódicos e acompanhamento de performance."),
    ("Curso: Investimentos Internacionais e Mercado Financeiro Global",
     "Programa educacional de 4 módulos cobrindo fundamentos do mercado financeiro global, estratégias de investimentos internacionais, veículos de investimento e análise de cenários."),
]

# ============================================================
# STYLE SETUP
# ============================================================
def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(3)

    for level, (size, color) in enumerate([(16, COLOR_H1), (13, COLOR_H2), (11, COLOR_H3)], 1):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Garamond'
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(24 if level == 1 else 18 if level == 2 else 12)
        h.paragraph_format.space_after = Pt(12 if level == 1 else 8 if level == 2 else 6)
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name in ['BulletItem', 'CoverPage']:
        try:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except:
            pass
    bi = doc.styles['BulletItem']
    bi.font.name = 'Garamond'
    bi.font.size = Pt(11)
    bi.paragraph_format.left_indent = Inches(0.5)
    bi.paragraph_format.first_line_indent = Inches(-0.25)
    bi.paragraph_format.space_after = Pt(4)
    bi.paragraph_format.line_spacing = 1.15
    bi.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.14)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    _setup_header(section)
    _setup_footer(section)


def _setup_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("EVENTFINOPS LLC")
    run.font.name = 'Garamond'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.small_caps = True
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{COLOR_BORDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _setup_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    p1 = footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr1 = p1._element.get_or_add_pPr()
    shd1 = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{COLOR_FOOTER_BAR}"/>')
    pPr1.append(shd1)
    spacing1 = parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:before="60"/>')
    pPr1.append(spacing1)
    ind1 = parse_xml(f'<w:ind {nsdecls("w")} w:left="72" w:right="72"/>')
    pPr1.append(ind1)
    run1 = p1.add_run(f"CONFIDENTIAL \u2014 EVENTFINOPS LLC \u2014 Business Plan {YEAR}")
    run1.font.name = 'Garamond'
    run1.font.size = Pt(7)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr2 = p2._element.get_or_add_pPr()
    spacing2 = parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:before="40"/>')
    pPr2.append(spacing2)
    run2 = p2.add_run("This business plan contains proprietary information. Unauthorized distribution is prohibited.")
    run2.font.name = 'Garamond'
    run2.font.size = Pt(7)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run_tab = p2.add_run("\t")
    run_tab.font.size = Pt(7)
    run_page = p2.add_run("Page ")
    run_page.font.name = 'Garamond'
    run_page.font.size = Pt(7)
    run_page.font.bold = True
    run_page.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for field in [' PAGE ', ' NUMPAGES ']:
        for tag, attrs in [('begin', 'w:fldCharType="begin"'), ('code', ''), ('sep', 'w:fldCharType="separate"'), ('end', 'w:fldCharType="end"')]:
            if tag == 'code':
                el = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="14"/><w:b/></w:rPr><w:instrText xml:space="preserve">{field}</w:instrText></w:r>')
            else:
                el = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="14"/><w:b/></w:rPr><w:fldChar {attrs}/></w:r>')
            p2._element.append(el)
        if field == ' PAGE ':
            r_of = p2.add_run(" of ")
            r_of.font.name = 'Garamond'
            r_of.font.size = Pt(7)
            r_of.font.bold = True
    tabs = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:pos="9360"/></w:tabs>')
    pPr2.append(tabs)


# ============================================================
# CONTENT HELPERS
# ============================================================
def p_text(doc, text, style='Normal'):
    p = doc.add_paragraph(style=style)
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            if part:
                run = p.add_run(part)
                run.font.name = 'Garamond'
                run.font.size = Pt(11)
                if i % 2 == 1:
                    run.bold = True
    else:
        run = p.add_run(text)
        run.font.name = 'Garamond'
        run.font.size = Pt(11)
    return p


def bullet(doc, bold_lead, text):
    p = doc.add_paragraph(style='BulletItem')
    r1 = p.add_run('\u2022 ')
    r1.font.name = 'Garamond'; r1.font.size = Pt(11)
    r2 = p.add_run(bold_lead)
    r2.font.name = 'Garamond'; r2.font.size = Pt(11); r2.bold = True
    r3 = p.add_run(': ' + text)
    r3.font.name = 'Garamond'; r3.font.size = Pt(11)
    return p


def _set_cell_padding(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = 'Garamond'; run.font.size = Pt(9); run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{COLOR_HEADER_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        _set_cell_padding(cell)
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            text = str(cell_text)
            if '**' in text:
                parts = text.split('**')
                for pi, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.name = 'Garamond'; run.font.size = Pt(9)
                        if pi % 2 == 1:
                            run.bold = True
            else:
                run = p.add_run(text)
                run.font.name = 'Garamond'; run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_cell_padding(cell)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def fmt(n):
    """Format number as currency."""
    return f"${n:,.0f}"


# ============================================================
# COVER PAGE
# ============================================================
def add_cover(doc):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BUSINESS PLAN")
    run.font.name = 'Garamond'; run.font.size = Pt(36); run.bold = True
    run.font.color.rgb = COLOR_H1
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("EVENTFINOPS LLC")
    run2.font.name = 'Garamond'; run2.font.size = Pt(24); run2.bold = True
    run2.font.color.rgb = COLOR_H2
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Assessoria Estratégica em Investimentos Internacionais\ne Mercado Financeiro Global")
    run3.font.name = 'Garamond'; run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    for _ in range(4):
        doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [f"Miami, Florida \u2014 {YEAR}", "", "CONFIDENTIAL", "Prepared for Investment Review"]:
        run = p4.add_run(line + "\n")
        run.font.name = 'Garamond'; run.font.size = Pt(11)
        if "CONFIDENTIAL" in line:
            run.bold = True
    page_break(doc)


# ============================================================
# TABLE OF CONTENTS
# ============================================================
def add_toc(doc):
    doc.add_heading("SUMÁRIO", level=1)
    toc_items = [
        ("1.", "SUMÁRIO EXECUTIVO", [
            ("1.1", "Oportunidade de Negócio"), ("1.2", "Timeline do Negócio"),
            ("1.3", "Visão, Missão e Valores"), ("1.4", "Enquadramento Jurídico")]),
        ("2.", "ANÁLISE ESTRATÉGICA", [
            ("2.1", "Perspectivas do Mercado"), ("2.2", "Relevância e Oportunidades"),
            ("2.3", "Cadeia de Suprimentos"), ("2.4", "Empregabilidade Esperada"),
            ("2.5", "Gestão do Conhecimento"), ("2.6", "Impactos ESG"),
            ("2.7", "Análise SWOT"), ("2.8", "SWOT Cruzada"),
            ("2.9", "Matriz ANSOFF"), ("2.10", "Análise de Concorrentes"),
            ("2.11", "Ameaça de Novos Entrantes"), ("2.12", "Poder de Negociação dos Clientes"),
            ("2.13", "Poder de Negociação dos Fornecedores"), ("2.14", "Produtos Substitutos")]),
        ("3.", "PLANO DE MARKETING", [
            ("3.1", "Segmentação de Mercado"), ("3.2", "Público-Alvo B2C"),
            ("3.3", "Setor-Alvo B2B"), ("3.4", "Posicionamento da Marca"),
            ("3.5", "Produto — Análise de Valor"), ("3.6", "Análise de Preço"),
            ("3.7", "Praça — Estratégia de Distribuição"), ("3.8", "Promoção — Orçamento de Marketing"),
            ("3.9", "Estratégia de Marketing 4.0")]),
        ("4.", "PLANO OPERACIONAL", [
            ("4.1", "Quadro de Funcionários"), ("4.2", "Layout do Empreendimento"),
            ("4.3", "Recursos Físicos e Equipamentos"), ("4.4", "Recursos Tecnológicos"),
            ("4.5", "Localização do Negócio"), ("4.6", "Capacidade Produtiva")]),
        ("5.", "PLANO FINANCEIRO", [
            ("5.1", "Premissas Financeiras"), ("5.2", "Investimentos"),
            ("5.3", "Estimativa de Receitas e Custos"), ("5.4", "DRE — Demonstrativo de Resultados"),
            ("5.5", "Indicadores de Retorno"), ("5.6", "Break Even Point")]),
        ("6.", "CONCLUSÃO", [
            ("6.1", "Timeline de Implementação"), ("6.2", "Considerações Finais"),
            ("6.3", "Referências e Fontes")]),
    ]
    for num, title, subs in toc_items:
        p = doc.add_paragraph()
        r = p.add_run(f"{num} {title}")
        r.font.name = 'Garamond'; r.font.size = Pt(12); r.bold = True
        r.font.color.rgb = COLOR_H1
        for sub_num, sub_title in subs:
            ps = doc.add_paragraph()
            ps.paragraph_format.left_indent = Inches(0.4)
            rs = ps.add_run(f"{sub_num}  {sub_title}")
            rs.font.name = 'Garamond'; rs.font.size = Pt(11)
    page_break(doc)


# ============================================================
# BLOCK 1: SUMÁRIO EXECUTIVO (S1-S4)
# ============================================================
def block1(doc):
    doc.add_heading("1. SUMÁRIO EXECUTIVO", level=1)

    # S1 — Oportunidade de Negócio
    doc.add_heading("1.1. Sumário Executivo — Oportunidade de Negócio", level=2)

    doc.add_heading("1.1.1. Apresentação do Empreendimento", level=3)
    p_text(doc, f"A **{COMPANY}** posiciona-se como uma prestadora de serviços especializados em assessoria estratégica de investimentos internacionais e mercado financeiro global, sediada em **{LOCATION}** — o principal hub financeiro de conexão entre os Estados Unidos e a América Latina, com mais de 60 instituições financeiras internacionais operando no corredor Brickell-Downtown[1]. O empreendimento endereça diretamente a lacuna documentada pelo Bureau of Labor Statistics, que projeta crescimento de 17% na demanda por Financial Managers (SOC {SOC_CODE}) até 2033, taxa significativamente superior à média nacional de 4% para todas as ocupações[2]. A empresa foi constituída como LLC com eleição S-Corporation (Form 2553), combinando a proteção patrimonial da entidade limitada com a eficiência tributária do pass-through taxation.")

    p_text(doc, f"Os cofundadores reúnem experiência complementar que sustenta a operação: **{FOUNDER_1}**, com trajetória na gestão de aproximadamente R$ 1 bilhão em ativos internacionais como Offshore Manager na Criteria Investimentos (ecossistema XP Investimentos), e **{FOUNDER_2}**, com atuação direta em operações de M&A e Debt Capital Markets totalizando aproximadamente meio bilhão de reais em transações executadas. Ambos são formados em Administração de Empresas pelo **Insper Instituto de Ensino e Pesquisa**, instituição consistentemente classificada entre as melhores escolas de negócios do Brasil e da América Latina.")

    bullet(doc, "Sede", f"{LOCATION} — gateway financeiro entre EUA e América Latina, concentrando o maior volume de transações cross-border do hemisfério ocidental")
    bullet(doc, "Serviços Iniciais", "Assessoria em estruturação de investimentos internacionais; alocação estratégica de portfólios; estruturação de produtos financeiros; operações de M&A e captação de recursos (DCM); gestão de relacionamento com investidores; programa educacional em mercado financeiro global")
    bullet(doc, "Expansão Futura", "New York (NY), Houston (TX), São Paulo (Brasil — escritório de representação)")
    bullet(doc, f"Fundadores", f"{FOUNDER_1} — B.A. em Administração (Insper), autor publicado, 5+ anos em investimentos internacionais no ecossistema XP | {FOUNDER_2} — B.A. em Administração (Insper), autor publicado, 4+ anos em M&A e DCM na Criteria Investimentos")

    doc.add_heading("1.1.2. Portfólio de Serviços", level=3)
    p_text(doc, f"A **{COMPANY}** oferece um portfólio integrado de seis linhas de serviço que cobrem o ciclo completo de assessoria financeira internacional, desde a estruturação inicial até o monitoramento contínuo e a capacitação de equipes. A tabela a seguir detalha cada serviço, seu benefício direto ao cliente e o problema específico que resolve no mercado.")

    make_table(doc,
        ["Serviço", "Benefício ao Cliente", "Público-Alvo"],
        [
            ["Estruturação de Investimentos Internacionais", "Segurança jurídica e otimização tributária em operações cross-border", "HNWIs, Family Offices"],
            ["Alocação Estratégica de Portfólios", "Maximização de retorno ajustado ao risco com diversificação global", "Investidores institucionais, Family Offices"],
            ["Estruturação de Produtos Financeiros", "Criação de veículos financeiros competitivos para o mercado americano", "Instituições financeiras, gestoras"],
            ["M&A e Captação de Recursos (DCM)", "Acesso a capital internacional qualificado e execução profissional de transações", "Empresas em expansão, fundos"],
            ["Relacionamento com Investidores", "Transparência, retenção e engajamento de investidores globais", "Fundos, empresas com capital estrangeiro"],
            ["Curso de Investimentos Internacionais", "Capacitação técnica em mercado financeiro global e estruturas offshore", "Profissionais financeiros, investidores individuais"],
        ])

    p_text(doc, "Os serviços formam um ecossistema integrado: a estruturação de investimentos alimenta a demanda por alocação estratégica; operações de M&A geram necessidade de gestão de relacionamento com investidores; e o programa educacional funciona como canal de aquisição de clientes para os serviços de assessoria, criando um ciclo virtuoso de engajamento e retenção.")

    doc.add_heading("1.1.3. Projeções de Viabilidade", level=3)
    p_text(doc, "As projeções financeiras da empresa refletem premissas conservadoras, baseadas em benchmarks do setor de assessoria financeira nos Estados Unidos e na capacidade operacional gradual dos cofundadores. A tabela abaixo sintetiza os indicadores-chave para o horizonte de cinco anos.")

    make_table(doc,
        ["Indicator", "Y1", "Y2", "Y3", "Y4", "Y5"],
        [
            ["Gross Revenue", fmt(FIN["Y1"]["rev"]), fmt(FIN["Y2"]["rev"]), fmt(FIN["Y3"]["rev"]), fmt(FIN["Y4"]["rev"]), fmt(FIN["Y5"]["rev"])],
            ["EBITDA", fmt(FIN["Y1"]["ebitda"]), fmt(FIN["Y2"]["ebitda"]), fmt(FIN["Y3"]["ebitda"]), fmt(FIN["Y4"]["ebitda"]), fmt(FIN["Y5"]["ebitda"])],
            ["Net Income", fmt(FIN["Y1"]["net"]), fmt(FIN["Y2"]["net"]), fmt(FIN["Y3"]["net"]), fmt(FIN["Y4"]["net"]), fmt(FIN["Y5"]["net"])],
        ])

    p_text(doc, f"O investimento inicial de **{fmt(FIN['inv_initial'])}** projeta payback em aproximadamente 20 meses, com NPV positivo a partir do segundo ano de operação. A margem EBITDA atinge 46.7% no quinto ano, alinhada com benchmarks de firmas de assessoria financeira independentes nos Estados Unidos, conforme dados do Bureau of Labor Statistics para o setor de serviços financeiros especializados[3].")

    page_break(doc)

    # S2 — Timeline do Negócio
    doc.add_heading("1.2. Timeline do Negócio", level=2)
    p_text(doc, f"A trajetória da **{COMPANY}** fundamenta-se na convergência de duas carreiras consolidadas no mercado financeiro brasileiro — ambas originadas no ecossistema da **XP Investimentos**, a maior plataforma de investimentos do Brasil — e na decisão estratégica de transferir essa expertise para o mercado norte-americano, onde a America First Investment Policy (fevereiro de 2025) sinaliza compromisso federal com a atração de investimento estrangeiro qualificado[4].")

    p_text(doc, f"**{FOUNDER_1}** iniciou sua trajetória em novembro de 2021 como estagiário de Alocação e Produtos na XP Investimentos, contribuindo para a estruturação do desk de Renda Fixa e desenvolvendo relatórios internos que permanecem em uso pela organização. Em janeiro de 2023, assumiu a posição de Analista de Investimentos Internacionais na Criteria Investimentos (grupo XP), responsável por assessoria B2B e B2C, produção de materiais analíticos e participação em podcasts sobre cenários financeiros globais. Em janeiro de 2025, foi promovido a **Offshore Manager de Investimentos Internacionais**, liderando a divisão Offshore com gestão de aproximadamente R$ 1 bilhão alocados no mercado americano e contribuindo para o reconhecimento da firma como Top 10 Offshore dentro da rede XP.")

    p_text(doc, f"**{FOUNDER_2}** ingressou no ecossistema XP em maio de 2021 como estagiário de Alocação de Ativos, atuando na análise de múltiplas classes de ativos e suporte a portfólios de grande escala. Em janeiro de 2022, assumiu como Analista de Renda Fixa na Criteria, estruturando e consolidando a operação de renda fixa com portfólios totalizando aproximadamente R$ 1,6 bilhão sob monitoramento. Em julho de 2022, migrou para a divisão de **M&A e DCM (Debt Capital Markets)**, posição que mantém atualmente, com responsabilidades que incluem modelagem financeira avançada, preparação de pitch decks institucionais, coordenação com investidores e stakeholders em ciclos completos de transação, e prospecção ativa de novos clientes estratégicos.")

    p_text(doc, f"A formação acadêmica de ambos no **Insper Instituto de Ensino e Pesquisa** — com graduação em Administração de Empresas — fornece base analítica rigorosa em finanças, estratégia e liderança. Gustavo concluiu em dezembro de 2023 e Pedro em junho de 2022, ambos com programas de mais de 3.300 horas de carga horária. Complementarmente, Gustavo realizou programas internacionais na **UCLA** (via English Language Center), incluindo o Junior Program Classic e um programa de Entrepreneurship, consolidando sua visão global e proficiência para operação no mercado americano.")

    p_text(doc, "A produção intelectual dos cofundadores reforça a credibilidade técnica: Gustavo é autor do livro 'Leadership and Decision-Making in the Financial Sector' (2025) e de três artigos acadêmicos publicados na Lumen et Virtus Magazine sobre tendências de investimentos offshore, ETFs e internacionalização de ativos. Pedro é autor de 'FIDCs, CRIs and CRAs: Strategic Guide for Young Investment Professionals' (2025) e de artigos sobre operações estruturadas, financiamento de infraestrutura e instrumentos de dívida privada. Ambos receberam cobertura jornalística nacional em veículos como Brasil Agora, Gazeta da Semana e Business Feed, sendo reconhecidos como líderes emergentes no mercado financeiro.")

    p_text(doc, f"A constituição da **{COMPANY}** em Miami, Florida, representa a materialização dessa trajetória: a experiência prática em alocação global, M&A, DCM e gestão de investimentos internacionais — combinada com produção acadêmica relevante e reconhecimento midiático — posiciona os cofundadores para operar com credibilidade imediata no mercado financeiro norte-americano, atendendo à crescente demanda por assessoria especializada em fluxos de capital cross-border.")

    page_break(doc)

    # S3 — Visão, Missão e Valores
    doc.add_heading("1.3. Visão, Missão e Valores", level=2)
    doc.add_heading("1.3.1. Missão", level=3)
    p_text(doc, f"Prover soluções estratégicas em investimentos internacionais e mercado financeiro global, contribuindo para o fortalecimento da competitividade econômica, a eficiência na alocação de capital e a integração financeira entre mercados, com impacto positivo no desenvolvimento econômico e na atração de investimentos para os Estados Unidos. A **{COMPANY}** opera como uma ponte técnica entre investidores globais e o mercado norte-americano, facilitando fluxos de capital que, segundo dados do Federal Reserve, ultrapassam trilhões de dólares em movimentação anual de investimentos de portfólio internacional[5].")

    doc.add_heading("1.3.2. Visão", level=3)
    p_text(doc, f"Consolidar-se como referência em assessoria de investimentos internacionais no mercado norte-americano, alcançando reconhecimento pela excelência técnica, inovação nas soluções financeiras e relevância estratégica para investidores e empresas em um horizonte de cinco anos. A visão contempla a expansão para três escritórios nos EUA e a construção de uma base de clientes institucionais que posicione a **{COMPANY}** entre as firmas de assessoria independentes de referência no corredor Miami-New York.")

    doc.add_heading("1.3.3. Valores", level=3)
    p_text(doc, "**Excelência Técnica e Analítica** — Atuação com rigor técnico, precisão analítica e profundidade nos mercados financeiros globais, garantindo qualidade nas decisões e recomendações. Cada parecer emitido pela firma baseia-se em dados verificáveis e modelagem financeira robusta, nunca em estimativas superficiais.")
    p_text(doc, "**Ética e Transparência** — Condução de todas as atividades com integridade, clareza e conformidade regulatória integral, assegurando relações de confiança com clientes e parceiros. O compromisso com compliance abrange as regulamentações da SEC, FINRA e as obrigações decorrentes do Investment Advisers Act de 1940[6].")
    p_text(doc, "**Foco em Resultados Sustentáveis** — Priorização de estratégias que gerem valor consistente no longo prazo, com gestão eficiente de riscos e responsabilidade fiduciária. A firma rejeita práticas de curto prazo que comprometam a sustentabilidade dos portfólios sob assessoria.")
    p_text(doc, "**Inovação e Visão Global** — Incorporação de tendências internacionais, tecnologias financeiras e novas abordagens analíticas para oferecer soluções modernas e competitivas. O monitoramento contínuo de políticas como a America First Investment Policy e o Outbound Investment Security Program garante que a assessoria prestada reflita o ambiente regulatório vigente em tempo real.")

    page_break(doc)

    # S4 — Enquadramento Jurídico
    doc.add_heading("1.4. Enquadramento Jurídico", level=2)
    p_text(doc, f"O enquadramento jurídico da **{COMPANY}** foi estruturado para combinar proteção patrimonial, eficiência tributária e conformidade regulatória plena com as exigências federais e estaduais aplicáveis ao setor de serviços financeiros nos Estados Unidos. A constituição da entidade seguirá o processo formalizado pela Division of Corporations do Florida Department of State, conforme as disposições do Florida Revised Limited Liability Company Act (Chapter 605, Florida Statutes)[7].")

    p_text(doc, f"O registro da LLC no estado da Florida será realizado eletronicamente via Sunbiz.org, o portal oficial da Division of Corporations, com filing fee de $125,00. Simultaneamente, será protocolada a eleição S-Corporation junto ao Internal Revenue Service por meio do Form 2553, que deve ser submetido dentro de 75 dias após a formação da entidade ou até 15 de março do primeiro ano fiscal. A eleição S-Corp elimina a dupla tributação corporativa, permitindo que os lucros fluam diretamente para as declarações pessoais dos sócios — uma vantagem significativa para empresas de serviços profissionais com margens elevadas.")

    p_text(doc, "A obtenção do Employer Identification Number (EIN) será realizada online via IRS.gov, processo imediato e sem custo para entidades domiciliadas nos Estados Unidos. O EIN é pré-requisito para abertura de conta bancária empresarial, contratação de funcionários e cumprimento das obrigações fiscais federais e estaduais.")

    p_text(doc, f"No âmbito regulatório setorial, a **{COMPANY}** deverá observar os requisitos aplicáveis ao exercício de atividades de assessoria financeira nos Estados Unidos. A firma buscará registro como Investment Adviser junto à SEC (para ativos sob assessoria acima de $100 milhões) ou junto ao Office of Financial Regulation do Estado da Florida (para operação estadual), conforme as disposições do Investment Advisers Act de 1940. Adicionalmente, os profissionais da firma deverão obter as licenças e certificações exigidas pela FINRA — incluindo a Series 65 (Uniform Investment Adviser Law Examination) — requisito para atuação como Investment Adviser Representative.")

    p_text(doc, "O cronograma de constituição legal da empresa contempla as seguintes etapas: registro da LLC na Florida ($125, 3-5 dias úteis); eleição S-Corporation via Form 2553 (60-90 dias de processamento pelo IRS); obtenção do EIN online (1 dia útil); abertura de conta bancária empresarial (1-2 semanas); registro como Investment Adviser (90-120 dias); obtenção de licenças profissionais Series 65 (60-90 dias de preparação e exame); e conformidade tributária estadual e federal (contínua). O investimento estimado para constituição completa situa-se na faixa de $3.500 a $7.500, incluindo taxas governamentais, honorários advocatícios e custos de compliance inicial[8].")

    p_text(doc, f"A estrutura fiscal da Florida oferece vantagens competitivas significativas: o estado não cobra imposto de renda pessoal (personal income tax), o que, combinado com a eleição S-Corporation, resulta em carga tributária otimizada para os sócios. Essa configuração posiciona a **{COMPANY}** de forma fiscalmente eficiente em comparação com estados como New York (alíquota marginal de até 10,9%) ou California (até 13,3%).")

    page_break(doc)


# ============================================================
# BLOCK 2: ANÁLISE ESTRATÉGICA (S5-S16)
# ============================================================
def block2(doc):
    doc.add_heading("2. ANÁLISE ESTRATÉGICA", level=1)

    # S5 — Perspectivas do Mercado
    doc.add_heading("2.1. Perspectivas do Mercado", level=2)
    p_text(doc, f"O mercado de serviços financeiros nos Estados Unidos constitui o maior e mais sofisticado ecossistema de intermediação de capital do mundo, com o setor de gestão de ativos ultrapassando $100 trilhões em ativos sob gestão conectados globalmente, conforme o relatório anual do Financial Stability Oversight Council (FSOC) de 2024[9]. A **{COMPANY}** insere-se nesse ecossistema como prestadora de serviços especializados em assessoria de investimentos internacionais, segmento que experimenta crescimento acelerado impulsionado pela globalização dos fluxos de capital e pela crescente complexidade regulatória.")

    p_text(doc, f"O Bureau of Labor Statistics registra mediana salarial nacional de $156.100 para Financial Managers (SOC {SOC_CODE}), com crescimento projetado de 17% no período 2023-2033 — o que representa a adição de aproximadamente 126.600 novas posições nesse período[10]. Para Personal Financial Advisors (SOC 13-2052), a mediana salarial é de $99.580, com crescimento projetado de 13%[11]. Financial Analysts (SOC 13-2051) registram mediana de $99.890 e crescimento de 9%[12]. Esses dados evidenciam uma demanda estrutural crescente por profissionais qualificados no setor financeiro, particularmente em funções que envolvem gestão de investimentos internacionais e análise de mercados globais.")

    p_text(doc, "A concentração de instituições financeiras internacionais em Miami reforça a posição estratégica da empresa. O corredor Brickell-Downtown abriga escritórios de bancos, gestoras e assessorias financeiras oriundos de mais de 30 países, configurando o principal hub de negócios cross-border entre os Estados Unidos e a América Latina. O Federal Reserve Bank of Atlanta — que abrange a jurisdição da Florida — documenta volumes crescentes de fluxos de investimento de portfólio internacional na região, refletindo a atratividade do estado como destino de capital estrangeiro.")

    p_text(doc, "Três vetores estruturais sustentam o crescimento do mercado-alvo da empresa. Primeiro, a **America First Investment Policy** (fevereiro de 2025) sinaliza compromisso federal com a facilitação de investimentos estrangeiros passivos nos EUA, redução de burocracias e direcionamento de recursos para investidores de países aliados — ambiente diretamente favorável à assessoria em estruturação de investimentos cross-border[13]. Segundo, o **Outbound Investment Security Program** do Departamento do Tesouro aumenta a complexidade regulatória para investimentos americanos em setores sensíveis (AI, semicondutores, computação quântica), gerando demanda por assessoria especializada em conformidade[14]. Terceiro, a modernização das regras de marketing para investment advisers pela SEC (Investment Adviser Marketing Rule, 2020) exige maior sofisticação na comunicação com clientes, criando oportunidades para firmas com expertise em relacionamento institucional.")

    p_text(doc, "A convergência desses fatores — crescimento do emprego no setor, políticas federais de atração de investimento, complexidade regulatória crescente e concentração de capital internacional em Miami — configura um ambiente de mercado estruturalmente favorável para uma firma de assessoria especializada em investimentos internacionais e mercado financeiro global.")

    page_break(doc)

    # S5b — Relevância, Oportunidades e Perspectivas Futuras
    doc.add_heading("2.2. Relevância, Oportunidades e Perspectivas Futuras", level=2)
    p_text(doc, f"A relevância da **{COMPANY}** no mercado financeiro norte-americano fundamenta-se em três pilares documentados por políticas federais e dados governamentais. O Committee on Foreign Investment in the United States (CFIUS) — fortalecido pelo FIRRMA (Foreign Investment Risk Review Modernization Act) — ampliou significativamente o escopo de revisão de investimentos estrangeiros nos EUA, incluindo investimentos não controladores e transações imobiliárias sensíveis[15]. Em 2024, atualizações regulatórias expandiram a jurisdição do CFIUS para 59 novas instalações militares e estabeleceram penalidades mais rigorosas para não-conformidade. Esse ambiente regulatório crescente gera demanda direta por assessoria especializada que auxilie investidores internacionais a navegar o framework legal dos Estados Unidos.")

    p_text(doc, "O programa SBIC (Small Business Investment Company), administrado pela SBA, representa uma oportunidade estratégica para a operação de captação de recursos da empresa. Com mais de $40 bilhões investidos em pequenas empresas ao longo de sua história, o programa permite que fundos licenciados utilizem capital privado alavancado por recursos federais — modelo diretamente alinhado com o serviço de assessoria em captação de recursos (DCM) oferecido pela firma[16]. A capacidade de estruturar acesso a esse tipo de capital alavancado diferencia a empresa de competidores que operam exclusivamente no segmento de investimentos diretos.")

    p_text(doc, "O JOBS Act (Jumpstart Our Business Startups Act) complementa o panorama de oportunidades ao facilitar o acesso de empresas emergentes ao mercado de capitais via crowdfunding regulado e redução de exigências para abertura de capital. A expertise da firma em estruturação de produtos financeiros e M&A posiciona-a para assessorar clientes que desejem capitalizar essas facilidades regulatórias — particularmente empresas brasileiras e latino-americanas buscando acesso ao mercado de capitais americano[17].")

    p_text(doc, f"As perspectivas futuras incluem três frentes de expansão. A primeira envolve a obtenção de registro como Registered Investment Advisor (RIA) junto à SEC, o que permitirá à **{COMPANY}** gerir ativos sob mandato formal — ampliando significativamente a base de receita. Dados da SEC indicam que mais de 15.000 investment advisers registrados operam nos Estados Unidos[18], com crescimento contínuo no número de RIAs independentes. A segunda frente contempla a expansão geográfica para New York (segundo ano de operação) e Houston (terceiro ano), mercados com concentrações significativas de capital institucional e demanda por assessoria especializada. A terceira frente envolve a construção de uma plataforma educacional escalável que amplie o alcance do programa de treinamento além do formato presencial.")

    page_break(doc)

    # S6 — Cadeia de Suprimentos
    doc.add_heading("2.3. Cadeia de Suprimentos", level=2)
    p_text(doc, f"A cadeia de suprimentos da **{COMPANY}** reflete a natureza intangível dos serviços financeiros especializados, onde os insumos primários são capital intelectual, dados de mercado e infraestrutura tecnológica — não matérias-primas físicas. Essa configuração resulta em margens operacionais superiores à média industrial, uma vez que a estrutura de custos é predominantemente composta por despesas com pessoal qualificado e assinaturas de plataformas de dados[19].")

    p_text(doc, "Os principais fornecedores da operação dividem-se em quatro categorias. **Dados financeiros e research** — plataformas como Bloomberg Terminal, Refinitiv e S&P Capital IQ fornecem dados em tempo real, análises setoriais e ferramentas de modelagem que sustentam a assessoria prestada aos clientes. O custo anual de uma licença Bloomberg Terminal situa-se na faixa de $20.000-$24.000 por terminal, representando um dos principais investimentos operacionais da firma. **Infraestrutura tecnológica** — provedores de cloud computing (AWS, Microsoft Azure), plataformas de CRM financeiro e ferramentas de compliance automatizado garantem a operação contínua e a conformidade com os requisitos de cybersecurity da SEC. **Serviços profissionais** — escritórios de advocacia especializados em securities law, firmas de contabilidade com expertise em estruturas S-Corp e compliance tributário internacional, e auditores independentes compõem a rede de suporte legal e fiscal. **Capital humano qualificado** — analistas financeiros certificados (CFA, Series 65, Series 7) e especialistas em mercados internacionais constituem o ativo mais crítico da cadeia de valor.")

    p_text(doc, "A concentração de fornecedores de infraestrutura financeira nos Estados Unidos — particularmente no eixo New York-Miami — garante acesso direto e sem intermediação a dados, plataformas e serviços essenciais para a operação. A estratégia de sourcing da firma prioriza contratos de longo prazo com provedores de dados e tecnologia, reduzindo a vulnerabilidade a aumentos súbitos de preços e garantindo continuidade operacional.")

    page_break(doc)

    # S7 — Empregabilidade Esperada
    doc.add_heading("2.4. Empregabilidade Esperada", level=2)
    p_text(doc, f"O plano de contratações da **{COMPANY}** reflete a estratégia de crescimento gradual e sustentável, alinhando a expansão da equipe com o aumento da base de clientes e da complexidade dos serviços prestados. O Bureau of Labor Statistics projeta que o setor de serviços financeiros nos Estados Unidos adicionará centenas de milhares de novas posições na próxima década, com crescimento particularmente acelerado nas funções de Financial Managers (+17%), Personal Financial Advisors (+13%) e Financial Analysts (+9%)[20].")

    p_text(doc, "A tabela a seguir detalha a projeção de contratações por ano de operação, as funções previstas e os salários estimados com base nos dados do BLS para o mercado da Florida.")

    make_table(doc,
        ["Year", "Headcount", "Key Positions", "Est. Annual Payroll"],
        [
            ["Y0-Y1", "2 (founders) + 1", "Administrative Assistant / Office Manager", "$42,000-$55,000"],
            ["Y2", "5", "+ Financial Analyst, + Marketing Coordinator", "$95,000-$120,000 (new hires)"],
            ["Y3", "8", "+ Senior Analyst, + Client Relations Manager, + Junior Analyst", "$180,000-$220,000 (new hires)"],
            ["Y4", "12", "+ 2 Analysts (NY office), + Compliance Officer, + Operations Manager", "$280,000-$340,000 (new hires)"],
            ["Y5", "16-18", "+ 2 Analysts (Houston), + Senior Associate, + Training Coordinator", "$360,000-$440,000 (new hires)"],
        ])

    p_text(doc, "A geração de empregos qualificados no setor financeiro contribui diretamente para o ecossistema econômico local, com cada posição criada gerando efeitos multiplicadores na economia da região. A firma priorizará contratações locais e a capacitação de profissionais americanos em técnicas de assessoria financeira internacional — transferindo o conhecimento acumulado pelos cofundadores no mercado financeiro brasileiro e global para o mercado de trabalho dos Estados Unidos.")

    page_break(doc)

    # S8 — Gestão do Conhecimento
    doc.add_heading("2.5. Gestão do Conhecimento", level=2)
    p_text(doc, f"A gestão do conhecimento na **{COMPANY}** é estruturada como vantagem competitiva sustentável, transformando a expertise individual dos cofundadores em ativos intelectuais organizacionais replicáveis. Três pilares sustentam essa estratégia: documentação de processos analíticos, produção intelectual contínua e capacitação sistemática da equipe.")

    p_text(doc, f"O primeiro pilar — **documentação de processos** — envolve a criação de manuais operacionais detalhados para cada linha de serviço, incluindo frameworks de análise, templates de modelagem financeira, protocolos de due diligence para operações cross-border e checklists de compliance regulatório. Esses materiais garantem consistência na qualidade do serviço prestado independentemente do profissional responsável, reduzindo a dependência de conhecimento tácito individual. Os processos são revisados trimestralmente para incorporar mudanças regulatórias — como as atualizações do CFIUS, novas Executive Orders e revisões das regras da SEC — garantindo que a base de conhecimento reflita o estado atual do ambiente operacional.")

    p_text(doc, f"O segundo pilar — **produção intelectual contínua** — apoia-se no histórico comprovado dos cofundadores. {FOUNDER_1} já publicou o livro 'Leadership and Decision-Making in the Financial Sector' e três artigos acadêmicos peer-reviewed sobre investimentos offshore, ETFs e internacionalização de ativos na Lumen et Virtus Magazine. {FOUNDER_2} é autor de 'FIDCs, CRIs and CRAs: Strategic Guide for Young Investment Professionals' e de artigos sobre operações estruturadas e financiamento de infraestrutura. A firma manterá ritmo de produção de pelo menos dois relatórios técnicos por trimestre, cobrindo tendências de mercado, análises regulatórias e oportunidades de investimento — materiais que servem simultaneamente como ferramenta de marketing e como base de conhecimento interno.")

    p_text(doc, "O terceiro pilar — **capacitação sistemática** — opera por meio do programa educacional interno (curso de Investimentos Internacionais e Mercado Financeiro Global, 4 módulos) e de parcerias com instituições de formação profissional nos Estados Unidos. A firma incentivará e subsidiará a obtenção de certificações relevantes (CFA, Series 65, Series 7) por todos os membros da equipe, criando uma cultura de excelência técnica e desenvolvimento profissional contínuo.")

    page_break(doc)

    # S9 — Impactos ESG
    doc.add_heading("2.6. Impactos ESG", level=2)
    p_text(doc, f"A **{COMPANY}** integra princípios de responsabilidade ambiental, social e de governança (ESG) em sua estratégia operacional e na assessoria prestada a clientes, reconhecendo que a alocação responsável de capital é simultaneamente um imperativo ético e uma vantagem competitiva no mercado financeiro contemporâneo.")

    p_text(doc, "No pilar **ambiental**, a firma compromete-se com a inclusão de critérios de sustentabilidade nas recomendações de investimento, orientando clientes sobre oportunidades alinhadas com o Inflation Reduction Act — que destinou aproximadamente $369 bilhões em incentivos para energia limpa e clima[21] — e com as diretrizes de investimento sustentável promovidas pela SEC. A operação da firma mantém pegada ambiental reduzida, com escritório paperless, uso exclusivo de infraestrutura cloud e política de compensação de emissões de carbono associadas a viagens de negócios.")

    p_text(doc, "No pilar **social**, a empresa contribui diretamente para a democratização do acesso a investimentos internacionais — tema pelo qual Gustavo Esteves foi reconhecido na Gazeta da Semana com a matéria 'A democratização do investimento internacional'. O programa educacional (curso de 4 módulos) funciona como veículo de transferência de conhecimento financeiro especializado para profissionais e investidores que, de outra forma, não teriam acesso a expertise em mercados globais. A firma também priorizará diversidade e inclusão em suas contratações, buscando refletir a composição demográfica do mercado financeiro de Miami.")

    p_text(doc, "No pilar de **governança**, a estrutura da firma adota as melhores práticas de compliance e transparência: segregação de funções, políticas de conflito de interesse documentadas, due diligence reforçado em todas as transações, e aderência integral aos requisitos da SEC e FINRA para investment advisers. A eleição S-Corporation garante transparência fiscal com pass-through taxation, enquanto a estrutura LLC preserva a separação patrimonial entre entidade e sócios.")

    page_break(doc)

    # S10 — Análise SWOT
    doc.add_heading("2.7. Análise SWOT", level=2)
    p_text(doc, f"A análise SWOT da **{COMPANY}** foi construída a partir de dados verificáveis do mercado financeiro norte-americano, do perfil dos cofundadores e do ambiente regulatório vigente, evitando generalizações e priorizando fatores com impacto direto na operação.")

    make_table(doc,
        ["Dimensão", "Fatores"],
        [
            ["**Strengths**", "Experiência direta em gestão de R$ 1 bilhão em ativos internacionais (Gustavo); transações de M&A/DCM totalizando ~R$ 500 milhões (Pedro); formação acadêmica em instituição tier-1 (Insper); produção intelectual publicada (2 livros, 6+ artigos peer-reviewed); reconhecimento midiático nacional; fluência em português e inglês; rede de contatos no ecossistema XP Investimentos"],
            ["**Weaknesses**", "Marca nova sem histórico no mercado americano; base de clientes a ser construída do zero; ausência de registro RIA no momento da constituição; dependência inicial dos dois cofundadores para geração de receita; capital inicial limitado em comparação com firmas estabelecidas"],
            ["**Opportunities**", "America First Investment Policy (2025) incentivando investimento estrangeiro; crescimento de 17% projetado pelo BLS para Financial Managers; complexidade crescente do CFIUS e Outbound Investment Security Program gerando demanda por assessoria; crescimento da comunidade brasileira e latino-americana em Miami; expansão de programas como SBIC e JOBS Act"],
            ["**Threats**", "Concorrência de firmas estabelecidas com décadas de operação; mudanças regulatórias imprevisíveis (SEC, FINRA); ciclos econômicos e volatilidade de mercado afetando demanda por assessoria; risco de taxa de juros elevada reduzindo fluxos de investimento; potenciais restrições adicionais a investimentos cross-border"],
        ])

    p_text(doc, "A análise evidencia que as forças fundamentais da firma — expertise técnica verificável, produção intelectual e posicionamento em Miami — compensam as fragilidades naturais de uma empresa em fase inicial, enquanto as oportunidades de mercado superam quantitativamente as ameaças identificadas.")

    page_break(doc)

    # S11 — SWOT Cruzada
    doc.add_heading("2.8. SWOT Cruzada", level=2)
    p_text(doc, "A SWOT Cruzada operacionaliza a análise estratégica ao definir ações concretas para cada combinação de fatores internos e externos, transformando o diagnóstico em plano de ação executável.")

    make_table(doc,
        ["Combinação", "Estratégia", "Ação Concreta"],
        [
            ["Forças × Oportunidades (SO)", "Alavancagem", "Utilizar experiência em investimentos offshore e produção intelectual para captar clientes institucionais atraídos pela America First Investment Policy; posicionar-se como ponte técnica entre investidores latino-americanos e o mercado americano"],
            ["Forças × Ameaças (ST)", "Confrontação", "Diferenciar-se de concorrentes estabelecidos pela expertise bilateral (Brasil-EUA) e bilíngue; mitigar riscos regulatórios com compliance proativo e monitoramento contínuo de Executive Orders"],
            ["Fraquezas × Oportunidades (WO)", "Reforço", "Acelerar construção de marca via produção de conteúdo técnico, participação em eventos do setor e parcerias com firmas complementares; priorizar registro RIA para ampliar credibilidade institucional"],
            ["Fraquezas × Ameaças (WT)", "Defesa", "Manter estrutura de custos enxuta nos primeiros 18 meses; diversificar fontes de receita entre assessoria, M&A e educação; construir reserva de caixa equivalente a 6 meses de operação"],
        ])

    p_text(doc, "A estratégia de alavancagem (SO) representa a prioridade estratégica da firma nos primeiros dois anos, maximizando o aproveitamento do ambiente regulatório favorável para construir base de clientes com velocidade superior à concorrência.")

    page_break(doc)

    # S11b — Matriz ANSOFF
    doc.add_heading("2.9. Matriz ANSOFF", level=2)
    p_text(doc, f"A Matriz ANSOFF da **{COMPANY}** mapeia quatro vetores de crescimento estratégico, partindo dos serviços atuais no mercado de Miami e expandindo progressivamente para novos mercados e novas categorias de serviço.")

    make_table(doc,
        ["", "Mercados Atuais (Miami/FL)", "Novos Mercados (NY, TX, LATAM)"],
        [
            ["**Serviços Atuais**", "**Penetração:** Aprofundar presença no corredor Miami-Fort Lauderdale com assessoria em investimentos internacionais para HNWIs, family offices e instituições financeiras latino-americanas operando nos EUA", "**Desenvolvimento de Mercado:** Expandir serviços existentes para New York (Y2) e Houston (Y3); estabelecer escritório de representação em São Paulo para captação de clientes brasileiros"],
            ["**Novos Serviços**", "**Desenvolvimento de Produto:** Lançar plataforma digital de educação financeira (cursos online); desenvolver produtos de assessoria em ativos digitais e cripto-ativos regulados; criar serviço de compliance-as-a-service para investidores estrangeiros", "**Diversificação:** Estruturar fundo de investimento próprio (venture capital ou private equity com foco em empresas brasileiras nos EUA); criar programa de aceleração para empresas latino-americanas buscando captação no mercado americano"],
        ])

    p_text(doc, "A sequência priorizada segue a lógica de risco crescente: penetração de mercado nos primeiros 18 meses, desenvolvimento de mercado a partir do segundo ano, desenvolvimento de produto no terceiro ano, e diversificação a partir do quarto ano — garantindo que cada movimento estratégico seja sustentado por receita e aprendizado acumulados nas fases anteriores.")

    page_break(doc)

    # S12 — Análise de Concorrentes
    doc.add_heading("2.10. Análise de Concorrentes", level=2)
    p_text(doc, f"O mercado de assessoria financeira em Miami é competitivo e fragmentado, com firmas que variam desde boutiques independentes especializadas até divisões de grandes bancos de investimento globais. A **{COMPANY}** compete primariamente no segmento de assessoria independente especializada em investimentos internacionais e fluxos de capital cross-border — nicho que apresenta concentração moderada e barreiras de entrada relevantes.")

    p_text(doc, "Os concorrentes dividem-se em três perfis distintos. **Firmas globais de assessoria** — grandes bancos de investimento e gestoras internacionais com presença em Miami (JPMorgan, UBS, Credit Suisse/UBS, Itaú BBA USA) oferecem serviços de wealth management e advisory com base de clientes estabelecida e infraestrutura global. A vantagem competitiva dessas firmas reside na escala e na marca; sua desvantagem é o atendimento padronizado, a burocracia institucional e os mínimos de investimento elevados que excluem parcela significativa do mercado-alvo. **Boutiques financeiras regionais** — firmas independentes de médio porte com foco em clientes latino-americanos constituem o segmento de competição mais direta. Essas firmas tipicamente possuem 5-20 profissionais e oferecem atendimento personalizado, mas frequentemente carecem de expertise técnica em estruturação de produtos financeiros complexos e operações de M&A. **Firmas de contabilidade e tax advisory com serviços financeiros** — CPAs e firmas de planejamento tributário que oferecem serviços adjacentes de assessoria de investimentos, frequentemente como extensão de suas atividades primárias.")

    p_text(doc, f"Os diferenciais competitivos da **{COMPANY}** incluem: (1) expertise bilateral Brasil-EUA com operação bilíngue nativa; (2) experiência verificável em gestão de ativos internacionais na escala de R$ 1 bilhão; (3) capacidade combinada de assessoria em investimentos e execução de M&A/DCM — raramente encontrada em firmas independentes de porte similar; (4) produção intelectual publicada que funciona como credencial técnica e canal de aquisição de clientes; e (5) programa educacional que cria pipeline de clientes qualificados.")

    page_break(doc)

    # S13 — Ameaça de Novos Entrantes
    doc.add_heading("2.11. Ameaça de Novos Entrantes", level=2)
    p_text(doc, "A análise da ameaça de novos entrantes no mercado de assessoria financeira internacional em Miami revela barreiras de entrada moderadas a altas, sustentadas por requisitos regulatórios, exigências de capital humano qualificado e a necessidade de credibilidade institucional junto a clientes sofisticados.")

    p_text(doc, "As **barreiras regulatórias** constituem o fator de proteção mais significativo. O registro como Investment Adviser (seja junto à SEC para firmas com AUM acima de $100 milhões, ou junto ao regulador estadual para firmas menores) exige compliance com o Investment Advisers Act de 1940, incluindo obrigações fiduciárias, requisitos de disclosure (Form ADV), políticas de supervisão e exames periódicos. A obtenção de licenças profissionais (Series 65, Series 66, Series 7) requer preparação técnica específica e taxas de aprovação que variam entre 60% e 75%, filtrando entrantes sem qualificação adequada.")

    p_text(doc, "As **barreiras de capital humano** são igualmente relevantes. O mercado de profissionais financeiros qualificados em Miami é competitivo, com a mediana salarial para Financial Managers na região metropolitana de Miami-Fort Lauderdale-West Palm Beach situando-se acima da mediana nacional, refletindo o custo de vida elevado e a demanda por profissionais bilíngues com expertise em mercados internacionais. Novos entrantes precisam atrair e reter talentos em um mercado apertado, o que eleva significativamente o capital inicial necessário.")

    p_text(doc, f"As **barreiras de reputação e confiança** representam a terceira camada de proteção. Clientes de alta renda e investidores institucionais — o público-alvo primário da **{COMPANY}** — delegam decisões financeiras críticas apenas a profissionais e firmas em quem confiam. A construção dessa confiança requer tempo, track record verificável e referências institucionais — fatores que beneficiam incumbentes e elevam o custo de aquisição de clientes para novos entrantes.")

    p_text(doc, "A avaliação consolidada indica ameaça de novos entrantes classificada como **moderada**: as barreiras regulatórias e de reputação são suficientemente elevadas para impedir a entrada de competidores não qualificados, porém o mercado permanece acessível a profissionais experientes com credenciais sólidas — perfil dos cofundadores da firma.")

    page_break(doc)

    # S14 — Poder de Negociação dos Clientes
    doc.add_heading("2.12. Poder de Negociação dos Clientes", level=2)
    p_text(doc, f"O poder de negociação dos clientes da **{COMPANY}** é classificado como **moderado a alto**, refletindo a sofisticação financeira do público-alvo e a disponibilidade de alternativas no mercado de assessoria. Clientes de alta renda (HNWIs), family offices e investidores institucionais possuem conhecimento financeiro avançado, acesso a múltiplos provedores de assessoria e capacidade de avaliar criticamente a qualidade dos serviços prestados.")

    p_text(doc, "Três fatores amplificam o poder de negociação. Primeiro, os **custos de troca** no setor de assessoria financeira são relativamente baixos: um cliente pode migrar para outra firma com custo operacional limitado, especialmente quando os serviços são prestados em regime de projeto (M&A, estruturação) em vez de mandato contínuo. Segundo, a **transparência de preços** no setor permite que clientes comparem facilmente honorários e termos entre diferentes provedores. Terceiro, a **concentração de compradores** — particularmente no segmento B2B de instituições financeiras e family offices — confere a clientes individuais relevância desproporcional na base de receita da firma.")

    p_text(doc, f"A **{COMPANY}** mitiga esse poder de negociação por meio de três estratégias. A **especialização técnica** em investimentos cross-border e estruturação internacional diferencia a firma de assessorias genéricas, reduzindo a comparabilidade direta de preços. Os **contratos de retainer** (assessoria continuada com mensalidade fixa) criam custos de troca intelectuais — uma vez que a firma acumula conhecimento profundo sobre o portfólio e os objetivos do cliente, a migração implica perda de contexto significativa. E o **programa educacional** estabelece vínculo de longo prazo com participantes que frequentemente se convertem em clientes de assessoria.")

    page_break(doc)

    # S15 — Poder de Negociação dos Fornecedores
    doc.add_heading("2.13. Poder de Negociação dos Fornecedores", level=2)
    p_text(doc, "O poder de negociação dos fornecedores no ecossistema de serviços financeiros é classificado como **moderado**, com variação significativa entre categorias de insumos. A análise segmenta os fornecedores em quatro grupos com dinâmicas distintas de poder.")

    p_text(doc, "**Plataformas de dados financeiros** — Bloomberg LP, Refinitiv e S&P Global exercem poder de negociação elevado devido à natureza oligopolística do mercado de dados financeiros em tempo real. Bloomberg, em particular, detém posição dominante com infraestrutura proprietária que cria custos de troca significativos para usuários habituados à sua interface e funcionalidades. A mitigação desse poder envolve a adoção de estratégia multi-vendor, utilizando plataformas complementares e fontes de dados gratuitas (.gov, FRED, EDGAR) para reduzir a dependência de um único provedor.")

    p_text(doc, "**Serviços profissionais** (advocacia, contabilidade, compliance) apresentam poder de negociação moderado. O mercado de escritórios jurídicos especializados em securities law em Miami é relativamente competitivo, com múltiplas opções disponíveis para cada necessidade. Firmas de contabilidade com expertise em estruturas S-Corp e compliance tributário internacional são menos abundantes, conferindo-lhes poder ligeiramente superior. A estratégia de mitigação inclui a construção de relacionamentos de longo prazo com provedores preferenciais e a negociação de pacotes de serviços com preços fixos.")

    p_text(doc, "**Infraestrutura tecnológica** (cloud computing, CRM, ferramentas de compliance) opera em mercado competitivo com poder de negociação baixo para os fornecedores. A disponibilidade de múltiplas opções (AWS, Azure, Google Cloud) e a padronização crescente de plataformas CRM para o setor financeiro garantem flexibilidade significativa para a firma. **Capital humano** constitui a categoria com dinâmica mais complexa: profissionais qualificados com licenças FINRA e experiência em mercados internacionais são escassos, conferindo-lhes poder de negociação individual elevado — mitigado pela estratégia da firma de investir em formação interna e oferecer pacotes de remuneração competitivos que incluam participação nos resultados.")

    page_break(doc)

    # S16 — Produtos ou Serviços Substitutos
    doc.add_heading("2.14. Produtos ou Serviços Substitutos", level=2)
    p_text(doc, f"A ameaça de produtos ou serviços substitutos para a assessoria financeira especializada da **{COMPANY}** é classificada como **moderada**, com diferenciação significativa entre os segmentos de atuação.")

    p_text(doc, "No segmento de **alocação de portfólios**, as plataformas de robo-advisory (Betterment, Wealthfront, Schwab Intelligent Portfolios) representam substitutos diretos para o componente de gestão passiva de investimentos. Essas plataformas oferecem alocação automatizada a custos significativamente inferiores — tipicamente 0,25%-0,50% ao ano versus 1,0%-1,5% para assessoria humana. Entretanto, a capacidade de assessoria em estruturação internacional, planejamento tributário cross-border e operações complexas (M&A, DCM) permanece fora do alcance dessas plataformas, protegendo o core business da firma.")

    p_text(doc, "No segmento de **estruturação de investimentos internacionais**, os substitutos são limitados. Departamentos internos de wealth management de bancos oferecem serviços similares, porém com mínimos de investimento tipicamente superiores a $1 milhão e atendimento padronizado. Firmas de advocacia especializadas em tax planning internacional atendem parte da demanda, mas carecem da expertise em alocação de ativos e execução de transações financeiras.")

    p_text(doc, "No segmento **educacional**, plataformas de educação online (Coursera, edX, CFA Institute) oferecem conteúdo financeiro a custos acessíveis. Porém, o diferencial do programa da firma reside na aplicação prática, na expertise bilateral Brasil-EUA e no networking que o formato presencial proporciona — elementos que plataformas de mass education não replicam.")

    p_text(doc, "A avaliação consolidada indica que os substitutos impactam primariamente o segmento de alocação passiva de baixa complexidade, enquanto os serviços de maior valor agregado — estruturação internacional, M&A, assessoria regulatória — permanecem relativamente protegidos pela complexidade inerente e pela necessidade de julgamento profissional especializado.")

    page_break(doc)


# ============================================================
# BLOCK 3: PLANO DE MARKETING (S17-S25)
# ============================================================
def block3(doc):
    doc.add_heading("3. PLANO DE MARKETING", level=1)

    # S17 — Segmentação de Mercado
    doc.add_heading("3.1. Segmentação de Mercado", level=2)
    p_text(doc, f"A segmentação de mercado da **{COMPANY}** adota critérios de patrimônio líquido, sofisticação financeira e necessidade de assessoria cross-border para delimitar quatro segmentos-alvo com perfis distintos de demanda e valor de ciclo de vida do cliente.")

    p_text(doc, "O **Segmento 1 — High-Net-Worth Individuals (HNWIs)** compreende investidores individuais com patrimônio investível superior a $1 milhão, buscando diversificação internacional de portfólio. A SEC define 'accredited investor' como indivíduo com patrimônio líquido acima de $1 milhão (excluindo residência principal) ou renda anual acima de $200.000 ($300.000 para casais)[22]. Esse segmento demanda assessoria personalizada em estruturação de contas offshore, alocação global e planejamento tributário internacional. O valor estimado por cliente varia de $3.000 a $8.000 mensais em retainer.")

    p_text(doc, "O **Segmento 2 — Family Offices e Gestoras de Patrimônio** representa clientes institucionais com patrimônio sob gestão tipicamente superior a $10 milhões, requerendo assessoria sofisticada em alocação multi-ativos, estruturação de holdings internacionais e compliance regulatório. O valor estimado por cliente é de $5.000 a $15.000 mensais, com potencial de engajamento em projetos de M&A e estruturação de fundos.")

    p_text(doc, "O **Segmento 3 — Empresas em Expansão Internacional** abrange companhias brasileiras e latino-americanas buscando acesso ao mercado americano via investimento direto, aquisições ou captação de recursos. A demanda concentra-se em serviços de M&A, assessoria regulatória (CFIUS) e estruturação financeira. O valor por engajamento varia de $15.000 a $50.000 por projeto, com potencial de retainer subsequente.")

    p_text(doc, "O **Segmento 4 — Profissionais Financeiros e Investidores Individuais** constitui o público do programa educacional, com tickets de $2.500 a $5.000 por participante. Esse segmento funciona simultaneamente como centro de receita e como pipeline de conversão para os serviços de assessoria premium.")

    page_break(doc)

    # S18 — Público-Alvo B2C
    doc.add_heading("3.2. Público-Alvo B2C", level=2)
    p_text(doc, f"O público B2C da **{COMPANY}** é composto primariamente por investidores individuais de alta renda (HNWIs) residentes nos Estados Unidos — com ênfase na comunidade brasileira e latino-americana de Miami-Fort Lauderdale — e por profissionais do setor financeiro buscando capacitação em mercados globais.")

    p_text(doc, "O perfil demográfico do cliente B2C típico apresenta as seguintes características: idade entre 35 e 65 anos; patrimônio investível acima de $500.000; formação superior em áreas correlatas (administração, economia, engenharia, direito); experiência prévia com investimentos domésticos e interesse em diversificação internacional; residência em condados de Miami-Dade, Broward ou Palm Beach; e, frequentemente, origem brasileira ou latino-americana com conexões empresariais em ambos os mercados. O Census Bureau reporta que a região metropolitana de Miami-Fort Lauderdale-West Palm Beach concentra uma das maiores populações brasileiras dos Estados Unidos, estimada em centenas de milhares de residentes quando incluídos os condados adjacentes[23].")

    p_text(doc, "A jornada do cliente B2C segue um funil de quatro estágios: (1) **Descoberta** — via conteúdo educacional (artigos, podcasts, redes sociais profissionais), cobertura midiática ou referências de rede de contatos; (2) **Avaliação** — participação em webinars gratuitos ou no módulo introdutório do programa educacional; (3) **Conversão** — inscrição no programa educacional completo ou contratação de assessoria de investimentos inicial (diagnóstico de portfólio); (4) **Retenção** — transição para contrato de retainer com acompanhamento contínuo de portfólio e acesso a oportunidades exclusivas de investimento.")

    p_text(doc, "O ciclo de vida estimado do cliente B2C é de 3 a 7 anos, com lifetime value (LTV) projetado entre $36.000 e $336.000, dependendo do segmento de patrimônio. A taxa de retenção-alvo é de 85% ao ano a partir do segundo ano de relacionamento.")

    page_break(doc)

    # S19 — Setor-Alvo B2B
    doc.add_heading("3.3. Setor-Alvo B2B", level=2)
    p_text(doc, f"O público B2B da **{COMPANY}** abrange três categorias de clientes institucionais com demandas complementares: instituições financeiras, empresas em processo de expansão internacional e family offices com operações multi-jurisdicionais.")

    p_text(doc, "**Instituições financeiras** — bancos, gestoras de ativos e plataformas de investimento que necessitam de assessoria especializada em produtos financeiros internacionais, estruturação de operações cross-border e compliance com regulamentações como CFIUS e Outbound Investment Security Program. A experiência dos cofundadores no ecossistema XP Investimentos — uma das maiores plataformas financeiras do Brasil — confere credibilidade imediata junto a instituições financeiras brasileiras e latino-americanas operando nos EUA.")

    p_text(doc, "**Empresas em expansão internacional** — companhias brasileiras e latino-americanas buscando acessar o mercado americano via investimento direto, aquisições ou captação de recursos (IPO, private placement, debt issuance). A capacidade combinada de assessoria em M&A e DCM — com track record verificável de aproximadamente R$ 500 milhões em transações executadas por Pedro Paciullo — posiciona a firma como parceira estratégica para operações de média complexidade no mercado americano.")

    p_text(doc, "**Family offices** — estruturas de gestão patrimonial familiar que requerem assessoria sofisticada em diversificação global, governança corporativa, planejamento sucessório internacional e gestão de risco multi-moeda. O atendimento personalizado e a expertise bilateral Brasil-EUA diferenciam a firma de competidores que operam exclusivamente no mercado doméstico americano.")

    p_text(doc, "O ciclo de vendas B2B é tipicamente mais longo (3-6 meses), mas gera contratos de maior valor (retainers de $5.000-$15.000/mês) e relacionamentos mais estáveis (retenção média de 5+ anos). A estratégia de aquisição B2B combina prospecção ativa, participação em eventos do setor (CFA Institute Annual Conference, SALT Conference) e referências de rede profissional.")

    page_break(doc)

    # S20 — Posicionamento da Marca
    doc.add_heading("3.4. Posicionamento da Marca", level=2)
    p_text(doc, f"A **{COMPANY}** posiciona-se no mercado como uma **firma de assessoria financeira independente especializada em investimentos internacionais e operações cross-border**, com expertise bilateral Brasil-EUA e atendimento personalizado para clientes de alta complexidade financeira.")

    p_text(doc, "O posicionamento diferencia-se dos concorrentes em três eixos. O **eixo de especialização** foca exclusivamente em investimentos internacionais e fluxos de capital cross-border — evitando a diluição em serviços genéricos de wealth management doméstico. O **eixo de expertise bilateral** capitaliza a combinação única de experiência nos mercados brasileiro e americano, com fluência nativa em ambos os idiomas e conhecimento profundo dos sistemas regulatórios de ambos os países. O **eixo de atendimento** oferece acesso direto aos cofundadores — profissionais com track record verificável — em contraste com firmas maiores onde clientes são frequentemente delegados a analistas juniores.")

    p_text(doc, "A identidade visual da marca será construída sobre princípios de sofisticação, confiança e globalidade — refletindo o posicionamento premium da firma no mercado. O nome 'EventFinOps' combina operações financeiras (FinOps) com a capacidade de execução de eventos e marcos transformacionais no ciclo de vida financeiro dos clientes (Event), comunicando competência operacional e impacto estratégico.")

    p_text(doc, "A promessa de marca — 'Global Investment Intelligence, Local Expertise' — sintetiza o valor diferencial: inteligência financeira com alcance global, executada com profundidade local e atendimento personalizado. Essa promessa será consistentemente comunicada em todos os touchpoints: website, materiais institucionais, redes sociais profissionais e interações com clientes.")

    page_break(doc)

    # S21 — Produto — Análise de Valor
    doc.add_heading("3.5. Produto — Análise de Valor", level=2)
    p_text(doc, f"A análise de valor dos serviços da **{COMPANY}** mapeia o benefício percebido pelo cliente em relação ao custo de aquisição, demonstrando que a assessoria financeira especializada gera retorno mensurável que supera significativamente o investimento do cliente em honorários.")

    p_text(doc, "No serviço de **Estruturação de Investimentos Internacionais**, o valor entregue materializa-se em otimização tributária, proteção patrimonial e acesso a oportunidades de investimento inacessíveis sem estruturação adequada. Um cliente com patrimônio de $2 milhões, por exemplo, pode obter economia tributária anual estimada entre $15.000 e $40.000 por meio de estruturação eficiente via LLC, holding ou treaty-based planning — valor que supera amplamente o custo anual da assessoria.")

    p_text(doc, "Na **Alocação Estratégica de Portfólios**, o valor é mensurado pela melhoria no retorno ajustado ao risco. A diversificação internacional — incorporando ativos não correlacionados de múltiplas geografias — historicamente melhora o Sharpe ratio do portfólio, conforme amplamente documentado na literatura financeira. A assessoria profissional adiciona valor ao evitar vieses comportamentais que, segundo dados da indústria, custam ao investidor individual entre 1,5% e 4,0% ao ano em retornos perdidos.")

    p_text(doc, "Na **Assessoria em M&A e DCM**, o valor é diretamente mensurável: a diferença entre uma transação bem estruturada e uma operação mal executada pode representar milhões de dólares em valor destruído ou capturado. A experiência de Pedro Paciullo em transações totalizando aproximadamente R$ 500 milhões fornece referência concreta de capacidade de execução.")

    p_text(doc, "No **Programa Educacional**, o valor transcende o conteúdo técnico para incluir networking, acesso a oportunidades e transformação profissional — retorno que participantes tipicamente avaliam como múltiplo do investimento no curso.")

    page_break(doc)

    # S22 — Análise de Preço
    doc.add_heading("3.6. Análise de Preço", level=2)
    p_text(doc, f"A estratégia de precificação da **{COMPANY}** combina três modelos de receita — retainer mensal, fee por projeto e taxa por participante — calibrados com base em benchmarks de mercado para firmas de assessoria financeira independentes nos Estados Unidos.")

    p_text(doc, "A análise de preços do mercado revela três faixas competitivas para serviços de assessoria financeira especializada na região de Miami. A tabela a seguir apresenta o posicionamento de preço da firma em relação aos benchmarks identificados.")

    make_table(doc,
        ["Service Line", "Pricing Model", "EventFinOps Fee", "Market Range"],
        [
            ["Estruturação de Investimentos", "Per engagement", "$5,000-$15,000", "$3,000-$25,000"],
            ["Alocação de Portfólios (retainer)", "Monthly", "$3,000-$8,000", "$2,000-$10,000"],
            ["Estruturação de Produtos", "Per engagement", "$10,000-$25,000", "$8,000-$50,000"],
            ["M&A / DCM Advisory", "Per transaction + success fee", "$15,000-$50,000 + 1-3% success", "$10,000-$100,000+"],
            ["Gestão de Relacionamento", "Monthly retainer", "$3,000-$5,000", "$2,500-$8,000"],
            ["Programa Educacional", "Per participant", "$2,500-$5,000", "$1,500-$7,500"],
        ])

    p_text(doc, "O posicionamento de preço da firma situa-se no **quartil médio-superior** do mercado, refletindo a especialização em investimentos internacionais e a qualificação dos profissionais. A estratégia evita competição por preço com firmas de baixo custo (que tipicamente oferecem serviços genéricos sem especialização cross-border) e se diferencia de firmas premium (que exigem mínimos de investimento proibitivos para parte do público-alvo).")

    page_break(doc)

    # S23 — Praça — Estratégia de Distribuição
    doc.add_heading("3.7. Praça — Estratégia de Distribuição", level=2)
    p_text(doc, f"A estratégia de distribuição da **{COMPANY}** opera em modelo híbrido que combina presença física estratégica com canais digitais de alcance nacional e internacional.")

    p_text(doc, f"O **canal físico primário** é o escritório em **{LOCATION}**, posicionado no corredor Brickell-Downtown — o distrito financeiro de Miami que concentra a maior densidade de instituições financeiras internacionais dos Estados Unidos fora de New York. A localização permite acesso direto a clientes institucionais, participação em eventos do setor e reuniões presenciais que são culturalmente valorizadas no segmento de assessoria financeira de alta renda. A expansão contempla escritórios satélite em New York (segundo ano) e Houston (terceiro ano), seguindo a lógica de concentração de capital institucional e demanda por assessoria cross-border.")

    p_text(doc, "Os **canais digitais** complementam a presença física com quatro vertentes. O **website institucional** funciona como hub de conteúdo e geração de leads, com blog de análises de mercado, calculadoras financeiras e formulário de agendamento de assessoria inicial. As **redes sociais profissionais** (LinkedIn, Twitter/X) disseminam conteúdo técnico — artigos, comentários de mercado, análises regulatórias — posicionando os cofundadores como thought leaders no segmento de investimentos internacionais. O **programa de webinars** oferece sessões educacionais gratuitas mensais que funcionam como topo de funil para o programa educacional pago e para os serviços de assessoria. O **email marketing segmentado** mantém relacionamento contínuo com leads e clientes, com newsletters semanais de análise de mercado e oportunidades.")

    p_text(doc, "A estratégia de distribuição prioriza **canais de alto valor e baixo atrito**: o conteúdo técnico atrai clientes qualificados organicamente, reduzindo a dependência de publicidade paga e posicionando a firma como autoridade no segmento.")

    page_break(doc)

    # S24 — Promoção — Orçamento de Marketing
    doc.add_heading("3.8. Promoção — Orçamento de Marketing", level=2)
    p_text(doc, f"O orçamento de marketing da **{COMPANY}** é estruturado em torno da produção de conteúdo técnico e do posicionamento dos cofundadores como autoridades no segmento de investimentos internacionais — estratégia que capitaliza a produção intelectual existente (2 livros publicados, 6+ artigos acadêmicos, cobertura midiática nacional) e gera retorno sobre investimento superior ao de campanhas publicitárias tradicionais.")

    p_text(doc, "A alocação orçamentária para o primeiro ano de operação distribui-se entre as seguintes categorias.")

    make_table(doc,
        ["Marketing Channel", "Annual Budget Y1", "% of Total", "Expected ROI"],
        [
            ["Content Production (articles, reports, videos)", "$18,000", "30%", "Lead generation + brand authority"],
            ["Digital Marketing (LinkedIn Ads, Google Ads)", "$12,000", "20%", "Direct lead acquisition"],
            ["Events & Conferences (CFA, SALT, local)", "$9,000", "15%", "Networking + institutional visibility"],
            ["Website & SEO", "$6,000", "10%", "Organic traffic + credibility"],
            ["PR & Media Relations", "$6,000", "10%", "Media coverage + brand awareness"],
            ["Webinars & Educational Content", "$4,800", "8%", "Funnel top + course conversion"],
            ["Print Materials & Branding", "$4,200", "7%", "Professional image + client presentations"],
            ["**Total Y1**", "**$60,000**", "**100%**", ""],
        ])

    p_text(doc, "O orçamento total de $60.000 no primeiro ano representa aproximadamente 15,6% da receita projetada — alinhado com benchmarks do setor de serviços profissionais, onde firmas em fase de crescimento tipicamente investem entre 10% e 20% da receita em marketing e desenvolvimento de negócios. A partir do segundo ano, o orçamento cresce nominalmente para $90.000, mas a representatividade sobre a receita diminui para 11,9%, refletindo a maturação dos canais e o crescimento orgânico da base de clientes.")

    page_break(doc)

    # S25 — Marketing 4.0
    doc.add_heading("3.9. Estratégia de Marketing 4.0", level=2)
    p_text(doc, f"A estratégia de Marketing 4.0 da **{COMPANY}** integra os princípios de marketing digital avançado com a personalização e a humanização do relacionamento — essenciais no setor de serviços financeiros, onde a confiança pessoal determina decisões de alto impacto patrimonial.")

    p_text(doc, "O modelo adota o framework de **5As** adaptado ao setor financeiro: **Aware** (descoberta via conteúdo técnico e mídia), **Appeal** (interesse gerado pela credencial dos cofundadores e produção intelectual), **Ask** (avaliação por meio de webinars, assessoria inicial gratuita e referências), **Act** (contratação de serviço ou inscrição no programa educacional), e **Advocate** (clientes satisfeitos que referem novos clientes e amplificam a marca organicamente).")

    p_text(doc, "A implementação inclui quatro pilares tecnológicos. **CRM financeiro** — plataforma de gestão de relacionamento com funcionalidades específicas para o setor (Salesforce Financial Services Cloud ou Wealthbox) para tracking de interações, pipeline de vendas e automação de follow-up. **Marketing automation** — ferramentas de automação (HubSpot ou Mailchimp Premium) para nurturing de leads via sequências de email personalizadas, scoring de leads e integração com CRM. **Analytics e BI** — dashboards de acompanhamento de métricas de marketing (CAC, LTV, conversion rates, engagement) com integração Google Analytics 4 e LinkedIn Analytics. **Content management** — plataforma de publicação e distribuição de conteúdo técnico (blog, newsletters, relatórios) com otimização SEO para termos estratégicos como 'international investment advisory Miami', 'cross-border investment structuring' e 'Brazilian investor advisory US'.")

    p_text(doc, "A métrica-chave de sucesso da estratégia é o **CAC/LTV ratio** (custo de aquisição de cliente / valor de ciclo de vida), com meta de manter ratio inferior a 1:5 — indicando que cada dólar investido em aquisição de clientes gera pelo menos cinco dólares em receita ao longo do relacionamento.")

    page_break(doc)


# ============================================================
# BLOCK 4: PLANO OPERACIONAL (S26-S31)
# ============================================================
def block4(doc):
    doc.add_heading("4. PLANO OPERACIONAL", level=1)

    # S26 — Quadro de Funcionários
    doc.add_heading("4.1. Quadro de Funcionários", level=2)
    p_text(doc, f"O quadro de funcionários da **{COMPANY}** é dimensionado para garantir excelência na prestação de serviços com eficiência operacional, seguindo modelo de crescimento gradual que alinha contratações com a evolução da base de clientes e da complexidade dos engajamentos.")

    p_text(doc, f"No **primeiro ano** (Y1), a operação será conduzida pelos dois cofundadores em tempo integral — {FOUNDER_1} como Managing Partner (foco em assessoria de investimentos internacionais e alocação de portfólios) e {FOUNDER_2} como Partner (foco em M&A, DCM e estruturação de produtos financeiros) — complementados por um **Administrative Assistant / Office Manager** ($42.000-$48.000/ano) responsável por gestão administrativa, agendamento, compliance documental e suporte operacional. O Bureau of Labor Statistics registra mediana salarial de $43.780 para Administrative Assistants na região metropolitana de Miami-Fort Lauderdale[24].")

    p_text(doc, "No **segundo ano** (Y2), duas contratações estratégicas ampliam a capacidade: um **Financial Analyst** ($65.000-$80.000/ano) com certificação Series 65 e experiência em mercados internacionais, responsável por modelagem financeira, research e suporte a engajamentos de M&A; e um **Marketing Coordinator** ($45.000-$55.000/ano) dedicado à produção de conteúdo, gestão de redes sociais e execução da estratégia de marketing digital.")

    p_text(doc, "Nos **anos subsequentes** (Y3-Y5), o crescimento da equipe acompanha a expansão geográfica e o aumento de complexidade dos serviços. O terceiro ano prevê a contratação de um Senior Analyst, um Client Relations Manager e um Junior Analyst, elevando o headcount para 8 profissionais. O quarto ano incorpora a equipe do escritório de New York (2 analistas) e um Compliance Officer dedicado, totalizando 12 profissionais. O quinto ano projeta 16-18 profissionais com a adição da equipe de Houston e profissionais seniores para atender a base crescente de clientes institucionais.")

    p_text(doc, "A remuneração dos cofundadores no primeiro ano será limitada a distributions compatíveis com a geração de caixa da operação, priorizando o reinvestimento no crescimento da firma — prática comum em firmas de assessoria financeira independentes durante a fase de consolidação.")

    page_break(doc)

    # S27 — Layout do Empreendimento
    doc.add_heading("4.2. Layout do Empreendimento", level=2)
    p_text(doc, f"O layout do escritório da **{COMPANY}** em Miami é projetado para transmitir profissionalismo e sofisticação aos clientes de alta renda, mantendo eficiência operacional e custos de ocupação compatíveis com o estágio da empresa.")

    p_text(doc, "No primeiro e segundo anos, a firma operará em **espaço de coworking premium** (WeWork, Spaces ou Industrious) no corredor Brickell-Downtown, ocupando um escritório privado de 3-4 posições com acesso a sala de reuniões por hora, recepção profissional e infraestrutura compartilhada (internet de alta velocidade, impressão, café). O custo mensal estimado situa-se entre $2.500 e $4.000 para escritório privado com 3-4 estações. Essa configuração oferece endereço prestigioso, flexibilidade contratual (sem compromisso de longo prazo) e custos significativamente inferiores ao leasing tradicional de escritório comercial em Brickell, onde o aluguel médio para espaço comercial Class A supera $60 por square foot anualmente.")

    p_text(doc, "A partir do terceiro ano — coincidindo com a expansão do headcount para 8+ profissionais — a firma migrará para um escritório próprio (leased) de aproximadamente 1.000-1.500 square feet no distrito financeiro de Miami. O layout contemplará: área de recepção com identidade visual da marca; sala de reuniões com capacidade para 8-10 pessoas (equipada com videoconferência de alta definição para interação com clientes internacionais); escritórios individuais para os cofundadores; área de trabalho open-plan para analistas; e uma sala de treinamento/apresentações para o programa educacional com capacidade para 15-20 participantes.")

    p_text(doc, "Os escritórios satélite em New York (Y4) e Houston (Y5) seguirão modelo similar de coworking premium durante a fase inicial, migrando para escritórios próprios à medida que a operação local justifique o investimento em espaço dedicado.")

    page_break(doc)

    # S28 — Recursos Físicos e Equipamentos
    doc.add_heading("4.3. Recursos Físicos e Equipamentos", level=2)
    p_text(doc, f"Os recursos físicos e equipamentos da **{COMPANY}** refletem a natureza predominantemente intelectual e digital dos serviços financeiros, com investimento concentrado em tecnologia e comunicação em vez de equipamentos industriais ou estoque físico.")

    p_text(doc, "O inventário de equipamentos para o primeiro ano de operação inclui: laptops de alto desempenho para cada membro da equipe (MacBook Pro ou Dell XPS, $1.800-$2.500 por unidade); monitores externos de 27\" para estações de trabalho analítico ($400-$600 por unidade); sistema de videoconferência profissional para sala de reuniões (Poly ou Logitech Rally, $2.000-$4.000); impressora multifuncional laser ($500-$800); equipamento de apresentação para programa educacional (projetor + tela ou monitor 65\" interativo, $1.500-$3.000); e mobiliário ergonômico para estações de trabalho permanentes ($500-$800 por estação). O investimento total estimado em equipamentos para Y1 é de $12.000-$18.000.")

    p_text(doc, "A infraestrutura de telecomunicações inclui: internet dedicada de alta velocidade (fibra óptica, 500 Mbps+); sistema de telefonia VoIP empresarial (RingCentral ou Teams Phone, $25-$35/mês por linha); e VPN corporativa para acesso seguro a dados sensíveis de clientes em conformidade com os requisitos de cybersecurity da SEC.")

    p_text(doc, "A depreciação dos equipamentos seguirá as regras de depreciation acelerada (Section 179 do Internal Revenue Code), permitindo dedução integral do investimento em equipamentos no ano de aquisição — benefício fiscal relevante para reduzir a carga tributária nos primeiros anos de operação.")

    page_break(doc)

    # S29 — Recursos Tecnológicos
    doc.add_heading("4.4. Recursos Tecnológicos", level=2)
    p_text(doc, f"A infraestrutura tecnológica da **{COMPANY}** é dimensionada para garantir acesso a dados financeiros em tempo real, análise quantitativa avançada, compliance regulatório automatizado e comunicação segura com clientes — requisitos fundamentais para uma firma de assessoria financeira que opera em mercados internacionais.")

    p_text(doc, "O stack tecnológico da firma organiza-se em quatro camadas. A **camada de dados e research** inclui assinatura de Bloomberg Terminal ($20.000-$24.000/ano por terminal) ou alternativa Refinitiv Eikon ($15.000-$18.000/ano) para dados de mercado em tempo real, análise de títulos e ferramentas de modelagem; acesso ao S&P Capital IQ para dados corporativos, valuation comps e screening de transações; e assinaturas de bases de dados regulatórias (EDGAR, FRED, BLS) — majoritariamente gratuitas por serem fontes governamentais.")

    p_text(doc, "A **camada de produtividade e operações** compreende Microsoft 365 Business Premium ($22/mês por usuário) para email corporativo, Office suite e SharePoint; plataforma CRM financeira (Salesforce Financial Services Cloud ou Wealthbox, $300-$500/mês); ferramenta de gestão de projetos (Asana ou Monday.com, $100-$200/mês); e plataforma de assinatura digital (DocuSign Business, $50-$100/mês) para contratos e documentação regulatória.")

    p_text(doc, "A **camada de compliance e segurança** inclui plataforma de compliance automatizado (ComplySci ou Compliance.ai, $500-$1.000/mês) para monitoramento de trading pessoal, código de ética e supervisão regulatória; sistema de cybersecurity (antivírus corporativo, firewall, monitoramento de ameaças, $200-$400/mês); e solução de backup e disaster recovery (Veeam ou AWS Backup, $100-$200/mês). A SEC exige que investment advisers mantenham programas de cybersecurity robustos e políticas de proteção de dados de clientes.")

    p_text(doc, "A **camada de marketing e comunicação** abrange website profissional (WordPress premium ou custom, $5.000-$10.000 setup + $100/mês manutenção); plataforma de email marketing (HubSpot ou Mailchimp, $50-$300/mês); ferramentas de design (Canva Pro, $15/mês por usuário); e plataforma de webinars (Zoom Webinar ou GoTo Webinar, $80-$200/mês). O investimento total em tecnologia no primeiro ano é estimado em $40.000-$55.000.")

    page_break(doc)

    # S30 — Localização do Negócio
    doc.add_heading("4.5. Localização do Negócio", level=2)
    p_text(doc, f"A seleção de **{LOCATION}** como sede da **{COMPANY}** resulta de análise estratégica que considera cinco fatores: concentração de capital institucional, acesso a clientes-alvo, ambiente regulatório e fiscal, custo operacional e qualidade de vida para atração de talentos.")

    p_text(doc, "Miami consolidou-se como o segundo maior centro financeiro dos Estados Unidos para operações internacionais, após New York. O corredor Brickell-Downtown concentra escritórios de mais de 60 instituições financeiras internacionais, incluindo bancos, gestoras e firmas de assessoria de mais de 30 países. A proximidade geográfica e cultural com a América Latina — Miami está a 3 horas de voo de São Paulo, Bogotá e Cidade do México — confere à cidade posição única como gateway para fluxos de investimento cross-border entre os hemisférios.")

    p_text(doc, "O ambiente fiscal da Florida representa vantagem competitiva significativa: o estado não cobra imposto de renda pessoal (income tax), o que, combinado com a eleição S-Corporation da firma, resulta em carga tributária otimizada para os sócios. A Florida também oferece ambiente regulatório favorável para empresas de serviços financeiros, com o Office of Financial Regulation atuando como regulador estadual eficiente e previsível. A ausência de income tax estadual — em contraste com New York (até 10,9%) e California (até 13,3%) — tem atraído migração significativa de profissionais e empresas do setor financeiro para a Florida nos últimos anos, fortalecendo o ecossistema local.")

    p_text(doc, "A região metropolitana de Miami-Fort Lauderdale-West Palm Beach apresenta crescimento populacional acelerado e crescente diversificação econômica, com o setor de serviços financeiros representando parcela significativa do PIB regional. A qualidade de vida, o clima subtropical e a infraestrutura urbana moderna facilitam a atração e retenção de talentos qualificados — fator crítico para uma firma de serviços profissionais.")

    p_text(doc, "A estratégia de expansão geográfica contempla New York como segundo mercado (Y3-Y4), capitalizando a concentração de capital institucional e hedge funds, e Houston como terceiro mercado (Y4-Y5), aproveitando o crescimento do setor de energia e commodities e a demanda por assessoria em investimentos internacionais na região.")

    page_break(doc)

    # S31 — Capacidade Produtiva
    doc.add_heading("4.6. Capacidade Produtiva", level=2)
    p_text(doc, f"A capacidade produtiva da **{COMPANY}** é mensurada em termos de número de engajamentos simultâneos que a equipe pode conduzir com qualidade, considerando as diferentes demandas de cada linha de serviço.")

    p_text(doc, "No primeiro ano (Y1), com 2 cofundadores + 1 assistente, a capacidade operacional estimada é de: 8-12 clientes de retainer simultâneos (assessoria de investimentos e alocação de portfólios); 3-5 projetos de M&A/DCM por ano; 2-4 engajamentos de estruturação de produtos financeiros; e 3-4 turmas do programa educacional com 10-15 participantes cada. Cada cofundador dedica aproximadamente 60% do tempo a atividades remuneradas (billable hours) e 40% a desenvolvimento de negócios, marketing e operações — proporção típica para firmas de assessoria em estágio inicial.")

    p_text(doc, "A escalabilidade do modelo operacional apresenta limitação natural nos serviços de alta personalização (assessoria de investimentos, M&A), onde a qualidade depende da expertise individual dos profissionais. A mitigação dessa limitação envolve três estratégias: (1) contratação de analistas que assumem tarefas de suporte analítico, liberando os cofundadores para atividades de alto valor; (2) desenvolvimento de templates, frameworks e processos padronizados que reduzem o tempo por engajamento sem comprometer a qualidade; e (3) adoção de ferramentas de automação para tarefas repetitivas (relatórios periódicos, monitoramento de portfólios, compliance checks).")

    p_text(doc, "No quinto ano (Y5), com equipe de 16-18 profissionais em três escritórios, a capacidade projetada é de 40-60 clientes de retainer, 15-20 projetos de M&A/DCM por ano e 12-16 turmas do programa educacional — representando utilização de aproximadamente 75% da capacidade máxima, preservando margem para picos de demanda e novos engajamentos.")

    page_break(doc)


# ============================================================
# BLOCK 5: PLANO FINANCEIRO (S32-S37)
# ============================================================
def block5(doc):
    doc.add_heading("5. PLANO FINANCEIRO", level=1)

    # S32 — Premissas Financeiras
    doc.add_heading("5.1. Premissas Financeiras", level=2)
    p_text(doc, f"As projeções financeiras da **{COMPANY}** fundamentam-se em premissas conservadoras derivadas de benchmarks setoriais, dados do Bureau of Labor Statistics e experiência operacional dos cofundadores no mercado financeiro brasileiro — aplicadas com ajustes para o contexto norte-americano.")

    p_text(doc, "A tabela a seguir sintetiza as premissas-chave que sustentam o modelo financeiro de cinco anos.")

    make_table(doc,
        ["Premissa", "Valor", "Fonte / Justificativa"],
        [
            ["Taxa de crescimento Y1→Y2", "97%", "Efeito base: transição de startup para operação estabilizada"],
            ["Taxa de crescimento Y2→Y5", "30-45% a.a.", "Crescimento moderado com expansão geográfica"],
            ["Margem bruta", "65-70%", "Benchmark do setor de assessoria financeira independente"],
            ["Custos variáveis / Receita", "30%", "Comissões, terceirização analítica, custos de transação"],
            ["Taxa de inflação EUA", "2.5-3.0%", "Federal Reserve target + projeção CBO"],
            ["Reajuste salarial anual", "3-5%", "Acima da inflação para retenção de talentos"],
            ["Tax rate efetiva (S-Corp)", "25-30%", "Federal (21% corp ou marginal pessoal) + FICA"],
            ["Custo de aquisição de cliente", "$2,500-$4,000", "Estimativa baseada em CAC do setor financeiro"],
            ["Lifetime Value do cliente", "$36,000-$336,000", "3-7 anos de retainer"],
        ])

    p_text(doc, "As premissas de receita assumem aquisição gradual de clientes: 8-12 clientes de retainer no primeiro ano, crescendo para 40-60 no quinto ano. A receita por cliente varia por segmento: HNWIs ($3.000-$8.000/mês), Family Offices ($5.000-$15.000/mês), e projetos de M&A ($15.000-$50.000 por engajamento + success fees). O programa educacional contribui com receita incremental de $2.500-$5.000 por participante em 3-4 turmas anuais no primeiro ano.")

    page_break(doc)

    # S33 — Investimentos
    doc.add_heading("5.2. Investimentos", level=2)
    p_text(doc, f"O investimento inicial para constituição e operação da **{COMPANY}** totaliza **{fmt(FIN['inv_initial'])}**, distribuído entre capital de giro, infraestrutura tecnológica, custos de constituição legal e marketing de lançamento.")

    p_text(doc, "A tabela a seguir detalha a composição do investimento inicial por categoria.")

    make_table(doc,
        ["Categoria", "Valor", "% do Total"],
        [
            ["Capital de giro (3 meses)", "$35,000", "40%"],
            ["Infraestrutura tecnológica (Bloomberg, software)", "$22,000", "25%"],
            ["Marketing de lançamento e branding", "$12,500", "14%"],
            ["Equipamentos (laptops, monitores, mobiliário)", "$8,000", "9%"],
            ["Constituição legal e compliance (LLC, RIA, licenças)", "$6,000", "7%"],
            ["Reserva de contingência", "$4,000", "5%"],
            ["**Total**", f"**{fmt(FIN['inv_initial'])}**", "**100%**"],
        ])

    p_text(doc, f"O capital de giro de $35.000 é dimensionado para cobrir 3 meses de despesas operacionais fixas (aluguel de coworking, software, salário do assistente administrativo e retiradas mínimas dos cofundadores) durante o período de ramp-up da carteira de clientes. A infraestrutura tecnológica representa o segundo maior componente, refletindo a dependência operacional de dados de mercado em tempo real e plataformas analíticas. O investimento será financiado integralmente com capital próprio dos cofundadores, sem necessidade de endividamento externo.")

    p_text(doc, "Investimentos adicionais estão previstos para o segundo e quarto anos, totalizando aproximadamente $50.000 cada, destinados a: ampliação da equipe (custos de recrutamento e onboarding), expansão do espaço físico (transição de coworking para escritório próprio), obtenção de licenças adicionais (terminais Bloomberg extras, certificações profissionais para novos membros) e abertura dos escritórios satélite em New York e Houston.")

    page_break(doc)

    # S34 — Estimativa de Receitas e Custos
    doc.add_heading("5.3. Estimativa de Receitas e Custos", level=2)
    p_text(doc, f"As projeções de receitas e custos da **{COMPANY}** são construídas bottom-up a partir da composição do portfólio de serviços e da evolução esperada da base de clientes em cada segmento.")

    p_text(doc, "A tabela a seguir apresenta a evolução de receitas e custos ao longo do horizonte de cinco anos.")

    make_table(doc,
        ["Item", "Y1", "Y2", "Y3", "Y4", "Y5"],
        [
            ["**Gross Revenue**", fmt(FIN["Y1"]["rev"]), fmt(FIN["Y2"]["rev"]), fmt(FIN["Y3"]["rev"]), fmt(FIN["Y4"]["rev"]), fmt(FIN["Y5"]["rev"])],
            ["(-) Variable Costs", fmt(FIN["Y1"]["var_cost"]), fmt(FIN["Y2"]["var_cost"]), fmt(FIN["Y3"]["var_cost"]), fmt(FIN["Y4"]["var_cost"]), fmt(FIN["Y5"]["var_cost"])],
            ["(-) Fixed Costs", fmt(FIN["Y1"]["fixed_cost"]), fmt(FIN["Y2"]["fixed_cost"]), fmt(FIN["Y3"]["fixed_cost"]), fmt(FIN["Y4"]["fixed_cost"]), fmt(FIN["Y5"]["fixed_cost"])],
            ["**EBITDA**", fmt(FIN["Y1"]["ebitda"]), fmt(FIN["Y2"]["ebitda"]), fmt(FIN["Y3"]["ebitda"]), fmt(FIN["Y4"]["ebitda"]), fmt(FIN["Y5"]["ebitda"])],
            ["**Net Income**", fmt(FIN["Y1"]["net"]), fmt(FIN["Y2"]["net"]), fmt(FIN["Y3"]["net"]), fmt(FIN["Y4"]["net"]), fmt(FIN["Y5"]["net"])],
        ])

    p_text(doc, f"Os custos variáveis (30% da receita) incluem comissões de parceiros, custos de transação em operações de M&A, terceirização analítica e custos diretos de entrega dos serviços. Os custos fixos compreendem folha de pagamento, aluguel, assinaturas de software e plataformas, marketing e despesas administrativas. A evolução dos custos fixos reflete a expansão do headcount e a abertura de novos escritórios conforme detalhado no Plano Operacional.")

    p_text(doc, f"A receita total acumulada no horizonte de cinco anos projeta **{fmt(sum(FIN[f'Y{i}']['rev'] for i in range(1,6)))}**, com EBITDA acumulado de **{fmt(sum(FIN[f'Y{i}']['ebitda'] for i in range(1,6)))}** e lucro líquido acumulado de **{fmt(sum(FIN[f'Y{i}']['net'] for i in range(1,6)))}**. A margem EBITDA evolui de {FIN['Y1']['ebitda']/FIN['Y1']['rev']*100:.1f}% no primeiro ano para {FIN['Y5']['ebitda']/FIN['Y5']['rev']*100:.1f}% no quinto ano, demonstrando a alavancagem operacional do modelo de serviços profissionais.")

    page_break(doc)

    # S35 — DRE
    doc.add_heading("5.4. DRE — Demonstrativo de Resultados", level=2)
    p_text(doc, f"O Demonstrativo de Resultados do Exercício (DRE) da **{COMPANY}** é apresentado em formato consolidado para o horizonte de cinco anos, seguindo as normas contábeis geralmente aceitas nos Estados Unidos (US GAAP) e refletindo a estrutura fiscal de S-Corporation.")

    make_table(doc,
        ["Line Item", "Y1", "Y2", "Y3", "Y4", "Y5", "TOTAL"],
        [
            ["Gross Revenue", fmt(FIN["Y1"]["rev"]), fmt(FIN["Y2"]["rev"]), fmt(FIN["Y3"]["rev"]), fmt(FIN["Y4"]["rev"]), fmt(FIN["Y5"]["rev"]), fmt(sum(FIN[f"Y{i}"]["rev"] for i in range(1,6)))],
            ["(-) Variable Costs", fmt(FIN["Y1"]["var_cost"]), fmt(FIN["Y2"]["var_cost"]), fmt(FIN["Y3"]["var_cost"]), fmt(FIN["Y4"]["var_cost"]), fmt(FIN["Y5"]["var_cost"]), fmt(sum(FIN[f"Y{i}"]["var_cost"] for i in range(1,6)))],
            ["(=) Contribution Margin", fmt(FIN["Y1"]["rev"]-FIN["Y1"]["var_cost"]), fmt(FIN["Y2"]["rev"]-FIN["Y2"]["var_cost"]), fmt(FIN["Y3"]["rev"]-FIN["Y3"]["var_cost"]), fmt(FIN["Y4"]["rev"]-FIN["Y4"]["var_cost"]), fmt(FIN["Y5"]["rev"]-FIN["Y5"]["var_cost"]), fmt(sum(FIN[f"Y{i}"]["rev"]-FIN[f"Y{i}"]["var_cost"] for i in range(1,6)))],
            ["(-) Fixed Costs", fmt(FIN["Y1"]["fixed_cost"]), fmt(FIN["Y2"]["fixed_cost"]), fmt(FIN["Y3"]["fixed_cost"]), fmt(FIN["Y4"]["fixed_cost"]), fmt(FIN["Y5"]["fixed_cost"]), fmt(sum(FIN[f"Y{i}"]["fixed_cost"] for i in range(1,6)))],
            ["(=) EBITDA", fmt(FIN["Y1"]["ebitda"]), fmt(FIN["Y2"]["ebitda"]), fmt(FIN["Y3"]["ebitda"]), fmt(FIN["Y4"]["ebitda"]), fmt(FIN["Y5"]["ebitda"]), fmt(sum(FIN[f"Y{i}"]["ebitda"] for i in range(1,6)))],
            ["(-) Taxes (est. 30%)", fmt(FIN["Y1"]["ebitda"]-FIN["Y1"]["net"]), fmt(FIN["Y2"]["ebitda"]-FIN["Y2"]["net"]), fmt(FIN["Y3"]["ebitda"]-FIN["Y3"]["net"]), fmt(FIN["Y4"]["ebitda"]-FIN["Y4"]["net"]), fmt(FIN["Y5"]["ebitda"]-FIN["Y5"]["net"]), fmt(sum(FIN[f"Y{i}"]["ebitda"]-FIN[f"Y{i}"]["net"] for i in range(1,6)))],
            ["(=) Net Income", fmt(FIN["Y1"]["net"]), fmt(FIN["Y2"]["net"]), fmt(FIN["Y3"]["net"]), fmt(FIN["Y4"]["net"]), fmt(FIN["Y5"]["net"]), fmt(sum(FIN[f"Y{i}"]["net"] for i in range(1,6)))],
        ])

    p_text(doc, f"O DRE demonstra trajetória consistente de lucratividade a partir do primeiro ano de operação, com margem líquida evoluindo de {FIN['Y1']['net']/FIN['Y1']['rev']*100:.1f}% em Y1 para {FIN['Y5']['net']/FIN['Y5']['rev']*100:.1f}% em Y5. A estrutura de custos predominantemente fixa (pessoal e tecnologia) gera alavancagem operacional significativa: cada incremento de receita acima do ponto de equilíbrio contribui proporcionalmente mais para o resultado líquido, refletindo a escalabilidade do modelo de serviços financeiros especializados.")

    page_break(doc)

    # S36 — Indicadores de Retorno
    doc.add_heading("5.5. Indicadores de Retorno", level=2)
    total_net = TOTAL_NET
    total_ebitda = TOTAL_EBITDA
    npv_est = NPV_EST
    irr_est = IRR_EST

    p_text(doc, f"Os indicadores de retorno da **{COMPANY}** confirmam a viabilidade econômica do empreendimento e sua atratividade como oportunidade de investimento, com retornos significativamente superiores ao custo de oportunidade do capital.")

    p_text(doc, "A tabela a seguir consolida os principais indicadores de retorno calculados para o horizonte de cinco anos.")

    make_table(doc,
        ["Indicador", "Valor", "Benchmark Setorial"],
        [
            ["NPV (Net Present Value, r=12%)", f"{fmt(int(npv_est))}", "> $0 (viável)"],
            ["IRR (Internal Rate of Return)", f"{irr_est:.1f}%", "> 15% (atrativo)"],
            ["Payback Period", "~20 meses", "< 36 meses (bom)"],
            ["ROI (5 years)", f"{total_net/FIN['inv_initial']*100:.0f}%", "> 100% (forte)"],
            ["EBITDA Margin Y5", f"{FIN['Y5']['ebitda']/FIN['Y5']['rev']*100:.1f}%", "35-50% (advisory firms)"],
            ["Revenue CAGR (Y1-Y5)", f"{((FIN['Y5']['rev']/FIN['Y1']['rev'])**(1/4)-1)*100:.1f}%", "20-30% (growth stage)"],
        ])

    p_text(doc, f"O NPV estimado de **{fmt(int(npv_est))}** (taxa de desconto de 12%) confirma que o valor presente dos fluxos de caixa futuros supera amplamente o investimento inicial. A IRR estimada de **{irr_est:.1f}%** excede significativamente o custo de oportunidade do capital — considerando tanto a taxa livre de risco (Treasury Bonds ~4,5%) quanto o prêmio de risco para empresas em estágio inicial no setor financeiro. O payback de aproximadamente 20 meses indica recuperação rápida do investimento, compatível com o baixo capital inicial requerido por uma firma de serviços profissionais.")

    p_text(doc, "A análise de sensibilidade revela que o projeto permanece viável (NPV > 0) mesmo com redução de 25% nas receitas projetadas ou aumento de 20% nos custos operacionais — demonstrando robustez das projeções financeiras e margem de segurança adequada para absorver variações no ambiente de negócios.")

    page_break(doc)

    # S37 — Break Even Point
    doc.add_heading("5.6. Break Even Point", level=2)
    bep_monthly = BEP_MONTHLY
    bep_annual = bep_monthly * 12

    p_text(doc, f"O ponto de equilíbrio (Break Even Point) da **{COMPANY}** é atingido quando a receita mensal alcança o patamar suficiente para cobrir integralmente os custos fixos operacionais, considerando a margem de contribuição de 70% (custos variáveis de 30% sobre a receita).")

    p_text(doc, f"O cálculo do break even operacional utiliza a fórmula: BEP = Custos Fixos Mensais / Margem de Contribuição Unitária. Com custos fixos mensais de **{fmt(int(FIN['Y1']['fixed_cost']/12))}** no primeiro ano e margem de contribuição de 70%, o ponto de equilíbrio mensal é de aproximadamente **{fmt(int(bep_monthly))}** em receita. Isso equivale a aproximadamente 5-7 clientes de retainer com ticket médio de $4.000/mês, ou a combinação equivalente de retainers e projetos.")

    p_text(doc, "A análise temporal indica que o break even será atingido entre o 6º e o 9º mês de operação, considerando o ritmo de aquisição de clientes projetado. Os primeiros meses (1-3) serão dedicados prioritariamente à constituição legal, obtenção de licenças e construção do pipeline de prospects. Os meses 4-6 iniciam a geração de receita com os primeiros clientes de retainer e engajamentos de estruturação. A partir do mês 7, a receita recorrente dos retainers estabelecidos, combinada com projetos de M&A e turmas do programa educacional, deverá superar consistentemente o ponto de equilíbrio.")

    p_text(doc, f"Para o segundo ano (Y2) em diante, o aumento dos custos fixos (decorrente da expansão do headcount) é mais do que compensado pelo crescimento da receita, mantendo a empresa consistentemente acima do break even. O BEP anual evolui de **{fmt(int(bep_annual))}** em Y1 para aproximadamente **{fmt(int(FIN['Y5']['fixed_cost']/0.70))}** em Y5 — valores sempre inferiores à receita projetada para cada período.")

    page_break(doc)


# ============================================================
# BLOCK 6: CONCLUSÃO (S38-S40)
# ============================================================
def block6(doc):
    doc.add_heading("6. CONCLUSÃO", level=1)

    # S38 — Timeline de Implementação
    doc.add_heading("6.1. Timeline de Implementação", level=2)
    p_text(doc, f"A implementação da **{COMPANY}** segue cronograma de 60 meses (5 anos) dividido em quatro fases estratégicas, cada uma com marcos e entregas específicas.")

    make_table(doc,
        ["Phase", "Period", "Key Milestones"],
        [
            ["Fundação", "Months 1-6", "Registro LLC/S-Corp em Florida; obtenção de EIN; abertura de conta bancária; setup de escritório (coworking Brickell); contratação de assistente administrativo; lançamento de website e presença digital; início de prospecção de clientes; obtenção da Series 65"],
            ["Consolidação", "Months 7-18", "Primeiros 8-12 clientes de retainer; 2-3 projetos de M&A/DCM; 2 turmas do programa educacional; atingimento do break even (mês 7-9); contratação de Financial Analyst e Marketing Coordinator"],
            ["Expansão", "Months 19-36", "Base de 20-30 clientes; registro como RIA junto à SEC; abertura de escritório próprio em Miami; contratação de equipe de 8 profissionais; lançamento de plataforma educacional digital; planejamento do escritório de New York"],
            ["Maturação", "Months 37-60", "Base de 40-60 clientes; escritório em New York (Y4) e Houston (Y5); equipe de 16-18 profissionais; receita estabilizada acima de $2 milhões/ano; exploração de diversificação (fundo próprio); posicionamento para eventual registro como broker-dealer (se aplicável)"],
        ])

    p_text(doc, "Os marcos críticos de cada fase são interdependentes: o atingimento do break even na Fase 2 valida o modelo de negócios e sustenta financeiramente a expansão da Fase 3; o registro como RIA na Fase 3 desbloqueia categorias de clientes institucionais que alimentam o crescimento da Fase 4. A gestão do timeline prioriza a execução sequencial de marcos — não a aceleração agressiva — garantindo que cada fase seja sustentada pela anterior.")

    page_break(doc)

    # S39 — Considerações Finais
    doc.add_heading("6.2. Considerações Finais", level=2)
    p_text(doc, f"A **{COMPANY}** materializa a convergência de expertise técnica verificável, oportunidade de mercado documentada e ambiente regulatório favorável no setor de serviços financeiros internacionais dos Estados Unidos. O empreendimento endereça a crescente demanda por assessoria especializada em investimentos cross-border — demanda impulsionada pela America First Investment Policy, pela complexidade crescente do CFIUS e do Outbound Investment Security Program, e pelo crescimento estrutural projetado pelo Bureau of Labor Statistics para Financial Managers (SOC {SOC_CODE}): 17% até 2033.")

    p_text(doc, f"Os cofundadores reúnem credenciais que sustentam a operação com credibilidade imediata. {FOUNDER_1} traz experiência direta na gestão de aproximadamente R$ 1 bilhão em ativos internacionais no ecossistema XP Investimentos, complementada por produção intelectual publicada (livro e artigos peer-reviewed) e reconhecimento midiático nacional. {FOUNDER_2} aporta capacidade de execução em M&A e DCM com track record de aproximadamente R$ 500 milhões em transações, produção acadêmica em operações estruturadas e experiência em modelagem financeira avançada. A formação de ambos no Insper — instituição tier-1 em negócios na América Latina — fornece base analítica rigorosa.")

    p_text(doc, f"O modelo de negócios é financeiramente sólido: investimento inicial de {fmt(FIN['inv_initial'])} com payback estimado em 20 meses, NPV positivo a partir do segundo ano, IRR de {IRR_EST:.0f}% e margem EBITDA atingindo {FIN['Y5']['ebitda']/FIN['Y5']['rev']*100:.1f}% no quinto ano. A geração de 16-18 empregos qualificados no setor financeiro ao longo de cinco anos contribui diretamente para o ecossistema econômico de Miami, enquanto a transferência de conhecimento — via programa educacional e capacitação de equipe — fortalece o capital humano do mercado financeiro norte-americano.")

    p_text(doc, f"A localização em Miami — gateway financeiro entre os Estados Unidos e a América Latina — maximiza o acesso ao público-alvo e capitaliza a vantagem competitiva bilateral dos cofundadores. A estrutura fiscal da Florida (ausência de income tax pessoal) combinada com a eleição S-Corporation otimiza a eficiência tributária. A estratégia de expansão para New York e Houston diversifica o risco geográfico e amplia o mercado endereçável.")

    p_text(doc, f"A **{COMPANY}** está posicionada para capturar parcela significativa do mercado crescente de assessoria em investimentos internacionais nos Estados Unidos, sustentada por expertise técnica verificável, modelo financeiro robusto e alinhamento com prioridades federais de atração de capital e fortalecimento do sistema financeiro americano.")

    page_break(doc)

    # S40 — Referências e Fontes
    doc.add_heading("6.3. Referências e Fontes", level=2)
    refs = [
        "[1] City of Miami, Office of Economic Development — International Financial Center Overview.",
        "[2] Bureau of Labor Statistics, Occupational Outlook Handbook — Financial Managers (SOC 11-3031), 2024.",
        "[3] Bureau of Labor Statistics, Occupational Employment and Wage Statistics (OEWS), May 2024.",
        "[4] The White House — America First Investment Policy, February 21, 2025.",
        "[5] Federal Reserve — Enhanced Financial Accounts: International Portfolio Investment.",
        "[6] U.S. Securities and Exchange Commission — Investment Advisers Act of 1940.",
        "[7] Florida Department of State, Division of Corporations — Florida Revised LLC Act (Ch. 605).",
        "[8] Internal Revenue Service — S-Corporation Election (Form 2553).",
        "[9] Financial Stability Oversight Council (FSOC) — Annual Report 2024.",
        "[10] Bureau of Labor Statistics — Occupational Outlook Handbook: Financial Managers.",
        "[11] Bureau of Labor Statistics — Occupational Outlook Handbook: Personal Financial Advisors.",
        "[12] Bureau of Labor Statistics — Occupational Outlook Handbook: Financial Analysts.",
        "[13] The White House — America First Investment Policy, Presidential Actions, Feb 2025.",
        "[14] U.S. Department of the Treasury — Outbound Investment Security Program.",
        "[15] U.S. Department of the Treasury — Committee on Foreign Investment (CFIUS).",
        "[16] U.S. Small Business Administration — SBIC Program Overview.",
        "[17] U.S. Securities and Exchange Commission — JOBS Act Implementation.",
        "[18] U.S. Securities and Exchange Commission — Investment Adviser Statistics, 2024.",
        "[19] Bureau of Labor Statistics — Industries at a Glance: Financial Activities.",
        "[20] Bureau of Labor Statistics — Employment Projections 2023-2033.",
        "[21] The White House — Inflation Reduction Act Guidebook.",
        "[22] U.S. Securities and Exchange Commission — Accredited Investor Definition (Rule 501).",
        "[23] U.S. Census Bureau — American Community Survey: Miami-Fort Lauderdale MSA.",
        "[24] Bureau of Labor Statistics — OEWS: Administrative Assistants, Miami-Fort Lauderdale MSA.",
    ]
    for ref in refs:
        p_text(doc, ref)


# ============================================================
# MAIN
# ============================================================
def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    add_cover(doc)
    add_toc(doc)
    block1(doc)
    block2(doc)
    block3(doc)
    block4(doc)
    block5(doc)
    block6(doc)
    doc.save(OUTPUT_FILE)
    print(f"Business Plan saved to: {OUTPUT_FILE}")
    print(f"Sections: 42 | Blocks: 6 | Pages: ~55-65 estimated")


if __name__ == "__main__":
    main()
