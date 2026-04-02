#!/usr/bin/env python3
"""
Insert SaaS Screenshots into DOCX — Premium Edition
Petition Engine

Two modes:
  1. PLACEHOLDER MODE: Find [SCREENSHOT_XX] in existing DOCX and replace
  2. SMART MODE: Auto-detect section headings and insert after relevant sections

Usage:
    python3 insert_saas_screenshots.py <docx_path> <screenshots_dir> [--map screenshot_map.json] [--premium]

Flags:
    --map      Path to screenshot_map.json (for descriptions)
    --premium  Apply premium styling (header, footer, gold dividers, justified text)
    --client   Client name for header/footer
    --visa     Visa type for footer (default: EB-2 NIW)
"""

import os
import sys
import re
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.text.paragraph import Paragraph

# Design tokens
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xC9, 0xA9, 0x6E)
ACCENT_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
FONT = 'Garamond'


def find_screenshots(screenshots_dir):
    """Map screenshot numbers to file paths."""
    mapping = {}
    if not os.path.exists(screenshots_dir):
        return mapping
    for f in sorted(os.listdir(screenshots_dir)):
        if f.startswith('SaaS_') and f.endswith('.png'):
            match = re.match(r'SaaS_(\d+)_(.+)\.png', f)
            if match:
                num = match.group(1)
                name = match.group(2).replace('_', ' ')
                mapping[num] = {'path': os.path.join(screenshots_dir, f), 'name': name, 'file': f}
    return mapping


def load_map(map_file):
    """Load screenshot_map.json for descriptions."""
    if not map_file or not os.path.exists(map_file):
        return {}
    with open(map_file) as f:
        data = json.load(f)
    result = {}
    for page in data.get('pages', []):
        num = str(page['number']).zfill(2)
        result[num] = {
            'display_name': page.get('display_name', ''),
            'description': page.get('description', ''),
            'route': page.get('route', ''),
        }
    return result


def apply_premium_header_footer(doc, client_name, visa_type='EB-2 NIW'):
    """Add premium header and footer with gold borders."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()

        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        r1 = hp.add_run(client_name.upper())
        r1.font.name = FONT
        r1.font.size = Pt(9)
        r1.font.color.rgb = NAVY
        r1.bold = True

        hp.add_run('\t\t')
        r3 = hp.add_run('SaaS Evidence Dossier — Platform Documentation')
        r3.font.name = FONT
        r3.font.size = Pt(8)
        r3.font.color.rgb = MED_GRAY
        r3.italic = True

        pPr = hp._element.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:color="C9A96E" w:space="4"/>'
            f'</w:pBdr>'
        ))
        pPr.append(parse_xml(
            f'<w:tabs {nsdecls("w")}>'
            f'  <w:tab w:val="right" w:pos="9360"/>'
            f'</w:tabs>'
        ))

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()

        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        fpPr = fp._element.get_or_add_pPr()
        fpPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="6" w:color="C9A96E" w:space="4"/>'
            f'</w:pBdr>'
        ))
        fpPr.append(parse_xml(
            f'<w:tabs {nsdecls("w")}>'
            f'  <w:tab w:val="right" w:pos="9360"/>'
            f'</w:tabs>'
        ))

        rf1 = fp.add_run(f'{visa_type} Petition — {client_name}')
        rf1.font.name = FONT
        rf1.font.size = Pt(8)
        rf1.font.color.rgb = NAVY
        rf1.bold = True

        fp.add_run('\t\t')
        rf3 = fp.add_run('CONFIDENTIAL')
        rf3.font.name = FONT
        rf3.font.size = Pt(7)
        rf3.font.color.rgb = GOLD
        rf3.bold = True
        rf3.italic = True

    print("  ✓ Premium header + footer applied")


def justify_all_text(doc):
    """Justify all body paragraphs."""
    count = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        if 'heading' not in style_name.lower() and 'title' not in style_name.lower():
            if para.text.strip() and len(para.text.strip()) > 30:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                count += 1
    print(f"  ✓ {count} paragraphs justified")


def insert_figure_block(doc, anchor_element, img_path, fig_num, title, description='', evidence_value=''):
    """Insert a premium figure block after the anchor element."""
    num_str = str(fig_num).zfill(2)

    # Build in reverse order (addnext inserts immediately after anchor)

    # 6. Bottom spacer
    p6 = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="200"/></w:pPr></w:p>')
    anchor_element.addnext(p6)

    # 5. Evidence value (if provided)
    if evidence_value:
        p5 = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:before="40" w:after="80"/><w:ind w:left="360" w:right="360"/><w:jc w:val="both"/></w:pPr></w:p>')
        anchor_element.addnext(p5)
        para5 = Paragraph(p5, doc)
        r5a = para5.add_run("Evidentiary Value: ")
        r5a.font.name = FONT; r5a.font.size = Pt(9); r5a.font.color.rgb = NAVY; r5a.bold = True; r5a.italic = True
        r5b = para5.add_run(evidence_value)
        r5b.font.name = FONT; r5b.font.size = Pt(9); r5b.font.color.rgb = ACCENT_BLUE; r5b.italic = True

    # 4. Description (if provided)
    if description:
        p4 = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:before="80" w:after="40"/><w:ind w:left="360" w:right="360"/><w:jc w:val="both"/></w:pPr></w:p>')
        anchor_element.addnext(p4)
        para4 = Paragraph(p4, doc)
        r4 = para4.add_run(description)
        r4.font.name = FONT; r4.font.size = Pt(10); r4.font.color.rgb = DARK_GRAY

    # 3. Caption
    p3 = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:before="80" w:after="40"/><w:jc w:val="center"/></w:pPr></w:p>')
    anchor_element.addnext(p3)
    para3 = Paragraph(p3, doc)
    r3a = para3.add_run(f"Figure {num_str}")
    r3a.font.name = FONT; r3a.font.size = Pt(10); r3a.font.color.rgb = GOLD; r3a.bold = True
    r3b = para3.add_run(f" — {title}")
    r3b.font.name = FONT; r3b.font.size = Pt(10); r3b.font.color.rgb = NAVY; r3b.bold = True

    # 2. Image
    p2 = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:before="80" w:after="40"/><w:jc w:val="center"/></w:pPr></w:p>')
    anchor_element.addnext(p2)
    para2 = Paragraph(p2, doc)
    run_img = para2.add_run()
    run_img.add_picture(img_path, width=Inches(6.0))

    # 1. Gold divider
    p1 = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:spacing w:before="240" w:after="60"/>'
        f'    <w:pBdr><w:bottom w:val="single" w:sz="8" w:color="C9A96E"/></w:pBdr>'
        f'  </w:pPr>'
        f'</w:p>'
    )
    anchor_element.addnext(p1)


def insert_placeholder_mode(doc, screenshots, map_data):
    """Mode 1: Find [SCREENSHOT_XX] placeholders and replace."""
    inserted = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        match = re.match(r'\[SCREENSHOT_(\d+)\s*[—\-]\s*(.+?)\]', text)
        if not match:
            continue

        num = match.group(1)
        name = match.group(2).strip()

        if num not in screenshots:
            print(f"  ⚠ Missing screenshot for SCREENSHOT_{num}")
            continue

        info = map_data.get(num, {})
        description = info.get('description', '')

        para.clear()
        anchor = para._element

        insert_figure_block(doc, anchor, screenshots[num]['path'],
                          int(num), name, description)
        inserted += 1
        print(f"  ✓ Replaced [SCREENSHOT_{num} — {name}]")

    return inserted


def insert_smart_mode(doc, screenshots, map_data):
    """Mode 2: Auto-detect section headings and insert screenshots after sections."""
    # Build heading → screenshot mapping by analyzing content
    headings = []
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ''
        if 'heading' in style_name.lower() or 'Heading' in style_name:
            headings.append((i, para.text.strip(), style_name))

    if not headings:
        print("  ⚠ No headings found — cannot auto-insert")
        return 0

    # Map screenshots to headings by keyword matching
    screenshot_list = sorted(screenshots.items(), key=lambda x: x[0])
    inserted = 0

    for num, info in screenshot_list:
        img_path = info['path']
        name = info['name']
        map_info = map_data.get(num, {})
        description = map_info.get('description', '')
        display_name = map_info.get('display_name', name)

        # Find best heading match for this screenshot
        keywords = name.lower().replace('_', ' ').split()
        best_heading_idx = None
        best_score = 0

        for h_idx, h_text, h_style in headings:
            h_lower = h_text.lower()
            score = sum(1 for kw in keywords if kw in h_lower and len(kw) > 3)
            if score > best_score:
                best_score = score
                best_heading_idx = h_idx

        if best_heading_idx is None:
            # Fallback: insert at end of document
            best_heading_idx = len(doc.paragraphs) - 1

        # Find end of this section
        insert_idx = best_heading_idx
        for j in range(best_heading_idx + 1, min(best_heading_idx + 20, len(doc.paragraphs))):
            next_style = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ''
            if 'heading' in next_style.lower() and 'Heading' in next_style:
                break
            if doc.paragraphs[j].text.strip():
                insert_idx = j

        anchor = doc.paragraphs[insert_idx]._element
        insert_figure_block(doc, anchor, img_path, int(num), display_name, description)
        inserted += 1
        print(f"  ✓ Figure {num} — {display_name[:50]}...")

    return inserted


def main():
    args = sys.argv[1:]

    if len(args) < 2:
        print('Usage: python3 insert_saas_screenshots.py <docx_path> <screenshots_dir> [--map file.json] [--premium] [--client name] [--visa type]')
        sys.exit(1)

    docx_path = args[0]
    screenshots_dir = args[1]

    # Parse flags
    map_file = None
    premium = False
    client_name = 'Cliente'
    visa_type = 'EB-2 NIW'

    i = 2
    while i < len(args):
        if args[i] == '--map' and i + 1 < len(args):
            map_file = args[i + 1]; i += 2
        elif args[i] == '--premium':
            premium = True; i += 1
        elif args[i] == '--client' and i + 1 < len(args):
            client_name = args[i + 1]; i += 2
        elif args[i] == '--visa' and i + 1 < len(args):
            visa_type = args[i + 1]; i += 2
        else:
            i += 1

    if not os.path.exists(docx_path):
        print(f'ERROR: DOCX not found: {docx_path}')
        sys.exit(1)

    screenshots = find_screenshots(screenshots_dir)
    if not screenshots:
        print(f'ERROR: No SaaS_XX_*.png in {screenshots_dir}')
        sys.exit(1)

    map_data = load_map(map_file)

    print(f"\n{'='*60}")
    print(f"  SaaS Screenshot Inserter — Premium Edition")
    print(f"{'='*60}")
    print(f"  DOCX: {os.path.basename(docx_path)}")
    print(f"  Screenshots: {len(screenshots)}")
    print(f"  Premium: {'YES' if premium else 'no'}")
    print(f"  Client: {client_name}")
    print(f"{'='*60}\n")

    doc = Document(docx_path)

    # Apply premium styling
    if premium:
        apply_premium_header_footer(doc, client_name, visa_type)
        justify_all_text(doc)

    # Check if DOCX has [SCREENSHOT_XX] placeholders
    has_placeholders = any('[SCREENSHOT_' in p.text for p in doc.paragraphs)

    if has_placeholders:
        print("  Mode: PLACEHOLDER (found [SCREENSHOT_XX] markers)")
        inserted = insert_placeholder_mode(doc, screenshots, map_data)
    else:
        print("  Mode: SMART (auto-detecting section headings)")
        inserted = insert_smart_mode(doc, screenshots, map_data)

    # Save
    output_path = docx_path.replace('.docx', '_PREMIUM.docx')
    doc.save(output_path)
    size_mb = round(os.path.getsize(output_path) / 1024 / 1024, 1)

    print(f"\n{'='*60}")
    print(f"  INSERÇÃO COMPLETA")
    print(f"{'='*60}")
    print(f"  Inserted: {inserted}/{len(screenshots)}")
    print(f"  Output: {output_path}")
    print(f"  Size: {size_mb}MB")
    print(f"{'='*60}\n")


if __name__ == '__main__':
    main()
