#!/usr/bin/env python3
"""
AIOS SQUAD RUNNER — Gera documento e roda Quality Gate em loop até aprovação.
Rodada 1: Gera → Quality → Auto-fix → Salva
Rodada 2: Re-valida o fix → Encontra novos problemas → Auto-fix → Salva
Rodada 3: Re-valida → Se limpo, APROVADO. Se não, reporta pro Paulo.
Max: 5 rodadas.
"""
import subprocess, json, re, os, sys
from datetime import datetime

BP_SCRIPT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "build_bp_v7.py")
OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output", "BP_EventFinOps_LLC_V7.docx")
RULES_PATH = "/Users/paulo1844/Documents/OMNI/_IMIGRAÇÃO/Sistema Automatizado/petition-engine/data/error_rules.json"
MAX_ROUNDS = 5

# ============================================================
# QUALITY AGENT (Python version of quality-local.ts)
# ============================================================
COT_PATTERNS = [
    r'Vou estruturar.*', r'Agora vou redigir.*', r'Vou redigir.*',
    r'Vou elaborar.*', r'Vou pesquisar.*', r'Vou apresentar.*',
    r'Contagem de palavras.*', r'Com base nos dados.*redigir.*',
    r'Excelente\.?\s*Tenho dados.*', r'anti-alucinação absoluta.*',
    r'seção de prosa,?\s*dados verificados.*', r'respeitando rigorosamente.*',
    r'português brasileiro 100%.*', r'com base nas informações pesquisadas.*',
    r'em português brasileiro.*com rigor.*', r'— — — —',
    r'Agora vou estruturar.*', r'---{3,}',
]

ORPHAN_PATTERNS = [r'^Introdução$', r'^Conclusão$', r'^Mapa Perceptual$', r'^#{1,3}\s+']

FORBIDDEN_TERMS = ['PROEX', 'Kortix', 'Carlos Avelino', 'prompt']


def load_rules():
    try:
        with open(RULES_PATH) as f:
            return [r for r in json.load(f) if r.get('active')]
    except:
        return []


def extract_docx_text(path):
    from docx import Document
    doc = Document(path)
    return '\n'.join(p.text for p in doc.paragraphs)


def apply_fixes_to_docx(path, cot_patterns, orphan_patterns):
    """Apply auto-fixes directly to the DOCX paragraphs."""
    from docx import Document
    doc = Document(path)
    fixes = 0

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue

        should_clear = False

        # CoT removal
        for pat in cot_patterns:
            if re.match(pat, t, re.IGNORECASE):
                should_clear = True
                break

        # Orphan heading removal
        if not should_clear:
            for pat in orphan_patterns:
                if re.match(pat, t, re.MULTILINE):
                    should_clear = True
                    break

        if should_clear:
            for run in p.runs:
                run.text = ''
            fixes += 1
            continue

        # Inline CoT cleanup
        original = t
        for pat in cot_patterns:
            t = re.sub(pat, '', t, flags=re.IGNORECASE)
        t = re.sub(r'\s{2,}', ' ', t).strip()

        if t != original and p.runs:
            p.runs[0].text = t
            for run in p.runs[1:]:
                run.text = ''
            fixes += 1

    # Remove excessive blank paragraphs (3+ consecutive empty)
    empty_streak = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            empty_streak += 1
            if empty_streak >= 3:
                for run in p.runs:
                    run.text = ''
                fixes += 1
        else:
            empty_streak = 0

    doc.save(path)
    return fixes


def run_quality(text, doc_type='business_plan', client_name=''):
    violations = []
    rules = load_rules()

    # 1. Error rules with regex
    applicable = [r for r in rules if not r.get('doc_type') or r.get('doc_type') == doc_type]
    for rule in applicable:
        pat = rule.get('rule_pattern')
        if pat:
            try:
                matches = re.findall(pat, text, re.IGNORECASE)
                if matches:
                    violations.append({
                        'severity': rule['severity'],
                        'rule': rule['rule_description'][:80],
                        'count': len(matches),
                        'autoFixable': rule.get('rule_action') == 'auto_fix',
                    })
            except:
                pass

    # 2. CoT detection
    for pat in COT_PATTERNS:
        matches = re.findall(pat, text, re.IGNORECASE | re.MULTILINE)
        if matches:
            violations.append({
                'severity': 'critical',
                'rule': f'CoT leak ({len(matches)}x)',
                'count': len(matches),
                'autoFixable': True,
            })

    # 3. Orphan headings
    for pat in ORPHAN_PATTERNS:
        matches = re.findall(pat, text, re.MULTILINE)
        if matches:
            violations.append({
                'severity': 'high',
                'rule': f'Heading orfao ({matches[0][:20]})',
                'count': len(matches),
                'autoFixable': True,
            })

    # 4. Forbidden terms
    for term in FORBIDDEN_TERMS:
        if re.search(rf'\b{term}\b', text, re.IGNORECASE):
            violations.append({
                'severity': 'critical',
                'rule': f'Termo proibido: {term}',
                'count': 1,
                'autoFixable': False,
            })

    # 5. Long paragraphs
    long = [p for p in text.split('\n') if len(p.strip()) > 1500]
    if long:
        violations.append({
            'severity': 'medium',
            'rule': f'Paragrafos longos >1500 chars',
            'count': len(long),
            'autoFixable': False,
        })

    # 6. Duplicate section numbers
    section_nums = re.findall(r'^\d+\.\d+\.?\s', text, re.MULTILINE)
    seen = {}
    for sn in section_nums:
        sn = sn.strip()
        seen[sn] = seen.get(sn, 0) + 1
    dupes = {k: v for k, v in seen.items() if v > 1}
    if dupes:
        violations.append({
            'severity': 'high',
            'rule': f'Secoes duplicadas: {list(dupes.keys())[:3]}',
            'count': len(dupes),
            'autoFixable': False,
        })

    # 7. Client name check
    if client_name:
        first = client_name.split(' ')[0].lower()
        if first not in text.lower():
            violations.append({
                'severity': 'medium',
                'rule': f'Nome do cliente nao encontrado: {client_name}',
                'count': 1,
                'autoFixable': False,
            })

    # Score
    crit = sum(1 for v in violations if v['severity'] == 'critical')
    high = sum(1 for v in violations if v['severity'] == 'high')
    med = sum(1 for v in violations if v['severity'] == 'medium')
    low = sum(1 for v in violations if v['severity'] == 'low')
    score = max(0, 100 - crit * 25 - high * 15 - med * 5 - low * 2)
    passed = score >= 80 and crit == 0

    return {
        'score': score,
        'passed': passed,
        'violations': violations,
        'counts': {'critical': crit, 'high': high, 'medium': med, 'low': low},
    }


# ============================================================
# MAIN LOOP
# ============================================================
def main():
    print("=" * 70)
    print("AIOS SQUAD RUNNER — Quality Gate Loop")
    print(f"Documento: {OUTPUT}")
    print(f"Max rodadas: {MAX_ROUNDS}")
    print(f"Regras ativas: {len(load_rules())}")
    print("=" * 70)

    if not os.path.exists(OUTPUT):
        print("\n[FASE 0] Gerando documento...")
        result = subprocess.run([sys.executable, BP_SCRIPT], capture_output=True, text=True)
        if result.returncode != 0:
            print(f"ERRO na geracao: {result.stderr[:500]}")
            return
        print(f"Documento gerado: {OUTPUT}")

    for round_num in range(1, MAX_ROUNDS + 1):
        print(f"\n{'─' * 70}")
        print(f"RODADA {round_num}/{MAX_ROUNDS}")
        print(f"{'─' * 70}")

        # Extract text
        text = extract_docx_text(OUTPUT)
        char_count = len(text)
        para_count = len([p for p in text.split('\n') if p.strip()])
        print(f"  Caracteres: {char_count:,} | Paragrafos: {para_count}")

        # Run quality
        result = run_quality(text, 'business_plan', 'Maria Amália Vita')

        print(f"\n  SCORE: {result['score']}/100 — {'APROVADO' if result['passed'] else 'REPROVADO'}")
        print(f"  Violacoes: {result['counts']['critical']} criticas, {result['counts']['high']} altas, {result['counts']['medium']} medias, {result['counts']['low']} baixas")

        if result['violations']:
            print(f"\n  PROBLEMAS ENCONTRADOS:")
            for v in result['violations']:
                marker = "X" if v['severity'] == 'critical' else "!" if v['severity'] == 'high' else "~"
                fix = " [AUTO-FIX]" if v.get('autoFixable') else ""
                print(f"    [{marker}] [{v['severity'].upper()}] {v['rule']} ({v['count']}x){fix}")

        # Apply auto-fixes
        fixable = [v for v in result['violations'] if v.get('autoFixable')]
        if fixable:
            print(f"\n  APLICANDO AUTO-FIXES ({len(fixable)} problemas auto-corrigiveis)...")
            fixes = apply_fixes_to_docx(OUTPUT, COT_PATTERNS, ORPHAN_PATTERNS)
            print(f"  Fixes aplicados no DOCX: {fixes}")
        else:
            if result['passed']:
                print(f"\n  DOCUMENTO APROVADO — Score {result['score']}/100")
                print(f"  Nenhum fix necessario.")

                # Report non-fixable issues for Paulo
                non_fixable = [v for v in result['violations'] if not v.get('autoFixable')]
                if non_fixable:
                    print(f"\n  ITENS PARA REVISAO DO PAULO ({len(non_fixable)}):")
                    for v in non_fixable:
                        print(f"    → [{v['severity'].upper()}] {v['rule']} ({v['count']}x)")
                    print(f"\n  Esses itens devem ser apontados na plataforma /erros")
                    print(f"  para que o Engine aprenda e nao repita na proxima geracao.")

                break
            else:
                print(f"\n  REPROVADO — Problemas nao auto-corrigiveis encontrados.")
                print(f"  Esses problemas precisam de intervencao:")
                for v in result['violations']:
                    if not v.get('autoFixable'):
                        print(f"    → [{v['severity'].upper()}] {v['rule']}")
                print(f"\n  Aponte na plataforma /erros para que vire regra de auto-correcao.")
                break

    print(f"\n{'=' * 70}")
    print(f"SQUAD RUNNER FINALIZADO")
    print(f"Documento final: {OUTPUT}")
    print(f"{'=' * 70}")


if __name__ == "__main__":
    main()
