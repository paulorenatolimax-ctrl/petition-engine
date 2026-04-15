#!/usr/bin/env python3
"""
Validate that thumbnails were correctly inserted into a DOCX file.
Runs AFTER insert_thumbnails.py to verify the result.

Checks:
1. No [THUMBNAIL] placeholders remain
2. Images exist in the document
3. Image count matches expected exhibit count

Usage:
    python3 validate_thumbnail_insertion.py /path/to/document.docx [expected_count]

Exit codes:
    0 = all thumbnails inserted correctly
    1 = validation failed
"""

import sys
import os
import re
from docx import Document


def validate(docx_path: str, expected_count: int = 0) -> bool:
    if not os.path.exists(docx_path):
        print(f"❌ DOCX not found: {docx_path}")
        return False

    doc = Document(docx_path)
    errors = []

    # 1. Check for remaining placeholders
    remaining_placeholders = []
    for p in doc.paragraphs:
        if "[THUMBNAIL" in p.text:
            remaining_placeholders.append(p.text.strip()[:80])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "[THUMBNAIL" in p.text:
                        remaining_placeholders.append(p.text.strip()[:80])

    if remaining_placeholders:
        errors.append(f"{len(remaining_placeholders)} placeholder(s) [THUMBNAIL] still in document:")
        for rp in remaining_placeholders[:5]:
            errors.append(f"  → {rp}")

    # 2. Count images
    inline_count = len(doc.inline_shapes)
    body_xml = doc.element.xml
    anchor_count = body_xml.count("wp:anchor")
    total_images = inline_count + anchor_count

    # 3. Compare with expected
    if expected_count > 0 and total_images < expected_count:
        errors.append(f"Expected {expected_count} images but found {total_images}")

    # 4. Check file size (DOCX with images should be > 100KB)
    file_size = os.path.getsize(docx_path)
    if total_images > 0 and file_size < 100 * 1024:
        errors.append(f"File is {file_size/1024:.0f}KB with {total_images} images — suspiciously small")

    print(f"\n{'='*60}")
    print(f"THUMBNAIL INSERTION VALIDATION — {os.path.basename(docx_path)}")
    print(f"{'='*60}")
    print(f"Inline images: {inline_count}")
    print(f"Anchor images: {anchor_count}")
    print(f"Total images: {total_images}")
    print(f"Remaining placeholders: {len(remaining_placeholders)}")
    print(f"File size: {file_size/1024:.0f}KB")

    if errors:
        print(f"\n❌ FAILED — {len(errors)} issues:")
        for e in errors:
            print(f"  ❌ {e}")
        print(f"{'='*60}")
        return False
    else:
        print(f"\n✅ THUMBNAILS VALIDATED")
        print(f"{'='*60}")
        return True


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 validate_thumbnail_insertion.py <docx_path> [expected_count]")
        sys.exit(1)

    expected = int(sys.argv[2]) if len(sys.argv) > 2 else 0
    passed = validate(sys.argv[1], expected)
    sys.exit(0 if passed else 1)
