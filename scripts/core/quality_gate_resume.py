#!/usr/bin/env python3
"""
Quality Gate — Résumé EB-2 NIW / EB-1A
Verifica mecanicamente se o DOCX gerado atende aos padrões mínimos.
Se FALHAR, imprime os erros e retorna exit code 1.
Se PASSAR, retorna exit code 0.

Uso:
    python3 quality_gate_resume.py /caminho/do/resume.docx
"""

import sys
import os
import re
from docx import Document

def check(doc_path):
    if not os.path.exists(doc_path):
        print(f"❌ ARQUIVO NÃO ENCONTRADO: {doc_path}")
        return False

    doc = Document(doc_path)
    text_paras = [p.text for p in doc.paragraphs]
    full_text = '\n'.join(text_paras)
    table_text = ''
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                table_text += cell.text + '\n'

    all_text = full_text + '\n' + table_text
    word_count = len(full_text.split())

    errors = []
    warnings = []

    # ═══════════════════════════════════════════════════════════
    # BLOQUEANTES (exit code 1)
    # ═══════════════════════════════════════════════════════════

    # 1. Mínimo de palavras
    if word_count < 3500:
        errors.append(f"PALAVRAS: {word_count} (mínimo 3.500)")

    # 2. "Exhibit" no corpo (exclusivo de Cover Letter)
    # Ignorar header (primeiros 500 chars)
    body_text = full_text[500:] + '\n' + table_text
    exhibit_count = len(re.findall(r'\bExhibit\s+\d', body_text))
    if exhibit_count > 0:
        errors.append(f"'Exhibit N' encontrado {exhibit_count}x no corpo (usar 'Evidência')")

    # 3. Placeholders
    placeholders = re.findall(r'\[(VERIFICAR|TODO|INSERIR|PENDENTE|TBD|XXX|COMPLETAR|PREENCHER)\]', all_text, re.I)
    if placeholders:
        errors.append(f"PLACEHOLDERS encontrados: {placeholders}")

    # 4. Termos proibidos
    prohibited = {
        'PROEX': r'\bPROEX\b',
        'Kortix': r'\bKortix\b',
        'Petition Engine': r'Petition Engine',
        'I believe': r'\bI\s+believe\b',
        'we believe': r'\bwe\s+believe\b',
        'in conclusion': r'\bin conclusion\b',
        'consultoria': r'\bconsultoria\b',
    }
    for term, pattern in prohibited.items():
        matches = re.findall(pattern, all_text, re.I)
        if matches:
            errors.append(f"TERMO PROIBIDO '{term}': {len(matches)}x")

    # 5. SOC code no corpo (só permitido no header)
    soc_in_body = re.findall(r'\b\d{2}-\d{4}(?:\.\d{2})?\b', body_text)
    if soc_in_body:
        errors.append(f"CÓDIGO SOC no corpo: {soc_in_body[:3]}... (só permitido no header)")

    # 6. Anti-Cristine
    anticristine = re.findall(r'\b(self-sustaining|operates autonomously|plug.and.play|scalable without|turnkey|white.label)\b', all_text, re.I)
    if anticristine:
        errors.append(f"ANTI-CRISTINE: {anticristine}")

    # 7. Cartas citadas no résumé
    carta_refs = re.findall(r'(?i)(conforme carta|como atestado|segundo recomenda|carta de recomenda|recommendation letter|as attested by)', all_text)
    if carta_refs:
        errors.append(f"CARTA CITADA NO RÉSUMÉ: {carta_refs}")

    # ═══════════════════════════════════════════════════════════
    # ESTRUTURAIS (exit code 1)
    # ═══════════════════════════════════════════════════════════

    # 8. Principais Responsabilidades
    resp_count = len(re.findall(r'(?i)principa(is|l)\s+responsabilidade', all_text))
    if resp_count == 0:
        errors.append("ZERO seções 'Principais Responsabilidades' (obrigatório por experiência)")

    # 9. Principais Resultados/Impactos
    result_count = len(re.findall(r'(?i)principa(is|l)\s+(resultado|impacto)', all_text))
    if result_count == 0:
        errors.append("ZERO seções 'Principais Resultados/Impactos' (obrigatório por experiência)")

    # 10. Síntese Profissional
    sintese = re.search(r'(?i)s[ií]ntese\s+profissional', all_text)
    if not sintese:
        errors.append("SEÇÃO 'Síntese Profissional' NÃO ENCONTRADA")

    # 11. Paginação (footer)
    has_footer = False
    for section in doc.sections:
        footer = section.footer
        if footer and footer.paragraphs:
            for p in footer.paragraphs:
                if p.text.strip():
                    has_footer = True
                    break
    if not has_footer:
        errors.append("FOOTER VAZIO — sem paginação 'Page X of Y'")

    # 12. Thumbnail placeholders (devem existir se é résumé com evidências)
    thumbnails = len(re.findall(r'\[THUMBNAIL', all_text))
    if thumbnails == 0:
        warnings.append("ZERO placeholders [THUMBNAIL] — résumé sem evidências visuais?")

    # ═══════════════════════════════════════════════════════════
    # WARNINGS (não bloqueiam)
    # ═══════════════════════════════════════════════════════════

    # 13. Fontes
    fonts = set()
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.name:
                fonts.add(r.font.name)
    if fonts and not any('Garamond' in f for f in fonts):
        warnings.append(f"FONTE: Garamond NÃO encontrado. Fontes: {fonts}")

    # 14. Inferências
    inference_words = len(re.findall(r'(?i)(demonstra|evidencia|comprova|transferível|relevante para|competência|capacidade)', all_text))
    if inference_words < 10:
        warnings.append(f"INFERÊNCIAS RASAS: apenas {inference_words} palavras de inferência (mínimo 10)")

    # ═══════════════════════════════════════════════════════════
    # RESULTADO
    # ═══════════════════════════════════════════════════════════

    print(f"\n{'='*60}")
    print(f"QUALITY GATE — {os.path.basename(doc_path)}")
    print(f"{'='*60}")
    print(f"Palavras: {word_count:,}")
    print(f"Parágrafos: {len(text_paras)}")
    print(f"Tabelas: {len(doc.tables)}")
    print(f"Thumbnails: {thumbnails}")
    print(f"Responsabilidades: {resp_count}")
    print(f"Resultados/Impactos: {result_count}")
    print(f"Fontes: {fonts if fonts else 'N/A'}")

    if errors:
        print(f"\n❌ REPROVADO — {len(errors)} erros bloqueantes:")
        for e in errors:
            print(f"  ❌ {e}")

    if warnings:
        print(f"\n⚠️ {len(warnings)} avisos:")
        for w in warnings:
            print(f"  ⚠️ {w}")

    if not errors:
        print(f"\n✅ APROVADO")

    print(f"{'='*60}\n")

    return len(errors) == 0


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Uso: python3 quality_gate_resume.py <caminho_do_docx>")
        sys.exit(1)

    passed = check(sys.argv[1])
    sys.exit(0 if passed else 1)
