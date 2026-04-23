#!/usr/bin/env python3
"""Validador determinístico bloqueante do dossier SaaS Evidence.

Uso:
    python3 validate_saas_evidence.py <docx> <lovable_spec.md> <evidence_summary.json> \
            <screenshots_dir> [report.md]

Exit 0 OK · exit 1 errors.

Checa:
- DOCX existe, bytes ≥ 300KB
- DOCX tem ≥10 inline_shapes (screenshots injetados)
- DOCX zero placeholders [SCREENSHOT_XX] / [TO BE VERIFIED] / [INSERIR]
- DOCX zero termos imigratórios (regra r141)
- DOCX zero termos anti-Prong-3 (ANTI-CRISTINE V2)
- DOCX zero "consulting" isolado (NAICS VIBE trigger)
- DOCX zero exposição infra (PROEX, Kortix, Petition Engine, Obsidian, RAG)
- LOVABLE_BUILD_SPEC.md existe + ≥200 linhas + tem as 7 seções mandatórias
- evidence_summary.json tem founder_dependency="critical"
- screenshots/ tem ≥10 PNGs, nenhum < 50KB
"""
import json
import os
import re
import sys
import zipfile

try:
    from docx import Document
except ImportError:
    sys.stderr.write("ERR: pip install python-docx\n")
    sys.exit(2)


IMMIGRATION_TERMS = [
    "petition", "petitioner", "peticionário", "USCIS", "visa", "green card",
    "I-140", "adjudicator", "Dhanasar", "Kazarian", "extraordinary ability",
    "National Interest Waiver", "NIW", "EB-2", "EB-1", "immigration",
    "Prong", "prong"
]

ANTI_CRISTINE_TERMS = [
    "standardized", "padronizado", "operates autonomously", "opera autonomamente",
    "self-sustaining", "auto-sustentável", "autossuficiente", "plug-and-play",
    "train-the-trainer", "white-label", "marca branca", "client autonomy",
    "autonomia do cliente", "founder dependency: low", "scalable without the founder",
    "scalable without founder", "replicable by any professional", "turnkey",
    "chave na mão", "chave-na-mão"
]

INFRA_EXPOSURE = [
    "PROEX", "Kortix", "Petition Engine", "Forjado por", "Obsidian",
    "RAG I", "RAG II", "RAG III", "gerado automaticamente",
    "gerado por Claude", "gerado por IA"
]

PLACEHOLDER_PATTERNS = [
    r"\[SCREENSHOT_\d+\]", r"\[TO BE VERIFIED[^\]]*\]", r"\[INSERIR[^\]]*\]",
    r"\[TODO[^\]]*\]", r"\[VERIFICAR[^\]]*\]", r"\[TBD[^\]]*\]",
    r"\[PENDENTE[^\]]*\]", r"\[XXX+[^\]]*\]", r"Lorem ipsum"
]

# Lovable spec required sections
LOVABLE_REQUIRED_SECTIONS = [
    r"WHAT TO BUILD",
    r"TECH STACK",
    r"COLOR PALETTE",
    r"TYPOGRAPHY",
    r"PAGES TO BUILD",
    r"(DATA TO SHOW|DESIGN PRINCIPLES)",
    r"CRITICAL RULES"
]


def _err(lst, msg):
    lst.append(msg)


def extract_all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        for container in (section.header, section.footer):
            if container:
                for p in container.paragraphs:
                    parts.append(p.text)
    return "\n".join(p for p in parts if p)


def check_docx_immigration(text, errors, passed):
    leaked = []
    for term in IMMIGRATION_TERMS:
        matches = re.findall(rf"\b{re.escape(term)}\b", text, re.IGNORECASE)
        if matches:
            leaked.append((term, len(matches)))
    if leaked:
        for term, n in leaked:
            _err(errors, f"immigration term leak: '{term}' ({n}x)")
    else:
        _err(passed, "zero immigration terminology")


def check_docx_anti_cristine(text, errors, passed):
    leaked = []
    for term in ANTI_CRISTINE_TERMS:
        if re.search(rf"\b{re.escape(term)}\b", text, re.IGNORECASE):
            leaked.append(term)
    if leaked:
        for l in leaked:
            _err(errors, f"anti-Cristine term leak: '{l}' (destroys Prong 3)")
    else:
        _err(passed, "zero anti-Prong-3 terminology")


def check_docx_consulting(text, errors, warnings, passed):
    # "consulting" isolado — mas permite "specialized technical consulting services"
    # Check if appears as standalone without qualifiers
    matches = re.findall(r"\bconsulting\b(?!\s+(?:services|tier|engagement|framework))", text, re.IGNORECASE)
    if matches:
        _err(warnings, f"'consulting' isolated {len(matches)}x — triggers USCIS VIBE NAICS flag")
    else:
        _err(passed, "no isolated 'consulting' (NAICS VIBE safe)")


def check_docx_infra(text, errors, passed):
    leaked = []
    for term in INFRA_EXPOSURE:
        if re.search(rf"\b{re.escape(term)}\b", text, re.IGNORECASE):
            leaked.append(term)
    if leaked:
        for l in leaked:
            _err(errors, f"infrastructure exposure: '{l}'")
    else:
        _err(passed, "zero internal infra exposure")


def check_docx_placeholders(text, errors, passed):
    found = []
    for pat in PLACEHOLDER_PATTERNS:
        m = re.search(pat, text)
        if m:
            found.append(m.group())
    if found:
        for f in found:
            _err(errors, f"unfilled placeholder: {f}")
    else:
        _err(passed, "zero placeholders (all content filled)")


def check_docx_images(docx_path, doc, errors, warnings, passed):
    imgs = len(doc.inline_shapes)
    if imgs < 10:
        _err(errors, f"inline_shapes={imgs} < 10 (screenshots não foram injetados ou capture falhou)")
    else:
        _err(passed, f"inline_shapes={imgs} ≥ 10 (screenshots presentes)")

    # Count via zip for redundancy
    try:
        with zipfile.ZipFile(docx_path) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/") and n.lower().endswith((".png", ".jpg", ".jpeg"))]
            if len(media) < 10:
                _err(warnings, f"word/media/ has {len(media)} files — expected ≥10")
    except Exception as e:
        _err(warnings, f"zip inspection failed: {e}")


def check_lovable_spec(spec_path, errors, passed):
    if not os.path.isfile(spec_path):
        _err(errors, f"LOVABLE_BUILD_SPEC.md not found at {spec_path}")
        return
    with open(spec_path, encoding="utf-8") as f:
        content = f.read()
    lines = content.splitlines()
    if len(lines) < 200:
        _err(errors, f"LOVABLE_BUILD_SPEC.md has {len(lines)} lines < 200 minimum (prompt too shallow)")
    else:
        _err(passed, f"LOVABLE_BUILD_SPEC.md has {len(lines)} lines ≥ 200")

    missing = []
    for pattern in LOVABLE_REQUIRED_SECTIONS:
        if not re.search(pattern, content, re.IGNORECASE):
            missing.append(pattern)
    if missing:
        for m in missing:
            _err(errors, f"LOVABLE_BUILD_SPEC.md missing required section: {m}")
    else:
        _err(passed, "LOVABLE_BUILD_SPEC.md has all 7 required sections")

    # Check color palette includes the brand hex
    if "#1B2A4A" in content or "1B2A4A" in content:
        _err(passed, "LOVABLE spec includes brand navy #1B2A4A")
    else:
        _err(errors, "LOVABLE spec missing brand navy #1B2A4A (Phase 7 V2 requirement)")


def check_evidence_summary(summary_path, errors, passed):
    if not os.path.isfile(summary_path):
        _err(errors, f"evidence_summary.json not found at {summary_path}")
        return
    try:
        data = json.load(open(summary_path, encoding="utf-8"))
    except Exception as e:
        _err(errors, f"evidence_summary.json unreadable: {e}")
        return

    platform_proof = data.get("platform_proof", {})
    fd = platform_proof.get("founder_dependency")
    if fd != "critical":
        _err(errors, f"evidence_summary.json founder_dependency='{fd}' — MUST be 'critical' (Anti-Cristine)")
    else:
        _err(passed, "founder_dependency = 'critical' ✓")


def check_screenshots_folder(folder, errors, warnings, passed):
    if not os.path.isdir(folder):
        _err(errors, f"screenshots/ folder not found at {folder}")
        return
    pngs = [f for f in os.listdir(folder) if f.lower().endswith(".png") and f.startswith("SaaS_")]
    if len(pngs) < 10:
        _err(errors, f"screenshots/ has {len(pngs)} SaaS_*.png files < 10 minimum")
    else:
        _err(passed, f"screenshots/ has {len(pngs)} SaaS_*.png files ≥ 10")
    # Size check — each ≥ 50KB
    too_small = [f for f in pngs if os.path.getsize(os.path.join(folder, f)) < 50_000]
    if too_small:
        _err(warnings, f"{len(too_small)} screenshots < 50KB (may be blank pages): {too_small[:3]}")


def main():
    if len(sys.argv) < 5:
        sys.stderr.write(
            "Usage: validate_saas_evidence.py <docx> <lovable_spec.md> "
            "<evidence_summary.json> <screenshots_dir> [report.md]\n"
        )
        sys.exit(2)

    docx_path = sys.argv[1]
    spec_path = sys.argv[2]
    summary_path = sys.argv[3]
    shots_dir = sys.argv[4]
    report_path = sys.argv[5] if len(sys.argv) > 5 else None

    errors, warnings, passed = [], [], []

    # DOCX checks
    if not os.path.isfile(docx_path):
        _err(errors, f"DOCX not found: {docx_path}")
    else:
        size = os.path.getsize(docx_path)
        if size < 300_000:
            _err(errors, f"DOCX size {size:,} < 300KB (no screenshots embedded?)")
        else:
            _err(passed, f"DOCX size {size:,} bytes ≥ 300KB")
        try:
            doc = Document(docx_path)
            text = extract_all_text(doc)
            check_docx_immigration(text, errors, passed)
            check_docx_anti_cristine(text, errors, passed)
            check_docx_consulting(text, errors, warnings, passed)
            check_docx_infra(text, errors, passed)
            check_docx_placeholders(text, errors, passed)
            check_docx_images(docx_path, doc, errors, warnings, passed)
        except Exception as e:
            _err(errors, f"DOCX open failed: {e}")

    check_lovable_spec(spec_path, errors, passed)
    check_evidence_summary(summary_path, errors, passed)
    check_screenshots_folder(shots_dir, errors, warnings, passed)

    status = "PASSED" if not errors else "FAILED"
    lines = [
        f"# SaaS Evidence Validator — {status}",
        "",
        f"**DOCX:** `{docx_path}`",
        f"**Checks:** {len(passed)} passed · {len(warnings)} warnings · {len(errors)} errors",
        ""
    ]
    if passed:
        lines.append("## ✓ Passed")
        for p in passed:
            lines.append(f"- {p}")
        lines.append("")
    if warnings:
        lines.append("## ⚠ Warnings (non-blocking)")
        for w in warnings:
            lines.append(f"- {w}")
        lines.append("")
    if errors:
        lines.append("## ✗ Errors (BLOCKING)")
        for e in errors:
            lines.append(f"- {e}")
        lines.append("")

    report = "\n".join(lines)
    print(report)
    if report_path:
        os.makedirs(os.path.dirname(os.path.abspath(report_path)) or ".", exist_ok=True)
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(report + "\n")

    sys.exit(1 if errors else 0)


if __name__ == "__main__":
    main()
