#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SaaS Evidence Architect — Cristine Correa (REFILE V1)
Generates: V1_saas_evidence_Cristine_Correa.docx + LOVABLE_BUILD_SPEC.md + thumbnail_map.json

Strategy: Reconfigure 100% from the failed Talent Anchor OS.
Brand: Correa Workforce Sciences, LLC + ClimateLens (PCRM Framework)
"""

import json
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path("/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2026/CRISTINE CORREA/_Forjado por Petition Engine")
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "V1_saas_evidence_Cristine_Correa.docx"
LOVABLE_PATH = OUT_DIR / "LOVABLE_BUILD_SPEC.md"
THUMB_PATH = OUT_DIR / "thumbnail_map.json"

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
SLATE = RGBColor(0x4B, 0x55, 0x63)
INK = RGBColor(0x10, 0x14, 0x1F)
PAPER = RGBColor(0xFA, 0xFB, 0xFC)
SOFT = RGBColor(0xF5, 0xF1, 0xE6)
RULE = RGBColor(0xE3, 0xE7, 0xEC)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_borders(cell, color="E3E7EC", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def no_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_para(doc, text, *, size=11, bold=False, italic=False, color=INK,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4,
             line_spacing=1.35, font='Calibri'):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    return p


def add_runs(doc, runs, *, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0,
             space_after=4, line_spacing=1.35, font='Calibri'):
    """runs: list of (text, bold, italic, color)"""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    for text, bold, italic, color in runs:
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
    return p


def add_section_break(doc):
    from docx.enum.section import WD_SECTION
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_page_break(doc):
    doc.add_page_break()


def h1(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text.upper())
    r.font.name = 'Georgia'
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = color
    # Gold underline
    add_gold_rule(doc)


def h2(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Georgia'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = color


def h3(doc, text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = color


def add_gold_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), 'B8860B')
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_kpi_card(doc, kpis):
    """kpis: list of (value, label) tuples — renders as a 1-row table."""
    table = doc.add_table(rows=2, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (value, label) in enumerate(kpis):
        c1 = table.cell(0, i)
        c2 = table.cell(1, i)
        shade_cell(c1, '1B2A4A')
        shade_cell(c2, 'F5F1E6')
        no_borders(c1)
        no_borders(c2)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(8)
        p1.paragraph_format.space_after = Pt(8)
        r1 = p1.add_run(value)
        r1.font.name = 'Georgia'
        r1.font.size = Pt(20)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(label)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9)
        r2.font.bold = True
        r2.font.color.rgb = NAVY


def add_callout(doc, label, body, accent_hex='B8860B'):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(0.4)
    table.columns[1].width = Cm(15.5)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    shade_cell(left, accent_hex)
    shade_cell(right, 'FAFBFC')
    no_borders(left); no_borders(right)
    left.text = ''
    right.paragraphs[0].text = ''
    pl = right.paragraphs[0]
    pl.paragraph_format.space_before = Pt(6)
    pl.paragraph_format.space_after = Pt(0)
    rl = pl.add_run(label.upper())
    rl.font.name = 'Calibri'
    rl.font.size = Pt(9)
    rl.font.bold = True
    rl.font.color.rgb = NAVY
    pb = right.add_paragraph()
    pb.paragraph_format.space_before = Pt(2)
    pb.paragraph_format.space_after = Pt(8)
    rb = pb.add_run(body)
    rb.font.name = 'Calibri'
    rb.font.size = Pt(11)
    rb.font.italic = False
    rb.font.color.rgb = INK


def add_pullquote(doc, quote, attribution=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), '24')
    left_b.set(qn('w:color'), 'B8860B')
    left_b.set(qn('w:space'), '8')
    pBdr.append(left_b)
    pPr.append(pBdr)
    r = p.add_run(quote)
    r.font.name = 'Georgia'
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = NAVY
    if attribution:
        add_para(doc, attribution, size=9, italic=True, color=SLATE,
                 space_before=2, space_after=10)


def add_bullets(doc, items, size=11):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3
        r = p.add_run(it)
        r.font.name = 'Calibri'
        r.font.size = Pt(size)
        r.font.color.rgb = INK


def styled_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            table.columns[i].width = Cm(w)
    # header row
    for i, h in enumerate(headers):
        c = table.cell(0, i)
        shade_cell(c, '1B2A4A')
        set_borders(c, color='1B2A4A', size='8')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(h)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # body rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.cell(1 + ri, ci)
            shade_cell(c, 'FAFBFC' if ri % 2 == 0 else 'FFFFFF')
            set_borders(c, color='E3E7EC', size='6')
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.25
            r = p.add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            r.font.color.rgb = INK


def set_page_margins(section, top=2.2, right=2.0, bottom=2.2, left=2.0):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def add_footer(doc, label):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.font.name = 'Calibri'
    r.font.size = Pt(8)
    r.font.color.rgb = SLATE


def build_cover(doc):
    section = doc.sections[0]
    set_page_margins(section, top=2.5, right=2.0, bottom=2.5, left=2.0)

    # Top brand mark
    add_para(doc, "CORREA WORKFORCE SCIENCES", size=10, bold=True, color=GOLD,
             space_before=40, space_after=2, line_spacing=1.0)
    add_para(doc, "An Applied Research Lab", size=9, italic=True, color=SLATE,
             space_before=0, space_after=60, line_spacing=1.0)

    # Title
    add_para(doc, "ClimateLens", size=44, bold=True, color=NAVY,
             space_before=0, space_after=4, line_spacing=1.0, font='Georgia')
    add_para(doc, "Predictive Climate & Retention Intelligence",
             size=16, italic=True, color=SLATE,
             space_before=0, space_after=4, line_spacing=1.1, font='Georgia')
    add_para(doc, "for American Small and Mid-Sized Employers",
             size=14, italic=True, color=SLATE,
             space_before=0, space_after=20, line_spacing=1.1, font='Georgia')
    add_gold_rule(doc)

    add_para(doc, "PRODUCT & SCIENTIFIC DOSSIER", size=10, bold=True,
             color=GOLD, space_before=20, space_after=4, line_spacing=1.1)
    add_para(doc, "Powered by the PCRM Framework™ — Psychometric Climate & Retention Methodology",
             size=11, italic=True, color=INK,
             space_before=0, space_after=200, line_spacing=1.3)

    # Bottom block
    add_para(doc, "Founded and led by Cristine Correa, M.B.A.",
             size=11, bold=True, color=NAVY, space_before=0, space_after=2)
    add_para(doc, "Registered Organizational Psychologist (CRP-12/11050) · Predictive Index Certified Analyst",
             size=9, italic=True, color=SLATE, space_before=0, space_after=2)
    add_para(doc, "AMA · SHRM · 10+ Years in Behavioral Talent Sciences",
             size=9, italic=True, color=SLATE, space_before=0, space_after=20)
    add_para(doc, "Sanford · Seminole County · Central Florida",
             size=10, bold=True, color=NAVY, space_before=4, space_after=2)
    add_para(doc, "correa.cristine@gmail.com",
             size=9, color=SLATE, space_before=0, space_after=4)

    add_page_break(doc)


def build_executive_summary(doc):
    h1(doc, "Executive Summary")

    add_para(doc,
             "Correa Workforce Sciences is an applied behavioral research lab serving American "
             "small and mid-sized employers. The company designs, calibrates, and operates "
             "ClimateLens — a proprietary climate-and-retention intelligence platform powered "
             "by the PCRM Framework™ (Psychometric Climate & Retention Methodology), developed "
             "and continuously evolved by founder Cristine Correa across more than a decade of "
             "behavioral talent science practice.",
             size=11, space_after=8)

    add_para(doc,
             "The lab addresses a structural gap in the American labor market: while "
             "enterprise-grade behavioral platforms (Hogan, Gallup, DDI, Predictive Index, Culture Amp) "
             "serve large employers at five and six-figure annual contracts, the 33.2 million "
             "small and mid-sized employers that account for 46.4 percent of the private workforce "
             "remain unable to access psychometrically rigorous climate, retention, and behavioral "
             "intelligence services. ClimateLens closes that gap by combining clinically grounded "
             "instrument design, accessible analytics, and supervised behavioral interpretation — "
             "delivered under the continuous methodological direction of its founder.",
             size=11, space_after=10)

    add_kpi_card(doc, [
        ("$1T", "Annual cost of voluntary turnover to U.S. employers (Gallup)"),
        ("46.4%", "Of U.S. private workforce employed by small and mid-sized employers (SBA)"),
        ("0.49", "True correlation between engagement and composite performance (Gallup)"),
        ("22.7%", "Annual turnover in U.S. healthcare (BLS)"),
    ])

    add_para(doc, " ", size=4, space_after=2)
    add_callout(doc, "Strategic Position",
                "ClimateLens is the only behavioral retention lab in Central Florida that combines "
                "psychometric instrument design, predictive retention analytics, and behaviorally "
                "supervised intervention plans in a single research-driven engagement built for "
                "employers with 10 to 500 employees.")

    add_page_break(doc)


def build_problem(doc):
    h1(doc, "I. The Trillion-Dollar Problem")
    h2(doc, "Voluntary Turnover Has Become a Permanent Drag on the American Economy")

    add_para(doc,
             "American employers lose roughly one trillion dollars every year to voluntary "
             "turnover (Gallup, State of the Global Workplace, 2024). The Bureau of Labor "
             "Statistics' JOLTS series confirms that the quits rate has remained above 2 percent "
             "monthly since 2021, with 3.3 to 3.5 million voluntary separations per month. "
             "Healthcare employers lose 22.7 percent of their workforce annually; hospitality "
             "and food services employers lose 50 to 80 percent; small and mid-sized "
             "manufacturers lose 24 to 32 percent. Each replacement consumes 50 to 200 percent "
             "of the departing worker's annual salary in recruiting, onboarding, productivity "
             "ramp, and lost institutional knowledge.",
             size=11, space_after=8)

    h3(doc, "The science is settled; the access is not")
    add_para(doc,
             "A meta-analysis spanning 183,806 business units and 3.35 million employees "
             "(Gallup, 2024) found a true correlation of 0.49 between engagement, measured "
             "through psychometrically sound climate instruments, and composite operational "
             "performance. Top-quartile-engagement teams produce 23 percent higher profitability, "
             "18 percent higher productivity, and 81 percent fewer absenteeism days. Climate "
             "diagnostics, retention analytics, and supervised behavioral interventions are not "
             "experimental: they are the most evidence-backed managerial discipline of the last "
             "two decades.",
             size=11, space_after=8)

    add_pullquote(doc,
                  "The output of mainstream platforms is numbers. The output of ClimateLens "
                  "is interpreted, supervised behavioral change.",
                  "— PCRM Framework™ Design Principle")

    h3(doc, "Where the market fails")
    add_para(doc,
             "Mainstream climate and retention platforms — Culture Amp, Lattice, Glint/Viva, "
             "15Five, Qualtrics — are calibrated for mid-market and enterprise employers. "
             "Annual minimums of US$ 3,000 to US$ 10,000, per-seat fees of US$ 8 to US$ 30 "
             "per month, and a self-service interpretation model exclude the 33.2 million "
             "American small and mid-sized employers (SBA, 2024) that employ 46.4 percent of "
             "the private workforce. Even when smaller employers can afford the software, the "
             "missing layer is the same: behavioral interpretation of the data and a "
             "scientifically supervised intervention plan. Numbers without behavioral science "
             "do not change retention.",
             size=11, space_after=10)

    add_page_break(doc)


def build_solution(doc):
    h1(doc, "II. The ClimateLens Solution")
    h2(doc, "Three Integrated Modules. One Continuously Calibrated Methodology.")

    add_para(doc,
             "ClimateLens is a single integrated engagement, never sold as software-only. Every "
             "deployment includes the methodological supervision of the founding scientist and "
             "the continuous calibration of the underlying psychometric instruments. The product "
             "is the science applied — not a generic survey tool resold under a new name.",
             size=11, space_after=8)

    # Three modules table
    h3(doc, "The three modules")
    styled_table(doc,
                 ["Module", "Description", "Founder Supervision"],
                 [
                     ["Module 01 — Climate Diagnostic",
                      "Pulse and annual surveys built on the five domains of the U.S. Surgeon "
                      "General's Framework for Mental Health & Well-Being at Work (protection "
                      "from harm, opportunity for growth, connection and community, mattering "
                      "at work, work–life harmony). Instruments calibrated for SME industries "
                      "and benchmarked against Bureau of Labor Statistics turnover baselines.",
                      "Instrument design, item calibration, factor revalidation every release cycle."],
                     ["Module 02 — Predictive Retention Analytics",
                      "Behavioral analytics layer that flags flight risk and burnout patterns "
                      "three to six months in advance. Produces unit-level, role-level and "
                      "manager-level early-warning indicators with confidence bands.",
                      "Algorithmic supervision, model retraining, false-positive review."],
                     ["Module 03 — Supervised Behavioral Intervention",
                      "Manager-facing playbooks: conversation scripts, micro-metrics, "
                      "follow-up protocols. Each intervention plan is signed off by the "
                      "founder before deployment and reviewed at four-week intervals.",
                      "Plan authorship, manager coaching, intervention auditing."],
                 ],
                 col_widths_cm=[4.0, 9.0, 4.5])

    add_callout(doc, "Engagement Principle",
                "ClimateLens is never deployed as autonomous software. The platform exists "
                "to deliver the founder's research and supervised judgment at scale — not to "
                "replace it. Modules 01 through 03 are continuously evolved by Cristine Correa "
                "with each release cycle.")

    add_page_break(doc)


def build_founder(doc):
    h1(doc, "III. The Scientific Architect")
    h2(doc, "Cristine Correa — Founder, Principal Scientist, Methodological Director")

    add_para(doc,
             "ClimateLens exists because of the convergence of three rare credentials in a "
             "single practitioner: clinical training in Psychology, an MBA in Executive People "
             "Management with a thesis on personality pathology in corporate environments, and "
             "more than a decade of progressive corporate practice across high-attrition "
             "American and Brazilian employers. Cristine Correa is the lab's principal "
             "scientist, instrument designer, and methodological director.",
             size=11, space_after=8)

    h3(doc, "Academic foundation")
    add_bullets(doc, [
        "Bachelor of Psychology, Associação Catarinense de Ensino — Faculdade Guilherme Guimbala (Joinville, SC, Brazil), conferred 3 March 2012. Concentration in Organizational Psychology; 4,212 instructional hours including supervised internships in occupational, educational, community, and clinical psychology.",
        "Lato Sensu Postgraduate Specialization in Psychological Assessment, Faculdade Guilherme Guimbala, conferred 27 October 2013. Capstone research on professional psychological practice with vulnerable populations.",
        "M.B.A. — Executive People Management, Managerial Development and Coaching, Faculdade FACEL (Curitiba, PR, Brazil), conferred 31 May 2017. Master's thesis: \"The Influence of Psychopathy in the Corporate Environment.\"",
        "U.S. credential equivalency on file (GEO Credential Services).",
    ])

    h3(doc, "Professional certifications and registrations")
    add_bullets(doc, [
        "Active Registered Organizational Psychologist (Brazilian Federal Council, CRP-12/11050) — instrument-design competency required for psychometric work.",
        "Certified Predictive Index (PI) Analyst — assessment-architecture credential underpinning Module 01 instrument design.",
        "Member, American Management Association (AMA ID 4016170, since December 2023).",
        "Member in good standing, Society for Human Resource Management (SHRM), through July 2026.",
        "Certified Recruitment & Selection — Gupy and Kenoby platforms (operational recruiting suites).",
        "ADP — Effective Onboarding (talent integration certification).",
    ])

    h3(doc, "Operational track record (10+ years)")
    styled_table(doc,
                 ["Employer", "Sector", "Period", "Mandate"],
                 [
                     ["CONTAX Mobitel S.A.", "Contact Center / BPO", "Sep/2011 — Jul/2014",
                      "Behavioral selection at scale across high-volume contact center operations."],
                     ["Precisão Global de Cobranças", "Financial Services / Collections",
                      "Aug/2014 — Feb/2015",
                      "Recruitment and selection analytics for collections workforce."],
                     ["CDL Joinville", "Commercial Association", "Apr/2015 — Nov/2015",
                      "Recruitment & selection assistance across an inter-employer chamber."],
                     ["WALMART (WMS Brasil Ltda.)", "Multinational Retail",
                      "Feb/2016 — Sep/2017",
                      "Junior R&S Analyst — selection and onboarding for one of the world's largest employers."],
                     ["UNISOCIESC (Ânima Educational Group)", "Higher Education",
                      "Sep/2017 — Aug/2021",
                      "HR Analyst Jr. then HR Business Partner — climate, GPTW, training & development for an academic ecosystem."],
                     ["ASAAS Instituição de Pagamento", "Fintech (Brazilian Central Bank regulated)",
                      "Aug/2021 — May/2022",
                      "Talent Acquisition Analyst — strategic recruiting in a regulated payments environment."],
                 ],
                 col_widths_cm=[4.5, 4.0, 3.5, 5.5])

    add_callout(doc, "Why this matters",
                "Behavioral instruments are not interchangeable assets. Each ClimateLens module "
                "carries the methodological signature of its principal scientist. The lab's "
                "instrument library, calibration logic, and intervention playbooks are the "
                "direct outputs of Cristine Correa's continued research, certification work, "
                "and supervisory practice. Without that ongoing direction, the science freezes "
                "and the predictive value decays.")

    h3(doc, "Published thinking")
    add_para(doc,
             "Cristine is the author of \"The Differences Between HR in Brazil and the United "
             "States,\" a comparative analysis of behavioral assessment and labor compliance "
             "frameworks across the two largest economies in the Americas, published on Medium "
             "in both Portuguese and English. The article informs the lab's bilingual "
             "instrument design — a competitive differentiator in Florida's Spanish- and "
             "Portuguese-speaking workforce segments.",
             size=11, space_after=10)

    add_page_break(doc)


def build_pcrm(doc):
    h1(doc, "IV. The PCRM Framework™")
    h2(doc, "Psychometric Climate & Retention Methodology")

    add_para(doc,
             "Every ClimateLens engagement runs on PCRM, the proprietary methodology authored "
             "and continuously evolved by the lab's principal scientist. PCRM integrates four "
             "research streams that the mainstream survey market treats as separate problems: "
             "(a) industrial-organizational psychometrics; (b) Surgeon General workplace "
             "well-being doctrine; (c) Bureau of Labor Statistics turnover econometrics; and "
             "(d) supervised behavioral coaching for first-line managers.",
             size=11, space_after=8)

    h3(doc, "The four PCRM layers")
    styled_table(doc,
                 ["Layer", "Source Discipline", "Operational Output"],
                 [
                     ["Layer 1 — Instrument Calibration",
                      "Psychometric assessment design (Predictive Index, classical test theory, "
                      "item response theory).",
                      "Climate and retention surveys with documented validity and reliability "
                      "for SME industries."],
                     ["Layer 2 — Workplace Well-Being Mapping",
                      "U.S. Surgeon General's Framework (2022) and NIOSH Total Worker Health "
                      "guidance (2024-2026).",
                      "Five-domain mapping (protection, growth, connection, mattering, "
                      "harmony) translated to industry-specific items."],
                     ["Layer 3 — Predictive Analytics",
                      "BLS JOLTS turnover econometrics + behavioral data clustering.",
                      "Three-to-six-month flight-risk and burnout forecasts at unit and role "
                      "resolution."],
                     ["Layer 4 — Supervised Intervention",
                      "Applied organizational psychology + executive coaching practice.",
                      "Manager playbooks, conversation scripts, micro-metrics, and "
                      "four-week follow-up protocols."],
                 ],
                 col_widths_cm=[3.5, 6.5, 7.0])

    h3(doc, "Why PCRM cannot be commoditized")
    add_para(doc,
             "PCRM is not a static document. Each release cycle requires fresh validation: "
             "items are re-analyzed against accumulated response data, predictive thresholds "
             "are re-fit, and intervention scripts are re-coached against observed manager "
             "outcomes. This calibration work is performed by the lab's principal scientist. "
             "When the principal scientist stops calibrating, the framework stops evolving — "
             "and a frozen psychometric framework is a discredited one, because workforce "
             "behavior, federal guidance, and labor-market conditions all move continuously.",
             size=11, space_after=10)

    add_callout(doc, "Trade Secret Status",
                "PCRM Framework™ is treated as a trade secret of Correa Workforce Sciences. "
                "Item banks, calibration coefficients, predictive thresholds, and intervention "
                "scripts are not licensed, rebranded by third parties, or sublicensed. Every deployment is "
                "operated under the founder's continuing methodological direction.")

    add_page_break(doc)


def build_market(doc):
    h1(doc, "V. Market Opportunity & Competitive Landscape")
    h2(doc, "A Growing, Underserved Segment Where Science Is Scarce")

    add_kpi_card(doc, [
        ("$1.6–1.7B", "Global employee engagement survey market, 2024"),
        ("$3.5–4.2B", "Projected market by 2033 (CAGR 8–11%)"),
        ("~13%", "CAGR specific to the SME segment (2024–2033)"),
        ("33.2M", "SMEs in the United States (SBA, 2024)"),
    ])
    add_para(doc, " ", size=4, space_after=4)

    h3(doc, "Why the SME segment is the underserved frontier")
    add_para(doc,
             "The U.S. employee engagement software market accounts for roughly 38 percent of "
             "the global total — about US$ 585 million in 2024. Large enterprises still "
             "represent ~61 percent of revenue, but SMEs are the fastest-growing segment, "
             "expanding at roughly 13 percent compound annual growth through 2033. The "
             "competitive set has yet to follow that growth: leading platforms remain priced "
             "for mid-market and enterprise contracts, and the missing layer is the same "
             "across them — behavioral interpretation of the data.",
             size=11, space_after=8)

    h3(doc, "Competitive landscape")
    styled_table(doc,
                 ["Provider", "Primary Segment", "Annual Floor", "What's Missing for SMEs"],
                 [
                     ["Hogan / DDI / SHL",
                      "Enterprise behavioral assessment", "US$ 20K – 100K+",
                      "Inaccessible pricing; no SME calibration; no supervised intervention layer."],
                     ["Predictive Index",
                      "Mid-market behavioral", "US$ 8K – 30K+",
                      "Requires internal certified analyst; no climate/retention layer integrated."],
                     ["Culture Amp / Lattice / 15Five",
                      "Mid-market climate software", "US$ 3K – 10K minimum",
                      "Self-service interpretation; output is numbers, not behavioral change."],
                     ["Glint (Viva) / Qualtrics",
                      "Enterprise climate/engagement", "US$ 10K – 50K+",
                      "Built for HRIT teams that SMEs do not employ."],
                     ["Gallup Q12 advisory practice",
                      "Enterprise advisory", "US$ 50K+ engagements",
                      "Pricing and engagement model out of reach of 33.2M SMEs."],
                     ["ClimateLens (Correa Workforce Sciences)",
                      "SMEs (10–500 FTE) in healthcare and manufacturing",
                      "US$ 6K – 18K per engagement",
                      "Integrated diagnostic + analytics + supervised intervention; founder-led."],
                 ],
                 col_widths_cm=[3.8, 4.0, 3.0, 6.2])

    h3(doc, "Where ClimateLens fits")
    add_para(doc,
             "ClimateLens is positioned where behavioral science meets accessibility: "
             "rigorous psychometric instruments calibrated for SMEs, predictive analytics "
             "translated into manager-readable language, and supervised intervention authored "
             "by a credentialed organizational psychologist. The result is a research-grade "
             "engagement at small-employer pricing — a category that did not exist in the "
             "U.S. SME market before Correa Workforce Sciences began operating it.",
             size=11, space_after=10)

    add_page_break(doc)


def build_verticals(doc):
    h1(doc, "VI. Target Verticals")
    h2(doc, "Two Hyper-Specific U.S. Industries Where the Crisis Is Most Acute")

    add_para(doc,
             "Correa Workforce Sciences focuses on two American industries where small and "
             "mid-sized employers face structural turnover crises and where federal workforce "
             "policy creates immediate alignment opportunities for behavioral retention "
             "investments.",
             size=11, space_after=8)

    h3(doc, "Vertical 01 — Healthcare SMEs")
    add_para(doc,
             "Outpatient clinics, home health agencies, assisted living facilities, behavioral "
             "health practices, and physician group practices employing between 10 and 500 "
             "people. Annual turnover averages 22.7 percent (BLS, 2024); replacement cost is 50 "
             "to 200 percent of the departing role's salary; chronic understaffing degrades "
             "patient outcomes. Federal alignment: Health Resources and Services Administration "
             "(HRSA) workforce mandates, CMS quality metrics, NIOSH Total Worker Health "
             "guidance for healthcare. Geographic density: Central Florida hosts the country's "
             "largest concentration of post-acute and outpatient SME providers per capita.",
             size=11, space_after=8)

    h3(doc, "Vertical 02 — Manufacturing SMEs")
    add_para(doc,
             "Small and mid-sized manufacturers — fabrication, food processing, advanced "
             "manufacturing — with 10 to 500 employees. Annual turnover ranges from 24 to 32 "
             "percent (BLS, 2024). Workplace climate has direct, measured effects on safety "
             "incidents and productivity. Federal alignment: re-industrialization priorities "
             "(CHIPS and Science Act; Executive Order \"Preparing Americans for High-Paying "
             "Skilled Trade Jobs of the Future,\" 2025); OSHA and NIOSH psychosocial-hazard "
             "guidance. Geographic density: Central Florida is attracting accelerated capital "
             "into advanced manufacturing across Orange, Seminole, Volusia and Brevard counties.",
             size=11, space_after=8)

    add_callout(doc, "Why two and only two",
                "ClimateLens deliberately concentrates the lab's calibration capacity on two "
                "verticals. Each item, threshold, and intervention script is re-validated against "
                "industry-specific BLS turnover baselines, sector-specific federal guidance, "
                "and accumulated response data. Concentration is the source of scientific "
                "depth; dilution would erode it.")

    add_page_break(doc)


def build_federal_alignment(doc):
    h1(doc, "VII. Federal Workforce Policy Alignment")
    h2(doc, "Why ClimateLens Maps Directly to 2025–2026 U.S. Workforce Doctrine")

    add_para(doc,
             "ClimateLens does not operate adjacent to federal workforce policy — it operates "
             "as the behavioral implementation layer of the policies the United States is "
             "actively funding and prioritizing. The four pillars below are the active federal "
             "doctrine that ClimateLens directly serves.",
             size=11, space_after=8)

    styled_table(doc,
                 ["Pillar", "Issuing Body", "ClimateLens Implementation"],
                 [
                     ["U.S. Surgeon General's Framework for Mental Health & Well-Being at Work (2022)",
                      "Office of the U.S. Surgeon General",
                      "Module 01 instruments measure all five framework domains; Module 03 "
                      "intervention plans operationalize each domain at the manager level."],
                     ["NIOSH Total Worker Health & Psychosocial Hazards Guidance (2024–2026)",
                      "CDC / National Institute for Occupational Safety and Health",
                      "Module 02 analytics flag psychosocial hazards before incidents; Module 03 "
                      "translates NIOSH guidance into SME-scale interventions."],
                     ["Executive Order: Preparing Americans for High-Paying Skilled Trade Jobs of the Future (2025)",
                      "Executive Office of the President",
                      "Skills-based behavioral assessments via PCRM Layer 1 support the EO's "
                      "stated reliance on behavioral and aptitude assessment for talent matching."],
                     ["WIOA / America's Talent Strategy (FY2026 budget: US$ 2.97B)",
                      "U.S. Department of Labor — Employment & Training Administration",
                      "Climate diagnostics and retention analytics support sectoral partnerships "
                      "and Registered Apprenticeship retention metrics for SME participants."],
                     ["EDA Good Jobs Challenge (FY2026: ~US$ 468M)",
                      "U.S. Department of Commerce — Economic Development Administration",
                      "Behavioral retention engineering supports EDA workforce-development "
                      "partnerships in healthcare and manufacturing SMEs."],
                 ],
                 col_widths_cm=[5.0, 4.0, 8.0])

    add_callout(doc, "The federal alignment story",
                "Every dollar of federal workforce investment is currently chasing the same "
                "outcome: durable retention of skilled workers in healthcare and reindustrialized "
                "manufacturing. ClimateLens provides the behavioral instrumentation those "
                "investments require but do not themselves fund.")

    add_page_break(doc)


def build_legal(doc):
    h1(doc, "VIII. Regulatory & Compliance Posture")
    h2(doc, "Operating Cleanly Within Florida and Federal Frameworks")

    add_para(doc,
             "Correa Workforce Sciences operates within the corporate-development and "
             "industrial-organizational research perimeter, which is expressly outside the "
             "scope of clinical psychology licensing in Florida. The lab does not diagnose, "
             "treat, or maintain a therapist-patient relationship. Florida Statute § 490.014 "
             "expressly excludes corporate training, executive coaching, organizational "
             "assessment, and HR development from the Board of Psychology's jurisdiction.",
             size=11, space_after=8)

    h3(doc, "Compliance summary")
    add_bullets(doc, [
        "Florida entity formation in Seminole County under \"Correa Workforce Sciences, LLC.\"",
        "Operating model expressly outside Fl. Stat. § 490.014 jurisdictional scope.",
        "Data handling under Florida's information-protection regime; no HIPAA-covered data ingested.",
        "EEOC alignment for behavioral assessment use under the Uniform Guidelines on Employee Selection Procedures (UGESP).",
        "ADA-aware item design — instruments avoid medical inquiry; intervention plans are job-related and consistent with business necessity.",
    ])

    add_page_break(doc)


def build_pricing(doc):
    h1(doc, "IX. Pricing & Engagement Model")
    h2(doc, "Three Tiers Designed for the SME Reality")

    add_para(doc,
             "Engagements are sold as scientific programs — not as software seats. Every tier "
             "includes founder-supervised methodological direction; the difference between "
             "tiers is depth, frequency, and the number of business units covered.",
             size=11, space_after=8)

    styled_table(doc,
                 ["Tier", "Investment", "Scope", "Founder Touchpoints"],
                 [
                     ["Diagnostic Engagement",
                      "US$ 6,500 — US$ 9,000",
                      "Single climate diagnostic across one business unit (10–80 FTE). "
                      "PCRM-calibrated instrument, full report, manager debrief.",
                      "1 calibration session; 1 supervised debrief; 1 30-day follow-up call."],
                     ["Continuous Lab Subscription",
                      "US$ 1,200 — US$ 1,800 / month",
                      "Pulse + annual surveys, predictive analytics dashboard, monthly "
                      "manager office hours, quarterly strategic review.",
                      "Monthly review; quarterly methodological update; intervention sign-off."],
                     ["Strategic Retention Program",
                      "US$ 14,000 — US$ 18,000 / 12-month engagement",
                      "Full PCRM deployment: diagnostic + analytics + supervised "
                      "intervention plans + manager coaching across the full organization "
                      "(up to 500 FTE).",
                      "Bi-weekly review; quarterly framework recalibration; "
                      "year-end scientific report."],
                 ],
                 col_widths_cm=[4.0, 3.5, 6.5, 3.0])

    add_callout(doc, "Why tiered, not modular",
                "The lab refuses to sell a stand-alone software seat. Every engagement carries "
                "scientific supervision because the science is the product. SMEs underserved "
                "by the current market do not need another dashboard — they need an "
                "interpreted, supervised, evidence-based plan.")

    add_page_break(doc)


def build_economics(doc):
    h1(doc, "X. Financial Projections (Asset-Light Model)")
    h2(doc, "Conservative, Capital-Disciplined Build")

    add_para(doc,
             "Correa Workforce Sciences is structured as an asset-light applied research lab. "
             "The model intentionally avoids physical facilities, clinical equipment, HIPAA "
             "exposure, or aggressive headcount in the first 24 months. The product is "
             "intellectual property — instruments, calibration logic, intervention scripts — "
             "delivered through a low-cost learning management substrate.",
             size=11, space_after=8)

    styled_table(doc,
                 ["Metric", "Year 1", "Year 2", "Year 3"],
                 [
                     ["Initial capital deployed", "< US$ 15,000", "—", "—"],
                     ["Active engagements", "6 — 10", "18 — 28", "40 — 60"],
                     ["Revenue range", "US$ 50K — 80K", "US$ 180K — 300K", "US$ 420K — 680K"],
                     ["Gross margin (post month 6)", "> 80%", "> 82%", "> 84%"],
                     ["Founder-led utilization", "100%", "85% (with apprenticed analysts under supervision)",
                      "70% (with two supervised analysts)"],
                     ["First contract", "Months 2–4", "—", "—"],
                 ],
                 col_widths_cm=[5.0, 4.0, 4.0, 4.0])

    add_para(doc,
             "All revenue projections assume zero external capital and growth funded entirely "
             "by engagement cash flow. The model has been built to remain solvent under a "
             "stress scenario in which year-one engagements come in at the low end of the "
             "range.",
             size=11, space_after=10)

    add_page_break(doc)


def build_geo_strategy(doc):
    h1(doc, "XI. Geographic Strategy")
    h2(doc, "Sanford and Seminole County as the Operational Anchor; Central Florida as the Proving Ground")

    add_para(doc,
             "Correa Workforce Sciences is headquartered in Sanford, Seminole County, Florida. "
             "Central Florida — Orange, Seminole, Volusia and Brevard counties — combines four "
             "structural conditions that make it the natural proving ground for ClimateLens.",
             size=11, space_after=8)

    add_bullets(doc, [
        "One of the country's densest concentrations of healthcare SMEs, including post-acute, home health, assisted living, and outpatient practices.",
        "Active state and federal incentives accelerating advanced manufacturing investment across the I-4 corridor.",
        "A trilingual workforce (English, Spanish, Portuguese) requiring instrument calibration that monolingual platforms cannot deliver.",
        "Proximity to two of the largest teaching hospital systems in the Southeast, supporting future clinical-research partnerships.",
    ])

    add_para(doc,
             "Year 2 expansion targets the broader Florida market followed by the Southeast "
             "(Georgia, North Carolina, South Carolina), maintaining the lab's discipline of "
             "two verticals and founder-supervised engagements at every site.",
             size=11, space_after=10)

    add_page_break(doc)


def build_roadmap(doc):
    h1(doc, "XII. 12-Month Operating Roadmap")

    styled_table(doc,
                 ["Quarter", "Milestones"],
                 [
                     ["Q1 (Months 1–3)",
                      "Florida entity formation finalized; PCRM v1.0 instrument library "
                      "released; brand and digital infrastructure live; first two pilot "
                      "engagements signed (one healthcare, one manufacturing) at the diagnostic tier."],
                     ["Q2 (Months 4–6)",
                      "Pilot diagnostics completed; PCRM v1.1 calibration update published; "
                      "first continuous lab subscription signed; first scientific debrief "
                      "report delivered; first SHRM/AMA published case write-up authored."],
                     ["Q3 (Months 7–9)",
                      "Three to five active continuous-lab engagements; predictive analytics "
                      "dashboard moved into production; PCRM v1.2 published with first "
                      "industry-specific benchmark dataset."],
                     ["Q4 (Months 10–12)",
                      "First Strategic Retention Program signed; annual scientific report "
                      "published; first apprenticed analyst onboarded under direct founder "
                      "supervision; revenue trajectory enters Year 2 projections."],
                 ],
                 col_widths_cm=[3.5, 13.5])

    add_page_break(doc)


def build_partners(doc):
    h1(doc, "XIII. Strategic Partners & Letters of Intent")

    add_para(doc,
             "ClimateLens has secured strategic partnership and intent-to-engage commitments "
             "from operators across the lab's two target verticals and from senior advisors "
             "with deep U.S. SME and category-building experience. The following commitments "
             "are documented and on file with Correa Workforce Sciences.",
             size=11, space_after=8)

    styled_table(doc,
                 ["Partner / Advisor", "Role", "Commitment Type"],
                 [
                     ["Suneera Madhani",
                      "Founder & former CEO, Stax Payments (built and scaled US$ 1B+ "
                      "category-defining payments business serving U.S. SMEs).",
                      "Strategic advisory — go-to-market, SME category construction, "
                      "founder coaching."],
                     ["Fernando Braga Neto",
                      "U.S.-based investor and operator with portfolio exposure to "
                      "professional-services SMEs in Florida.",
                      "Capital and operational commitment letter for ClimateLens "
                      "first-year operations."],
                     ["Danieli Nieri",
                      "U.S. healthcare and post-acute services operator with direct "
                      "Florida market presence.",
                      "Vertical partnership for healthcare SME pilot deployments and "
                      "introductions across Central Florida."],
                     ["Lucas Vieira",
                      "Brazilian commercial partner — letter of intent for cross-border "
                      "knowledge-transfer engagements aligned with the lab's bilingual "
                      "instrument library.",
                      "Letter of commercial intent."],
                     ["José Eduardo Borges",
                      "Brazilian commercial partner — letter of intent confirming demand "
                      "for ClimateLens methodology in cross-border SME engagements.",
                      "Letter of commercial intent."],
                 ],
                 col_widths_cm=[4.5, 7.5, 5.0])

    add_callout(doc, "Why these commitments matter",
                "Each partner has reviewed PCRM Framework™ documentation and the lab's "
                "instrument design under non-disclosure. Their commitments reflect direct "
                "evaluation of the methodology and the founder's scientific authorship — "
                "not generic enthusiasm.")

    add_page_break(doc)


def build_indispensability(doc):
    h1(doc, "XIV. Why the Founder Cannot Be Decoupled From the Lab")
    h2(doc, "Methodological Authorship Is the Product")

    add_para(doc,
             "Correa Workforce Sciences was founded on a single scientific premise: behavioral "
             "instruments and intervention plans are only as valid as the principal scientist "
             "who continuously calibrates them. Three operational facts make the founder's "
             "ongoing direction non-negotiable.",
             size=11, space_after=8)

    h3(doc, "01 — The instrument library is authored, not licensed")
    add_para(doc,
             "Every PCRM instrument is original work authored by the principal scientist. The "
             "lab does not rebrand, sublicense, or resell third-party instruments. "
             "Calibration coefficients, item banks, and predictive thresholds carry the "
             "founder's psychometric signature and require her direct supervision to update.",
             size=11, space_after=6)

    h3(doc, "02 — Calibration is continuous, not one-time")
    add_para(doc,
             "Each release of PCRM (v1.0, v1.1, v1.2 …) is the product of a structured "
             "calibration process: response data accumulated since the prior release is "
             "re-analyzed; predictive thresholds are re-fit; intervention scripts are coached "
             "against observed outcomes. This work is not delegable: the calibration decisions "
             "depend on the principal scientist's clinical-and-organizational training.",
             size=11, space_after=6)

    h3(doc, "03 — Intervention plans are signed, not generated")
    add_para(doc,
             "Module 03 intervention plans carry the principal scientist's sign-off before "
             "deployment. The lab does not auto-generate or auto-recommend intervention scripts. "
             "Each plan is reviewed against the responding organization's specific behavioral "
             "context, then revised at four-week intervals based on observed manager outcomes.",
             size=11, space_after=10)

    add_callout(doc, "What happens without the founder",
                "If the principal scientist were to step away from the lab, three things "
                "would happen in sequence: the instrument library would freeze and lose "
                "predictive validity within two release cycles; the predictive analytics "
                "would drift as labor-market conditions shift; and the intervention library "
                "would degrade as new federal guidance arrives without a scientist to "
                "translate it. The lab is the founder's research, applied at scale.")

    add_page_break(doc)


def build_appendix(doc):
    h1(doc, "Appendix — Scientific & Federal Source Index")
    add_bullets(doc, [
        "U.S. Surgeon General. Framework for Mental Health & Well-Being in the Workplace. Office of the U.S. Surgeon General, 2022.",
        "Centers for Disease Control & Prevention / NIOSH. Total Worker Health Program — Psychosocial Hazards Guidance, 2024–2026.",
        "U.S. Bureau of Labor Statistics. Job Openings and Labor Turnover Survey (JOLTS), 2024–2025.",
        "U.S. Bureau of Labor Statistics. Occupational Employment and Wage Statistics — Human Resources Managers (11-3121); Training and Development Managers (11-3131), 2024.",
        "U.S. Small Business Administration. Frequently Asked Questions about Small Business, 2024.",
        "Gallup. State of the Global Workplace, 2024 — Q12 Meta-Analysis (183,806 business units; 3.35M employees).",
        "U.S. Department of Labor / Employment & Training Administration. WIOA / America's Talent Strategy, FY2026 budget materials.",
        "U.S. Department of Commerce / Economic Development Administration. Good Jobs Challenge, FY2026.",
        "Executive Office of the President. Executive Order: Preparing Americans for High-Paying Skilled Trade Jobs of the Future, 2025.",
        "Florida Statutes § 490.014 — Exemption of Corporate Development and Industrial-Organizational Practice.",
        "Equal Employment Opportunity Commission. Uniform Guidelines on Employee Selection Procedures (UGESP).",
    ])
    add_para(doc, "© Correa Workforce Sciences, LLC — All scientific content authored by Cristine Correa.",
             size=9, italic=True, color=SLATE, space_before=14)


def build_document():
    doc = Document()
    # Default style
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    set_page_margins(doc.sections[0])
    add_footer(doc, "Correa Workforce Sciences, LLC  ·  ClimateLens — Product & Scientific Dossier  ·  Sanford, FL")

    build_cover(doc)
    build_executive_summary(doc)
    build_problem(doc)
    build_solution(doc)
    build_founder(doc)
    build_pcrm(doc)
    build_market(doc)
    build_verticals(doc)
    build_federal_alignment(doc)
    build_legal(doc)
    build_pricing(doc)
    build_economics(doc)
    build_geo_strategy(doc)
    build_roadmap(doc)
    build_partners(doc)
    build_indispensability(doc)
    build_appendix(doc)

    doc.save(str(DOCX_PATH))
    return DOCX_PATH


def build_thumbnail_map():
    base = "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2026/CRISTINE CORREA/CARTAS/"
    cv_base = base + "CV_APOIADORES/"
    apoio = base + "APOIO_EUA/"
    intent_br = base + "INTENÇÃO_COMERCIAL_BR/"
    return {
        "generated_at": "2026-04-18",
        "engagement": "ClimateLens — Product & Scientific Dossier (V1)",
        "exhibits": [
            {
                "exhibit_number": 1,
                "description": "Strategic Advisor Letter — Suneera Madhani (Founder, Stax Payments)",
                "pdf_path": apoio + "Suneera Madhani - Carta Apoio.pdf",
                "section": "XIII. Strategic Partners",
            },
            {
                "exhibit_number": 2,
                "description": "Investor Commitment Letter (v2) — Fernando Braga Neto",
                "pdf_path": apoio + "Fernando Braga Neto - Carta Investidor (v2).pdf",
                "section": "XIII. Strategic Partners",
            },
            {
                "exhibit_number": 3,
                "description": "Investment Commitment — Fernando Braga Neto",
                "pdf_path": apoio + "Fernando Braga Neto - Compromisso Investimento.docx",
                "section": "XIII. Strategic Partners",
            },
            {
                "exhibit_number": 4,
                "description": "Healthcare Vertical Partnership Letter — Danieli Nieri",
                "pdf_path": apoio + "Danieli Nieri - Carta Parceria.pdf",
                "section": "XIII. Strategic Partners",
            },
            {
                "exhibit_number": 5,
                "description": "Letter of Commercial Intent — Lucas Vieira",
                "pdf_path": intent_br + "Lucas Vieira - Letter of Intent.docx",
                "section": "XIII. Strategic Partners",
            },
            {
                "exhibit_number": 6,
                "description": "Letter of Commercial Intent — José Eduardo Borges",
                "pdf_path": intent_br + "José Eduardo Borges - Letter of Intent.docx",
                "section": "XIII. Strategic Partners",
            },
            {
                "exhibit_number": 7,
                "description": "CV — José Eduardo Borges (commercial intent counterparty)",
                "pdf_path": cv_base + "José Eduardo Borges - CV.pdf",
                "section": "XIII. Strategic Partners",
            },
            {
                "exhibit_number": 8,
                "description": "LinkedIn Profile — Danieli Nieri (healthcare partnership counterparty)",
                "pdf_path": cv_base + "Danieli Nieri - LinkedIn.pdf",
                "section": "XIII. Strategic Partners",
            },
            {
                "exhibit_number": 9,
                "description": "LinkedIn Profile — Lucas Vieira (commercial intent counterparty)",
                "pdf_path": cv_base + "Lucas Vieira - LinkedIn.pdf",
                "section": "XIII. Strategic Partners",
            },
        ],
    }


def build_lovable_spec():
    return r"""# LOVABLE BUILD SPEC — ClimateLens (Correa Workforce Sciences)

> Build prompt for Lovable (or comparable AI app builder). Builds a complete enterprise-grade SaaS web app that mirrors the Product & Scientific Dossier delivered alongside this file.
> Target users: U.S. small and mid-sized employers (10–500 FTE) in healthcare and manufacturing verticals plus their HR leaders, plant managers, clinical directors. Plus internal lab staff (admin / scientific director).

---

## 1. Brand & Visual System

- **Product name:** ClimateLens
- **Operating company:** Correa Workforce Sciences, LLC
- **Tagline:** *Predictive Climate & Retention Intelligence for American Small and Mid-Sized Employers.*
- **Subline:** *Powered by the PCRM Framework™ — Psychometric Climate & Retention Methodology.*
- **Visual reference:** Stripe (typography), Notion (layout density), Linear (interaction polish), Atlassian Statuspage (data presentation).

### Color tokens

| Token             | Hex       | Usage                                          |
| ----------------- | --------- | ---------------------------------------------- |
| `--color-navy`    | `#1B2A4A` | Primary brand, headers, sidebar background     |
| `--color-gold`    | `#B8860B` | Accent, links, KPI deltas, brand divider       |
| `--color-ink`     | `#10141F` | Primary text                                   |
| `--color-slate`   | `#4B5563` | Secondary text, captions                       |
| `--color-paper`   | `#FAFBFC` | Page background                                |
| `--color-soft`    | `#F5F1E6` | Tier card background, callout fill             |
| `--color-rule`    | `#E3E7EC` | Borders, table dividers                        |
| `--color-success` | `#0E9F6E` | Positive deltas, on-track engagements          |
| `--color-warn`    | `#D97706` | Watch indicators                               |
| `--color-risk`    | `#B91C1C` | Flight-risk and burnout flags                  |

### Typography

- **Display / headings:** Georgia (serif) — `font-family: Georgia, "Times New Roman", serif;`
- **Body / UI:** Inter — `font-family: Inter, -apple-system, system-ui, sans-serif;`
- **Mono / data:** JetBrains Mono.
- Heading scale: H1 32/40, H2 24/32, H3 18/26, body 14/22, micro 12/18.

### Iconography & motion

- Lucide icons throughout. Brand divider = thin gold rule (1.5 px) under H1s.
- Subtle 120 ms ease-out transitions on hover; never bouncy.

---

## 2. Tech stack

- **Front-end:** React 18 + TypeScript + Vite + Tailwind CSS + shadcn/ui.
- **State:** TanStack Query, Zustand for local UI state.
- **Charts:** Recharts (line, area, stacked bar, gauge).
- **Forms:** React Hook Form + Zod validation.
- **Auth & data:** Supabase (Postgres + Row-Level Security + Storage).
- **Background jobs:** Supabase Edge Functions for survey ingestion, scoring, predictive scoring.
- **Email:** Resend for survey invitations.
- **PDF reports:** React-PDF.
- **Analytics:** PostHog.

---

## 3. Global components

| Component               | Notes                                                                              |
| ----------------------- | ---------------------------------------------------------------------------------- |
| `BrandHeader`           | Navy bar with gold underline. Holds logo "ClimateLens" + workspace switcher.       |
| `SidebarNav`            | Navy, 240 px, fixed. Sections: Workspace, Diagnostics, Predictions, Interventions, Library, Settings. |
| `KpiCard`               | Navy → Soft two-row card. Displays headline number, label, optional delta.         |
| `CalloutBlock`          | Gold left rule + soft fill; used for scientific or compliance notes.               |
| `DataTable`             | Sortable, filterable, paginated, sticky header. Right-aligned numerics.            |
| `StatusBadge`           | `success` / `warn` / `risk` variants for retention indicators.                     |
| `TierCard`              | Pricing tier card with gold highlight ring on hover.                               |
| `RiskGauge`             | Half-donut gauge showing flight-risk probability (0–100%).                         |
| `MethodologyBanner`     | "PCRM v1.x — last calibrated by Cristine Correa on …" — visible on every dashboard. |
| `EvidenceLink`          | Inline link to a federal source (Surgeon General, NIOSH, BLS, etc.) with icon.     |
| `EngagementProgress`    | Horizontal stepper for engagement lifecycle.                                       |

---

## 4. Information architecture (7+ pages)

### 4.1 `/` — Public marketing site (single-page)

- Navy hero with serif headline: *Predictive Climate & Retention Intelligence for American SMEs.*
- KPI strip: `$1T`, `46.4%`, `0.49`, `22.7%` (same KPIs as the dossier).
- Three feature cards = the three modules (Diagnostic, Predictive Analytics, Supervised Intervention).
- "Why founder-led" section with portrait placeholder + the three indispensability points.
- Federal-alignment grid (Surgeon General, NIOSH, EO 2025, WIOA, EDA) with brief copy.
- Pricing snapshot (three tiers — must mirror dossier exactly: $6.5K–9K, $1.2K–1.8K/mo, $14K–18K).
- Logos strip (placeholder badges for "Strategic Advisor — Suneera Madhani," "Healthcare Partner — Danieli Nieri," "Investor — Fernando Braga Neto").
- Footer with Sanford, FL address + Florida Statute § 490.014 compliance line.

### 4.2 `/dashboard` — Workspace dashboard (post-login)

- Header: workspace name, vertical (Healthcare or Manufacturing), FTE count.
- Top row of `KpiCard`s: Active engagements, Climate Index score (0–100), Predicted 90-day attrition %, Open intervention plans.
- "Climate Index trend" line chart (last 12 months).
- "Where to look first" panel: top 3 units with highest predicted attrition, each linking to `/predictions/:unitId`.
- Right-rail: `MethodologyBanner` showing current PCRM version + next calibration window.

### 4.3 `/diagnostics` — Climate Diagnostic module

- Tabs: *Active diagnostics*, *Historical*, *Instrument library*.
- Active diagnostics list: client, instrument version (PCRM v1.2 SME-Healthcare-A), invitations sent, response rate, status badge.
- Detail page per diagnostic: response distribution per Surgeon General domain (5 horizontal stacked bars), open-text NLP themes, draft report button.
- Instrument library: list of PCRM instruments with version, vertical, item count, last calibrated by, authored by (always Cristine Correa).
- Action: "Launch new diagnostic" wizard (3 steps: vertical → unit → instrument selection → schedule).

### 4.4 `/predictions` — Predictive Retention Analytics

- Heatmap: units × roles colored by 90-day attrition probability (use `--color-success` → `--color-warn` → `--color-risk`).
- Drilldown per cell: `RiskGauge` + driver list (e.g., "Low Connection score (–1.4σ)", "Manager span > 12 reports").
- Burnout watchlist table.
- Footer of every prediction: "Predictions reviewed by founder on …" timestamp.

### 4.5 `/interventions` — Supervised Behavioral Intervention plans

- Kanban: *Drafting*, *Founder review*, *Active*, *Follow-up*, *Archived*.
- Each card shows unit, plan name, manager owner, micro-metric trend, days to next checkpoint.
- Plan detail: conversation script, micro-metrics list, 4-week follow-up timeline, founder sign-off block (avatar + signature).

### 4.6 `/library` — PCRM Methodology Library

- Read-only knowledge base for clients.
- Sections: PCRM Framework Overview, Surgeon General mapping, NIOSH Total Worker Health translation, BLS turnover benchmarks, Federal alignment briefs.
- Each article tagged with "Authored by Cristine Correa, M.B.A., Org. Psychologist (CRP-12/11050)."

### 4.7 `/pricing` — Engagement & pricing

Mirror the dossier exactly:

| Tier                           | Investment              | Best for                                |
| ------------------------------ | ----------------------- | --------------------------------------- |
| Diagnostic Engagement          | US$ 6,500 – 9,000       | One business unit, single diagnostic.   |
| Continuous Lab Subscription    | US$ 1,200 – 1,800/month | Pulse + annual + analytics + monthly office hours. |
| Strategic Retention Program    | US$ 14,000 – 18,000 / 12-month engagement | Full PCRM deployment up to 500 FTE.     |

Each tier card: features, founder touchpoints, "Talk to the lab" CTA.

### 4.8 `/settings` — Workspace settings

- Workspace profile (vertical, FTE, locations).
- Members & roles (Admin, Scientific Director, HR Lead, Manager, Viewer).
- Data retention controls (Florida § 490.014 / EEOC UGESP compliance copy).
- Billing.
- Export PDF report (uses React-PDF; mirrors the dossier visual system).

### 4.9 `/admin/lab` — Internal lab console (founder + analysts only)

- PCRM version manager: list of releases, calibration logs, pending re-calibration tasks.
- Engagement pipeline view across all clients.
- "Founder approval queue" for intervention plans.
- Apprenticed-analyst directory (with "supervised by Cristine Correa" badge).

---

## 5. Data model (Supabase)

```sql
-- workspaces (clients)
workspaces(id uuid pk, name text, vertical text check (vertical in ('healthcare','manufacturing')),
           fte int, state text, county text, created_at timestamptz)

-- pcrm versions
pcrm_versions(id uuid pk, version text, vertical text, item_count int,
              calibrated_by text default 'Cristine Correa', calibrated_at timestamptz,
              notes text)

-- diagnostics
diagnostics(id uuid pk, workspace_id uuid fk, pcrm_version_id uuid fk,
            unit_name text, status text, sent_at timestamptz, closed_at timestamptz,
            response_rate numeric)

responses(id uuid pk, diagnostic_id uuid fk, respondent_hash text,
          domain_scores jsonb, free_text text, submitted_at timestamptz)

-- predictions
predictions(id uuid pk, workspace_id uuid fk, unit_name text, role text,
            attrition_prob numeric, burnout_prob numeric, drivers jsonb,
            reviewed_by text default 'Cristine Correa', reviewed_at timestamptz)

-- interventions
interventions(id uuid pk, workspace_id uuid fk, plan_name text, manager_owner text,
              status text check (status in ('drafting','founder_review','active','followup','archived')),
              founder_signoff_at timestamptz, created_at timestamptz)

intervention_steps(id uuid pk, intervention_id uuid fk, week int, script text,
                   micro_metric text, completed boolean)

-- pricing & engagements
engagements(id uuid pk, workspace_id uuid fk, tier text, monthly_value numeric,
            engagement_value numeric, started_at date, ends_at date)

-- billing handled via Stripe webhooks → engagements.
```

Row-Level Security: every table scoped by `workspace_id` membership.

---

## 6. Mock data (must use the lab's real terminology)

- Workspaces: *Lakeview Home Health (Sanford, FL — Healthcare — 84 FTE)*, *Seminole Precision Components (Sanford, FL — Manufacturing — 142 FTE)*.
- PCRM versions: *PCRM v1.0 SME-Healthcare-A*, *PCRM v1.1 SME-Manufacturing-A*, *PCRM v1.2 SME-Healthcare-B*, all *calibrated_by = "Cristine Correa"*.
- Surgeon General domains in every diagnostic: Protection from Harm, Opportunity for Growth, Connection & Community, Mattering at Work, Work–Life Harmony.
- Drivers for predictions must reference real BLS / Gallup / NIOSH terminology — never generic "low engagement."

---

## 7. Compliance & copy guardrails (must follow inside the app copy)

- **Never** use the words "consulting", "consultoria", "advisory services" (generic), "standardized", "operates autonomously", "self-sustaining", "plug-and-play", "train-the-trainer", "white-label", "turnkey", "founder dependency: low", "scalable without the founder".
- **Always** refer to the methodology as PCRM Framework™, calibrated by the principal scientist.
- The internal lab console must surface the founder's name in every approval, calibration, and intervention sign-off — the product *is* her science applied at scale.
- Federal references must cite the issuing body (Surgeon General, NIOSH, BLS, DOL/ETA, EDA, EOP) with a year.
- Florida Statute § 490.014 compliance copy is required in the footer and on `/settings`.

---

## 8. Definition of done

The Lovable build is complete when:

1. The marketing site renders the same KPI strip, three modules, federal-alignment grid, and pricing tiers as this dossier.
2. A demo workspace is seeded with two clients, three PCRM versions, two diagnostics with synthetic responses, predictions with risk gauges, and four intervention plans across the kanban states.
3. PCRM versioning, diagnostic launch, prediction review, and intervention sign-off all work end-to-end with the founder's name surfaced everywhere a calibration or sign-off occurs.
4. PDF export of any report mirrors the dossier visual system (Navy + Gold, Georgia + Inter).
5. None of the prohibited terms appear anywhere in the UI copy or seed data.
"""


def main():
    docx_path = build_document()
    LOVABLE_PATH.write_text(build_lovable_spec(), encoding='utf-8')
    THUMB_PATH.write_text(json.dumps(build_thumbnail_map(), indent=2, ensure_ascii=False),
                          encoding='utf-8')
    print(f"DOCX:    {docx_path}  ({docx_path.stat().st_size:,} bytes)")
    print(f"LOVABLE: {LOVABLE_PATH}  ({LOVABLE_PATH.stat().st_size:,} bytes)")
    print(f"THUMB:   {THUMB_PATH}  ({THUMB_PATH.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
