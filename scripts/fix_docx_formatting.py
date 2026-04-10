#!/usr/bin/env python3
"""
fix_docx_formatting.py — Post-processing obrigatório para TODOS os DOCX gerados.

Corrige os 4 problemas recorrentes:
1. Spacing inflado (2000+ páginas → ~80 páginas)
2. Imagens inline estouradas → anchor + wrapSquare
3. Parágrafos vazios excessivos
4. Imagens gigantes → max 1.5 inches

Uso: python3 fix_docx_formatting.py input.docx [output.docx]
Se output não fornecido, sobrescreve o input.
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from lxml import etree
import copy

WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'


def fix_spacing(doc):
    """Comprimir spacing de todos os parágrafos e tabelas."""
    fixed = 0
    for p in doc.paragraphs:
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        pf.space_before = Pt(2)
        pf.line_spacing = Pt(13)
        fixed += 1
        # Remove page breaks excessivos (manter só em headers bold >13pt)
        for run in p.runs:
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    is_header = run.bold and run.font.size and run.font.size >= Pt(13)
                    if not is_header:
                        br.getparent().remove(br)
                        fixed += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.space_after = Pt(2)
                    pf.space_before = Pt(1)
                    pf.line_spacing = Pt(12)
    return fixed


def remove_empty_paragraphs(doc):
    """Remover parágrafos vazios consecutivos (manter max 1)."""
    body = doc.element.body
    prev_empty = False
    to_remove = []
    for p_elem in body.findall(qn('w:p')):
        text = ''
        for r in p_elem.findall(qn('w:r')):
            for t in r.findall(qn('w:t')):
                text += (t.text or '')
        is_empty = len(text.strip()) == 0
        has_drawing = len(p_elem.findall('.//' + qn('w:drawing'))) > 0
        if is_empty and not has_drawing:
            if prev_empty:
                to_remove.append(p_elem)
            prev_empty = True
        else:
            prev_empty = False
    for elem in to_remove:
        elem.getparent().remove(elem)
    return len(to_remove)


def resize_images(doc, max_width_inches=1.5):
    """Redimensionar imagens maiores que max_width."""
    max_cx = int(Inches(max_width_inches))
    img_fixed = 0
    body = doc.element.body
    for extent in body.iter('{%s}extent' % WP_NS):
        cx = int(extent.get('cx', 0))
        cy = int(extent.get('cy', 0))
        if cx > max_cx and cy > 0:
            ratio = max_cx / cx
            extent.set('cx', str(max_cx))
            extent.set('cy', str(int(cy * ratio)))
            img_fixed += 1
    return img_fixed


def inline_to_anchor(doc):
    """Converter imagens inline em table cells para anchor + wrapSquare."""
    body = doc.element.body
    fixed = 0

    for table in body.iter('{%s}tbl' % W_NS):
        for tc in table.iter('{%s}tc' % W_NS):
            for p in tc.findall('{%s}p' % W_NS):
                for r in p.findall('{%s}r' % W_NS):
                    for drawing in r.findall('{%s}drawing' % W_NS):
                        inlines = drawing.findall('{%s}inline' % WP_NS)
                        for inline in inlines:
                            extent = inline.find('{%s}extent' % WP_NS)
                            cx = int(extent.get('cx', 0)) if extent is not None else 0
                            cy = int(extent.get('cy', 0)) if extent is not None else 0

                            max_w = int(Inches(1.5))
                            if cx > max_w and cy > 0:
                                ratio = max_w / cx
                                cx = max_w
                                cy = int(cy * ratio)

                            graphic = inline.find('{%s}graphic' % A_NS)
                            docPr = inline.find('{%s}docPr' % WP_NS)

                            if graphic is None:
                                continue

                            anchor = etree.SubElement(drawing, '{%s}anchor' % WP_NS)
                            anchor.set('distT', '0')
                            anchor.set('distB', '0')
                            anchor.set('distL', '114300')
                            anchor.set('distR', '114300')
                            anchor.set('simplePos', '0')
                            anchor.set('relativeHeight', '0')
                            anchor.set('behindDoc', '0')
                            anchor.set('locked', '0')
                            anchor.set('layoutInCell', '1')
                            anchor.set('allowOverlap', '1')

                            simplePos = etree.SubElement(anchor, '{%s}simplePos' % WP_NS)
                            simplePos.set('x', '0')
                            simplePos.set('y', '0')

                            posH = etree.SubElement(anchor, '{%s}positionH' % WP_NS)
                            posH.set('relativeFrom', 'column')
                            posHOffset = etree.SubElement(posH, '{%s}posOffset' % WP_NS)
                            posHOffset.text = '0'

                            posV = etree.SubElement(anchor, '{%s}positionV' % WP_NS)
                            posV.set('relativeFrom', 'paragraph')
                            posVOffset = etree.SubElement(posV, '{%s}posOffset' % WP_NS)
                            posVOffset.text = '0'

                            new_extent = etree.SubElement(anchor, '{%s}extent' % WP_NS)
                            new_extent.set('cx', str(cx))
                            new_extent.set('cy', str(cy))

                            new_ee = etree.SubElement(anchor, '{%s}effectExtent' % WP_NS)
                            new_ee.set('l', '0')
                            new_ee.set('t', '0')
                            new_ee.set('r', '0')
                            new_ee.set('b', '0')

                            wrapSquare = etree.SubElement(anchor, '{%s}wrapSquare' % WP_NS)
                            wrapSquare.set('wrapText', 'bothSides')

                            if docPr is not None:
                                anchor.append(copy.deepcopy(docPr))

                            anchor.append(copy.deepcopy(graphic))

                            drawing.remove(inline)
                            fixed += 1
    return fixed


def fix_margins(doc):
    """Ajustar margens para o padrão PROEX."""
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.55)


def main():
    if len(sys.argv) < 2:
        print("Uso: python3 fix_docx_formatting.py input.docx [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else input_path

    if not os.path.exists(input_path):
        print(f"Arquivo nao encontrado: {input_path}")
        sys.exit(1)

    print(f"{'='*60}")
    print(f"FIX DOCX FORMATTING")
    print(f"{'='*60}")
    print(f"Input:  {input_path}")
    print(f"Output: {output_path}")

    doc = Document(input_path)

    # 1. Spacing
    spacing_fixes = fix_spacing(doc)
    print(f"[1/5] Spacing comprimido: {spacing_fixes} fixes")

    # 2. Empty paragraphs
    empty_removed = remove_empty_paragraphs(doc)
    print(f"[2/5] Paragrafos vazios removidos: {empty_removed}")

    # 3. Resize images
    img_resized = resize_images(doc)
    print(f"[3/5] Imagens redimensionadas: {img_resized}")

    # 4. Inline → Anchor
    anchor_fixes = inline_to_anchor(doc)
    print(f"[4/5] Inline → Anchor+wrapSquare: {anchor_fixes}")

    # 5. Margins
    fix_margins(doc)
    print(f"[5/5] Margens ajustadas")

    doc.save(output_path)

    # Stats
    doc2 = Document(output_path)
    chars = sum(len(p.text) for p in doc2.paragraphs)
    print(f"\nResultado: {len(doc2.paragraphs)} paras | {len(doc2.tables)} tabelas | ~{chars//3000} paginas")
    print(f"Salvo: {output_path}")
    print(f"{'='*60}")


if __name__ == '__main__':
    main()
