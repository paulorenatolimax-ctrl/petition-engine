#!/usr/bin/env python3
"""QA determinístico do DOCX gerado pelo pipeline IMPACTO®.

Valida o DOCX final pós-builder via python-docx. Exit 0 OK, 1 FAIL.
Gera relatório markdown PASS/FAIL com métricas e diff.

Uso:
    python3 qa_impacto_docx.py <generated.docx> [qa_report.md]

Baseline: Luciano VF (325 paras, 10 tables) — tolerância generosa
porque o builder v3 pode variar com richness do config.
"""
import hashlib
import io
import os
import re
import sys
import zipfile

try:
    from docx import Document
except ImportError:
    sys.stderr.write("ERR: python-docx not installed (pip install python-docx)\n")
    sys.exit(2)

# Baseline VF Luciano (canônico aprovado pelo CEO)
VF_PARAS = 325
VF_TABLES = 10
PARAS_MIN_RATIO = 0.5   # >= 163 paras
PARAS_MAX_RATIO = 1.5   # <= 487 paras
TABLES_MIN_RATIO = 0.5  # >= 5 tables
MIN_IMAGES = 3          # logo PROEX + 2 charts
MIN_BYTES = 200_000     # 200 KB

IMMIGRATION_TERMS = [
    "EB-2", "EB2", "EB-1", "NIW", "Dhanasar", "Prong", "USCIS",
    "petitioner", "peticionário", "I-140", "waiver", "waive",
    "National Interest Waiver", "8 CFR", "immigration", "imigração",
    "labor certification", "adjudicator", "Kazarian",
    "extraordinary ability"
]

AUTHORITATIVE_DOMAINS = ("bea.gov", "bls.gov", "census.gov", "epi.org",
                         "siccode.com", "naics.com", "whitehouse.gov",
                         "floridajobs.org")

SOURCES_MARKERS = ["Fontes Primárias Consultadas", "Primary Sources Consulted", "12.6"]


def extract_all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    # Header / footer
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for p in header.paragraphs:
                    parts.append(p.text)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for p in footer.paragraphs:
                    parts.append(p.text)
    return "\n".join(p for p in parts if p)


def extract_embedded_images(docx_path):
    """Return list of (name, bytes) from word/media/ inside docx."""
    out = []
    with zipfile.ZipFile(docx_path) as z:
        for name in z.namelist():
            if name.startswith("word/media/") and name.lower().endswith((".png", ".jpg", ".jpeg")):
                out.append((os.path.basename(name), z.read(name)))
    return out


def is_image_mostly_uniform(img_bytes):
    """Detect charts with all-zero / near-constant data.
    Uses PIL if available; else falls back to byte-entropy heuristic."""
    try:
        from PIL import Image
        im = Image.open(io.BytesIO(img_bytes)).convert("L")
        px = list(im.getdata())
        if not px:
            return True
        mean = sum(px) / len(px)
        variance = sum((p - mean) ** 2 for p in px) / len(px)
        # Low variance = uniform (chart all-zeros renders mostly white)
        return variance < 50
    except ImportError:
        # Fallback: file size < 15KB is very suspicious for a chart PNG
        return len(img_bytes) < 15_000


def run_checks(docx_path):
    errors, warnings, passed = [], [], []

    if not os.path.isfile(docx_path):
        return [f"file not found: {docx_path}"], [], []

    size = os.path.getsize(docx_path)
    if size < MIN_BYTES:
        errors.append(f"file size {size:,} < minimum {MIN_BYTES:,} bytes")
    else:
        passed.append(f"file size {size:,} bytes ≥ {MIN_BYTES:,}")

    try:
        d = Document(docx_path)
    except Exception as e:
        errors.append(f"python-docx failed to open: {e}")
        return errors, warnings, passed

    paras = len(d.paragraphs)
    tables = len(d.tables)
    imgs = len(d.inline_shapes)

    paras_min = int(VF_PARAS * PARAS_MIN_RATIO)
    paras_max = int(VF_PARAS * PARAS_MAX_RATIO)
    tables_min = int(VF_TABLES * TABLES_MIN_RATIO)

    if paras < paras_min:
        errors.append(f"paragraphs={paras} < {paras_min} (min {PARAS_MIN_RATIO:.0%} of VF={VF_PARAS})")
    elif paras > paras_max:
        warnings.append(f"paragraphs={paras} > {paras_max} (max {PARAS_MAX_RATIO:.0%} of VF) — may be too verbose")
    else:
        passed.append(f"paragraphs={paras} in range [{paras_min}, {paras_max}]")

    if tables < tables_min:
        errors.append(f"tables={tables} < {tables_min} (min {TABLES_MIN_RATIO:.0%} of VF={VF_TABLES})")
    else:
        passed.append(f"tables={tables} ≥ {tables_min}")

    if imgs < MIN_IMAGES:
        errors.append(f"inline images={imgs} < {MIN_IMAGES} required (logo PROEX + 2 charts)")
    else:
        passed.append(f"inline images={imgs} ≥ {MIN_IMAGES}")

    text = extract_all_text(d)
    words = len(text.split())

    # Immigration terms
    leaked = []
    for term in IMMIGRATION_TERMS:
        matches = re.findall(rf"\b{re.escape(term)}\b", text, re.IGNORECASE)
        if matches:
            leaked.append((term, len(matches)))
    if leaked:
        for term, n in leaked:
            errors.append(f"immigration term leak: '{term}' occurs {n}x in DOCX text")
    else:
        passed.append("zero immigration terminology")

    # Sources section
    if not any(m in text for m in SOURCES_MARKERS):
        errors.append("sources section (12.6 / Primary Sources Consulted) not found in document")
    else:
        passed.append("sources section present")

    # Authoritative domain mentions
    auth_hits = sum(1 for d_ in AUTHORITATIVE_DOMAINS if d_ in text)
    if auth_hits < 2:
        errors.append(f"only {auth_hits} authoritative domain(s) mentioned; minimum 2")
    else:
        passed.append(f"authoritative domains mentioned: {auth_hits}")

    # Embedded images quality (detect all-zero charts)
    images = extract_embedded_images(docx_path)
    if not images:
        errors.append("no embedded images in word/media/ — logo/charts missing")
    else:
        uniform = [name for name, buf in images if is_image_mostly_uniform(buf)]
        # Logo (small icon) is expected to be relatively uniform; we only fail if 2+ uniform
        if len(uniform) > 1:
            errors.append(f"multiple images appear uniform (likely all-zero charts): {uniform}")
        elif len(uniform) == 1:
            warnings.append(f"one image appears uniform (probably the logo — OK): {uniform}")
        else:
            passed.append(f"{len(images)} embedded images, all with content variation")

    # Metrics summary for report
    metrics = {
        "paragraphs": paras, "tables": tables, "inline_images": imgs,
        "words": words, "bytes": size,
        "embedded_media_files": len(images),
    }
    return errors, warnings, passed, metrics


def main():
    if len(sys.argv) < 2:
        sys.stderr.write("Usage: qa_impacto_docx.py <generated.docx> [qa_report.md]\n")
        sys.exit(2)
    docx_path = sys.argv[1]
    report_path = sys.argv[2] if len(sys.argv) > 2 else None

    result = run_checks(docx_path)
    if len(result) == 3:  # file not found
        errors, warnings, passed = result
        metrics = {}
    else:
        errors, warnings, passed, metrics = result

    status = "PASSED" if not errors else "FAILED"
    lines = [
        f"# IMPACTO® QA Report — {status}",
        "",
        f"**File:** `{docx_path}`",
        "",
        "## Metrics",
        ""
    ]
    for k, v in metrics.items():
        lines.append(f"- {k}: {v:,}" if isinstance(v, int) else f"- {k}: {v}")

    lines.append("")
    lines.append(f"## Checks — {len(passed)} passed · {len(warnings)} warnings · {len(errors)} errors")
    lines.append("")
    if passed:
        lines.append("### ✓ Passed")
        for p in passed:
            lines.append(f"- {p}")
        lines.append("")
    if warnings:
        lines.append("### ⚠ Warnings (non-blocking)")
        for w in warnings:
            lines.append(f"- {w}")
        lines.append("")
    if errors:
        lines.append("### ✗ Errors (BLOCKING)")
        for e in errors:
            lines.append(f"- {e}")
        lines.append("")

    report = "\n".join(lines)
    print(report)
    if report_path:
        os.makedirs(os.path.dirname(os.path.abspath(report_path)) or ".", exist_ok=True)
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(report + "\n")

    sys.exit(1 if errors else 0)


if __name__ == "__main__":
    main()
