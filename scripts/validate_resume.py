#!/usr/bin/env python3
"""
Validador automático de résumé pós-geração.
Implementa checks de QUALITY_REVIEWER + regras do sistema V4.
Deve ser chamado AUTOMATICAMENTE após cada geração.

Uso:
    python validate_resume.py /path/to/resume.docx [eb2_niw|eb1a]

    from validate_resume import ResumeValidator
    report = ResumeValidator('resume.docx', 'eb2_niw').validate()
"""
import re
import sys
from pathlib import Path
from docx import Document


class ResumeValidator:

    # Conteúdo proibido com regex
    CHECKS = {
        'exhibit_references': {
            'pattern': r'\[Exhibit\s*\d+\]|\bExhibit\s+\d+\b|\bEvidence\s+\d+\b',
            'severity': 'CRITICAL',
            'msg': 'Referência numérica a Exhibit/Evidence — não pertence ao Résumé',
            'applies': ['eb2_niw', 'eb1a'],
        },
        'recommendation_letters': {
            'pattern': r'recommendation letter|letter from (?:Dr\.|Prof\.|Mr\.|Ms\.)',
            'severity': 'CRITICAL',
            'msg': 'Menção a carta de recomendação — Résumé demonstra via evidências, não opiniões',
            'applies': ['eb2_niw', 'eb1a'],
        },
        'exhibit_list': {
            'pattern': r'(?:List of|Index of|Summary of)\s*(?:Exhibit|Evidence)',
            'severity': 'CRITICAL',
            'msg': 'Lista de Exhibits — pertence à Cover Letter',
            'applies': ['eb2_niw', 'eb1a'],
        },
        'opinion_language': {
            'pattern': r'\bI believe\b|\bwe believe\b|\bI think\b|\bin my opinion\b',
            'severity': 'HIGH',
            'msg': 'Linguagem opinativa — Résumé deve ser factual',
            'applies': ['eb2_niw', 'eb1a'],
        },
        'thumbnail_placeholder': {
            'pattern': r'\[THUMBNAIL',
            'severity': 'CRITICAL',
            'msg': 'Placeholder de thumbnail — deve ser imagem real do PDF',
            'applies': ['eb2_niw', 'eb1a'],
        },
        # EB-2 NIW específicos
        'dhanasar_in_resume': {
            'pattern': r'\bDhanasar\b|three[- ]prong|Matter of Dhanasar',
            'severity': 'CRITICAL',
            'msg': 'Framework Dhanasar — pertence EXCLUSIVAMENTE à Cover Letter',
            'applies': ['eb2_niw'],
        },
        'criteria_c1_c10': {
            'pattern': r'\bC[1-9]0?\b.*(?:criterion|critério)',
            'severity': 'CRITICAL',
            'msg': 'Critérios C1-C10 — estrutura EB-1A, não EB-2 NIW',
            'applies': ['eb2_niw'],
        },
        'petitioner_beneficiary': {
            'pattern': r'\bpetitioner\b|\bbeneficiary\b',
            'severity': 'HIGH',
            'msg': 'Termos "petitioner"/"beneficiary" — usar nome do profissional',
            'applies': ['eb2_niw', 'eb1a'],
        },
    }

    def __init__(self, docx_path: str, doc_type: str = 'eb2_niw'):
        self.docx_path = docx_path
        self.doc_type = doc_type
        self.doc = Document(docx_path)
        self.full_text = self._extract_text()
        self.issues = []

    def _extract_text(self) -> str:
        parts = [p.text for p in self.doc.paragraphs]
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return '\n'.join(parts)

    def check_forbidden_content(self):
        for name, check in self.CHECKS.items():
            if self.doc_type not in check['applies']:
                continue
            matches = re.findall(check['pattern'], self.full_text, re.IGNORECASE)
            if matches:
                self.issues.append({
                    'check': name,
                    'severity': check['severity'],
                    'msg': check['msg'],
                    'count': len(matches),
                    'samples': matches[:3],
                })

    def check_font(self):
        for para in self.doc.paragraphs[:50]:
            for run in para.runs:
                if run.font.name and run.font.name not in ('Garamond', None):
                    self.issues.append({
                        'check': 'font_not_garamond',
                        'severity': 'HIGH',
                        'msg': f'Fonte "{run.font.name}" encontrada — deve ser 100% Garamond',
                        'count': 1,
                    })
                    return

    def check_images(self):
        img_count = sum(1 for r in self.doc.part.rels.values() if "image" in r.reltype)
        if img_count < 3:
            self.issues.append({
                'check': 'low_image_count',
                'severity': 'HIGH',
                'msg': f'Apenas {img_count} imagens — benchmarks têm 37-61 thumbnails',
                'count': img_count,
            })

    def check_page_count(self):
        # Estimativa: ~400 palavras por página
        words = len(self.full_text.split())
        est_pages = words / 400
        if est_pages < 15:
            self.issues.append({
                'check': 'low_page_count',
                'severity': 'HIGH',
                'msg': f'~{int(est_pages)} páginas estimadas ({words} palavras) — mínimo esperado: 20+',
                'count': int(est_pages),
            })

    def validate(self) -> dict:
        self.check_forbidden_content()
        self.check_font()
        self.check_images()
        self.check_page_count()

        critical = [i for i in self.issues if i['severity'] == 'CRITICAL']
        high = [i for i in self.issues if i['severity'] == 'HIGH']

        verdict = 'PASS'
        if critical:
            verdict = 'FAIL'
        elif high:
            verdict = 'CONCERNS'

        report = {
            'file': self.docx_path,
            'doc_type': self.doc_type,
            'verdict': verdict,
            'critical': len(critical),
            'high': len(high),
            'issues': self.issues,
        }

        print(f"\n{'='*60}")
        print(f"QUALITY REVIEW — {Path(self.docx_path).name}")
        print(f"Verdict: {verdict} | {len(critical)} CRITICAL | {len(high)} HIGH")
        print(f"{'='*60}")
        for issue in self.issues:
            icon = 'X' if issue['severity'] == 'CRITICAL' else '!'
            print(f"  [{icon}] {issue['check']}: {issue['msg']}")
            if 'samples' in issue:
                print(f"      Exemplos: {issue['samples'][:2]}")

        return report


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python validate_resume.py <arquivo.docx> [eb2_niw|eb1a]")
        sys.exit(1)
    path = sys.argv[1]
    dtype = sys.argv[2] if len(sys.argv) > 2 else 'eb2_niw'
    ResumeValidator(path, dtype).validate()
