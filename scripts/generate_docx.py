#!/usr/bin/env python3
"""
Petition Engine — DOCX Generator
Converts structured text (markdown-like) into professionally formatted DOCX.
"""

import sys
import os
import json
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


STYLES = {
    'font_family': 'Times New Roman',
    'body_size': Pt(12),
    'heading1_size': Pt(14),
    'heading2_size': Pt(13),
    'heading3_size': Pt(12),
    'line_spacing': 1.15,
    'paragraph_spacing_after': Pt(6),
    'margin': Inches(1),
}


def generate_docx(content: str, output_path: str, template_path: str = None, metadata: dict = None) -> dict:
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()
        setup_default_styles(doc)

    for section in doc.sections:
        section.top_margin = STYLES['margin']
        section.bottom_margin = STYLES['margin']
        section.left_margin = STYLES['margin']
        section.right_margin = STYLES['margin']

    if metadata:
        doc.core_properties.title = metadata.get('title', '')
        doc.core_properties.author = metadata.get('author', 'Petition Engine')
        doc.core_properties.subject = metadata.get('subject', '')

    lines = content.split('\n')
    current_paragraph = []
    word_count = 0

    for line in lines:
        stripped = line.strip()

        if stripped == '---':
            if current_paragraph:
                word_count += flush_paragraph(doc, current_paragraph)
                current_paragraph = []
            doc.add_page_break()
            continue

        if stripped.startswith('# ') and not stripped.startswith('## '):
            if current_paragraph:
                word_count += flush_paragraph(doc, current_paragraph)
                current_paragraph = []
            add_heading(doc, stripped[2:], level=1)
            continue

        if stripped.startswith('## ') and not stripped.startswith('### '):
            if current_paragraph:
                word_count += flush_paragraph(doc, current_paragraph)
                current_paragraph = []
            add_heading(doc, stripped[3:], level=2)
            continue

        if stripped.startswith('### '):
            if current_paragraph:
                word_count += flush_paragraph(doc, current_paragraph)
                current_paragraph = []
            add_heading(doc, stripped[4:], level=3)
            continue

        if stripped == '':
            if current_paragraph:
                word_count += flush_paragraph(doc, current_paragraph)
                current_paragraph = []
            continue

        current_paragraph.append(stripped)

    if current_paragraph:
        word_count += flush_paragraph(doc, current_paragraph)

    doc.save(output_path)
    page_estimate = max(1, word_count // 250)

    return {
        "status": "ok",
        "output": output_path,
        "word_count": word_count,
        "page_count_estimate": page_estimate,
        "size_bytes": os.path.getsize(output_path),
    }


def setup_default_styles(doc: Document):
    style = doc.styles['Normal']
    font = style.font
    font.name = STYLES['font_family']
    font.size = STYLES['body_size']
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = STYLES['paragraph_spacing_after']
    paragraph_format.line_spacing = STYLES['line_spacing']


def add_heading(doc: Document, text: str, level: int):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text.upper() if level == 1 else text)
    run.bold = True
    sizes = {1: STYLES['heading1_size'], 2: STYLES['heading2_size'], 3: STYLES['heading3_size']}
    run.font.size = sizes.get(level, STYLES['body_size'])
    run.font.color.rgb = RGBColor(0, 0, 0)


def flush_paragraph(doc: Document, lines: list) -> int:
    text = ' '.join(lines)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    word_count = len(text.split())

    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)

    return word_count


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(json.dumps({"error": "Uso: generate_docx.py <content_file> <output.docx> [template.docx] [metadata_json]"}))
        sys.exit(1)

    content_file = sys.argv[1]
    output_path = sys.argv[2]
    template = sys.argv[3] if len(sys.argv) > 3 else None
    metadata = json.loads(sys.argv[4]) if len(sys.argv) > 4 else None

    with open(content_file, 'r', encoding='utf-8') as f:
        content = f.read()

    result = generate_docx(content, output_path, template, metadata)
    print(json.dumps(result, indent=2))
