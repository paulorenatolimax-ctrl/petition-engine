#!/usr/bin/env python3
"""
Résumé EB-2 NIW — Ricardo Augusto Borges Porfirio Pereira
100% Português Brasileiro | Garamond | Navy/Teal palette
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# CONSTANTS
# ============================================================
NAVY = RGBColor(0x2D, 0x3E, 0x50)
TEAL = RGBColor(0x34, 0x98, 0xA2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)

FONT_NAME = "Garamond"
FONT_HEADER_NAME = Pt(20)
FONT_SECTION = Pt(11)
FONT_SUBSECTION = Pt(10)
FONT_BODY = Pt(10.5)
FONT_SMALL = Pt(9.5)
FONT_CONTACT = Pt(9)

META_WIDTH = 5760
THUMB_WIDTH = 4320

OUTPUT_DIR = (
    "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/"
    "_2. MEUS CASOS/2024/"
    "Ricardo Augusto Borges Porfirio Pereira (EB-2NIW)/"
    "_Forjado por Petition Engine/"
)
OUTPUT_FILE = "resume_eb2_niw_Ricardo_Augusto_Borges_Porfirio_Pereira.docx"

BENEFICIARY_NAME = "RICARDO AUGUSTO BORGES PORFIRIO PEREIRA"
BENEFICIARY_EMAIL = "ricardoa.pereira@hotmail.com"
BENEFICIARY_PHONE = "+55 (62) 99114-6810"
BENEFICIARY_LOCATION = "Goiânia, Goiás — Brasil"
SOC_CODE = "11-9021.00"


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def set_cell_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    tcPr.append(tcW)


def set_cell_margins(cell, top=0, bottom=0, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_full_width(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)


def add_run(paragraph, text, font_name=FONT_NAME, size=FONT_BODY, color=BLACK,
            bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rPr.append(parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" '
        f'w:eastAsia="{font_name}" w:cs="{font_name}"/>'
    ))
    return run


def add_body_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       space_after=Pt(6), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = Pt(14)
    add_run(p, text, size=FONT_BODY, color=BLACK)
    return p


def add_body_paragraph_rich(doc, segments, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_after=Pt(6)):
    """Add paragraph with mixed formatting. segments = list of (text, bold, italic)."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    for text, bold, italic in segments:
        add_run(p, text, size=FONT_BODY, color=BLACK, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)
    p.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix:
        add_run(p, f"• {bold_prefix}", size=FONT_BODY, color=BLACK, bold=True)
        add_run(p, f" {text}", size=FONT_BODY, color=BLACK)
    else:
        add_run(p, f"• {text}", size=FONT_BODY, color=BLACK)
    return p


def add_navy_section_header(doc, title, page_break=True):
    if page_break:
        doc.add_page_break()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "2D3E50")
    set_cell_margins(cell, top=60, bottom=60, left=200, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title.upper(), size=FONT_SECTION, color=WHITE, bold=True)
    remove_table_borders(table)
    set_table_full_width(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_teal_sub_header(doc, title):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "3498A2")
    set_cell_margins(cell, top=40, bottom=40, left=200, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title, size=FONT_SUBSECTION, color=WHITE, bold=True)
    remove_table_borders(table)
    set_table_full_width(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_evidence_block(doc, exhibit_num, metadata_lines, impact_text):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    set_cell_width(left_cell, META_WIDTH)
    set_cell_width(right_cell, THUMB_WIDTH)
    set_cell_borders(left_cell, "CCCCCC", 4)
    set_cell_borders(right_cell, "CCCCCC", 4)
    set_cell_margins(left_cell, top=80, bottom=80, left=120, right=120)
    set_cell_margins(right_cell, top=80, bottom=80, left=80, right=80)

    first = True
    for label, value in metadata_lines:
        if first:
            p = left_cell.paragraphs[0]
            first = False
        else:
            p = left_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        add_run(p, f"{label}: ", size=Pt(10), color=BLACK, bold=True)
        add_run(p, value, size=Pt(10), color=BLACK)

    if impact_text:
        p_label = left_cell.add_paragraph()
        p_label.paragraph_format.space_before = Pt(6)
        p_label.paragraph_format.space_after = Pt(2)
        add_run(p_label, "Descrição e Impacto", size=FONT_SMALL, color=NAVY, bold=True)

        p_impact = left_cell.add_paragraph()
        p_impact.paragraph_format.space_after = Pt(4)
        p_impact.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p_impact, impact_text, size=FONT_SMALL, color=DARK_GRAY, italic=True)

    p_thumb = right_cell.paragraphs[0]
    p_thumb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_run(p_thumb, f"[THUMBNAIL — Exhibit {exhibit_num}]",
            size=FONT_CONTACT, color=MED_GRAY, italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_company_box(doc, company, role, period, location, sector):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")
    set_cell_borders(cell, "CCCCCC", 4)
    set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
    set_table_full_width(table)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_run(p, company, size=FONT_BODY, color=NAVY, bold=True)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    add_run(p2, role, size=FONT_BODY, color=BLACK, bold=True)
    add_run(p2, f"  |  {period}", size=FONT_SMALL, color=DARK_GRAY)

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    add_run(p3, f"{location}  |  {sector}", size=FONT_SMALL, color=MED_GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ============================================================
# DOCUMENT SETUP
# ============================================================

def setup_document():
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.header_distance = Inches(0)
    return doc


def build_header(doc):
    header_section = doc.sections[0].header
    header_section.is_linked_to_previous = False

    table = header_section.add_table(rows=3, cols=2, width=Inches(7.2))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)
    remove_table_borders(table)

    # Row 0: Navy — Name + Contact
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "2D3E50")
    set_cell_margins(table.cell(0, 0), top=120, bottom=40, left=200, right=100)
    set_cell_margins(table.cell(0, 1), top=120, bottom=40, left=100, right=200)

    p_name = table.cell(0, 0).paragraphs[0]
    add_run(p_name, BENEFICIARY_NAME, size=FONT_HEADER_NAME, color=WHITE, bold=True)

    p_contact = table.cell(0, 1).paragraphs[0]
    p_contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_contact, BENEFICIARY_EMAIL, size=FONT_CONTACT, color=WHITE)
    p_contact2 = table.cell(0, 1).add_paragraph()
    p_contact2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_contact2, BENEFICIARY_PHONE, size=FONT_CONTACT, color=WHITE)
    p_contact3 = table.cell(0, 1).add_paragraph()
    p_contact3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_contact3, BENEFICIARY_LOCATION, size=FONT_CONTACT, color=WHITE)

    # Row 1: Navy — RESUME + SOC + EB-2 NIW
    for cell in table.rows[1].cells:
        set_cell_shading(cell, "2D3E50")
    set_cell_margins(table.cell(1, 0), top=20, bottom=60, left=200, right=100)
    set_cell_margins(table.cell(1, 1), top=20, bottom=60, left=100, right=200)

    p_resume = table.cell(1, 0).paragraphs[0]
    add_run(p_resume, "RÉSUMÉ", size=FONT_SECTION, color=WHITE, bold=True)
    add_run(p_resume, f"  |  SOC/O*Net {SOC_CODE}  |  EB-2 NIW",
            size=FONT_CONTACT, color=WHITE)

    p_soc_desc = table.cell(1, 1).paragraphs[0]
    p_soc_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_soc_desc, "Construction Managers", size=FONT_CONTACT, color=WHITE, italic=True)

    # Row 2: Teal accent bar
    for cell in table.rows[2].cells:
        set_cell_shading(cell, "3498A2")
    set_row_height(table.rows[2], 4)

    for p in header_section.paragraphs:
        if p.text == "":
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)


def build_footer(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False

    table = footer.add_table(rows=1, cols=1, width=Inches(7.2))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "2D3E50")
    set_cell_margins(cell, top=40, bottom=40, left=200, right=200)
    set_table_full_width(table)
    remove_table_borders(table)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = p.add_run("Page ")
    run1.font.name = FONT_NAME
    run1.font.size = FONT_CONTACT
    run1.font.color.rgb = WHITE

    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r1 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r1.append(fldChar1)
    p._p.append(r1)

    instrText = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr>'
        f'<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
    )
    p._p.append(instrText)

    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r2 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r2.append(fldChar2)
    p._p.append(r2)

    run_of = p.add_run(" of ")
    run_of.font.name = FONT_NAME
    run_of.font.size = FONT_CONTACT
    run_of.font.color.rgb = WHITE

    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r3 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r3.append(fldChar3)
    p._p.append(r3)

    instrText2 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr>'
        f'<w:instrText xml:space="preserve"> NUMPAGES </w:instrText></w:r>'
    )
    p._p.append(instrText2)

    fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r4 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r4.append(fldChar4)
    p._p.append(r4)

    for p in footer.paragraphs:
        if p.text == "":
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)


# ============================================================
# SECTION BUILDERS
# ============================================================

def build_synthesis(doc):
    add_navy_section_header(doc, "SÍNTESE PROFISSIONAL", page_break=False)

    add_body_paragraph(doc,
        "Ricardo Augusto Borges Porfirio Pereira é Engenheiro Civil Estrutural com mais de "
        "14 anos de experiência consolidada na liderança de projetos complexos e equipes "
        "multidisciplinares na indústria da construção civil. Com formação avançada — Bacharelado "
        "em Engenharia Civil pela Pontifícia Universidade Católica de Goiás (PUC-GO) e "
        "pós-graduação em Estruturas de Concreto Armado e Fundações pelo INBEC —, o peticionário "
        "reúne sólida base acadêmica e expertise técnica reconhecida em análise estrutural de "
        "concreto armado, estruturas metálicas, materiais mistos e reforço estrutural avançado "
        "com fibra de carbono e concreto protendido."
    )

    add_body_paragraph(doc,
        "Como fundador e CEO da RBP Construtora EIRELI desde 2017, o peticionário demonstrou "
        "capacidade excepcional de gestão e expansão operacional, conduzindo a empresa a operar "
        "simultaneamente em sete estações de trabalho distribuídas por cinco cidades, com uma "
        "equipe de 90 profissionais e crescimento exponencial que triplicou a escala de operações "
        "em um único exercício. Essa trajetória de expansão é sustentada por domínio técnico que "
        "atraiu grandes clientes corporativos multinacionais e nacionais, incluindo JBS S/A, "
        "Brookfield Engenharia S.A., Cosmed Indústria de Cosméticos e Medicamentos S/A, Flora "
        "Produtos de Higiene e Limpeza S.A. e Halexistar Indústria Farmacêutica S.A."
    )

    add_body_paragraph(doc,
        "No campo técnico-científico, o peticionário é reconhecido por contribuições significativas "
        "em reforço estrutural avançado. O projeto de reforço do tanque de aeração da Cosmed, "
        "desenvolvido durante sua atuação na Hirata e Associados — Projetos Estruturais, utilizou "
        "vigas protendidas para estabilização de estrutura comprometida e foi apresentado no 58º "
        "Congresso Brasileiro do Concreto (IBRACON, 2016), além de receber indicação a prêmio pelo "
        "PTI Journal (Post-Tensioning Institute), Volume 13, Número 2, de dezembro de 2017 — "
        "publicação internacional de referência em pós-tensionamento."
    )

    add_body_paragraph(doc,
        "A expertise forense do peticionário é evidenciada pela atuação como perito técnico "
        "contratado pelo TRT da 18ª Região para avaliar a integridade estrutural do Complexo "
        "Trabalhista de Goiânia após sinistro de incêndio, demandando análise de estrutura de "
        "concreto armado e protendido em 8.863,81 m². Complementarmente, os múltiplos contratos "
        "de recuperação estrutural executados para a JBS S/A — em unidades industriais de Goiás, "
        "Mato Grosso e Minas Gerais, com intervenção em mais de 1.000 m² de lajes, vigas e "
        "pilares — demonstram capacidade comprovada de executar obras de alta complexidade em "
        "ambientes industriais operacionais."
    )

    add_body_paragraph(doc,
        "O peticionário é autor de três artigos científicos publicados na Revista FT (ISSN "
        "1678-0817, Qualis B2), abordando tecnologias avançadas de reforço estrutural, "
        "reabilitação de edificações históricas e inovações em estruturas de concreto armado. "
        "Publicou, ainda, o e-book \"Gestão de Projetos de Engenharia Civil em Ambientes "
        "Multiculturais: Estratégias para o Mercado Global\" (Golden Int., ISBN 978-65-986494-2-5, "
        "2025), com abordagem sobre metodologias ágeis, BIM, inteligência artificial e gestão "
        "intercultural em projetos internacionais de engenharia."
    )

    add_body_paragraph(doc,
        "Com experiência internacional que inclui formação nos Estados Unidos (University of "
        "California San Diego — Business English), Canadá (Burnaby South Secondary School, "
        "Vancouver) e Espanha (Colegio Nestor Almendros, Sevilha), além de atuação profissional "
        "recente como Senior Project Manager e Project Engineer na Karins Engineering em Orlando, "
        "Flórida, o peticionário demonstra capacidade de operar em contextos multiculturais e "
        "regulatórios diversificados. Eleito Presidente da Associação dos Produtores de Borracha "
        "Natural dos Estados de Goiás e Tocantins (APROB-GO/TO) para o biênio 2019-2021, "
        "evidencia liderança institucional e capacidade de articulação setorial que transcende "
        "o campo técnico da engenharia."
    )


def build_timeline(doc):
    add_navy_section_header(doc, "HISTÓRICO PROFISSIONAL")

    timeline_data = [
        ("2025", "Senior Project Manager / Project Engineer",
         "Karins Engineering", "Orlando, FL — EUA"),
        ("2017 — Presente", "Fundador e CEO",
         "RBP Construtora EIRELI", "Goiânia, GO — Brasil"),
        ("2017 — 2020", "Administrador e Assessor Técnico",
         "Corporio Empreendimentos Imobiliários", "Bruxelas / Goiânia"),
        ("2014 — 2017", "Engenheiro Estrutural",
         "Hirata e Associados — Projetos Estruturais", "Goiânia, GO — Brasil"),
        ("2012 — 2014", "Engenheiro Civil",
         "TCI Construtora e Incorporadora", "Goiânia, GO — Brasil"),
        ("2011 — 2012", "Estágio de Engenharia Civil",
         "Solar Botafogo", "Goiânia, GO — Brasil"),
        ("2010 — Presente", "Gestor",
         "Agrotex Agropecuária", "Aurilândia, GO — Brasil"),
    ]

    table = doc.add_table(rows=len(timeline_data) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)

    headers = ["Período", "Posição", "Organização", "Local"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "2D3E50")
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, size=FONT_SMALL, color=WHITE, bold=True)

    for row_idx, (periodo, posicao, org, local) in enumerate(timeline_data):
        bg = "FFFFFF" if row_idx % 2 == 0 else "F5F5F5"
        for col_idx, text in enumerate([periodo, posicao, org, local]):
            cell = table.cell(row_idx + 1, col_idx)
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=30, bottom=30, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            bold = col_idx == 1
            add_run(p, text, size=FONT_SMALL, color=BLACK if col_idx < 3 else DARK_GRAY,
                    bold=bold)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_experience(doc):
    add_navy_section_header(doc, "EXPERIÊNCIA PROFISSIONAL")

    # --- Karins Engineering ---
    add_company_box(doc,
        "Karins Engineering",
        "Senior Project Manager / Project Engineer",
        "Junho — Setembro 2025",
        "Orlando, Flórida — EUA",
        "Engenharia Estrutural e Administração de Obras"
    )
    add_body_paragraph(doc,
        "Atuação em administração de obras, levantamentos, inspeções estruturais, projeto e "
        "assessoria de restauração estrutural em oito projetos simultâneos sob a liderança do "
        "Presidente e CEO David Karins. Aplicação de metodologias avançadas de gerenciamento "
        "incluindo ArenaSoft Estimating e Oracle Primavera, em contexto regulatório norte-americano. "
        "Experiência direta com padrões construtivos do mercado dos Estados Unidos, integração com "
        "equipes multidisciplinares americanas e execução de projetos de restauração estrutural "
        "e administração de construção no estado da Flórida."
    )

    # --- RBP Construtora ---
    add_company_box(doc,
        "RBP Construtora EIRELI",
        "Fundador e CEO",
        "Julho 2017 — Presente",
        "Goiânia, Goiás — Brasil",
        "Engenharia Civil, Reforço Estrutural e Infraestrutura"
    )
    add_body_paragraph(doc,
        "Fundação e condução da RBP Construtora como empresa especializada em reforço estrutural, "
        "recuperação de edificações e execução de obras civis, estruturais e de infraestrutura. "
        "Expansão da operação para sete estações de trabalho simultâneas em cinco cidades (Goiânia, "
        "Senador Canedo, Aparecida de Goiânia, Barra do Garças e Luziânia), com equipe de 90 "
        "profissionais e crescimento que triplicou a escala operacional em um único exercício."
    )
    add_body_paragraph(doc,
        "Conquista de clientes corporativos de grande porte, incluindo JBS S/A (múltiplas unidades "
        "industriais), Brookfield Engenharia S.A., Cosmed Indústria de Cosméticos e Medicamentos, "
        "Flora Produtos de Higiene e Limpeza, Halexistar Indústria Farmacêutica, Plaza D'Oro "
        "Shopping e Condomínio Edifício Village Royal. Portfólio de 19 Anotações de "
        "Responsabilidade Técnica (ARTs) registradas no CREA-GO, abrangendo reforço estrutural com "
        "fibra de carbono, recuperação de concreto armado, projetos de fundações profundas, "
        "estruturas metálicas, drenagem e sistemas hidrossanitários."
    )

    # --- Corporio ---
    add_company_box(doc,
        "Corporio Empreendimentos Imobiliários",
        "Administrador e Assessor Técnico",
        "Novembro 2017 — Maio 2020",
        "Bruxelas, Bélgica / Goiânia, Brasil",
        "Incorporação e Desenvolvimento Imobiliário"
    )
    add_body_paragraph(doc,
        "Diretor da empresa no Brasil, responsável por estudos de viabilidade para incorporação "
        "e loteamentos, assessoria técnica e estratégica, e gestão de contratos. Coordenação de "
        "projetos de empreendimentos residenciais, incluindo o Loteamento Master Ville (Juína-MT) "
        "e o Residencial Morro dos Ventos (Rosário Oeste-MT), com foco em otimização de processos "
        "e redução de não-conformidades em 25%. Atuação de coordenação remota entre a sede "
        "europeia e as operações brasileiras, com aumento de 20% na eficiência da gestão remota."
    )

    # --- Hirata ---
    add_company_box(doc,
        "Hirata e Associados — Projetos Estruturais",
        "Engenheiro Estrutural",
        "Julho 2014 — Julho 2017",
        "Goiânia, Goiás — Brasil",
        "Projetos Estruturais e Reforço Avançado"
    )
    add_body_paragraph(doc,
        "Desenvolvimento de soluções inovadoras e econômicas para estabilização e reforço de "
        "sistemas estruturais existentes, incluindo o premiado projeto de reforço do tanque de "
        "aeração da Cosmed utilizando concreto protendido. Análise de estruturas existentes para "
        "identificação de fragilidades e recomendação de estratégias de remediação. Realização "
        "de ensaios de resistência e carga para determinação da segurança e integridade de "
        "estruturas. Inspeção de canteiros de obras para garantia de conformidade com planos e "
        "especificações. Domínio de softwares especializados: SAP2000, ETABS e TQS. Atuação sob "
        "a mentoria de Ademar Toyonori Hirata, engenheiro estrutural com mais de 50 anos de "
        "experiência e referência nacional no setor."
    )

    # --- TCI ---
    add_company_box(doc,
        "TCI Construtora e Incorporadora",
        "Engenheiro Civil",
        "Setembro 2012 — Julho 2014",
        "Goiânia, Goiás — Brasil",
        "Construção e Incorporação — Torres Residenciais de Alto Porte"
    )
    add_body_paragraph(doc,
        "Gerenciamento de duas torres residenciais de 44 e 42 andares com equipe de 300 "
        "trabalhadores. Coordenação com empreiteiros, fornecedores e partes interessadas para "
        "garantia de conclusão bem-sucedida dos projetos. Utilização de softwares avançados — "
        "AutoCAD, Revit e MS Project — para elaboração de plantas detalhadas e documentos "
        "técnicos. Execução de análises de solo e estruturais para identificação de riscos e "
        "desenvolvimento de medidas de mitigação. Colaboração interdepartamental para garantia "
        "de consistência e qualidade do projeto."
    )


def build_contributions(doc):
    add_navy_section_header(doc, "CONTRIBUIÇÕES TÉCNICAS E PROFISSIONAIS")

    # ── TEMA 1: Fibra de Carbono e Protendido ──
    add_teal_sub_header(doc,
        "Reforço Estrutural com Fibra de Carbono e Concreto Protendido")

    add_body_paragraph(doc,
        "O peticionário desenvolveu expertise técnica diferenciada em tecnologias avançadas de "
        "reforço estrutural, com foco na aplicação de fibra de carbono e sistemas de concreto "
        "protendido para recuperação de estruturas comprometidas. Esta especialização, "
        "desenvolvida inicialmente na Hirata e Associados e consolidada na RBP Construtora, "
        "combina análise computacional avançada com técnicas de intervenção que preservam a "
        "integridade estrutural original enquanto ampliam a capacidade de carga e a vida útil "
        "das edificações. A abordagem do peticionário integra normas brasileiras (ABNT NBR 6118, "
        "NBR 8800) e internacionais (ACI 318, ACI 440.2R, Eurocode 2), permitindo soluções "
        "compatíveis com padrões globais de engenharia estrutural."
    )

    add_evidence_block(doc, 1, [
        ("Evidência", "Brookfield Engenharia S.A. — Reforço com Fibra de Carbono"),
        ("Tipo", "Anotação de Responsabilidade Técnica (ART)"),
        ("Registro CREA-GO", "ART 1020180116216"),
        ("Data", "Fevereiro a Abril de 2018"),
        ("Contratante", "Brookfield Engenharia S.A."),
        ("Local", "Edifício The Expression, Setor Bueno, Goiânia-GO"),
        ("Escopo", "3.000 metros de laminados de fibra de carbono + 100 m² de reparo estrutural"),
    ],
        "Projeto de reforço estrutural em edifício residencial de 30 andares com aplicação de "
        "3.000 metros de laminados de fibra de carbono em modelos CUT-IN e bonded-to-beam, além "
        "de 36 metros lineares de laminados colados em vigas. A intervenção incluiu aumento de "
        "seção de concreto em vigas, lajes e pilares para correção de falha de projeto original "
        "causada por excesso de protensão, que havia gerado fissuração extensiva e deformações "
        "estruturais significativas. O domínio da tecnologia de fibra de carbono com módulo de "
        "elasticidade de 210 MPa demonstra capacidade técnica em soluções de vanguarda."
    )

    add_evidence_block(doc, 2, [
        ("Evidência", "Cosmed Indústria — Reforço com Concreto Protendido"),
        ("Tipo", "Anotação de Responsabilidade Técnica (ART)"),
        ("Registro CREA-GO", "ART 1020150213056"),
        ("Data", "Outubro a Dezembro de 2015"),
        ("Contratante", "Cosmed Indústria de Cosméticos e Medicamentos S/A"),
        ("Local", "Polo Coureiro, Senador Canedo-GO"),
        ("Escopo", "1.181,24 m² de estrutura protendida — Tanque de aeração da ETE"),
    ],
        "Projeto de reforço estrutural em concreto armado e protendido no tanque de aeração da "
        "estação de tratamento de efluentes (ETE) da Cosmed. A solução técnica envolveu o "
        "dimensionamento de vigas protendidas para criar alta inércia e estabilizar a deformação "
        "das paredes do tanque observada após o primeiro uso. Este projeto alcançou reconhecimento "
        "internacional ao ser apresentado no 58º Congresso Brasileiro do Concreto (IBRACON) e no "
        "III Seminário Latino-Americano de Protendido (SELAP), além de receber indicação a prêmio "
        "pelo Post-Tensioning Institute (PTI)."
    )

    add_evidence_block(doc, 3, [
        ("Evidência", "PTI Journal — Indicação a Prêmio Internacional"),
        ("Tipo", "Publicação Internacional — Reconhecimento por Entidade de Prestígio"),
        ("Publicação", "PTI Journal, Volume 13, Número 2, Dezembro de 2017"),
        ("Instituição", "Post-Tensioning Institute (PTI) — EUA"),
        ("Projeto", "Reforço do Tanque de Aeração — Senador Canedo, GO"),
    ],
        "O projeto de reforço do tanque de aeração da Cosmed foi destacado e indicado a prêmio "
        "na edição de dezembro de 2017 do PTI Journal, publicação internacional do "
        "Post-Tensioning Institute, entidade de referência mundial em sistemas de pós-tensionamento. "
        "A indicação valida a qualidade técnica e o caráter inovador da solução desenvolvida, "
        "posicionando o peticionário entre os profissionais reconhecidos internacionalmente no "
        "campo do concreto protendido. O PTI é a principal organização técnica dedicada a "
        "promover o avanço e a aplicação do concreto protendido."
    )

    # ── TEMA 2: Recuperação Industrial ──
    add_teal_sub_header(doc,
        "Recuperação Estrutural em Ambientes Industriais")

    add_body_paragraph(doc,
        "A atuação do peticionário em recuperação estrutural de instalações industriais demonstra "
        "capacidade técnica para executar intervenções complexas em ambientes operacionais onde a "
        "continuidade da produção é requisito inegociável. Os projetos executados para a JBS S/A "
        "— maior processadora de proteínas do mundo — em múltiplas unidades industriais, exigiram "
        "planejamento logístico rigoroso para garantir que reparos em pilares, vigas e lajes "
        "fossem realizados sem interrupção das operações fabris. A metodologia aplicada inclui "
        "projeção de grout, técnicas de remoção por pressão de água e elevação estrutural, "
        "conjugadas com análise de patologias em estruturas metálicas e de concreto armado em "
        "ambientes de alta agressividade."
    )

    add_evidence_block(doc, 4, [
        ("Evidência", "JBS S/A — Recuperação Estrutural de Grande Porte (Senador Canedo-GO)"),
        ("Tipo", "Anotação de Responsabilidade Técnica (ART)"),
        ("Registros CREA-GO", "ART 1020190205125 e ART 1020200049717"),
        ("Data", "Outubro 2019 — Março 2020 (fases contínuas)"),
        ("Contratante", "JBS S.A. — Distrito Agro Industrial, Senador Canedo-GO"),
        ("Escopo", "Mais de 400 unidades estruturais (pilares, vigas e lajes) em concreto armado"),
    ],
        "Execução de recuperação estrutural em duas fases na unidade industrial frigorífica da "
        "JBS S/A em Senador Canedo, abrangendo a substituição de 1.000 m² de lajes, vigas e "
        "pilares sem interrupção das operações industriais. As técnicas empregadas incluíram "
        "projeção de grout, remoção por pressão de água e elevação estrutural. O projeto "
        "representou uma das maiores intervenções de recuperação estrutural em planta industrial "
        "operacional na região Centro-Oeste, exigindo coordenação entre equipes de engenharia, "
        "produção e segurança do trabalho para manter a conformidade regulatória."
    )

    add_evidence_block(doc, 5, [
        ("Evidência", "JBS S/A — Análise Estrutural Multi-site (GO, MG, MT)"),
        ("Tipo", "Anotação de Responsabilidade Técnica (ART) — Pareceres Técnicos"),
        ("Registros CREA-GO", "ARTs 1020210057219, 1020210073792, 1020210092949, 1020210073922"),
        ("Data", "Março a Maio de 2021"),
        ("Contratante", "JBS S.A. — Goiânia, Ituiutaba, Iturama, Senador Canedo"),
        ("Escopo", "48 câmaras de resfriamento — análise de estruturas metálicas e concreto armado"),
    ],
        "Emissão de pareceres técnicos estruturais para análise das estruturas de apoio das "
        "câmaras de resfriamento de carcaças em quatro unidades industriais da JBS S/A, "
        "distribuídas nos estados de Goiás, Minas Gerais e Mato Grosso. A avaliação abrangeu "
        "estruturas metálicas e de concreto armado, totalizando 48 câmaras de resfriamento "
        "analisadas. A amplitude geográfica e o volume de unidades avaliadas demonstram a "
        "confiança depositada pela JBS S/A na expertise técnica do peticionário para diagnóstico "
        "de infraestrutura crítica em escala nacional."
    )

    add_evidence_block(doc, 6, [
        ("Evidência", "Flora Produtos de Higiene e Limpeza — Recuperação Industrial"),
        ("Tipo", "Anotação de Responsabilidade Técnica (ART)"),
        ("Registro CREA-GO", "ART 1020180234609"),
        ("Data", "Outubro a Dezembro de 2018"),
        ("Contratante", "Flora Produtos de Higiene e Limpeza S.A."),
        ("Local", "Luziânia-GO"),
        ("Escopo", "442,40 m² de laje de cobertura — tratamento anticorrosivo e impermeabilização"),
    ],
        "Recuperação estrutural da laje de cobertura da área de produção de sabonetes, incluindo "
        "tratamento de armaduras em processo de corrosão, remoção de manta asfáltica antiga e "
        "instalação de nova membrana PVC para impermeabilização. O projeto exigiu execução em "
        "ambiente industrial com requisitos específicos de higiene e segurança alimentar, "
        "demonstrando capacidade do peticionário de adaptar procedimentos de engenharia "
        "estrutural às exigências regulatórias de cada setor industrial."
    )

    # ── TEMA 3: Perícia Forense ──
    add_teal_sub_header(doc,
        "Perícia Estrutural e Avaliação Forense")

    add_body_paragraph(doc,
        "O peticionário é reconhecido como profissional qualificado para emissão de pareceres "
        "técnicos em contextos de alta responsabilidade, incluindo avaliações pós-sinistro para "
        "instituições públicas e análises de patologias estruturais em edificações residenciais e "
        "industriais. A capacidade de avaliar a integridade de estruturas comprometidas por "
        "incêndio, corrosão ou falhas de projeto, e de propor soluções tecnicamente fundamentadas "
        "e economicamente viáveis, posiciona o peticionário como referência em perícia estrutural "
        "na região Centro-Oeste do Brasil."
    )

    add_evidence_block(doc, 7, [
        ("Evidência", "TRT da 18ª Região — Avaliação Estrutural Pós-Incêndio"),
        ("Tipo", "Anotação de Responsabilidade Técnica (ART) — Parecer Técnico"),
        ("Registro CREA-GO", "ART 1020160020103 (coautoria da ART 1020160019862)"),
        ("Data", "Dezembro 2015 — Fevereiro 2016"),
        ("Contratante", "Instituição Pública Federal — Goiânia-GO"),
        ("Escopo", "8.863,81 m² de concreto armado e protendido — avaliação pós-sinistro"),
    ],
        "Avaliação estrutural e elaboração de parecer técnico sobre a estrutura de concreto "
        "armado e protendido do Complexo Trabalhista de Goiânia, afetada por sinistro de "
        "incêndio. O contexto de avaliação pós-sinistro exigiu expertise forense para emitir "
        "diagnóstico preciso e confiável sobre a integridade estrutural após evento danoso. "
        "A contratação por instituição pública federal para este tipo de avaliação técnica "
        "altamente especializada constitui reconhecimento objetivo da competência profissional "
        "do peticionário em perícia estrutural de alta complexidade."
    )

    add_evidence_block(doc, 8, [
        ("Evidência", "Condomínio Edifício Village Royal — Reparação por Corrosão"),
        ("Tipo", "Anotação de Responsabilidade Técnica (ART)"),
        ("Registro CREA-GO", "ART 1020170179486"),
        ("Data", "Outubro 2017 — Julho 2018"),
        ("Contratante", "Condomínio Edifício Village Royal — Setor Oeste, Goiânia-GO"),
        ("Escopo", "540 m² de reparo e reforço estrutural — subsolo e térreo"),
    ],
        "Desenvolvimento de projeto de reparo e reforço de estrutura em processo de corrosão "
        "no subsolo e térreo de edifício residencial, com execução de reforma completa da "
        "guarita e área de lazer incluindo estrutura metálica, alvenaria e troca de acabamentos. "
        "O projeto abrangeu diagnóstico de patologias, dimensionamento de reforço e execução "
        "integral das intervenções, demonstrando capacidade do peticionário de conduzir projetos "
        "de reabilitação estrutural desde o diagnóstico até a entrega final."
    )

    # ── TEMA 4: Grande Porte ──
    add_teal_sub_header(doc,
        "Projetos de Grande Porte e Infraestrutura Diversificada")

    add_body_paragraph(doc,
        "A trajetória do peticionário inclui a execução de projetos de engenharia civil de escala "
        "significativa, abrangendo complexos hoteleiros, empreendimentos comerciais, infraestrutura "
        "urbana e edificações industriais. A capacidade de gerenciar simultaneamente projetos "
        "diversificados em múltiplas localidades, coordenando equipes multidisciplinares e "
        "cumprindo prazos e especificações técnicas rigorosas, demonstra maturidade profissional "
        "e competência gerencial compatíveis com padrões internacionais de gestão de projetos "
        "de engenharia."
    )

    add_evidence_block(doc, 9, [
        ("Evidência", "SPE Olímpia Q27 — Complexo Hoteleiro de Grande Porte"),
        ("Tipo", "Anotação de Responsabilidade Técnica (ART)"),
        ("Registro CREA-GO", "ART 1020160094667"),
        ("Data", "Setembro a Novembro de 2015"),
        ("Contratante", "SPE Olímpia Q27 Empreendimentos Imobiliários S/A"),
        ("Local", "Olímpia-SP"),
        ("Escopo", "83.088,07 m² de estrutura de concreto armado — seções A a S"),
    ],
        "Análise técnica da estrutura de concreto armado e projeto de reforço estrutural em "
        "concreto armado e fibra de carbono para complexo hoteleiro de 83.088,07 m², abrangendo "
        "as seções A até S. O projeto de escala extraordinária — equivalente a mais de 10 campos "
        "de futebol em área construída — demonstra a capacidade do peticionário de lidar com "
        "estruturas de alta complexidade e grande envergadura, mantendo rigor técnico e "
        "conformidade com normas de segurança estrutural."
    )

    add_evidence_block(doc, 10, [
        ("Evidência", "Hiperposto 3 Amigos — Complexo Comercial Multifuncional"),
        ("Tipo", "Anotação de Responsabilidade Técnica (ART)"),
        ("Registro CREA-GO", "ART 1020170187151"),
        ("Data", "Maio 2017 — Maio 2018"),
        ("Contratante", "Hiperposto 3 Amigos Ltda."),
        ("Local", "Papillon Park, Aparecida de Goiânia-GO"),
        ("Escopo", "2.376,91 m² — fundação, concreto armado, estrutura metálica e terraplenagem"),
    ],
        "Projeto e execução completa de fundação e estrutura de concreto armado, cobertura "
        "metálica e piso para posto de combustível, além de duas edificações comerciais. "
        "O escopo abrangeu terraplenagem (4.210 m³), estrutura, alvenaria, coberturas, acabamentos "
        "e instalações prediais. A integração de múltiplas disciplinas de engenharia — fundações, "
        "concreto armado, estruturas metálicas e infraestrutura — em um único projeto comercial "
        "evidencia a versatilidade técnica e a capacidade de coordenação multidisciplinar do "
        "peticionário."
    )


def build_publications(doc):
    add_navy_section_header(doc, "PUBLICAÇÕES E ARTIGOS")

    add_body_paragraph(doc,
        "O peticionário mantém produção acadêmica e técnica ativa, com três artigos científicos "
        "publicados em periódico com revisão por pares e um livro técnico com ISBN, abordando "
        "temas de reforço estrutural, reabilitação de edificações históricas, inovação em concreto "
        "armado e gestão de projetos internacionais de engenharia civil."
    )

    add_evidence_block(doc, 11, [
        ("Publicação", "\"Tecnologias Avançadas de Reforço Estrutural: Soluções Inovadoras para a "
         "Preservação de Grandes Infraestruturas\""),
        ("Periódico", "Revista FT (ISSN 1678-0817, Qualis B2)"),
        ("Volume/Edição", "Volume 26, Edição 111, Páginas 28-29"),
        ("Data", "25 de Junho de 2022"),
        ("DOI", "10.69849/revistaft/cs10202206251428"),
    ],
        "Artigo científico com revisão por pares que aborda as tecnologias mais avançadas "
        "disponíveis para reforço estrutural de infraestruturas de grande porte, incluindo fibra "
        "de carbono, concreto protendido e técnicas de aumento de seção. O trabalho sistematiza "
        "a experiência profissional do peticionário e contribui para a disseminação de "
        "conhecimento técnico especializado no campo da engenharia estrutural brasileira."
    )

    add_evidence_block(doc, 12, [
        ("Publicação", "\"Abordagens Modernas em Reabilitação Estrutural de Edificações "
         "Históricas: Preservação e Sustentabilidade\""),
        ("Periódico", "Revista FT (ISSN 1678-0817, Qualis B2)"),
        ("Volume/Edição", "Volume 27, Edição 118, Páginas 46-47"),
        ("Data", "15 de Janeiro de 2023"),
        ("DOI", "10.69849/revistaft/cs10202301151446"),
    ],
        "Artigo científico multidisciplinar que analisa abordagens modernas na reabilitação "
        "estrutural de edificações históricas, destacando a compatibilidade entre materiais "
        "tradicionais e novas tecnologias, a aplicação de técnicas sustentáveis e a necessidade "
        "de diagnósticos estruturais detalhados. O trabalho inclui análise de casos como o "
        "Museu de Arte do Rio, Pinacoteca de São Paulo, Casa Burguesa do Porto e Biblioteca "
        "Joanina de Coimbra, demonstrando visão internacional e interdisciplinar."
    )

    add_evidence_block(doc, 13, [
        ("Publicação", "\"Inovações e Métodos Avançados para Reforço de Estruturas de "
         "Concreto Armado\""),
        ("Periódico", "Revista FT (ISSN 1678-0817, Qualis B2)"),
        ("Volume/Edição", "Volume 28, Edição 131, Páginas 58-59"),
        ("Data", "19 de Fevereiro de 2024"),
        ("DOI", "10.69849/revistaft/cs10202402191458"),
    ],
        "Artigo científico que apresenta inovações recentes e métodos avançados para reforço de "
        "estruturas de concreto armado, consolidando as contribuições técnicas do peticionário ao "
        "campo. A publicação contínua em periódico indexado com Qualis B2 demonstra compromisso "
        "com a disseminação do conhecimento técnico e participação ativa na comunidade acadêmica "
        "e profissional da engenharia estrutural."
    )

    add_evidence_block(doc, 14, [
        ("Publicação", "\"Gestão de Projetos de Engenharia Civil em Ambientes Multiculturais: "
         "Estratégias para o Mercado Global\""),
        ("Tipo", "E-book (Livro Técnico)"),
        ("Editora", "Golden Int., Ribeirão Preto-SP"),
        ("ISBN", "978-65-986494-2-5"),
        ("CDD", "624"),
        ("Ano", "2025"),
        ("Páginas", "33 páginas, 5 capítulos"),
    ],
        "Livro técnico que aborda a gestão de projetos de engenharia civil em contextos "
        "internacionais e multiculturais, incluindo discussão sobre metodologias ágeis (Scrum, "
        "Kanban), BIM (Building Information Modeling), inteligência artificial, sustentabilidade, "
        "Lean Construction e gestão intercultural. A publicação com ISBN e catalogação "
        "bibliográfica formal (CDD 624) demonstra capacidade do peticionário de produzir e "
        "disseminar conhecimento técnico sistematizado, com visão estratégica para o mercado "
        "global de engenharia."
    )


def build_education(doc):
    add_navy_section_header(doc, "FORMAÇÃO ACADÊMICA")

    # Pós-graduação
    table1 = doc.add_table(rows=1, cols=1)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table1)
    cell1 = table1.cell(0, 0)
    set_cell_shading(cell1, "F5F5F5")
    set_cell_borders(cell1, "CCCCCC", 4)
    set_cell_margins(cell1, top=80, bottom=80, left=150, right=150)

    p1 = cell1.paragraphs[0]
    add_run(p1, "Pós-Graduação em Estruturas de Concreto Armado e Fundações",
            size=FONT_BODY, color=NAVY, bold=True)
    p1b = cell1.add_paragraph()
    add_run(p1b, "INBEC — Instituto Brasileiro de Educação Continuada",
            size=FONT_BODY, color=BLACK)
    p1c = cell1.add_paragraph()
    add_run(p1c, "Especialização Lato Sensu", size=FONT_SMALL, color=DARK_GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Graduação
    table2 = doc.add_table(rows=1, cols=1)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table2)
    cell2 = table2.cell(0, 0)
    set_cell_shading(cell2, "F5F5F5")
    set_cell_borders(cell2, "CCCCCC", 4)
    set_cell_margins(cell2, top=80, bottom=80, left=150, right=150)

    p2 = cell2.paragraphs[0]
    add_run(p2, "Bacharelado em Engenharia Civil",
            size=FONT_BODY, color=NAVY, bold=True)
    p2b = cell2.add_paragraph()
    add_run(p2b, "Pontifícia Universidade Católica de Goiás (PUC-GO)",
            size=FONT_BODY, color=BLACK)
    p2c = cell2.add_paragraph()
    add_run(p2c, "Janeiro 2010 — Janeiro 2015  |  Goiânia, Goiás — Brasil",
            size=FONT_SMALL, color=DARK_GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_courses(doc):
    add_navy_section_header(doc, "CURSOS, PALESTRAS E CONFERÊNCIAS")

    courses = [
        ("58º Congresso Brasileiro do Concreto (IBRACON) e III SELAP",
         "Instituto Brasileiro do Concreto (IBRACON)",
         "Outubro 2016  |  Centro de Convenções Minascentro, Belo Horizonte-MG",
         "Apresentação do projeto de reforço do tanque de aeração da Cosmed com concreto "
         "protendido, posteriormente indicado a prêmio pelo PTI Journal. Congresso de referência "
         "nacional com participação de especialistas internacionais."),
        ("Business English Program",
         "University of California San Diego (UCSD) — EUA",
         "Julho — Agosto 2013",
         "Programa intensivo de inglês para negócios na UCSD, uma das 15 melhores universidades "
         "de pesquisa dos Estados Unidos, proporcionando imersão em comunicação empresarial "
         "em contexto acadêmico norte-americano."),
        ("International Students Program",
         "Burnaby South Secondary School — Vancouver, Canadá",
         "Julho — Agosto 2007",
         "Programa de intercâmbio internacional com foco em imersão cultural e linguística "
         "em ambiente educacional canadense."),
        ("Primer Bachillerato",
         "Colegio Nestor Almendros — Sevilha, Espanha",
         "Agosto 2008 — Janeiro 2009",
         "Programa acadêmico em instituição educacional espanhola, proporcionando fluência em "
         "espanhol e experiência em sistema educacional europeu."),
    ]

    for title, institution, period, description in courses:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_full_width(table)
        cell = table.cell(0, 0)
        set_cell_shading(cell, "F5F5F5")
        set_cell_borders(cell, "CCCCCC", 4)
        set_cell_margins(cell, top=60, bottom=60, left=150, right=150)

        p = cell.paragraphs[0]
        add_run(p, title, size=FONT_BODY, color=NAVY, bold=True)
        p2 = cell.add_paragraph()
        add_run(p2, institution, size=FONT_BODY, color=BLACK)
        p3 = cell.add_paragraph()
        add_run(p3, period, size=FONT_SMALL, color=DARK_GRAY)
        p4 = cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p4, description, size=FONT_SMALL, color=DARK_GRAY, italic=True)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_recommendations(doc):
    add_navy_section_header(doc, "CARTAS DE RECOMENDAÇÃO")

    add_body_paragraph(doc,
        "O peticionário apresenta cinco cartas de recomendação de profissionais qualificados que "
        "atestam suas competências técnicas, capacidade de liderança e contribuições ao campo da "
        "engenharia civil e estrutural, incluindo profissional estabelecido nos Estados Unidos."
    )

    recommenders = [
        ("Ademar Toyonori Hirata",
         "Fundador e Proprietário — Hirata e Associados, Projetos Estruturais",
         "Belo Horizonte, MG — Brasil",
         "Engenheiro estrutural com mais de 50 anos de experiência, referência nacional em "
         "projetos estruturais. Atesta a expertise do peticionário em reforço com fibra de "
         "carbono, concreto protendido e análise estrutural pós-sinistro, com base em "
         "colaboração profissional desde 2012."),
        ("Carlos Eduardo Rocha de Assis",
         "Engenheiro Associado — Hirata e Associados, Projetos Estruturais",
         "Goiânia, GO — Brasil",
         "Engenheiro estrutural que colaborou diretamente com o peticionário por mais de uma "
         "década, atestando redução de 15% nos índices de retrabalho, diminuição de 20% nos "
         "prazos de entrega e a qualidade técnica dos projetos de reforço estrutural."),
        ("Antônio Claret Gama",
         "Engenheiro de Planejamento — Trinus Co.",
         "Rio de Janeiro, RJ — Brasil",
         "Engenheiro de planejamento de fundo de investimento imobiliário que contratou a "
         "RBP Construtora para projetos de loteamento e desenvolvimento imobiliário, atestando "
         "redução de 25% em não-conformidades e economia de 15% nos custos dos projetos."),
        ("Thiago Avelino",
         "TCI Construtora e Incorporadora",
         "Goiânia, GO — Brasil",
         "Profissional da TCI Construtora que supervisionou a atuação do peticionário no "
         "gerenciamento de torres residenciais de 42 e 44 andares, atestando capacidade de "
         "liderança, visão estratégica e excelência técnica em projetos de alta complexidade."),
        ("David Karins",
         "Presidente e CEO — Karins Engineering",
         "Orlando, Flórida — EUA",
         "Presidente e CEO de empresa norte-americana de engenharia que empregou o peticionário "
         "como Senior Project Manager e Project Engineer em oito projetos de administração de "
         "obras, inspeções estruturais e restauração na Flórida. Atesta competência técnica, "
         "capacidade gerencial e adaptação ao mercado norte-americano."),
    ]

    for name, role, location, description in recommenders:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_full_width(table)
        cell = table.cell(0, 0)
        set_cell_shading(cell, "F5F5F5")
        set_cell_borders(cell, "CCCCCC", 4)
        set_cell_margins(cell, top=60, bottom=60, left=150, right=150)

        p = cell.paragraphs[0]
        add_run(p, name, size=FONT_BODY, color=NAVY, bold=True)
        p2 = cell.add_paragraph()
        add_run(p2, role, size=FONT_BODY, color=BLACK)
        p3 = cell.add_paragraph()
        add_run(p3, location, size=FONT_SMALL, color=DARK_GRAY)
        p4 = cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p4, description, size=FONT_SMALL, color=DARK_GRAY, italic=True)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ============================================================
# MAIN
# ============================================================

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = setup_document()
    build_header(doc)
    build_footer(doc)

    # Sections
    build_synthesis(doc)
    build_timeline(doc)
    build_experience(doc)
    build_contributions(doc)
    build_publications(doc)
    build_education(doc)
    build_courses(doc)
    build_recommendations(doc)

    # Save
    output_path = os.path.join(OUTPUT_DIR, OUTPUT_FILE)
    doc.save(output_path)
    print(f"Documento gerado com sucesso: {output_path}")
    print(f"Seções: 8 | Evidence blocks: 14 | Recomendações: 5")


if __name__ == "__main__":
    main()
