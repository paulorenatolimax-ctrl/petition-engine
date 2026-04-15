#!/usr/bin/env python3
"""
Quality Reviewer — Résumé EB-2 NIW Cristine Correa
Based on QUALITY_REVIEWER.md specifications
"""

import re
import sys
from docx import Document
from docx.shared import Inches, Emu

DOC_PATH = "/Users/paulo1844/Documents/_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2026/CRISTINE CORREA/_Forjado por Petition Engine/resume_eb2_niw_Cristine_Correa.docx"
DOC_TYPE = "eb2niw"

# Allowed colors (hex, lowercase, no #)
ALLOWED_COLORS = {
    "2d3e50", "3498a2", "ffffff", "000000", "333333",
    "666666", "f5f5f5", "cccccc", "fafafa", "auto", ""
}

# Forbidden words
FORBIDDEN_PATTERNS = {
    "R\\$": "S1 — Moeda brasileira no documento",
    r"\b(I|we)\s+believe\b": "CRITICAL — 'I/we believe' proibido",
    r"\b(I|we)\s+think\b": "HIGH — 'I/we think' proibido",
    r"\b(in conclusion|to summarize)\b": "HIGH — 'in conclusion/to summarize' proibido",
    r"\bprompt\b": "CRITICAL — Palavra 'prompt' é termo interno",
    r"(PROEX|Kortix|Carlos Avelino)": "CRITICAL — Referência interna proibida",
    r"(Version \d|Generated:|SaaS Evidence Architect|Petition Engine)": "CRITICAL — Termo interno do sistema",
    r"\b(standardized|padronizado|operates autonomously|self-sustaining|auto-sustent|plug.and.play|train.the.trainer|white.label|marca branca|client autonomy|founder dependency|scalable without|replicable by any|turnkey|chave.na.m)\b": "CRITICAL — Anti-Cristine V2: termo que destrói Prong 3",
    r"\b(consultoria|consulting)\b": "HIGH — Dispara flag VIBE/Dun & Bradstreet",
    r"\b(denial|negativa anterior|RFE anterior|previous petition|prior filing|refile|segunda tentativa|nova submissão|petição anterior|corrected approach|lessons learned from denial)\b": "CRITICAL — Histórico processual proibido",
    r"(23-1011|29-1069|17-201[1-9]|13-2011)": "CRITICAL — SOC code que exige validação de diploma",
    r"proposed\s+(venture|business)": "MEDIUM — Usar 'proposed endeavor'",
    r"\b(introducao|peticao|informacao|certificacao|formacao|avaliacao|ocupacao|operacao|integracao|migracao|capacitacao|micropigmentacao)\b": "CRITICAL — Acentuação faltante",
    r"Correa LLC": "CRITICAL — Referência ao endeavor negado",
}


def check_fonts(doc):
    """Check 1: 100% Garamond. S0 if Arial found."""
    issues = []
    font_counts = {}

    # Check paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            fname = run.font.name
            if fname:
                font_counts[fname] = font_counts.get(fname, 0) + 1
                if fname.lower() not in ("garamond", "garamond premier pro"):
                    if fname.lower() in ("arial", "calibri"):
                        issues.append(f"S0 CRITICAL: Fonte '{fname}' encontrada em parágrafo: '{run.text[:50]}...'")
                    else:
                        issues.append(f"S2 MODERATE: Fonte não-Garamond '{fname}' encontrada: '{run.text[:40]}...'")

    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        fname = run.font.name
                        if fname:
                            font_counts[fname] = font_counts.get(fname, 0) + 1
                            if fname.lower() not in ("garamond", "garamond premier pro"):
                                if fname.lower() in ("arial", "calibri"):
                                    issues.append(f"S0 CRITICAL: Fonte '{fname}' em tabela: '{run.text[:40]}...'")

    return issues, font_counts


def check_page_setup(doc):
    """Check 2: US Letter, 0.65" margins."""
    issues = []
    for section in doc.sections:
        w = section.page_width
        h = section.page_height

        # US Letter: 8.5" x 11"
        expected_w = Inches(8.5)
        expected_h = Inches(11)
        if abs(w - expected_w) > Inches(0.1):
            issues.append(f"S1 SEVERE: Largura da página {w/914400:.2f}\" (esperado 8.5\")")
        if abs(h - expected_h) > Inches(0.1):
            issues.append(f"S1 SEVERE: Altura da página {h/914400:.2f}\" (esperado 11\")")

        # Margins
        left = section.left_margin
        right = section.right_margin
        expected_margin = Inches(0.65)
        tolerance = Inches(0.03)

        if abs(left - expected_margin) > tolerance:
            issues.append(f"S1 SEVERE: Margem esquerda {left/914400:.2f}\" (esperado 0.65\")")
        if abs(right - expected_margin) > tolerance:
            issues.append(f"S1 SEVERE: Margem direita {right/914400:.2f}\" (esperado 0.65\")")

    return issues


def check_forbidden_words(doc):
    """Check 3: Forbidden patterns."""
    issues = []

    # Collect all text
    all_text = []
    for para in doc.paragraphs:
        all_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_text.append(para.text)

    merged = " ".join(all_text)

    for pattern, desc in FORBIDDEN_PATTERNS.items():
        matches = re.findall(pattern, merged, re.IGNORECASE)
        if matches:
            # Deduplicate
            unique = list(set([m if isinstance(m, str) else m[0] for m in matches]))[:3]
            issues.append(f"{desc}: encontrado {len(matches)}x — exemplos: {unique}")

    return issues


def check_images(doc):
    """Check 4: Images count."""
    issues = []
    image_count = 0

    for para in doc.paragraphs:
        for run in para.runs:
            if run._r.findall(f'.//{{{run._r.nsmap.get("w", "")}}}drawing'):
                image_count += 1

    # Check inline shapes
    inline_count = len(doc.inline_shapes)
    total = image_count + inline_count

    if total == 0:
        issues.append("S1 SEVERE: Zero imagens/thumbnails no documento (thumbnails pendentes)")

    # Check for [THUMBNAIL] placeholders
    placeholder_count = 0
    for para in doc.paragraphs:
        if "[THUMBNAIL]" in para.text or "[Thumbnail]" in para.text:
            placeholder_count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "[THUMBNAIL]" in para.text or "[Thumbnail]" in para.text:
                        placeholder_count += 1

    if placeholder_count > 0:
        issues.append(f"S2 MODERATE: {placeholder_count} placeholders de thumbnail pendentes")

    return issues, total, placeholder_count


def check_evidence_blocks(doc):
    """Check 5: Evidence blocks have impact inside."""
    issues = []
    evidence_count = 0
    impact_inside = 0

    for table in doc.tables:
        # Check for 2-column, 1-row tables (evidence blocks)
        if len(table.columns) == 2 and len(table.rows) == 1:
            evidence_count += 1
            left_text = table.cell(0, 0).text.lower()
            if "impacto" in left_text or "impact" in left_text or "descrição" in left_text:
                impact_inside += 1

    if evidence_count > 0 and impact_inside == 0:
        issues.append("S1 SEVERE: Nenhum evidence block contém impacto DENTRO")

    return issues, evidence_count, impact_inside


def check_required_sections(doc):
    """Check 8: Required sections for EB-2 NIW."""
    issues = []
    all_text = " ".join([p.text for p in doc.paragraphs])

    # Also check table text (section headers are in tables)
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text.append(cell.text)
    all_table = " ".join(table_text)
    combined = all_text + " " + all_table
    combined_lower = combined.lower()

    # Required sections
    checks = {
        "Síntese/Summary": any(x in combined_lower for x in ["síntese", "sintese", "synthesis", "executive summary"]),
        "Histórico/Timeline": any(x in combined_lower for x in ["histórico", "historico", "timeline", "gantt"]),
        "Proposed Endeavors": any(x in combined_lower for x in ["proposed endeavor", "empreendimento proposto"]),
        "Contribuições Técnicas": any(x in combined_lower for x in ["contribuições", "contribuicoes", "contributions"]),
        "Formação Acadêmica": any(x in combined_lower for x in ["formação", "formacao", "education", "acadêmica"]),
    }

    for section, found in checks.items():
        if not found:
            issues.append(f"S1 SEVERE: Seção obrigatória ausente: {section}")

    # Check Dhanasar reference
    if "dhanasar" not in combined_lower:
        issues.append("S2 MODERATE: Referência ao framework Dhanasar ausente")

    # Check BLS code pattern
    bls_pattern = r"\d{2}-\d{4}"
    if not re.search(bls_pattern, combined):
        issues.append("S2 MODERATE: Nenhum código BLS/SOC (XX-XXXX) encontrado")

    # Check for C1-C10 sections (FORBIDDEN in EB-2 NIW)
    c_pattern = r"\bC[1-9]0?\b.*?(criterion|critério)"
    if re.search(c_pattern, combined, re.IGNORECASE):
        issues.append("S1 SEVERE: Seções C1-C10 encontradas em documento EB-2 NIW (formato errado)")

    # Check EB-2 NIW label
    if "eb-2 niw" not in combined_lower and "eb2 niw" not in combined_lower:
        issues.append("S0 CRITICAL: Label 'EB-2 NIW' não encontrada no documento")

    return issues


def check_paragraph_length(doc):
    """Check 6: Paragraph length."""
    issues = []
    short_count = 0
    total_paras = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if len(text) > 30:  # Only check substantial paragraphs
            total_paras += 1
            # Estimate lines: ~85 chars per line at Garamond 10.5pt, 0.65" margins
            est_lines = len(text) / 85
            if est_lines < 2.5:
                short_count += 1

    if short_count > 5:
        issues.append(f"S2 MODERATE: {short_count} parágrafos potencialmente curtos (<2.5 linhas estimadas)")

    return issues, short_count, total_paras


def check_header_footer(doc):
    """Check 7: Header and footer."""
    issues = []

    for section in doc.sections:
        header = section.header
        footer = section.footer

        if not header or not header.tables:
            issues.append("S1 SEVERE: Header ausente ou sem tabela (barra navy)")

        if not footer or not footer.tables:
            issues.append("S1 SEVERE: Footer ausente ou sem tabela (barra navy)")

    return issues


def count_stats(doc):
    """Count document statistics."""
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)

    # Count total text
    all_text = " ".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += " " + cell.text

    word_count = len(all_text.split())

    return para_count, table_count, word_count


def main():
    print("=" * 70)
    print("QUALITY REVIEWER — Résumé EB-2 NIW — Cristine Correa")
    print("=" * 70)

    doc = Document(DOC_PATH)

    all_issues = []

    # Stats
    para_count, table_count, word_count = count_stats(doc)
    print(f"\n📊 ESTATÍSTICAS DO DOCUMENTO:")
    print(f"   Parágrafos: {para_count}")
    print(f"   Tabelas: {table_count}")
    print(f"   Palavras: {word_count:,}")

    # Check 1: Fonts
    print("\n[1/8] Verificando fontes...")
    font_issues, font_counts = check_fonts(doc)
    all_issues.extend(font_issues)
    print(f"   Fontes encontradas: {font_counts}")

    # Check 2: Page setup
    print("[2/8] Verificando configuração de página...")
    page_issues = check_page_setup(doc)
    all_issues.extend(page_issues)

    # Check 3: Forbidden words
    print("[3/8] Verificando palavras proibidas...")
    forbidden_issues = check_forbidden_words(doc)
    all_issues.extend(forbidden_issues)

    # Check 4: Images
    print("[4/8] Verificando imagens...")
    img_issues, img_count, placeholder_count = check_images(doc)
    all_issues.extend(img_issues)
    print(f"   Imagens: {img_count}, Placeholders: {placeholder_count}")

    # Check 5: Evidence blocks
    print("[5/8] Verificando evidence blocks...")
    ev_issues, ev_count, impact_count = check_evidence_blocks(doc)
    all_issues.extend(ev_issues)
    print(f"   Evidence blocks: {ev_count}, Com impacto: {impact_count}")

    # Check 6: Paragraph length
    print("[6/8] Verificando comprimento de parágrafos...")
    para_issues, short_count, total_paras = check_paragraph_length(doc)
    all_issues.extend(para_issues)
    print(f"   Parágrafos analisados: {total_paras}, Curtos: {short_count}")

    # Check 7: Header/footer
    print("[7/8] Verificando header/footer...")
    hf_issues = check_header_footer(doc)
    all_issues.extend(hf_issues)

    # Check 8: Required sections
    print("[8/8] Verificando seções obrigatórias...")
    section_issues = check_required_sections(doc)
    all_issues.extend(section_issues)

    # VERDICT
    print("\n" + "=" * 70)

    s0_count = sum(1 for i in all_issues if "S0 CRITICAL" in i or "CRITICAL" in i)
    s1_count = sum(1 for i in all_issues if "S1 SEVERE" in i)
    s2_count = sum(1 for i in all_issues if "S2 MODERATE" in i)
    s3_count = sum(1 for i in all_issues if "S3 MINOR" in i)

    print(f"\n📋 RESULTADO DA REVISÃO:")
    print(f"   S0 CRITICAL: {s0_count}")
    print(f"   S1 SEVERE:   {s1_count}")
    print(f"   S2 MODERATE: {s2_count}")
    print(f"   S3 MINOR:    {s3_count}")

    if all_issues:
        print(f"\n📝 DETALHAMENTO ({len(all_issues)} issues):")
        for i, issue in enumerate(all_issues, 1):
            print(f"   {i}. {issue}")

    if s0_count > 0 or s1_count > 0:
        verdict = "REPROVADO"
        emoji = "❌"
    elif s2_count > 0:
        verdict = "APROVADO COM RESSALVAS"
        emoji = "⚠️"
    else:
        verdict = "APROVADO"
        emoji = "✅"

    print(f"\n{emoji} VEREDITO: {verdict}")
    print("=" * 70)

    return verdict, all_issues


if __name__ == "__main__":
    main()
