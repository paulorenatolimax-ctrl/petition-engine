#!/usr/bin/env python3
"""
Résumé EB-2 NIW — Daniele de Sousa Manfredi Cedini
Sistema Resume EB-2 NIW V2.0 (DNA Visual V4)
100% Português Brasileiro | Garamond | Navy/Teal palette
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# CONSTANTS — DNA Visual V4
# ============================================================
NAVY = RGBColor(0x2D, 0x3E, 0x50)
TEAL = RGBColor(0x34, 0x98, 0xA2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)

FONT_NAME = "Garamond"
FONT_HEADER_NAME = Pt(20)
FONT_SECTION = Pt(11)
FONT_SUBSECTION = Pt(10)
FONT_BODY = Pt(10.5)
FONT_SMALL = Pt(9.5)
FONT_CONTACT = Pt(9)
FONT_COMPACT = Pt(9)
FONT_TABLE_SMALL = Pt(9.5)

META_WIDTH = 5760
THUMB_WIDTH = 4320
COMPACT_META = 6480
COMPACT_THUMB = 3600

OUTPUT_DIR = "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2025/DANIELLE CEDINI (Refile - EB-2 NIW - DIRETO)/_Forjado por Petition Engine/"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "resume_eb2_niw_DANIELLE_CEDINI.docx")

BENEFICIARY_NAME = "DANIELE DE SOUSA MANFREDI CEDINI"
BENEFICIARY_EMAIL = "danism27@yahoo.com.br"
SOC_CODE = "11-3031"


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def set_cell_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    tcPr.append(tcW)


def set_cell_margins(cell, top=0, bottom=0, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_full_width(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)


def add_run(paragraph, text, font_name=FONT_NAME, size=FONT_BODY, color=BLACK,
            bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rPr.append(parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" '
        f'w:eastAsia="{font_name}" w:cs="{font_name}"/>'
    ))
    return run


def add_body_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       space_after=Pt(6), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = Pt(14)
    add_run(p, text, size=FONT_BODY, color=BLACK)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)
    if level == 0:
        p.paragraph_format.left_indent = Inches(0.3)
    else:
        p.paragraph_format.left_indent = Inches(0.6)

    if bold_prefix:
        add_run(p, f"● {bold_prefix}", size=FONT_BODY if level == 0 else Pt(10),
                color=BLACK, bold=True)
        add_run(p, f" {text}", size=FONT_BODY if level == 0 else Pt(10),
                color=BLACK if level == 0 else DARK_GRAY)
    else:
        sz = FONT_BODY if level == 0 else Pt(10)
        clr = BLACK if level == 0 else DARK_GRAY
        marker = "●" if level == 0 else "○"
        add_run(p, f"{marker} {text}", size=sz, color=clr)
    return p


def add_navy_section_header(doc, title, page_break=True):
    if page_break:
        doc.add_page_break()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "2D3E50")
    set_cell_margins(cell, top=60, bottom=60, left=200, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title.upper(), size=FONT_SECTION, color=WHITE, bold=True)
    set_table_full_width(table)
    remove_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_teal_sub_header(doc, title):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "3498A2")
    set_cell_margins(cell, top=40, bottom=40, left=200, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title, size=FONT_SUBSECTION, color=WHITE, bold=True)
    set_table_full_width(table)
    remove_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_evidence_block(doc, metadata_lines, impact_text, compact=False):
    mw = COMPACT_META if compact else META_WIDTH
    tw = COMPACT_THUMB if compact else THUMB_WIDTH
    impact_sz = FONT_COMPACT if compact else FONT_SMALL

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    set_cell_width(left_cell, mw)
    set_cell_width(right_cell, tw)
    set_cell_borders(left_cell, "CCCCCC", 4)
    set_cell_borders(right_cell, "CCCCCC", 4)
    set_cell_margins(left_cell, top=80, bottom=80, left=120, right=120)
    set_cell_margins(right_cell, top=80, bottom=80, left=80, right=80)

    first = True
    for label, value in metadata_lines:
        if first:
            p = left_cell.paragraphs[0]
            first = False
        else:
            p = left_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        add_run(p, f"{label}: ", size=Pt(10), color=BLACK, bold=True)
        add_run(p, value, size=Pt(10), color=BLACK)

    if impact_text:
        p_label = left_cell.add_paragraph()
        p_label.paragraph_format.space_before = Pt(6)
        p_label.paragraph_format.space_after = Pt(2)
        add_run(p_label, "Descrição e Impacto", size=FONT_SMALL, color=NAVY, bold=True)

        p_impact = left_cell.add_paragraph()
        p_impact.paragraph_format.space_after = Pt(4)
        p_impact.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p_impact, impact_text, size=impact_sz, color=DARK_GRAY, italic=True)

    p_thumb = right_cell.paragraphs[0]
    p_thumb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_run(p_thumb, "[Evidência documental em anexo]", size=FONT_CONTACT, color=MED_GRAY, italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_company_box(doc, company, role, period, location, sector):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")
    set_cell_borders(cell, "CCCCCC", 4)
    set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
    set_table_full_width(table)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_run(p, company, size=FONT_BODY, color=NAVY, bold=True)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    add_run(p2, role, size=FONT_BODY, color=BLACK, bold=True)
    add_run(p2, f"  |  {period}", size=FONT_SMALL, color=DARK_GRAY)

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    add_run(p3, f"{location}  |  {sector}", size=FONT_SMALL, color=MED_GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


# ============================================================
# DOCUMENT SETUP
# ============================================================

def setup_document():
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.header_distance = Inches(0)
    return doc


def build_header(doc):
    header_section = doc.sections[0].header
    header_section.is_linked_to_previous = False

    table = header_section.add_table(rows=3, cols=2, width=Inches(7.2))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)
    remove_table_borders(table)

    # Row 0: Navy — Name
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "2D3E50")
    set_cell_margins(table.cell(0, 0), top=120, bottom=40, left=200, right=100)
    set_cell_margins(table.cell(0, 1), top=120, bottom=40, left=100, right=200)
    p_name = table.cell(0, 0).paragraphs[0]
    add_run(p_name, BENEFICIARY_NAME, size=FONT_HEADER_NAME, color=WHITE, bold=True)

    # Row 1: Navy — RÉSUMÉ + SOC + EB-2 NIW | Email
    for cell in table.rows[1].cells:
        set_cell_shading(cell, "2D3E50")
    set_cell_margins(table.cell(1, 0), top=20, bottom=60, left=200, right=100)
    set_cell_margins(table.cell(1, 1), top=20, bottom=60, left=100, right=200)
    p_resume = table.cell(1, 0).paragraphs[0]
    add_run(p_resume, "RÉSUMÉ", size=FONT_SECTION, color=WHITE, bold=True)
    add_run(p_resume, f"  |  SOC/O*Net {SOC_CODE}  |  EB-2 NIW",
            size=FONT_CONTACT, color=WHITE)
    p_email = table.cell(1, 1).paragraphs[0]
    p_email.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_email, BENEFICIARY_EMAIL, size=FONT_CONTACT, color=WHITE)

    # Row 2: Teal accent
    for cell in table.rows[2].cells:
        set_cell_shading(cell, "3498A2")
    set_row_height(table.rows[2], 4)

    for p in header_section.paragraphs:
        if p.text == "":
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)


def build_footer(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False

    table = footer.add_table(rows=1, cols=1, width=Inches(7.2))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "2D3E50")
    set_cell_margins(cell, top=40, bottom=40, left=200, right=200)
    set_table_full_width(table)
    remove_table_borders(table)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = p.add_run("Page ")
    run1.font.name = FONT_NAME
    run1.font.size = FONT_CONTACT
    run1.font.color.rgb = WHITE

    # PAGE field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r1 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r1.append(fldChar1)
    p._p.append(r1)

    instrText = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr>'
        f'<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
    )
    p._p.append(instrText)

    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r2 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r2.append(fldChar2)
    p._p.append(r2)

    run_of = p.add_run(" of ")
    run_of.font.name = FONT_NAME
    run_of.font.size = FONT_CONTACT
    run_of.font.color.rgb = WHITE

    # NUMPAGES field
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r3 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r3.append(fldChar3)
    p._p.append(r3)

    instrText2 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr>'
        f'<w:instrText xml:space="preserve"> NUMPAGES </w:instrText></w:r>'
    )
    p._p.append(instrText2)

    fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r4 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r4.append(fldChar4)
    p._p.append(r4)

    for p in footer.paragraphs:
        if p.text == "":
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)


# ============================================================
# SECTION 1: SÍNTESE PROFISSIONAL
# ============================================================

def build_sintese(doc):
    add_navy_section_header(doc, "SÍNTESE PROFISSIONAL", page_break=False)

    add_body_paragraph(doc,
        "Daniele de Sousa Manfredi Cedini é uma profissional com mais de 18 anos de experiência "
        "consolidada nas áreas de gestão fiscal, contabilidade tributária e planejamento financeiro "
        "corporativo, com atuação em empresas de médio e grande porte nos setores de importação e "
        "exportação de alimentos, varejo de grande superfície e prestação de serviços técnicos "
        "especializados. Ao longo de sua trajetória profissional, desenvolveu expertise aprofundada "
        "em compliance tributário, implementação de sistemas integrados de gestão empresarial (ERP) "
        "e reestruturação de processos fiscais complexos, atuando diretamente na interface entre a "
        "regulamentação tributária brasileira e as operações de comércio internacional em múltiplos "
        "estados e portos aduaneiros."
    )

    add_body_paragraph(doc,
        "Sua contribuição mais relevante foi a liderança técnica integral na condução do processo "
        "de recuperação judicial do Grupo La Rioja, um conglomerado de quatro empresas atuantes nos "
        "segmentos de importação, exportação e comércio. Na função de Gerente Fiscal e Contábil, "
        "Daniele estruturou e executou o plano de recuperação em colaboração direta com o Perito "
        "Judicial nomeado pela Justiça, tendo o plano sido aprovado sem objeções dos credores — "
        "resultado que atesta tanto a solidez técnica da proposta quanto a capacidade de negociação, "
        "gestão de crise e visão estratégica da peticionária em cenários de extrema complexidade "
        "regulatória e financeira."
    )

    add_body_paragraph(doc,
        "Os resultados quantificáveis de sua gestão incluem redução de 30% nos custos operacionais "
        "mediante reestruturação dos processos fiscais e contábeis, aumento de 30% na produtividade "
        "interna através da implementação de sistemas integrados de controle tributário, e elevação "
        "de 60% nos índices de satisfação de clientes internos e externos. Daniele foi responsável "
        "pela implementação de sistemas tributários completos — abrangendo módulos de importação, "
        "financeiro, contas a pagar e receber, compras, vendas, emissão de notas fiscais, estoque e tributos — "
        "operando em seis plataformas ERP distintas: Procwork, AS400, SAP, Contmatic, JD Oracle e "
        "Mega ERP."
    )

    add_body_paragraph(doc,
        "Bacharela em Ciências Contábeis pelo Centro Universitário Anhanguera de São Paulo, com "
        "diploma avaliado pela GEO Evaluations como equivalente ao bacharelado norte-americano em "
        "Contabilidade, Daniele acumula mais de uma década de experiência progressiva e especializada "
        "que a qualifica como profissional com habilidade excepcional equivalente a grau avançado "
        "nos termos da classificação EB-2 NIW. "
        "Sua formação é complementada por mais de 30 certificações e cursos de aperfeiçoamento "
        "profissional nas áreas de tributação, gestão fiscal, comércio exterior, planejamento "
        "financeiro e gestão empresarial integrada, demonstrando compromisso contínuo com a "
        "atualização técnica e a excelência profissional."
    )

    add_body_paragraph(doc,
        "Em reconhecimento à sua trajetória e expertise, Daniele é membro ativo de quatro "
        "associações profissionais norte-americanas de prestígio: a American Management Association "
        "(AMA), a Association for Financial Professionals (AFP), a Government Finance Officers "
        "Association (GFOA) e o Global Entrepreneurship Institute (GCASE). Essas afiliações "
        "demonstram seu compromisso com os padrões internacionais de excelência em gestão financeira "
        "e empresarial, além de sua integração ativa no ecossistema profissional dos Estados Unidos."
    )

    add_body_paragraph(doc,
        "Ao longo de sua carreira, Daniele desenvolveu conhecimento especializado em regimes "
        "tributários brasileiros — incluindo Lucro Real, Lucro Presumido e Simples Nacional —, "
        "bem como em obrigações acessórias de alta complexidade como SPED Fiscal, SPED Contábil, "
        "ECF, EFD, DCTF, DIRF, GIA e SISCOMEX. Sua atuação em operações de importação e exportação "
        "envolveu cálculos de importação, processos de desembaraço aduaneiro nos principais portos "
        "brasileiros (São Paulo, Rio de Janeiro, Paraná, Santa Catarina e Alagoas), registro no "
        "Radar SISCOMEX e gestão de representantes aduaneiros, conferindo-lhe uma perspectiva "
        "abrangente sobre as dinâmicas do comércio bilateral entre Brasil e Estados Unidos."
    )


# ============================================================
# SECTION 2: HISTÓRICO PROFISSIONAL (GANTT)
# ============================================================

def build_historico(doc):
    add_navy_section_header(doc, "HISTÓRICO PROFISSIONAL")

    # Gantt timeline: years as columns
    years = ["'04", "'05", "'06", "'07", "'08", "'09", "'10", "'11",
             "'12", "'13", "'14", "'15", "'16", "'17", "'18", "'19",
             "'20", "'21", "'22", "'23", "'24", "'25"]
    roles = [
        ("Wel Assessoria Empresarial", 0, 2),      # 2004-2006
        ("Dia Brasil (Carrefour Group)", 2, 4),     # 2006-2008
        ("Mercantil Farmed Ltda", 4, 5),            # 2008-2009
        ("Grupo La Rioja — Cantareira", 5, 17),     # 2009-2021
        ("Desenvolvimento Prof. — EUA", 19, 22),    # 2023-2025
    ]

    n_cols = len(years) + 1
    n_rows = len(roles) + 1
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)

    # Header row: Navy
    for j, yr in enumerate(["Atividade"] + years):
        cell = table.cell(0, j)
        set_cell_shading(cell, "2D3E50")
        set_cell_borders(cell, "2D3E50", 2)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        sz = Pt(7) if j > 0 else Pt(8)
        add_run(p, yr, size=sz, color=WHITE, bold=True)

    # Data rows
    for i, (role_name, start, end) in enumerate(roles):
        row_idx = i + 1
        for j in range(n_cols):
            cell = table.cell(row_idx, j)
            set_cell_borders(cell, "CCCCCC", 2)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

            if j == 0:
                # Role name cell
                add_run(p, role_name, size=Pt(7.5), color=BLACK, bold=False)
                set_cell_width(cell, 2800)
            else:
                yr_idx = j - 1
                if start <= yr_idx < end:
                    set_cell_shading(cell, "3498A2")
                elif row_idx % 2 == 0:
                    set_cell_shading(cell, "F5F5F5")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_body_paragraph(doc,
        "O histórico profissional acima demonstra uma trajetória ascendente e contínua ao longo de "
        "mais de 18 anos, com progressão natural de funções analíticas para posições de gestão "
        "integral, culminando em 12 anos de liderança fiscal e contábil no Grupo La Rioja, onde "
        "Daniele gerenciou simultaneamente quatro empresas do conglomerado em operações que "
        "abrangiam importação, exportação, comércio e indústria alimentícia."
    )


# ============================================================
# SECTION 3: EXPERIÊNCIA PROFISSIONAL DETALHADA
# ============================================================

def build_experiencia(doc):
    add_navy_section_header(doc, "EXPERIÊNCIA PROFISSIONAL DETALHADA")

    # --- La Rioja (principal) ---
    add_company_box(doc,
        "Comercial Importação e Exportação La Rioja Ltda (Grupo Cantareira)",
        "Gerente Fiscal e Contábil",
        "Abril 2009 — Março 2021  (12 anos)",
        "São Paulo, SP, Brasil",
        "Importação, Exportação, Comércio de Alimentos"
    )

    add_bullet(doc, bold_prefix="Gestão Integral do Grupo:",
        text="responsável pela gestão fiscal, contábil, de tesouraria e compliance de quatro "
             "empresas do conglomerado: Comercial La Rioja Ltda, Comercial FEAGRO Importação e "
             "Exportação Eireli, Comercial Porto Vitória Ltda e Imperial Import and Export LTDA, "
             "atuantes nos segmentos de importação de alimentos, construção civil, transporte, "
             "bebidas, exportação de laranja, indústria alimentícia, concessionárias Honda e "
             "importação de pneus e artigos esportivos.")

    add_bullet(doc, bold_prefix="Recuperação Judicial:",
        text="liderou integralmente a estruturação contábil e financeira do plano de recuperação "
             "judicial da La Rioja, em colaboração direta com o Perito Judicial nomeado pelo "
             "Justiça. Elaborou demonstrativos contábeis (balanços patrimoniais, demonstrações de "
             "resultado, fluxos de caixa), negociou diretamente com credores e estruturou as "
             "estratégias operacionais e financeiras que resultaram na aprovação do plano sem "
             "objeções dos credores.")

    add_bullet(doc, bold_prefix="Implementação de Sistemas Tributários:",
        text="implementou sistemas tributários do zero, abrangendo módulos completos de importação, "
             "financeiro, contas a pagar e receber, compras, vendas, emissão de notas fiscais, estoque e "
             "tributos nas plataformas Mega ERP e Contmatic, integrando todas as operações fiscais "
             "do grupo em um único ecossistema digital.")

    add_bullet(doc, bold_prefix="Operações de Importação e Exportação:",
        text="conduziu cálculos de importação, processos de desembaraço aduaneiro em portos "
             "de São Paulo, Rio de Janeiro, Paraná, Santa Catarina e Alagoas, registro no Radar "
             "SISCOMEX, gestão de representantes aduaneiros e registro junto à Agência Nacional de "
             "Vigilância Sanitária, registro de marcas e patentes.")

    add_bullet(doc, bold_prefix="Resultados Mensuráveis:",
        text="redução de 30% nos custos operacionais, aumento de 30% na produtividade interna "
             "e elevação de 60% nos índices de satisfação de clientes. Estabeleceu novo "
             "departamento que impulsionou crescimento de aproximadamente 30% nas operações.")

    add_bullet(doc, bold_prefix="Compliance e Obrigações Acessórias:",
        text="gestão completa das obrigações acessórias do grupo: SPED Fiscal, SPED Contábil, "
             "ECF, EFD, DCTF, DIRF, DIPJ, PERDCOMP, DACON, SINTEGRA, GIA, GIA ST, DES, "
             "contemplando os regimes de Lucro Real, Lucro Presumido e Simples Nacional.")

    add_bullet(doc, bold_prefix="Sistema de Cruzamento de Obrigações:",
        text="implementou sistema de cruzamento de obrigações fiscais e tributárias que reduziu "
             "significativamente o risco de auditoria fiscal e garantiu conformidade integral com "
             "a legislação tributária brasileira em vigor.")

    # --- Mercantil Farmed ---
    add_company_box(doc,
        "Mercantil Farmed Ltda",
        "Analista Fiscal",
        "Outubro 2008 — Abril 2009  (6 meses)",
        "São Paulo, SP, Brasil",
        "Distribuição Farmacêutica"
    )

    add_bullet(doc,
        text="Gestão de todas as obrigações acessórias: DACON, DCTF, SINTEGRA, DIRF, GIA, "
             "DAPI, DES, GIA ST — garantindo conformidade fiscal completa da operação.")

    add_bullet(doc,
        text="Cálculo e conciliação tributária com foco em ICMS Substituição Tributária, "
             "integração com DANFE (NF-e) e manutenção de livros fiscais no ERP JD Oracle.")

    add_bullet(doc,
        text="Apuração de tributos federais e estaduais, conciliação de contas fiscais e "
             "preparação de documentação para auditorias externas.")

    # --- Dia Brasil ---
    add_company_box(doc,
        "Dia Brasil Sociedade Ltda (Grupo Carrefour)",
        "Analista Fiscal Sênior",
        "Julho 2006 — Setembro 2008  (2 anos e 2 meses)",
        "São Paulo, SP, Brasil",
        "Varejo de Grande Superfície"
    )

    add_bullet(doc,
        text="Classificação de notas fiscais e cálculo tributário completo: DES, NFS, IRRF, "
             "INSS, CSLL, PIS, Cofins, IRPJ, com atuação integrada nos ERPs Procwork, SAP e "
             "AS400 do Grupo Carrefour.")

    add_bullet(doc,
        text="Cálculo diário de ICMS Substituição Tributária, lacração de impressoras fiscais, "
             "serviços de importação e preparação de declarações fiscais: DCTF, DACON, DIRF, "
             "DIPJ e GIA.")

    add_bullet(doc,
        text="Suporte técnico fiscal a mais de 400 lojas da rede Dia Brasil em território "
             "nacional, garantindo uniformidade no tratamento tributário em múltiplas jurisdições "
             "estaduais.")

    # --- Wel Assessoria ---
    add_company_box(doc,
        "Wel Assessoria Empresarial LTDA",
        "Supervisora do Departamento Fiscal",
        "Junho 2004 — Julho 2006  (2 anos e 1 mês)",
        "São Paulo, SP, Brasil",
        "Serviços Técnicos Contábeis"
    )

    add_bullet(doc,
        text="Supervisão do departamento fiscal responsável por aproximadamente 120 clientes "
             "entre microempresas e empresas de pequeno porte, incluindo empresas ativas e "
             "inativas em diferentes regimes tributários.")

    add_bullet(doc,
        text="Cálculo e recolhimento de tributos: DARF, GPS, GARE e ISS, com interação "
             "direta com clientes para orientação tributária e atualização sobre mudanças na "
             "legislação fiscal vigente.")

    add_bullet(doc,
        text="Operação do ERP Contmatic para escrituração fiscal, gerando experiência formativa "
             "que seria ampliada nos anos seguintes em plataformas de maior complexidade como SAP, "
             "Oracle e Mega ERP.")


# ============================================================
# SECTION 4: CONTRIBUIÇÕES TÉCNICAS E PROFISSIONAIS
# ============================================================

def build_contribuicoes(doc):
    add_navy_section_header(doc, "CONTRIBUIÇÕES TÉCNICAS E PROFISSIONAIS")

    # --- Tema 1: Reestruturação Corporativa ---
    add_teal_sub_header(doc, "Reestruturação Corporativa e Recuperação Judicial")

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Liderança Técnica em Recuperação Judicial"),
            ("Entidade", "Grupo La Rioja — Cantareira (4 empresas)"),
            ("Período", "2013 — 2018"),
            ("Perito Judicial", "José Vanderlei Masson dos Santos, CRC/SP 124747"),
        ],
        impact_text=(
            "Daniele liderou integralmente a estruturação contábil e financeira do plano de "
            "recuperação judicial do Grupo La Rioja, conglomerado que enfrentava crise de "
            "liquidez severa. Em colaboração direta com o Perito Judicial nomeado pela Justiça, "
            "José Vanderlei Masson dos Santos — presidente da APEJESP e vice-presidente da "
            "FEBRAPAM —, elaborou balanços patrimoniais, demonstrações de resultado e projeções "
            "de projeções operacionais que fundamentaram o plano de recuperação. O plano foi aprovado "
            "sem objeções dos credores, resultado excepcional que evidencia tanto a solidez "
            "técnica da proposta quanto a capacidade de negociação e gestão de crise. A "
            "reestruturação restaurou liquidez, estabilidade operacional e sustentabilidade "
            "financeira do grupo, preservando empregos e relações comerciais."
        )
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Estratégia de Reestruturação Operacional"),
            ("Entidade", "Grupo La Rioja — Todas as unidades"),
            ("Resultado", "Redução de 30% nos custos operacionais"),
        ],
        impact_text=(
            "Após a aprovação do plano de recuperação judicial, Daniele implementou programa "
            "de reestruturação operacional que envolveu a revisão integral dos processos fiscais "
            "e contábeis das quatro empresas do grupo. A iniciativa resultou em redução de 30% "
            "nos custos operacionais através da eliminação de redundâncias processuais, "
            "renegociação de contratos com fornecedores de serviços contábeis e automação de "
            "rotinas tributárias que anteriormente demandavam processamento manual. O novo modelo "
            "operacional tornou-se referência interna para as demais unidades do conglomerado."
        )
    )

    # --- Tema 2: Implementação de Sistemas Tributários ---
    add_teal_sub_header(doc, "Implementação de Sistemas Tributários e Compliance Fiscal")

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Implementação de ERP Tributário Integrado"),
            ("Plataformas", "Mega ERP, Contmatic, SAP, Procwork, JD Oracle, AS400"),
            ("Módulos", "Importação, Financeiro, Contas a Pagar/Receber, Compras, Vendas, "
                        "Emissão de NF, Estoque, Tributos"),
        ],
        impact_text=(
            "Daniele implementou sistemas tributários completos do zero em múltiplas plataformas "
            "ERP ao longo de sua carreira, abrangendo a totalidade dos módulos operacionais: "
            "importação, financeiro, contas a pagar e receber, compras, vendas, emissão de notas fiscais, "
            "estoque e tributos. Essa experiência transversal em seis plataformas distintas — "
            "Procwork e SAP no Grupo Carrefour, JD Oracle na Mercantil Farmed, Contmatic na Wel "
            "Assessoria e Mega ERP na La Rioja — confere à peticionária domínio técnico raro "
            "sobre a arquitetura de sistemas fiscais integrados, capacitando-a a diagnosticar "
            "ineficiências operacionais e propor soluções de integração em ambientes corporativos "
            "de alta complexidade."
        )
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Sistema de Cruzamento de Obrigações Fiscais"),
            ("Entidade", "Grupo La Rioja — Cantareira"),
            ("Resultado", "Redução significativa do risco de auditoria fiscal"),
        ],
        impact_text=(
            "Desenvolveu e implementou sistema de cruzamento de obrigações fiscais e tributárias "
            "que permitiu a verificação automatizada da consistência entre declarações federais, "
            "estaduais e municipais. O sistema confrontava dados de SPED Fiscal, SPED Contábil, "
            "DCTF, DIRF e GIA, identificando divergências antes da transmissão às autoridades "
            "fiscais. A iniciativa reduziu significativamente o risco de autuações e otimizou "
            "o processo de resposta a fiscalizações, sendo adotada como padrão operacional por "
            "todas as quatro empresas do grupo."
        )
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Gestão de Obrigações Acessórias Complexas"),
            ("Abrangência", "SPED Fiscal, SPED Contábil, ECF, EFD, DCTF, DIRF, GIA, "
                            "SISCOMEX, LALUR, FCONT, EFD REINF"),
            ("Regimes", "Lucro Real, Lucro Presumido, Simples Nacional"),
        ],
        impact_text=(
            "Ao longo de mais de uma década de gestão tributária, Daniele acumulou domínio "
            "operacional completo sobre o ecossistema de obrigações acessórias brasileiras, "
            "incluindo escriturações digitais de última geração (SPED Fiscal, SPED Contábil, ECF, "
            "EFD, EFD REINF), declarações federais (DCTF, DIRF, DIPJ, PERDCOMP, DACON) e "
            "estaduais (GIA, GIA ST, SINTEGRA, DES). Essa expertise abrangia os três regimes "
            "tributários vigentes e múltiplas jurisdições estaduais, capacitando-a a oferecer "
            "orientação técnica precisa em cenários de crescente complexidade regulatória."
        )
    )

    # --- Tema 3: Comércio Exterior e Operações de Importação ---
    add_teal_sub_header(doc, "Comércio Exterior e Operações de Importação")

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Gestão de Operações de Importação e Exportação"),
            ("Entidade", "Grupo La Rioja — 4 empresas"),
            ("Portos", "São Paulo, Rio de Janeiro, Paraná, Santa Catarina, Alagoas"),
            ("Sistemas", "SISCOMEX, Radar, ANVISA, INPI"),
        ],
        impact_text=(
            "Daniele conduziu integralmente as operações de comércio exterior do Grupo La Rioja, "
            "abrangendo cálculos de importação, processos de desembaraço aduaneiro em cinco "
            "estados brasileiros, registro e manutenção do Radar SISCOMEX, cadastramento de "
            "representantes aduaneiros e registro junto à Agência Nacional de Vigilância "
            "Sanitária. Gerenciou operações de importação de alimentos, bebidas, pneus e "
            "artigos esportivos, além de exportações de laranja e derivados, acumulando "
            "experiência prática em regulamentação sanitária, classificação tarifária e "
            "logística portuária internacional. Essa vivência em comércio bilateral conferiu "
            "à peticionária perspectiva técnica e operacional sobre as dinâmicas de intercâmbio "
            "comercial entre Brasil e Estados Unidos."
        )
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Registro de Marcas e Propriedade Industrial"),
            ("Entidade", "Grupo La Rioja"),
            ("Órgão", "INPI — Instituto Nacional da Propriedade Industrial"),
        ],
        impact_text=(
            "Conduziu processos de registro de marcas e patentes junto ao Instituto Nacional "
            "da Propriedade Industrial (INPI) para as empresas do Grupo La Rioja, gerenciando "
            "o portfólio de propriedade intelectual do conglomerado. Essa atuação demonstra "
            "capacidade de gestão de ativos intangíveis e compreensão das dinâmicas de proteção "
            "de marca em operações comerciais internacionais, competência que se estende à "
            "assessoria sobre registro de marcas no mercado norte-americano."
        ),
        compact=True
    )

    # --- Tema 4: Otimização de Processos ---
    add_teal_sub_header(doc, "Otimização de Processos e Eficiência Operacional")

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Programa de Otimização Operacional"),
            ("Entidade", "Grupo La Rioja — Cantareira"),
            ("Resultados", "30% produtividade ↑ | 60% satisfação ↑ | 30% crescimento"),
        ],
        impact_text=(
            "Daniele concebeu e implementou programa abrangente de otimização dos processos "
            "fiscais, contábeis e operacionais do Grupo La Rioja, alcançando aumento de 30% "
            "na produtividade interna e elevação de 60% nos índices de satisfação de clientes "
            "internos e externos. O programa incluiu a criação de um departamento inteiramente "
            "novo dedicado à gestão integrada dos processos fiscais das quatro empresas, cuja "
            "implementação impulsionou crescimento de aproximadamente 30% nas operações do "
            "grupo. As métricas foram aferidas por meio de indicadores internos de desempenho "
            "implementados pela própria peticionária, garantindo mensuração objetiva dos "
            "resultados alcançados."
        )
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Abertura de Filiais em Território Nacional"),
            ("Entidade", "Grupo La Rioja"),
            ("Abrangência", "Múltiplos estados brasileiros"),
        ],
        impact_text=(
            "Responsável técnica pela abertura de novas filiais do Grupo La Rioja em múltiplos "
            "estados brasileiros, liderando todos os aspectos fiscais e regulatórios do processo: "
            "inscrição estadual, cadastro municipal, habilitação no SISCOMEX, registro na "
            "Vigilância Sanitária e adequação tributária à legislação de cada jurisdição. Cada "
            "abertura de filial exigiu análise detalhada dos incentivos fiscais estaduais "
            "disponíveis, configuração de regimes tributários específicos e implementação dos "
            "módulos de ERP correspondentes, demonstrando capacidade de execução em ambientes "
            "regulatórios multijurisdicionais."
        ),
        compact=True
    )


# ============================================================
# SECTION 5: PUBLICAÇÕES, PALESTRAS E EVENTOS
# ============================================================

def build_publicacoes(doc):
    add_navy_section_header(doc, "PUBLICAÇÕES, PALESTRAS E EVENTOS PROFISSIONAIS")

    add_teal_sub_header(doc, "Palestras e Participações em Eventos")

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Palestra em Encontro de Contabilistas"),
            ("Evento", "Encontro de Contabilistas — Juventus"),
            ("Ano", "2018"),
            ("Tema", "Atualizações tributárias e impacto nas empresas"),
        ],
        impact_text=(
            "Daniele apresentou palestra sobre as atualizações tributárias mais recentes e seus "
            "impactos nas operações empresariais, com ênfase em modelos digitais, automação, "
            "inteligência artificial e aprendizado de máquina aplicados à análise de dados "
            "e transações fiscais. A apresentação abordou tendências de digitalização fiscal "
            "que se materializaram nos anos seguintes com a expansão do SPED e a implementação "
            "da Nota Fiscal Eletrônica em âmbito nacional."
        )
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Convite Especial — Encontro de Especialistas Tributários"),
            ("Evento", "Business Talk — Lançamento de Livro"),
            ("Data", "29 de setembro de 2016"),
            ("Natureza", "Convite especial como especialista reconhecida no setor"),
        ],
        impact_text=(
            "Recebeu convite especial para participar de evento de lançamento de livro que "
            "reuniu especialistas tributários e econômicos de destaque no cenário paulistano. "
            "O convite evidencia o reconhecimento da comunidade profissional pela contribuição "
            "técnica de Daniele ao campo da gestão fiscal e tributária, posicionando-a como "
            "referência entre seus pares em um segmento altamente especializado."
        ),
        compact=True
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Apresentação em Convenção de Contabilistas"),
            ("Evento", "Convenção de Contabilistas"),
            ("Ano", "2013"),
        ],
        impact_text=(
            "Apresentação anual dos indicadores contábeis no contexto da convenção profissional "
            "de 2013, período que coincidiu com o início dos trabalhos de recuperação judicial "
            "do Grupo La Rioja. A participação ativa em fóruns profissionais durante esse período "
            "demonstra a capacidade de manter engajamento técnico e contribuição ao campo mesmo "
            "em cenários de elevada demanda operacional."
        ),
        compact=True
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Entrevista — Programa Mulheres de Sucesso"),
            ("Veículo", "SBT — Sistema Brasileiro de Televisão"),
            ("Ano", "2011"),
            ("Seleção", "Representante dos Contabilistas de São Paulo"),
        ],
        impact_text=(
            "Selecionada para representar os contabilistas de São Paulo em programa piloto "
            "da emissora SBT dedicado a destacar mulheres que se sobressaíram em suas "
            "respectivas profissões. A seleção reflete reconhecimento público da trajetória "
            "profissional de Daniele em um campo historicamente de predominância masculina, "
            "evidenciando sua posição de destaque no cenário contábil e fiscal paulistano."
        ),
        compact=True
    )

    add_teal_sub_header(doc, "Publicações e Produção Técnica")

    publications = [
        ("Avaliação da Gestão na Propriedade Rural", "2023",
         "Análise abrangente dos processos de gestão aplicados à propriedade rural, com foco "
         "em indicadores de desempenho, controle financeiro e sustentabilidade operacional."),
        ("Gestão Financeira", "2023",
         "Estudo técnico sobre princípios e práticas de gestão financeira corporativa, abordando "
         "planejamento orçamentário, análise de investimentos e controle de custos."),
        ("Como Elaborar um Plano de Negócio", "2023",
         "Metodologia estruturada para elaboração de planos de negócios, contemplando análise de "
         "mercado, projeções financeiras, estruturação organizacional e estratégias de crescimento."),
        ("Responsabilidade Corporativa e Comercial", "2023",
         "Produção técnica sobre responsabilidade corporativa e social no contexto empresarial "
         "contemporâneo, realizada em parceria com a Cisco Networking Academy."),
        ("Como Tornar Sua Empresa uma Franquia", "2023",
         "Análise técnica sobre o processo de franqueamento empresarial, abrangendo aspectos "
         "regulatórios, operacionais e de expansão de marca."),
        ("Gestão Empresarial Integrada", "2023",
         "Estudo técnico sobre integração de processos empresariais com foco em eficiência "
         "operacional, gestão por indicadores e sistemas integrados de informação."),
    ]

    for title, year, desc in publications:
        add_evidence_block(doc,
            metadata_lines=[
                ("Título", title),
                ("Ano", year),
                ("Idioma", "Português (com tradução certificada para inglês)"),
            ],
            impact_text=desc,
            compact=True
        )


# ============================================================
# SECTION 6: FORMAÇÃO ACADÊMICA
# ============================================================

def build_formacao(doc):
    add_navy_section_header(doc, "FORMAÇÃO ACADÊMICA")

    add_evidence_block(doc,
        metadata_lines=[
            ("Grau", "Bacharelado em Ciências Contábeis"),
            ("Instituição", "Centro Universitário Anhanguera de São Paulo"),
            ("Conclusão", "27 de janeiro de 2014"),
            ("Carga Horária", "3.380 horas-aula"),
            ("Diploma Expedido", "11 de março de 2021"),
        ],
        impact_text=(
            "Formação acadêmica em Ciências Contábeis com carga horária de 3.380 horas-aula, "
            "abrangendo disciplinas de contabilidade geral e avançada, direito tributário, "
            "auditoria, controladoria, análise de custos, matemática financeira e legislação "
            "comercial. O bacharelado, complementado por mais de uma década de experiência "
            "progressiva em posições de crescente responsabilidade, qualifica a peticionária "
            "como profissional com habilidade excepcional equivalente a grau avançado, conforme "
            "reconhecido pela avaliação de credenciais internacional."
        )
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Avaliação de Credencial Internacional"),
            ("Agência", "GEO Evaluations — Global Education Office"),
            ("Equivalência", "Bacharelado norte-americano em Contabilidade (Accounting)"),
            ("País de Origem", "Brasil"),
        ],
        impact_text=(
            "Avaliação oficial emitida pela GEO Evaluations atesta que o diploma brasileiro em "
            "Ciências Contábeis de Daniele é equivalente ao bacharelado (Bachelor's degree) em "
            "Contabilidade nos Estados Unidos, confirmando a compatibilidade de sua formação "
            "acadêmica com os padrões educacionais norte-americanos e habilitando-a ao exercício "
            "profissional no mercado financeiro e contábil dos Estados Unidos."
        )
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Tipo", "Ensino de Idiomas"),
            ("Instituição", "Liberty Language"),
            ("Status", "Estudante ativa"),
            ("Idioma", "Inglês"),
        ],
        impact_text=(
            "Investimento contínuo no aperfeiçoamento do idioma inglês através de programa "
            "de estudos na Liberty Language, demonstrando comprometimento com a integração "
            "linguística e cultural nos Estados Unidos e capacidade de comunicação efetiva no "
            "ambiente empresarial norte-americano."
        ),
        compact=True
    )


# ============================================================
# SECTION 7: CURSOS E CERTIFICAÇÕES PROFISSIONAIS
# ============================================================

def build_cursos(doc):
    add_navy_section_header(doc, "CURSOS E CERTIFICAÇÕES PROFISSIONAIS")

    # --- Certificações Internacionais ---
    add_teal_sub_header(doc, "Certificações e Afiliações Profissionais Internacionais")

    certifications = [
        ("American Management Association (AMA)", "Membro desde setembro de 2023",
         "ID: 4052642",
         "Afiliação à principal associação norte-americana de gestão empresarial, com acesso a "
         "programas de desenvolvimento em liderança, estratégia e gestão financeira."),
        ("Association for Financial Professionals (AFP)", "Membro desde janeiro de 2024",
         "ID: 808711",
         "Afiliação à associação que representa mais de 16.000 profissionais financeiros nos "
         "Estados Unidos, focada em gestão de tesouraria, planejamento financeiro e análise de "
         "riscos corporativos."),
        ("Government Finance Officers Association (GFOA)", "Membro desde janeiro de 2024",
         "ID: 300281927",
         "Afiliação à associação que promove excelência em gestão financeira governamental, "
         "com enfoque em contabilidade pública, orçamento e compliance regulatório."),
        ("Global Entrepreneurship Institute (GCASE)", "Membro desde setembro de 2023",
         "Registro ativo",
         "Afiliação ao instituto internacional dedicado ao fomento do empreendedorismo, inovação "
         "e desenvolvimento de novos negócios com alcance global."),
    ]

    for name, date, id_info, desc in certifications:
        add_evidence_block(doc,
            metadata_lines=[
                ("Certificação/Afiliação", name),
                ("Data", date),
                ("Identificação", id_info),
            ],
            impact_text=desc,
            compact=True
        )

    add_teal_sub_header(doc, "Certificações Técnicas")

    tech_certs = [
        ("International Business Finance Certification Practice Test",
         "Brainmeasures — Janeiro de 2024",
         "Aprovada com pontuação de 75/100 em avaliação internacional de finanças empresariais, "
         "abrangendo análise financeira internacional, gestão de câmbio e planejamento tributário "
         "transfronteiriço."),
        ("Finance and Insurance Test",
         "Brainmeasures — Janeiro de 2024",
         "Aprovada com 30 de 40 questões corretas (75%) em avaliação abrangente de finanças e "
         "seguros, cobrindo instrumentos financeiros, análise de risco e fundamentos de seguros "
         "corporativos."),
        ("Get Connected — Cisco Networking Academy",
         "Cisco — Setembro de 2023 — 24 horas",
         "Certificação em tecnologia da informação e conectividade digital pela Cisco Networking "
         "Academy, ampliando competências em infraestrutura tecnológica aplicada à gestão "
         "empresarial."),
        ("Corporate and Social Responsibility",
         "Cisco Networking Academy — 2023",
         "Certificação em responsabilidade corporativa e social, abordando sustentabilidade "
         "empresarial, governança e impacto social no ambiente de negócios contemporâneo."),
    ]

    for name, date, desc in tech_certs:
        add_evidence_block(doc,
            metadata_lines=[
                ("Curso", name),
                ("Instituição/Data", date),
            ],
            impact_text=desc,
            compact=True
        )

    # --- Cursos Técnicos em Gestão ---
    add_teal_sub_header(doc, "Cursos Técnicos em Gestão Empresarial")

    mgmt_courses = [
        ("Como Elaborar um Plano de Negócio", "SEBRAE — Outubro 2023 — 2h"),
        ("Avaliação da Gestão na Propriedade Rural", "SEBRAE — Setembro 2023 — 8h"),
        ("Gestão Financeira", "SEBRAE — Setembro 2023 — 3h"),
        ("Como Tornar Sua Empresa uma Franquia", "SEBRAE — Setembro 2023 — 5h"),
        ("Gestão Empresarial Integrada", "SEBRAE — Setembro 2023 — 15h"),
    ]

    for name, detail in mgmt_courses:
        add_bullet(doc, bold_prefix=f"{name} —", text=detail)

    # --- Cursos Técnicos em Tributação ---
    add_teal_sub_header(doc, "Cursos Técnicos em Tributação e Contabilidade")

    tax_courses = [
        ("DIPJ Digital — ECF (Regras para Geração de Arquivo Digital)",
         "CENOFISCO — Junho 2015 — 8h"),
        ("Preparação de Assistente Contábil", "Econet — Setembro 2014 — 32h"),
        ("DIRF 2014 — Regras de Retenção e Preenchimento", "Econet — Janeiro 2014 — 8h"),
        ("Cálculo do Lucro Real para 2014", "Econet — Janeiro 2014 — 8h"),
        ("SPED Fiscal", "IOB — Módulo completo"),
        ("SPED Contábil", "IOB — Módulo completo"),
        ("SPED Contribuições PIS/COFINS", "IOB — Módulo completo"),
        ("ECF — Escrituração Contábil Fiscal", "Aduaneiras — Módulo completo"),
        ("FCONT/ELALUR/RTT", "IOB — Módulo completo"),
        ("Análise e Aspectos Fiscais", "IOB — Módulo completo"),
        ("Retenções de Serviços", "IOB — Módulo completo"),
        ("Contabilidade Geral", "IOB — Módulo completo"),
        ("Importações", "Aduaneiras — Módulo completo"),
        ("Mercadoria Importada Alíquota 4% Interestadual", "Econet — Módulo completo"),
        ("FCI — Ficha de Conteúdo de Importação", "Econet — Módulo completo"),
        ("Substituição Tributária", "Econet — Módulo completo"),
        ("MP 627 — Impacto Lucro Real e Presumido", "Econet — Módulo completo"),
        ("Analista Contábil", "Econet — Módulo completo"),
        ("ICMS — Aspectos Fiscais", "CEAD-Contimatic — Módulo completo"),
        ("SINTEGRA", "CEAD-Contimatic — Módulo completo"),
        ("Escrita Fiscal", "CEAD-Contimatic — Módulo completo"),
        ("Emissão de Nota Fiscal", "CEAD-Contimatic — Módulo completo"),
        ("ICMS ST com Mercadorias", "PRODEP — 2010"),
        ("Sistema de Desenvolvimento NF-e", "Serasa — Módulo completo"),
        ("Retenções de IR na Fonte", "Verba Net — Módulo completo"),
        ("EFD REINF", "Sindicato dos Contabilistas — Módulo completo"),
        ("Auditoria e Compliance Fiscal", "CENOFISCO — Módulo completo"),
        ("Excel Avançado", "Impacta Tecnologia — Módulo completo"),
        ("Manutenção de Computadores", "SENAC — Módulo completo"),
    ]

    for name, detail in tax_courses:
        add_bullet(doc, bold_prefix=f"{name} —", text=detail)

    add_body_paragraph(doc,
        "O extenso portfólio de mais de 30 cursos e certificações técnicas reflete o compromisso "
        "contínuo de Daniele com a atualização profissional em um campo caracterizado por mudanças "
        "regulatórias frequentes, demonstrando capacidade adaptativa e investimento sistemático em "
        "desenvolvimento técnico ao longo de duas décadas de carreira.",
        space_before=Pt(8)
    )


# ============================================================
# SECTION 8: CARTAS DE RECOMENDAÇÃO
# ============================================================

def build_cartas(doc):
    add_navy_section_header(doc, "CARTAS DE RECOMENDAÇÃO")

    add_body_paragraph(doc,
        "As cartas de recomendação a seguir foram emitidas por profissionais que atuaram "
        "diretamente com a peticionária em diferentes capacidades ao longo de sua carreira. "
        "Os signatários ocupam posições de destaque em seus respectivos campos de atuação e "
        "oferecem testemunho direto das contribuições técnicas e profissionais de Daniele."
    )

    # Recommendation letters table
    headers = ["Signatário", "Cargo / Instituição", "Relação Profissional", "Data"]
    data = [
        ["José Vanderlei Masson dos Santos",
         "Contador, CRC/SP 124747\nPresidente APEJESP (2014-15)\nVice-Presidente FEBRAPAM",
         "Perito Judicial no processo de\nrecuperação judicial da La Rioja.\nColaboração desde maio de 2013.",
         "Nov/2023"],
        ["Paulo Reszka",
         "Inspetor Estadual — Secretaria\nda Fazenda do Estado de São Paulo\nBacharel em Economia (PUC-SP)",
         "Parceria profissional desde 2010\nna La Rioja. Serviços para a\nSecretaria da Fazenda de SP.",
         "Nov/2023"],
        ["Gustavo Vieira Ribeiro",
         "Franco Montoro e Peixoto\nAdvogados Associados\nDesde 2003",
         "Colaboração profissional na\nLa Rioja desde 2009.\nPlanejamento tributário conjunto.",
         "Set/2023"],
        ["Luiz Carlos Perez",
         "Empresário e Diretor\nAdministrativo — Grupo La Rioja\nCom. Fegaro Import. Export. LTDA",
         "Empregador direto e colaborador\nnas quatro empresas do grupo.\nRelação de trabalho de 12 anos.",
         "Dez/2023"],
        ["Solon Andrade",
         "Bacharel em Direito (UNICSUL)\nSócio-Gerente — L&S Contabilidade\nEmpresarial Ltda",
         "Colaboração profissional na\nLa Rioja e após a saída de\nDaniele em março de 2021.",
         "Dez/2023"],
    ]

    n_rows = len(data) + 1
    n_cols = len(headers)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)

    # Header row
    for j, header_text in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, "2D3E50")
        set_cell_borders(cell, "2D3E50", 2)
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header_text, size=FONT_SMALL, color=WHITE, bold=True)

    # Data rows
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i + 1, j)
            set_cell_borders(cell, "CCCCCC", 2)
            set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
            if (i + 1) % 2 == 0:
                set_cell_shading(cell, "F5F5F5")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            sz = Pt(8.5)
            add_run(p, cell_text, size=sz, color=BLACK if j == 0 else DARK_GRAY,
                    bold=(j == 0))

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Brief synthesis of each letter
    add_teal_sub_header(doc, "Síntese das Recomendações")

    add_body_paragraph(doc,
        "José Vanderlei Masson dos Santos — Perito Judicial — atesta que a atuação de Daniele "
        "na gestão da recuperação judicial da La Rioja constitui exemplo emblemático de competência "
        "e resiliência no mundo empresarial. Confirma que, juntos, estruturaram as informações "
        "contábeis que fundamentaram o plano de reorganização aprovado sem objeções dos credores, "
        "e que a peticionária introduziu práticas de gestão mais eficientes e sustentáveis que "
        "asseguraram tanto a sobrevivência imediata quanto a viabilidade de longo prazo do grupo."
    )

    add_body_paragraph(doc,
        "Paulo Reszka — Inspetor Estadual da Fazenda de São Paulo — destaca que Daniele foi "
        "instrumental na implementação de estratégias de compliance tributário que resultaram em "
        "melhorias significativas na eficiência operacional e na acurácia dos processos fiscais. "
        "Ressalta sua capacidade excepcional de processar solicitações de isenção e imunidade "
        "tributária, e sua visão estratégica em orientar empresas brasileiras no processo de "
        "internacionalização."
    )

    add_body_paragraph(doc,
        "Gustavo Vieira Ribeiro — Associado sênior em escritório de serviços especializados — "
        "afirma que a proficiência de Daniele em planejamento tributário e estruturação corporativa "
        "foi notável, tendo suas contribuições sido fundamentais na redução de obrigações "
        "tributárias e custos operacionais. Expressa intenção de estabelecer vínculo de "
        "colaboração comercial com o empreendimento futuro de Daniele nos Estados Unidos."
    )

    add_body_paragraph(doc,
        "Luiz Carlos Perez — Diretor Administrativo do Grupo La Rioja — confirma o "
        "conhecimento aprofundado e a experiência prática de Daniele em gestão tributária e "
        "contábil, destacando sua excelência como coordenadora, comunicadora e implementadora "
        "de sistemas de controle tributário e de estoque. Atesta que seu empreendimento "
        "contribuirá para geração de empregos, inovação e crescimento econômico bilateral."
    )

    add_body_paragraph(doc,
        "Solon Andrade — Sócio-gerente de escritório de contabilidade — ressalta as "
        "habilidades excepcionais de Daniele na reformulação e otimização de processos gerenciais, "
        "fiscais e contábeis, destacando que trouxe clareza aos sistemas de tributação e "
        "controles de estoque PEPS, contribuições cruciais durante a adaptação a novas regulações "
        "tributárias. Confirma que, após a saída de Daniele em março de 2021, manteve colaboração "
        "estendendo sua expertise a outras empresas."
    )


# ============================================================
# SECTION 9: PERFIL PROFISSIONAL — O*NET
# ============================================================

def build_perfil_onet(doc):
    add_navy_section_header(doc, "PERFIL PROFISSIONAL E APTIDÃO OCUPACIONAL")

    add_body_paragraph(doc,
        "Os resultados de avaliações profissionais padronizadas confirmam o alinhamento entre o "
        "perfil de competências de Daniele e as exigências da classificação ocupacional de "
        "Financial Manager (SOC 11-3031), reforçando sua aptidão para atuação nos setores de "
        "gestão financeira e empresarial nos Estados Unidos."
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Avaliação", "O*NET Interest Profiler"),
            ("Fonte", "Departamento do Trabalho dos Estados Unidos — O*NET OnLine"),
            ("Perfil Dominante", "Social (37) | Enterprising (35) | Conventional (34)"),
        ],
        impact_text=(
            "O perfil ocupacional de Daniele, conforme apurado pelo O*NET Interest Profiler do "
            "Departamento do Trabalho dos Estados Unidos, revela pontuações mais altas nas "
            "dimensões Social (37 pontos), Enterprising (35 pontos) e Conventional (34 pontos), "
            "seguidas por Investigative (31 pontos). Esse padrão é altamente consistente com as "
            "competências exigidas para a classificação SOC 11-3031 (Financial Manager), que "
            "demanda capacidade de liderança de equipes (Social), tomada de decisão empresarial "
            "(Enterprising) e rigor analítico em processos normatizados (Conventional). As "
            "pontuações nas dimensões Realistic (22) e Artistic (22) completam um perfil "
            "equilibrado orientado à gestão."
        )
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Avaliação", "123Test Career Aptitude Test"),
            ("Tipo de Personalidade", "SCEIRA"),
            ("Perfil Dominante", "Social (23%) | Entrepreneurial (18%) | Investigative (18%)"),
        ],
        impact_text=(
            "A avaliação de aptidão profissional conduzida pela 123Test confirma o perfil Social "
            "como dimensão predominante (23%), seguido pelas dimensões Entrepreneurial e "
            "Investigative (18% cada). O tipo de personalidade SCEIRA indica profissional com "
            "forte orientação para interação interpessoal, empreendedorismo e capacidade "
            "investigativa — combinação particularmente adequada para gestão financeira corporativa "
            "e assessoria empresarial especializada."
        ),
        compact=True
    )


# ============================================================
# MAIN — BUILD DOCUMENT
# ============================================================

def main():
    print("=" * 60)
    print("RESUME EB-2 NIW — Daniele de Sousa Manfredi Cedini")
    print("Sistema Resume EB-2 NIW V2.0 (DNA Visual V4)")
    print("=" * 60)

    # Setup
    print("\n[1/11] Configurando documento...")
    doc = setup_document()

    print("[2/11] Construindo header...")
    build_header(doc)

    print("[3/11] Construindo footer...")
    build_footer(doc)

    # Content sections
    print("[4/11] Seção: Síntese Profissional...")
    build_sintese(doc)

    print("[5/11] Seção: Histórico Profissional (Gantt)...")
    build_historico(doc)

    print("[6/11] Seção: Experiência Profissional...")
    build_experiencia(doc)

    print("[7/11] Seção: Contribuições Técnicas...")
    build_contribuicoes(doc)

    print("[8/11] Seção: Publicações e Eventos...")
    build_publicacoes(doc)

    print("[9/11] Seção: Formação Acadêmica...")
    build_formacao(doc)

    print("[10/11] Seção: Cursos e Certificações...")
    build_cursos(doc)

    print("[11/11] Seção: Cartas de Recomendação + Perfil O*NET...")
    build_cartas(doc)
    build_perfil_onet(doc)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)

    file_size = os.path.getsize(OUTPUT_FILE)
    print(f"\n{'=' * 60}")
    print(f"DOCUMENTO GERADO COM SUCESSO")
    print(f"Arquivo: {OUTPUT_FILE}")
    print(f"Tamanho: {file_size:,} bytes")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
