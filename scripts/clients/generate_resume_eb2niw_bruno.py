#!/usr/bin/env python3
"""
Résumé EB-2 NIW — Bruno da Silva Ucella
Gerado pelo Petition Engine v2.0
100% Português Brasileiro | Garamond | Navy/Teal palette
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# CONSTANTS
# ============================================================
NAVY = RGBColor(0x2D, 0x3E, 0x50)
TEAL = RGBColor(0x34, 0x98, 0xA2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)

FONT_NAME = "Garamond"
FONT_HEADER_NAME = Pt(20)
FONT_SECTION = Pt(11)
FONT_SUBSECTION = Pt(10)
FONT_BODY = Pt(10.5)
FONT_SMALL = Pt(9.5)
FONT_CONTACT = Pt(9)
FONT_COMPACT = Pt(9)
FONT_TABLE_SMALL = Pt(8)

META_WIDTH = 5760
THUMB_WIDTH = 4320
COMPACT_META = 6480
COMPACT_THUMB = 3600

OUTPUT_DIR = (
    "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/"
    "_2. MEUS CASOS/2026/Bruno da Silva Ucella (EB-2 NIW)/"
    "_Forjado por Petition Engine/"
)
OUTPUT_FILE = "resume_eb2_niw_Bruno_da_Silva_Ucella.docx"

BENEFICIARY_NAME = "BRUNO DA SILVA UCELLA"
SOC_CODE = "11-9041"


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def set_cell_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    tcPr.append(tcW)


def set_cell_margins(cell, top=0, bottom=0, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def add_run(paragraph, text, font_name=FONT_NAME, size=FONT_BODY, color=BLACK,
            bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rPr.append(parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" '
        f'w:hAnsi="{font_name}" w:eastAsia="{font_name}" w:cs="{font_name}"/>'
    ))
    return run


def add_body_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       space_after=Pt(6), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = Pt(14)
    add_run(p, text, size=FONT_BODY, color=BLACK)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)
    if level == 0:
        p.paragraph_format.left_indent = Inches(0.3)
    else:
        p.paragraph_format.left_indent = Inches(0.6)

    sz = FONT_BODY if level == 0 else Pt(10)
    clr = BLACK if level == 0 else DARK_GRAY
    if bold_prefix:
        add_run(p, f"● {bold_prefix}", size=sz, color=BLACK, bold=True)
        add_run(p, f" {text}", size=sz, color=clr)
    else:
        add_run(p, f"● {text}", size=sz, color=clr)
    return p


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_full_width(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)


def add_navy_section_header(doc, title, page_break=True):
    if page_break:
        doc.add_page_break()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "2D3E50")
    set_cell_margins(cell, top=60, bottom=60, left=200, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title.upper(), size=FONT_SECTION, color=WHITE, bold=True)
    remove_table_borders(table)
    set_table_full_width(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_teal_sub_header(doc, title):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "3498A2")
    set_cell_margins(cell, top=40, bottom=40, left=200, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title, size=FONT_SUBSECTION, color=WHITE, bold=True)
    remove_table_borders(table)
    set_table_full_width(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_evidence_block(doc, metadata_lines, impact_text, exhibit_num, compact=False):
    """Evidence block: 2-col table. Left=metadata+impact. Right=thumbnail."""
    mw = COMPACT_META if compact else META_WIDTH
    tw = COMPACT_THUMB if compact else THUMB_WIDTH
    impact_sz = FONT_COMPACT if compact else FONT_SMALL

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    set_cell_width(left_cell, mw)
    set_cell_width(right_cell, tw)
    set_cell_borders(left_cell, "CCCCCC", 4)
    set_cell_borders(right_cell, "CCCCCC", 4)
    set_cell_margins(left_cell, top=80, bottom=80, left=120, right=120)
    set_cell_margins(right_cell, top=80, bottom=80, left=80, right=80)

    # Left cell: metadata
    first = True
    for label, value in metadata_lines:
        if first:
            p = left_cell.paragraphs[0]
            first = False
        else:
            p = left_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        add_run(p, f"{label}: ", size=Pt(10), color=BLACK, bold=True)
        add_run(p, value, size=Pt(10), color=BLACK)

    # Impact inside left cell
    if impact_text:
        p_label = left_cell.add_paragraph()
        p_label.paragraph_format.space_before = Pt(6)
        p_label.paragraph_format.space_after = Pt(2)
        add_run(p_label, "Descrição e Impacto", size=FONT_SMALL, color=NAVY, bold=True)

        p_impact = left_cell.add_paragraph()
        p_impact.paragraph_format.space_after = Pt(4)
        p_impact.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p_impact, impact_text, size=impact_sz, color=DARK_GRAY, italic=True)

    # Right cell: thumbnail placeholder
    p_thumb = right_cell.paragraphs[0]
    p_thumb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_run(p_thumb, f"[THUMBNAIL — Exhibit {exhibit_num}]",
            size=FONT_CONTACT, color=MED_GRAY, italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_company_box(doc, company, role, period, location, sector):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")
    set_cell_borders(cell, "CCCCCC", 4)
    set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
    set_table_full_width(table)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_run(p, company, size=FONT_BODY, color=NAVY, bold=True)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    add_run(p2, role, size=FONT_BODY, color=BLACK, bold=True)
    add_run(p2, f"  |  {period}", size=FONT_SMALL, color=DARK_GRAY)

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    add_run(p3, f"{location}  |  {sector}", size=FONT_SMALL, color=MED_GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


# ============================================================
# DOCUMENT SETUP
# ============================================================

def setup_document():
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.header_distance = Inches(0)

    # Default style
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_BODY
    return doc


def build_header(doc):
    header_section = doc.sections[0].header
    header_section.is_linked_to_previous = False

    table = header_section.add_table(rows=3, cols=2, width=Inches(7.2))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)
    remove_table_borders(table)

    # Row 0: Navy — Name
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "2D3E50")
    set_cell_margins(table.cell(0, 0), top=120, bottom=40, left=200, right=100)
    set_cell_margins(table.cell(0, 1), top=120, bottom=40, left=100, right=200)

    p_name = table.cell(0, 0).paragraphs[0]
    add_run(p_name, BENEFICIARY_NAME, size=FONT_HEADER_NAME, color=WHITE, bold=True)

    # Row 1: Navy — RÉSUMÉ + EB-2 NIW
    for cell in table.rows[1].cells:
        set_cell_shading(cell, "2D3E50")
    set_cell_margins(table.cell(1, 0), top=20, bottom=60, left=200, right=100)
    set_cell_margins(table.cell(1, 1), top=20, bottom=60, left=100, right=200)

    p_resume = table.cell(1, 0).paragraphs[0]
    add_run(p_resume, "RÉSUMÉ", size=FONT_SECTION, color=WHITE, bold=True)

    p_soc = table.cell(1, 1).paragraphs[0]
    p_soc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_soc, f"SOC/O*Net {SOC_CODE}  |  EB-2 NIW",
            size=FONT_CONTACT, color=WHITE)

    # Row 2: Teal accent
    for cell in table.rows[2].cells:
        set_cell_shading(cell, "3498A2")
    set_row_height(table.rows[2], 4)

    for p in header_section.paragraphs:
        if p.text == "":
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)


def build_footer(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False

    table = footer.add_table(rows=1, cols=1, width=Inches(7.2))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "2D3E50")
    set_cell_margins(cell, top=40, bottom=40, left=200, right=200)
    set_table_full_width(table)
    remove_table_borders(table)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = p.add_run("Page ")
    run1.font.name = FONT_NAME
    run1.font.size = FONT_CONTACT
    run1.font.color.rgb = WHITE

    # PAGE field
    r1 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" '
        f'w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/>'
        f'<w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r1.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    p._p.append(r1)

    instrPage = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" '
        f'w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/>'
        f'<w:color w:val="FFFFFF"/></w:rPr>'
        f'<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
    )
    p._p.append(instrPage)

    r2 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" '
        f'w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/>'
        f'<w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r2.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    p._p.append(r2)

    run_of = p.add_run(" of ")
    run_of.font.name = FONT_NAME
    run_of.font.size = FONT_CONTACT
    run_of.font.color.rgb = WHITE

    # NUMPAGES field
    r3 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" '
        f'w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/>'
        f'<w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r3.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    p._p.append(r3)

    instrNum = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" '
        f'w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/>'
        f'<w:color w:val="FFFFFF"/></w:rPr>'
        f'<w:instrText xml:space="preserve"> NUMPAGES </w:instrText></w:r>'
    )
    p._p.append(instrNum)

    r4 = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" '
        f'w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/>'
        f'<w:color w:val="FFFFFF"/></w:rPr></w:r>'
    )
    r4.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    p._p.append(r4)

    for p in footer.paragraphs:
        if p.text == "":
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)


# ============================================================
# SECTION BUILDERS
# ============================================================

def build_sintese(doc):
    """Síntese Profissional."""
    add_navy_section_header(doc, "SÍNTESE PROFISSIONAL", page_break=False)

    add_body_paragraph(doc,
        "Bruno da Silva Ucella é engenheiro mecânico com pós-graduação lato sensu em "
        "Engenharia de Materiais — com ênfase em Engenharia, Produção e Construção — "
        "especializado em tratamento de superfícies, revestimentos protetivos e pintura "
        "aeronáutica. Com quase duas décadas de experiência na Embraer S.A., terceira maior "
        "fabricante de jatos comerciais do mundo, desenvolveu competência excepcional em "
        "processos de acabamento que garantem a integridade estrutural, proteção contra "
        "corrosão e longevidade operacional de aeronaves comerciais e executivas que operam "
        "em condições ambientais extremas nos cinco continentes."
    )

    add_body_paragraph(doc,
        "Ao longo de sua trajetória profissional na Embraer, Bruno assumiu responsabilidades "
        "progressivamente complexas que incluem a execução de atividades especializadas de "
        "pintura em aeronaves de alto desempenho, a orientação técnica de profissionais em "
        "formação e o domínio operacional de processos regulados simultaneamente pela Agência "
        "Nacional de Aviação Civil (ANAC) e pela Federal Aviation Administration (FAA). Essa "
        "trajetória progressiva evidencia capacidade excepcional para coordenar operações de "
        "engenharia em ambientes de conformidade regulatória multinacional, onde cada etapa do "
        "processo produtivo é auditada e certificada segundo padrões internacionais de "
        "aeronavegabilidade."
    )

    add_body_paragraph(doc,
        "Sua formação acadêmica avançada combina a base analítica da Engenharia Mecânica — "
        "obtida na Faculdade Anhanguera de São José — com especialização em materiais de "
        "última geração pela Universidade Pitágoras Unopar Anhanguera, abrangendo "
        "nanomateriais, materiais metálicos, poliméricos, cerâmicos e vítreos, técnicas de "
        "caracterização e tecnologias de manufatura avançada. Essa combinação habilita-o a "
        "avaliar, selecionar e aplicar sistemas de revestimento que atendam simultaneamente "
        "requisitos de desempenho aerodinâmico, proteção anticorrosiva e conformidade com "
        "regulamentações internacionais de segurança da aviação."
    )

    add_body_paragraph(doc,
        "Bruno acumulou mais de 28 certificações técnicas especializadas concedidas pela "
        "Embraer em áreas estratégicas como tratamento de superfícies, anodização crômica, "
        "prevenção de corrosão, conformidade com regulamentos da FAA, certificação ANAC "
        "RBAC-145, gestão da qualidade pelo método 5S, manufatura de peças primárias de "
        "aeronaves, controle de exportação de tecnologia aeronáutica e gestão de qualificação "
        "de pessoas — um portfólio de qualificações que poucos profissionais do setor "
        "aeronáutico conseguem reunir ao longo de uma carreira inteira. Esse acervo demonstra "
        "não apenas profundidade técnica, mas também amplitude multidisciplinar que o posiciona "
        "como profissional de capacidade excepcional em sua área de atuação."
    )


def build_historico(doc):
    """Histórico Profissional — Gantt Timeline."""
    add_navy_section_header(doc, "HISTÓRICO PROFISSIONAL")

    years = [2006, 2008, 2010, 2012, 2014, 2016, 2018, 2020, 2022, 2024, 2026]
    activities = [
        ("Embraer S.A. — Pintor de Aeronaves", 2006, 2026),
        ("Eng. Mecânica — Anhanguera", 2013, 2018),
        ("Pós-Grad. Materiais — Unopar", 2022, 2023),
    ]

    n_cols = len(years) + 1
    n_rows = len(activities) + 1
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)

    # Header row
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, "2D3E50")
        set_cell_margins(cell, top=40, bottom=40, left=15, right=15)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            add_run(p, "Atividade", size=FONT_TABLE_SMALL, color=WHITE, bold=True)
        else:
            add_run(p, str(years[i - 1]), size=Pt(7), color=WHITE, bold=True)

    # Data rows
    for row_idx, (name, start, end) in enumerate(activities):
        row = table.rows[row_idx + 1]
        bg = "F5F5F5" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, cell in enumerate(row.cells):
            set_cell_margins(cell, top=30, bottom=30, left=15, right=15)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if col_idx == 0:
                set_cell_shading(cell, bg)
                add_run(p, name, size=Pt(7), color=BLACK)
            else:
                year = years[col_idx - 1]
                if start <= year <= end:
                    set_cell_shading(cell, "3498A2")
                else:
                    set_cell_shading(cell, bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def build_experiencia(doc):
    """Experiência Profissional Detalhada."""
    add_navy_section_header(doc, "EXPERIÊNCIA PROFISSIONAL DETALHADA")

    add_company_box(doc,
        company="Embraer S.A.",
        role="Pintor de Aeronaves (Aircraft Painter)",
        period="Dezembro de 2006 — Presente",
        location="São José dos Campos, SP, Brasil",
        sector="Indústria Aeronáutica"
    )

    add_body_paragraph(doc,
        "A Embraer S.A. é a terceira maior fabricante de jatos comerciais do mundo e líder "
        "global no segmento de aviação executiva, com presença em mais de 90 países e uma "
        "frota de mais de 8.000 aeronaves em operação. Com sede em São José dos Campos, "
        "Brasil, a empresa opera sob os mais rigorosos padrões de qualidade e segurança da "
        "indústria, mantendo certificações simultâneas de ANAC, FAA e EASA."
    )

    add_bullet(doc,
        "Executa processos completos de pintura e revestimento protetor em aeronaves "
        "comerciais e executivas, envolvendo preparação de superfícies metálicas e "
        "compósitas, aplicação de primers anticorrosivos, topcoats de alta performance e "
        "acabamentos especiais que atendem a especificações técnicas de engenharia e "
        "padrões internacionais de aeronavegabilidade estabelecidos por ANAC, FAA e EASA."
    )

    add_bullet(doc,
        "Orienta tecnicamente profissionais menos experientes, transmitindo conhecimentos "
        "especializados sobre técnicas avançadas de preparação e tratamento de superfícies "
        "aeronáuticas, métodos de aplicação de sistemas de pintura multicamadas, "
        "procedimentos de inspeção de acabamentos e interpretação de documentação técnica "
        "— função que demonstra liderança técnica e capacidade de coordenar a formação de "
        "equipes especializadas dentro do processo produtivo."
    )

    add_bullet(doc,
        "Domina procedimentos operacionais regulados simultaneamente por ANAC e FAA, "
        "assegurando que cada etapa do processo — desde a preparação da superfície até a "
        "inspeção final do revestimento — atenda integralmente aos requisitos documentados "
        "nos manuais de serviço, boletins de manutenção e especificações de material "
        "aprovadas pelas autoridades aeronáuticas competentes."
    )

    add_bullet(doc,
        "Aplica conhecimentos especializados em prevenção de corrosão, tratamento de "
        "superfícies e anodização crômica para garantir a proteção estrutural de "
        "componentes críticos de aeronaves, contribuindo diretamente para a segurança de "
        "voo e a extensão da vida útil operacional das aeronaves."
    )

    add_bullet(doc,
        "Participa de processos de controle de qualidade aderentes às metodologias "
        "Lean Manufacturing e 5S implementadas na linha de produção aeronáutica, "
        "contribuindo para a redução de desperdícios, a padronização de processos e a "
        "manutenção de ambientes de trabalho organizados e seguros em conformidade com "
        "normas aeronáuticas."
    )

    # Evidence block: Employer Declaration
    add_evidence_block(doc,
        metadata_lines=[
            ("Documento", "Declaração do Empregador — Embraer S.A."),
            ("Emitido por", "Willians Alves Silva, Supervisor de Folha de Pagamento"),
            ("Data", "20 de julho de 2023"),
            ("Cargo", "Pintor de Aviões (Aircraft Painter)"),
            ("Admissão", "04 de dezembro de 2006"),
        ],
        impact_text=(
            "Declaração oficial da Embraer S.A. que confirma o vínculo empregatício "
            "contínuo de Bruno da Silva Ucella desde dezembro de 2006, totalizando quase "
            "duas décadas de experiência ininterrupta na maior empresa aeronáutica da "
            "América Latina. O documento atesta que Bruno ocupa o cargo de Pintor de "
            "Aviões e desempenha atividades especializadas de pintura em aeronaves, além "
            "de orientar tecnicamente profissionais menos experientes — evidenciando "
            "expertise técnica consolidada e papel de liderança na formação de equipes "
            "dentro do processo produtivo aeronáutico."
        ),
        exhibit_num=1
    )


def build_contribuicoes(doc):
    """Contribuições Técnicas e Profissionais (por tema)."""
    add_navy_section_header(doc, "CONTRIBUIÇÕES TÉCNICAS E PROFISSIONAIS")

    add_body_paragraph(doc,
        "As contribuições técnicas de Bruno da Silva Ucella estão organizadas por área "
        "temática, refletindo a amplitude e profundidade de suas qualificações no setor "
        "aeronáutico. Cada certificação representa competência operacional validada pela "
        "Embraer S.A. em processos que impactam diretamente a segurança, durabilidade e "
        "conformidade regulatória de aeronaves comerciais e executivas."
    )

    # ── Theme 1: Surface Treatment ──
    add_teal_sub_header(doc, "Tratamento de Superfícies e Revestimentos Aeronáuticos")

    add_body_paragraph(doc,
        "O tratamento de superfícies constitui uma das disciplinas mais críticas da "
        "manufatura aeronáutica, envolvendo processos químicos e mecânicos que preparam "
        "substratos metálicos e compósitos para receber sistemas de revestimento protetor. "
        "A expertise de Bruno nesta área abrange desde a anodização crômica de ligas de "
        "alumínio até a aplicação de sistemas completos de pintura multicamadas, passando "
        "por métodos avançados de prevenção de corrosão essenciais para a operação segura "
        "de aeronaves em ambientes atmosféricos agressivos."
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "Pintura Aeronáutica (Aeronautical Painting)"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação Técnica Especializada"),
        ],
        impact_text=(
            "Certificação que valida a competência de Bruno na execução de processos de "
            "pintura em aeronaves segundo padrões industriais da Embraer. Esta qualificação "
            "demonstra domínio de técnicas de preparação de superfícies, aplicação de primers "
            "e topcoats aeronáuticos, controle de espessura de película, aderência e "
            "acabamento superficial. A pintura aeronáutica constitui camada crítica de "
            "proteção contra corrosão, radiação ultravioleta, erosão por partículas e "
            "variações térmicas extremas que aeronaves enfrentam em operação — competência "
            "que evidencia domínio técnico em processos regulados de alta responsabilidade."
        ),
        exhibit_num=2
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "Tratamento de Superfícies (Surface Treatment)"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação Técnica Especializada"),
        ],
        impact_text=(
            "Certificação que comprova qualificação em processos de tratamento químico e "
            "mecânico de superfícies metálicas aeronáuticas. O tratamento de superfícies é "
            "etapa preparatória fundamental que determina a qualidade e durabilidade de todo "
            "o sistema de revestimento subsequente. Esta competência demonstra conhecimento "
            "avançado em processos como desengraxe, decapagem, conversão química e preparação "
            "para primers — operações que exigem controle rigoroso de parâmetros como "
            "temperatura, concentração, tempo de imersão e qualidade da água de lavagem, "
            "com impacto direto na integridade estrutural da aeronave."
        ),
        exhibit_num=3
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "Anodização Crômica (Chromic Anodizing)"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação Técnica Especializada"),
        ],
        impact_text=(
            "Certificação em anodização crômica — processo eletroquímico crítico que cria "
            "uma camada de óxido protetora sobre ligas de alumínio aeronáuticas. Este "
            "processo é especificamente requerido por especificações militares e comerciais "
            "para componentes estruturais de aeronaves, oferecendo proteção superior contra "
            "corrosão intergranular e sob tensão. A competência de Bruno neste processo "
            "demonstra domínio de operações que exigem controle preciso de parâmetros "
            "eletroquímicos, em conformidade com especificações amplamente adotadas na "
            "indústria aeronáutica global para proteção de ligas de alta resistência."
        ),
        exhibit_num=4
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "Prevenção de Corrosão (Corrosion Prevention)"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação Técnica Especializada"),
        ],
        impact_text=(
            "Certificação que valida expertise em métodos de prevenção e controle de "
            "corrosão em estruturas aeronáuticas. A corrosão é uma das principais ameaças "
            "à integridade estrutural de aeronaves e sua prevenção eficaz é requisito "
            "mandatório das autoridades aeronáuticas internacionais. Esta qualificação "
            "demonstra competência para identificar, avaliar e prevenir diferentes tipos "
            "de corrosão — incluindo galvânica, filiforme, intergranular e sob tensão — "
            "aplicando tratamentos protetivos que garantem a segurança operacional das "
            "aeronaves ao longo de todo o seu ciclo de vida útil."
        ),
        exhibit_num=5
    )

    # ── Theme 2: Regulatory Compliance ──
    add_navy_section_header(doc, "CONTRIBUIÇÕES TÉCNICAS E PROFISSIONAIS (CONT.)")
    add_teal_sub_header(doc, "Conformidade Regulatória e Certificação Internacional")

    add_body_paragraph(doc,
        "A indústria aeronáutica opera sob um dos marcos regulatórios mais rigorosos do "
        "mundo, com requisitos de certificação impostos por autoridades como FAA, ANAC e "
        "EASA. As certificações de Bruno nesta área demonstram domínio operacional das "
        "normas que governam a fabricação, manutenção e reparo de aeronaves — conhecimento "
        "indispensável para profissionais que coordenam atividades de engenharia em "
        "organizações de produção certificadas."
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "FAA (Federal Aviation Administration)"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação de Conformidade Regulatória"),
        ],
        impact_text=(
            "Certificação que comprova conhecimento sobre os regulamentos e requisitos "
            "da Federal Aviation Administration, autoridade aeronáutica dos Estados Unidos. "
            "O domínio das normas FAA é requisito fundamental para profissionais que "
            "supervisionam a fabricação de aeronaves destinadas ao mercado norte-americano, "
            "uma vez que a FAA exerce jurisdição sobre a certificação de tipo e a "
            "aeronavegabilidade continuada de todas as aeronaves operando nos EUA. Esta "
            "qualificação demonstra aptidão para operar em conformidade com o marco "
            "regulatório mais influente da aviação mundial."
        ),
        exhibit_num=6
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "Organização e Regulamentação FAA"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação de Conformidade Regulatória"),
        ],
        impact_text=(
            "Certificação complementar que aprofunda o conhecimento sobre a estrutura "
            "organizacional da FAA e o arcabouço regulatório que governa a fabricação e "
            "manutenção de aeronaves. Este treinamento abrange a compreensão dos Federal "
            "Aviation Regulations (FARs), Advisory Circulars, Airworthiness Directives e "
            "o sistema de certificação de organizações de produção — conhecimentos que "
            "permitem interpretar e aplicar corretamente as exigências regulatórias em cada "
            "etapa do processo produtivo, assegurando conformidade contínua com os padrões "
            "da autoridade aeronáutica dos Estados Unidos."
        ),
        exhibit_num=7
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "ANAC RBAC-145"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação de Conformidade Regulatória Nacional"),
        ],
        impact_text=(
            "Certificação no Regulamento Brasileiro da Aviação Civil nº 145 (RBAC-145), "
            "que estabelece os requisitos para organizações de manutenção de aeronaves. "
            "Este regulamento, harmonizado com os padrões da ICAO e da FAA Part 145, "
            "define as condições sob as quais uma organização pode executar serviços de "
            "manutenção, reparo e revisão em aeronaves e componentes aeronáuticos. A "
            "certificação demonstra aptidão para atuar em ambientes regulados por padrões "
            "de segurança aeronáutica reconhecidos internacionalmente, coordenando "
            "processos que exigem rastreabilidade completa."
        ),
        exhibit_num=8
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "Controles de Comércio Exterior dos EUA"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação de Conformidade em Exportação"),
        ],
        impact_text=(
            "Certificação que comprova conhecimento sobre os regulamentos de controle de "
            "exportação dos Estados Unidos, incluindo o International Traffic in Arms "
            "Regulations (ITAR) e os Export Administration Regulations (EAR). Na indústria "
            "aeronáutica, o domínio destes regulamentos é essencial dado que tecnologias, "
            "materiais e processos de fabricação de aeronaves são frequentemente "
            "classificados como itens controlados. Esta qualificação demonstra capacidade "
            "para operar em conformidade com as leis de exportação que regem a "
            "transferência de tecnologia aeronáutica sensível entre países."
        ),
        exhibit_num=9
    )

    # ── Theme 3: Manufacturing and Quality ──
    add_navy_section_header(doc, "CONTRIBUIÇÕES TÉCNICAS E PROFISSIONAIS (CONT.)")
    add_teal_sub_header(doc, "Manufatura Aeronáutica e Gestão da Qualidade")

    add_body_paragraph(doc,
        "A manufatura aeronáutica exige níveis de precisão e controle de qualidade que "
        "superam a maioria das indústrias, dado que qualquer defeito em componentes ou "
        "processos pode comprometer a segurança de voo. As certificações de Bruno em "
        "manufatura de peças primárias, metodologias Lean e sistemas de gestão da "
        "qualidade demonstram capacidade de supervisionar e otimizar processos produtivos "
        "que atendem aos padrões mais exigentes da indústria."
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "Manufatura de Peças Primárias de Aeronaves"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação Técnica de Manufatura"),
        ],
        impact_text=(
            "Certificação que valida competência em processos de fabricação de peças "
            "primárias — componentes estruturais de aeronaves que suportam cargas "
            "aerodinâmicas e mecânicas durante o voo. A manufatura de peças primárias "
            "exige conformidade rigorosa com especificações de material, tolerâncias "
            "dimensionais e requisitos de acabamento superficial que são auditados e "
            "aprovados pelas autoridades aeronáuticas. Esta qualificação demonstra "
            "conhecimento dos processos de conformação, usinagem e acabamento de "
            "componentes críticos para a aeronavegabilidade das aeronaves."
        ),
        exhibit_num=10
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "Lean Thinking"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação em Gestão de Processos"),
        ],
        impact_text=(
            "Certificação na metodologia Lean Thinking, que demonstra capacidade para "
            "identificar e eliminar desperdícios em processos produtivos aeronáuticos, "
            "otimizar fluxos de valor e promover melhoria contínua. Na indústria "
            "aeronáutica, a aplicação de princípios Lean resulta não apenas em eficiência "
            "operacional, mas também em maior consistência de qualidade — fator crítico "
            "quando cada componente deve atender a especificações precisas de engenharia. "
            "Esta competência posiciona Bruno como profissional capaz de contribuir para "
            "a excelência operacional em ambientes de manufatura avançada."
        ),
        exhibit_num=11
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "Metodologia 5S"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação em Gestão da Qualidade"),
        ],
        impact_text=(
            "Certificação na metodologia 5S (Seiri, Seiton, Seisō, Seiketsu, Shitsuke), "
            "sistema de organização e padronização de ambientes de trabalho amplamente "
            "adotado na indústria aeronáutica para garantir segurança, eficiência e "
            "rastreabilidade de processos. Em linhas de produção aeronáuticas, a "
            "implementação rigorosa do 5S é requisito de conformidade que impacta "
            "diretamente a qualidade do produto final e a prevenção de Foreign Object "
            "Debris (FOD), uma das preocupações mais críticas da manufatura aeronáutica "
            "com impacto direto na segurança de voo."
        ),
        exhibit_num=12
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Certificação", "Qualidade Sempre (Quality Always)"),
            ("Emissor", "Embraer S.A."),
            ("Tipo", "Certificação em Gestão da Qualidade"),
        ],
        impact_text=(
            "Certificação no programa de qualidade total da Embraer, que integra os "
            "princípios de garantia da qualidade aos processos cotidianos de produção. "
            "Este programa enfatiza a responsabilidade individual pela qualidade em cada "
            "etapa do processo produtivo, desde o recebimento de materiais até a entrega "
            "final. A certificação demonstra que Bruno internalizou a cultura de qualidade "
            "que posiciona a Embraer como referência mundial em excelência de fabricação "
            "aeronáutica, com índices de confiabilidade reconhecidos por operadores aéreos "
            "e autoridades regulatórias em todo o mundo."
        ),
        exhibit_num=13
    )


def build_formacao(doc):
    """Formação Acadêmica."""
    add_navy_section_header(doc, "FORMAÇÃO ACADÊMICA")

    add_body_paragraph(doc,
        "A formação acadêmica de Bruno da Silva Ucella combina a base sólida da engenharia "
        "mecânica com especialização avançada em ciência e engenharia de materiais, "
        "proporcionando fundamentação teórica e analítica que complementa sua extensa "
        "experiência prática na indústria aeronáutica. O grau avançado de pós-graduação "
        "atende ao requisito de formação acadêmica qualificada para a classificação EB-2 NIW."
    )

    add_teal_sub_header(doc, "Pós-Graduação Lato Sensu")

    add_evidence_block(doc,
        metadata_lines=[
            ("Curso", "Engenharia de Materiais — Engenharia, Produção e Construção"),
            ("Instituição", "Universidade Pitágoras Unopar Anhanguera"),
            ("Período", "Outubro de 2022 — Abril de 2023"),
            ("Carga Horária", "360 horas"),
            ("Grau Conferido", "Especialização Lato Sensu (Pós-Graduação)"),
        ],
        impact_text=(
            "O programa de pós-graduação em Engenharia de Materiais abrangeu disciplinas "
            "de fronteira tecnológica incluindo bio e nanomateriais (conceito 10,0), "
            "materiais metálicos (8,0), poliméricos (7,0), cerâmicos e vítreos (7,0), "
            "ciência dos materiais (10,0), técnicas de caracterização (8,0), tecnologias "
            "de manufatura avançada (8,0), embalagem (8,0) e reciclagem de materiais "
            "(7,0). Esta formação avançada proporciona conhecimento aprofundado sobre "
            "propriedades, comportamento e aplicações de materiais fundamentais para a "
            "seleção e aplicação de sistemas de revestimento na indústria aeronáutica, "
            "onde a compatibilidade entre materiais e a resistência a ambientes "
            "operacionais extremos são requisitos inegociáveis."
        ),
        exhibit_num=14
    )

    add_teal_sub_header(doc, "Graduação")

    add_evidence_block(doc,
        metadata_lines=[
            ("Curso", "Engenharia Mecânica (Bacharelado)"),
            ("Instituição", "Faculdade Anhanguera de São José / UNIDERP"),
            ("Conclusão", "Dezembro de 2017 — Colação em Janeiro de 2018"),
            ("Grau Conferido", "Bacharel em Engenharia Mecânica"),
            ("Diploma Emitido", "30 de julho de 2018, Campo Grande, MS"),
        ],
        impact_text=(
            "O Bacharelado em Engenharia Mecânica forneceu a base científica e analítica "
            "essencial para compreender os fenômenos físicos, químicos e mecânicos que "
            "governam o comportamento de materiais e estruturas em aplicações aeronáuticas. "
            "A formação abrange termodinâmica, mecânica dos fluidos, resistência dos "
            "materiais, processos de fabricação e controle de qualidade — disciplinas "
            "diretamente aplicáveis à atuação na Embraer em processos de tratamento de "
            "superfícies e revestimentos que devem resistir a tensões térmicas, mecânicas "
            "e químicas em condições operacionais de voo."
        ),
        exhibit_num=15
    )

    add_evidence_block(doc,
        metadata_lines=[
            ("Documento", "Histórico Escolar — Engenharia Mecânica"),
            ("Instituição", "Faculdade Anhanguera de São José"),
            ("Tipo", "Registro Acadêmico Oficial"),
        ],
        impact_text=(
            "O histórico escolar documenta o percurso acadêmico completo de Bruno durante "
            "o curso de Engenharia Mecânica, incluindo todas as disciplinas cursadas, "
            "cargas horárias, notas obtidas e aprovações em cada etapa do currículo. Este "
            "documento comprova a integralização de todos os créditos exigidos para a "
            "obtenção do grau de Bacharel em Engenharia Mecânica, demonstrando formação "
            "abrangente nas áreas fundamentais e aplicadas da engenharia."
        ),
        exhibit_num=16,
        compact=True
    )


def build_cursos(doc):
    """Cursos e Certificações Técnicas."""
    add_navy_section_header(doc, "CURSOS E CERTIFICAÇÕES TÉCNICAS")

    add_body_paragraph(doc,
        "Além das certificações apresentadas nas seções de contribuições técnicas, Bruno "
        "acumulou um amplo portfólio de certificações especializadas concedidas pela "
        "Embraer S.A. em áreas complementares de segurança operacional, instrumentação, "
        "gestão de pessoas e conformidade industrial. Este conjunto de qualificações "
        "evidencia formação multidisciplinar e comprometimento contínuo com o "
        "desenvolvimento profissional ao longo de quase duas décadas de carreira."
    )

    certificates = [
        ("Gerenciamento de Projetos (MS Project)", "Gestão de Projetos"),
        ("Instrumentos Básicos de Medição", "Instrumentação e Metrologia"),
        ("Manuseio e Armazenamento de Produtos Perigosos", "Segurança Operacional"),
        ("Comportamento Seguro e Uso de EPI", "Segurança Operacional"),
        ("Espaço Confinado (Confined Space)", "Segurança Operacional"),
        ("Leitura e Interpretação de Desenhos Embraer", "Documentação Técnica"),
        ("Eliminação de Objetos Estranhos (FOE)", "Qualidade e Segurança"),
        ("Proteção Auditiva (Hearing Protection)", "Saúde Ocupacional"),
        ("Veículos Industriais (Industrial Vehicles)", "Operação Industrial"),
        ("Gestão da Qualificação de Pessoas", "Gestão de Pessoas"),
        ("Responsabilidade Civil (Liability)", "Conformidade"),
        ("Segurança de Máquinas (Machine Safety)", "Segurança Operacional"),
        ("Cultura Aeronáutica Preventiva", "Cultura de Segurança"),
        ("Responsabilidade do Produto (Product Liability)", "Conformidade"),
        ("Ergonomia na Produção", "Saúde Ocupacional"),
        ("Proteção Respiratória (Respiratory Protection)", "Saúde Ocupacional"),
        ("Manual do Centro de Serviço (Service Center Manual)", "Documentação Técnica"),
    ]

    n_rows = len(certificates) + 1
    table = doc.add_table(rows=n_rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_full_width(table)

    # Header row
    headers = ["Certificação", "Área", "Emissor"]
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "2D3E50")
        set_cell_margins(cell, top=50, bottom=50, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h_text, size=FONT_CONTACT, color=WHITE, bold=True)

    # Data rows
    for row_idx, (cert_name, area) in enumerate(certificates):
        row = table.rows[row_idx + 1]
        bg = "F5F5F5" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, cell in enumerate(row.cells):
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=35, bottom=35, left=100, right=100)
            p = cell.paragraphs[0]
            if col_idx == 0:
                add_run(p, cert_name, size=FONT_CONTACT, color=BLACK)
            elif col_idx == 1:
                add_run(p, area, size=FONT_CONTACT, color=DARK_GRAY)
            else:
                add_run(p, "Embraer S.A.", size=FONT_CONTACT, color=DARK_GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ============================================================
# MAIN
# ============================================================

def main():
    doc = setup_document()
    build_header(doc)
    build_footer(doc)

    build_sintese(doc)
    build_historico(doc)
    build_experiencia(doc)
    build_contribuicoes(doc)
    build_formacao(doc)
    build_cursos(doc)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, OUTPUT_FILE)
    doc.save(output_path)
    print(f"Résumé salvo em: {output_path}")
    print(f"Tamanho: {os.path.getsize(output_path):,} bytes")
    return output_path


if __name__ == "__main__":
    main()
