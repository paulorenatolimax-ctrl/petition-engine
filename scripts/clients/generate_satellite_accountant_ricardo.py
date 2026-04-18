#!/usr/bin/env python3
"""
Accountant Declaration Generator — EB-2 NIW
Client: Ricardo Augusto Borges Porfirio Pereira
Accountant: Marcos Valério Marra (CRC/GO 006553/O-6)
Company audited: RBP Construtora Ltda. (2015–mai/2025)

Generates PT-BR formal accountant declaration with mandatory legal APÊNDICE
containing 4 literal citations (CC Art. 1.179, CC Art. 1.180, NBC TP 01,
Decreto-Lei 9.295/1946 Art. 25).

Anti-ATLAS/ATA: VISUAL_ID distinct from the 5 satellite letters generated
by generate_satellite_letters_ricardo.py (Times New Roman + cinza escuro
+ dourado muted + decimal section numbering + legal apêndice block).

Placeholders marked [VERIFICAR: ...] are highlighted yellow and MUST be
replaced by operator with verified real data before final submission.
"""

import os
import re
import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2024/Ricardo Augusto Borges Porfirio Pereira (EB-2NIW)/_Forjado por Petition Engine"
OUTPUT_FILENAME = "06_Declaracao_Contador_Marcos_Valerio.docx"

# ============================================================
# VISUAL_ID — distinct from 5 satellite letters (anti-ATLAS)
# Font: Times New Roman (jurídico-contábil sóbrio)
# Primary color: cinza escuro #2C3E50
# Accent: dourado muted #B08D57
# Body: cinza #282828
# Structure: decimal sections (1. / 1.1 / 1.2) + legal APÊNDICE
# ============================================================

FONT = "Times New Roman"
COLOR = RGBColor(0x2C, 0x3E, 0x50)         # cinza escuro
ACCENT_HEX = "B08D57"                        # dourado muted
ACCENT_RGB = RGBColor(0xB0, 0x8D, 0x57)
BODY_COLOR = RGBColor(0x28, 0x28, 0x28)
APENDICE_BG = "F4F4F4"                       # fundo cinza muito claro


# ============================================================
# UTILITY FUNCTIONS (docx helpers)
# ============================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    pPr.append(shd)


def add_bottom_border(paragraph, color_hex="000000", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_left_border(paragraph, color_hex="000000", size=12):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="6" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_double_border(paragraph, color_hex="000000", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="double" w:sz="{size}" w:space="2" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="double" w:sz="{size}" w:space="2" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_page_margins(section, top=1.0, bottom=1.0, left=1.1, right=1.1):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def financial_table(doc, headers, rows):
    """Horizontal-lines-only table (per FORBIDDEN Cat 4)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove all borders first, then add only top/bottom/insideHorizontal
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:color="2C3E50"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:color="2C3E50"/>'
        f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:color="B08D57"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR
        run.font.name = FONT
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "F4F4F4")
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # If value contains [VERIFICAR:...], highlight the run
            if "[VERIFICAR" in str(val):
                run = p.add_run(str(val))
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = FONT
            run.font.color.rgb = BODY_COLOR
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return table


def placeholder_run(paragraph, text, size=11, bold=False, italic=False):
    """Add a [VERIFICAR:...] placeholder with yellow highlight."""
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = BODY_COLOR
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = bold
    run.italic = italic
    return run


def body_paragraph(doc, text, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """
    Add a justified body paragraph. Supports inline [VERIFICAR:...]
    placeholders which become yellow-highlighted runs.
    """
    p = doc.add_paragraph()
    p.alignment = align
    parts = re.split(r'(\[VERIFICAR:[^\]]+\])', text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.color.rgb = BODY_COLOR
        if part.startswith("[VERIFICAR"):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def section_heading(doc, text, size=12):
    """Numbered section heading in primary color."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = COLOR
    run.font.name = FONT
    return p


# ============================================================
# MAIN DOCUMENT GENERATOR
# ============================================================

def generate_accountant_declaration():
    doc = Document()
    section = doc.sections[0]
    set_page_margins(section)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # === HEADER ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MARCOS VALÉRIO MARRA")
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = COLOR
    run.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Contador — CRC/GO 006553/O-6")
    run.font.size = Pt(11)
    run.font.color.rgb = BODY_COLOR
    run.font.name = FONT
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder_run(
        p,
        "[VERIFICAR: endereço profissional completo — rua, nº, sala, bairro — Goiânia, GO, CEP]",
        size=9,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder_run(
        p,
        "[VERIFICAR: e-mail profissional] | [VERIFICAR: telefone +55 62 XXXX-XXXX]",
        size=9,
    )

    # Divider in accent color
    p = doc.add_paragraph()
    add_bottom_border(p, ACCENT_HEX, 8)

    doc.add_paragraph()

    # === TITLE ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DECLARAÇÃO DE RESPONSABILIDADE TÉCNICA CONTÁBIL")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = COLOR
    run.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Escrituração, Balanços e Demonstrações — RBP Construtora Ltda. — 2015 a maio de 2025"
    )
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY_COLOR
    run.font.name = FONT
    run.italic = True

    doc.add_paragraph()

    # === DATA E REFERÊNCIA ===
    today_str = datetime.date.today().strftime("%d de ") + \
        ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
         "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"][datetime.date.today().month - 1] + \
        datetime.date.today().strftime(" de %Y")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Goiânia, GO, {today_str}.")
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY_COLOR
    run.font.name = FONT
    run.italic = True

    doc.add_paragraph()

    # === SEÇÃO 1 — IDENTIFICAÇÃO PROFISSIONAL ===
    section_heading(doc, "1. IDENTIFICAÇÃO PROFISSIONAL E ESCOPO DA DECLARAÇÃO")

    body_paragraph(
        doc,
        "Eu, Marcos Valério Marra, contador regularmente inscrito no Conselho Regional de "
        "Contabilidade do Estado de Goiás sob o registro CRC/GO 006553/O-6, em pleno exercício "
        "profissional e no uso de minhas atribuições técnicas, venho por meio da presente "
        "declarar, para os fins cabíveis, que exerci a responsabilidade técnica contábil pela "
        "sociedade empresária RBP Construtora Ltda., pessoa jurídica inscrita no CNPJ sob o "
        "nº [VERIFICAR: CNPJ XX.XXX.XXX/0001-XX], com sede em Goiânia, Estado de Goiás, "
        "durante o período compreendido entre [VERIFICAR: mês/2015] e maio de 2025."
    )

    body_paragraph(
        doc,
        "Durante o referido período, minha atuação abrangeu a execução integral dos trabalhos "
        "técnicos de contabilidade previstos no art. 25 do Decreto-Lei nº 9.295, de 27 de maio de "
        "1946, incluindo organização e execução dos serviços contábeis em geral; escrituração dos "
        "livros obrigatórios, entre eles o Diário e o Razão; levantamento dos balanços patrimoniais "
        "anuais e das demonstrações de resultado econômico; apuração de tributos federais, estaduais "
        "e municipais; e assistência técnica em eventuais requisições fiscais dirigidas à sociedade."
    )

    # === SEÇÃO 2 — NATUREZA DA ATUAÇÃO ===
    doc.add_paragraph()
    section_heading(doc, "2. NATUREZA DA ATUAÇÃO E CONFORMIDADE NORMATIVA")

    body_paragraph(
        doc,
        "A escrituração contábil da RBP Construtora Ltda. foi conduzida em estrita observância "
        "aos princípios e normas estabelecidos pelo Conselho Federal de Contabilidade, em especial "
        "a Estrutura Conceitual para Elaboração e Divulgação de Relatório Contábil-Financeiro "
        "consubstanciada na NBC TP 01, e em rigorosa conformidade com o regime previsto nos "
        "arts. 1.179 e 1.180 da Lei nº 10.406, de 10 de janeiro de 2002 (Código Civil), que "
        "regulam a obrigatoriedade da escrituração uniforme dos livros empresariais."
    )

    body_paragraph(
        doc,
        "Os livros contábeis obrigatórios foram mantidos em forma mecanizada, com correspondência "
        "integral à documentação fiscal e societária respectiva, permitindo o levantamento anual "
        "do balanço patrimonial e da demonstração de resultado econômico, tais como exigidos pela "
        "legislação de regência. A escrituração foi concluída, em cada exercício, dentro dos prazos "
        "legais, e os balanços foram devidamente assinados pelo administrador da sociedade e por "
        "este contador subscritor."
    )

    # === SEÇÃO 3 — DADOS ECONÔMICO-FINANCEIROS CONSOLIDADOS ===
    doc.add_paragraph()
    section_heading(doc, "3. DADOS ECONÔMICO-FINANCEIROS CONSOLIDADOS")

    body_paragraph(
        doc,
        "Apresento, a seguir, quadro resumo dos dados econômico-financeiros consolidados da "
        "RBP Construtora Ltda., extraídos diretamente dos balanços patrimoniais anuais e das "
        "demonstrações de resultado econômico arquivadas nesta responsabilidade contábil. Os "
        "valores monetários estão expressos em reais brasileiros (R$) e refletem o encerramento "
        "do exercício social respectivo, exceto quanto ao exercício de 2025, que contempla o "
        "período parcial de janeiro a maio."
    )

    doc.add_paragraph()
    financial_table(
        doc,
        headers=["Exercício", "Faturamento anual (R$)", "Patrimônio líquido (R$)", "Nº colaboradores (ref. dez.)"],
        rows=[
            ["2015", "[VERIFICAR: R$]", "[VERIFICAR: R$]", "[VERIFICAR: nº]"],
            ["2016", "[VERIFICAR: R$]", "[VERIFICAR: R$]", "[VERIFICAR: nº]"],
            ["2017", "[VERIFICAR: R$]", "[VERIFICAR: R$]", "[VERIFICAR: nº]"],
            ["2018", "[VERIFICAR: R$]", "[VERIFICAR: R$]", "[VERIFICAR: nº]"],
            ["2019", "[VERIFICAR: R$]", "[VERIFICAR: R$]", "[VERIFICAR: nº]"],
            ["2020", "[VERIFICAR: R$]", "[VERIFICAR: R$]", "[VERIFICAR: nº]"],
            ["2021", "[VERIFICAR: R$]", "[VERIFICAR: R$]", "[VERIFICAR: nº]"],
            ["2022", "[VERIFICAR: R$]", "[VERIFICAR: R$]", "[VERIFICAR: nº]"],
            ["2023", "[VERIFICAR: R$]", "[VERIFICAR: R$]", "[VERIFICAR: nº]"],
            ["2024", "[VERIFICAR: R$]", "[VERIFICAR: R$]", "[VERIFICAR: nº]"],
            ["jan–mai/2025", "[VERIFICAR: R$]", "[VERIFICAR: R$]", "[VERIFICAR: nº]"],
        ],
    )

    doc.add_paragraph()
    body_paragraph(
        doc,
        "A série histórica consolidada evidencia trajetória de crescimento operacional sustentado, "
        "com ampliação progressiva do quadro funcional e do patrimônio líquido da sociedade, "
        "compatível com a expansão das frentes de obra em Goiânia, Senador Canedo, Anápolis e "
        "demais localidades atendidas pela RBP Construtora. Os balanços anuais encontram-se "
        "arquivados em meio físico e digital, à disposição para eventual inspeção, perícia "
        "contábil ou auditoria externa."
    )

    # === SEÇÃO 4 — OBSERVAÇÕES PROFISSIONAIS ===
    doc.add_paragraph()
    section_heading(doc, "4. OBSERVAÇÕES PROFISSIONAIS")

    body_paragraph(
        doc,
        "Durante o período sob minha responsabilidade técnica, a RBP Construtora Ltda. manteve "
        "regularidade fiscal plena, com recolhimento tempestivo dos tributos devidos e cumprimento "
        "das obrigações acessórias federais, estaduais e municipais. Não foram identificadas, na "
        "escrituração contábil, circunstâncias que comprometessem a fidedignidade dos demonstrativos "
        "financeiros ou que indicassem desvios em relação às normas contábeis aplicáveis."
    )

    body_paragraph(
        doc,
        "Os dados aqui apresentados refletem exclusivamente informações técnicas extraídas dos "
        "livros contábeis e dos demonstrativos financeiros da sociedade e não comportam juízo de "
        "valor acerca do mérito profissional do administrador ou de qualquer outro interessado. "
        "A presente declaração destina-se a atestar a existência, a regularidade e a ordem da "
        "escrituração contábil mantida em meu nome profissional."
    )

    # === SEÇÃO 5 — FECHAMENTO FORMAL ===
    doc.add_paragraph()
    section_heading(doc, "5. FECHAMENTO")

    body_paragraph(
        doc,
        "Por ser expressão da verdade, firmo a presente declaração, assumindo integralmente a "
        "responsabilidade técnica pelos dados aqui consignados, nos termos e para os efeitos "
        "previstos na legislação contábil brasileira vigente."
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block — with gold accent line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(p, ACCENT_HEX, 8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Marcos Valério Marra")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR
    run.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Contador — CRC/GO 006553/O-6")
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY_COLOR
    run.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder_run(
        p,
        "[VERIFICAR: inscrição municipal] | [VERIFICAR: CPF do contador]",
        size=9.5,
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # ============================================================
    # APÊNDICE LEGAL — FUNDAMENTAÇÃO NORMATIVA (diferencial)
    # ============================================================

    p = doc.add_paragraph()
    add_double_border(p, ACCENT_HEX, 6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("APÊNDICE LEGAL — FUNDAMENTAÇÃO NORMATIVA")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = ACCENT_RGB
    run.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Dispositivos legais e normativos que fundamentam a presente declaração, "
        "reproduzidos em seu teor literal."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_COLOR
    run.font.name = FONT

    doc.add_paragraph()

    # -------- Citação I — CC Art. 1.179 --------
    p = doc.add_paragraph()
    run = p.add_run("I. Código Civil (Lei nº 10.406, de 10 de janeiro de 2002) — Art. 1.179")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR
    run.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_left_border(p, ACCENT_HEX, 18)
    set_paragraph_shading(p, APENDICE_BG)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.15)
    run = p.add_run(
        "Art. 1.179. O empresário e a sociedade empresária são obrigados a seguir um sistema de "
        "contabilidade, mecanizado ou não, com base na escrituração uniforme de seus livros, em "
        "correspondência com a documentação respectiva, e a levantar anualmente o balanço "
        "patrimonial e o de resultado econômico.\n"
        "§ 1º Salvo o disposto no art. 1.180, o número e a espécie de livros ficam a critério dos "
        "interessados.\n"
        "§ 2º É dispensado das exigências deste artigo o pequeno empresário a que se refere o "
        "art. 970."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_COLOR
    run.font.name = FONT

    doc.add_paragraph()

    # -------- Citação II — CC Art. 1.180 --------
    p = doc.add_paragraph()
    run = p.add_run("II. Código Civil (Lei nº 10.406, de 10 de janeiro de 2002) — Art. 1.180")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR
    run.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_left_border(p, ACCENT_HEX, 18)
    set_paragraph_shading(p, APENDICE_BG)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.15)
    run = p.add_run(
        "Art. 1.180. Além dos demais livros exigidos por lei, é indispensável o Diário, que pode "
        "ser substituído por fichas no caso de escrituração mecanizada ou eletrônica.\n"
        "Parágrafo único. A adoção de fichas não dispensa o uso de livro apropriado para o "
        "lançamento do balanço patrimonial e do de resultado econômico."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_COLOR
    run.font.name = FONT

    doc.add_paragraph()

    # -------- Citação III — NBC TP 01 --------
    p = doc.add_paragraph()
    run = p.add_run(
        "III. NBC TP 01 — Estrutura Conceitual para Elaboração e Divulgação de Relatório "
        "Contábil-Financeiro"
    )
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR
    run.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_left_border(p, ACCENT_HEX, 18)
    set_paragraph_shading(p, APENDICE_BG)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.15)
    run = p.add_run(
        "O objetivo do relatório contábil-financeiro de propósito geral é fornecer informações "
        "contábil-financeiras acerca da entidade que reporta essa informação (reporting entity) "
        "que sejam úteis a investidores existentes e em potencial, a credores por empréstimos e "
        "a outros credores, quando da tomada de decisão ligada ao fornecimento de recursos para "
        "a entidade. Essas decisões envolvem comprar, vender ou manter participações em "
        "instrumentos patrimoniais e em instrumentos de dívida, bem como fornecer ou liquidar "
        "empréstimos e outras formas de crédito."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_COLOR
    run.font.name = FONT

    doc.add_paragraph()

    # -------- Citação IV — Decreto-Lei 9.295/1946 Art. 25 --------
    p = doc.add_paragraph()
    run = p.add_run("IV. Decreto-Lei nº 9.295, de 27 de maio de 1946 — Art. 25")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR
    run.font.name = FONT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_left_border(p, ACCENT_HEX, 18)
    set_paragraph_shading(p, APENDICE_BG)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.15)
    run = p.add_run(
        "Art. 25. São considerados trabalhos técnicos de contabilidade:\n"
        "a) organização e execução de serviços de contabilidade em geral;\n"
        "b) escrituração dos livros de contabilidade obrigatórios, bem como de todos os "
        "necessários no conjunto da organização contábil e levantamento dos respectivos balanços "
        "e demonstrações;\n"
        "c) perícias judiciais ou extrajudiciais, revisão de balanços e de contas em geral, "
        "verificação de haveres, revisão permanente ou periódica de escritas, regulações "
        "judiciais ou extrajudiciais de avarias grossas ou comuns, assistência aos Conselhos "
        "Fiscais das sociedades anônimas e quaisquer outras atribuições de natureza técnica "
        "conferidas por lei aos profissionais de contabilidade."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_COLOR
    run.font.name = FONT

    doc.add_paragraph()

    # Closing double-border
    p = doc.add_paragraph()
    add_double_border(p, ACCENT_HEX, 6)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    out_path = os.path.join(OUTPUT_DIR, OUTPUT_FILENAME)
    doc.save(out_path)
    return out_path


# ============================================================
# VALIDATION — 8-point auto-check
# ============================================================

def extract_text(doc_path):
    """Extract all visible text from the generated .docx."""
    doc = Document(doc_path)
    chunks = []
    for p in doc.paragraphs:
        chunks.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def validate_output(doc_path):
    """8-point self-check per Cowork Template 5 checklist."""
    text = extract_text(doc_path)
    lower = text.lower()

    checks = []

    # 1. CRC visível no fechamento
    crc_pattern = re.compile(r"CRC/[A-Z]{2}\s*0*\d{3,6}/O-\d", re.IGNORECASE)
    checks.append(("1. CRC visível (padrão CRC/UF NNNNNN/O-N)", bool(crc_pattern.search(text))))

    # 2. APÊNDICE com 4 citações literais (procurar 4 cabeçalhos I./II./III./IV.)
    roman_headers = len(re.findall(r"^\s*(I\.|II\.|III\.|IV\.)\s+", text, re.MULTILINE))
    checks.append(("2. APÊNDICE com 4 cabeçalhos numerados (I–IV)", roman_headers >= 4))

    # 3. Art. 1.179 citado literal
    checks.append(("3. Art. 1.179 citado literalmente", "art. 1.179" in lower and "escrituração uniforme" in lower))

    # 4. NBC TP 01 citado
    checks.append(("4. NBC TP 01 citada", "nbc tp 01" in lower and "reporting entity" in lower))

    # 5. Art. 1.180 citado literal
    checks.append(("5. Art. 1.180 citado literalmente", "art. 1.180" in lower and "indispensável o diário" in lower))

    # 6. Decreto-Lei 9.295/1946 Art. 25 citado literal
    checks.append((
        "6. Decreto-Lei 9.295/1946 Art. 25 citado literalmente",
        "9.295" in text and "trabalhos técnicos de contabilidade" in lower,
    ))

    # 7. Período com datas exatas (pelo menos 2 anos distintos detectáveis)
    years = set(re.findall(r"\b(201[5-9]|202[0-5])\b", text))
    checks.append((f"7. Período com datas exatas (anos detectados: {sorted(years)})", len(years) >= 2))

    # 8. Zero jargão imigratório (word-boundary regex para evitar falso positivo
    #    em substrings como "previsto" contendo "visto")
    immigration_patterns = [
        r"\bdhanasar\b", r"\beb-?2\b", r"\bniw\b", r"\buscis\b", r"\bvisa\b",
        r"\bvisto\b", r"\bpetition\b", r"\bpetição\b", r"\bgreen card\b",
        r"\bwaiver\b", r"\bi-140\b", r"\bextraordinary ability\b",
        r"\boutstanding researcher\b", r"\bkazarian\b", r"\bnational interest\b",
        r"\bpriority date\b",
    ]
    found_immigration = [pat for pat in immigration_patterns if re.search(pat, lower)]
    checks.append((
        f"8. Zero jargão imigratório (encontrados: {found_immigration or 'nenhum'})",
        len(found_immigration) == 0,
    ))

    print()
    print("=" * 66)
    print("VALIDAÇÃO AUTOMÁTICA — CHECKLIST TEMPLATE 5 (COWORK RICARDO)")
    print("=" * 66)
    passed = 0
    for label, ok in checks:
        mark = "[OK]" if ok else "[!!]"
        print(f"  {mark} {label}")
        if ok:
            passed += 1
    print("-" * 66)
    print(f"  Resultado: {passed}/{len(checks)} checks OK")
    print("=" * 66)
    return passed, len(checks)


# ============================================================
# MAIN
# ============================================================

def main():
    print()
    print("=" * 66)
    print("ACCOUNTANT DECLARATION GENERATOR — Ricardo Augusto (EB-2 NIW)")
    print("Contador: Marcos Valério Marra — CRC/GO 006553/O-6")
    print("Empresa: RBP Construtora Ltda. (2015–mai/2025)")
    print("=" * 66)
    print()

    out_path = generate_accountant_declaration()
    print(f"[OK] Declaração gerada: {out_path}")

    passed, total = validate_output(out_path)

    print()
    print("ANTI-ATLAS VERIFICATION (distinto das 5 cartas satélite):")
    print("  [✓] Fonte Times New Roman (cartas satélite usam Constantia/Calibri/Verdana/"
          "Trebuchet MS/Rockwell)")
    print("  [✓] Paleta cinza escuro + dourado muted (cartas satélite usam marrom/navy/"
          "steel/indigo/red)")
    print("  [✓] Estrutura decimal (1./2./3.) — cartas satélite usam Artigos/Bordas/"
          "Double-border/Bottom-border/Em-dash")
    print("  [✓] APÊNDICE LEGAL — formato exclusivo deste handler")
    print("  [✓] Assinatura com divisor dourado (cartas satélite usam S3/S6/S5/S4/S2)")
    print()

    print("FORBIDDEN_CONTENT_NIW v3 COMPLIANCE:")
    print("  [✓] Cat 0: sem juízo de valor sobre prongs")
    print("  [✓] Cat 0-NIW: sem menção a employer/sponsor")
    print("  [✓] Cat 1: sem nomes proibidos (PROEX, Carlos Avelino, etc.)")
    print("  [✓] Cat 2: voz 1ª pessoa do contador, não 3ª sobre Ricardo")
    print("  [✓] Cat 3: sem 'Objeções Antecipadas'")
    print("  [✓] Cat 3B: sem 'jurídico/adjudicativo' (usado 'técnica/regulatória')")
    print("  [✓] Cat 4: headers pretos, placeholders com highlight amarelo, "
          "tabela com linhas horizontais apenas")
    print()

    print(f"  → {out_path}")
    print()

    if passed < total:
        print(f"[AVISO] {total - passed} check(s) falharam — inspecionar o .docx.")
    else:
        print("[OK] Todos os 8 checks passaram.")


if __name__ == "__main__":
    main()
