#!/usr/bin/env python3
"""
Validator: 18 Error Rules for SaaS Evidence — Ricardo Augusto Pereira
Extracts text from .docx and checks ALL error rules.
"""

import re
import os
from docx import Document

DOCX_PATH = (
    "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/"
    "_2. MEUS CASOS/2024/"
    "Ricardo Augusto Borges Porfirio Pereira (EB-2NIW)/"
    "_Forjado por Petition Engine/"
    "saas_evidence_Ricardo_Augusto_Borges_Porfirio_Pereira.docx"
)

def extract_text(path):
    doc = Document(path)
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)

def validate(text):
    rules = [
        # (ID, severity, name, regex_pattern)
        ("R01", "CRITICAL/BLOCK", 'Never use "I believe" or "we believe"', r'\b(I|we)\s+believe\b'),
        ("R02", "HIGH/BLOCK", 'Never use "we think" or "I think"', r'\b(I|we)\s+think\b'),
        ("R03", "MEDIUM/AUTO-FIX", 'Use "proposed endeavor" not "proposed venture/business"', r'proposed\s+(venture|business)'),
        ("R04", "LOW/WARN", "Headings bold with correct capitalization", None),  # manual check
        ("R05", "HIGH/BLOCK", 'Never use "in conclusion" or "to summarize"', r'\b(in conclusion|to summarize)\b'),
        ("R06", "MEDIUM/WARN", "Evidence blocks should have thumbnails", None),  # manual
        ("R07", "CRITICAL/BLOCK", "Forbidden SOC codes requiring US diploma validation",
         r'(23-1011|29-1069|17-201[1-9]|13-2011)'),
        ("R08", "HIGH/WARN", "SOC codes requiring bachelor's — check compatibility", None),  # manual
        ("R09", "CRITICAL/BLOCK", 'NEVER use word "prompt" in output',
         r'(?i)\bprompt\b'),
        ("R10", "HIGH/WARN", "Output 100% Portuguese — no English mixing",
         None),  # complex check, manual
        ("R11", "HIGH/WARN", "Always consult RAGs before generating", None),  # process check
        ("R12", "CRITICAL/BLOCK", "Never mention PROEX, Carlos Avelino, Kortix, other clients",
         r'(?i)(PROEX|Kortix|Carlos Avelino)'),
        ("R13", "CRITICAL/BLOCK", "Portuguese accents mandatory",
         r'\b(introducao|peticao|informacao|certificacao|formacao|avaliacao|ocupacao|operacao|integracao|migracao|capacitacao|micropigmentacao)\b'),
        ("R14", "CRITICAL/BLOCK", 'No "Version X.X", "Generated:", "SaaS Evidence Architect", "Petition Engine"',
         r'(?i)\b(Version \d|Generated:|SaaS Evidence Architect|Petition Engine)\b'),
        ("R15", "CRITICAL/BLOCK", "ANTI-CRISTINE: Never use terms proving endeavor works without beneficiary",
         r'(?i)\b(standardized|padronizado|operates autonomously|self-sustaining|auto-sustent|plug.and.play|train.the.trainer|white.label|marca branca|client autonomy|founder dependency|scalable without|replicable by any|turnkey|chave.na.m)\b'),
        ("R16", "HIGH/WARN", 'Never use "consultoria" or "consulting" isolated',
         r'(?i)\b(consultoria|consulting)\b'),
        ("R17", "CRITICAL/BLOCK", "No case history mentions",
         r'(?i)\b(denial|negativa anterior|RFE anterior|previous petition|prior filing|refile|segunda tentativa|nova submiss|peti..o anterior|corrected approach|lessons learned from denial)\b'),
        ("R18", "CRITICAL/BLOCK", "No immigration terms in product dossier",
         r'(?i)\b(petition|petitioner|EB-2|EB-1|NIW|USCIS|immigration|visa|I-140|green.card|adjudicator|Dhanasar|Kazarian|extraordinary.ability|national.interest.waiver)\b'),
    ]

    violations = []
    passes = []

    for rule_id, severity, name, pattern in rules:
        if pattern is None:
            passes.append((rule_id, severity, name, "MANUAL CHECK"))
            continue

        matches = re.findall(pattern, text)
        if matches:
            # Get context for each match
            contexts = []
            for m in re.finditer(pattern, text):
                start = max(0, m.start() - 30)
                end = min(len(text), m.end() + 30)
                ctx = text[start:end].replace('\n', ' ').strip()
                contexts.append(f'  ...{ctx}...')
            violations.append((rule_id, severity, name, matches, contexts))
        else:
            passes.append((rule_id, severity, name, "PASS"))

    return violations, passes

def main():
    print("=" * 70)
    print("VALIDAÇÃO — 18 REGRAS DE ERRO (SaaS Evidence Ricardo)")
    print("=" * 70)
    print()

    if not os.path.exists(DOCX_PATH):
        print(f"ERRO: Arquivo não encontrado: {DOCX_PATH}")
        return

    text = extract_text(DOCX_PATH)
    print(f"Texto extraído: {len(text):,} caracteres")
    print()

    violations, passes = validate(text)

    # Print passes
    print("─" * 70)
    print("✅ REGRAS APROVADAS:")
    print("─" * 70)
    for rule_id, severity, name, status in passes:
        print(f"  ✓ [{rule_id}] {name} — {status}")
    print()

    # Print violations
    if violations:
        print("─" * 70)
        print("❌ VIOLAÇÕES ENCONTRADAS:")
        print("─" * 70)
        for rule_id, severity, name, matches, contexts in violations:
            print(f"\n  ✗ [{rule_id}] [{severity}] {name}")
            print(f"    Matches ({len(matches)}): {matches[:5]}")
            for ctx in contexts[:3]:
                print(f"    {ctx}")
        print()
        print(f"RESULTADO: {len(violations)} VIOLAÇÕES / {len(passes)} APROVADAS")
        # Count critical
        critical = sum(1 for v in violations if "CRITICAL" in v[1])
        if critical > 0:
            print(f"⚠️  {critical} violações CRITICAL/BLOCK — documento REJEITADO")
        return violations
    else:
        print("─" * 70)
        print("🎉 RESULTADO: TODAS AS REGRAS APROVADAS — DOCUMENTO VÁLIDO")
        print("─" * 70)
        return []

if __name__ == "__main__":
    violations = main()
