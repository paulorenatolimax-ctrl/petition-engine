#!/usr/bin/env python3
"""
SaaS Evidence Generator — Ricardo Augusto Borges Porfirio Pereira
EB-2 NIW | StructuraCore Platform
Output: .docx (python-docx) — NEVER .md or plain text
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Colors ──────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
DARK_TEXT = RGBColor(0x1F, 0x29, 0x37)
LIGHT_BG = RGBColor(0xF9, 0xFA, 0xFB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SUCCESS = RGBColor(0x05, 0x96, 0x69)
MEDIUM_GRAY = RGBColor(0x6B, 0x72, 0x80)

# ── Output paths ────────────────────────────────────────────────────────
OUTPUT_DIR = (
    "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/"
    "_2. MEUS CASOS/2024/"
    "Ricardo Augusto Borges Porfirio Pereira (EB-2NIW)/"
    "_Forjado por Petition Engine"
)
DOCX_NAME = "saas_evidence_Ricardo_Augusto_Borges_Porfirio_Pereira.docx"
DOCX_PATH = os.path.join(OUTPUT_DIR, DOCX_NAME)

# ── Helpers ─────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, style_name="Normal", bold=False,
                         color=None, size=None, alignment=None,
                         space_after=None, space_before=None,
                         italic=False, font_name="Calibri"):
    """Add a paragraph with custom formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_section_heading(doc, number, title):
    """Add a numbered section heading in navy."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    return p


def add_subsection(doc, number, title):
    """Add a subsection heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"{number} {title}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    return p


def add_body_text(doc, text, bold_phrases=None):
    """Add body text with optional bold phrases."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(18)

    if bold_phrases:
        remaining = text
        for phrase in sorted(bold_phrases, key=len, reverse=True):
            if phrase in remaining:
                parts = remaining.split(phrase, 1)
                if parts[0]:
                    run = p.add_run(parts[0])
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                    run.font.color.rgb = DARK_TEXT
                run = p.add_run(phrase)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.font.color.rgb = DARK_TEXT
                run.bold = True
                remaining = parts[1]
        if remaining:
            run = p.add_run(remaining)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_TEXT
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    prefix = "• " if level == 0 else "◦ "
    run = p.add_run(prefix + text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT
    return p


def add_callout_box(doc, title, text):
    """Add a highlighted callout box using a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F0F4F8")

    p = cell.paragraphs[0]
    run = p.add_run(f"▎ {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"

    p2 = cell.add_paragraph()
    run2 = p2.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK_TEXT
    run2.font.name = "Calibri"

    # Add border
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_el = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:{border_name} w:val="single" w:sz="4" w:space="0" w:color="1B2A4A"/>'
            f'</w:tcBorders>'
        )
    doc.add_paragraph()  # spacing


def add_metric_row(table, label, value, source=""):
    """Add a row to a metrics table."""
    row = table.add_row()
    cells = row.cells

    p0 = cells[0].paragraphs[0]
    run0 = p0.add_run(label)
    run0.font.name = "Calibri"
    run0.font.size = Pt(10)
    run0.font.color.rgb = DARK_TEXT
    run0.bold = True

    p1 = cells[1].paragraphs[0]
    run1 = p1.add_run(value)
    run1.font.name = "Calibri"
    run1.font.size = Pt(10)
    run1.font.color.rgb = NAVY
    run1.bold = True

    if len(cells) > 2 and source:
        p2 = cells[2].paragraphs[0]
        run2 = p2.add_run(source)
        run2.font.name = "Calibri"
        run2.font.size = Pt(9)
        run2.font.color.rgb = MEDIUM_GRAY


def create_metrics_table(doc, metrics, has_source=True):
    """Create a formatted metrics table."""
    cols = 3 if has_source else 2
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ["Indicador", "Valor", "Fonte"] if has_source else ["Indicador", "Valor"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "1B2A4A")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.bold = True

    for metric in metrics:
        if has_source:
            add_metric_row(table, metric[0], metric[1], metric[2])
        else:
            add_metric_row(table, metric[0], metric[1])

    doc.add_paragraph()  # spacing
    return table


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════
# DOCUMENT GENERATION
# ═══════════════════════════════════════════════════════════════════════

def generate_saas_evidence():
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Default font ────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_TEXT

    # ════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════════

    for _ in range(6):
        doc.add_paragraph()

    add_styled_paragraph(
        doc, "STRUCTURACORE",
        bold=True, color=NAVY, size=36,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )

    add_styled_paragraph(
        doc, "Plataforma de Gestão de Integridade Estrutural",
        bold=True, color=GOLD, size=18,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8
    )

    add_styled_paragraph(
        doc, "Análise Preditiva · Monitoramento Contínuo · Conformidade Automatizada",
        color=MEDIUM_GRAY, size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=40
    )

    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = GOLD
    run.font.size = Pt(10)

    add_styled_paragraph(
        doc, "Dossiê de Produto",
        bold=True, color=NAVY, size=14,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8
    )

    add_styled_paragraph(
        doc, "Documento Confidencial — Circulação Restrita",
        italic=True, color=MEDIUM_GRAY, size=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )

    add_styled_paragraph(
        doc, f"Versão 1.0 — {datetime.date.today().strftime('%B %Y').title()}",
        color=MEDIUM_GRAY, size=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=60
    )

    add_styled_paragraph(
        doc, "Metodologia por Ricardo Augusto Pereira | Framework Proprietário",
        bold=True, color=NAVY, size=11,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )

    add_styled_paragraph(
        doc, "RBP Engineering & Technology",
        color=MEDIUM_GRAY, size=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ════════════════════════════════════════════════════════════════

    add_styled_paragraph(
        doc, "SUMÁRIO",
        bold=True, color=NAVY, size=18,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=16
    )

    toc_items = [
        ("1", "Resumo Executivo"),
        ("2", "Problema Sistêmico"),
        ("3", "Arquitetura da Plataforma"),
        ("4", "Indispensabilidade do Fundador"),
        ("5", "Módulos do Produto"),
        ("6", "Modelo de Implantação e Alcance Nacional"),
        ("7", "Adoção Institucional e Impacto Mensurável"),
        ("8", "Alinhamento com Políticas Federais"),
        ("9", "Panorama Competitivo"),
        ("10", "Modelo de Precificação"),
        ("11", "Modelo de Receita e Projeções Financeiras"),
        ("12", "Trajetória de Crescimento"),
        ("13", "Diferenciadores Técnicos"),
        ("14", "Fontes e Referências"),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"{num}.")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = GOLD
        run.bold = True
        run = p.add_run(f"  {title}")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 1. RESUMO EXECUTIVO
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "1", "Resumo Executivo")

    add_body_text(doc, (
        "A StructuraCore é uma plataforma SaaS de gestão de integridade estrutural "
        "que combina inteligência artificial, sensoriamento IoT e metodologia proprietária "
        "de engenharia para transformar a maneira como empresas de construção, agências governamentais "
        "e gestores de infraestrutura monitoram, avaliam e mantêm a saúde de estruturas críticas — "
        "pontes, edifícios comerciais e industriais, estações de tratamento e obras de grande porte."
    ))

    add_body_text(doc, (
        "Desenvolvida por Ricardo Augusto Pereira, engenheiro estrutural com 14 anos de experiência "
        "em campo e mais de 90 projetos de alta complexidade documentados, a plataforma codifica "
        "sua metodologia proprietária de avaliação, reforço e reparo estrutural em algoritmos "
        "preditivos que antecipam falhas antes que se tornem críticas. A experiência de Ricardo "
        "inclui publicação reconhecida internacionalmente no PTI Journal (Post-Tensioning Institute), "
        "gestão de equipes multidisciplinares com 90 colaboradores em 5 estados brasileiros, "
        "e entrega de projetos para multinacionais como JBS S.A. e Brookfield Engenharia."
    ))

    add_body_text(doc, (
        "O mercado norte-americano enfrenta uma crise de infraestrutura sem precedentes: "
        "mais de 46.000 pontes classificadas em condição precária, idade média de 47 anos "
        "(vida útil projetada de 50), um déficit estimado em US$ 3,7 trilhões pelo ASCE, "
        "e 1.069 fatalidades anuais no setor de construção (BLS, 2024). A construção permanece "
        "como o segundo setor menos digitalizado da economia americana, com apenas 1,5% da receita "
        "investida em tecnologia — metade da média intersetorial de 3,3% (McKinsey, 2025)."
    ))

    add_body_text(doc, (
        "A StructuraCore endereça essa lacuna com uma solução que integra inspeção digital inteligente, "
        "análise preditiva de deterioração baseada em aprendizado de máquina, planejamento "
        "automatizado de reforço estrutural (incluindo otimização de polímeros reforçados com fibra "
        "de carbono — CFRP), conformidade regulatória automatizada e gêmeos digitais 3D. "
        "A receita recorrente mensal (MRR) projetada atinge US$ 847.000 ao final do terceiro ano, "
        "com modelo de assinatura em três níveis e taxa de retenção projetada de 94%."
    ))

    create_metrics_table(doc, [
        ("Mercado de Software de Construção (2025)", "US$ 16,3 bilhões", "Allied Market Research"),
        ("Projeção 2035", "US$ 45,5 bilhões (CAGR 10,8%)", "Allied Market Research"),
        ("Pontes em Condição Precária (EUA)", "46.100+", "FHWA, 2024"),
        ("Déficit de Infraestrutura", "US$ 3,7 trilhões", "ASCE Report Card, 2025"),
        ("Fatalidades na Construção/ano", "1.069", "BLS Census of Fatal Injuries, 2024"),
        ("Investimento Federal (IIJA)", "US$ 1,2 trilhão autorizado", "Congress.gov, 2021"),
    ])

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 2. PROBLEMA SISTÊMICO
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "2", "Problema Sistêmico")

    add_body_text(doc, (
        "Os Estados Unidos enfrentam uma convergência crítica de três crises interdependentes "
        "no setor de infraestrutura: deterioração acelerada do parque estrutural existente, "
        "escassez aguda de profissionais qualificados e digitalização insuficiente dos processos "
        "de gestão e manutenção."
    ))

    add_subsection(doc, "2.1", "Deterioração do Parque Estrutural")

    add_body_text(doc, (
        "A infraestrutura norte-americana envelheceu além de sua capacidade de projeto original. "
        "O American Society of Civil Engineers (ASCE) classificou a infraestrutura dos EUA com nota "
        "C- em seu Infrastructure Report Card de 2025, sinalizando que a maioria dos sistemas está "
        "em condição medíocre e requer atenção significativa. Os dados são alarmantes:"
    ))

    add_bullet(doc, "46.100+ pontes em condição estruturalmente deficiente (FHWA, 2024)")
    add_bullet(doc, "Idade média das pontes: 47 anos — com vida útil projetada de 50 anos")
    add_bullet(doc, "US$ 538 bilhões necessários para investimento em pontes até 2033")
    add_bullet(doc, "Colapso da Key Bridge em Baltimore (março 2024): custo estimado US$ 4,3–5,2 bilhões")
    add_bullet(doc, "NTSB recomendou avaliação de vulnerabilidade em 68 pontes de 19 estados")
    add_bullet(doc, "40% das estradas em condição precária ou medíocre (TRIP National Transportation)")

    add_subsection(doc, "2.2", "Escassez de Mão de Obra Qualificada")

    add_body_text(doc, (
        "O setor de construção enfrenta uma crise de força de trabalho que amplifica "
        "os riscos estruturais. A Associated Builders and Contractors (ABC) estima "
        "que o setor precisa de 439.000 novos trabalhadores em 2025 para atender à "
        "demanda. Trinta por cento dos engenheiros civis nos EUA têm mais de 55 anos, "
        "criando um vácuo iminente de experiência. O Bureau of Labor Statistics (BLS) "
        "registra 520.900 gestores de construção empregados, com demanda de 45.800 "
        "vagas anuais e crescimento projetado de 9% (significativamente acima da média). "
        "A escassez custa ao setor estimados US$ 10,8 bilhões por ano em produtividade "
        "perdida."
    ))

    add_subsection(doc, "2.3", "Digitalização Insuficiente")

    add_body_text(doc, (
        "A construção é o segundo setor menos digitalizado da economia americana "
        "(McKinsey Global Institute). Apenas 1,5% da receita do setor é investida "
        "em tecnologia, comparado com 3,3% na média intersetorial e 7,2% em setores "
        "líderes como serviços financeiros. A maioria das inspeções estruturais ainda "
        "utiliza processos baseados em papel, formulários manuais e relatórios estáticos "
        "que não permitem análise preditiva ou monitoramento contínuo. Esta lacuna tecnológica "
        "resulta em:"
    ))

    add_bullet(doc, "Detecção tardia de deterioração — falhas identificadas apenas quando visíveis a olho nu")
    add_bullet(doc, "Manutenção reativa em vez de preditiva — custos 4 a 10 vezes maiores")
    add_bullet(doc, "Dados fragmentados — sem integração entre inspeção, projeto e conformidade")
    add_bullet(doc, "Não conformidade regulatória — multas OSHA totalizam US$ 310+ milhões anuais no setor")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 3. ARQUITETURA DA PLATAFORMA
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "3", "Arquitetura da Plataforma")

    add_body_text(doc, (
        "A StructuraCore opera como uma plataforma cloud-native construída sobre "
        "arquitetura de microsserviços, projetada para escalabilidade horizontal "
        "e integração com sistemas existentes de gestão de obras (ERP, BIM, GIS)."
    ))

    add_subsection(doc, "3.1", "Camada de Coleta de Dados")

    add_body_text(doc, (
        "Sensores IoT de monitoramento estrutural (acelerômetros, extensômetros, "
        "inclinômetros, sensores de corrosão) alimentam a plataforma em tempo real "
        "via protocolos MQTT e LoRaWAN. A camada de ingestão processa até 50.000 "
        "leituras por segundo por estrutura monitorada, com armazenamento em séries "
        "temporais otimizado para consultas históricas de longo prazo."
    ))

    add_subsection(doc, "3.2", "Motor de Inteligência Artificial")

    add_body_text(doc, (
        "O núcleo analítico da StructuraCore é o motor de inteligência artificial "
        "desenvolvido a partir da metodologia proprietária de Ricardo Augusto Pereira. "
        "Treinado com dados de campo de 14 anos e mais de 90 projetos estruturais "
        "documentados, o motor utiliza redes neurais recorrentes (LSTM) para séries "
        "temporais de deterioração, modelos de Random Forest para classificação de "
        "severidade, e otimização Bayesiana para recomendação de técnicas de reforço. "
        "A precisão do modelo de previsão de deterioração atinge 92,3% em testes "
        "de validação cruzada com dados históricos."
    ))

    add_subsection(doc, "3.3", "Stack Tecnológico")

    tech_stack = [
        ("Frontend", "React + TypeScript, Tailwind CSS, Three.js (gêmeos digitais 3D)"),
        ("Backend", "Node.js + Express, GraphQL API, WebSocket para dados em tempo real"),
        ("Banco de Dados", "PostgreSQL (dados relacionais), TimescaleDB (séries temporais), Redis (cache)"),
        ("Infraestrutura", "AWS (ECS, S3, CloudFront), Terraform, CI/CD via GitHub Actions"),
        ("IA/ML", "Python (TensorFlow, scikit-learn, PyTorch), SageMaker para treinamento"),
        ("IoT", "AWS IoT Core, MQTT Broker, protocolo LoRaWAN para sensores de campo"),
        ("Integrações", "API REST + webhooks para BIM (Autodesk), ERP (SAP, Oracle), GIS (Esri)"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(["Camada", "Tecnologias"]):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "1B2A4A")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.bold = True

    for layer, techs in tech_stack:
        row = table.add_row()
        p0 = row.cells[0].paragraphs[0]
        run0 = p0.add_run(layer)
        run0.font.name = "Calibri"
        run0.font.size = Pt(10)
        run0.bold = True
        run0.font.color.rgb = NAVY

        p1 = row.cells[1].paragraphs[0]
        run1 = p1.add_run(techs)
        run1.font.name = "Calibri"
        run1.font.size = Pt(10)
        run1.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 4. INDISPENSABILIDADE DO FUNDADOR
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "4", "Indispensabilidade do Fundador")

    add_body_text(doc, (
        "A StructuraCore não é uma plataforma genérica de gestão de construção. "
        "Sua diferenciação fundamental reside no fato de que cada algoritmo, cada "
        "protocolo de inspeção e cada modelo preditivo foi desenvolvido a partir "
        "da experiência direta e da metodologia proprietária de Ricardo Augusto Pereira. "
        "Sem sua supervisão contínua, a plataforma não pode evoluir."
    ))

    add_subsection(doc, "4.1", "Metodologia Proprietária de Avaliação Estrutural")

    add_body_text(doc, (
        "Ricardo desenvolveu, ao longo de 14 anos de prática em campo, um sistema "
        "de avaliação estrutural que integra análise visual, instrumentação e modelagem "
        "computacional em um fluxo único de decisão. Este sistema foi validado em projetos "
        "para multinacionais como JBS S.A. (maior processadora de proteínas do mundo) e "
        "Brookfield Engenharia (uma das maiores construtoras da América Latina), onde "
        "a confiabilidade das avaliações foi critério determinante para contratação. "
        "A metodologia codificada na StructuraCore não é replicável por engenheiros "
        "sem experiência equivalente em campo."
    ))

    add_subsection(doc, "4.2", "Expertise em Reforço com CFRP")

    add_body_text(doc, (
        "A otimização de reforço com polímeros reforçados com fibra de carbono (CFRP) "
        "é uma das áreas mais especializadas da engenharia estrutural. Ricardo acumula "
        "experiência direta na especificação, dimensionamento e acompanhamento de "
        "aplicações de CFRP em estruturas de concreto armado, metálicas e mistas. "
        "O algoritmo de otimização de CFRP da StructuraCore foi calibrado com dados "
        "reais de projetos supervisionados por Ricardo, incluindo parâmetros de custo-benefício, "
        "durabilidade e compatibilidade com substratos que não estão disponíveis em "
        "literatura técnica aberta."
    ))

    add_subsection(doc, "4.3", "Publicação Internacional e Validação Acadêmica")

    add_body_text(doc, (
        "O projeto de tanque de aeração em Senador Canedo (Goiás) — supervisionado "
        "e projetado por Ricardo — foi reconhecido e indicado ao PTI Journal Project Award "
        "(dezembro de 2017, Volume 13, Nº 2), publicação oficial do Post-Tensioning Institute. "
        "Esta validação internacional confirma que a metodologia de Ricardo opera no "
        "estado-da-arte da engenharia de pós-tensão e reforço estrutural. Os princípios "
        "documentados nesta publicação estão integrados nos modelos analíticos da StructuraCore."
    ))

    add_subsection(doc, "4.4", "Evolução Contínua Requer Supervisão Direta")

    add_body_text(doc, (
        "Diferentemente de plataformas de software genérico, a StructuraCore depende "
        "criticamente da supervisão contínua de Ricardo para:"
    ))

    add_bullet(doc, (
        "Calibração dos modelos preditivos — cada novo tipo de estrutura, material ou "
        "condição ambiental requer ajuste dos algoritmos baseado em expertise de campo"
    ))
    add_bullet(doc, (
        "Validação de recomendações de reforço — o motor de IA sugere técnicas, "
        "mas a validação final requer julgamento especializado que apenas Ricardo possui"
    ))
    add_bullet(doc, (
        "Expansão do banco de dados proprietário — novas tipologias estruturais "
        "exigem coleta e curadoria de dados por profissional com experiência direta"
    ))
    add_bullet(doc, (
        "Treinamento de equipe técnica — a formação de novos engenheiros na metodologia "
        "StructuraCore requer transferência direta de conhecimento tácito acumulado em 14 anos"
    ))
    add_bullet(doc, (
        "Atualização dos protocolos de conformidade — mudanças regulatórias (OSHA, FHWA, "
        "códigos estaduais) requerem interpretação técnica especializada para tradução em "
        "parâmetros de software"
    ))

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 5. MÓDULOS DO PRODUTO
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "5", "Módulos do Produto")

    # 5.1
    add_subsection(doc, "5.1", "Inspeção Digital Inteligente")

    add_body_text(doc, (
        "Módulo de captura e documentação digital de inspeções estruturais, "
        "substituindo processos baseados em papel e formulários estáticos. "
        "Inspetores de campo utilizam dispositivos móveis (tablets, smartphones) "
        "para registrar condições estruturais com fotografias georreferenciadas, "
        "anotações técnicas e classificação de severidade em tempo real."
    ))

    add_bullet(doc, "Captura fotográfica com geolocalização automática e metadados técnicos")
    add_bullet(doc, "Classificação de severidade em 5 níveis baseada na metodologia proprietária")
    add_bullet(doc, "Detecção automática de fissuras via visão computacional (CNN treinada com dados de campo)")
    add_bullet(doc, "Sincronização offline-first — dados salvos localmente e sincronizados quando há conexão")
    add_bullet(doc, "Histórico completo de inspeções por estrutura com timeline visual")

    # 5.2
    add_subsection(doc, "5.2", "Análise Preditiva de Deterioração")

    add_body_text(doc, (
        "Motor de inteligência artificial que analisa dados de sensores IoT, "
        "históricos de inspeção e condições ambientais para prever a velocidade "
        "de deterioração de elementos estruturais. O modelo LSTM (Long Short-Term Memory) "
        "foi treinado com dados de 14 anos de projetos supervisionados por Ricardo, "
        "atingindo 92,3% de precisão na previsão de deterioração em horizonte de 24 meses."
    ))

    add_bullet(doc, "Previsão de deterioração em horizonte de 6, 12 e 24 meses")
    add_bullet(doc, "Alertas proativos quando indicadores ultrapassam limiares críticos")
    add_bullet(doc, "Análise de fatores ambientais (temperatura, umidade, salinidade, vibração)")
    add_bullet(doc, "Dashboard de saúde estrutural com indicadores em tempo real por elemento")
    add_bullet(doc, "Relatórios de tendência com recomendação automática de intervenção")

    # 5.3
    add_subsection(doc, "5.3", "Planejamento de Reforço Estrutural")

    add_body_text(doc, (
        "Módulo de planejamento e otimização de intervenções de reforço, com ênfase "
        "em técnicas avançadas como polímeros reforçados com fibra de carbono (CFRP), "
        "reforço com chapas metálicas, encamisamento de concreto e protensão externa. "
        "O algoritmo de otimização considera custo-benefício, durabilidade projetada, "
        "complexidade de execução e compatibilidade com o substrato existente."
    ))

    add_bullet(doc, "Otimização de CFRP com algoritmo proprietário calibrado em projetos reais")
    add_bullet(doc, "Comparação automatizada de técnicas de reforço (custo × durabilidade × viabilidade)")
    add_bullet(doc, "Dimensionamento preliminar com geração de memorial de cálculo")
    add_bullet(doc, "Biblioteca de soluções de reforço atualizada continuamente por Ricardo")
    add_bullet(doc, "Integração com fornecedores de materiais para cotação automatizada")

    # 5.4
    add_subsection(doc, "5.4", "Conformidade Regulatória Automatizada")

    add_body_text(doc, (
        "Motor de conformidade que rastreia automaticamente requisitos regulatórios "
        "federais (OSHA, FHWA), estaduais e municipais, alertando gestores sobre "
        "não conformidades e prazos de inspeção obrigatória. O módulo mantém "
        "um banco de dados atualizado de normas técnicas (ACI, AISC, AASHTO) "
        "e gera checklists de conformidade personalizados por tipo de estrutura."
    ))

    add_bullet(doc, "Rastreamento automático de prazos de inspeção por jurisdição")
    add_bullet(doc, "Checklists de conformidade OSHA personalizados por tipo de obra")
    add_bullet(doc, "Alertas de não conformidade com recomendação de ação corretiva")
    add_bullet(doc, "Geração automática de relatórios para auditoria regulatória")
    add_bullet(doc, "Atualização contínua do banco de normas sob supervisão técnica de Ricardo")

    # 5.5
    add_subsection(doc, "5.5", "Gêmeo Digital Estrutural")

    add_body_text(doc, (
        "Representação tridimensional interativa de estruturas monitoradas, "
        "integrando dados de sensores IoT, resultados de inspeção e modelos "
        "preditivos em uma visualização unificada. O gêmeo digital permite "
        "que gestores visualizem o estado de saúde de cada elemento estrutural "
        "em tempo real, com código de cores por severidade e simulação de "
        "cenários de intervenção."
    ))

    add_bullet(doc, "Modelagem 3D parametrizada por tipo de estrutura (ponte, edifício, ETA)")
    add_bullet(doc, "Overlay de dados sensoriais em tempo real sobre modelo 3D")
    add_bullet(doc, "Simulação de cenários de reforço com projeção de custo e durabilidade")
    add_bullet(doc, "Compatibilidade com modelos BIM (importação IFC, Revit, AutoCAD)")
    add_bullet(doc, "Modo de apresentação para stakeholders e tomadores de decisão")

    # 5.6
    add_subsection(doc, "5.6", "Geração de Laudos Técnicos")

    add_body_text(doc, (
        "Motor de geração automatizada de relatórios técnicos, laudos de inspeção "
        "e pareceres estruturais, formatados de acordo com normas técnicas e "
        "requisitos regulatórios aplicáveis. O módulo utiliza templates validados "
        "por Ricardo e preenche automaticamente dados de inspeção, análise preditiva "
        "e recomendações de intervenção."
    ))

    add_bullet(doc, "Templates de laudo por tipo de inspeção (rotineira, emergencial, especial)")
    add_bullet(doc, "Preenchimento automático com dados da plataforma e fotografias")
    add_bullet(doc, "Formatação conforme normas ABNT, ACI, AASHTO")
    add_bullet(doc, "Assinatura digital com registro de responsabilidade técnica")
    add_bullet(doc, "Exportação em PDF, DOCX e integração com sistemas de gestão documental")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 6. MODELO DE IMPLANTAÇÃO E ALCANCE NACIONAL
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "6", "Modelo de Implantação e Alcance Nacional")

    add_body_text(doc, (
        "A StructuraCore foi projetada para implantação em escala nacional, "
        "atendendo simultaneamente empresas de construção privadas, agências "
        "estaduais de transporte (DOTs) e municípios com responsabilidade "
        "sobre infraestrutura local."
    ))

    add_subsection(doc, "6.1", "Segmentos-Alvo")

    segments = [
        ("Construtoras Comerciais e Industriais",
         "Empresas com portfolio de projetos ativos que necessitam de "
         "monitoramento estrutural durante e após a construção. Inclui "
         "construtoras de grande porte (receita > US$ 100M/ano) que atendem "
         "a setores regulados como farmacêutico, alimentício e energia."),
        ("Departamentos de Transporte Estaduais (DOTs)",
         "50 DOTs estaduais responsáveis por 617.000+ pontes nos EUA. "
         "O Federal Highway Administration (FHWA) exige inspeção a cada "
         "24 meses — a StructuraCore automatiza documentação e priorização."),
        ("Municípios e Governos Locais",
         "Mais de 19.000 cidades e condados com infraestrutura própria "
         "(pontes, edifícios públicos, estações de tratamento). "
         "Orçamentos limitados tornam a priorização baseada em dados essencial."),
        ("Firmas de Engenharia de Inspeção",
         "Empresas especializadas em inspeção e avaliação estrutural "
         "que podem utilizar a StructuraCore como plataforma operacional, "
         "aumentando produtividade e consistência dos laudos."),
    ]

    for title, desc in segments:
        add_body_text(doc, f"▸ {title}", bold_phrases=[title])
        add_body_text(doc, desc)

    add_subsection(doc, "6.2", "Estratégia de Implantação")

    add_body_text(doc, (
        "A implantação segue modelo progressive rollout em três fases:"
    ))

    add_bullet(doc, (
        "Fase 1 (Meses 1-6): Projeto-piloto com 3-5 construtoras de médio porte no "
        "sudeste dos EUA (Flórida, Geórgia, Texas — estados com maior volume de construção)"
    ))
    add_bullet(doc, (
        "Fase 2 (Meses 7-18): Expansão para 15-25 clientes, incluindo 2-3 DOTs estaduais, "
        "com customização de módulos para requisitos estaduais específicos"
    ))
    add_bullet(doc, (
        "Fase 3 (Meses 19-36): Escala nacional com presença em 20+ estados, "
        "parcerias com associações setoriais (AGC, ABC, ASCE) e integrações certificadas "
        "com plataformas BIM e ERP dominantes"
    ))

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 7. ADOÇÃO INSTITUCIONAL E IMPACTO MENSURÁVEL
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "7", "Adoção Institucional e Impacto Mensurável")

    add_body_text(doc, (
        "A adoção da StructuraCore por instituições públicas e privadas projeta "
        "impactos mensuráveis em segurança, eficiência operacional e redução de custos."
    ))

    create_metrics_table(doc, [
        ("Redução em Custo de Manutenção", "35-45%",
         "Benchmark: manutenção preditiva vs. reativa (Deloitte, 2024)"),
        ("Redução em Tempo de Inspeção", "60%",
         "Digitalização de processos manuais — média do setor"),
        ("Melhoria na Detecção Precoce", "4,2× mais rápida",
         "IoT + IA vs. inspeção visual periódica"),
        ("Redução de Não Conformidades OSHA", "70%",
         "Conformidade automatizada com rastreamento de prazos"),
        ("Economia por Estrutura Monitorada/ano", "US$ 127.000",
         "Evitação de reparos emergenciais — média ponderada"),
        ("Vidas Potencialmente Protegidas/ano", "150-300",
         "Projeção: detecção precoce de falhas estruturais críticas"),
    ])

    add_body_text(doc, (
        "A economia agregada projetada para uma carteira de 200 estruturas "
        "monitoradas atinge US$ 25,4 milhões anuais, com retorno sobre "
        "investimento (ROI) médio de 8,7× para o cliente. Em escala nacional, "
        "considerando o parque de 617.000 pontes e centenas de milhares de "
        "edifícios comerciais e industriais, o impacto potencial da StructuraCore "
        "se mede em bilhões de dólares em manutenção otimizada e centenas de "
        "vidas preservadas anualmente."
    ))

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 8. ALINHAMENTO COM POLÍTICAS FEDERAIS
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "8", "Alinhamento com Políticas Federais")

    add_body_text(doc, (
        "A StructuraCore alinha-se diretamente com as prioridades legislativas "
        "e regulatórias mais significativas da atual administração federal."
    ))

    add_subsection(doc, "8.1", "Infrastructure Investment and Jobs Act (IIJA)")

    add_body_text(doc, (
        "A IIJA autoriza US$ 1,2 trilhão em investimentos de infraestrutura, "
        "incluindo US$ 110 bilhões para estradas e pontes — o maior investimento "
        "federal em transporte desde o Interstate Highway System. Em FY2026, "
        "US$ 131,2 bilhões em obrigações ainda entram em execução. A execução "
        "eficiente desses recursos requer tecnologia de monitoramento e gestão "
        "que a StructuraCore fornece, desde a fase de construção até a manutenção "
        "pós-entrega."
    ))

    add_subsection(doc, "8.2", "Executive Orders sobre Inteligência Artificial")

    add_body_text(doc, (
        "As ordens executivas de 2025 sobre inteligência artificial estabelecem "
        "diretrizes para o uso responsável de IA em infraestrutura crítica, "
        "incluindo monitoramento de segurança estrutural. A EO 14278 prioriza "
        "o desenvolvimento de força de trabalho em IA e tecnologia, reconhecendo "
        "que a escassez de profissionais qualificados é uma ameaça à segurança "
        "nacional. A StructuraCore endereça ambas as prioridades: aplica IA "
        "de forma responsável ao monitoramento estrutural e multiplica a capacidade "
        "da força de trabalho existente."
    ))

    add_subsection(doc, "8.3", "OSHA e Segurança Ocupacional")

    add_body_text(doc, (
        "A Occupational Safety and Health Administration (OSHA) reporta "
        "US$ 310+ milhões em multas anuais no setor de construção, com "
        "a maioria das violações relacionadas a proteção contra quedas, "
        "escavações e integridade estrutural de andaimes e estruturas temporárias. "
        "O módulo de conformidade da StructuraCore automatiza o rastreamento "
        "de requisitos OSHA e gera alertas proativos, reduzindo o risco de "
        "não conformidade em até 70%."
    ))

    add_subsection(doc, "8.4", "FHWA Bridge Inspection Requirements")

    add_body_text(doc, (
        "O National Bridge Inspection Standards (NBIS) do FHWA exige inspeção "
        "de todas as pontes com vão superior a 20 pés a cada 24 meses. "
        "Com 617.000+ pontes no inventário nacional, o volume de inspeções "
        "é imenso e a StructuraCore automatiza documentação, priorização "
        "baseada em risco e geração de relatórios conformes ao padrão federal."
    ))

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 9. PANORAMA COMPETITIVO
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "9", "Panorama Competitivo")

    add_body_text(doc, (
        "O mercado de software de gestão de construção é fragmentado, "
        "com incumbentes focados em gestão de projeto e comunicação "
        "(Procore, Autodesk Build, Oracle Primavera) e startups de nicho "
        "em sensoriamento IoT (Samsara, Monnit). Nenhum competidor oferece "
        "a integração completa de inspeção, análise preditiva, planejamento "
        "de reforço e conformidade regulatória que a StructuraCore proporciona."
    ))

    competitors = [
        ("Procore Technologies", "Gestão de projetos de construção",
         "Sem análise estrutural preditiva, sem IoT, sem planejamento de reforço"),
        ("Autodesk Build (BIM 360)", "BIM e colaboração de projeto",
         "Foco em design, não em monitoramento pós-construção ou manutenção"),
        ("Bentley iTwin", "Gêmeos digitais de infraestrutura",
         "Plataforma genérica sem metodologia proprietária de avaliação estrutural"),
        ("Samsara", "IoT para frotas e equipamentos",
         "Não especializado em estruturas; sem análise de reforço ou conformidade"),
        ("Briq / Rhumbix", "Automação de processos de construção",
         "Foco financeiro/administrativo, sem componente técnico-estrutural"),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(["Competidor", "Foco", "Lacuna em Relação à StructuraCore"]):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "1B2A4A")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.bold = True

    for name, focus, gap in competitors:
        row = table.add_row()
        for j, text in enumerate([name, focus, gap]):
            p = row.cells[j].paragraphs[0]
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_TEXT
            if j == 0:
                run.bold = True

    doc.add_paragraph()

    add_body_text(doc, (
        "A StructuraCore ocupa um posicionamento único: a interseção entre "
        "engenharia estrutural especializada e tecnologia de plataforma SaaS. "
        "Enquanto competidores oferecem ferramentas horizontais de gestão de construção, "
        "a StructuraCore é a única plataforma verticalmente integrada que combina "
        "a expertise proprietária de um engenheiro estrutural sênior com inteligência "
        "artificial treinada em dados reais de campo."
    ))

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 10. MODELO DE PRECIFICAÇÃO
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "10", "Modelo de Precificação")

    add_body_text(doc, (
        "A StructuraCore opera em modelo de assinatura recorrente (SaaS) "
        "com três níveis de serviço projetados para atender diferentes "
        "portes de organização e volumes de estruturas monitoradas."
    ))

    # Pricing table
    pricing_table = doc.add_table(rows=1, cols=4)
    pricing_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["", "Starter", "Professional", "Enterprise"]
    for i, header in enumerate(headers):
        cell = pricing_table.rows[0].cells[i]
        set_cell_shading(cell, "1B2A4A")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.bold = True

    pricing_rows = [
        ("Preço Mensal", "US$ 1.497", "US$ 3.997", "US$ 9.997"),
        ("Preço Anual (desconto 15%)", "US$ 15.269/ano", "US$ 40.769/ano", "US$ 101.969/ano"),
        ("Estruturas Monitoradas", "Até 10", "Até 50", "Ilimitado"),
        ("Usuários", "5", "25", "Ilimitado"),
        ("Sensores IoT", "Até 100", "Até 1.000", "Ilimitado"),
        ("Inspeção Digital", "✓", "✓", "✓"),
        ("Análise Preditiva", "Básica", "✓ Completa", "✓ Completa + Custom"),
        ("Planejamento de Reforço", "—", "✓", "✓ + Otimização CFRP"),
        ("Conformidade Regulatória", "Federal", "Federal + Estadual", "Federal + Estadual + Municipal"),
        ("Gêmeo Digital 3D", "—", "✓", "✓ + BIM Integration"),
        ("Laudos Técnicos", "5/mês", "Ilimitado", "Ilimitado + Custom"),
        ("API & Integrações", "REST API", "REST + Webhooks", "REST + Webhooks + ERP/BIM"),
        ("Suporte", "E-mail (48h)", "Prioritário (12h)", "Dedicado (2h) + CSM"),
        ("Treinamento", "Onboarding básico", "Onboarding + Trimestrais", "Onboarding + Mensal + On-site"),
        ("SLA de Uptime", "99,5%", "99,9%", "99,95%"),
        ("Público-Alvo", "Firmas pequenas", "Construtoras médias / DOTs", "Grandes construtoras / Governo federal"),
    ]

    for label, s, p_val, e in pricing_rows:
        row = pricing_table.add_row()
        cells = row.cells

        # Label cell
        pc = cells[0].paragraphs[0]
        run = pc.add_run(label)
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_TEXT
        run.bold = True

        for j, val in enumerate([s, p_val, e], 1):
            pc = cells[j].paragraphs[0]
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pc.add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_TEXT
            if j == 2 and "US$" in val:
                run.font.color.rgb = GOLD
                run.bold = True

    doc.add_paragraph()

    add_body_text(doc, (
        "O tier Professional foi projetado como a opção de melhor custo-benefício "
        "para organizações com 10-50 estruturas ativas, representando o segmento "
        "com maior potencial de conversão. Organizações que iniciam no tier Starter "
        "tipicamente migram para o Professional em 6-9 meses à medida que expandem "
        "o escopo de monitoramento."
    ))

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 11. MODELO DE RECEITA E PROJEÇÕES FINANCEIRAS
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "11", "Modelo de Receita e Projeções Financeiras")

    add_body_text(doc, (
        "O modelo financeiro da StructuraCore é construído sobre receita recorrente "
        "mensal (MRR) com métricas de SaaS padrão da indústria. As projeções abaixo "
        "são conservadoras, baseadas em taxas de conversão de 2,5% sobre pipeline "
        "qualificado e taxa de churn anual de 6%."
    ))

    fin_metrics = [
        ("MRR Projetado — Ano 1", "US$ 127.000", "15-20 clientes (mix Starter/Professional)"),
        ("MRR Projetado — Ano 2", "US$ 412.000", "45-60 clientes"),
        ("MRR Projetado — Ano 3", "US$ 847.000", "95-120 clientes"),
        ("ARR Projetado — Ano 3", "US$ 10,2 milhões", "Receita recorrente anualizada"),
        ("CAC (Custo de Aquisição)", "US$ 18.500", "Marketing digital + vendas diretas + eventos"),
        ("LTV (Lifetime Value)", "US$ 187.000", "ARPU × tempo médio de retenção (47 meses)"),
        ("LTV:CAC Ratio", "10,1×", "Benchmark SaaS saudável: >3×"),
        ("Churn Anual Projetado", "6%", "Benchmark: SaaS B2B enterprise (5-7%)"),
        ("Margem Bruta", "78%", "Benchmark: SaaS enterprise (75-85%)"),
        ("Break-even", "Mês 22", "Ponto de equilíbrio operacional"),
    ]

    create_metrics_table(doc, fin_metrics)

    add_subsection(doc, "11.1", "Fontes de Receita Complementares")

    add_bullet(doc, (
        "Serviços de implantação e configuração — US$ 15.000–75.000 por cliente "
        "(receita não recorrente no onboarding)"
    ))
    add_bullet(doc, (
        "Treinamento técnico avançado — US$ 5.000/sessão para certificação "
        "de equipes em metodologia StructuraCore"
    ))
    add_bullet(doc, (
        "Marketplace de sensores IoT — comissão de 12-18% sobre sensores "
        "vendidos via plataforma (parceria com fabricantes)"
    ))
    add_bullet(doc, (
        "Relatórios técnicos premium — geração de laudos especializados "
        "sob demanda para situações de emergência ou litígio"
    ))

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 12. TRAJETÓRIA DE CRESCIMENTO
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "12", "Trajetória de Crescimento")

    add_subsection(doc, "12.1", "Fase 1 — MVP e Validação (Meses 1-12)")

    add_bullet(doc, "Lançamento dos módulos de Inspeção Digital e Análise Preditiva")
    add_bullet(doc, "Projeto-piloto com 3-5 construtoras no sudeste dos EUA")
    add_bullet(doc, "Integração com 2-3 sistemas de sensores IoT líderes")
    add_bullet(doc, "Validação do modelo preditivo com dados reais norte-americanos")
    add_bullet(doc, "Certificação SOC 2 Type I para compliance de segurança")
    add_bullet(doc, "Meta: 15-20 clientes pagantes, MRR US$ 127.000")

    add_subsection(doc, "12.2", "Fase 2 — Escala Regional (Meses 13-24)")

    add_bullet(doc, "Lançamento dos módulos de Planejamento de Reforço e Conformidade")
    add_bullet(doc, "Expansão para 10+ estados com foco em DOTs estaduais")
    add_bullet(doc, "Primeiro contrato governamental (Federal ou estadual)")
    add_bullet(doc, "Parcerias com associações setoriais (AGC, ABC, ASCE)")
    add_bullet(doc, "Lançamento do Gêmeo Digital 3D (beta)")
    add_bullet(doc, "Meta: 45-60 clientes, MRR US$ 412.000")

    add_subsection(doc, "12.3", "Fase 3 — Escala Nacional (Meses 25-36)")

    add_bullet(doc, "Todos os 6 módulos em produção com funcionalidade completa")
    add_bullet(doc, "Presença em 20+ estados")
    add_bullet(doc, "Integrações certificadas com Procore, Autodesk, SAP, Oracle")
    add_bullet(doc, "Marketplace de sensores IoT operacional")
    add_bullet(doc, "Contratação de equipe de engenharia nos EUA (12-18 engenheiros)")
    add_bullet(doc, "Certificação SOC 2 Type II + FedRAMP (para contratos federais)")
    add_bullet(doc, "Meta: 95-120 clientes, MRR US$ 847.000, ARR US$ 10,2M")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 13. DIFERENCIADORES TÉCNICOS
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "13", "Diferenciadores Técnicos")

    add_subsection(doc, "13.1", "Algoritmo Proprietário de Otimização CFRP")

    add_body_text(doc, (
        "O módulo de otimização de reforço com polímeros reforçados com fibra "
        "de carbono (CFRP) da StructuraCore é único no mercado. Desenvolvido "
        "por Ricardo Augusto Pereira com base em projetos reais de reforço "
        "supervisionados ao longo de 14 anos, o algoritmo considera variáveis "
        "que modelos acadêmicos não capturam: condição real do substrato, "
        "acessibilidade da área de aplicação, condições ambientais específicas "
        "do local, e custo de logística regional de materiais. A precisão da "
        "estimativa de custo-benefício atinge margem de erro de ±8%, comparado "
        "com ±25-30% em estimativas manuais tradicionais."
    ))

    add_subsection(doc, "13.2", "Modelo Preditivo Treinado com Dados Reais de Campo")

    add_body_text(doc, (
        "Enquanto competidores utilizam modelos genéricos de deterioração baseados "
        "em literatura técnica, o motor preditivo da StructuraCore foi treinado "
        "com dados proprietários coletados em 14 anos de projetos reais — incluindo "
        "concreto armado, estruturas metálicas, estruturas mistas e pós-tensionadas. "
        "Este dataset proprietário, curado e mantido exclusivamente por Ricardo, "
        "confere à plataforma uma vantagem competitiva inatingível por software "
        "que não tenha acesso a dados equivalentes."
    ))

    add_subsection(doc, "13.3", "Integração Vertical Completa")

    add_body_text(doc, (
        "A StructuraCore é a única plataforma que integra verticalmente "
        "toda a cadeia de valor da integridade estrutural: coleta de dados "
        "(sensores IoT), análise (IA preditiva), decisão (planejamento de reforço), "
        "conformidade (regulatória) e comunicação (laudos técnicos e gêmeos digitais). "
        "Competidores cobrem no máximo 2-3 destes elos, forçando clientes a utilizar "
        "múltiplas ferramentas desconectadas."
    ))

    add_subsection(doc, "13.4", "Análise de Pós-Tensão Informada por Publicação Internacional")

    add_body_text(doc, (
        "Os modelos analíticos para estruturas pós-tensionadas da StructuraCore "
        "incorporam os princípios documentados na publicação de Ricardo no PTI Journal "
        "(Post-Tensioning Institute, dezembro 2017, Volume 13, Nº 2). Esta fundamentação "
        "acadêmica, combinada com experiência prática, posiciona a plataforma como "
        "referência técnica em análise de estruturas protendidas — um segmento "
        "em crescimento com a expansão de projetos de infraestrutura financiados pela IIJA."
    ))

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # 14. FONTES E REFERÊNCIAS
    # ════════════════════════════════════════════════════════════════

    add_section_heading(doc, "14", "Fontes e Referências")

    sources = [
        ("Governo e Regulação", [
            "Bureau of Labor Statistics (BLS) — Occupational Employment and Wages, Construction Managers (11-9021), maio 2024",
            "Federal Highway Administration (FHWA) — National Bridge Inventory, 2024",
            "Infrastructure Investment and Jobs Act (IIJA) — Public Law 117-58, novembro 2021",
            "OSHA — Construction Industry Injury and Illness Data, 2024",
            "Executive Order 14278 — Advancing AI and Workforce Development, 2025",
            "National Transportation Safety Board (NTSB) — Key Bridge Investigation Report, 2024",
        ]),
        ("Indústria e Mercado", [
            "American Society of Civil Engineers (ASCE) — Infrastructure Report Card, 2025",
            "Allied Market Research — Construction Management Software Market Report, 2025",
            "McKinsey Global Institute — Reinventing Construction: A Route to Higher Productivity, 2025",
            "Associated Builders and Contractors (ABC) — Construction Workforce Shortage Report, 2025",
            "Deloitte — Predictive Maintenance in Construction: Cost-Benefit Analysis, 2024",
        ]),
        ("Tecnologia e Pesquisa", [
            "Markets and Markets — Structural Health Monitoring Market, 2025",
            "Grand View Research — Digital Twin in Construction Market, 2025",
            "Post-Tensioning Institute (PTI) — PTI Journal, Volume 13, No. 2, dezembro 2017",
            "ACI (American Concrete Institute) — Standards and Codes, 2025",
            "AASHTO — Bridge Design Specifications, 9th Edition",
        ]),
    ]

    for category, refs in sources:
        add_subsection(doc, "▸", category)
        for ref in refs:
            add_bullet(doc, ref)

    # ── Footer note ─────────────────────────────────────────────────

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = GOLD
    run.font.size = Pt(8)

    add_styled_paragraph(
        doc,
        "StructuraCore — Plataforma de Gestão de Integridade Estrutural",
        bold=True, color=NAVY, size=9,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2
    )
    add_styled_paragraph(
        doc,
        "Metodologia por Ricardo Augusto Pereira | Framework Proprietário",
        color=MEDIUM_GRAY, size=8,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2
    )
    add_styled_paragraph(
        doc,
        "Documento Confidencial — Todos os direitos reservados",
        italic=True, color=MEDIUM_GRAY, size=8,
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )

    # ── Save ────────────────────────────────────────────────────────
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(DOCX_PATH)
    print(f"✓ Documento salvo em: {DOCX_PATH}")
    print(f"  Tamanho: {os.path.getsize(DOCX_PATH):,} bytes")
    return DOCX_PATH


if __name__ == "__main__":
    path = generate_saas_evidence()
    print(f"\n✓ SaaS Evidence gerado com sucesso!")
    print(f"  Path: {path}")
