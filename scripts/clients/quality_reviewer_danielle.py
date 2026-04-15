#!/usr/bin/env python3
"""
QUALITY REVIEWER — Verificação Automatizada de Resume V4
Baseado em QUALITY_REVIEWER.md do EB2_NIW_RESUME_SYSTEM
"""
import sys
import os
import re
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ALLOWED_COLORS = {
    "2D3E50", "3498A2", "FFFFFF", "000000",
    "333333", "666666", "F5F5F5", "CCCCCC", "FAFAFA",
}
ALLOWED_FONTS = {"garamond"}
FORBIDDEN_WORDS = ["R$"]

# Error rules from generation prompt
ERROR_RULES_REGEX = [
    (r'\b(I|we)\s+believe\b', "CRITICAL/BLOCK: 'I/we believe' encontrado"),
    (r'\b(I|we)\s+think\b', "HIGH/BLOCK: 'I/we think' encontrado"),
    (r'proposed\s+(venture|business)', "MEDIUM: 'proposed venture/business' em vez de 'proposed endeavor'"),
    (r'\b(in conclusion|to summarize)\b', "HIGH/BLOCK: 'in conclusion/to summarize' encontrado"),
    (r'(23-1011|29-1069|17-201[1-9]|13-2011)', "CRITICAL/BLOCK: SOC code que exige validação de diploma"),
    (r'\bprompt\b', "CRITICAL/BLOCK: Palavra 'prompt' encontrada"),
    (r'(PROEX|Kortix|Carlos Avelino)', "CRITICAL/BLOCK: Referência interna encontrada"),
    (r'\b(introducao|peticao|informacao|certificacao|formacao|avaliacao|ocupacao|operacao|integracao|migracao|capacitacao|micropigmentacao)\b',
     "CRITICAL/BLOCK: Palavra sem acento encontrada"),
    (r'\b(Version \d|Generated:|SaaS Evidence Architect|Petition Engine)\b',
     "CRITICAL/BLOCK: Metadado interno encontrado"),
    (r'\b(standardized|padronizado|operates autonomously|self-sustaining|auto-sustent|plug.and.play|train.the.trainer|white.label|marca branca|client autonomy|founder dependency|scalable without|replicable by any|turnkey|chave.na.m)\b',
     "CRITICAL/BLOCK: Anti-Cristine V2 — termo que destrói Prong 3"),
    (r'\b(consultoria|consulting)\b',
     "HIGH/WARN: 'consultoria/consulting' isolado — flag VIBE/D&B"),
    (r'\b(denial|negativa anterior|RFE anterior|previous petition|prior filing|refile|segunda tentativa|nova submissão|petição anterior|corrected approach|lessons learned from denial)\b',
     "CRITICAL/BLOCK: Referência a histórico processual"),
    (r'\b(equipe jur[ií]dica|advogado|escrit[oó]rio de advocacia|representa[çc][aã]o legal|assessoria jur[ií]dica)\b',
     "CRITICAL/BLOCK: Terminologia jurídica/advocatícia"),
    (r'\b(tribunal|ju[ií]z|senten[çc]a|julgamento|vara|processo judicial)\b',
     "CRITICAL/BLOCK: Terminologia judicial"),
    (r'\b(tradu[çc][aã]o juramentada|tradutor juramentado|tradutor p[uú]blico)\b',
     "CRITICAL/BLOCK: 'Tradução juramentada' em vez de 'tradução certificada'"),
    (r'\b(RAG I|RAG II|RAG III|RAG \d|RAGs|repositório de argumentação)\b',
     "CRITICAL/BLOCK: Referência a RAGs internos"),
    (r'\b(Petition Engine|Forjado por|gerado automaticamente|gerado por (IA|AI|Claude|sistema))\b',
     "CRITICAL/BLOCK: Referência ao sistema de geração"),
    (r'\b(Obsidian|formato \.md|markdown|arquivo \.md)\b',
     "CRITICAL/BLOCK: Referência a ferramentas internas"),
    (r'(Vers[aã]o:? \d|V\d\.\d|Descontaminad|Separation of Concerns|SoC aplicado|Para Revis[aã]o|Rascunho Interno|DOCUMENTO INTERNO)',
     "CRITICAL/BLOCK: Meta-informação de versionamento"),
    (r'\b(proposed endeavor|endeavor\s*(1|2|3)|three endeavor|tres endeavor|opcao de endeavor|endeavor potencial|endeavors? alternativos?)\b',
     "CRITICAL/BLOCK: RESUME não deve ter seção de Proposed Endeavors"),
    (r'\b(revenue|receita bruta|faturamento|projecao financeira|lucro liquido|net income|gross revenue|ROI|EBITDA|fluxo de caixa|cash flow|margem de lucro|profit margin|financial projection|Y1|Y2|Year 1|Year 2)\b',
     "CRITICAL/BLOCK: Dados financeiros no resume"),
    (r'\b(codigo SOC alternativ|SOC code.*(1|2|3)|modelo de receita|publico.alvo.*(B2B|B2C)|quadro.comparativ|pre.projeto|anteprojeto)\b',
     "CRITICAL/BLOCK: Conteúdo de Anteprojeto no resume"),
]


def load_document(path):
    doc = Document(path)
    para_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + "\n"
    merged = para_text + "\n" + table_text
    return doc, merged


def check_fonts(doc):
    issues = []
    arial_count = 0
    other_fonts = set()
    total_runs = 0

    def check_runs(paragraphs):
        nonlocal arial_count, total_runs
        for para in paragraphs:
            for run in para.runs:
                total_runs += 1
                fname = run.font.name
                if fname and fname.lower() not in ALLOWED_FONTS and fname.lower() != "garamond":
                    if "arial" in fname.lower():
                        arial_count += 1
                    other_fonts.add(fname)

    check_runs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                check_runs(cell.paragraphs)

    if arial_count > 0:
        issues.append(("S0", f"ARIAL ENCONTRADO: {arial_count} runs"))
    cleaned = {f for f in other_fonts if f and f.lower() != "garamond"}
    if cleaned:
        issues.append(("S1", f"Fontes não-Garamond: {cleaned}"))
    return issues, total_runs


def check_page_setup(doc):
    issues = []
    section = doc.sections[0]
    w = round(section.page_width.inches, 1)
    h = round(section.page_height.inches, 1)
    if w != 8.5 or h != 11.0:
        issues.append(("S1", f"Paper size ERRADO: {w}\"x{h}\""))
    lm = round(section.left_margin.inches, 2)
    rm = round(section.right_margin.inches, 2)
    if abs(lm - 0.65) > 0.03:
        issues.append(("S1", f"Margem esquerda: {lm}\" (esperado: 0.65\")"))
    if abs(rm - 0.65) > 0.03:
        issues.append(("S1", f"Margem direita: {rm}\" (esperado: 0.65\")"))
    return issues


def check_forbidden_words(merged_text):
    issues = []
    for word in FORBIDDEN_WORDS:
        if word in merged_text:
            issues.append(("S1", f"Palavra proibida: '{word}'"))
    return issues


def check_error_rules(merged_text):
    issues = []
    for pattern, msg in ERROR_RULES_REGEX:
        matches = re.findall(pattern, merged_text, re.IGNORECASE)
        if matches:
            severity = "S0" if "CRITICAL" in msg else "S1" if "HIGH" in msg else "S2"
            sample = matches[0] if isinstance(matches[0], str) else str(matches[0])
            issues.append((severity, f"{msg} — match: '{sample}'"))
    return issues


def check_images(doc):
    issues = []
    img_count = len(doc.inline_shapes)
    if img_count == 0:
        issues.append(("S2", "Zero imagens (thumbnails não inseridos — esperado em V1)"))
    placeholder_count = 0
    for para in doc.paragraphs:
        if "[THUMBNAIL]" in para.text or "[Thumbnail]" in para.text:
            placeholder_count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "[THUMBNAIL]" in cell.text or "[Thumbnail]" in cell.text:
                    placeholder_count += 1
    return issues, img_count, placeholder_count


def check_evidence_blocks(doc):
    issues = []
    tables_with_impact = 0
    tables_2col = 0
    for table in doc.tables:
        if len(table.columns) == 2 and len(table.rows) == 1:
            tables_2col += 1
            left_text = table.rows[0].cells[0].text
            if "Impact" in left_text or "Impacto" in left_text or "Descrição" in left_text:
                tables_with_impact += 1
    if tables_2col > 0 and tables_with_impact == 0:
        issues.append(("S1", f"NENHUM evidence block tem impacto dentro ({tables_2col} encontrados)"))
    return issues, tables_2col, tables_with_impact


def check_header_footer(doc):
    issues = []
    section = doc.sections[0]
    header = section.header
    if not header.paragraphs and not header.tables:
        issues.append(("S1", "HEADER AUSENTE"))
    elif len(header.tables) == 0:
        issues.append(("S2", "Header sem tabela navy"))
    footer = section.footer
    if not footer.paragraphs and not footer.tables:
        issues.append(("S1", "FOOTER AUSENTE"))
    return issues


def check_sections_eb2niw(merged_text):
    issues = []
    upper = merged_text.upper()
    required = [
        ("Síntese Profissional", ["SÍNTESE PROFISSIONAL", "SINTESE PROFISSIONAL", "EXECUTIVE SUMMARY"]),
        ("Histórico Profissional", ["HISTÓRICO PROFISSIONAL", "HISTORICO PROFISSIONAL", "CAREER TIMELINE"]),
        ("Experiência Profissional", ["EXPERIÊNCIA PROFISSIONAL", "EXPERIENCIA PROFISSIONAL"]),
        ("Contribuições Técnicas", ["CONTRIBUIÇÕES TÉCNICAS", "CONTRIBUICOES TECNICAS", "CONTRIBUI"]),
        ("Formação Acadêmica", ["FORMAÇÃO ACADÊMICA", "FORMACAO ACADEMICA"]),
        ("Cartas de Recomendação", ["CARTAS DE RECOMENDAÇÃO", "CARTAS DE RECOMENDACAO"]),
    ]
    for label, needles in required:
        found = any(n in upper for n in needles)
        if not found:
            issues.append(("S1", f"Seção ausente: {label}"))

    # Should NOT have C1-C10
    if "CRITERION 1" in upper or "CRITERION 2" in upper or "CRITÉRIO" in upper:
        issues.append(("S0", "Seções por Critério (EB-1A) encontradas em documento EB-2 NIW"))

    # Check for EB-2 NIW in header area
    if "EB-2 NIW" not in merged_text and "EB-2NIW" not in upper:
        issues.append(("S1", "Referência a 'EB-2 NIW' não encontrada"))

    # Should NOT have Proposed Endeavors section
    if "PROPOSED ENDEAVOR" in upper and ("PROJETO EB-2 NIW" in upper or "PROPOSED ENDEAVORS" in upper):
        issues.append(("S1", "Seção de Proposed Endeavors presente — NÃO deveria estar no resume"))

    return issues


def check_colors_in_xml(doc):
    issues = []
    unexpected = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill and fill.upper() not in ALLOWED_COLORS and fill != "auto":
                            unexpected.add(fill)
    if unexpected:
        issues.append(("S2", f"Cores inesperadas: {unexpected}"))
    return issues


def run_review(docx_path):
    print("=" * 70)
    print(f"QUALITY REVIEWER — {os.path.basename(docx_path)}")
    print(f"Tipo: EB-2 NIW")
    print(f"Tamanho: {os.path.getsize(docx_path):,} bytes")
    print("=" * 70)

    doc, merged = load_document(docx_path)
    all_issues = []

    print("\n[1/9] Verificando fontes...")
    font_issues, total_runs = check_fonts(doc)
    all_issues.extend(font_issues)
    print(f"      {total_runs} runs verificados")

    print("[2/9] Verificando page setup...")
    all_issues.extend(check_page_setup(doc))

    print("[3/9] Verificando palavras proibidas...")
    all_issues.extend(check_forbidden_words(merged))

    print("[4/9] Verificando error rules (27 regras)...")
    all_issues.extend(check_error_rules(merged))

    print("[5/9] Verificando imagens...")
    img_issues, img_count, ph_count = check_images(doc)
    all_issues.extend(img_issues)
    print(f"      {img_count} imagens, {ph_count} placeholders")

    print("[6/9] Verificando evidence blocks...")
    eb_issues, t2col, t_imp = check_evidence_blocks(doc)
    all_issues.extend(eb_issues)
    print(f"      {t2col} evidence blocks, {t_imp} com impacto dentro")

    print("[7/9] Verificando header e footer...")
    all_issues.extend(check_header_footer(doc))

    print("[8/9] Verificando seções obrigatórias EB-2 NIW...")
    all_issues.extend(check_sections_eb2niw(merged))

    print("[9/9] Verificando cores...")
    all_issues.extend(check_colors_in_xml(doc))

    # Report
    print("\n" + "=" * 70)
    print("RELATÓRIO DE QUALIDADE")
    print("=" * 70)
    print(f"\nParágrafos: {len(doc.paragraphs)}")
    print(f"Tabelas: {len(doc.tables)}")
    print(f"Imagens: {img_count}")
    print(f"Evidence Blocks (2-col): {t2col}")
    print(f"  com impacto dentro: {t_imp}")
    print(f"Runs verificados: {total_runs}")

    s0 = [(s, m) for s, m in all_issues if s == "S0"]
    s1 = [(s, m) for s, m in all_issues if s == "S1"]
    s2 = [(s, m) for s, m in all_issues if s == "S2"]
    s3 = [(s, m) for s, m in all_issues if s == "S3"]

    print(f"\n--- ISSUES ---")
    print(f"S0 GRAVÍSSIMO: {len(s0)}")
    print(f"S1 GRAVE:      {len(s1)}")
    print(f"S2 MODERADO:   {len(s2)}")
    print(f"S3 MENOR:      {len(s3)}")

    if s0:
        print(f"\n!!! S0 — GRAVÍSSIMO (BLOQUEIA ENTREGA) !!!")
        for _, msg in s0:
            print(f"  [S0] {msg}")
    if s1:
        print(f"\n!! S1 — GRAVE (BLOQUEIA ENTREGA) !!")
        for _, msg in s1:
            print(f"  [S1] {msg}")
    if s2:
        print(f"\n! S2 — MODERADO !")
        for _, msg in s2:
            print(f"  [S2] {msg}")
    if s3:
        print(f"\nS3 — MENOR")
        for _, msg in s3:
            print(f"  [S3] {msg}")

    print(f"\n{'=' * 70}")
    if s0 or s1:
        print(f"VEREDICTO: REPROVADO — {len(s0)} gravíssimos + {len(s1)} graves")
    elif s2:
        print(f"VEREDICTO: APROVADO COM RESSALVAS — {len(s2)} moderados")
    else:
        print("VEREDICTO: APROVADO — Documento pronto")
    print("=" * 70)

    return len(s0), len(s1), len(s2), len(s3)


if __name__ == "__main__":
    docx_path = "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2025/DANIELLE CEDINI (Refile - EB-2 NIW - DIRETO)/_Forjado por Petition Engine/resume_eb2_niw_DANIELLE_CEDINI.docx"
    s0, s1, s2, s3 = run_review(docx_path)
    sys.exit(1 if (s0 + s1) > 0 else 0)
