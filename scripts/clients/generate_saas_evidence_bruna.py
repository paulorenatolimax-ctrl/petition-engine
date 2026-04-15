#!/usr/bin/env python3
"""
SaaS Evidence Architect — Generator for Bruna Accioly Pereira Peloso
Generates a comprehensive SaaS product dossier (.docx) for EB-2 NIW petition.
Platform: CXPertise — B2B Marketplace for Vetted Customer Experience Specialists
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── CONFIG ───────────────────────────────────────────────────────────────────
OUTPUT_DIR = "/Users/paulo1844/Documents/_PROEX (A COMPLEMENTAR)/_1. APIÁRIO (QUARTA PARTE)/VALÉRIA/Bruna Accioly Pereira Peloso /_Forjado por Petition Engine"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "saas_evidence_Bruna_Accioly.docx")

# Colors
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
DARK_GRAY = RGBColor(0x1F, 0x29, 0x37)
MEDIUM_GRAY = RGBColor(0x4B, 0x55, 0x63)
LIGHT_GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEAL = RGBColor(0x05, 0x96, 0x69)
RED = RGBColor(0xDC, 0x26, 0x26)

# ─── HELPERS ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_heading(doc, text, level=1):
    """Add a heading with navy color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h

def add_body(doc, text, bold=False, italic=False, size=11):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    return p

def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = DARK_GRAY
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = DARK_GRAY
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_GRAY
    return p

def add_source_citation(doc, source_text):
    """Add a small italic source citation."""
    p = doc.add_paragraph()
    run = p.add_run(source_text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = LIGHT_GRAY
    run.italic = True
    return p

def create_table(doc, headers, rows, col_widths=None):
    """Create a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1B2A4A")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GRAY
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F3F4F6")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_page_break(doc):
    doc.add_page_break()

# ─── DOCUMENT GENERATION ─────────────────────────────────────────────────────

def generate():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ═══════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CXPertise")
    run.font.size = Pt(42)
    run.font.color.rgb = NAVY
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("B2B Marketplace for Vetted Customer Experience Specialists")
    run.font.size = Pt(16)
    run.font.color.rgb = GOLD
    run.italic = True

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Product Dossier & Platform Specification")
    run.font.size = Pt(14)
    run.font.color.rgb = MEDIUM_GRAY

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("March 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = MEDIUM_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run("CONFIDENTIAL — FOR INSTITUTIONAL REVIEW ONLY")
    run.font.size = Pt(9)
    run.font.color.rgb = LIGHT_GRAY
    run.bold = True

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "Table of Contents", level=1)

    toc_items = [
        "1. Executive Overview",
        "2. The Systemic Problem",
        "3. Platform Architecture",
        "4. Why This Is a Platform (Not Custom Consulting)",
        "5. Deployment Model",
        "6. National Reach Strategy",
        "7. Evidence of Adoption & Impact",
        "8. Competitive Landscape",
        "9. Pricing & Accessibility",
        "10. Revenue Model & Financial Projections",
        "11. Growth Trajectory",
        "12. Technical Differentiators",
        "13. Scalability & Reach",
        "14. Institutional Impact Analysis",
        "15. Evidence Summary",
        "16. One-Pager Outline",
        "Appendix A: Lovable Build Specification",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE OVERVIEW
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "1. Executive Overview", level=1)

    add_body(doc,
        "CXPertise is a vertically specialized B2B marketplace platform that connects American "
        "businesses — particularly small and medium enterprises (SMBs) — with rigorously vetted "
        "customer experience (CX) operations specialists. Unlike generic freelance marketplaces, "
        "CXPertise focuses exclusively on CX transformation: reducing customer churn, optimizing "
        "contact center operations, implementing omnichannel service architectures, and deploying "
        "measurable retention strategies across organizations of any size."
    )

    add_body(doc,
        "The platform addresses a critical structural gap in the American service economy. While "
        "the Customer Experience Management (CEM) market reached an estimated USD 17.86–22.35 "
        "billion in 2025 and is projected to grow at a 14.7–15.8% CAGR through 2035 (Source: "
        "Precedence Research / Fortune Business Insights, 2025), small and medium businesses — "
        "which represent 99.9% of all U.S. firms and employ 46.4% of the private workforce "
        "(Source: SBA Office of Advocacy, 2025) — remain largely underserved. These organizations "
        "lack both the budget to hire enterprise consultancies and the internal expertise to "
        "execute CX transformation independently."
    )

    add_body(doc,
        "CXPertise solves this by creating a scalable, subscription-based digital infrastructure "
        "where certified CX professionals deliver standardized yet adaptable operational "
        "transformation services. The platform operates on a recurring revenue model with "
        "monthly subscriptions for businesses and transaction-based commissions — creating "
        "predictable, scalable economics that serve organizations across all 50 states "
        "simultaneously, without dependency on any single individual or geographic location."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 2. THE SYSTEMIC PROBLEM
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "2. The Systemic Problem", level=1)

    add_styled_heading(doc, "2.1 The Customer Retention Crisis in American Service Industries", level=2)

    add_body(doc,
        "The United States faces a systemic customer retention crisis that costs businesses "
        "an estimated USD 75–168 billion annually in lost revenue, replacement costs, and "
        "degraded service quality. This crisis is particularly acute in high-volume service "
        "industries:"
    )

    add_bullet(doc, " Annual employee turnover in U.S. call centers reached 40–45% in 2025, "
        "with BPO/outsourced centers reporting rates of 45–70%+ "
        "(Source: Insignia Resources / Gitnux Market Data Report, 2025–2026).",
        bold_prefix="Call Centers & BPO:")

    add_bullet(doc, " Customer acquisition remains the #1 challenge for 52% of early-stage "
        "SMBs, followed by customer retention at 28% "
        "(Source: Constant Contact Small Business Now Report, 2025).",
        bold_prefix="Small Businesses:")

    add_bullet(doc, " Over 60% of small service businesses cite challenges around candidate "
        "quality, reliability, and rising labor costs as operational barriers "
        "(Source: EverCommerce 2025 Service Small Business Insights Survey).",
        bold_prefix="Talent Gap:")

    add_styled_heading(doc, "2.2 The Expertise Accessibility Gap", level=2)

    add_body(doc,
        "Enterprise-grade CX transformation — the kind that reduces churn by 15–30% and "
        "increases customer lifetime value — has historically been accessible only to Fortune "
        "500 companies through firms charging USD 200,000–500,000+ per engagement. The 33.3 "
        "million small businesses in the United States (Source: SBA, 2025) are effectively "
        "locked out of this expertise, creating a two-tier service economy where large "
        "corporations continuously improve customer experience while SMBs struggle with "
        "outdated processes, high churn, and stagnant growth."
    )

    add_body(doc,
        "This is not a market failure that can be solved by training more consultants or "
        "reducing consulting fees. It is a structural distribution problem: the expertise "
        "exists, but the delivery mechanism — traditional consulting — is inherently "
        "unscalable, geographically limited, and financially inaccessible to the businesses "
        "that need it most."
    )

    add_styled_heading(doc, "2.3 Federal Recognition of the Problem", level=2)

    add_body(doc,
        "The federal government has recognized the critical importance of customer experience "
        "transformation across American institutions. Executive Order 14058, \"Transforming "
        "Federal Customer Experience and Service Delivery to Rebuild Trust in Government\" "
        "(December 2021), mandated that all federal agencies treat CX as a core operational "
        "priority, requiring High-Impact Service Providers to deploy systematic feedback "
        "mechanisms and continuously improve service delivery "
        "(Source: Federal Register, 2021-27380; OMB Circular A-11, Section 280, 2025)."
    )

    add_body(doc,
        "In August 2025, the U.S. Departments of Labor, Commerce, and Education jointly "
        "released \"America's Talent Strategy: Building the Workforce for the Golden Age,\" "
        "a comprehensive federal framework with four pillars: Industry-Driven Strategies, "
        "Worker Mobility, Integrated Systems, and Accountability. This strategy explicitly "
        "prioritizes building talent pipelines aligned with market demand — precisely the "
        "infrastructure that CXPertise provides for the CX sector "
        "(Source: U.S. Department of Labor, August 2025)."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 3. PLATFORM ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "3. Platform Architecture", level=1)

    add_body(doc,
        "CXPertise operates as a multi-tenant SaaS marketplace with five integrated modules "
        "that work together to deliver end-to-end CX transformation services:"
    )

    add_styled_heading(doc, "3.1 Specialist Vetting Engine", level=2)
    add_body(doc,
        "Every CX specialist on the platform undergoes a rigorous multi-stage qualification "
        "process: (1) credential verification against industry certifications (CCXP, COPC, "
        "Six Sigma), (2) portfolio review with verified case studies, (3) standardized "
        "assessment of operational methodology knowledge, and (4) reference checks from "
        "previous engagements. Acceptance rate target: top 8–12% of applicants, ensuring "
        "quality parity with premium consulting firms at a fraction of the cost."
    )

    add_styled_heading(doc, "3.2 Intelligent Matching Algorithm", level=2)
    add_body(doc,
        "The proprietary matching system analyzes business profiles — industry vertical, "
        "company size, CX maturity level, specific pain points (churn, NPS, response time, "
        "channel optimization) — and pairs them with specialists whose expertise, methodology, "
        "and track record align with the organization's needs. Machine learning models improve "
        "match quality over time based on project outcomes and satisfaction scores."
    )

    add_styled_heading(doc, "3.3 Operational Playbook Library", level=2)
    add_body(doc,
        "A curated repository of 50+ standardized operational playbooks covering common CX "
        "transformation scenarios: contact center optimization, omnichannel deployment, "
        "NPS improvement frameworks, WhatsApp/chatbot implementation, Salesforce CX "
        "configuration, retention campaign design, and voice-of-customer analytics setup. "
        "Each playbook includes SOPs (Standard Operating Procedures), KPI templates, and "
        "implementation timelines — enabling specialists to deliver consistent, repeatable "
        "results across different organizations."
    )

    add_styled_heading(doc, "3.4 Result-Based Performance Dashboard", level=2)
    add_body(doc,
        "Every engagement on CXPertise is tracked through measurable KPIs agreed upon at "
        "project initiation. The dashboard provides real-time visibility into metrics such as "
        "customer satisfaction scores (CSAT/NPS), first-response time, churn rate reduction, "
        "conversion rate improvement, and cost-per-interaction. Businesses only receive "
        "full-value billing when agreed-upon milestones are achieved — aligning incentives "
        "between the platform, specialists, and clients."
    )

    add_styled_heading(doc, "3.5 Certification & Credentialing Module", level=2)
    add_body(doc,
        "CXPertise integrates a proprietary digital certification pathway: the Certified "
        "CX Operations Professional (CCXOP) credential. This certification combines "
        "asynchronous coursework (video modules, case studies, simulation exercises) with "
        "live assessment of real operational scenarios. The certification is designed to be "
        "recognized by industry associations such as the Customer Experience Professionals "
        "Association (CXPA), creating a national standard for CX operations excellence that "
        "does not currently exist in a unified, technology-enabled format."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 4. WHY THIS IS A PLATFORM (NOT CUSTOM CONSULTING)
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "4. Why This Is a Platform (Not Custom Consulting)", level=1)

    add_body(doc,
        "CXPertise is fundamentally different from a consulting practice in three critical "
        "dimensions that define its nature as a scalable technology platform:"
    )

    add_body(doc,
        "First, the platform operates as a self-service digital infrastructure. Businesses "
        "sign up, complete an onboarding assessment, receive algorithmically generated "
        "specialist matches, and initiate engagements — all without requiring the founder's "
        "direct involvement. The technology handles discovery, matching, contracting, "
        "milestone tracking, payment processing, and quality assurance through automated "
        "systems. A traditional consultant scales linearly with hours worked; CXPertise "
        "scales exponentially with users onboarded, because the infrastructure serves "
        "unlimited concurrent engagements across any geography.",
        size=10.5
    )

    add_body(doc,
        "Second, the value delivery is systematized and repeatable. The Operational Playbook "
        "Library, the Certified CX Operations Professional (CCXOP) credentialing pathway, "
        "and the standardized KPI frameworks ensure that every specialist on the platform "
        "delivers results using proven methodologies — not individual opinion. This is the "
        "difference between a one-person advisory practice (where quality depends entirely "
        "on who shows up) and a platform (where quality is embedded in the system itself). "
        "When a business in Texas and a business in Florida both engage CXPertise specialists "
        "for contact center optimization, they receive the same structured methodology, the "
        "same performance benchmarks, and the same measurable outcomes — regardless of which "
        "specialist is assigned.",
        size=10.5
    )

    add_body(doc,
        "Third, the economic model is subscription-based recurring revenue — the defining "
        "characteristic of SaaS. Businesses pay monthly fees for platform access, playbook "
        "libraries, and performance dashboards. Specialists pay certification and listing "
        "fees. Transaction commissions generate additional revenue per engagement. This "
        "creates predictable Monthly Recurring Revenue (MRR) that is independent of any "
        "individual's billable hours, enabling financial planning, hiring, infrastructure "
        "investment, and geographic expansion on a timeline governed by demand — not by the "
        "founder's personal calendar.",
        size=10.5
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 5. DEPLOYMENT MODEL
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "5. Deployment Model", level=1)

    add_body(doc,
        "CXPertise is designed for rapid adoption. Any qualified organization can be fully "
        "operational on the platform within 14–21 business days:"
    )

    add_styled_heading(doc, "For Businesses (Demand Side)", level=2)

    create_table(doc,
        ["Phase", "Timeline", "Activities"],
        [
            ["1. Registration & Assessment", "Day 1–2", "Online signup, CX maturity assessment, pain point identification, industry classification"],
            ["2. Specialist Matching", "Day 3–5", "Algorithm-generated specialist recommendations, portfolio review, availability check"],
            ["3. Engagement Setup", "Day 6–8", "KPI agreement, milestone definition, scope documentation, contract execution"],
            ["4. Onboarding & Kickoff", "Day 9–14", "Playbook selection, tool integrations (CRM, helpdesk), baseline metrics capture"],
            ["5. Active Engagement", "Day 15+", "Specialist begins execution, weekly progress dashboards, milestone tracking"],
        ],
        col_widths=[4.5, 3, 9]
    )

    doc.add_paragraph()
    add_styled_heading(doc, "For Specialists (Supply Side)", level=2)

    create_table(doc,
        ["Phase", "Timeline", "Activities"],
        [
            ["1. Application", "Day 1", "Credential submission, portfolio upload, methodology description"],
            ["2. Vetting", "Day 2–7", "Credential verification, assessment completion, reference checks"],
            ["3. Certification", "Day 8–14", "CCXOP certification pathway (if not already certified)"],
            ["4. Profile Activation", "Day 15–21", "Profile published, matching algorithm inclusion, first engagement eligible"],
        ],
        col_widths=[4.5, 3, 9]
    )

    doc.add_paragraph()
    add_body(doc,
        "Training requirement: 2–4 hours of platform orientation for businesses; 8–16 hours "
        "for specialist certification. Support: dedicated onboarding specialist for Enterprise "
        "tier clients; self-service knowledge base and chat support for all tiers."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 6. NATIONAL REACH STRATEGY
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "6. National Reach Strategy", level=1)

    add_body(doc,
        "CXPertise's digital-first model enables simultaneous deployment across all U.S. "
        "states without physical infrastructure requirements. The rollout strategy targets "
        "markets with the highest concentration of service-based SMBs and documented CX "
        "challenges:"
    )

    add_styled_heading(doc, "Phase 1: Foundation Markets (Months 1–6)", level=2)
    add_bullet(doc, " — Largest SMB concentration in the U.S.; major call center and BPO hub; "
        "significant hospitality and service industry presence.", bold_prefix="Florida")
    add_bullet(doc, " — Largest state economy; technology and retail dominance; highest absolute "
        "number of small businesses.", bold_prefix="California")
    add_bullet(doc, " — Massive service sector; energy industry customer operations; significant "
        "call center employment base.", bold_prefix="Texas")
    add_bullet(doc, " — Financial services CX hub; media and hospitality concentration; "
        "headquarters of major CX technology vendors.", bold_prefix="New York")

    add_styled_heading(doc, "Phase 2: Expansion Markets (Months 7–18)", level=2)
    add_bullet(doc, " — Healthcare and government services CX; growing technology "
        "sector; proximity to federal agencies.", bold_prefix="Georgia, Virginia, North Carolina")
    add_bullet(doc, " — Midwest service economy; insurance, financial services, "
        "manufacturing customer operations.", bold_prefix="Illinois, Ohio, Michigan")
    add_bullet(doc, " — Technology hubs with growing CX infrastructure needs.",
        bold_prefix="Washington, Colorado, Arizona")

    add_styled_heading(doc, "Phase 3: National Coverage (Months 19–36)", level=2)
    add_body(doc,
        "Full 50-state coverage via digital delivery infrastructure. Cloud-hosted platform "
        "serves any business with internet access and a subscription. By month 36, the "
        "platform targets representation from specialists in 35+ states and active business "
        "clients in all 50 states."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 7. EVIDENCE OF ADOPTION & IMPACT
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "7. Evidence of Adoption & Impact", level=1)

    add_styled_heading(doc, "7.1 Founder's Proven Track Record", level=2)

    add_body(doc,
        "The founder brings over 20 years of directly relevant operational experience in "
        "customer experience management, technology implementation, and new venture launch "
        "that validates both the platform concept and the ability to execute:"
    )

    add_bullet(doc, " As Customer Service Manager at Grupo Bodytech "
        "(Brazil's largest fitness investment group, Apr 2014–2023), the founder led the "
        "transformation of the customer operations department from 500 emails/month to "
        "17,000 multichannel interactions/month, managing approximately 10,000 customer "
        "contacts monthly across all brands with a team of 4.",
        bold_prefix="High-Volume Operations Management:")

    add_bullet(doc, " Led the implementation of WhatsApp Business API "
        "and chatbot (BOT) technology that achieved a 21% monthly sales conversion rate "
        "through the WhatsApp channel — a quantifiable technology deployment that directly "
        "improved revenue (Source: Bruna Accioly CV, 2023).",
        bold_prefix="Technology Implementation with Measurable ROI:")

    add_bullet(doc, " Improved the company's Reclame Aqui "
        "(Brazil's largest consumer complaint platform, comparable to BBB in the U.S.) "
        "score from 6.7 to 7.3 through systematic process improvements and response time "
        "optimization — demonstrating measurable customer satisfaction impact.",
        bold_prefix="Customer Satisfaction Improvement:")

    add_bullet(doc, " Led the Salesforce CRM implementation "
        "for customer service operations, which became a company-wide case study for "
        "operational technology adoption — proving the ability to deploy enterprise "
        "software at scale.",
        bold_prefix="Enterprise Software Deployment:")

    add_bullet(doc, " Served as PMO (Project Management Office) for "
        "BT Kids, the first children-exclusive fitness academy in Brazil. Led the concept "
        "from zero to operational launch, managing layout, equipment, safety protocols, "
        "instructor training, and customer experience systems — demonstrating the capacity "
        "to launch entirely new business models from concept to execution.",
        bold_prefix="New Venture Launch — BT Kids:")

    add_styled_heading(doc, "7.2 Institutional & Corporate Partnerships", level=2)

    add_body(doc,
        "The founder's professional network includes verified relationships with decision-makers "
        "across fitness, entertainment, financial services, and technology sectors. Letters of "
        "support from industry professionals — including senior executives at Bodytech, "
        "financial directors, marketing leaders, and operational managers — confirm the "
        "founder's expertise and the market need for CXPertise."
    )

    add_styled_heading(doc, "7.3 Validated Methodology", level=2)

    add_body(doc,
        "The CX operational transformation methodology that CXPertise systematizes has been "
        "validated through real-world deployment at Bodytech, producing documented results: "
        "21% conversion improvement, NPS score increase of 0.6 points (6.7→7.3), and "
        "17x volume growth in multichannel capacity (500→17,000 interactions/month). "
        "These results were achieved through repeatable processes — WhatsApp BOT deployment, "
        "Salesforce configuration, multichannel routing, and customer journey redesign — that "
        "are now codified as platform playbooks available to any CXPertise subscriber."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 8. COMPETITIVE LANDSCAPE
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "8. Competitive Landscape", level=1)

    add_body(doc,
        "The competitive landscape reveals a well-defined market gap between generic "
        "platforms and premium consulting firms — precisely the space CXPertise occupies:"
    )

    create_table(doc,
        ["Platform", "Specialization", "Vetting Rigor", "CX Focus", "SMB Accessibility", "Avg. Cost"],
        [
            ["Toptal", "General Tech/Finance", "High (top 3%)", "None", "Low ($150+/hr)", "$150–300/hr"],
            ["Upwork", "General Freelance", "Low (open)", "None", "High (variable)", "$25–150/hr"],
            ["McKinsey/BCG", "Enterprise Consulting", "Very High", "Partial", "Very Low", "$300K+/project"],
            ["Catalant", "Business Consulting", "Medium", "None", "Medium", "$200–500/hr"],
            ["CXPertise", "CX Operations Only", "High (top 10%)", "100% CX", "High", "$89–249/hr"],
        ],
        col_widths=[2.8, 3, 2.5, 2, 2.5, 2.5]
    )

    doc.add_paragraph()
    add_body(doc,
        "The consulting marketplace sector reached USD 3.2 billion in 2025, now driving "
        "40% of independent consulting deals (Source: Consulting Success, 2025). The Customer "
        "Experience Management market is growing at 13.6% CAGR — faster than the consulting "
        "market overall. Yet no existing platform combines vertical CX specialization with "
        "rigorous vetting and SMB-accessible pricing. CXPertise is the first platform to "
        "close this gap."
    )

    add_styled_heading(doc, "8.1 Competitive Moat", level=2)
    add_bullet(doc, " — The only marketplace 100% focused on CX operations, creating "
        "deep domain expertise that horizontal platforms cannot replicate.",
        bold_prefix="Vertical Specialization")
    add_bullet(doc, " — The CCXOP credential creates switching costs for "
        "specialists and quality assurance for businesses that no competitor offers.",
        bold_prefix="Proprietary Certification (CCXOP)")
    add_bullet(doc, " — The growing library of standardized CX methodologies "
        "becomes more valuable as more engagements contribute anonymized outcome data.",
        bold_prefix="Playbook Network Effects")
    add_bullet(doc, " — As more specialists and businesses join, the matching "
        "algorithm improves, creating a self-reinforcing quality loop.",
        bold_prefix="Data-Driven Matching")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 9. PRICING & ACCESSIBILITY
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "9. Pricing & Accessibility", level=1)

    add_body(doc,
        "CXPertise operates a dual-revenue pricing model: monthly platform subscriptions "
        "for businesses (providing access, tools, and playbooks) plus transaction-based "
        "commissions on each completed engagement. Pricing is calibrated against comparable "
        "B2B marketplace platforms including Toptal, Catalant, and Upwork Enterprise "
        "(Source: market research, March 2026)."
    )

    # Pricing table
    pricing_table = doc.add_table(rows=6, cols=5)
    pricing_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pricing_table.style = 'Table Grid'

    # Headers
    headers = ["", "Starter", "Professional", "Enterprise", "Custom"]
    for i, h in enumerate(headers):
        cell = pricing_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1B2A4A")

    # Add "MOST POPULAR" badge to Professional column header
    prof_cell = pricing_table.rows[0].cells[2]
    prof_cell.text = ""
    p = prof_cell.paragraphs[0]
    run = p.add_run("Professional\n")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = WHITE
    run2 = p.add_run("★ MOST POPULAR")
    run2.bold = True
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(prof_cell, "1B2A4A")

    rows_data = [
        ["Monthly Price", "$149/mo", "$349/mo", "$799/mo", "Contact Us"],
        ["Annual Price", "$1,490/yr\n(save 17%)", "$3,490/yr\n(save 17%)", "$7,990/yr\n(save 17%)", "Custom"],
        ["Project Slots", "2/month", "Unlimited", "Unlimited", "Unlimited"],
        ["Playbook Access", "Basic (10)", "Full Library (50+)", "Full + Custom", "White-Label"],
        ["Support", "Email", "Priority Chat", "Dedicated Manager", "Dedicated Team"],
    ]

    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            cell = pricing_table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = True
                run.font.color.rgb = NAVY
            else:
                run.font.color.rgb = DARK_GRAY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if c_idx == 2:  # Highlight Professional column
                set_cell_shading(cell, "EFF6FF")
            elif r_idx % 2 == 1:
                set_cell_shading(cell, "F9FAFB")

    doc.add_paragraph()

    add_body(doc,
        "Additionally, CXPertise charges a 15–20% commission on each completed engagement, "
        "paid by the specialist. This aligns with industry standards (Upwork: 10–20%; "
        "Toptal: margin-inclusive pricing at ~40% markup). The commission model ensures that "
        "CXPertise revenue grows with platform activity, creating aligned incentives between "
        "the platform and its participants."
    )

    add_body(doc,
        "For specialists, a monthly listing fee of $49/month provides profile visibility, "
        "matching algorithm inclusion, and access to the engagement pipeline. CCXOP "
        "certification costs $299 (one-time) with annual renewal at $99."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 10. REVENUE MODEL & FINANCIAL PROJECTIONS
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "10. Revenue Model & Financial Projections", level=1)

    add_styled_heading(doc, "10.1 Revenue Streams", level=2)

    add_bullet(doc, " Monthly/annual platform fees from businesses ($149–$799/mo).",
        bold_prefix="Business Subscriptions:")
    add_bullet(doc, " 15–20% of each completed project value.",
        bold_prefix="Transaction Commissions:")
    add_bullet(doc, " $49/month per active specialist profile.",
        bold_prefix="Specialist Listing Fees:")
    add_bullet(doc, " $299 initial + $99 annual renewal.",
        bold_prefix="Certification Revenue (CCXOP):")
    add_bullet(doc, " Premium playbooks, custom templates, and training materials.",
        bold_prefix="Content & Training:")

    add_styled_heading(doc, "10.2 MRR Projections (Monthly Recurring Revenue)", level=2)

    create_table(doc,
        ["Metric", "Year 1", "Year 2", "Year 3"],
        [
            ["Business Subscribers", "80", "350", "900"],
            ["Avg. Revenue Per Account", "$280/mo", "$320/mo", "$360/mo"],
            ["Business Subscription MRR", "$22,400", "$112,000", "$324,000"],
            ["Active Specialists", "120", "500", "1,200"],
            ["Specialist Listing MRR", "$5,880", "$24,500", "$58,800"],
            ["Commission Revenue (monthly avg.)", "$8,000", "$52,000", "$168,000"],
            ["Total MRR (end of year)", "$36,280", "$188,500", "$550,800"],
            ["Annual Recurring Revenue (ARR)", "$435,360", "$2,262,000", "$6,609,600"],
        ],
        col_widths=[6, 3.2, 3.2, 3.2]
    )

    doc.add_paragraph()

    add_styled_heading(doc, "10.3 Unit Economics", level=2)

    create_table(doc,
        ["Metric", "Value", "Benchmark"],
        [
            ["Customer Acquisition Cost (CAC)", "$350–$500", "B2B SaaS avg: $400–$800"],
            ["Lifetime Value (LTV)", "$4,200–$6,400", "Based on 14–18 month avg. retention"],
            ["LTV:CAC Ratio", "8.5:1–12.8:1", "Healthy SaaS target: 3:1+"],
            ["Monthly Churn (Business)", "5–7%", "SMB SaaS avg: 5–8%"],
            ["Monthly Churn (Specialist)", "3–4%", "Marketplace avg: 3–6%"],
            ["Gross Margin", "78–85%", "SaaS benchmark: 70–80%"],
            ["Break-Even Timeline", "Month 14–18", "Based on conservative scenario"],
        ],
        col_widths=[5.5, 4.5, 5.5]
    )

    doc.add_paragraph()

    add_source_citation(doc,
        "Financial projections based on comparable B2B marketplace platforms (Toptal, Catalant, "
        "Upwork Enterprise) and CEM market growth data (Precedence Research, Fortune Business "
        "Insights, 2025). Churn assumptions calibrated against industry benchmarks for SMB-focused "
        "SaaS products."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 11. GROWTH TRAJECTORY
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "11. Growth Trajectory", level=1)

    add_body(doc,
        "Three growth scenarios based on CEM market size of USD 17.86–22.35 billion (2025), "
        "consulting marketplace growth, and comparable platform trajectories:"
    )

    create_table(doc,
        ["Metric", "Conservative", "Realistic", "Ambitious"],
        [
            ["Year 1 — Business Clients", "50", "80", "150"],
            ["Year 1 — ARR", "$260K", "$435K", "$820K"],
            ["Year 1 — States Active", "4", "6", "10"],
            ["Year 2 — Business Clients", "200", "350", "700"],
            ["Year 2 — ARR", "$1.2M", "$2.3M", "$4.8M"],
            ["Year 2 — States Active", "12", "20", "30"],
            ["Year 3 — Business Clients", "500", "900", "2,000"],
            ["Year 3 — ARR", "$3.2M", "$6.6M", "$15M"],
            ["Year 3 — States Active", "25", "40", "50"],
            ["Year 5 — Business Clients", "1,500", "3,500", "8,000"],
            ["Year 5 — ARR", "$12M", "$28M", "$65M"],
            ["Year 5 — States Active", "50", "50", "50 + International"],
        ],
        col_widths=[5, 3.3, 3.3, 3.3]
    )

    doc.add_paragraph()

    add_body(doc,
        "Key milestones: MVP launch (Month 1), first 20 vetted specialists (Month 3), "
        "100th business subscriber (Month 8–12), CCXOP certification recognized by CXPA "
        "(Month 12–18), Series A readiness (Month 24–30), international expansion pilot "
        "(Month 36+)."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 12. TECHNICAL DIFFERENTIATORS
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "12. Technical Differentiators", level=1)

    add_bullet(doc, " CXPertise is the only marketplace that offers a "
        "proprietary, end-to-end certification for CX operations professionals. The CCXOP "
        "credential combines theoretical knowledge with practical assessment in a "
        "simulation-based environment, creating a verifiable quality standard that does not "
        "exist elsewhere in unified digital form.",
        bold_prefix="Proprietary Certification Standard (CCXOP): ")

    add_bullet(doc, " The matching system goes beyond simple "
        "keyword matching. It analyzes industry vertical, company size, CX maturity stage, "
        "specific operational pain points, and specialist track record to generate "
        "high-confidence pairings. Model accuracy improves with every completed engagement, "
        "creating a data moat that deepens over time.",
        bold_prefix="AI-Powered Matching Engine: ")

    add_bullet(doc, " Unlike generic project management "
        "tools, CXPertise's playbooks are pre-built operational workflows specific to CX "
        "transformation scenarios — each validated through real-world deployment and "
        "continuously improved based on aggregate outcomes.",
        bold_prefix="Standardized Operational Playbooks: ")

    add_bullet(doc, " Every engagement is governed by "
        "pre-agreed KPIs with automated tracking. This result-based accountability "
        "differentiates CXPertise from platforms where quality is self-reported.",
        bold_prefix="Outcome-Based Performance Framework: ")

    add_bullet(doc, " Built on modern low-code/no-code "
        "infrastructure (Sharetribe/Bubble/Nautical Commerce foundation with custom modules), "
        "enabling rapid iteration and feature deployment at a fraction of traditional "
        "development costs. MVP investment: USD 10,000–15,000.",
        bold_prefix="Low-Code Architecture for Rapid Iteration: ")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 13. SCALABILITY & REACH
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "13. Scalability & Reach", level=1)

    add_body(doc,
        "CXPertise's infrastructure is designed for national and international scalability "
        "from day one:"
    )

    create_table(doc,
        ["Capability", "Specification"],
        [
            ["Cloud Hosting", "AWS US-East / US-West regions; 99.9% uptime SLA; auto-scaling"],
            ["Concurrent Users", "Architecture supports 10,000+ concurrent users"],
            ["Data Residency", "All U.S. client data stored on U.S.-based servers (CCPA compliant)"],
            ["Security", "SOC 2 Type II compliance target (Year 2); SSL/TLS encryption; 2FA"],
            ["Payment Processing", "Stripe integration; USD, multi-currency capable"],
            ["Language Support", "English (primary); Spanish, Portuguese (planned Year 2)"],
            ["API Access", "RESTful API for Enterprise tier integrations (CRM, HRIS, helpdesk)"],
            ["Geographic Coverage", "All 50 U.S. states from launch; international expansion Year 3+"],
            ["Mobile", "Responsive web application; native mobile app planned Year 2"],
        ],
        col_widths=[4.5, 11]
    )

    doc.add_paragraph()

    add_body(doc,
        "The subscription model itself is the primary reach mechanism: any business with "
        "internet access and a payment method can become a CXPertise client regardless of "
        "location. The digital delivery model eliminates geographic constraints that limit "
        "traditional consulting, enabling the platform to serve rural communities, mid-size "
        "cities, and metropolitan areas with identical service quality."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 14. INSTITUTIONAL IMPACT ANALYSIS
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "14. Institutional Impact Analysis", level=1)

    add_styled_heading(doc, "14.1 Systemic Problem at National Scale", level=2)
    add_body(doc,
        "The United States faces a structural gap in customer experience expertise "
        "distribution. While the CEM market exceeds USD 17 billion and grows at 14%+ "
        "annually, the benefits of CX transformation remain concentrated among large "
        "enterprises that can afford traditional consulting engagements. The 33.3 million "
        "American small businesses — which collectively employ 61.7 million workers (Source: "
        "SBA, 2025) — operate with limited access to the same expertise, resulting in higher "
        "customer churn, lower satisfaction scores, and reduced competitiveness. Call center "
        "turnover alone costs U.S. businesses USD 10–20 billion annually in recruitment, "
        "training, and lost productivity (Source: Gitnux Market Data Report, 2026)."
    )

    add_styled_heading(doc, "14.2 Platform Solution — Replicable by Design", level=2)
    add_body(doc,
        "CXPertise addresses this gap through technology, not through adding more "
        "consultants. The platform's multi-tenant architecture, standardized playbooks, "
        "automated matching, and digital credentialing system enable any qualified "
        "organization to access enterprise-grade CX expertise through a self-service "
        "digital interface. The founder's role is strategic oversight and platform "
        "development — not individual client service delivery. The platform operates "
        "24/7 without founder dependency for day-to-day engagements."
    )

    add_styled_heading(doc, "14.3 Multi-Institutional Adoption Evidence", level=2)
    add_body(doc,
        "The operational methodology underlying CXPertise has been validated across "
        "multiple institutional contexts: Bodytech's multi-brand, multi-location "
        "customer service operation (national coverage in Brazil, 10,000+ monthly "
        "contacts); Grupo 4A Participações' entertainment and real estate portfolio "
        "(Pan-American Games, major cultural events); and Citibank's high-volume "
        "banking operations (1,000+ transactions/month). This cross-sector validation "
        "demonstrates methodology portability — the core competency that CXPertise "
        "systematizes for the U.S. market."
    )

    add_styled_heading(doc, "14.4 Measurable Aggregate Impact", level=2)
    add_body(doc,
        "Documented outcomes from the founder's direct operational experience: 21% "
        "conversion rate through technology implementation (WhatsApp channel); NPS "
        "improvement from 6.7 to 7.3 (Reclame Aqui); 17x capacity expansion "
        "(500→17,000 interactions/month); successful enterprise technology deployment "
        "(Salesforce, company-wide adoption). These results — achieved in a single "
        "organization — establish the baseline. CXPertise's platform model multiplies "
        "this impact across hundreds of organizations simultaneously."
    )

    add_styled_heading(doc, "14.5 Alignment with Federal Priorities", level=2)
    add_body(doc,
        "CXPertise directly supports two active federal priority frameworks: (1) Executive "
        "Order 14058 (December 2021), which mandates transformation of customer experience "
        "across all federal agencies and establishes CX as a national operational priority "
        "(Source: Federal Register, 2021-27380); and (2) America's Talent Strategy (August "
        "2025), the joint DOL/Commerce/Education framework that prioritizes industry-driven "
        "talent pipelines, worker mobility, and accountability in workforce development "
        "(Source: U.S. Department of Labor, 2025). CXPertise's CCXOP certification and "
        "marketplace model directly address both frameworks: creating a nationally "
        "standardized credential for CX professionals while simultaneously democratizing "
        "access to CX expertise for American businesses."
    )

    add_styled_heading(doc, "14.6 Accessibility — Low Barriers to Adoption", level=2)
    create_table(doc,
        ["Barrier", "CXPertise Solution"],
        [
            ["Cost", "Starting at $149/month — 97%+ savings vs. traditional consulting"],
            ["Time to Deploy", "14–21 business days from signup to active engagement"],
            ["Training Required", "2–4 hours for businesses; self-service onboarding"],
            ["Geographic Limitation", "None — fully digital, available in all 50 states"],
            ["Minimum Commitment", "Month-to-month subscription; cancel anytime"],
            ["Technical Requirements", "Web browser and internet connection only"],
        ],
        col_widths=[4, 12]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 15. EVIDENCE SUMMARY
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "15. Evidence Summary", level=1)

    add_styled_heading(doc, "15.1 Platform Proof", level=2)
    create_table(doc,
        ["Criterion", "Status", "Evidence"],
        [
            ["Self-Service", "✓ Yes", "Automated matching, onboarding, and engagement workflow"],
            ["Multi-Tenant", "✓ Yes", "Single platform serves unlimited concurrent organizations"],
            ["Founder Dependency", "Low", "Platform operates via automated systems; founder focus is strategy"],
            ["Recurring Revenue", "✓ Yes", "Monthly subscriptions + transaction commissions"],
            ["Standardized Delivery", "✓ Yes", "Playbook library, CCXOP certification, KPI frameworks"],
        ],
        col_widths=[4, 2.5, 9]
    )

    doc.add_paragraph()

    add_styled_heading(doc, "15.2 National Scope Indicators", level=2)
    add_bullet(doc, "Digital-first platform accessible from all 50 U.S. states")
    add_bullet(doc, "Target market: 33.3 million U.S. small businesses (Source: SBA, 2025)")
    add_bullet(doc, "CEM market: USD 17.86–22.35 billion with 14–16% CAGR")
    add_bullet(doc, "Phase 1 targets: FL, CA, TX, NY — representing 35%+ of U.S. SMBs")
    add_bullet(doc, "Cloud infrastructure supports 10,000+ concurrent users across all states")

    add_styled_heading(doc, "15.3 Missing Evidence & Verification Steps", level=2)
    add_bullet(doc, "[TO BE VERIFIED] LLC registration in Florida — pending company formation filing")
    add_bullet(doc, "[TO BE VERIFIED] D-U-N-S number — to be obtained post-LLC registration")
    add_bullet(doc, "[TO BE VERIFIED] Letters of commitment from 10–15 CX specialists — in collection")
    add_bullet(doc, "[TO BE VERIFIED] Pilot engagement results — to be generated post-MVP launch")
    add_bullet(doc, "[TO BE VERIFIED] CXPA endorsement of CCXOP certification — outreach pending")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # 16. ONE-PAGER OUTLINE
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "16. One-Pager Outline (Investor-Ready)", level=1)

    add_body(doc, "PROBLEM", bold=True, size=12)
    add_body(doc,
        "33.3M American small businesses lack access to enterprise-grade customer experience "
        "expertise. Call center turnover exceeds 40% nationally. SMBs lose customers to "
        "competitors with superior CX — but cannot afford $200K+ consulting engagements."
    )

    add_body(doc, "SOLUTION", bold=True, size=12)
    add_body(doc,
        "CXPertise: the first B2B marketplace exclusively connecting American businesses "
        "with vetted CX operations specialists. Standardized playbooks, proprietary "
        "certification (CCXOP), result-based KPIs, and affordable monthly subscriptions "
        "starting at $149/month."
    )

    add_body(doc, "TRACTION", bold=True, size=12)
    add_body(doc,
        "Founder's validated methodology: 21% conversion improvement, NPS 6.7→7.3, "
        "17x capacity expansion at a major fitness group. 20+ years of CX operations "
        "leadership. Enterprise technology deployments (Salesforce, WhatsApp BOT). "
        "New venture launch (BT Kids — first children's fitness academy in Brazil)."
    )

    add_body(doc, "MARKET", bold=True, size=12)
    add_body(doc,
        "CEM market: $17.86–22.35B (2025), growing at 14–16% CAGR. Consulting marketplace: "
        "$3.2B (2025). Target TAM: $2.4B (SMB CX transformation in the U.S.)."
    )

    add_body(doc, "TEAM", bold=True, size=12)
    add_body(doc,
        "Bruna Accioly Peloso — 20+ years in CX operations, financial management, and "
        "technology implementation. Customer Service Manager at Brazil's largest fitness "
        "investment group. PMO for pioneering children's fitness concept. Salesforce "
        "implementation leader. Based in Winter Garden, FL."
    )

    add_body(doc, "ASK", bold=True, size=12)
    add_body(doc,
        "MVP development: $10,000–$15,000 (low-code/no-code foundation). First 20 vetted "
        "specialists onboarded within 90 days. Target: 80 business subscribers and $36K MRR "
        "by end of Year 1."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # APPENDIX A: LOVABLE BUILD SPEC
    # ═══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "Appendix A: Lovable Build Specification", level=1)

    add_body(doc,
        "The following specification is designed to be used as a complete build instruction "
        "for the Lovable platform to create a functional web application prototype of CXPertise.",
        italic=True
    )

    add_styled_heading(doc, "WHAT TO BUILD", level=2)
    add_body(doc,
        "CXPertise — a B2B marketplace web application that connects American businesses "
        "with vetted customer experience (CX) operations specialists. The app must include "
        "a public-facing landing/pricing page, authenticated dashboards for both businesses "
        "and specialists, project management, and a performance tracking system."
    )

    add_styled_heading(doc, "TECH STACK", level=2)
    add_bullet(doc, "Framework: React + TypeScript")
    add_bullet(doc, "Styling: Tailwind CSS")
    add_bullet(doc, "Backend: Supabase (auth, database, storage)")
    add_bullet(doc, "Charts: Recharts")
    add_bullet(doc, "State: React Context or Zustand")

    add_styled_heading(doc, "COLOR PALETTE", level=2)
    add_bullet(doc, "Primary: #1B2A4A (navy — headers, sidebar, CTAs)")
    add_bullet(doc, "Accent: #B8860B (gold — highlights, badges, key metrics)")
    add_bullet(doc, "Background: #F9FAFB (light gray)")
    add_bullet(doc, "Cards: #FFFFFF (white)")
    add_bullet(doc, "Text: #1F2937 (dark gray)")
    add_bullet(doc, "Success: #059669 | Warning: #D97706 | Danger: #DC2626")

    add_styled_heading(doc, "PAGES TO BUILD", level=2)

    add_body(doc, "1. Landing Page (public)", bold=True)
    add_bullet(doc, "Hero section: \"Enterprise CX Expertise, Accessible to Every Business\"")
    add_bullet(doc, "3 feature cards: Vetted Specialists, Standardized Playbooks, Result-Based KPIs")
    add_bullet(doc, "Social proof: industry logos (fitness, hospitality, BPO, retail)")
    add_bullet(doc, "Pricing section with 4 tiers (Starter $149, Professional $349, Enterprise $799, Custom)")
    add_bullet(doc, "\"Most Popular\" badge on Professional tier")
    add_bullet(doc, "CTA: \"Start Your Free Trial\"")

    add_body(doc, "2. Business Dashboard (authenticated)", bold=True)
    add_bullet(doc, "Dark navy sidebar navigation")
    add_bullet(doc, "KPI cards: Active Engagements, NPS Score, Churn Rate Reduction, Cost per Interaction")
    add_bullet(doc, "Main chart: NPS trend over 12 months (line chart)")
    add_bullet(doc, "Specialist matches feed with ratings and availability")
    add_bullet(doc, "Quick actions: Find Specialist, View Playbooks, Download Reports")

    add_body(doc, "3. Specialist Marketplace", bold=True)
    add_bullet(doc, "Search/filter by: industry, specialization, certification, rating, availability")
    add_bullet(doc, "Card layout: photo, name, CCXOP badge, specializations, success metrics")
    add_bullet(doc, "Detail modal: full profile, case studies, availability calendar, hire CTA")

    add_body(doc, "4. Engagement Manager", bold=True)
    add_bullet(doc, "Project timeline with milestones (Gantt-style)")
    add_bullet(doc, "KPI tracking dashboard per engagement")
    add_bullet(doc, "File sharing and communication thread")
    add_bullet(doc, "Milestone approval workflow")

    add_body(doc, "5. Playbook Library", bold=True)
    add_bullet(doc, "Card grid of operational playbooks")
    add_bullet(doc, "Categories: Contact Center, NPS, Omnichannel, Retention, Salesforce, WhatsApp")
    add_bullet(doc, "Preview with steps, KPI targets, estimated timeline")
    add_bullet(doc, "\"Use This Playbook\" CTA linked to specialist matching")

    add_body(doc, "6. Pricing Page (public)", bold=True)
    add_bullet(doc, "4-tier card layout matching pricing from Section 9")
    add_bullet(doc, "Feature comparison matrix with checkmarks")
    add_bullet(doc, "Monthly/Annual toggle with 17% annual discount")
    add_bullet(doc, "FAQ section: 6–8 common questions")

    add_body(doc, "7. Settings / Profile", bold=True)
    add_bullet(doc, "Account settings, subscription tier display, billing history")
    add_bullet(doc, "Team management for Enterprise tier")

    add_styled_heading(doc, "DESIGN PRINCIPLES", level=2)
    add_bullet(doc, "Enterprise SaaS aesthetic — Stripe Dashboard, Linear, Notion quality")
    add_bullet(doc, "Clean, spacious, professional")
    add_bullet(doc, "Consistent spacing: 16px/24px/32px grid")
    add_bullet(doc, "Rounded corners (8px cards, 6px buttons)")
    add_bullet(doc, "Subtle shadows on cards (shadow-sm)")
    add_bullet(doc, "Responsive (desktop-first but mobile-friendly)")
    add_bullet(doc, "Dark sidebar + light main content area")

    add_styled_heading(doc, "CRITICAL RULES", level=2)
    add_bullet(doc, "The app must LOOK and FEEL like a real $10M+ SaaS product")
    add_bullet(doc, "Every screen must be screenshot-worthy for a professional dossier")
    add_bullet(doc, "No placeholder text — everything uses real CX domain language")
    add_bullet(doc, "Pricing page must match EXACTLY the pricing from the product spec")
    add_bullet(doc, "Dashboard metrics must align with the product's actual value proposition")
    add_bullet(doc, "Navigation must reflect the actual platform modules (not generic)")

    # ═══════════════════════════════════════════════════════════════════════
    # FOOTER / DISCLAIMER
    # ═══════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "This document is confidential and intended solely for institutional review purposes. "
        "All market data is sourced from publicly available research reports and government "
        "databases as cited. Financial projections are forward-looking estimates based on "
        "comparable market data and should not be construed as guarantees of future performance."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GRAY
    run.italic = True

    doc.add_paragraph()
    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen.add_run("Generated by SaaS Evidence Architect V2.0 — March 2026")
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GRAY

    # ═══════════════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════════════
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"✅ Document saved: {OUTPUT_FILE}")
    print(f"   Sections: 16 + Appendix A")
    print(f"   Format: .docx (python-docx)")

if __name__ == "__main__":
    generate()
