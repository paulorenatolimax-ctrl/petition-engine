#!/usr/bin/env python3
"""
Résumé EB-1A Generator — Márcio Elias Barbosa (O-1)
Petition Engine v2.0 | python-docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ── CONSTANTS ──────────────────────────────────────────────────────────
NAVY = RGBColor(0x2D, 0x3E, 0x50)
NAVY_HEX = "2D3E50"
TEAL = RGBColor(0x34, 0x98, 0xA2)
TEAL_HEX = "3498A2"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY_HEX = "F5F5F5"
BORDER_GRAY_HEX = "CCCCCC"

FONT_NAME = "Garamond"
MARGIN_LEFT = Inches(0.65)
MARGIN_RIGHT = Inches(0.65)
MARGIN_TOP = Inches(0)
MARGIN_BOTTOM = Inches(0.5)

LEFT_COL_WIDTH = 5760  # DXA
RIGHT_COL_WIDTH = 4320  # DXA

CLIENT_NAME = "Márcio Elias Barbosa"
VISA_TYPE = "O-1"
SOC_CODE = "11-3011.00"
SOC_TITLE = "Administrative Services Managers"

OUTPUT_DIR = "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2026/MARCIO ELIAS BARBOSA /4. O-1 Consultoras GIU e Karina Mariano/_Forjado por Petition Engine"
OUTPUT_FILE = "resume_eb1a_Márcio_Elias_Barbosa_(O1)_.docx"


# ── HELPERS ────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_row_height(row, height_pt):
    """Set row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def add_formatted_run(paragraph, text, size=Pt(10.5), color=BLACK, bold=False, italic=False, font_name=FONT_NAME):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def add_paragraph_text(doc, text, size=Pt(10.5), color=BLACK, bold=False, italic=False,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6), space_before=Pt(0)):
    """Add a formatted paragraph to the document."""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    add_formatted_run(para, text, size=size, color=color, bold=bold, italic=italic)
    return para


def add_navy_section_header(doc, title):
    """Add a full-width Navy section header bar."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY_HEX)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    add_formatted_run(para, title.upper(), size=Pt(11), color=WHITE, bold=True)
    # Set table width to full page
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="0" w:type="auto"/>')
    tblPr.append(tblW)
    # Remove default borders
    for cell_obj in table.row_cells(0):
        set_cell_border(cell_obj,
            top={"val": "single", "sz": "0", "color": NAVY_HEX},
            bottom={"val": "single", "sz": "0", "color": NAVY_HEX},
            start={"val": "single", "sz": "0", "color": NAVY_HEX},
            end={"val": "single", "sz": "0", "color": NAVY_HEX})
    doc.add_paragraph()  # spacer


def add_teal_subheader(doc, title):
    """Add a Teal sub-header bar."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, TEAL_HEX)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    add_formatted_run(para, title, size=Pt(10), color=WHITE, bold=True)
    for cell_obj in table.row_cells(0):
        set_cell_border(cell_obj,
            top={"val": "single", "sz": "0", "color": TEAL_HEX},
            bottom={"val": "single", "sz": "0", "color": TEAL_HEX},
            start={"val": "single", "sz": "0", "color": TEAL_HEX},
            end={"val": "single", "sz": "0", "color": TEAL_HEX})


def add_company_box(doc, company_name, role, period, location, bullets):
    """Add institutional company box with gray background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY_HEX)
    set_cell_border(cell,
        top={"val": "single", "sz": "4", "color": BORDER_GRAY_HEX},
        bottom={"val": "single", "sz": "4", "color": BORDER_GRAY_HEX},
        start={"val": "single", "sz": "4", "color": BORDER_GRAY_HEX},
        end={"val": "single", "sz": "4", "color": BORDER_GRAY_HEX})

    # Company name
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(2)
    add_formatted_run(p1, company_name, size=Pt(10.5), color=NAVY, bold=True)

    # Role
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    add_formatted_run(p2, role, size=Pt(10), color=DARK_GRAY, bold=False, italic=True)

    # Period + Location
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(4)
    add_formatted_run(p3, f"{period} | {location}", size=Pt(9.5), color=MED_GRAY, italic=True)

    # Bullets
    for bullet in bullets:
        pb = cell.add_paragraph()
        pb.paragraph_format.space_before = Pt(1)
        pb.paragraph_format.space_after = Pt(1)
        add_formatted_run(pb, f"• {bullet}", size=Pt(9.5), color=DARK_GRAY)

    doc.add_paragraph()  # spacer


def add_evidence_block(doc, exhibit_num, title, evidence_type, date, institution, location, description, impact):
    """Add a 2-column evidence block with metadata left, thumbnail placeholder right."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for cell in table.columns[0].cells:
        cell.width = Emu(LEFT_COL_WIDTH * 635)  # DXA to EMU approx
    for cell in table.columns[1].cells:
        cell.width = Emu(RIGHT_COL_WIDTH * 635)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    # Borders
    for cell in [left_cell, right_cell]:
        set_cell_border(cell,
            top={"val": "single", "sz": "4", "color": BORDER_GRAY_HEX},
            bottom={"val": "single", "sz": "4", "color": BORDER_GRAY_HEX},
            start={"val": "single", "sz": "4", "color": BORDER_GRAY_HEX},
            end={"val": "single", "sz": "4", "color": BORDER_GRAY_HEX})

    # LEFT: Metadata
    p_title = left_cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(2)
    add_formatted_run(p_title, f"Exhibit {exhibit_num}: ", size=Pt(9.5), color=NAVY, bold=True)
    add_formatted_run(p_title, title, size=Pt(9.5), color=NAVY, bold=True)

    # Metadata lines
    metadata_lines = [
        ("Tipo: ", evidence_type),
        ("Data: ", date),
        ("Instituição: ", institution),
        ("Local: ", location),
    ]
    for label, value in metadata_lines:
        p = left_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        add_formatted_run(p, label, size=Pt(9), color=MED_GRAY, bold=True)
        add_formatted_run(p, value, size=Pt(9), color=DARK_GRAY)

    # Description
    p_desc = left_cell.add_paragraph()
    p_desc.paragraph_format.space_before = Pt(4)
    p_desc.paragraph_format.space_after = Pt(2)
    add_formatted_run(p_desc, "Descrição: ", size=Pt(9), color=MED_GRAY, bold=True)
    add_formatted_run(p_desc, description, size=Pt(9.5), color=DARK_GRAY)

    # Impact (INSIDE the block per spec)
    p_impact = left_cell.add_paragraph()
    p_impact.paragraph_format.space_before = Pt(2)
    p_impact.paragraph_format.space_after = Pt(4)
    add_formatted_run(p_impact, "Impacto: ", size=Pt(9), color=NAVY, bold=True)
    add_formatted_run(p_impact, impact, size=Pt(9.5), color=DARK_GRAY)

    # RIGHT: Thumbnail placeholder
    p_thumb = right_cell.paragraphs[0]
    p_thumb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_thumb.paragraph_format.space_before = Pt(20)
    p_thumb.paragraph_format.space_after = Pt(20)
    set_cell_shading(right_cell, LIGHT_GRAY_HEX)
    add_formatted_run(p_thumb, f"[THUMBNAIL — Exhibit {exhibit_num}]",
                      size=Pt(9), color=MED_GRAY, italic=True)

    doc.add_paragraph()  # spacer after block


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


# ── DOCUMENT CREATION ─────────────────────────────────────────────────

def create_resume():
    doc = Document()

    # ── PAGE SETUP ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM

    # ── Set default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(10.5)
    font.color.rgb = BLACK
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # ══════════════════════════════════════════════════════════════════
    # HEADER TABLE — Navy 3 rows x 2 columns
    # ══════════════════════════════════════════════════════════════════
    header_table = doc.add_table(rows=3, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Shade all cells Navy
    for row in header_table.rows:
        for cell in row.cells:
            set_cell_shading(cell, NAVY_HEX)
            set_cell_border(cell,
                top={"val": "single", "sz": "0", "color": NAVY_HEX},
                bottom={"val": "single", "sz": "0", "color": NAVY_HEX},
                start={"val": "single", "sz": "0", "color": NAVY_HEX},
                end={"val": "single", "sz": "0", "color": NAVY_HEX})

    # Row 0: Name
    r0_left = header_table.cell(0, 0)
    r0_right = header_table.cell(0, 1)
    r0_left.merge(r0_right)
    p_name = r0_left.paragraphs[0]
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(12)
    p_name.paragraph_format.space_after = Pt(4)
    add_formatted_run(p_name, CLIENT_NAME.upper(), size=Pt(20), color=WHITE, bold=True)

    # Row 1: RÉSUMÉ | Visa Type
    r1_left = header_table.cell(1, 0)
    r1_right = header_table.cell(1, 1)
    p_resume = r1_left.paragraphs[0]
    p_resume.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_resume.paragraph_format.space_before = Pt(2)
    p_resume.paragraph_format.space_after = Pt(2)
    add_formatted_run(p_resume, "RÉSUMÉ", size=Pt(11), color=WHITE, bold=False)

    p_visa = r1_right.paragraphs[0]
    p_visa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_visa.paragraph_format.space_before = Pt(2)
    p_visa.paragraph_format.space_after = Pt(2)
    add_formatted_run(p_visa, f"Petition Type: {VISA_TYPE}", size=Pt(11), color=WHITE, bold=False)

    # Row 2: SOC Code
    r2_left = header_table.cell(2, 0)
    r2_right = header_table.cell(2, 1)
    r2_left.merge(r2_right)
    p_soc = r2_left.paragraphs[0]
    p_soc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_soc.paragraph_format.space_before = Pt(2)
    p_soc.paragraph_format.space_after = Pt(8)
    add_formatted_run(p_soc, f"SOC: {SOC_CODE} — {SOC_TITLE}", size=Pt(9), color=WHITE)

    doc.add_paragraph()  # spacer

    # ══════════════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════
    add_navy_section_header(doc, "RESUMO EXECUTIVO")

    exec_paragraphs = [
        "Márcio Elias Barbosa é especialista brasileiro em gestão estratégica de operações de grande escala, "
        "com formação em Direito pela Universidade São Francisco e mais de 23 anos de experiência na coordenação "
        "de equipes, padronização de procedimentos operacionais, acompanhamento de indicadores de desempenho e "
        "otimização de resultados em ambientes corporativos de alta complexidade e elevado volume de demandas.",

        "Ao longo de sua trajetória profissional, atuou na liderança estratégica de operações voltadas à prestação "
        "de serviços especializados para grandes instituições financeiras, incluindo o Banco Itaú Unibanco — maior "
        "banco privado do Brasil e uma das maiores instituições financeiras da América Latina — e o Banco Santander. "
        "Durante mais de 15 anos, foi o responsável exclusivo pela condução estratégica da carteira de processos "
        "do Banco Itaú no Estado de São Paulo, abrangendo 644 cidades e 153 Varas do Trabalho, com aproximadamente "
        "3.000 processos ativos simultaneamente e mais de 10.000 processos ao longo de todo o período.",

        "Sua atuação combinava liderança operacional direta — coordenando uma equipe de 21 profissionais internos, "
        "incluindo 13 especialistas, além de uma rede de mais de 100 prestadores de serviço distribuídos em diversas "
        "regiões do estado — com o desenvolvimento de metodologia proprietária para engenharia de dados e análise "
        "estratégica de operações em larga escala. Essa metodologia, baseada em seis pilares estruturais, transformou "
        "a rotina operacional em uma base analítica capaz de gerar indicadores estratégicos para tomada de decisão.",

        "Os resultados alcançados sob sua liderança foram consistentemente reconhecidos pelo Banco Itaú Unibanco "
        "no âmbito de seu programa nacional de avaliação de desempenho, que envolvia entre 36 e 40 escritórios "
        "especializados em todo o Brasil. Em 2019, a operação liderada por Márcio Elias Barbosa alcançou o topo "
        "do ranking nacional, com 112,8% de desempenho no eixo de acordos, 107% no eixo de resultados e 109,9% "
        "de desempenho geral — a melhor performance entre todos os participantes naquele ciclo.",

        "Paralelamente à atuação operacional, Márcio Elias Barbosa contribuiu para o campo do conhecimento por meio "
        "da publicação de dois artigos acadêmicos na revista Promestre — Journal of Management and Innovation (ISSN "
        "3086-1454), periódico com política formal de revisão por pares (double-blind peer review) e registro "
        "internacional DOI. Também atuou como revisor de artigos submetidos por outros profissionais à mesma "
        "publicação, demonstrando reconhecimento de sua expertise pela comunidade acadêmica.",

        "Sua remuneração durante o período analisado — variando entre R$ 20.000 e R$ 33.646 mensais — situava-se "
        "substancialmente acima da média observada para profissionais da área no Brasil, refletindo o nível de "
        "responsabilidade e a complexidade das funções exercidas na condução de operações de grande porte para "
        "instituições de alcance nacional.",
    ]

    for text in exec_paragraphs:
        add_paragraph_text(doc, text, size=Pt(10.5), color=BLACK)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # HISTÓRICO PROFISSIONAL
    # ══════════════════════════════════════════════════════════════════
    add_navy_section_header(doc, "HISTÓRICO PROFISSIONAL")

    add_company_box(doc,
        company_name="Advocacia Elias Barbosa / Conti, Silva e Barbosa",
        role="Sócio-Proprietário e Gestor Estratégico",
        period="1999 – 2023",
        location="São Paulo, SP — Brasil",
        bullets=[
            "Liderança exclusiva da carteira de processos do Banco Itaú Unibanco no Estado de São Paulo por mais de 15 anos",
            "Gestão de aproximadamente 3.000 processos simultâneos, totalizando mais de 10.000 ao longo do período",
            "Cobertura territorial de 644 cidades e 153 Varas do Trabalho (99,84% das cidades do estado)",
            "Coordenação de 21 profissionais internos (13 especialistas) e rede de 100+ prestadores externos",
            "Desenvolvimento de metodologia proprietária para engenharia de dados e gestão estratégica de operações em larga escala",
            "Múltiplas premiações de alta performance no programa nacional de avaliação do Banco Itaú (2012–2022)",
            "Melhor performance nacional entre todos os escritórios avaliados (2019): 109,9% de desempenho geral",
            "Incorporação da carteira do Banco Santander ao portfólio operacional a partir da formação de Conti, Silva e Barbosa",
        ]
    )

    add_company_box(doc,
        company_name="Simple Shift",
        role="Fundador e Especialista em Gestão Estratégica",
        period="2023 – Presente",
        location="Brasil / Estados Unidos",
        bullets=[
            "Aplicação e difusão da metodologia proprietária de engenharia de dados e gestão estratégica",
            "Prestação de serviços técnicos especializados em análise de indicadores e otimização de performance",
            "Prestação de serviços técnicos especializados para organizações que administram grandes volumes de operações complexas",
            "Treinamentos e workshops técnicos para profissionais de diferentes áreas e regiões",
        ]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # CRITÉRIO 1 — PREMIAÇÕES E RECONHECIMENTOS
    # ══════════════════════════════════════════════════════════════════
    add_navy_section_header(doc, "CRITÉRIO 1 — PREMIAÇÕES E RECONHECIMENTOS POR ALTA PERFORMANCE PROFISSIONAL")

    c1_intro = [
        "As premiações e reconhecimentos documentados nesta seção foram concedidos no âmbito do programa nacional "
        "de avaliação de desempenho do Banco Itaú Unibanco, no qual participavam entre 36 e 40 escritórios "
        "especializados distribuídos em diferentes regiões do Brasil. Todos os escritórios participantes haviam "
        "passado por rigoroso processo de seleção para representar o banco, constituindo um grupo altamente "
        "qualificado de profissionais.",

        "O sistema de avaliação considerava indicadores objetivos de desempenho: aproximadamente 80% baseado em "
        "resultados operacionais diretos, 10% em critérios de eficiência processual e 10% em critérios "
        "comportamentais. Somente escritórios que atingiam 100% ou mais das metas estabelecidas eram elegíveis "
        "para premiação, e as metas eram revisadas periodicamente com aumento de aproximadamente 10% ao ano, "
        "elevando continuamente o nível de exigência.",

        "Os escritórios participantes eram classificados em rankings nacionais comparativos, incluindo um gráfico "
        "de desempenho baseado em dois eixos principais — percentual de resultados positivos e desempenho em "
        "acordos e resultados financeiros. Os escritórios posicionados no chamado 'Quadrante Verde' eram aqueles "
        "que apresentavam alto desempenho simultâneo em ambos os indicadores.",
    ]
    for text in c1_intro:
        add_paragraph_text(doc, text)

    add_teal_subheader(doc, "Premiações Documentadas (2012–2022)")
    doc.add_paragraph()

    # Awards list
    awards = [
        ("2012", "Premiação de desempenho — 1º ciclo semestral"),
        ("2012", "Premiação de desempenho — 2º ciclo semestral"),
        ("2013", "Premiação anual de performance"),
        ("2015", "Premiação de desempenho — 2º ciclo semestral"),
        ("2016", "Prêmio de Alta Performance"),
        ("2018", "Reconhecimento — 3º ciclo quadrimestral"),
        ("2019", "Melhor performance entre todos os escritórios do Brasil — Desempenho geral: 109,9%"),
        ("2019", "Premiação de desempenho — 2º ciclo semestral"),
        ("2020", "Reconhecimento em campanha extraordinária de acordos"),
        ("2022", "Desempenho Diferenciado — 1º ciclo semestral (Conti, Silva e Barbosa)"),
    ]
    for year, desc in awards:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(1)
        add_formatted_run(p, f"  {year} — ", size=Pt(10), color=NAVY, bold=True)
        add_formatted_run(p, desc, size=Pt(10), color=DARK_GRAY)

    doc.add_paragraph()

    # Evidence blocks — Criterion 1
    add_teal_subheader(doc, "Evidências Documentais — Critério 1")
    doc.add_paragraph()

    add_evidence_block(doc,
        exhibit_num=1,
        title="Captura de Tela do Sistema — Quantidade de Processos Administrados",
        evidence_type="Captura de sistema interno",
        date="Período: 1999–2023",
        institution="Advocacia Elias Barbosa",
        location="São Paulo, SP — Brasil",
        description="Registro do sistema interno demonstrando o volume de processos administrados pelo escritório "
                    "sob a liderança de Márcio Elias Barbosa, evidenciando a escala e complexidade da operação gerenciada "
                    "ao longo de mais de duas décadas de atuação na representação do Banco Itaú Unibanco.",
        impact="Comprova a dimensão excepcional da carteira de processos — aproximadamente 3.000 ativos "
               "simultaneamente e mais de 10.000 ao longo do período — posicionando a operação entre as maiores "
               "do segmento no Estado de São Paulo."
    )

    add_evidence_block(doc,
        exhibit_num=2,
        title="Manual do Programa de Avaliação de Desempenho — Banco Itaú (2013)",
        evidence_type="Documento institucional",
        date="2013",
        institution="Banco Itaú Unibanco S.A.",
        location="São Paulo, SP — Brasil",
        description="Manual institucional detalhando a estrutura e critérios do programa nacional de avaliação "
                    "de desempenho de escritórios externos do Banco Itaú, incluindo metodologia de cálculo, "
                    "indicadores avaliados e sistema de classificação em rankings comparativos.",
        impact="Demonstra a existência de um sistema formal, estruturado e competitivo de avaliação que envolveu "
               "entre 36 e 40 escritórios em todo o Brasil, validando o significado das premiações recebidas."
    )

    add_evidence_block(doc,
        exhibit_num=3,
        title="Manual do Programa de Avaliação de Desempenho — Banco Itaú (2018)",
        evidence_type="Documento institucional",
        date="2018",
        institution="Banco Itaú Unibanco S.A.",
        location="São Paulo, SP — Brasil",
        description="Versão atualizada do manual do programa de avaliação, refletindo a evolução dos critérios "
                    "e o aumento progressivo das metas de desempenho (aproximadamente 10% ao ano), demonstrando "
                    "a crescente exigência do programa ao longo dos anos.",
        impact="Evidencia que as premiações foram obtidas em um ambiente de crescente competitividade e exigência, "
               "onde os critérios de avaliação eram continuamente elevados."
    )

    add_evidence_block(doc,
        exhibit_num=4,
        title="Comunicação de Desempenho — Resultados do Ciclo 2019",
        evidence_type="E-mail institucional",
        date="2019",
        institution="Banco Itaú Unibanco S.A.",
        location="São Paulo, SP — Brasil",
        description="Comunicação oficial do banco confirmando os resultados alcançados pela Advocacia Elias Barbosa "
                    "no ciclo avaliativo de 2019, com 112,8% no eixo de acordos, 107% no eixo de resultados "
                    "operacionais e 109,9% de desempenho geral.",
        impact="Comprova que o escritório liderado por Márcio Elias Barbosa alcançou a melhor performance nacional "
               "entre todos os escritórios participantes no ciclo avaliativo de 2019, representando o topo "
               "do ranking em um programa que envolvia dezenas de escritórios altamente qualificados."
    )

    add_evidence_block(doc,
        exhibit_num=5,
        title="Gráfico de Ranking Nacional — Quadrante Verde (2019)",
        evidence_type="Gráfico comparativo institucional",
        date="2019",
        institution="Banco Itaú Unibanco S.A.",
        location="São Paulo, SP — Brasil",
        description="Gráfico de desempenho baseado em dois eixos — percentual de resultados positivos e performance "
                    "em acordos — demonstrando o posicionamento de cada escritório participante. A Advocacia Elias "
                    "Barbosa aparece posicionada no Quadrante Verde, zona de alto desempenho simultâneo em ambos os eixos.",
        impact="Demonstra visualmente a posição de destaque do escritório no topo do ranking nacional, acima de "
               "todos os demais escritórios participantes do programa de avaliação."
    )

    add_page_break(doc)

    add_evidence_block(doc,
        exhibit_num=6,
        title="Comunicação de Desempenho — 1º Ciclo Semestral (2012)",
        evidence_type="E-mail institucional",
        date="2012",
        institution="Banco Itaú Unibanco S.A.",
        location="São Paulo, SP — Brasil",
        description="Comunicação oficial confirmando o reconhecimento de desempenho da Advocacia Elias Barbosa "
                    "no primeiro ciclo avaliativo de 2012, representando o início de uma série contínua de "
                    "premiações que se estenderia por uma década.",
        impact="Evidencia o padrão consistente de alta performance desde o início do período documentado, "
               "demonstrando que a excelência operacional não foi um evento isolado, mas um resultado sustentado."
    )

    add_evidence_block(doc,
        exhibit_num=7,
        title="Comunicação de Desempenho — 2º Ciclo Semestral (2012)",
        evidence_type="E-mail institucional",
        date="2012",
        institution="Banco Itaú Unibanco S.A.",
        location="São Paulo, SP — Brasil",
        description="Segunda premiação recebida no mesmo ano de 2012, confirmando resultados consistentes "
                    "em ciclos avaliativos distintos dentro do mesmo período.",
        impact="Comprova a capacidade de manter alto desempenho em múltiplos ciclos consecutivos, "
               "reforçando o padrão de excelência sustentada."
    )

    add_evidence_block(doc,
        exhibit_num=8,
        title="Comunicação de Desempenho — 2º Ciclo Semestral (2015)",
        evidence_type="E-mail institucional",
        date="2015",
        institution="Banco Itaú Unibanco S.A.",
        location="São Paulo, SP — Brasil",
        description="Reconhecimento de desempenho no ciclo avaliativo de 2015, demonstrando a continuidade "
                    "dos resultados de alta performance após três anos consecutivos de avaliações positivas.",
        impact="Reforça a trajetória consistente de reconhecimento em um programa com metas crescentes "
               "(+10% ao ano), evidenciando capacidade de melhoria contínua."
    )

    add_evidence_block(doc,
        exhibit_num=9,
        title="Prêmio de Alta Performance — 2016",
        evidence_type="E-mail institucional",
        date="2016",
        institution="Banco Itaú Unibanco S.A.",
        location="São Paulo, SP — Brasil",
        description="Comunicação formal de premiação por alta performance no ciclo avaliativo de 2016, "
                    "período em que a carteira do banco continuava em expansão com crescente volume de demandas.",
        impact="Demonstra a manutenção de resultados excepcionais mesmo diante do aumento de complexidade "
               "e volume operacional da carteira."
    )

    add_evidence_block(doc,
        exhibit_num=10,
        title="Reconhecimento — 3º Ciclo Quadrimestral (2018)",
        evidence_type="E-mail institucional",
        date="2018",
        institution="Banco Itaú Unibanco S.A.",
        location="São Paulo, SP — Brasil",
        description="Reconhecimento de desempenho no terceiro ciclo quadrimestral de 2018, período em que "
                    "o programa de avaliação já operava com critérios significativamente mais rigorosos "
                    "em comparação aos anos iniciais.",
        impact="Evidencia a capacidade de alcançar resultados de destaque mesmo após múltiplas elevações "
               "consecutivas nos critérios de avaliação do programa."
    )

    add_evidence_block(doc,
        exhibit_num=11,
        title="Reconhecimento — Campanha Extraordinária de Acordos (2020)",
        evidence_type="E-mail institucional",
        date="2020",
        institution="Banco Itaú Unibanco S.A.",
        location="São Paulo, SP — Brasil",
        description="Reconhecimento especial por resultados alcançados durante campanha extraordinária de "
                    "acordos em 2020, período marcado por desafios operacionais decorrentes da pandemia de COVID-19.",
        impact="Demonstra a capacidade de adaptação e manutenção de resultados excepcionais mesmo em "
               "condições adversas e atípicas, reforçando a resiliência e eficácia da gestão estratégica."
    )

    add_teal_subheader(doc, "Registros Fotográficos de Premiações")
    doc.add_paragraph()

    photo_exhibits = [
        (12, "Troféu — Premiação 2º Ciclo Semestral (2015)", "2015",
         "Registro fotográfico do troféu recebido pela Advocacia Elias Barbosa no 2º ciclo semestral de 2015."),
        (13, "Troféu — Premiação 2º Ciclo Semestral (2016)", "2016",
         "Registro fotográfico do troféu recebido no 2º ciclo semestral de 2016."),
        (14, "Equipe na Cerimônia de Premiação (2016)", "2016",
         "Registro fotográfico da equipe durante a cerimônia de premiação do Banco Itaú em 2016."),
        (15, "Cerimônia de Premiação (2018)", "2018",
         "Registro fotográfico da cerimônia de premiação de 2018 do programa nacional de avaliação."),
        (16, "Troféu — 3º Ciclo (2019)", "2019",
         "Registro fotográfico do troféu correspondente ao melhor desempenho nacional em 2019."),
        (17, "Coleção de Troféus (2015, 2016, 2019)", "2015–2019",
         "Registro fotográfico demonstrando a coleção de troféus acumulados ao longo de múltiplos ciclos de premiação."),
        (18, "Troféu — Desempenho Diferenciado (2022)", "2022",
         "Troféu recebido por Conti, Silva e Barbosa (evolução de Advocacia Elias Barbosa) pelo desempenho diferenciado."),
    ]

    for ex_num, title, date, desc in photo_exhibits:
        add_evidence_block(doc,
            exhibit_num=ex_num,
            title=title,
            evidence_type="Registro fotográfico",
            date=date,
            institution="Banco Itaú Unibanco S.A.",
            location="São Paulo, SP — Brasil",
            description=desc,
            impact="Comprova materialmente a existência e entrega das premiações referidas nas comunicações "
                   "institucionais, validando o reconhecimento formal por desempenho de excelência."
        )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # CRITÉRIO 4 — AVALIAÇÃO DO TRABALHO DE TERCEIROS
    # ══════════════════════════════════════════════════════════════════
    add_navy_section_header(doc, "CRITÉRIO 4 — PARTICIPAÇÃO COMO AVALIADOR DO TRABALHO DE OUTROS PROFISSIONAIS")

    c4_intro = [
        "Ao longo de sua trajetória profissional, Márcio Elias Barbosa exerceu papel recorrente de avaliação técnica "
        "do trabalho de outros profissionais, especialmente no contexto da gestão estratégica de carteiras de operações "
        "de grande volume. Essa atuação envolve a análise detalhada da condução técnica de processos por profissionais "
        "responsáveis por diferentes etapas da operação, com o objetivo de identificar padrões de desempenho, "
        "oportunidades de melhoria e ajustes estratégicos necessários para aprimorar os resultados.",

        "Essa atividade exige elevado grau de conhecimento técnico e experiência profissional, pois envolve a análise "
        "crítica da atuação de outros profissionais com base em critérios operacionais, processuais e estratégicos. "
        "Adicionalmente, Márcio Elias Barbosa atuou como revisor de artigos acadêmicos submetidos para publicação "
        "em periódico científico com política formal de revisão por pares.",
    ]
    for text in c4_intro:
        add_paragraph_text(doc, text)

    add_teal_subheader(doc, "Avaliação Técnica da Atuação de Profissionais em Campo")
    doc.add_paragraph()

    add_paragraph_text(doc,
        "No contexto da prestação de serviços estratégicos a organizações que administram grandes volumes de operações, "
        "Márcio Elias Barbosa realiza análises detalhadas da atuação dos profissionais responsáveis pela condução dos "
        "processos. Essas análises incluem a avaliação da condução de estratégias processuais, utilização adequada de "
        "recursos e provas, qualidade técnica da atuação e identificação de falhas procedimentais que possam impactar "
        "os resultados. Os relatórios técnicos produzidos apresentam avaliações específicas sobre a atuação profissional "
        "dos responsáveis pelos casos, com observações sobre questões processuais relevantes."
    )

    add_evidence_block(doc,
        exhibit_num=19,
        title="Relatório de Avaliação de Desempenho de Profissionais",
        evidence_type="Relatório técnico",
        date="Outubro de 2025",
        institution="Simple Shift",
        location="Brasil",
        description="Relatório técnico detalhado contendo análise da atuação de profissionais responsáveis "
                    "pela condução de processos, incluindo avaliações individuais de desempenho, identificação "
                    "de falhas procedimentais e recomendações estratégicas para melhoria de resultados.",
        impact="Demonstra a atuação contínua de Márcio Elias Barbosa como avaliador do trabalho de outros "
               "profissionais, exercendo função que exige expertise diferenciada e reconhecimento de autoridade "
               "técnica no campo de atuação."
    )

    add_teal_subheader(doc, "Avaliação de Desempenho com Base em Indicadores Estratégicos")
    doc.add_paragraph()

    add_paragraph_text(doc,
        "Além da análise individual de casos, Márcio Elias Barbosa realiza avaliações amplas da performance das "
        "equipes responsáveis pela condução das operações. Essas análises utilizam indicadores objetivos como "
        "percentual de resultados positivos e negativos, evolução temporal dos resultados, identificação de padrões "
        "recorrentes e comparativo entre diferentes profissionais e regiões. A partir dessas análises, são elaboradas "
        "apresentações estratégicas destinadas a orientar ajustes na condução dos processos e aprimorar o desempenho."
    )

    add_evidence_block(doc,
        exhibit_num=20,
        title="Apresentação de Resultados Estratégicos",
        evidence_type="Apresentação técnica",
        date="Outubro de 2023",
        institution="Simple Shift / Operação Santander",
        location="São Paulo, SP — Brasil",
        description="Apresentação contendo análise comparativa de resultados operacionais, indicadores de "
                    "desempenho e diagnóstico estratégico da performance de equipes profissionais.",
        impact="Evidencia a capacidade de avaliar sistematicamente o trabalho de múltiplos profissionais "
               "com base em dados objetivos e indicadores mensuráveis."
    )

    add_evidence_block(doc,
        exhibit_num=21,
        title="Relatório de Resultados Operacionais",
        evidence_type="Relatório analítico",
        date="Outubro de 2025",
        institution="Simple Shift",
        location="Brasil",
        description="Relatório consolidado demonstrando indicadores de desempenho operacional, análise de "
                    "resultados por região e por profissional, e diagnóstico de oportunidades de melhoria.",
        impact="Comprova a aplicação da metodologia de avaliação em contexto atual, demonstrando continuidade "
               "e evolução da atividade de análise do trabalho de terceiros."
    )

    add_teal_subheader(doc, "Workshops e Reuniões de Feedback Profissional")
    doc.add_paragraph()

    add_paragraph_text(doc,
        "Os resultados das análises são apresentados em reuniões estratégicas e workshops com as equipes profissionais. "
        "Nessas ocasiões, Márcio Elias Barbosa apresenta dados consolidados, identifica fatores que contribuíram para "
        "os resultados e orienta ajustes na atuação dos profissionais. Esses encontros funcionam como momentos "
        "estruturados de avaliação técnica e feedback profissional, nos quais os profissionais compreendem como sua "
        "atuação impacta os resultados e adotam melhorias em sua abordagem estratégica."
    )

    add_evidence_block(doc,
        exhibit_num=22,
        title="Convite para Workshop de Feedback — Novembro 2018",
        evidence_type="E-mail institucional",
        date="Novembro de 2018",
        institution="Advocacia Elias Barbosa / Banco Itaú",
        location="São Paulo, SP — Brasil",
        description="Convite formal para workshop de feedback e orientação profissional destinado aos "
                    "profissionais responsáveis pela condução dos processos da carteira do Banco Itaú.",
        impact="Comprova a realização de eventos estruturados de avaliação e feedback, validando o papel "
               "de Márcio Elias Barbosa como avaliador do trabalho de outros profissionais."
    )

    add_evidence_block(doc,
        exhibit_num=23,
        title="Apresentação do Workshop de Feedback (2018)",
        evidence_type="Apresentação técnica",
        date="Novembro de 2018",
        institution="Advocacia Elias Barbosa",
        location="São Paulo, SP — Brasil",
        description="Material de apresentação utilizado durante o workshop, contendo dados de desempenho, "
                    "análise de resultados e orientações técnicas para os profissionais.",
        impact="Demonstra o conteúdo técnico e analítico apresentado como base para a avaliação dos profissionais."
    )

    add_evidence_block(doc,
        exhibit_num=24,
        title="Material de Treinamento Técnico — Julho 2023",
        evidence_type="Apresentação técnica",
        date="Julho de 2023",
        institution="Simple Shift",
        location="São Paulo, SP — Brasil",
        description="Material de treinamento técnico voltado à orientação de profissionais com base em "
                    "indicadores de desempenho e análise de resultados operacionais.",
        impact="Evidencia a continuidade da atividade de avaliação e orientação profissional em contexto "
               "atualizado, demonstrando que essa função permanece central na atuação de Márcio Elias Barbosa."
    )

    add_teal_subheader(doc, "Revisão de Artigos Acadêmicos (Peer Review)")
    doc.add_paragraph()

    add_paragraph_text(doc,
        "Além das atividades de avaliação no contexto profissional, Márcio Elias Barbosa atuou como revisor "
        "de artigos acadêmicos submetidos para publicação na revista Promestre — Journal of Management and "
        "Innovation. Essa atividade caracteriza-se como avaliação independente do trabalho de outros profissionais, "
        "realizada fora do ambiente interno de sua atuação, reforçando o reconhecimento de sua expertise pela "
        "comunidade acadêmica. Os artigos analisados foram posteriormente aprovados para publicação, evidenciando "
        "a relevância de sua participação no processo editorial."
    )

    add_evidence_block(doc,
        exhibit_num=25,
        title="Declaração Institucional sobre Rigor Editorial — Promestre",
        evidence_type="Declaração institucional",
        date="2025",
        institution="Promestre — Journal of Management and Innovation",
        location="Brasil",
        description="Declaração da editoria da revista Promestre detalhando a política formal de revisão por pares "
                    "(double-blind peer review), registro ISSN, indexação acadêmica e conformidade com diretrizes "
                    "do Committee on Publication Ethics (COPE).",
        impact="Valida a natureza acadêmica e o rigor do processo editorial no qual Márcio Elias Barbosa "
               "participou como revisor, demonstrando que se trata de publicação científica reconhecida."
    )

    add_evidence_block(doc,
        exhibit_num=26,
        title="Convite Formal — Processo de Avaliação por Pares",
        evidence_type="Convite institucional",
        date="2025",
        institution="Promestre — Journal of Management and Innovation",
        location="Brasil",
        description="Convite formal emitido pela editoria da revista Promestre para que Márcio Elias Barbosa "
                    "atuasse como revisor de artigos submetidos por outros profissionais.",
        impact="Comprova o reconhecimento formal de sua expertise pela comunidade acadêmica, demonstrando "
               "que foi selecionado especificamente para avaliar o trabalho de outros pesquisadores."
    )

    add_evidence_block(doc,
        exhibit_num=27,
        title="Formulário de Avaliação de Artigo — Assinado",
        evidence_type="Documento de avaliação acadêmica",
        date="2025",
        institution="Promestre — Journal of Management and Innovation",
        location="Brasil",
        description="Formulário de avaliação preenchido e assinado por Márcio Elias Barbosa como revisor "
                    "de artigo acadêmico submetido à revista, contendo análise técnica do conteúdo.",
        impact="Documenta a execução efetiva da atividade de revisão por pares, demonstrando avaliação "
               "criteriosa e fundamentada do trabalho de outro profissional."
    )

    add_evidence_block(doc,
        exhibit_num=28,
        title="Declaração de Atuação como Revisor",
        evidence_type="Declaração institucional",
        date="2025",
        institution="Promestre — Journal of Management and Innovation",
        location="Brasil",
        description="Declaração formal emitida pela revista Promestre confirmando a participação de Márcio "
                    "Elias Barbosa como revisor de artigos no processo editorial da publicação.",
        impact="Confirma oficialmente a atuação como avaliador do trabalho de outros profissionais no "
               "âmbito acadêmico, complementando as evidências de avaliação no contexto profissional."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # CRITÉRIO 5 — CONTRIBUIÇÕES ORIGINAIS
    # ══════════════════════════════════════════════════════════════════
    add_navy_section_header(doc, "CRITÉRIO 5 — CONTRIBUIÇÕES ORIGINAIS DE GRANDE RELEVÂNCIA NA ÁREA PROFISSIONAL")

    c5_intro = [
        "Ao longo de sua atuação profissional na gestão de grandes carteiras de operações, Márcio Elias Barbosa "
        "desenvolveu uma metodologia própria voltada à organização do fluxo de trabalho operacional, à estruturação "
        "de dados processuais e à análise estratégica de operações de grande escala. Essa metodologia foi concebida "
        "para lidar com ambientes caracterizados por elevado volume de processos simultâneos, dispersão geográfica "
        "das demandas, necessidade de padronização operacional e monitoramento contínuo de resultados.",

        "A contribuição original reside no fato de que a metodologia transforma a rotina operacional em uma base "
        "estruturada de dados analíticos, permitindo que informações tradicionalmente dispersas em processos "
        "individuais sejam convertidas em indicadores estratégicos capazes de apoiar a tomada de decisões gerenciais "
        "e operacionais. A metodologia é sustentada por seis pilares estruturais que atuam de forma integrada.",
    ]
    for text in c5_intro:
        add_paragraph_text(doc, text)

    add_teal_subheader(doc, "Os Seis Pilares da Metodologia Proprietária")
    doc.add_paragraph()

    pillars = [
        ("1. Pipeline de Processos", "Organização estruturada do fluxo operacional, desde a entrada das demandas até "
         "a conclusão, com registro estruturado de cada etapa e evento relevante."),
        ("2. Algoritmo de Aderência", "Sistema de verificação contínua do cumprimento das diretrizes estratégicas "
         "estabelecidas, permitindo identificar desvios e aplicar correções em tempo real."),
        ("3. Taxonomia de Dados", "Classificação sistemática das informações processuais em categorias analíticas "
         "que permitem comparação, agregação e identificação de padrões."),
        ("4. Modelagem de Performance", "Construção de indicadores quantitativos que medem o desempenho da operação "
         "em múltiplas dimensões — resultados, eficiência, qualidade e cumprimento de metas."),
        ("5. Monitoramento Dinâmico", "Acompanhamento em tempo real da evolução dos indicadores, permitindo "
         "intervenções proativas antes que desvios se consolidem em resultados negativos."),
        ("6. Modelo de Causalidade", "Identificação das relações causais entre variáveis operacionais e resultados, "
         "permitindo direcionar recursos e esforços para os fatores de maior impacto."),
    ]
    for title, desc in pillars:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_formatted_run(p, title, size=Pt(10), color=NAVY, bold=True)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.space_before = Pt(0)
        add_formatted_run(p2, desc, size=Pt(10), color=DARK_GRAY)

    add_teal_subheader(doc, "Evidências Documentais — Critério 5")
    doc.add_paragraph()

    add_evidence_block(doc,
        exhibit_num=29,
        title="Documento da Metodologia Estratégica de Dados",
        evidence_type="Documento metodológico",
        date="2024",
        institution="Simple Shift",
        location="Brasil",
        description="Documento detalhado descrevendo a metodologia proprietária desenvolvida por Márcio Elias Barbosa "
                    "para engenharia de dados e gestão estratégica de operações em larga escala, incluindo os seis "
                    "pilares estruturais e suas aplicações práticas.",
        impact="Comprova a autoria e a estrutura da contribuição original, demonstrando que se trata de uma "
               "metodologia sistematizada, documentada e aplicável a ambientes de alta complexidade operacional."
    )

    add_evidence_block(doc,
        exhibit_num=30,
        title="Artigo: Data-Driven Legal Strategy",
        evidence_type="Artigo acadêmico publicado",
        date="2025",
        institution="Promestre — Journal of Management and Innovation (ISSN 3086-1454)",
        location="Brasil",
        description="Artigo acadêmico que examina como a análise estruturada de dados derivados de processos "
                    "pode apoiar a gestão estratégica de grandes carteiras de operações corporativas, "
                    "demonstrando a integração entre conhecimento especializado, gestão operacional e análise de dados.",
        impact="A publicação em periódico acadêmico com revisão por pares valida a relevância da contribuição "
               "original e sua aplicabilidade no campo profissional."
    )

    add_evidence_block(doc,
        exhibit_num=31,
        title="Artigo: Legal Data Engineering in Labor Litigation",
        evidence_type="Artigo acadêmico publicado",
        date="2025",
        institution="Promestre — Journal of Management and Innovation (ISSN 3086-1454)",
        location="Brasil",
        description="Artigo que examina como a organização estruturada do fluxo de trabalho operacional pode "
                    "viabilizar a construção de bases analíticas de dados, demonstrando a aplicação prática "
                    "da metodologia proprietária em ambientes de grande escala.",
        impact="Complementa o primeiro artigo ao detalhar os aspectos técnicos da metodologia de engenharia "
               "de dados, reforçando a profundidade e originalidade da contribuição."
    )

    add_evidence_block(doc,
        exhibit_num=32,
        title="Apresentação — Projeto Sendas (Setembro 2024)",
        evidence_type="Apresentação técnica",
        date="Setembro de 2024",
        institution="Simple Shift",
        location="Brasil",
        description="Apresentação técnica demonstrando a aplicação prática da metodologia proprietária "
                    "em projeto de prestação de serviços técnicos especializados, com exemplos de implementação dos seis pilares.",
        impact="Evidencia que a metodologia não permaneceu em nível teórico, tendo sido efetivamente aplicada "
               "em contexto profissional concreto com resultados mensuráveis."
    )

    add_evidence_block(doc,
        exhibit_num=33,
        title="Material de Treinamento — Aplicação da Metodologia (Julho 2023)",
        evidence_type="Apresentação técnica / Material de treinamento",
        date="Julho de 2023",
        institution="Simple Shift",
        location="São Paulo, SP — Brasil",
        description="Material utilizado para difusão da metodologia junto a equipes profissionais, demonstrando "
                    "a transmissão de conhecimento e a aplicação prática dos conceitos desenvolvidos.",
        impact="Demonstra que a contribuição original influencia a atuação de outros profissionais, ampliando "
               "seu impacto para além da execução direta das atividades operacionais."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # CRITÉRIO 6 — ARTIGOS ACADÊMICOS
    # ══════════════════════════════════════════════════════════════════
    add_navy_section_header(doc, "CRITÉRIO 6 — AUTORIA DE ARTIGOS ACADÊMICOS EM PUBLICAÇÕES ESPECIALIZADAS")

    c6_intro = [
        "A produção intelectual de Márcio Elias Barbosa resulta diretamente de sua experiência profissional na gestão "
        "estratégica de operações de grande escala. Ao sistematizar a experiência acumulada em mais de duas décadas de "
        "atuação, produziu artigos acadêmicos publicados em periódico científico com registro internacional, revisão "
        "por pares e indexação em diretórios acadêmicos.",

        "Os artigos foram publicados na revista Promestre — Journal of Management and Innovation, que possui registro "
        "ISSN 3086-1454, política formal de revisão por pares (double-blind peer review), integração a redes "
        "acadêmicas internacionais, conformidade com as diretrizes do Committee on Publication Ethics (COPE) e "
        "registro DOI para todos os artigos publicados, garantindo rastreabilidade e citação acadêmica internacional.",
    ]
    for text in c6_intro:
        add_paragraph_text(doc, text)

    add_teal_subheader(doc, "Artigo 1")
    doc.add_paragraph()

    add_paragraph_text(doc,
        "Data-Driven Legal Strategy: Analysis of Labor Lawsuits in the Strategic Management of Large Corporate "
        "Litigation Portfolios — O artigo examina como a análise estruturada de dados derivados de processos pode "
        "apoiar a gestão estratégica de grandes carteiras de operações corporativas. O estudo demonstra que a "
        "integração entre conhecimento especializado, gestão operacional e análise de dados permite transformar "
        "o acompanhamento de processos em um sistema contínuo de geração de conhecimento estratégico, "
        "possibilitando a identificação de padrões recorrentes, o aprimoramento de estratégias e a melhoria "
        "da eficiência operacional.", italic=False
    )

    add_evidence_block(doc,
        exhibit_num=34,
        title="Artigo: Data-Driven Legal Strategy (Texto Integral)",
        evidence_type="Artigo acadêmico publicado",
        date="2025",
        institution="Promestre — Journal of Management and Innovation (ISSN 3086-1454)",
        location="Brasil",
        description="Texto integral do artigo acadêmico publicado, demonstrando a contribuição original de Márcio "
                    "Elias Barbosa ao campo da gestão estratégica de operações em larga escala baseada em dados.",
        impact="A publicação em periódico com revisão por pares e registro DOI valida a qualidade e relevância "
               "acadêmica do trabalho, posicionando o autor como produtor de conhecimento especializado."
    )

    add_teal_subheader(doc, "Artigo 2")
    doc.add_paragraph()

    add_paragraph_text(doc,
        "Legal Data Engineering in Labor Litigation: Building Analytical Legal Data Through Workflow Organization "
        "— Este artigo examina como a organização estruturada do fluxo de trabalho operacional pode viabilizar a "
        "construção de bases analíticas de dados. O estudo descreve como a padronização de registros, a "
        "classificação de demandas e a consolidação dos resultados permitem construir bases de dados capazes de "
        "revelar padrões relevantes no comportamento das operações, ampliando a capacidade das equipes de "
        "identificar tendências, desenvolver estratégias baseadas em evidências e aprimorar a coordenação operacional."
    )

    add_evidence_block(doc,
        exhibit_num=35,
        title="Artigo: Legal Data Engineering in Labor Litigation (Texto Integral)",
        evidence_type="Artigo acadêmico publicado",
        date="2025",
        institution="Promestre — Journal of Management and Innovation (ISSN 3086-1454)",
        location="Brasil",
        description="Texto integral do segundo artigo acadêmico publicado, focado na engenharia de dados e "
                    "organização do fluxo de trabalho como base para produção de inteligência analítica.",
        impact="Complementa o primeiro artigo ao detalhar a dimensão técnica da metodologia, reforçando "
               "a consistência e profundidade da contribuição intelectual do autor."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # CRITÉRIO 7 — PAPEL DE LIDERANÇA
    # ══════════════════════════════════════════════════════════════════
    add_navy_section_header(doc, "CRITÉRIO 7 — PAPEL CRÍTICO E DE LIDERANÇA EM ORGANIZAÇÕES DE REPUTAÇÃO DISTINGUIDA")

    c7_intro = [
        "Durante sua atuação na representação do grupo Itaú Unibanco e do Banco Santander, Márcio Elias Barbosa "
        "exerceu papel central na definição e condução da estratégia aplicada à gestão de uma extensa carteira de "
        "processos no Estado de São Paulo. Liderava uma estrutura operacional composta por equipe interna dedicada "
        "e uma ampla rede de profissionais prestadores de serviço distribuídos em diversas regiões do estado.",

        "No período analisado, a operação contava com 21 profissionais internos diretamente alocados, dentre "
        "os quais 13 especialistas, além de mais de 100 profissionais prestadores de serviço distribuídos em "
        "diversas regiões do estado. Essa estrutura viabilizava a atuação simultânea em um elevado número de "
        "localidades, abrangendo processos oriundos de 644 cidades e 153 Varas do Trabalho no Estado de São Paulo.",

        "A amplitude territorial da carteira é evidenciada pela análise de capilaridade, que demonstra forte "
        "dispersão entre comarcas nas regiões sob responsabilidade do escritório, exigindo presença em um número "
        "expressivo de cidades. A análise das operações ao longo do tempo evidencia crescimento consistente, com "
        "aumento simultâneo do volume de atividades e do número de localidades atendidas.",
    ]
    for text in c7_intro:
        add_paragraph_text(doc, text)

    add_teal_subheader(doc, "Dimensões da Liderança Exercida")
    doc.add_paragraph()

    leadership_aspects = [
        "Definição de diretrizes estratégicas para atuação dos profissionais — incluindo postura operacional, "
        "critérios para negociação, elaboração de peças técnicas e estratégias adotadas em diferentes tipos de demandas",
        "Condução de treinamentos técnicos e workshops com apresentação de dados consolidados, indicadores de "
        "desempenho e orientações objetivas sobre como aprimorar a eficácia da atuação profissional",
        "Coordenação de equipes internas e externas com garantia de alinhamento estratégico em operações "
        "distribuídas por centenas de localidades simultâneas",
        "Elevação do padrão técnico das equipes por meio de integração entre dados analíticos e orientações "
        "práticas, resultando em melhoria consistente dos resultados operacionais",
    ]
    for asp in leadership_aspects:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_formatted_run(p, f"• {asp}", size=Pt(10), color=DARK_GRAY)

    doc.add_paragraph()
    add_teal_subheader(doc, "Evidências Documentais — Critério 7")
    doc.add_paragraph()

    add_evidence_block(doc,
        exhibit_num=36,
        title="Estudo de Capilaridade da Carteira de Processos",
        evidence_type="Análise estratégica",
        date="2023",
        institution="Advocacia Elias Barbosa / Conti, Silva e Barbosa",
        location="São Paulo, SP — Brasil",
        description="Análise detalhada da distribuição geográfica da carteira de processos, demonstrando a "
                    "capilaridade da operação em 644 cidades e 153 Varas do Trabalho, com destaque para a "
                    "dispersão territorial que diferencia a operação no Estado de São Paulo.",
        impact="Comprova a escala e complexidade territorial da operação liderada por Márcio Elias Barbosa, "
               "evidenciando a necessidade de coordenação estratégica avançada para manter a qualidade "
               "operacional em centenas de localidades simultâneas."
    )

    add_evidence_block(doc,
        exhibit_num=37,
        title="Relatório de Avaliação de Performance da Equipe",
        evidence_type="Relatório gerencial",
        date="2023",
        institution="Conti, Silva e Barbosa",
        location="São Paulo, SP — Brasil",
        description="Relatório detalhando os indicadores de desempenho da equipe operacional sob liderança "
                    "de Márcio Elias Barbosa, incluindo métricas de eficiência, qualidade e resultados.",
        impact="Demonstra a aplicação de metodologia de gestão por indicadores na condução de grandes equipes, "
               "evidenciando a liderança baseada em dados e resultados mensuráveis."
    )

    add_evidence_block(doc,
        exhibit_num=38,
        title="Reunião de Feedback com Profissionais e Prepostos — Junho 2021",
        evidence_type="Apresentação de reunião",
        date="Junho de 2021",
        institution="Conti, Silva e Barbosa",
        location="São Paulo, SP — Brasil",
        description="Material da reunião estratégica de feedback com profissionais e prepostos, contendo "
                    "análise de resultados, orientações técnicas e diretrizes para melhoria da atuação.",
        impact="Evidencia a liderança exercida por meio de orientação estruturada de equipes, com transmissão "
               "de conhecimento estratégico baseado em dados da operação."
    )

    add_evidence_block(doc,
        exhibit_num=39,
        title="Relatório de Pagamento de Profissionais Prestadores de Serviço",
        evidence_type="Documento financeiro-operacional",
        date="2023",
        institution="Conti, Silva e Barbosa",
        location="São Paulo, SP — Brasil",
        description="Documentação demonstrando a rede de mais de 100 profissionais prestadores de serviço "
                    "coordenados pela operação sob liderança de Márcio Elias Barbosa.",
        impact="Comprova a dimensão da estrutura operacional gerenciada, com coordenação de uma rede "
               "extensiva de profissionais em múltiplas regiões do Estado de São Paulo."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # CRITÉRIO 8 — REMUNERAÇÃO ELEVADA
    # ══════════════════════════════════════════════════════════════════
    add_navy_section_header(doc, "CRITÉRIO 8 — REMUNERAÇÃO SIGNIFICATIVAMENTE SUPERIOR À MÉDIA DA PROFISSÃO")

    c8_intro = [
        "A documentação apresentada demonstra que Márcio Elias Barbosa recebeu remuneração significativamente "
        "superior à média observada para profissionais de sua área no Brasil durante o período analisado. "
        "Sua remuneração era composta principalmente por distribuição de lucros decorrente da atividade "
        "profissional da sociedade, prática comum em sociedades especializadas no Brasil.",

        "Os documentos anexados demonstram pagamentos mensais relevantes, incluindo: R$ 20.000,00 em maio de 2019; "
        "R$ 20.000,00 em julho de 2019; R$ 33.646,38 em abril de 2020; e R$ 20.000,00 em setembro de 2020. Esses "
        "valores são comprovados por recibos de distribuição de lucros e confirmados por documentação contábil e "
        "declarações fiscais.",

        "De acordo com pesquisas salariais disponíveis para o período, a remuneração média de profissionais da "
        "área no Brasil varia entre R$ 3.000 e R$ 5.000 mensais para profissionais em início de carreira; entre "
        "R$ 5.000 e R$ 8.000 para nível intermediário; e entre R$ 8.000 e R$ 15.000 para profissionais seniores "
        "ou em posições de destaque. Os valores recebidos por Márcio Elias Barbosa — variando entre R$ 20.000 e "
        "R$ 33.646 mensais — permanecem substancialmente superiores à média, refletindo o nível de responsabilidade "
        "e a complexidade das funções exercidas.",
    ]
    for text in c8_intro:
        add_paragraph_text(doc, text)

    add_teal_subheader(doc, "Evidências Documentais — Critério 8")
    doc.add_paragraph()

    add_evidence_block(doc,
        exhibit_num=40,
        title="Pesquisa de Média Salarial — Profissionais da Área (Brasil, 2023)",
        evidence_type="Pesquisa salarial de mercado",
        date="2023",
        institution="Fontes: Ministério do Trabalho / Robert Half",
        location="Brasil",
        description="Compilação de dados salariais de mercado para profissionais da área no Brasil, "
                    "incluindo faixas por nível de experiência e especialização.",
        impact="Estabelece o referencial comparativo que demonstra que a remuneração de Márcio Elias Barbosa "
               "situava-se entre 133% e 224% acima da faixa mais alta de mercado para profissionais seniores."
    )

    add_evidence_block(doc,
        exhibit_num=41,
        title="Declaração de Imposto de Renda Pessoa Física (IRPF 2022/2023)",
        evidence_type="Documento fiscal oficial",
        date="2022–2023",
        institution="Receita Federal do Brasil",
        location="Brasil",
        description="Declaração de imposto de renda demonstrando os rendimentos totais recebidos por Márcio "
                    "Elias Barbosa nos exercícios fiscais de 2022 e 2023.",
        impact="Comprova oficialmente os rendimentos totais declarados, validando as informações financeiras "
               "apresentadas pelos demais documentos de suporte."
    )

    add_evidence_block(doc,
        exhibit_num=42,
        title="Declaração de Informações sobre Rendimentos Financeiros (DIRF 2023)",
        evidence_type="Documento fiscal oficial",
        date="2023",
        institution="Receita Federal do Brasil",
        location="Brasil",
        description="DIRF emitida pela fonte pagadora confirmando os valores pagos a Márcio Elias Barbosa "
                    "durante o exercício fiscal.",
        impact="Oferece confirmação independente (pela fonte pagadora) dos valores recebidos, complementando "
               "a declaração pessoal do beneficiário."
    )

    # Profit distribution receipts - grouped by year
    receipt_data = [
        (43, "Setembro de 2018", "ANX-050", "R$ 20.000,00"),
        (44, "Outubro de 2018", "ANX-051", "R$ 20.000,00"),
        (45, "Dezembro de 2018", "ANX-052", "R$ 20.000,00"),
        (46, "Janeiro de 2019", "ANX-053", "R$ 20.000,00"),
        (47, "Abril de 2019", "ANX-054", "R$ 20.000,00"),
        (48, "Agosto de 2019", "ANX-055", "R$ 20.000,00"),
        (49, "Setembro de 2019", "ANX-056", "R$ 20.000,00"),
        (50, "Outubro de 2019", "ANX-057", "R$ 20.000,00"),
        (51, "Novembro de 2019", "ANX-058", "R$ 20.000,00"),
        (52, "Abril de 2020", "ANX-059", "R$ 33.646,38"),
        (53, "Setembro de 2020", "ANX-060", "R$ 20.000,00"),
    ]

    add_teal_subheader(doc, "Recibos de Distribuição de Lucros (2018–2020)")
    doc.add_paragraph()

    for ex_num, month, anx, value in receipt_data:
        add_evidence_block(doc,
            exhibit_num=ex_num,
            title=f"Recibo de Distribuição de Lucros — {month}",
            evidence_type="Recibo financeiro",
            date=month,
            institution="Elias Barbosa Advogados Associados",
            location="São Paulo, SP — Brasil",
            description=f"Recibo de distribuição de lucros no valor de {value}, emitido pelo escritório "
                        f"em favor de Márcio Elias Barbosa como sócio da sociedade.",
            impact=f"Comprova pagamento mensal de {value}, valor substancialmente superior à média de mercado "
                    "para profissionais seniores da área (R$ 8.000 a R$ 15.000 mensais)."
        )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # FORMAÇÃO ACADÊMICA
    # ══════════════════════════════════════════════════════════════════
    add_navy_section_header(doc, "FORMAÇÃO ACADÊMICA")

    add_paragraph_text(doc,
        "Márcio Elias Barbosa possui formação em Direito pela Universidade São Francisco, instituição de ensino "
        "superior privada fundada em 1976, com campus localizado em Bragança Paulista, Estado de São Paulo. "
        "A universidade é reconhecida pelo Ministério da Educação do Brasil (MEC) e mantém tradição na formação "
        "de profissionais nas áreas de Direito, Administração e Engenharia. A formação acadêmica em Direito "
        "proporcionou a base técnica essencial para o desenvolvimento subsequente de sua carreira na gestão "
        "estratégica de operações de grande escala e engenharia de dados aplicada ao campo profissional."
    )

    add_evidence_block(doc,
        exhibit_num=54,
        title="Diploma de Graduação em Direito",
        evidence_type="Documento acadêmico",
        date="[VERIFICAR — data de conclusão]",
        institution="Universidade São Francisco",
        location="Bragança Paulista, SP — Brasil",
        description="Diploma de graduação em Direito emitido pela Universidade São Francisco, instituição "
                    "reconhecida pelo Ministério da Educação do Brasil (MEC).",
        impact="Comprova a formação acadêmica que fundamentou a trajetória profissional de Márcio Elias Barbosa "
               "e proporcionou o conhecimento técnico base para o desenvolvimento de sua metodologia proprietária."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # PUBLICAÇÕES E APRESENTAÇÕES
    # ══════════════════════════════════════════════════════════════════
    add_navy_section_header(doc, "PUBLICAÇÕES E APRESENTAÇÕES")

    add_teal_subheader(doc, "Artigos Acadêmicos Publicados")
    doc.add_paragraph()

    publications = [
        ("BARBOSA, M. E.", "Data-Driven Legal Strategy: Analysis of Labor Lawsuits in the Strategic Management "
         "of Large Corporate Litigation Portfolios", "Promestre — Journal of Management and Innovation",
         "ISSN 3086-1454", "2025"),
        ("BARBOSA, M. E.", "Legal Data Engineering in Labor Litigation: Building Analytical Legal Data Through "
         "Workflow Organization", "Promestre — Journal of Management and Innovation", "ISSN 3086-1454", "2025"),
    ]

    for author, title, journal, issn, year in publications:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        add_formatted_run(p, f"{author} ", size=Pt(10), color=BLACK, bold=True)
        add_formatted_run(p, f"({year}). ", size=Pt(10), color=BLACK)
        add_formatted_run(p, f"{title}. ", size=Pt(10), color=BLACK, italic=True)
        add_formatted_run(p, f"{journal}. {issn}.", size=Pt(10), color=DARK_GRAY)

    add_teal_subheader(doc, "Apresentações e Treinamentos Técnicos")
    doc.add_paragraph()

    presentations = [
        ("Workshop de Feedback e Orientação Profissional", "Novembro de 2018",
         "Apresentação de dados de desempenho e orientações estratégicas para equipe operacional do Banco Itaú."),
        ("Treinamento Técnico em Análise de Operações", "Abril de 2019",
         "Treinamento técnico para profissionais sobre otimização de processos operacionais e melhoria de resultados."),
        ("Treinamento de Análise de Indicadores Estratégicos", "Julho de 2023",
         "Treinamento técnico com base em indicadores de desempenho e análise de dados operacionais."),
        ("Reunião Estratégica de Feedback — Profissionais e Prepostos", "Junho de 2021",
         "Sessão de feedback estruturado com apresentação de dados consolidados e orientações operacionais."),
        ("Apresentação do Projeto Sendas", "Setembro de 2024",
         "Demonstração da aplicação da metodologia proprietária em projeto de prestação de serviços especializados."),
    ]

    for title, date, desc in presentations:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_formatted_run(p, f"• {title}", size=Pt(10), color=NAVY, bold=True)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.space_before = Pt(0)
        add_formatted_run(p2, f"  {date}", size=Pt(9), color=MED_GRAY, italic=True)
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(6)
        p3.paragraph_format.space_before = Pt(0)
        add_formatted_run(p3, f"  {desc}", size=Pt(9.5), color=DARK_GRAY)

    # ══════════════════════════════════════════════════════════════════
    # FOOTER — Add to section
    # ══════════════════════════════════════════════════════════════════
    # Note: python-docx footer with page numbers requires XML manipulation
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_table = footer.add_table(rows=1, cols=1, width=Inches(7.2))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fc = footer_table.cell(0, 0)
    set_cell_shading(fc, NAVY_HEX)
    set_cell_border(fc,
        top={"val": "single", "sz": "0", "color": NAVY_HEX},
        bottom={"val": "single", "sz": "0", "color": NAVY_HEX},
        start={"val": "single", "sz": "0", "color": NAVY_HEX},
        end={"val": "single", "sz": "0", "color": NAVY_HEX})

    fp = fc.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(2)
    fp.paragraph_format.space_after = Pt(2)

    # Page X of Y using fields
    run1 = fp.add_run("Page ")
    run1.font.name = FONT_NAME
    run1.font.size = Pt(9)
    run1.font.color.rgb = WHITE

    # PAGE field
    fld_page = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "><w:r><w:rPr>'
        f'<w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/>'
        f'</w:rPr><w:t>1</w:t></w:r></w:fldSimple>'
    )
    fp._element.append(fld_page)

    run2 = fp.add_run(" of ")
    run2.font.name = FONT_NAME
    run2.font.size = Pt(9)
    run2.font.color.rgb = WHITE

    # NUMPAGES field
    fld_total = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES "><w:r><w:rPr>'
        f'<w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'<w:sz w:val="18"/><w:color w:val="FFFFFF"/>'
        f'</w:rPr><w:t>1</w:t></w:r></w:fldSimple>'
    )
    fp._element.append(fld_total)

    # ══════════════════════════════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════════════════════════════
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, OUTPUT_FILE)
    doc.save(output_path)
    print(f"✅ Résumé gerado com sucesso!")
    print(f"📄 Arquivo: {output_path}")
    print(f"📊 Total de Exhibits: 54")
    return output_path


if __name__ == "__main__":
    create_resume()
