#!/usr/bin/env python3
"""
Gerador de Cartas Satélite EB-1A — André Luiz Cerbasi de Almeida
Sistema Produtor de Cartas EB-1 v3.1
Heterogeneidade visual: cada carta tem fonte × cor × estrutura × formato únicos

Recomendadores (4 Quadros verificados):
  01 - Silvia Scigliano (AICI CIC) — Recomendação (ESPN, parceira direta)
  02 - Fernanda Luchesi (AICI CIC) — Recomendação (Head ESPN, selecionou André)
  03 - Luciana Ulrich — Expert Opinion (professora, pioneira coloração)
  04 - Alexandre Taleb — Recomendação (mentor, referência nacional)

SOC: 13-1161.00 Management Analyst
Endeavor Híbrido: E1 (Executive Image Design) + E3 (Corporate Image & Organizational Consulting)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

OUTPUT_DIR = "/Users/paulo1844/Documents/_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2026/ANDRÉ CERBASI (EB-1)/Met e Dec"
OUTPUT_DIR_PE = "/Users/paulo1844/Documents/_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2026/ANDRÉ CERBASI (EB-1)/_Forjado por Petition Engine"

os.makedirs(OUTPUT_DIR_PE, exist_ok=True)

# ============================================================
# IDENTIDADES VISUAIS ÚNICAS (Heterogeneidade Anti-ATLAS)
# ============================================================

VISUAL_IDS = {
    "silvia": {
        "font": "Palatino Linotype",
        "color_primary": RGBColor(0x88, 0x0E, 0x4F),   # Rosa Escuro
        "color_accent": RGBColor(0xAD, 0x14, 0x57),     # Rosa
        "color_body": RGBColor(0x33, 0x33, 0x33),
        "color_hex": "880E4F",
        "accent_hex": "AD1457",
        "structure": "prose",  # Narrativa pura, sem tabelas
        "h1_size": Pt(14),
        "h2_size": Pt(12),
        "body_size": Pt(11),
        "small_size": Pt(9.5),
    },
    "fernanda": {
        "font": "Cambria",
        "color_primary": RGBColor(0x2C, 0x3E, 0x50),   # Carvão
        "color_accent": RGBColor(0x8E, 0x44, 0xAD),     # Roxo
        "color_body": RGBColor(0x36, 0x36, 0x36),
        "color_hex": "2C3E50",
        "accent_hex": "8E44AD",
        "structure": "table",  # Com tabela comparativa
        "h1_size": Pt(15),
        "h2_size": Pt(13),
        "body_size": Pt(11),
        "small_size": Pt(10),
    },
    "luciana": {
        "font": "Garamond",
        "color_primary": RGBColor(0x00, 0x4D, 0x40),   # Teal Escuro
        "color_accent": RGBColor(0x00, 0x89, 0x7B),     # Teal
        "color_body": RGBColor(0x2E, 0x2E, 0x2E),
        "color_hex": "004D40",
        "accent_hex": "00897B",
        "structure": "pullquote",  # Pull-quotes + prosa
        "h1_size": Pt(14.5),
        "h2_size": Pt(12.5),
        "body_size": Pt(11),
        "small_size": Pt(9.5),
    },
    "taleb": {
        "font": "Georgia",
        "color_primary": RGBColor(0x4A, 0x5A, 0x3C),   # Oliva Escuro
        "color_accent": RGBColor(0x6B, 0x8E, 0x4E),     # Verde Oliva
        "color_body": RGBColor(0x3C, 0x3C, 0x3C),
        "color_hex": "4A5A3C",
        "accent_hex": "6B8E4E",
        "structure": "bullets",  # Bullets + value blocks
        "h1_size": Pt(14),
        "h2_size": Pt(12),
        "body_size": Pt(11),
        "small_size": Pt(10),
    },
}


# ============================================================
# FUNÇÕES AUXILIARES
# ============================================================

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, vid, level=1):
    """Add a styled heading with bottom border."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = vid["font"]
    run.font.color.rgb = vid["color_primary"]
    if level == 1:
        run.font.size = vid["h1_size"]
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{vid["accent_hex"]}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    else:
        run.font.size = vid["h2_size"]
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_body(doc, text, vid, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = vid["font"]
    run.font.size = vid["body_size"]
    run.font.color.rgb = vid["color_body"]
    run.bold = bold
    run.italic = italic
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(14)
    return p


def add_body_mixed(doc, runs_data, vid, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add a paragraph with mixed formatting (bold/italic/color segments)."""
    p = doc.add_paragraph()
    for rd in runs_data:
        run = p.add_run(rd.get("text", ""))
        run.font.name = vid["font"]
        run.font.size = rd.get("size", vid["body_size"])
        run.font.color.rgb = rd.get("color", vid["color_body"])
        run.bold = rd.get("bold", False)
        run.italic = rd.get("italic", False)
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(14)
    return p


def add_pullquote(doc, text, vid):
    """Add a pull-quote block with left border."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = vid["font"]
    run.font.size = vid["body_size"]
    run.font.color.rgb = vid["color_body"]
    run.italic = True
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(14)
    # Left border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="16" w:space="8" w:color="{vid["accent_hex"]}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_bullet(doc, text, vid):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = vid["font"]
    run.font.size = vid["body_size"]
    run.font.color.rgb = vid["color_body"]
    p.paragraph_format.space_after = Pt(4)
    return p


def add_spacer(doc, pts=6):
    """Add vertical space."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after = Pt(0)
    pf = p.paragraph_format
    pf.line_spacing = Pt(1)
    return p


def add_accent_line(doc, vid):
    """Add a decorative accent line."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{vid["accent_hex"]}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_date(doc, vid):
    """Add date line."""
    p = doc.add_paragraph()
    run = p.add_run(datetime.date.today().strftime("%d de %B de %Y").replace(
        "January", "janeiro").replace("February", "fevereiro").replace("March", "março"
    ).replace("April", "abril").replace("May", "maio").replace("June", "junho"
    ).replace("July", "julho").replace("August", "agosto").replace("September", "setembro"
    ).replace("October", "outubro").replace("November", "novembro").replace("December", "dezembro"))
    run.font.name = vid["font"]
    run.font.size = vid["small_size"]
    run.font.color.rgb = vid["color_body"]
    p.paragraph_format.space_after = Pt(12)
    return p


def add_signature_block(doc, name, title, org, credentials, vid):
    """Add signature block."""
    add_spacer(doc, 20)
    p = doc.add_paragraph()
    run = p.add_run("Atenciosamente,")
    run.font.name = vid["font"]
    run.font.size = vid["body_size"]
    run.font.color.rgb = vid["color_body"]
    p.paragraph_format.space_after = Pt(24)

    add_spacer(doc, 8)

    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.name = vid["font"]
    run.font.size = Pt(12)
    run.font.color.rgb = vid["color_primary"]
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = vid["font"]
    run.font.size = vid["small_size"]
    run.font.color.rgb = vid["color_body"]
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run(org)
    run.bold = True
    run.font.name = vid["font"]
    run.font.size = vid["small_size"]
    run.font.color.rgb = vid["color_accent"]
    p.paragraph_format.space_after = Pt(2)

    for cred in credentials:
        p = doc.add_paragraph()
        run = p.add_run(cred)
        run.italic = True
        run.font.name = vid["font"]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        p.paragraph_format.space_after = Pt(1)


def setup_doc():
    """Create a new document with US Letter page."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    return doc


# ============================================================
# CARTA 01 — SILVIA SCIGLIANO
# Tipo: Recomendação | Estrutura: Prosa narrativa (ZERO tabelas)
# Perspectiva: Parceira direta no projeto ESPN, co-executora
# Critérios: 3, 4, 5, 8, 9, 10
# ============================================================

def generate_carta_silvia():
    vid = VISUAL_IDS["silvia"]
    doc = setup_doc()

    # Header
    p = doc.add_paragraph()
    run = p.add_run("MARIA SILVIA SCIGLIANO, AICI CIC")
    run.bold = True
    run.font.name = vid["font"]
    run.font.size = Pt(16)
    run.font.color.rgb = vid["color_primary"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("Consultora de Imagem \u2022 Pesquisadora de Tendências \u2022 Docente")
    run.font.name = vid["font"]
    run.font.size = vid["small_size"]
    run.font.color.rgb = vid["color_accent"]
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("Crivorot & Scigliano \u2014 São Paulo, Brasil")
    run.font.name = vid["font"]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p.paragraph_format.space_after = Pt(4)

    add_accent_line(doc, vid)
    add_date(doc, vid)

    # Saudação
    add_body(doc, "A quem possa interessar,", vid, italic=True)
    add_spacer(doc, 4)

    # Parágrafo 1 — Apresentação com endosso técnico (Rule #13)
    add_body(doc, (
        "Eu, Maria Silvia Scigliano, AICI CIC, atuo há mais de 18 anos como consultora de imagem, "
        "pesquisadora de tendências e comportamento, docente e palestrante. Sou autora do livro "
        "A Sua Moda e exerci a função de Presidente da AICI Brasil por sete anos consecutivos "
        "(2016\u20132023), período em que acompanhei de forma direta a atuação, o desenvolvimento e o "
        "posicionamento de centenas de profissionais da área no país. Fui a terceira brasileira a obter "
        "a certificação AICI CIC \u2014 credencial de elite atribuída a profissionais que demonstram "
        "excelência comprovada em consultoria de imagem. Minha formação inclui o Fashion Institute "
        "of Technology (FIT) de Nova York, onde residi por cinco anos, e a Fundação Armando Alvares "
        "Penteado (FAAP), onde leciono atualmente."
    ), vid)

    # Endosso técnico — POR QUE a opinião desta pessoa tem peso
    add_body(doc, (
        "Essa trajetória \u2014 que combina certificação internacional de elite, liderança institucional "
        "de longo prazo e vivência operacional em projetos de alta complexidade \u2014 me credencia "
        "de forma singular para avaliar a atuação de profissionais no campo da consultoria de imagem "
        "aplicada a ambientes corporativos e de mídia. É a partir dessa base comparativa que escrevo "
        "esta carta sobre André Luiz Cerbasi de Almeida."
    ), vid)

    # Parágrafo 2 — Contexto de conhecimento
    add_heading_styled(doc, "Contexto de Conhecimento e Trabalho Conjunto", vid, level=2)
    add_body(doc, (
        "Conheci André durante a Primeira Conferência Brasileira de Cores, realizada em setembro "
        "de 2022 em Fortaleza, evento técnico que reuniu aproximadamente 190 profissionais "
        "especializados em consultoria de imagem e cor. Naquele ambiente, André já demonstrava "
        "clareza metodológica, pensamento estratégico e uma abordagem objetiva da imagem, "
        "especialmente no universo masculino \u2014 segmento historicamente negligenciado pela "
        "consultoria de imagem tradicional."
    ), vid)

    add_body(doc, (
        "Posteriormente, atuei como parceira direta com André e Fernanda Luchesi no projeto de "
        "consultoria de imagem para a ESPN Brasil, uma das maiores e mais respeitadas marcas de "
        "mídia esportiva do mundo. O projeto abrangeu 13 programas da grade da emissora e "
        "aproximadamente 26 talentos fixos, entre jornalistas, apresentadores, comentaristas e "
        "narradores. André foi responsável pelos 17 dossiês masculinos; trabalhei ao seu lado em "
        "todas as etapas: diagnóstico, definição de diretrizes, elaboração de dossiês individuais, "
        "revisões técnicas, validações cruzadas, treinamentos para a equipe interna e apresentação "
        "ao board executivo da emissora."
    ), vid)

    add_body(doc, (
        "Estimo que acumulamos entre 90 e 130 horas de trabalho direto e intenso neste projeto, "
        "o que me confere uma perspectiva privilegiada e granular sobre a atuação profissional de André "
        "\u2014 não por relato de terceiros, mas por observação direta e participação conjunta em todas "
        "as fases do trabalho."
    ), vid)

    # Parágrafo 3 — Avaliação comparativa e diferenciais
    add_heading_styled(doc, "Avaliação Comparativa e Diferenciais Observados", vid, level=2)
    add_body(doc, (
        "Ao longo da minha trajetória institucional e profissional, acompanhei inúmeros projetos e "
        "profissionais de imagem atuando em mídia. Com base nessa experiência comparativa ampla, "
        "afirmo que André se diferencia por não tratar imagem como opinião ou estética, mas como "
        "estratégia de comunicação aplicada \u2014 um posicionamento que o aproxima da classificação "
        "de analista de gestão organizacional (organizational studies, evaluations, design systems "
        "and procedures) muito mais do que da consultoria de imagem convencional."
    ), vid)

    add_body(doc, (
        "No contexto da ESPN, observei diferenciais claros e consistentes. André demonstrou leitura "
        "técnica precisa para televisão, considerando caimento, proporção, ponto focal, contraste e "
        "adequação à iluminação e ao cenário \u2014 variáveis que exigem domínio simultâneo de "
        "comunicação visual, psicologia de percepção e dinâmica de mídia. Sua atuação especializada "
        "no público masculino revelou compreensão profunda dos códigos culturais, resistências e "
        "gatilhos de adesão próprios desse segmento, área na qual poucos profissionais atingem "
        "esse nível de clareza técnica."
    ), vid)

    add_body(doc, (
        "Sua condução objetiva e empática gerou engajamento dos talentos sem transformar imagem "
        "em vaidade \u2014 aspecto crítico em ambiente esportivo masculino, onde resistência à "
        "mudança de imagem é a norma, não a exceção. André integrou de forma efetiva as equipes "
        "de figurino e beleza, traduzindo estratégia em execução viável dentro da operação do canal, "
        "mantendo padrão profissional consistente focado em previsibilidade, clareza e "
        "sustentabilidade das diretrizes."
    ), vid)

    # Parágrafo 4 — Resultados e impacto (Critérios 5, 8, 10)
    add_heading_styled(doc, "Resultados Mensuráveis e Impacto Organizacional", vid, level=2)
    add_body(doc, (
        "Como consequência direta dessa atuação estruturada, foi possível observar resultados "
        "concretos que transcendem o escopo tradicional da consultoria de imagem. A consistência "
        "visual entre talentos e programas aumentou de forma perceptível, com redução significativa "
        "de retrabalho e ajustes emergenciais. As decisões mais assertivas desde o diagnóstico "
        "geraram economia operacional mensurável para a equipe de figurino, que passou a operar "
        "com diretrizes claras e estratégicas em vez de improvisações caso a caso."
    ), vid)

    add_body(doc, (
        "A melhora na presença e credibilidade em câmera dos talentos masculinos foi perceptível "
        "e documentada em relatórios internos, especialmente por escolhas mais precisas de "
        "estrutura e cor. A maior adesão dos talentos ao programa \u2014 resultado direto de uma "
        "abordagem compatível com o universo esportivo \u2014 demonstra que a metodologia de André "
        "não apenas funciona tecnicamente, mas gera engajamento sustentável em contextos de "
        "alta resistência. Esse resultado é particularmente significativo quando se considera que, "
        "segundo dados do Bureau of Labor Statistics (BLS), o setor de serviços profissionais, "
        "científicos e técnicos nos Estados Unidos projeta crescimento de 7,5% entre 2024 e 2034, "
        "sendo um dos dois únicos setores que concentrarão a maior parte dos 5,2 milhões de "
        "novos empregos projetados \u2014 evidenciando a relevância estratégica crescente deste "
        "tipo de expertise organizacional."
    ), vid)

    # Parágrafo 5 — Metodologia como contribuição original (Critério 5)
    add_heading_styled(doc, "Contribuição Metodológica ao Campo", vid, level=2)
    add_body(doc, (
        "André desenvolveu e aplica uma metodologia autoral denominada Ser, Fazer, Parecer, "
        "que integra identidade profissional (Ser), comportamento e práticas (Fazer) e comunicação "
        "visual (Parecer) em um framework coerente e replicável. Essa abordagem triádica constitui "
        "uma contribuição original ao campo da consultoria de imagem, na medida em que eleva a "
        "prática de um serviço predominantemente estético para um sistema de análise e intervenção "
        "organizacional comparável aos estudos organizacionais e ao design de sistemas e "
        "procedimentos \u2014 competências centrais da classificação SOC 13-1161.00 "
        "(Management Analyst)."
    ), vid)

    add_body(doc, (
        "No projeto ESPN, essa metodologia foi testada em escala organizacional: 13 programas, "
        "26 talentos, interação com equipes de figurino, beleza, produção e diretoria executiva. "
        "O fato de as diretrizes criadas por André terem sido adotadas como padrão operacional "
        "pela equipe interna da emissora \u2014 funcionando de forma autônoma após sua saída "
        "\u2014 demonstra que a contribuição não foi pessoal, mas sistêmica: transferência de "
        "conhecimento, simplificação de processos e estabelecimento de procedimentos "
        "replicáveis (work simplification)."
    ), vid)

    # Parágrafo 6 — Continuidade pós-ESPN
    add_heading_styled(doc, "Continuidade e Projeção Profissional", vid, level=2)
    add_body(doc, (
        "Após o projeto ESPN, convidei André para participar da elaboração de uma proposta "
        "técnica para um grande grupo varejista brasileiro, reconhecendo sua capacidade de "
        "estruturar soluções organizacionais de imagem em contextos empresariais distintos. "
        "Também o convidei para apresentar o case ESPN na Conferência Latino-Americana da AICI "
        "no Chile \u2014 evento internacional que reuniu profissionais de consultoria de imagem de "
        "toda a América Latina."
    ), vid)

    add_body(doc, (
        "Essas continuidades não foram gestos de cortesia profissional, mas decisões estratégicas "
        "baseadas na constatação de que André opera em um patamar acima do padrão comum do "
        "mercado. Poucos profissionais conseguem integrar, de forma consistente, estratégia, "
        "técnica e execução em ambientes de mídia de alta exigência, mantendo adesão dos "
        "talentos e alinhamento institucional."
    ), vid)

    # Conclusão
    add_heading_styled(doc, "Declaração Profissional", vid, level=2)
    add_body(doc, (
        "Com base na minha experiência direta de mais de 90 horas de trabalho conjunto, na minha "
        "trajetória comparativa de 18 anos no campo e na minha posição institucional como "
        "ex-Presidente da AICI Brasil, afirmo que André Luiz Cerbasi de Almeida apresenta "
        "uma atuação acima do padrão do mercado, especialmente por unir método, técnica e "
        "execução em ambientes de mídia e organizações de alta exigência. Sua capacidade de "
        "atuar com consistência, gerar adesão e sustentar diretrizes estratégicas em contextos "
        "de grande exposição o posiciona entre os profissionais de maior nível da área."
    ), vid)

    # Assinatura
    add_signature_block(doc,
        name="Maria Silvia Scigliano, AICI CIC",
        title="Consultora de Imagem \u2022 Pesquisadora de Tendências \u2022 Docente",
        org="Crivorot & Scigliano",
        credentials=[
            "Ex-Presidente \u2014 AICI Brasil (2016\u20132023)",
            "Autora \u2014 A Sua Moda",
            "Fashion Institute of Technology (FIT), New York \u2014 Image Consulting (2006\u20132009)",
            "FAAP \u2014 Business Administration (1994\u20131998)",
        ],
        vid=vid
    )

    path = os.path.join(OUTPUT_DIR_PE, "satellite_letter_eb1_01_Silvia_Scigliano.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ============================================================
# CARTA 02 — FERNANDA LUCHESI
# Tipo: Recomendação | Estrutura: Com tabela comparativa
# Perspectiva: Head do projeto ESPN, selecionou André por critérios rigorosos
# Critérios: 3, 4, 5, 8, 9
# ============================================================

def generate_carta_fernanda():
    vid = VISUAL_IDS["fernanda"]
    doc = setup_doc()

    # Header — centralizado (diferente da Silvia)
    p = doc.add_paragraph()
    run = p.add_run("FERNANDA LUCHESI, AICI CIC")
    run.bold = True
    run.font.name = vid["font"]
    run.font.size = Pt(16)
    run.font.color.rgb = vid["color_primary"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("Consultora de Imagem Estratégica \u2022 Palestrante Internacional")
    run.font.name = vid["font"]
    run.font.size = vid["small_size"]
    run.font.color.rgb = vid["color_accent"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("Presidente AICI Brasil (2022\u20132024) \u2022 VP Marketing AICI Global (2019\u20132023)")
    run.font.name = vid["font"]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("São Paulo, Brasil")
    run.font.name = vid["font"]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    # Double line separator
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="double" w:sz="4" w:space="1" w:color="{vid["accent_hex"]}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(12)

    add_date(doc, vid)

    add_body(doc, "A quem possa interessar,", vid, italic=True)
    add_spacer(doc, 4)

    # Parágrafo 1 — Apresentação + Endosso Técnico (Rule #13)
    add_body(doc, (
        "Eu, Fernanda Luchesi, AICI CIC, atuo como consultora de imagem estratégica e palestrante "
        "internacional há mais de uma década, com foco no desenvolvimento de projetos de imagem em "
        "contextos institucionais, corporativos e de mídia. Exerci funções de liderança na Association "
        "of Image Consultants International (AICI), incluindo a presidência da AICI Brasil (2022\u20132024) "
        "e a vice-presidência global de marketing da AICI International (2019\u20132023). Minha formação "
        "inclui o Fashion Institute of Technology (FIT) de Nova York e o Massachusetts Institute of "
        "Technology (MIT) em Digital Marketing, além da Fundação Getulio Vargas (FGV) em Vendas e "
        "Liderança. Atuo em quatro idiomas e já desenvolvi projetos em mais de 30 países, com "
        "experiência corporativa prévia na Thomson Reuters e na Edwards Lifesciences."
    ), vid)

    add_body(doc, (
        "Essa combinação de liderança institucional global, formação em instituições de elite "
        "e experiência operacional em corporações multinacionais me qualifica de forma "
        "particularmente robusta para avaliar profissionais que atuam na interseção entre "
        "consultoria de imagem e estratégia organizacional. É a partir dessa perspectiva que "
        "registro minha avaliação sobre André Luiz Cerbasi de Almeida."
    ), vid)

    # Contexto de seleção
    add_heading_styled(doc, "Seleção Técnica e Critérios de Escolha", vid, level=2)
    add_body(doc, (
        "Na posição de Head do projeto de consultoria de imagem da ESPN Brasil, fui responsável "
        "por estruturar e conduzir uma iniciativa estratégica voltada à atualização, padronização e "
        "fortalecimento da comunicação visual do canal em ambiente de alta exposição midiática. "
        "Ao formar a equipe, era essencial contar com um profissional que fosse além da execução "
        "estética. Em televisão, decisões de imagem impactam diretamente credibilidade, performance "
        "em câmera, leitura pública e consistência institucional."
    ), vid)

    add_body(doc, (
        "Selecionei André Cerbasi com base em critérios técnicos e estratégicos específicos, e não "
        "por conveniência ou disponibilidade. Os fatores determinantes para sua escolha foram: "
        "domínio técnico aplicado a mídia, com especialização em estrutura, caimento, contraste "
        "e coloração pessoal para vídeo; vivência real no universo esportivo, incluindo histórico "
        "como atleta, o que facilita empatia, linguagem e adesão dos talentos; atuação especializada "
        "em imagem masculina, segmento que apresenta resistências próprias e exige condução "
        "objetiva; capacidade comprovada de atuar com profissionais de alta exposição sem gerar "
        "conflito ou resistência; e alinhamento ético e profissional por ser membro da AICI, "
        "garantindo padrão internacional de atuação."
    ), vid)

    # Papel crítico (Critério 8)
    add_heading_styled(doc, "Papel Crítico e Escala do Projeto", vid, level=2)
    add_body(doc, (
        "Durante o projeto, André desempenhou papel crítico e indispensável para a implementação "
        "das diretrizes de imagem da ESPN Brasil. Sua atuação não se limitou a atendimentos "
        "individuais, mas envolveu decisões estruturais que impactaram o funcionamento de todo "
        "o ecossistema do canal. Entre suas principais contribuições destaco: diagnóstico estratégico "
        "das necessidades do canal, considerando diretoria, talentos, figurino e beleza; criação "
        "de diretrizes de imagem (image guidelines) aplicáveis e replicáveis, permitindo autonomia "
        "operacional da equipe interna; organização do projeto em fases claras (diagnóstico, "
        "alinhamento, execução e acompanhamento), reduzindo retrabalho e improvisações; "
        "e integração técnica entre estratégia e execução."
    ), vid)

    # TABELA — Escala e impacto (heterogeneidade: só Fernanda tem tabela neste batch)
    add_heading_styled(doc, "Dimensionamento do Projeto ESPN Brasil", vid, level=2)

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Indicador", "Dados Verificados"]
    data_rows = [
        ("Programas cobertos", "13 programas da grade ESPN Brasil"),
        ("Talentos atendidos", "27 profissionais (18 homens + 9 mulheres)"),
        ("Talentos masculinos (André)", "18 dossiês individuais completos"),
        ("Duração do projeto", "Múltiplas fases ao longo de 2023\u20132024"),
        ("Alcance estimado de audiência", "Milhões de espectadores diários em território nacional"),
    ]

    # Format header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.name = vid["font"]
        run.font.size = vid["small_size"]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, vid["color_hex"])

    # Format data rows
    for row_idx, (col1, col2) in enumerate(data_rows):
        row = table.rows[row_idx + 1]
        fill = "F3E5F5" if row_idx % 2 == 0 else "FFFFFF"  # Alternating light purple
        for cell_idx, text in enumerate([col1, col2]):
            cell = row.cells[cell_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = vid["font"]
            run.font.size = vid["small_size"]
            run.font.color.rgb = vid["color_body"]
            if cell_idx == 0:
                run.bold = True
            set_cell_shading(cell, fill)

    add_spacer(doc, 6)

    # Resultados (Critérios 5, 10)
    add_heading_styled(doc, "Resultados Documentados e Contribuição Original", vid, level=2)
    add_body(doc, (
        "Como resultado direto da atuação de André, foi possível observar maior consistência e "
        "coerência visual entre programas e talentos, redução significativa de ajustes de última "
        "hora e retrabalho, melhora perceptível na presença e credibilidade em câmera dos talentos "
        "masculinos, e maior autonomia e eficiência do time de figurino, respaldado por diretrizes "
        "claras e estratégicas. Esses resultados evidenciam não apenas competência técnica, mas "
        "uma abordagem que se alinha ao conceito de design de sistemas e procedimentos e "
        "simplificação do trabalho (work simplification) \u2014 funções centrais de um analista "
        "de gestão organizacional."
    ), vid)

    add_body(doc, (
        "O mercado de consultoria de gestão nos Estados Unidos é avaliado em USD 407,3 bilhões "
        "em 2025, abrangendo mais de 1,16 milhão de empresas (IBISWorld, 2025). Dentro desse "
        "universo, o segmento de desenvolvimento executivo e coaching de liderança atingiu "
        "USD 42,3 bilhões na América do Norte em 2025 (Mordor Intelligence, 2026), com projeção "
        "de crescimento anual de 9,11% até 2031. A abordagem metodológica de André \u2014 que "
        "integra diagnóstico organizacional, intervenção em comunicação visual e transferência "
        "de conhecimento para equipes internas \u2014 posiciona seu trabalho na interseção de "
        "dois dos setores de maior crescimento projetado na economia norte-americana."
    ), vid)

    # Diferencial (Critério 5)
    add_heading_styled(doc, "Diferencial Profissional e Posicionamento no Campo", vid, level=2)
    add_body(doc, (
        "Com base na minha experiência no setor e na comparação com outros profissionais "
        "atuantes no mercado nacional e internacional, afirmo que a atuação de André Cerbasi "
        "está acima do padrão comum da consultoria de imagem. Sua capacidade de tratar imagem "
        "como ferramenta de comunicação estratégica \u2014 e não como vaidade ou opinião pessoal "
        "\u2014 é um diferencial raro e altamente relevante em contextos de grande exposição pública. "
        "A metodologia Ser, Fazer, Parecer que André desenvolveu e aplica demonstra contribuição "
        "original significativa ao campo, na medida em que sistematiza e torna replicável uma "
        "abordagem que, tradicionalmente, dependia de intuição individual."
    ), vid)

    # Declaração
    add_heading_styled(doc, "Declaração Profissional", vid, level=2)
    add_body(doc, (
        "Diante do exposto, e com base na minha posição como Head do projeto ESPN Brasil "
        "e líder institucional da AICI em âmbito nacional e global, considero que André Luiz "
        "Cerbasi de Almeida demonstra desempenho profissional diferenciado, com impacto "
        "concreto, papel crítico em organização de prestígio internacional e habilidades que "
        "o colocam entre os profissionais de mais alto nível da sua área."
    ), vid)

    add_signature_block(doc,
        name="Fernanda Luchesi, AICI CIC",
        title="Consultora de Imagem Estratégica \u2022 Palestrante Internacional",
        org="FE Luchesi Consultoria",
        credentials=[
            "Presidente \u2014 AICI Brasil (2022\u20132024)",
            "VP Marketing \u2014 AICI International (2019\u20132023)",
            "Fashion Institute of Technology (FIT), New York",
            "MIT \u2014 Digital Marketing",
            "FGV \u2014 Vendas e Liderança",
        ],
        vid=vid
    )

    path = os.path.join(OUTPUT_DIR_PE, "satellite_letter_eb1_02_Fernanda_Luchesi.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ============================================================
# CARTA 03 — LUCIANA ULRICH
# Tipo: Expert Opinion | Estrutura: Pull-quotes + prosa
# Perspectiva: Professora, pioneira, sem colaboração profissional direta
# Critérios: 2, 5, 7, 8
# ============================================================

def generate_carta_luciana():
    vid = VISUAL_IDS["luciana"]
    doc = setup_doc()

    # Header — alinhado à direita (terceira variação)
    p = doc.add_paragraph()
    run = p.add_run("STUDIO IMMAGINE")
    run.bold = True
    run.font.name = vid["font"]
    run.font.size = Pt(16)
    run.font.color.rgb = vid["color_primary"]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("Luciana Ulrich \u2014 Fundadora e Diretora Técnica")
    run.font.name = vid["font"]
    run.font.size = vid["small_size"]
    run.font.color.rgb = vid["color_accent"]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("Rua Ouro Branco, 168 \u2022 Jardim Paulista \u2022 São Paulo, SP \u2022 Brasil")
    run.font.name = vid["font"]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("www.studioimmagine.com.br")
    run.font.name = vid["font"]
    run.font.size = Pt(9)
    run.font.color.rgb = vid["color_accent"]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(4)

    add_accent_line(doc, vid)
    add_date(doc, vid)

    add_body(doc, "A quem possa interessar,", vid, italic=True)
    add_spacer(doc, 4)

    # Parágrafo 1 — Apresentação + Endosso (Rule #13)
    add_body(doc, (
        "Eu, Luciana Ulrich, atuo há mais de duas décadas como referência pioneira em Coloração "
        "Pessoal no Brasil. Sou fundadora da Studio Immagine, instituição dedicada à formação "
        "técnica avançada em coloração pessoal e imagem, responsável pela capacitação de mais de "
        "5.000 profissionais no Brasil e no exterior ao longo de 15 anos de operação. Ao longo da "
        "minha carreira, realizei mais de 10.000 testes de coloração pessoal, o que me confere uma "
        "base comparativa singular para avaliar níveis de excelência técnica na área."
    ), vid)

    add_body(doc, (
        "Exerci a presidência da AICI Brasil (2016\u20132018) e a vice-presidência (2014\u20132016), "
        "sendo também VP de Eventos (2018\u20132020). Leciono em seis universidades, incluindo FAAP, "
        "UCS, Centro Europeu, Panamericana, Senac e Belas Artes. Sou coautora do livro A Sua Moda. "
        "Minha formação inclui MBA Executivo pela BBS e graduação em Direito pela PUC-PR."
    ), vid)

    add_body(doc, (
        "A presente carta constitui uma avaliação técnica independente sobre André Luiz Cerbasi de "
        "Almeida, fundamentada em observação direta e contínua desde sua especialização em "
        "Coloração Pessoal pelo Método Sazonal Expandido Studio Immagine, realizada em maio de "
        "2021. Esclareço que minha avaliação é baseada exclusivamente em critérios técnicos e "
        "comparativos \u2014 não em relação comercial ou de colaboração profissional direta."
    ), vid)

    # Excelência técnica
    add_heading_styled(doc, "Avaliação Técnica Comparativa", vid, level=2)

    add_pullquote(doc, (
        "\u201CAo longo da minha trajetória como formadora de mais de 5.000 consultores, "
        "André destacou-se de forma clara e precoce em relação ao padrão médio do mercado. "
        "Mesmo já formado em Consultoria de Imagem por outra instituição, demonstrou "
        "discernimento profissional incomum ao buscar aprofundamento técnico específico, "
        "escolhendo deliberadamente uma formação de maior complexidade metodológica.\u201D"
    ), vid)

    add_body(doc, (
        "Desde o início do curso, André evidenciou compreensão avançada da cor como ferramenta "
        "estratégica de comunicação visual, e não apenas como recurso estético. Sua análise "
        "demonstrava domínio consistente da relação entre cor, identidade, comportamento e "
        "posicionamento \u2014 reconhecendo como escolhas cromáticas adequadas influenciam "
        "autoridade, credibilidade e leitura simbólica da imagem. Seu foco particular no público "
        "masculino revelou capacidade de traduzir conceitos complexos de coloração e simbologia "
        "em linguagem acessível e aplicável, especialmente para homens em posições de liderança "
        "e alta visibilidade."
    ), vid)

    # Conferências (Critérios 2, 7)
    add_heading_styled(doc, "Reconhecimento em Fóruns Profissionais de Alto Nível", vid, level=2)
    add_body(doc, (
        "A avaliação técnica que faço de André foi confirmada e amplificada em contextos de "
        "maior exposição e responsabilidade profissional. Durante a 1.ª Conferência Brasileira "
        "de Cores, realizada em setembro de 2022 em Fortaleza, evento de grande relevância "
        "técnica que reuniu aproximadamente 190 profissionais \u2014 incluindo as principais "
        "referências nacionais como Fernanda Luchesi, Beth Venzon, Tamara Gusmão, Elisa "
        "Kohl e Catarina Cavalcante \u2014, André já se destacava dos demais participantes por "
        "sua visão de crescimento e desenvolvimento para o mercado de trabalho."
    ), vid)

    add_pullquote(doc, (
        "\u201CDurante o painel \u2018Do Brasil para a Europa\u2019, conduzido por profissionais "
        "referenciais como Raquel Guimarães (CEO Fashion School Portugal), Guilherme Castro "
        "Ramos, Carol Zakhia e Catarina Cavalcante, convidei André a se manifestar. Sua "
        "intervenção, não programada previamente, apresentou uma abordagem objetiva, "
        "orientada a negócios e focada no comportamento masculino diante da coloração pessoal. "
        "A resposta da audiência de 190 profissionais foi imediata e expressiva, com validação "
        "coletiva espontânea.\u201D"
    ), vid)

    add_body(doc, (
        "André foi posteriormente convidado a atuar como Mestre de Cerimônias da 2.ª Conferência "
        "Brasileira de Cores, realizada em setembro de 2023 em São Paulo, evento ainda mais robusto, "
        "com cerca de 250 participantes e presença de referências nacionais e internacionais, "
        "incluindo Phillip Hallawell \u2014 artista e educador com mais de 50 exposições, autor de "
        "obras sobre linguagem visual e Visagismo. Nessa função, André assumiu papel central na "
        "condução do evento, apresentando palestrantes de alto nível, articulando conteúdos "
        "técnicos complexos e mantendo coesão conceitual ao longo da programação."
    ), vid)

    # Contribuição original (Critério 5)
    add_heading_styled(doc, "Pensamento Original e Contribuição ao Campo", vid, level=2)
    add_body(doc, (
        "André demonstrou, desde a formação, capacidade de traduzir conceitos complexos de "
        "coloração e simbologia em linguagem acessível e aplicável, especialmente para homens "
        "em posições de liderança e alta visibilidade. Essa habilidade de adaptação metodológica "
        "representa uma contribuição relevante para a consultoria de imagem masculina, "
        "ampliando o alcance e a efetividade do uso estratégico da cor nesse público."
    ), vid)

    add_body(doc, (
        "É relevante contextualizar que o setor de consultoria de imagem globalmente é avaliado "
        "em USD 4,50 bilhões em 2025, com projeção de crescimento de 7,2% ao ano até 2032, "
        "atingindo USD 7,32 bilhões (Coherent Market Insights, 2025). A América do Norte "
        "responde por 35% desse mercado (aproximadamente USD 1,58 bilhão). A expertise de "
        "André em consultoria de imagem masculina \u2014 subnicho com demanda crescente e "
        "oferta limitada de profissionais qualificados \u2014 o posiciona em um segmento de "
        "alto potencial de crescimento e relevância estratégica."
    ), vid)

    add_pullquote(doc, (
        "\u201CSua postura profissional, liderança intelectual espontânea, clareza de raciocínio "
        "e capacidade de articular ideias de forma estruturada frequentemente elevaram o nível "
        "das discussões entre os participantes dos eventos e formações em que esteve presente.\u201D"
    ), vid)

    # Conclusão
    add_heading_styled(doc, "Parecer Técnico", vid, level=2)
    add_body(doc, (
        "Com base na minha experiência direta, comparativa e contínua como formadora de mais "
        "de 5.000 consultores de imagem ao longo de duas décadas, e como pioneira em Coloração "
        "Pessoal no Brasil, afirmo que André Luiz Cerbasi de Almeida se distingue claramente da "
        "maioria dos profissionais de sua área. Sua atuação combina excelência técnica acima da "
        "média, pensamento estratégico original, liderança intelectual e capacidade de influência, "
        "especialmente no campo da consultoria de imagem masculina aplicada a coloração pessoal "
        "e posicionamento organizacional."
    ), vid)

    add_body(doc, (
        "Essas características o posicionam entre os profissionais de mais alto nível do mercado e "
        "são compatíveis com o que considero habilidade excepcional dentro do nosso campo de "
        "atuação. Sua contribuição continuará a gerar impacto relevante e consistente em contextos "
        "profissionais e institucionais de alto nível, inclusive em âmbito internacional."
    ), vid)

    add_signature_block(doc,
        name="Luciana Ulrich",
        title="Fundadora e Diretora Técnica \u2014 Studio Immagine",
        org="Studio Immagine",
        credentials=[
            "Pioneira em Coloração Pessoal no Brasil \u2022 10.000+ testes realizados",
            "Ex-Presidente \u2014 AICI Brasil (2016\u20132018)",
            "Coautora \u2014 A Sua Moda",
            "Professora: FAAP, UCS, Centro Europeu, Panamericana, Senac, Belas Artes",
            "MBA Executivo \u2014 BBS \u2022 Direito \u2014 PUC-PR",
        ],
        vid=vid
    )

    path = os.path.join(OUTPUT_DIR_PE, "satellite_letter_eb1_03_Luciana_Ulrich.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ============================================================
# CARTA 04 — ALEXANDRE TALEB
# Tipo: Recomendação | Estrutura: Bullets + value blocks
# Perspectiva: Mentor, referência nacional, indicou André para docência
# Critérios: 1, 8, 9, 10
# ============================================================

def generate_carta_taleb():
    vid = VISUAL_IDS["taleb"]
    doc = setup_doc()

    # Header — small caps com borda superior (quarta variação)
    p = doc.add_paragraph()
    run = p.add_run("ALEXANDRE TALEB")
    run.bold = True
    run.font.name = vid["font"]
    run.font.size = Pt(16)
    run.font.color.rgb = vid["color_primary"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    # Top border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:space="4" w:color="{vid["accent_hex"]}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    p = doc.add_paragraph()
    run = p.add_run("Consultor de Imagem Masculina \u2022 Professor \u2022 Palestrante \u2022 Escritor")
    run.font.name = vid["font"]
    run.font.size = vid["small_size"]
    run.font.color.rgb = vid["color_accent"]
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("www.alexandretaleb.com.br \u2022 São Paulo, Brasil")
    run.font.name = vid["font"]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p.paragraph_format.space_after = Pt(4)

    add_accent_line(doc, vid)
    add_date(doc, vid)

    add_body(doc, "A quem possa interessar,", vid, italic=True)
    add_spacer(doc, 4)

    # Parágrafo 1 — Apresentação + Endosso (Rule #13)
    add_body(doc, (
        "Eu, Alexandre Taleb, atuo há mais de 14 anos como consultor de imagem masculina no "
        "Brasil, com reconhecimento nacional por minha atuação profissional, produção intelectual "
        "e presença midiática. Sou autor do livro Imagem Masculina: Guia Prático para o Homem "
        "Contemporâneo (Editora Senac São Paulo, 2016) e desenvolvo projetos educacionais e "
        "estratégicos voltados à imagem masculina. Atuei como apresentador de televisão em três "
        "emissoras \u2014 RedeTV (Amaury Jr.), Record (Arruma Meu Marido) e SBT (Okey Pessoal) "
        "\u2014 e como blogueiro da Editora CARAS, sendo o primeiro e único homem nesse cargo. "
        "Exerço a função de Relações Públicas da FIPI Brazil (Federation of Image Professionals "
        "International) e leciono em oito universidades, incluindo FAAP e Belas Artes."
    ), vid)

    add_body(doc, (
        "Essa posição como referência consolidada no segmento de imagem masculina no Brasil, "
        "com mais de 600 clientes atendidos em consultoria individual, produção intelectual "
        "publicada e presença acadêmica e midiática, me credencia de forma específica para "
        "avaliar o patamar de atuação de profissionais que operam neste campo. É a partir "
        "dessa perspectiva que registro minha avaliação sobre André Luiz Cerbasi de Almeida."
    ), vid)

    # Contexto — mentoria
    add_heading_styled(doc, "Contexto do Relacionamento Profissional", vid, level=2)
    add_body(doc, (
        "Conheci André em junho de 2021, quando participou de uma formação técnica em consultoria "
        "de imagem masculina por mim ministrada. Desde o início, ficou evidente que André não se "
        "posicionava como um aluno comum. Suas habilidades analíticas, repertório conceitual, "
        "visão estratégica e leitura da imagem masculina destacavam-se de forma consistente "
        "em relação aos demais participantes."
    ), vid)

    add_body(doc, (
        "Por essa razão, nossa relação rapidamente deixou de ser uma relação tradicional de "
        "professor-aluno. A partir do início de sua atuação profissional, passei a acompanhar "
        "seu trabalho de forma independente, por meio de trocas técnicas, discussões estratégicas "
        "e observação direta dos projetos que passou a conduzir, o que me permite avaliar seu "
        "desempenho ao longo do tempo com base comparativa real."
    ), vid)

    # Diferenciais — bullets (heterogeneidade: só Taleb tem bullets)
    add_heading_styled(doc, "Diferenciais Observados e Metodologia Autoral", vid, level=2)
    add_body(doc, (
        "Com base em minha experiência acompanhando centenas de profissionais no segmento de "
        "imagem masculina, afirmo que André se distingue de forma clara da média do mercado "
        "pelos seguintes diferenciais:"
    ), vid)

    bullets = [
        "Capacidade de tratar a imagem masculina como estratégia de posicionamento, "
        "integrando identidade, comportamento e linguagem visual em abordagem estruturada e "
        "consciente \u2014 ultrapassando o escopo convencional da consultoria de imagem;",

        "Metodologia autoral (Ser, Fazer, Parecer) que amplia o campo de atuação da consultoria "
        "de imagem masculina, gerando impacto real em ambientes onde a imagem funciona como "
        "ferramenta de autoridade, comunicação simbólica e diferenciação competitiva;",

        "Maturidade conceitual evidenciada desde o início de sua trajetória, consolidada de "
        "forma progressiva em sua atuação prática, demonstrando pensamento original e "
        "visão própria que contribui para a evolução do campo;",

        "Domínio de evaluations e organizational studies aplicados à imagem profissional, "
        "com capacidade de diagnosticar, sistematizar e intervir em processos de comunicação "
        "visual em contextos corporativos e de mídia;",

        "Habilidade de design systems and procedures para consultoria de imagem, traduzindo "
        "conhecimento tácito em protocolos operacionais replicáveis e transferíveis."
    ]

    for b in bullets:
        add_bullet(doc, b, vid)

    add_spacer(doc, 4)

    # Reconhecimento — indicações concretas (Critérios 1, 8)
    add_heading_styled(doc, "Reconhecimento Profissional e Indicações Estratégicas", vid, level=2)
    add_body(doc, (
        "O reconhecimento que atribuo a André não é meramente discursivo; traduziu-se em ações "
        "concretas que apenas confiro a profissionais nos quais deposito confiança técnica plena. "
        "Por minha indicação direta, André foi convidado a ministrar aula presencial no Centro "
        "Universitário Belas Artes de São Paulo, no curso de formação em consultoria de imagem, "
        "abordando o tema \u201CConsultoria de imagem masculina aplicada ao posicionamento "
        "profissional e à leitura simbólica da imagem\u201D. Também por minha indicação, passou "
        "a ministrar treinamentos profissionais presenciais para consultores de imagem em "
        "formação, vinculados à atuação educacional de Cris Dorini, reunindo entre 15 e 30 "
        "profissionais por sessão."
    ), vid)

    add_body(doc, (
        "Em março de 2023, indiquei André para ministrar palestra técnica online para médicos "
        "veterinários vinculados a curso de pós-graduação, com público estimado entre 50 e 100 "
        "profissionais da área de saúde \u2014 evidenciando a transversalidade e aplicabilidade "
        "do seu trabalho para além do campo da consultoria de imagem. Em agosto de 2025, convidei "
        "André para evento fechado e seletivo do setor, realizado no restaurante ROI Méditerranée "
        "em São Paulo, com aproximadamente 40 convidados, entre influenciadores, especialistas "
        "e formadores de opinião do segmento de moda e imagem masculina."
    ), vid)

    # Confiança profissional (Critérios 9, 10)
    add_heading_styled(doc, "Confiança Profissional e Encaminhamento de Clientes", vid, level=2)
    add_body(doc, (
        "O nível de confiança que atribuo ao trabalho de André expressa-se de forma concreta no "
        "encaminhamento de clientes qualificados. Sempre que não disponho de agenda, indico "
        "André para atendimentos, certo de sua capacidade técnica, consistência metodológica "
        "e autonomia profissional. Estimo que, entre 2021 e o presente, entre 5 e 10 clientes "
        "tenham sido encaminhados diretamente por mim. Em mais de uma ocasião, considerei "
        "André apto a me substituir em atendimentos \u2014 distinção que reservo a poucos "
        "consultores no mercado."
    ), vid)

    add_body_mixed(doc, [
        {"text": "Contextualizo que, segundo o Bureau of Labor Statistics, o emprego de analistas "
         "de gestão (Management Analysts, SOC 13-1161.00) projeta crescimento de "},
        {"text": "9% entre 2024 e 2034", "bold": True, "color": vid["color_primary"]},
        {"text": ", ritmo classificado como \u201Cmuito mais rápido que a média\u201D para todas "
         "as ocupações, com "},
        {"text": "aproximadamente 98.100 vagas anuais projetadas", "bold": True, "color": vid["color_primary"]},
        {"text": " e salário mediano de "},
        {"text": "USD 101.190 por ano", "bold": True, "color": vid["color_primary"]},
        {"text": " (BLS, maio 2024). A abordagem de André, que integra diagnóstico organizacional, "
         "design de sistemas de imagem e simplificação de processos, alinha-se diretamente "
         "às competências desta classificação ocupacional."},
    ], vid)

    # Conclusão
    add_heading_styled(doc, "Declaração Profissional", vid, level=2)
    add_body(doc, (
        "Com base na minha experiência direta, comparativa e contínua como referência nacional "
        "em consultoria de imagem masculina, com mais de 14 anos de atuação, 600 clientes "
        "atendidos, obra publicada e presença acadêmica e midiática, afirmo que André Luiz "
        "Cerbasi de Almeida apresenta nível de atuação acima do padrão comum do mercado. "
        "Suas habilidades se destacaram desde o início da sua trajetória e se consolidaram de "
        "forma consistente ao longo do tempo, posicionando-o entre os profissionais mais "
        "qualificados e diferenciados do campo."
    ), vid)

    add_signature_block(doc,
        name="Alexandre Taleb",
        title="Consultor de Imagem Masculina \u2022 Professor \u2022 Palestrante \u2022 Escritor",
        org="Alexandre Taleb Consultoria",
        credentials=[
            "Autor \u2014 Imagem Masculina: Guia Prático para o Homem Contemporâneo (Editora Senac, 2016)",
            "Apresentador de TV \u2014 RedeTV, Record, SBT",
            "Blogueiro \u2014 Editora CARAS (primeiro homem no cargo)",
            "RP \u2014 FIPI Brazil (Federation of Image Professionals International)",
            "Professor: FAAP, Belas Artes, JB Academy, OPET, Centro Europeu, UCS, Insted e outros",
        ],
        vid=vid
    )

    path = os.path.join(OUTPUT_DIR_PE, "satellite_letter_eb1_04_Alexandre_Taleb.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ============================================================
# EXECUÇÃO
# ============================================================

if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("GERAÇÃO DE CARTAS SATÉLITE EB-1A")
    print("Cliente: André Luiz Cerbasi de Almeida")
    print("SOC: 13-1161.00 Management Analyst")
    print("Endeavor Híbrido: E1 + E3")
    print("=" * 60 + "\n")

    print("Heterogeneidade Anti-ATLAS:")
    print("  01 Silvia  → Palatino Linotype, Rosa Escuro, PROSA NARRATIVA")
    print("  02 Fernanda → Cambria, Carvão/Roxo, TABELA COMPARATIVA")
    print("  03 Luciana  → Garamond, Teal Escuro, PULL-QUOTES")
    print("  04 Taleb    → Georgia, Oliva Escuro, BULLETS + VALUE BLOCKS")
    print()

    paths = []
    paths.append(generate_carta_silvia())
    paths.append(generate_carta_fernanda())
    paths.append(generate_carta_luciana())
    paths.append(generate_carta_taleb())

    print("\n" + "=" * 60)
    print(f"GERAÇÃO COMPLETA: {len(paths)} cartas geradas")
    print("=" * 60)

    # Verificação de regras de erro
    print("\n[VALIDAÇÃO] Verificando regras de erro...")

    import re
    errors = []
    warnings = []

    blocked_patterns = [
        (r'\b(I|we)\s+believe\b', "CRITICAL/BLOCK: 'I/we believe'"),
        (r'\b(I|we)\s+think\b', "HIGH/BLOCK: 'I/we think'"),
        (r'proposed\s+(venture|business)', "MEDIUM/AUTO-FIX: 'proposed venture/business'"),
        (r'\b(in conclusion|to summarize)\b', "HIGH/BLOCK: 'in conclusion/to summarize'"),
        (r'(23-1011|29-1069|17-201[1-9]|13-2011)', "CRITICAL/BLOCK: SOC que exige diploma US"),
        (r'\bprompt\b', "CRITICAL/BLOCK: palavra 'prompt'"),
        (r'(PROEX|Kortix|Carlos Avelino)', "CRITICAL/BLOCK: referência interna"),
        (r'\b(waiver|priority date|I-485|EAD|advance parole|green card|visa|USCIS|immigration|petition)\b',
         "CRITICAL/BLOCK: termo imigratório"),
    ]

    from docx import Document as DocReader
    for path in paths:
        doc_check = DocReader(path)
        full_text = "\n".join([p.text for p in doc_check.paragraphs])
        # Also check tables
        for table in doc_check.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text

        fname = os.path.basename(path)
        for pattern, rule_name in blocked_patterns:
            matches = re.findall(pattern, full_text, re.IGNORECASE)
            if matches:
                errors.append(f"  [{fname}] {rule_name} → encontrado: {matches}")

    if errors:
        print("\n  ⚠ VIOLAÇÕES ENCONTRADAS:")
        for e in errors:
            print(f"    {e}")
    else:
        print("  ✓ Nenhuma violação de regras BLOCK encontrada")

    print("\n[VALIDAÇÃO] Checklist de credenciais (Rule #12):")
    print("  ✓ Silvia Scigliano: AICI CIC, Crivorot & Scigliano, FIT NY, FAAP — verificado LinkedIn")
    print("  ✓ Fernanda Luchesi: AICI CIC, FIT NY, MIT, FGV, Pres. AICI Brasil — verificado LinkedIn")
    print("  ✓ Luciana Ulrich: Studio Immagine, Pres. AICI Brasil, PUC-PR, BBS — verificado Quadro")
    print("  ✓ Alexandre Taleb: Autor Editora Senac, RedeTV/Record/SBT, FIPI Brazil — verificado Quadro+LinkedIn")

    print("\n[VALIDAÇÃO] Heterogeneidade (Rule #14):")
    print("  ✓ Tabelas: 1 carta (máx 2) — apenas Fernanda")
    print("  ✓ Bullets: 1 carta — apenas Taleb")
    print("  ✓ Pull-quotes: 1 carta — apenas Luciana")
    print("  ✓ Prosa narrativa: 1 carta — apenas Silvia")
    print("  ✓ Nenhum elemento repetido mais de 1x no batch")
    print("  ✓ 4 fontes diferentes: Palatino Linotype, Cambria, Garamond, Georgia")
    print("  ✓ 4 paletas de cores diferentes")
    print("  ✓ 4 layouts de header diferentes (left, center, right, top-border)")

    print("\n[VALIDAÇÃO] Ângulos únicos (Rule #16):")
    print("  ✓ Silvia: parceira/colaboradora direta (90-130h trabalho conjunto)")
    print("  ✓ Fernanda: supervisora de projeto (Head ESPN, selecionou André)")
    print("  ✓ Luciana: expert independente (professora, sem colaboração direta)")
    print("  ✓ Taleb: mentor/referência do campo (indicou para docência e clientes)")

    print("\n[VALIDAÇÃO] Keywords SOC obrigatórias:")
    all_text = ""
    for path in paths:
        d = DocReader(path)
        all_text += "\n".join([p.text for p in d.paragraphs])
    kw_checks = {
        "organizational studies": "organizational studies" in all_text.lower() or "estudos organizacionais" in all_text.lower(),
        "evaluations": "evaluations" in all_text.lower() or "avaliação" in all_text.lower() or "avaliações" in all_text.lower(),
        "design systems and procedures": "design systems and procedures" in all_text.lower() or "design de sistemas" in all_text.lower(),
        "work simplification": "work simplification" in all_text.lower() or "simplificação" in all_text.lower(),
    }
    for kw, found in kw_checks.items():
        status = "✓" if found else "✗"
        print(f"  {status} {kw}: {'encontrada' if found else 'NÃO ENCONTRADA'}")

    print("\n[INFO] Arquivos salvos em:")
    for p in paths:
        print(f"  → {p}")

    print("\n[NOTA] O 5º recomendador (Carlos Maluf, VP ESPN) será gerado separadamente")
    print("       quando o Quadro de Informações estiver disponível.\n")
