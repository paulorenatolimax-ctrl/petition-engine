#!/usr/bin/env python3
"""
DentalShield Ops — Comprehensive Product & Technology Dossier
Client: Camilla Santana Pereira Paes de Barros
Output: Professional .docx (investor/customer/institutional grade)
Data source: https://camilla-sparkle-charm.lovable.app/ (authoritative)
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ─── CONFIG ───────────────────────────────────────────────────────
OUTPUT_DIR = "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2025/CAMILLA/_Forjado por Petition Engine"
OUTPUT_FILE = "saas_evidence_Camilla_Santana_Pereira_Paes_de_Barros.docx"
OUTPUT_PATH = os.path.join(OUTPUT_DIR, OUTPUT_FILE)

# Colors (Navy + Gold brand palette)
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x1F, 0x29, 0x37)
MEDIUM_GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_NAVY = RGBColor(0xE8, 0xEB, 0xF0)

FONT_MAIN = "Calibri"
FONT_HEADING = "Calibri"

# ─── HELPERS ──────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_text(doc, text, font_size=11, bold=False, color=DARK_TEXT,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6),
             space_before=Pt(0), italic=False):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = FONT_MAIN
    run.font.italic = italic
    return p

def add_heading_styled(doc, text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = FONT_HEADING
    return h

def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500" * 60)
    run.font.color.rgb = GOLD
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def add_kpi_table(doc, kpis):
    table = doc.add_table(rows=2, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(kpis):
        cell_val = table.cell(0, i)
        cell_val.text = ""
        p = cell_val.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(value))
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run.font.name = FONT_HEADING
        set_cell_shading(cell_val, "F9FAFB")
        cell_lab = table.cell(1, i)
        cell_lab.text = ""
        p2 = cell_lab.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = MEDIUM_GRAY
        r2.font.name = FONT_MAIN
        set_cell_shading(cell_lab, "F9FAFB")
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)
    return table

def add_data_table(doc, headers, rows, highlight_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.name = FONT_MAIN
        set_cell_shading(cell, "1B2A4A")
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_TEXT
            run.font.name = FONT_MAIN
            if highlight_col is not None and c_idx == highlight_col:
                run.font.bold = True
                run.font.color.rgb = NAVY
            bg = "FFFFFF" if r_idx % 2 == 0 else "F3F4F6"
            set_cell_shading(cell, bg)
    return table

def add_rich_para(doc, segments, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6)):
    """Add paragraph with mixed formatting. segments = list of (text, bold, italic, color, size)."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    for seg in segments:
        text, bold, italic, color, size = seg
        run = p.add_run(text)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        run.font.size = Pt(size)
        run.font.name = FONT_MAIN
    return p


# ─── MAIN GENERATION ─────────────────────────────────────────────

def generate():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_MAIN
    font.size = Pt(11)
    font.color.rgb = DARK_TEXT

    # ══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DENTALSHIELD OPS")
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.name = FONT_HEADING

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INTEGRATED COMPLIANCE PLATFORM")
    run.font.size = Pt(18)
    run.font.color.rgb = GOLD
    run.font.name = FONT_HEADING
    run.font.bold = True

    doc.add_paragraph()

    add_text(doc,
        "Product & Technology Dossier",
        font_size=16, color=DARK_TEXT,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    add_text(doc,
        "The only all-in-one solution combining physical IoT compliance\n"
        "infrastructure with cloud operational software for dental practices.",
        font_size=12, color=MEDIUM_GRAY, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    for _ in range(5):
        doc.add_paragraph()

    add_text(doc,
        "DentalShield Systems, LLC\nLakeland, Florida | Central Florida Operations Hub",
        font_size=10, color=MEDIUM_GRAY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    add_text(doc,
        "Founded by Camilla Santana Pereira Paes de Barros\n"
        "Licensed Dentist (CRO-DF, Brazil) | US-Certified Dental Professional (DAPA, FL)",
        font_size=10, color=MEDIUM_GRAY, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

    add_text(doc,
        "CONFIDENTIAL \u2014 FOR INSTITUTIONAL AND INVESTOR REVIEW ONLY",
        font_size=9, bold=True, color=RGBColor(0xE7, 0x4C, 0x3C),
        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Overview",
        "2. The Systemic Problem: Dental Compliance in America",
        "3. Platform Architecture: Hardware + Software Integration",
        "4. The 156-Point Proprietary Compliance Framework",
        "5. Founder Profile & Technical Authority",
        "6. Competitive Landscape",
        "7. Pricing & Revenue Model",
        "8. Financial Projections",
        "9. Go-to-Market Strategy",
        "10. National Impact & Regulatory Alignment",
        "11. Technical Specifications",
        "12. Sources & Regulatory Citations",
    ]
    for item in toc_items:
        add_text(doc, item, font_size=11, color=NAVY, space_after=Pt(4),
                 alignment=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. EXECUTIVE OVERVIEW
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "1. Executive Overview", level=1)

    add_kpi_table(doc, [
        ("U.S. Dental Market", "$165B"),
        ("Market by 2034", "$196.5B"),
        ("Americans in Dental HPSAs", "63.7M"),
        ("Y5 Revenue Target", "$7.86M"),
    ])

    doc.add_paragraph()

    add_text(doc,
        "DentalShield Ops is a vertically integrated dental compliance platform that combines "
        "three proprietary IoT hardware devices with five cloud-connected software modules to "
        "address the structural compliance gap exposing American dental practices to regulatory "
        "penalties, patient safety failures, and operational inefficiency. Headquartered in "
        "Lakeland, Florida, the company targets Central Florida\u2019s high-need dental corridors "
        "before scaling to Miami-Fort Lauderdale and national markets.")

    add_text(doc,
        "The U.S. dental industry represents a $165 billion market projected to reach $196.5 "
        "billion by 2034, growing at a 7.9% CAGR (Mordor Intelligence, 2024). Within this market, "
        "63.7 million Americans reside in federally designated dental Health Professional Shortage "
        "Areas (HPSAs), and 10,744 additional dentists are required to eliminate these designations "
        "(HRSA, December 2025). Practices in these underserved regions face compounding pressure: "
        "workforce shortages, rising OSHA enforcement\u2014with penalties reaching $16,550 per serious "
        "violation and $165,514 per willful violation as of January 2025\u2014and an increasingly "
        "complex web of CDC, HIPAA, and EPA compliance requirements.")

    add_text(doc,
        "Unlike competitors that offer software-only dashboards (Compliancy Group, iLearn Dental) "
        "or generic advisory services (SafeLink), DentalShield Ops is the only platform that "
        "bridges the physical-digital divide with purpose-built IoT hardware for dental sterilization "
        "environments. The platform operates on a tiered subscription model\u2014Starter at $299/month, "
        "Professional at $499/month, Elite at $799/month, and custom Enterprise pricing for dental "
        "service organizations (DSOs)\u2014with implementation fees ranging from $1,500 to $5,000. "
        "Year 1 projections target 50 new clinic implementations generating $940,000 in total revenue, "
        "scaling to 780 active subscriptions and $7.86 million in annual revenue by Year 5.")

    add_text(doc,
        "The entire system was architected by Camilla Santana Pereira Paes de Barros, whose dual "
        "clinical background\u2014licensed dentist with four years of direct sterilization protocol "
        "oversight in Brazil (CRO-DF, UNIEURO 2021) and US-certified dental professional through "
        "DAPA (Florida, 2025)\u2014gives her the rare technical authority to design compliance workflows "
        "that bridge international best practices with American regulatory requirements. Her "
        "proprietary 156-point audit protocol, developed through hands-on clinical experience at "
        "Odonto e Arte (Bras\u00edlia, 2022\u20132026), is the intellectual foundation of every module "
        "in the platform.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. THE SYSTEMIC PROBLEM
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "2. The Systemic Problem: Dental Compliance in America", level=1)

    add_heading_styled(doc, "2.1 The Workforce Crisis", level=2)

    add_text(doc,
        "The American dental workforce faces a structural deficit that compounds patient safety "
        "risks. As of December 2025, the Health Resources and Services Administration (HRSA) "
        "reports 7,054 dental HPSAs across the country, affecting 63.7 million people\u2014roughly "
        "one in five Americans. Eliminating these shortage designations would require 10,744 "
        "additional dentists, a gap that is widening as the profession confronts an accelerating "
        "retirement wave: the average dentist retirement age has risen to 68.7 years (up from "
        "64.7 in 2001), and 21% of practice managers plan to retire within six years (ADA "
        "Workforce Report, 2025; Pearl AI, 2026).")

    add_text(doc,
        "The rural-urban disparity is severe: rural areas average 32.7 dentists per 100,000 "
        "residents compared to 64.7 in urban areas\u2014a 2:1 gap. Sixty-six percent of all dental "
        "HPSAs are located in rural communities, affecting approximately 19.4 million rural "
        "residents (RuralHealthInfoHub, January 2026). Florida alone has 274 designated dental "
        "HPSAs\u2014the fourth highest in the nation.")

    add_heading_styled(doc, "National Dental Workforce Indicators", level=3)

    add_data_table(doc,
        headers=["Indicator", "Value", "Source"],
        rows=[
            ["Dental HPSAs (nationwide)", "7,054", "HRSA, March 2025"],
            ["Population in dental HPSAs", "63.7 million", "HRSA, December 2025"],
            ["Additional dentists needed", "10,744", "HRSA, December 2025"],
            ["Active dentists nationwide", "202,485", "ADA, 2024"],
            ["Dental Assistants (SOC 29-1292.00)", "342,000", "BLS OES, 2024"],
            ["Annual DA openings projected", "52,900/year", "BLS, 2024\u20132034"],
            ["Rural vs. urban density", "32.7 vs. 64.7 per 100K", "HRSA, 2025"],
            ["Dental hygienist growth outlook", "7% (faster than average)", "BLS, 2024\u20132034"],
            ["Practice managers retiring \u226410 yrs", "43%", "ADA, 2025"],
        ])

    doc.add_paragraph()

    add_heading_styled(doc, "2.2 The Regulatory Burden", level=2)

    add_text(doc,
        "Dental practices must simultaneously comply with multiple overlapping regulatory "
        "frameworks: OSHA\u2019s Bloodborne Pathogens Standard (29 CFR 1910.1030), the CDC\u2019s "
        "Guidelines for Infection Control in Dental Health-Care Settings (MMWR 52:RR-17), "
        "HIPAA\u2019s Security Rule (45 CFR 164), EPA waste disposal requirements, and state-specific "
        "sterilization monitoring mandates. The 2026 OSHA compliance updates add new requirements "
        "for aerosol-generating procedures\u2014ultrasonic scaling, high-speed drilling, and air "
        "polishing\u2014further expanding the compliance surface practices must manage.")

    add_text(doc,
        "Non-compliance carries severe financial consequences. As of January 15, 2025, OSHA "
        "penalties reach $16,550 per serious violation and $165,514 per willful or repeated "
        "violation (OSHA Annual Adjustments, 2025). OSHA enforcement data from 2024 shows that "
        "40% of small dental clinics receiving a major OSHA fine close within two years. For a "
        "two-chair practice generating $500,000 annually, a single willful violation represents "
        "over 33% of annual revenue\u2014a potentially fatal financial event.")

    add_heading_styled(doc, "2.3 The Infection Control Gap", level=2)

    add_text(doc,
        "The CDC estimates over 8,000 dental procedure-related infections annually in the United "
        "States. Dental Assistants face 10 times higher bloodborne pathogen exposure risk than "
        "the general population (OSHA, 2024). The CDC\u2019s 2003 Guidelines for Infection Control "
        "remain the standard of practice, supplemented by the updated \u201cSummary of Infection "
        "Prevention Practices in Dental Settings: Basic Expectations for Safe Care\u201d and the "
        "Association for Dental Safety\u2019s 2025 Edition implementation guide. Despite these "
        "resources, most small practices lack the staffing and systems to maintain continuous "
        "compliance across all required protocols.")

    add_text(doc,
        "Published research demonstrates sterilizer failure rates as high as 64.7% during "
        "biological indicator testing (PubMed PMID 2370583). A 2024 retrospective study found "
        "chemical vapor sterilizer failure rates of 43.8% and steam sterilizer failure rates "
        "of 6.5%, with common causes including overloading, improper instrument loading, and "
        "incorrect biological indicator placement (PMC, 2024). These failures are invisible "
        "without real-time monitoring\u2014the exact capability that DentalShield Ops\u2019 "
        "SteriSensor hardware provides.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. PLATFORM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "3. Platform Architecture: Hardware + Software Integration", level=1)

    add_text(doc,
        "DentalShield Ops is built on a three-tier architecture: physical IoT devices at the "
        "clinic level, a cloud-based data ingestion and processing layer, and a client-facing "
        "SaaS dashboard. All data flows from IoT hardware through encrypted channels to the "
        "cloud platform, where it is processed, stored, and surfaced through the ComplianceScreen "
        "interface and web/mobile dashboards.")

    add_heading_styled(doc, "3.1 Hardware Layer: Three Proprietary IoT Devices", level=2)

    devices = [
        ("SteriSensor", "Real-Time Autoclave Monitoring",
         "Continuously tracks temperature, pressure, and cycle duration for every sterilization "
         "run via wireless IoT sensors. Automatically generates digital audit logs compliant with "
         "29 CFR 1910.1030, replacing manual paper-based spore testing records. Each sensor "
         "transmits data to the ComplianceScreen dashboard, creating an unbroken chain of custody "
         "for sterilization verification. Connectivity: Wi-Fi/BLE wireless."),

        ("BiohazardBox", "Smart Biomedical Waste Tracking",
         "RFID chain-of-custody tracking for sharps and biomedical waste disposal. Each disposal "
         "event is logged with operator ID, timestamp, waste category, and container fill level. "
         "Automatically triggers replacement alerts and generates EPA-compliant waste manifests, "
         "eliminating the manual tracking errors that account for a significant share of OSHA "
         "waste-management citations. Connectivity: RFID + Wi-Fi."),

        ("ComplianceScreen", "Dedicated Live Dashboard",
         "Wall-mounted tablet displaying real-time compliance status using red/yellow/green "
         "indicators for each compliance domain (sterilization, waste, PPE, documentation, "
         "training). Push notifications and email alerts fire when protocol deviations are "
         "detected. Designed for always-on visibility in the clinical workspace. "
         "Connectivity: Wi-Fi."),
    ]

    for name, subtitle, desc in devices:
        add_rich_para(doc, [
            (f"{name} ", True, False, NAVY, 12),
            (f"\u2014 {subtitle}", False, True, GOLD, 10),
        ], space_after=Pt(2))
        add_text(doc, desc, font_size=10.5, space_after=Pt(10))

    add_heading_styled(doc, "3.2 Software Layer: Five Integrated Modules", level=2)

    modules = [
        ("SmartSupply", "Inventory Management",
         "Unified inventory dashboard with predictive alerts based on historical usage patterns. "
         "API integration with Henry Schein and Patterson Dental supply chains. Eliminates "
         "15\u201325% of supply waste through automated reorder optimization and expiration date "
         "tracking. Transforms reactive purchasing into data-driven procurement."),

        ("SteriTrack", "Sterilization Compliance (Core Module)",
         "Complete digital sterilization cycle record: temperature, pressure, duration, and "
         "operator ID for every autoclave run. Auto-generated audit logs meeting OSHA 29 CFR "
         "1910.1030 and CDC MMWR 52:RR-17. The only module backed by physical IoT hardware "
         "(SteriSensor), creating a closed-loop verification system that eliminates the gap "
         "between protocol and proof. Targets 70\u201380% reduction in sterilization-related "
         "compliance violations."),

        ("EquipMax", "Preventive Maintenance",
         "Equipment maintenance scheduling based on manufacturer-specified intervals and "
         "historical performance data. Repair-versus-replacement cost analysis for capital "
         "equipment decisions. Targets the 12\u201318 hours per month of chair time lost to "
         "unplanned equipment failure in average dental practices, directly recovering revenue "
         "through uptime optimization."),

        ("TeamSync", "Workforce Management",
         "Equitable workload distribution scheduling integrated with practice management "
         "systems (PMS). Addresses the 35\u201340% annual Dental Assistant turnover rate through "
         "balanced shift allocation and workload visibility. Tracks certification expiration, "
         "required training hours, and continuing education compliance for all clinical staff."),

        ("OSHA Audit Protection Guarantee", "Financial Coverage",
         "Financial coverage of up to $15,000 for covered OSHA violations, active while "
         "subscription is current and IoT logs confirm correct system usage. Professional tier "
         "covers up to $10,000; Elite tier covers up to $15,000. This guarantee is made possible "
         "because the IoT monitoring data provides verifiable proof of compliance\u2014the "
         "platform\u2019s hardware creates the audit trail that underwrites the financial protection."),
    ]

    for name, subtitle, desc in modules:
        add_rich_para(doc, [
            (f"{name} ", True, False, NAVY, 11),
            (f"({subtitle})", False, True, GOLD, 10),
        ], space_after=Pt(2))
        add_text(doc, desc, font_size=10.5, space_after=Pt(10))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. 156-POINT FRAMEWORK
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "4. The 156-Point Proprietary Compliance Framework", level=1)

    add_text(doc,
        "The intellectual core of DentalShield Ops is a proprietary 156-point compliance audit "
        "protocol developed by Camilla Santana Pereira Paes de Barros through four years of "
        "direct clinical sterilization management at Odonto e Arte (Bras\u00edlia, 2022\u20132026). "
        "This framework is not a checklist downloaded from a regulatory agency\u2014it is an "
        "integrated assessment methodology that maps every physical touchpoint in a dental "
        "practice to its corresponding regulatory requirement across OSHA, CDC, HIPAA, and "
        "EPA domains.")

    add_text(doc,
        "The 156 audit points are organized across six compliance domains, each informed by "
        "the founder\u2019s clinical expertise in infection control, sterilization protocol design, "
        "and biomedical waste management. The framework cross-references Brazilian ANVISA "
        "standards\u2014among the most rigorous sterilization protocols in Latin America\u2014with "
        "U.S. regulatory requirements, identifying gaps that neither system addresses in "
        "isolation. This dual-regulatory perspective is what makes the framework uniquely "
        "comprehensive.")

    add_heading_styled(doc, "156-Point Audit Protocol: Domain Breakdown", level=3)

    add_data_table(doc,
        headers=["Compliance Domain", "Points", "Key Regulations", "IoT Module"],
        rows=[
            ["Sterilization & Infection Control", "42", "29 CFR 1910.1030, MMWR 52:RR-17", "SteriSensor"],
            ["Biomedical Waste Management", "28", "EPA RCRA, state HW regulations", "BiohazardBox"],
            ["PPE & Exposure Control", "24", "29 CFR 1910.132\u2013138, OSHA BPS", "ComplianceScreen"],
            ["Documentation & Record-Keeping", "22", "HIPAA 45 CFR 164, OSHA logs", "All modules"],
            ["Equipment Safety & Maintenance", "21", "ADA Standards, MFR specs", "EquipMax"],
            ["Staff Training & Certification", "19", "OSHA annual training, CE reqs", "TeamSync"],
        ])

    doc.add_paragraph()

    add_text(doc,
        "Each audit point has a defined assessment procedure, pass/fail criteria, remediation "
        "protocol, and ongoing monitoring frequency. The framework is not static: it evolves "
        "continuously as the founder incorporates new regulatory updates, field observations "
        "from active implementations, and emerging compliance requirements (such as the 2026 "
        "aerosol-generating procedure mandates). This living methodology requires the founder\u2019s "
        "ongoing clinical judgment and regulatory interpretation to remain current, accurate, "
        "and defensible.")

    add_text(doc,
        "The dual-market expertise\u2014Brazilian ANVISA (RDC 15/2012, RDC 222/2018) combined with "
        "American OSHA/CDC frameworks\u2014cannot be replicated through hiring or training alone. "
        "It represents years of hands-on clinical practice in both regulatory environments, "
        "synthesized into a unified framework that incorporates the most protective elements "
        "of each system.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. FOUNDER PROFILE
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "5. Founder Profile & Technical Authority", level=1)

    add_text(doc,
        "Camilla Santana Pereira Paes de Barros brings a combination of clinical depth and "
        "regulatory fluency that positions her as the indispensable technical authority behind "
        "DentalShield Ops. Her qualifications span two regulatory ecosystems\u2014Brazilian ANVISA "
        "standards and American OSHA/CDC frameworks\u2014giving her the cross-jurisdictional insight "
        "to design compliance systems that exceed any single regulatory baseline.")

    add_heading_styled(doc, "Academic & Clinical Credentials", level=3)

    add_data_table(doc,
        headers=["Credential", "Institution", "Detail"],
        rows=[
            ["Doctor of Dental Surgery (DDS)", "UNIEURO, Bras\u00edlia", "CRO-DF Licensed, 2021"],
            ["Post-Graduate: Endodontics", "IOA \u2014 885 hours", "Completed 2021\u20132022"],
            ["Specialization: Orthodontics", "IOA", "In progress, 2023\u20132026"],
            ["US Dental Professional Cert.", "DAPA, Orlando, FL", "March 2025 \u2014 Lic. #5705"],
            ["Advanced Surgical Rehabilitation", "Clinical training", "Prosthetic & functional"],
            ["Cosmetic Procedures Certification", "Clinical training", "Botox, harmonization"],
        ])

    doc.add_paragraph()

    add_heading_styled(doc, "Clinical Leadership: Odonto e Arte (2022\u20132026)", level=3)

    add_text(doc,
        "As Technical Director of Odonto e Arte in Bras\u00edlia for four years, Camilla directly "
        "supervised all sterilization protocols, biohazard management procedures, and regulatory "
        "documentation for the practice. This hands-on leadership role is where the 156-point "
        "audit protocol was developed through iterative refinement of real-world infection "
        "control workflows. She identified systematic gaps in existing compliance approaches\u2014"
        "documentation delays, inconsistent sterilization verification, manual waste tracking "
        "errors\u2014and designed integrated solutions that became the foundation of DentalShield Ops.")

    add_heading_styled(doc, "The Dual-Market Bridge", level=3)

    add_text(doc,
        "Camilla\u2019s unique value lies in her ability to bridge two distinct regulatory ecosystems. "
        "Brazilian dental regulations (ANVISA RDC 15/2012, RDC 222/2018) mandate some of the "
        "most detailed sterilization and waste management protocols in the Western Hemisphere. "
        "American regulations (OSHA 29 CFR 1910.1030, CDC MMWR 52:RR-17) approach the same "
        "problems through a different regulatory philosophy. By holding clinical credentials in "
        "both systems, Camilla designs compliance frameworks that incorporate the most protective "
        "elements of each\u2014creating a standard that exceeds what either system requires individually.")

    add_text(doc,
        "This cross-regulatory expertise is not transferable through documentation alone. The "
        "ongoing evolution of the 156-point framework\u2014incorporating new OSHA aerosol-generating "
        "procedure requirements, updated CDC infection prevention guidance, and state-specific "
        "sterilization monitoring mandates\u2014requires Camilla\u2019s direct clinical judgment and "
        "regulatory interpretation. The platform\u2019s compliance accuracy depends on her continued "
        "technical leadership.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. COMPETITIVE LANDSCAPE
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "6. Competitive Landscape", level=1)

    add_text(doc,
        "The dental compliance market currently lacks a comprehensive solution that bridges "
        "physical infrastructure and digital management. Existing competitors address fragments "
        "of the problem, leaving critical gaps that expose practices to regulatory risk.")

    add_data_table(doc,
        headers=["Capability", "DentalShield Ops", "SafeLink", "Compliancy Group", "iLearn Dental"],
        rows=[
            ["IoT Hardware Installation", "\u2713", "\u2717", "\u2717", "\u2717"],
            ["OSHA BPS (29 CFR 1910.1030)", "\u2713", "Partial", "\u2717", "Partial"],
            ["CDC Infection Control", "\u2713", "\u2717", "\u2717", "\u2717"],
            ["HIPAA Security Rule", "\u2713", "\u2717", "\u2713", "\u2717"],
            ["Physical Sterilization Setup", "\u2713", "\u2717", "\u2717", "\u2717"],
            ["Real-Time IoT Monitoring", "\u2713", "\u2717", "\u2717", "\u2717"],
            ["OSHA Audit Guarantee", "\u2713", "\u2717", "\u2717", "\u2717"],
            ["Dental-Specific Clinical Founder", "\u2713", "\u2717", "\u2717", "Partial"],
            ["Compliance Breadth", "OSHA+CDC+HIPAA+EPA", "Multi-sector", "HIPAA only", "Training only"],
            ["Physical + Digital", "\u2713", "\u2717", "\u2717", "\u2717"],
        ])

    doc.add_paragraph()

    add_text(doc,
        "SafeLink offers multi-sector compliance advisory but lacks dental specificity, physical "
        "implementation capability, IoT hardware, and audit protection guarantees. Compliancy Group "
        "focuses exclusively on HIPAA\u2014leaving OSHA, CDC, and EPA compliance entirely unaddressed. "
        "iLearn Dental provides online training content, but training without physical compliance "
        "infrastructure leaves practices non-compliant in sterilization, waste management, and PPE "
        "protocols. DentalShield Ops is the only entrant that addresses the full compliance surface "
        "with integrated hardware and software.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. PRICING & REVENUE MODEL
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "7. Pricing & Revenue Model", level=1)

    add_text(doc,
        "DentalShield Ops uses a tiered all-inclusive bundle model. Each tier includes "
        "implementation, hardware, software access, training, maintenance, and the OSHA Audit "
        "Protection Guarantee. The tier structure is informed by behavioral economics research "
        "(Ariely, 2008) to maximize adoption of the Professional tier, which delivers the "
        "highest value-to-cost ratio for mid-sized practices.")

    add_heading_styled(doc, "Subscription Tiers", level=3)

    add_data_table(doc,
        headers=["", "Starter", "Professional", "Elite", "Enterprise"],
        rows=[
            ["Monthly Fee", "$299", "$499", "$799", "Custom"],
            ["Setup Cost", "$1,500", "$3,000", "$5,000", "Custom"],
            ["Target Size", "1\u20132 dentists", "3\u20135 dentists", "5+ dentists", "DSOs 50\u2013500 loc."],
            ["SteriSensors", "1", "3", "5", "Per location"],
            ["BiohazardBoxes", "1", "2", "3", "Per location"],
            ["ComplianceScreens", "1", "2", "3", "Per location"],
            ["SmartSupply", "Basic", "Full", "Full", "Full + custom"],
            ["EquipMax", "Basic", "Full", "Full", "Full + custom"],
            ["TeamSync", "\u2014", "\u2713", "\u2713", "\u2713"],
            ["OSHA Guarantee", "Included", "Up to $10K", "Up to $15K", "Custom"],
            ["On-Site Visits/Year", "2", "4", "6", "Custom"],
            ["Business Review", "\u2014", "Quarterly", "Monthly", "Dedicated AM"],
            ["Training Access", "Basic", "Full (20 hrs)", "Full + Priority", "Custom"],
            ["Support SLA", "Standard", "Standard", "2-hour", "Dedicated"],
        ])

    doc.add_paragraph()

    add_heading_styled(doc, "Unit Economics", level=3)

    add_data_table(doc,
        headers=["Metric", "Value"],
        rows=[
            ["Customer Acquisition Cost (CAC)", "$1,500\u2013$3,000 per clinic"],
            ["Lifetime Value (Professional, 5yr @ 10% churn)", "~$24,000"],
            ["LTV:CAC Ratio", "8:1 \u2192 10:1 target"],
            ["Payback Period", "<12 months"],
            ["Implementation Avg. Ticket", "$14,000/clinic"],
            ["Professional Monthly Subscription", "$499/month ($5,988/year)"],
        ])

    doc.add_paragraph()

    add_text(doc,
        "Revenue is generated through two complementary streams: one-time implementation fees "
        "(assessment, physical setup, documentation, training\u2014averaging $14,000 per clinic) "
        "and recurring monthly subscriptions. The subscription model ensures predictable, "
        "compounding revenue as the installed base grows. Conversion hypotheses by tier: "
        "Starter 15%, Professional 65%, Elite 5%, Enterprise 15% of total revenue.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 8. FINANCIAL PROJECTIONS
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "8. Financial Projections", level=1)

    add_text(doc,
        "Financial projections are based on a phased geographic expansion beginning in Central "
        "Florida (Lakeland, Polk County) and scaling to Miami-Fort Lauderdale and national "
        "markets by Year 3\u20135. All figures derive from the current pricing model, implementation "
        "capacity assumptions, and observed market demand in the target geography.")

    add_heading_styled(doc, "Five-Year Revenue Projections", level=3)

    add_data_table(doc,
        headers=["Metric", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
        rows=[
            ["New Clinic Implementations", "50", "120", "180", "220", "260"],
            ["Active Subscriptions", "40", "150", "320", "530", "780"],
            ["ARR (Annual Recurring)", "$240K", "$900K", "$1.92M", "$3.18M", "$4.68M"],
            ["Implementation Revenue", "$700K", "$1.26M", "$1.92M", "$2.38M", "$3.18M"],
            ["Total Annual Revenue", "$940K", "$2.16M", "$3.84M", "$5.56M", "$7.86M"],
        ],
        highlight_col=5)

    doc.add_paragraph()

    add_heading_styled(doc, "Revenue Growth Trajectory", level=3)

    add_data_table(doc,
        headers=["Indicator", "Value"],
        rows=[
            ["Year 1 Total Revenue", "$940,000"],
            ["Year 5 Total Revenue", "$7,860,000"],
            ["Year 5 Active Subscriptions", "780"],
            ["Year 5 ARR", "$4,680,000"],
            ["5-Year Cumulative Revenue", "$20.36M (projected)"],
            ["ARR as % of Total Revenue by Y5", "~60%"],
        ])

    doc.add_paragraph()

    add_text(doc,
        "The financial model demonstrates a clear path from initial market validation (50 clinics, "
        "Year 1) to scaled operations (780 active subscriptions, Year 5). The growing proportion "
        "of recurring subscription revenue\u2014from 25% of total revenue in Year 1 to 60% by "
        "Year 5\u2014creates an increasingly predictable and valuable revenue base. The shift from "
        "implementation-heavy to subscription-heavy revenue is a key indicator of platform maturity "
        "and enterprise value.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 9. GO-TO-MARKET STRATEGY
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "9. Go-to-Market Strategy", level=1)

    add_text(doc,
        "DentalShield Ops follows a four-phase geographic expansion, beginning in the "
        "highest-need corridor of Central Florida and scaling methodically based on validated "
        "demand and operational capacity.")

    phases = [
        ("Phase 1: VALIDATE", "Months 1\u201312",
         "Deploy to 20\u201330 clinics in HPSA-dense Polk County, where the dentist-to-population "
         "ratio reaches 1:2,747. Target corridors: Lakeland, Davenport, Clermont, Leesburg, "
         "and Apopka. Revenue focus: implementation fees ($14,000 average) plus first "
         "maintenance subscriptions. Goal: 15+ documented compliance success stories as case "
         "studies for subsequent phases."),

        ("Phase 2: INTEGRATE", "Months 6\u201318",
         "Develop API integrations between IoT hardware and the DentOps SaaS platform. Beta "
         "rollout with 5\u201310 existing DentalShield clients. ComplianceScreen evolves from "
         "standalone dashboard into the full platform interface. This phase transforms the "
         "hardware installation business into a recurring-revenue SaaS operation."),

        ("Phase 3: SCALE FLORIDA", "Months 12\u201324",
         "Launch full bundle (hardware + SaaS) for new clients. Execute upsell campaign to "
         "convert existing hardware-only clients to subscription model. Expand to Miami-Fort "
         "Lauderdale MSA\u2014the largest dental HPSA deficit nationally, with 1,259 additional "
         "dentists needed. Target: 150 active subscriptions by end of Phase 3."),

        ("Phase 4: NATIONAL", "Months 18\u201360",
         "Geographic expansion to three additional metro areas selected by HPSA density and "
         "population growth: Houston-The Woodlands, TX (264 active HPSAs, +98,700 "
         "residents/year); Atlanta-Sandy Springs, GA (184 HPSAs, 8th fastest-growing US metro). "
         "Target: 500+ active monitoring subscriptions by Year 4, reaching 780 by Year 5."),
    ]

    for name, timeline, desc in phases:
        add_rich_para(doc, [
            (f"{name} ", True, False, NAVY, 12),
            (f"({timeline})", False, True, GOLD, 10),
        ], space_after=Pt(2))
        add_text(doc, desc, font_size=10.5, space_after=Pt(12))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 10. NATIONAL IMPACT
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "10. National Impact & Regulatory Alignment", level=1)

    add_text(doc,
        "DentalShield Ops directly addresses documented federal priorities related to healthcare "
        "access, workforce development, patient safety, and regulatory compliance technology "
        "adoption.")

    add_heading_styled(doc, "10.1 Addressing the National Dental Shortage", level=2)

    add_text(doc,
        "With 63.7 million Americans living in dental HPSAs and 7,054 designated shortage areas "
        "nationwide, the dental workforce crisis is a documented federal priority. DentalShield "
        "Ops does not train or place new dentists\u2014but it directly amplifies the capacity of "
        "the existing dental workforce by automating the compliance burden that consumes an "
        "estimated 15\u201320% of practice administrative time. By freeing clinical staff to focus "
        "on patient care rather than regulatory paperwork, the platform effectively increases "
        "the productive capacity of every practice it serves.")

    add_heading_styled(doc, "10.2 Reducing OSHA Violation Rates", level=2)

    add_text(doc,
        "OSHA enforcement data shows that sterilization and bloodborne pathogen violations under "
        "29 CFR 1910.1030 remain among the most frequently cited categories in dental inspections. "
        "With penalties reaching $165,514 per willful violation and 40% of small clinics closing "
        "within two years of a major fine, the financial impact cascades beyond the individual "
        "practice to community dental access. DentalShield Ops\u2019 SteriTrack module targets a "
        "70\u201380% reduction in sterilization-related violations through continuous IoT monitoring "
        "and automated audit log generation.")

    add_heading_styled(doc, "10.3 Patient Safety at Scale", level=2)

    add_text(doc,
        "The CDC estimates over 8,000 dental procedure-related infections annually. Dental "
        "Assistants face 10x higher bloodborne pathogen exposure risk than the general population. "
        "DentalShield Ops addresses these risks at the infrastructure level\u2014not through training "
        "materials or software dashboards alone, but through physical IoT sensors that verify "
        "sterilization actually occurred at the correct parameters. This hardware-verified approach "
        "creates an evidence standard that software-only solutions cannot match.")

    add_heading_styled(doc, "10.4 Regulatory Framework Coverage", level=2)

    add_data_table(doc,
        headers=["Regulatory Framework", "DentalShield Ops Coverage"],
        rows=[
            ["OSHA 29 CFR 1910.1030 (Bloodborne Pathogens)", "SteriTrack + BiohazardBox monitoring"],
            ["CDC MMWR 52:RR-17 (Infection Control)", "156-point audit protocol + SteriSensor"],
            ["HIPAA 45 CFR 164 (Security Rule)", "ComplianceScreen documentation automation"],
            ["EPA RCRA (Hazardous Waste)", "BiohazardBox chain-of-custody tracking"],
            ["OSHA 29 CFR 1910.132\u2013138 (PPE)", "ComplianceScreen PPE compliance tracking"],
            ["State-specific sterilization mandates", "Configurable per jurisdiction"],
            ["OSHA 2026 aerosol procedures", "Updated protocol integration"],
        ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 11. TECHNICAL SPECIFICATIONS
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "11. Technical Specifications", level=1)

    add_heading_styled(doc, "11.1 IoT Hardware", level=2)

    add_data_table(doc,
        headers=["Device", "Function", "Data Captured", "Connectivity"],
        rows=[
            ["SteriSensor", "Autoclave monitoring", "Temp, pressure, cycle time", "Wi-Fi/BLE"],
            ["BiohazardBox", "Waste tracking", "Operator ID, timestamp, type, fill", "RFID + Wi-Fi"],
            ["ComplianceScreen", "Dashboard display", "Aggregated compliance status", "Wi-Fi tablet"],
        ])

    doc.add_paragraph()

    add_heading_styled(doc, "11.2 Software Platform", level=2)

    specs = [
        ("Frontend:", "React + TypeScript with responsive design"),
        ("Database:", "PostgreSQL with Row Level Security (RLS) policies"),
        ("API Layer:", "RESTful API with real-time WebSocket for IoT data"),
        ("Supply Chain:", "Henry Schein + Patterson Dental API integration"),
        ("PMS:", "Practice Management System integration for TeamSync"),
        ("Security:", "TLS 1.3 encryption, HIPAA-compliant handling, RBAC"),
        ("Reporting:", "Automated OSHA audit docs, training certs, compliance reports"),
    ]

    for label, desc in specs:
        add_rich_para(doc, [
            (label + " ", True, False, NAVY, 10.5),
            (desc, False, False, DARK_TEXT, 10.5),
        ], space_after=Pt(4))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 12. SOURCES
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, "12. Sources & Regulatory Citations", level=1)

    sources = [
        "Mordor Intelligence (2024). United States Dental Services Market. "
        "$165B (2024) \u2192 $196.5B (2034), CAGR 7.9%.",

        "Health Resources and Services Administration (HRSA). Health Professional Shortage "
        "Areas Dashboard. December 2025. 63.7M people in dental HPSAs; 10,744 additional "
        "dentists needed.",

        "HRSA Quarterly Report, March 2025. 7,054 Dental HPSAs; 59.7M people affected.",

        "Bureau of Labor Statistics (BLS). Occupational Employment Statistics, May 2024. "
        "SOC 29-1292.00: 342,000 Dental Assistants; 52,900 annual openings projected.",

        "BLS Occupational Outlook Handbook, 2024\u20132034. Dental Hygienists: 7% growth, "
        "median wage $94,260. Dentists: 4% growth, median wage $179,210.",

        "OSHA Annual Adjustments to Civil Penalties, January 7, 2025. $16,550/serious; "
        "$165,514/willful violation.",

        "OSHA Enforcement Data, 2024. 40% of small dental clinics receiving major fines "
        "close within 2 years.",

        "CDC. Guidelines for Infection Control in Dental Health-Care Settings, 2003 "
        "(MMWR 52:RR-17).",

        "CDC. Summary of Infection Prevention Practices in Dental Settings: Basic "
        "Expectations for Safe Care. Updated 2024.",

        "Association for Dental Safety (ADS). From Policy to Practice: 2025 Edition.",

        "ADA Workforce Report, 2025. 202,485 active dentists; 59.5 per 100K population.",

        "RuralHealthInfoHub, January 2026. 66% of dental HPSAs in rural areas; 19.4M "
        "affected. Rural: 32.7/100K vs. urban: 64.7/100K.",

        "Pearl AI. Dentist Workforce Statistics 2026. Retirement age: 68.7 years; "
        "DSO affiliation: 16.1%.",

        "PubMed PMID 2370583. Sterilizer failure rates up to 64.7% during biological "
        "indicator testing.",

        "PMC, 2024. Chemical vapor sterilizer failure: 43.8%; steam: 6.5%.",

        "Ariely, D. (2008). Predictably Irrational. HarperCollins.",

        "OSHA 29 CFR 1910.1030 \u2014 Bloodborne Pathogens Standard.",

        "CDC MMWR 52:RR-17 \u2014 Infection Control in Dental Health-Care Settings.",

        "HIPAA 45 CFR 164 \u2014 Security Rule.",

        "EPA Resource Conservation and Recovery Act (RCRA).",

        "OSHA 29 CFR 1910.132\u2013138 \u2014 PPE Standards.",

        "ANVISA RDC 15/2012, RDC 222/2018 \u2014 Brazilian sterilization and waste standards.",
    ]

    for i, source in enumerate(sources, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run_num = p.add_run(f"[{i}] ")
        run_num.font.bold = True
        run_num.font.size = Pt(9)
        run_num.font.color.rgb = NAVY
        run_num.font.name = FONT_MAIN
        run_txt = p.add_run(source)
        run_txt.font.size = Pt(9)
        run_txt.font.color.rgb = MEDIUM_GRAY
        run_txt.font.name = FONT_MAIN

    # ── FOOTER ────────────────────────────────────────────────────

    add_separator(doc)

    add_text(doc,
        "This document is confidential and intended for institutional and investor review "
        "purposes only. Reproduction or distribution without written consent from "
        "DentalShield Systems, LLC is prohibited.",
        font_size=8, color=MEDIUM_GRAY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    add_text(doc,
        "DentalShield Systems, LLC | Lakeland, Florida | dentalshieldops.com",
        font_size=8, bold=True, color=NAVY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ── SAVE ──────────────────────────────────────────────────────

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"SUCCESS: Document saved to {OUTPUT_PATH}")
    print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    generate()
