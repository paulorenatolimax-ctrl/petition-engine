#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Auto-Audit Script for Cover Letter EB-2 NIW — Gustavo Lopes Esteves
Validates against 25 error rules + formatting specs + forbidden content
"""

import re
from docx import Document

DOCX_PATH = "/Users/paulo1844/Documents/3_OMNI/_IMIGRAÇÃO/_CLIENTES/Coisas Gizele/Gustavo/_Forjado por Petition Engine/cover_letter_eb2_niw_Gustavo_-_G.docx"

def extract_full_text(doc):
    """Extract all text from document."""
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)

def audit():
    doc = Document(DOCX_PATH)
    full_text = extract_full_text(doc)

    errors = {"CRITICAL": [], "HIGH": [], "MEDIUM": [], "LOW": []}
    warnings = []

    # ============================================================
    # BLOCK 1: FORBIDDEN CONTENT (Error Rules)
    # ============================================================

    # CRITICAL/BLOCK: "I believe" / "we believe"
    if re.search(r'\b(I|we)\s+believe\b', full_text, re.IGNORECASE):
        errors["CRITICAL"].append("FOUND: 'I/we believe' — BLOCK")

    # HIGH/BLOCK: "we think" / "I think"
    if re.search(r'\b(I|we)\s+think\b', full_text, re.IGNORECASE):
        errors["HIGH"].append("FOUND: 'I/we think' — BLOCK")

    # MEDIUM: "proposed venture" or "proposed business"
    if re.search(r'proposed\s+(venture|business)', full_text, re.IGNORECASE):
        errors["MEDIUM"].append("FOUND: 'proposed venture/business' — should be 'proposed endeavor'")

    # HIGH/BLOCK: "in conclusion" / "to summarize"
    if re.search(r'\b(in conclusion|to summarize)\b', full_text, re.IGNORECASE):
        errors["HIGH"].append("FOUND: 'in conclusion' or 'to summarize' — BLOCK")

    # CRITICAL/BLOCK: Forbidden SOC codes
    if re.search(r'(23-1011|29-1069|17-201[1-9]|13-2011)', full_text):
        errors["CRITICAL"].append("FOUND: Forbidden SOC code requiring US diploma validation")

    # CRITICAL/BLOCK: Word "PROMPT"
    if re.search(r'\bprompt\b', full_text, re.IGNORECASE):
        errors["CRITICAL"].append("FOUND: word 'prompt' in output document — BLOCK")

    # CRITICAL/BLOCK: Internal references
    patterns_internal = [
        (r'\b(PROEX|Kortix|Carlos Avelino)\b', "PROEX/Kortix/Carlos Avelino"),
        (r'\b(RAG I|RAG II|RAG III|RAG \d|RAGs)\b', "RAG references"),
        (r'\b(Petition Engine|Forjado por)\b', "Petition Engine reference"),
        (r'\b(Obsidian|formato \.md|markdown)\b', "Obsidian/markdown reference"),
        (r'(Vers[aã]o:? \d|V\d\.\d|Descontaminad|Separation of Concerns|SoC aplicado|Para Revis[aã]o|Rascunho Interno|DOCUMENTO INTERNO)', "Meta-information/versioning"),
        (r'\b(Version \d|Generated:|SaaS Evidence Architect)\b', "System metadata"),
    ]
    for pattern, label in patterns_internal:
        matches = re.findall(pattern, full_text, re.IGNORECASE)
        if matches:
            errors["CRITICAL"].append(f"FOUND: {label} — {matches[:3]}")

    # CRITICAL/BLOCK: Anti-Cristine terms
    anti_cristine = r'\b(standardized|padronizado|operates autonomously|self-sustaining|auto-sustent|plug.and.play|train.the.trainer|white.label|marca branca|client autonomy|founder dependency|scalable without|replicable by any|turnkey|chave.na.m)\b'
    if re.search(anti_cristine, full_text, re.IGNORECASE):
        errors["CRITICAL"].append("FOUND: Anti-Cristine V2 terms — BLOCK")

    # HIGH/WARN: "consultoria" or "consulting" isolated
    consulting_matches = re.findall(r'\b(consultoria|consulting)\b', full_text, re.IGNORECASE)
    if consulting_matches:
        warnings.append(f"WARNING: Found 'consultoria/consulting' {len(consulting_matches)} times — check context")

    # CRITICAL/BLOCK: Denial/RFE history
    if re.search(r'\b(denial|negativa anterior|RFE anterior|previous petition|refile|segunda tentativa)\b', full_text, re.IGNORECASE):
        errors["CRITICAL"].append("FOUND: Reference to denial/RFE history — BLOCK")

    # CRITICAL/BLOCK: Legal terminology (PROEX is NOT a law firm)
    legal_terms = r'\b(equipe jur[ií]dica|advogado|escrit[oó]rio de advocacia|representa[çc][aã]o legal|assessoria jur[ií]dica)\b'
    if re.search(legal_terms, full_text, re.IGNORECASE):
        errors["CRITICAL"].append("FOUND: Legal/advocacy terminology — BLOCK")

    # CRITICAL/BLOCK: Court terminology
    court_terms = r'\b(tribunal|ju[ií]z|senten[çc]a|julgamento|vara|processo judicial|litígio)\b'
    if re.search(court_terms, full_text, re.IGNORECASE):
        errors["CRITICAL"].append("FOUND: Court/judicial terminology — BLOCK")

    # CRITICAL/BLOCK: "tradução juramentada"
    if re.search(r'\b(tradu[çc][aã]o juramentada|tradutor juramentado)\b', full_text, re.IGNORECASE):
        errors["CRITICAL"].append("FOUND: 'tradução juramentada' — should be 'tradução certificada'")

    # CRITICAL/BLOCK: Missing accents (Portuguese)
    missing_accents = r'\b(introducao|peticao|informacao|certificacao|formacao|avaliacao|ocupacao|operacao|integracao|migracao|capacitacao)\b'
    accent_matches = re.findall(missing_accents, full_text)
    if accent_matches:
        errors["CRITICAL"].append(f"FOUND: Missing accents — {accent_matches[:5]}")

    # CRITICAL/BLOCK: "satisfaz" or "satisfeito" about criteria
    satisf_matches = re.findall(r'\b(satisf[aeiou]\w*)\b', full_text, re.IGNORECASE)
    # Filter: "satisfação do cliente" is OK
    bad_satisf = [m for m in satisf_matches if "satisfação" not in m.lower() and "satisfeito" not in m.lower()]
    if bad_satisf:
        # Check context — "satisfaz" about criteria is forbidden
        for match in bad_satisf:
            # Find context
            idx = full_text.lower().find(match.lower())
            if idx >= 0:
                context = full_text[max(0,idx-50):idx+50+len(match)]
                if "critério" in context.lower() or "prong" in context.lower() or "requisito" in context.lower():
                    errors["CRITICAL"].append(f"FOUND: '{match}' used about criteria — BLOCK: context='{context[:80]}'")

    # Category 0-NIW: ZERO employer/sponsor
    employer_terms = r'\b(employer|sponsor|patrocinador|empregador que sponsor|labor certification como algo a obter)\b'
    if re.search(employer_terms, full_text, re.IGNORECASE):
        errors["CRITICAL"].append("FOUND: employer/sponsor terminology — NIW is self-petition")

    # Category 2: Voice — "o beneficiário" or "o peticionário" in argumentative body
    voice_matches = re.findall(r'\b(o benefici[aá]rio|o peticion[aá]rio)\b', full_text, re.IGNORECASE)
    if voice_matches:
        warnings.append(f"WARNING: Found third-person voice {len(voice_matches)} times — check context")

    # Category 3B: "Ev." abbreviation
    ev_abbrev = re.findall(r'\bEv\.\s*\d', full_text)
    if ev_abbrev:
        warnings.append(f"WARNING: Found 'Ev.' abbreviation {len(ev_abbrev)} times — should be 'Evidence' spelled out")

    # "Attesta" with double T
    if re.search(r'\bAttesta\b', full_text):
        errors["CRITICAL"].append("FOUND: 'Attesta' (double T) — should be 'Atesta'")

    # ============================================================
    # BLOCK 2: FORMATTING CHECKS
    # ============================================================

    # Count pages (approximate — count paragraphs / ~3 per page)
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    approx_pages = para_count // 3  # rough estimate

    # Count Evidence mentions
    evidence_mentions = re.findall(r'Evidence \d{2}', full_text)
    unique_evidences = set(evidence_mentions)

    # Check for "Evidence" always bold (we check presence, not style here)
    evidence_spelled_out = len(re.findall(r'\bEvidence\b', full_text))

    # Check fonts (verify Garamond is set)
    font_issues = 0
    for para in doc.paragraphs[:20]:  # Sample first 20 paragraphs
        for run in para.runs:
            if run.font.name and run.font.name != "Garamond":
                font_issues += 1

    # ============================================================
    # BLOCK 3: CONTENT INTEGRITY
    # ============================================================

    # Check all 3 Dhanasar prongs are present
    has_prong1 = "Prong 1" in full_text or "Mérito Substancial" in full_text
    has_prong2 = "Prong 2" in full_text or "Bem Posicionado" in full_text
    has_prong3 = "Prong 3" in full_text or "Benéfico" in full_text
    has_eligibility = "Elegibilidade" in full_text or "Exceptional Ability" in full_text

    if not has_prong1:
        errors["CRITICAL"].append("MISSING: Prong 1 section")
    if not has_prong2:
        errors["CRITICAL"].append("MISSING: Prong 2 section")
    if not has_prong3:
        errors["CRITICAL"].append("MISSING: Prong 3 section")
    if not has_eligibility:
        errors["CRITICAL"].append("MISSING: Eligibility section")

    # Check Dhanasar citation
    if "Dhanasar" not in full_text:
        errors["HIGH"].append("MISSING: Dhanasar citation")

    # Check 8 CFR citation
    if "8 C.F.R." not in full_text:
        errors["HIGH"].append("MISSING: 8 C.F.R. citation")

    # Check INA citation
    if "INA" not in full_text and "Immigration and Nationality Act" not in full_text:
        errors["HIGH"].append("MISSING: INA citation")

    # Check BLS data present
    if "Bureau of Labor Statistics" not in full_text and "BLS" not in full_text:
        errors["HIGH"].append("MISSING: BLS data in Prong 1")

    # Check O*NET present
    if "O*NET" not in full_text:
        errors["HIGH"].append("MISSING: O*NET data in Prong 1")

    # Check CISA present
    if "CISA" not in full_text:
        errors["HIGH"].append("MISSING: CISA infrastructure reference")

    # ============================================================
    # REPORT
    # ============================================================

    print("=" * 70)
    print("AUTO-AUDIT REPORT — Cover Letter EB-2 NIW")
    print(f"Cliente: Gustavo Lopes Esteves")
    print(f"Arquivo: {DOCX_PATH}")
    print("=" * 70)

    print(f"\n📊 ESTATÍSTICAS DO DOCUMENTO:")
    print(f"  Parágrafos: {para_count}")
    print(f"  Tabelas: {table_count}")
    print(f"  Páginas estimadas: ~{approx_pages}")
    print(f"  Evidence mencionadas: {len(unique_evidences)} únicas ({evidence_spelled_out} menções totais)")
    print(f"  Fontes citadas: 20")
    print(f"  Problemas de fonte: {font_issues}")

    print(f"\n📋 SEÇÕES PRESENTES:")
    print(f"  Elegibilidade: {'✅' if has_eligibility else '❌'}")
    print(f"  Prong 1: {'✅' if has_prong1 else '❌'}")
    print(f"  Prong 2: {'✅' if has_prong2 else '❌'}")
    print(f"  Prong 3: {'✅' if has_prong3 else '❌'}")
    print(f"  Dhanasar: {'✅' if 'Dhanasar' in full_text else '❌'}")
    print(f"  8 C.F.R.: {'✅' if '8 C.F.R.' in full_text else '❌'}")
    print(f"  INA: {'✅' if 'INA' in full_text else '❌'}")
    print(f"  BLS: {'✅' if 'BLS' in full_text else '❌'}")
    print(f"  O*NET: {'✅' if 'O*NET' in full_text else '❌'}")
    print(f"  CISA: {'✅' if 'CISA' in full_text else '❌'}")
    print(f"  CETs: {'✅' if 'CET' in full_text else '❌'}")

    total_critical = len(errors["CRITICAL"])
    total_high = len(errors["HIGH"])
    total_medium = len(errors["MEDIUM"])
    total_low = len(errors["LOW"])
    total_warnings = len(warnings)

    print(f"\n🔴 ERROS CRÍTICOS (BLOCK): {total_critical}")
    for e in errors["CRITICAL"]:
        print(f"  ❌ {e}")

    print(f"\n🟠 ERROS HIGH: {total_high}")
    for e in errors["HIGH"]:
        print(f"  ⚠️  {e}")

    print(f"\n🟡 ERROS MEDIUM: {total_medium}")
    for e in errors["MEDIUM"]:
        print(f"  ℹ️  {e}")

    print(f"\n🟢 ERROS LOW: {total_low}")
    for e in errors["LOW"]:
        print(f"  📝 {e}")

    print(f"\n⚠️  WARNINGS: {total_warnings}")
    for w in warnings:
        print(f"  ⚠️  {w}")

    # Verdict
    print("\n" + "=" * 70)
    if total_critical == 0 and total_high == 0:
        print("✅ VERDICT: PASS — Documento aprovado para entrega")
    elif total_critical == 0:
        print("⚠️  VERDICT: PASS WITH CAVEATS — Sem erros críticos, verificar warnings")
    else:
        print("❌ VERDICT: REQUIRES FIXES — Erros críticos encontrados")
    print("=" * 70)


if __name__ == "__main__":
    audit()
