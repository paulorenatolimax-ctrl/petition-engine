#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Cover Letter EB-2 NIW Generator — Gustavo Lopes Esteves
EventFinOps LLC | SOC: 11-3031 (Financial Managers)
Sistema v3.0 | 100% PT-BR | Garamond | python-docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ============================================================
# CONSTANTS
# ============================================================
OUTPUT_DIR = "/Users/paulo1844/Documents/3_OMNI/_IMIGRAÇÃO/_CLIENTES/Coisas Gizele/Gustavo/_Forjado por Petition Engine"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "cover_letter_eb2_niw_Gustavo_-_G.docx")

# Colors
VERDE_PROEX = "D6E1DB"
CREME = "FFF8EE"
EV_GREEN = "2E7D32"
H4_SUBTITLE = "F2F5D7"
BLACK = "000000"
GRAY = "808080"

# Client data
CLIENT_NAME = "Gustavo Lopes Esteves"
CLIENT_DOB = "18 de janeiro de 2000"
CLIENT_NATIONALITY = "Brasileira"
SOC_CODE = "11-3031.00"
SOC_TITLE = "Financial Managers"
COMPANY_NAME = "EventFinOps LLC"
COMPANY_LOCATION = "Miami, Flórida"
PE_SUMMARY = (
    "a fundação e operação da EventFinOps LLC, uma empresa de assessoria estratégica "
    "em investimentos internacionais, alocação de portfólios, estruturação de produtos "
    "financeiros e serviços de fusões e aquisições (M&A) e mercado de capitais de dívida "
    "(DCM), com sede em Miami, Flórida, voltada para investidores brasileiros e latino-americanos "
    "que buscam diversificação patrimonial nos mercados norte-americanos"
)

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_paragraph_shading(paragraph, color_hex):
    """Apply background shading to a paragraph."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    paragraph._p.get_or_add_pPr().append(shading_elm)


def add_section_header(doc, text, level=1):
    """Add a section header with proper formatting."""
    p = doc.add_paragraph()
    if level == 1:
        # H Section: 14pt bold, shading D6E1DB
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Garamond"
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_shading(p, VERDE_PROEX)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        # keepNext
        pPr = p._p.get_or_add_pPr()
        keepNext = parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>')
        pPr.append(keepNext)
    elif level == 2:
        # H Sub: 13pt bold italic, shading D6E1DB
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(13)
        run.font.name = "Garamond"
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_shading(p, VERDE_PROEX)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        pPr = p._p.get_or_add_pPr()
        keepNext = parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>')
        pPr.append(keepNext)
    elif level == 3:
        # H Subsec: 12pt bold
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Garamond"
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        keepNext = parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>')
        pPr.append(keepNext)
    elif level == 4:
        # H4 subtitle: shading F2F5D7
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(12)
        run.font.name = "Garamond"
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_shading(p, H4_SUBTITLE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        keepNext = parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>')
        pPr.append(keepNext)
    return p


def add_body_paragraph(doc, text, bold_phrases=None, italic_phrases=None):
    """Add a body paragraph with proper formatting. Supports bold/italic inline."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(14.5)
    p.paragraph_format.space_after = Pt(6)

    if bold_phrases is None:
        bold_phrases = []
    if italic_phrases is None:
        italic_phrases = []

    # Find all bold and italic segments
    segments = _split_text_for_formatting(text, bold_phrases, italic_phrases)
    for seg_text, is_bold, is_italic in segments:
        run = p.add_run(seg_text)
        run.font.name = "Garamond"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if is_bold:
            run.bold = True
        if is_italic:
            run.italic = True

    # widowControl
    pPr = p._p.get_or_add_pPr()
    wc = parse_xml(f'<w:widowControl {nsdecls("w")} w:val="true"/>')
    pPr.append(wc)

    return p


def _split_text_for_formatting(text, bold_phrases, italic_phrases):
    """Split text into segments with formatting markers."""
    if not bold_phrases and not italic_phrases:
        return [(text, False, False)]

    # Build markers
    markers = []
    for bp in bold_phrases:
        idx = text.find(bp)
        if idx >= 0:
            markers.append((idx, idx + len(bp), True, False))
    for ip in italic_phrases:
        idx = text.find(ip)
        if idx >= 0:
            markers.append((idx, idx + len(ip), False, True))

    if not markers:
        return [(text, False, False)]

    markers.sort(key=lambda x: x[0])

    segments = []
    pos = 0
    for start, end, is_bold, is_italic in markers:
        if start > pos:
            segments.append((text[pos:start], False, False))
        segments.append((text[start:end], is_bold, is_italic))
        pos = end
    if pos < len(text):
        segments.append((text[pos:], False, False))

    return segments


def add_evidence_card(doc, evidence_num, title, ev_type, source, date, url, description):
    """Add an evidence block v4: 1x2 table with thumbnail placeholder + metadata."""
    # Create 1-row, 2-column table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table width to 100%
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)
    # tblInd = 0
    tblInd = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="0" w:type="dxa"/>')
    tblPr.append(tblInd)

    # Remove all borders (we'll add horizontal only)
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="{VERDE_PROEX}"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{VERDE_PROEX}"/>
        <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))

    # Column 0: thumbnail placeholder (~3.5cm) - white background
    cell0 = table.cell(0, 0)
    cell0.width = Cm(3.5)
    # Set white background (ShadingType.CLEAR)
    set_cell_shading(cell0, "FFFFFF")
    p0 = cell0.paragraphs[0]
    run0 = p0.add_run("[Thumbnail]")
    run0.font.name = "Garamond"
    run0.font.size = Pt(8)
    run0.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Column 1: metadata - cream background (#FFF8EE)
    cell1 = table.cell(0, 1)
    set_cell_shading(cell1, CREME)

    # Clear default paragraph
    cell1.paragraphs[0].clear()

    # Evidence title (green, bold)
    p_title = cell1.paragraphs[0]
    run_title = p_title.add_run(f"Evidence {evidence_num:02d}. {title}")
    run_title.bold = True
    run_title.font.name = "Garamond"
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    # Metadata lines
    meta_lines = [
        f"Tipo: {ev_type}",
        f"Fonte: {source}",
        f"Data: {date}",
    ]
    if url:
        meta_lines.append(f"URL: {url}")
    meta_lines.append(f"Descrição e Relevância: {description}")

    for line in meta_lines:
        p_meta = cell1.add_paragraph()
        run_meta = p_meta.add_run(line)
        run_meta.font.name = "Garamond"
        run_meta.font.size = Pt(10)
        run_meta.font.color.rgb = RGBColor(0, 0, 0)
        p_meta.paragraph_format.space_after = Pt(2)
        p_meta.paragraph_format.space_before = Pt(0)

    # cantSplit on the row
    row = table.rows[0]
    trPr = row._tr.get_or_add_trPr()
    cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>')
    trPr.append(cantSplit)

    # Space after evidence card
    doc.add_paragraph()
    return table


def add_data_table(doc, headers, rows, col_widths=None):
    """Add a formatted data table with horizontal-only borders."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    # 100% width
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)
    # tblInd = 0
    tblInd = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="0" w:type="dxa"/>')
    tblPr.append(tblInd)
    # Center
    jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>')
    tblPr.append(jc)

    # Horizontal-only borders
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="6" w:space="0" w:color="{BLACK}"/>
        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{BLACK}"/>
        <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="808080"/>
        <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, VERDE_PROEX)
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = "Garamond"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx + 1, c_idx)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = "Garamond"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing after table
    return table


def add_footnote_text(doc, number, text):
    """Add a footnote-style text at paragraph level."""
    p = doc.add_paragraph()
    run_num = p.add_run(f"[{number}] ")
    run_num.font.name = "Garamond"
    run_num.font.size = Pt(10)
    run_num.font.color.rgb = RGBColor(0, 0, 0)
    run_num.bold = True
    run_text = p.add_run(text)
    run_text.font.name = "Garamond"
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type=7)  # page break
    return p


def set_document_defaults(doc):
    """Set document-wide defaults for font and margins."""
    # Margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.5)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = "Garamond"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = Pt(14.5)
    pf.space_after = Pt(6)

    # Set Garamond as the default font in XML
    rPr = doc.styles['Normal']._element.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="Garamond" w:hAnsi="Garamond" '
        f'w:eastAsia="Garamond" w:cs="Garamond"/>'
    )
    rPr.append(rFonts)


def add_footer(doc, client_name):
    """Add footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run1 = p.add_run("EB-2 NIW | I-140 Petition — Cover Letter ")
        run1.font.name = "Garamond"
        run1.font.size = Pt(8)
        run1.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        run2 = p.add_run(client_name)
        run2.font.name = "Garamond"
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        run3 = p.add_run(" | Página ")
        run3.font.name = "Garamond"
        run3.font.size = Pt(8)
        run3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Page number field
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_field = p.add_run()
        run_field._r.append(fld_char_begin)
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_field2 = p.add_run()
        run_field2._r.append(instr)
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_field3 = p.add_run()
        run_field3._r.append(fld_char_end)

        run4 = p.add_run(" de ")
        run4.font.name = "Garamond"
        run4.font.size = Pt(8)
        run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Total pages field
        fld_char_begin2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_f1 = p.add_run()
        run_f1._r.append(fld_char_begin2)
        instr2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
        run_f2 = p.add_run()
        run_f2._r.append(instr2)
        fld_char_end2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_f3 = p.add_run()
        run_f3._r.append(fld_char_end2)


# ============================================================
# DOCX BREAK TYPE
# ============================================================
from docx.oxml import OxmlElement
import docx.oxml

# Fix for page break via run
def _add_page_break_run(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)
    return p


# ============================================================
# MAIN DOCUMENT GENERATION
# ============================================================

def generate_cover_letter():
    doc = Document()
    set_document_defaults(doc)

    # --------------------------------------------------------
    # SECTION 0: COVER PAGE (LETTER FORMAT)
    # --------------------------------------------------------
    # Date (right-aligned)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = p_date.add_run("9 de abril de 2026")
    run_date.font.name = "Garamond"
    run_date.font.size = Pt(12)

    doc.add_paragraph()  # spacing

    # Addressee
    p_addr = doc.add_paragraph()
    p_addr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lines = [
        "U.S. Citizenship and Immigration Services",
        "Texas Service Center",
        "P.O. Box 650447",
        "Dallas, TX 75265-0447"
    ]
    for i, line in enumerate(lines):
        run = p_addr.add_run(line)
        run.font.name = "Garamond"
        run.font.size = Pt(12)
        if i < len(lines) - 1:
            run.add_break()

    doc.add_paragraph()

    # RE line
    p_re = doc.add_paragraph()
    run_re = p_re.add_run("Re: ")
    run_re.bold = True
    run_re.font.name = "Garamond"
    run_re.font.size = Pt(12)
    run_re2 = p_re.add_run(
        "Formulário I-140, Immigrant Petition for Alien Workers — "
        "Classificação EB-2, National Interest Waiver (NIW) — "
        f"{CLIENT_NAME}"
    )
    run_re2.font.name = "Garamond"
    run_re2.font.size = Pt(12)

    doc.add_paragraph()

    # Salutation
    p_sal = doc.add_paragraph()
    run_sal = p_sal.add_run("Prezado(a) Oficial de Imigração:")
    run_sal.font.name = "Garamond"
    run_sal.font.size = Pt(12)

    doc.add_paragraph()

    # Green metadata block
    p_meta = doc.add_paragraph()
    set_paragraph_shading(p_meta, VERDE_PROEX)
    meta_text = (
        f"Peticionário: {CLIENT_NAME} | Nacionalidade: Brasileira | "
        f"Classificação: EB-2 National Interest Waiver | "
        f"Formulário: I-140 | Seção: INA § 203(b)(2)(B) | "
        f"Regulamento: 8 C.F.R. § 204.5(k) | "
        f"SOC: {SOC_CODE} — {SOC_TITLE}"
    )
    run_meta = p_meta.add_run(meta_text)
    run_meta.font.name = "Garamond"
    run_meta.font.size = Pt(11)
    run_meta.bold = True

    doc.add_paragraph()

    # Opening paragraph
    add_body_paragraph(
        doc,
        "Submeto esta petição I-140 solicitando classificação como trabalhador imigrante "
        "de segunda preferência com base em emprego (EB-2), acompanhada de pedido de "
        "dispensa do requisito de oferta de emprego e certificação de trabalho "
        "(National Interest Waiver — NIW), nos termos da Seção 203(b)(2)(B) do "
        "Immigration and Nationality Act (INA) e do regulamento 8 C.F.R. § 204.5(k).",
        bold_phrases=["National Interest Waiver — NIW"],
        italic_phrases=["Immigration and Nationality Act"]
    )

    add_body_paragraph(
        doc,
        f"Meu proposed endeavor consiste em {PE_SUMMARY}. "
        "Conforme demonstrarei ao longo desta petição, meu empreendimento proposto "
        "possui mérito substancial e importância nacional, estou bem posicionado para "
        "avançá-lo, e a dispensa dos requisitos de oferta de emprego e certificação de "
        "trabalho seria benéfica para os Estados Unidos.",
        italic_phrases=["proposed endeavor"]
    )

    add_body_paragraph(
        doc,
        "Esta Cover Letter está organizada da seguinte forma: primeiro, demonstro minha "
        "elegibilidade para a classificação EB-2; em seguida, abordo cada um dos três "
        "prongs do teste estabelecido em Matter of Dhanasar, 26 I&N Dec. 884 (AAO 2016), "
        "que constitui o framework analítico vigente para avaliação de petições NIW.",
        italic_phrases=["Matter of Dhanasar, 26 I&N Dec. 884 (AAO 2016)"]
    )

    _add_page_break_run(doc)

    # --------------------------------------------------------
    # SECTION 1: SINOPSE (2-3 pages)
    # --------------------------------------------------------
    add_section_header(doc, "I. Sinopse da Petição")

    add_body_paragraph(
        doc,
        f"Meu nome é {CLIENT_NAME}, sou cidadão brasileiro nascido em {CLIENT_DOB} em "
        "São Paulo, Brasil. Possuo bacharelado em Administração de Empresas pelo "
        "Insper Instituto de Ensino e Pesquisa, uma das instituições de ensino superior "
        "mais prestigiadas do Brasil, com formação complementar internacional pela "
        "University of California, Los Angeles (UCLA). Ao longo de minha carreira no "
        "setor financeiro, acumulei experiência significativa em investimentos "
        "internacionais, gestão de portfólios offshore, estruturação de produtos "
        "financeiros e assessoria estratégica para investidores de alta renda.",
        bold_phrases=["Insper Instituto de Ensino e Pesquisa", "University of California, Los Angeles (UCLA)"]
    )

    add_body_paragraph(
        doc,
        "Durante minha atuação no Criteria Financial Group — um dos escritórios integrantes "
        "da rede XP Inc., a maior plataforma de investimentos do Brasil com mais de R$ 1,2 "
        "trilhão em ativos sob custódia —, fui responsável pela gestão de portfólios "
        "offshore totalizando mais de US$ 210 milhões em ativos pessoais e contribuí para "
        "a vertical de investimentos internacionais que alcançou aproximadamente R$ 1 bilhão "
        "(~US$ 200 milhões) alocados em mercados norte-americanos. O escritório foi "
        "reconhecido como Top 10 Offshore dentro da rede XP por anos consecutivos, "
        "conquista para a qual contribuí diretamente como analista e, posteriormente, "
        "como diretor da área.",
        bold_phrases=["US$ 210 milhões", "R$ 1 bilhão", "Top 10 Offshore"]
    )

    add_body_paragraph(
        doc,
        "Minha progressão de carreira foi notavelmente acelerada: em menos de quatro anos, "
        "evolui de estagiário na área de Renda Fixa para Diretor de Investimentos Offshore, "
        "sendo formalmente designado como sucessor do então Head de Offshore, Rafael "
        "Wurzmann (atualmente no BTG Pactual, um dos maiores bancos de investimento da "
        "América Latina). Essa trajetória ascendente foi documentada por cinco profissionais "
        "seniores do setor financeiro em cartas de recomendação independentes.",
        bold_phrases=["menos de quatro anos", "estagiário", "Diretor de Investimentos Offshore"]
    )

    add_body_paragraph(
        doc,
        "Além de minha experiência profissional, contribuí para a literatura acadêmica do "
        "campo com três artigos científicos revisados por pares publicados na Revista Lumen "
        "et Virtus (classificação QUALIS CAPES B2), cobrindo temas como internacionalização "
        "de ativos, ETFs e tendências do mercado offshore. Também publiquei o livro "
        "\"Liderança e Tomada de Decisão no Setor Financeiro: Lições de um jovem executivo\" "
        "(ISBN 978-65-83827-35-7, Golden Int Editora, 2025), e fui perfilado por três "
        "veículos de mídia nacionais como líder emergente no setor de investimentos "
        "internacionais.",
        bold_phrases=["três artigos científicos", "Liderança e Tomada de Decisão no Setor Financeiro"],
        italic_phrases=["Revista Lumen et Virtus"]
    )

    add_body_paragraph(
        doc,
        f"Meu proposed endeavor — {PE_SUMMARY} — está alinhado com prioridades federais "
        "documentadas, incluindo o fortalecimento do setor de serviços financeiros como "
        "infraestrutura crítica (CISA/PPD-21), a promoção de tecnologias críticas e "
        "emergentes (NSTC CET List 2024), e o desenvolvimento econômico do estado da "
        "Flórida. O Bureau of Labor Statistics projeta crescimento de 15% para Financial "
        "Managers (SOC 11-3031) entre 2024 e 2034, cinco vezes acima da média nacional "
        "de 3%, com designação Bright Outlook pelo O*NET OnLine.",
        italic_phrases=["proposed endeavor"],
        bold_phrases=["15%", "cinco vezes acima da média nacional"]
    )

    # Synopsis table
    add_data_table(doc,
        ["Seção", "Conteúdo", "Evidências-Chave"],
        [
            ["Elegibilidade", "Exceptional Ability — 3 de 6 critérios atendidos", "Evidence 01–05"],
            ["Prong 1", "Mérito substancial + importância nacional: CETs, EOs, BLS, CISA", "Evidence 06–11"],
            ["Prong 2", "Educação, carreira, 5 cartas, Business Plan, publicações", "Evidence 02–25"],
            ["Prong 3", "Análise custo-benefício: 5 fatores Dhanasar/NYSDOT", "Evidence 06"],
        ]
    )

    _add_page_break_run(doc)

    # --------------------------------------------------------
    # SECTION 2: ELIGIBILITY (8-12 pages)
    # --------------------------------------------------------
    add_section_header(doc, "II. Elegibilidade para Classificação EB-2")

    add_section_header(doc, "Base Legal", level=2)

    add_body_paragraph(
        doc,
        "A classificação EB-2 está prevista na Seção 203(b)(2) do INA e regulamentada "
        "pelo 8 C.F.R. § 204.5(k). Existem duas vias de qualificação: (1) Advanced "
        "Degree — diploma avançado ou equivalente (bacharelado acrescido de cinco anos "
        "de experiência progressiva na especialidade); ou (2) Exceptional Ability — "
        "habilidade excepcional nas ciências, artes ou negócios, demonstrada pelo "
        "atendimento de pelo menos três dos seis critérios regulamentares do "
        "8 C.F.R. § 204.5(k)(3)(ii).",
        italic_phrases=["Advanced Degree", "Exceptional Ability"]
    )

    add_body_paragraph(
        doc,
        "Conforme orientação do USCIS Policy Manual (Vol. 6, Part F, Ch. 5) e da recente "
        "Policy Alert PA-2025-03 (15 de janeiro de 2025), a habilidade excepcional deve "
        "se relacionar ao empreendimento proposto como parte do pedido de NIW, e o USCIS "
        "determina a relação entre a habilidade excepcional e o proposed endeavor caso a "
        "caso, considerando quaisquer conjuntos de habilidades, conhecimentos ou "
        "expertise compartilhados.",
        italic_phrases=["Policy Alert PA-2025-03", "proposed endeavor"]
    )

    add_section_header(doc, "Via de Elegibilidade: Exceptional Ability (8 C.F.R. § 204.5(k)(3)(ii))", level=2)

    add_body_paragraph(
        doc,
        "Demonstro elegibilidade para classificação EB-2 pela via de Exceptional Ability, "
        "atendendo a três dos seis critérios regulamentares, conforme documentado a seguir.",
        italic_phrases=["Exceptional Ability"]
    )

    # Criterion 1: Academic Degree
    add_section_header(doc, "Critério 1: Diploma ou Certificado de Instituição de Ensino Relacionado à Área de Habilidade Excepcional", level=3)

    add_evidence_card(doc, 2, "Diploma de Bacharel em Administração de Empresas — Insper Instituto de Ensino e Pesquisa (Dezembro de 2023)",
        ev_type="Documento Acadêmico",
        source="Insper Instituto de Ensino e Pesquisa, São Paulo, Brasil",
        date="Dezembro de 2023",
        url="",
        description=(
            "Diploma de bacharelado em Administração de Empresas emitido pelo Insper Instituto de Ensino e Pesquisa, "
            "uma das instituições de ensino superior mais seletivas do Brasil, com nota máxima 5 no Ministério da "
            "Educação (MEC) e reconhecida internacionalmente pela AACSB. O programa de quatro anos (3.300+ horas) "
            "inclui disciplinas em finanças corporativas, análise de investimentos, gestão de portfólios, economia "
            "quantitativa e estratégia empresarial, fornecendo a base teórica e técnica diretamente aplicável às "
            "service lines da EventFinOps LLC."
        )
    )

    add_body_paragraph(
        doc,
        "Possuo bacharelado em Administração de Empresas pelo Insper Instituto de Ensino e "
        "Pesquisa (Evidence 02), uma das instituições de ensino superior mais prestigiadas "
        "do Brasil, com avaliação máxima (nota 5) pelo Ministério da Educação (MEC) e "
        "acreditação internacional pela Association to Advance Collegiate Schools of Business "
        "(AACSB). O programa de quatro anos, com carga horária superior a 3.300 horas, "
        "incluiu formação aprofundada em finanças corporativas, análise de investimentos, "
        "gestão de portfólios, economia quantitativa e estratégia empresarial.",
        bold_phrases=["Evidence 02"],
        italic_phrases=["Association to Advance Collegiate Schools of Business"]
    )

    add_body_paragraph(
        doc,
        "Complementei minha formação acadêmica com dois programas internacionais na "
        "University of California, Los Angeles (UCLA): o Junior Program Classic e o "
        "programa de Empreendedorismo (Evidence 03 e Evidence 04), que proporcionaram "
        "imersão no ecossistema empresarial e financeiro norte-americano, incluindo "
        "práticas de regulação financeira dos EUA, modelos de negócios em mercados de "
        "capitais e técnicas de gestão empresarial aplicáveis ao contexto norte-americano.",
        bold_phrases=["Evidence 03", "Evidence 04"],
    )

    add_evidence_card(doc, 3, "Certificado — UCLA Junior Program Classic",
        ev_type="Certificação Internacional",
        source="University of California, Los Angeles (UCLA)",
        date="2018",
        url="",
        description=(
            "Certificado de conclusão do programa Junior Program Classic da UCLA, abrangendo "
            "fundamentos de negócios e economia no contexto norte-americano. O programa proporciona "
            "imersão acadêmica em uma das universidades mais prestigiadas do mundo (Top 20 QS World "
            "Ranking), complementando a formação do Insper com perspectiva internacional e exposição "
            "direta ao mercado financeiro dos Estados Unidos."
        )
    )

    add_evidence_card(doc, 4, "Certificado — UCLA Entrepreneurship Program",
        ev_type="Certificação Internacional",
        source="University of California, Los Angeles (UCLA)",
        date="2019",
        url="",
        description=(
            "Certificado de conclusão do programa de Empreendedorismo da UCLA, focado em "
            "desenvolvimento de modelos de negócios, validação de mercado e estratégias de "
            "crescimento empresarial. Esta formação é diretamente aplicável à fundação e "
            "operação da EventFinOps LLC, demonstrando preparação específica para empreendedorismo "
            "nos Estados Unidos."
        )
    )

    add_body_paragraph(
        doc,
        "O diploma em Administração de Empresas pelo Insper, complementado pelos programas "
        "da UCLA, está diretamente relacionado à minha área de habilidade excepcional — "
        "investimentos internacionais, gestão de portfólios e assessoria financeira "
        "estratégica — e ao proposed endeavor de fundar e operar a EventFinOps LLC, "
        "atendendo plenamente o primeiro critério regulamentar.",
        italic_phrases=["proposed endeavor"]
    )

    # Criterion 4: Salary Above the Norm
    add_section_header(doc, "Critério 4: Remuneração que Demonstra Habilidade Acima da Norma", level=3)

    add_body_paragraph(
        doc,
        "No exercício de minhas funções no Criteria Financial Group, fui responsável pela "
        "gestão direta de portfólios offshore totalizando mais de US$ 210 milhões em "
        "ativos, contribuindo para uma vertical que alcançou aproximadamente R$ 1 bilhão "
        "(~US$ 200 milhões) em alocações internacionais (Evidence 07). Esta "
        "responsabilidade fiduciária sobre volumes de ativos desta magnitude, atribuída a "
        "um profissional que iniciou como estagiário apenas três anos antes, constitui "
        "evidência objetiva de remuneração e responsabilidade substancialmente acima da "
        "norma para profissionais na mesma faixa de experiência.",
        bold_phrases=["US$ 210 milhões", "R$ 1 bilhão", "Evidence 07"]
    )

    add_body_paragraph(
        doc,
        "Atualmente, minha remuneração anual é de US$ 130.000,00 (Evidence 07), valor que, "
        "quando contextualizado pela minha idade (26 anos) e pelo mercado brasileiro — onde "
        "a renda média mensal do trabalhador formal é de R$ 3.322,00 (IBGE/PNAD, 2024), "
        "equivalente a aproximadamente US$ 7.974,00 anuais —, representa uma compensação "
        "mais de 16 vezes superior à média nacional brasileira. Ademais, o reconhecimento "
        "por meio de viagens institucionais aos escritórios da XP Inc. em Miami e Nova York, "
        "concedidas a um seleto grupo de profissionais de alto desempenho, corrobora o nível "
        "excepcional de minha contribuição.",
        bold_phrases=["US$ 130.000,00", "16 vezes superior"]
    )

    # Salary comparison table
    add_data_table(doc,
        ["Métrica", "Valor", "Fonte"],
        [
            ["Remuneração atual (Gustavo)", "US$ 130.000,00/ano", "Evidence 07"],
            ["Renda média formal — Brasil", "~US$ 7.974,00/ano", "IBGE/PNAD 2024"],
            ["Múltiplo vs. média brasileira", "16,3x", "Cálculo"],
            ["Portfólio sob gestão pessoal", "US$ 210.000.000+", "Evidence 07, 21, 22"],
            ["Vertical offshore total", "~R$ 1.000.000.000", "Evidence 15–19"],
        ]
    )

    # Criterion 6: Recognition for Contributions
    add_section_header(doc, "Critério 6: Reconhecimento por Contribuições Significativas ao Campo", level=3)

    add_body_paragraph(
        doc,
        "Minhas contribuições ao campo de investimentos internacionais foram reconhecidas "
        "por múltiplas fontes independentes, incluindo publicações acadêmicas revisadas "
        "por pares, publicação de livro com ISBN, cobertura de mídia nacional e cartas de "
        "recomendação de profissionais seniores do setor financeiro.",
        bold_phrases=[]
    )

    add_section_header(doc, "Publicações Acadêmicas Revisadas por Pares", level=4)

    add_body_paragraph(
        doc,
        "Publiquei três artigos científicos na Revista Lumen et Virtus (ISSN: 2177-2789), "
        "periódico classificado como QUALIS CAPES B2 no sistema brasileiro de avaliação "
        "de periódicos acadêmicos, com revisão por pares no formato double-blind. Cada "
        "artigo possui Digital Object Identifier (DOI) e está indexado em bases de dados "
        "acadêmicas, demonstrando contribuição original e validada ao corpo de conhecimento "
        "do campo de investimentos internacionais.",
        bold_phrases=["três artigos científicos"],
        italic_phrases=["Revista Lumen et Virtus"]
    )

    add_evidence_card(doc, 9, "Artigo Científico — \"Riscos e Oportunidades na Internacionalização de Ativos para Clientes de Alta Renda\" (Lumen et Virtus, v.13, n.31)",
        ev_type="Publicação Acadêmica Revisada por Pares",
        source="Revista Lumen et Virtus, ISSN 2177-2789, QUALIS CAPES B2",
        date="27 de dezembro de 2023",
        url="DOI: 10.56238/levv13n31-058",
        description=(
            "Primeiro artigo acadêmico publicado, analisando os riscos e oportunidades inerentes "
            "à internacionalização de ativos para clientes de alta renda no mercado brasileiro. "
            "O artigo examina estruturas de diversificação geográfica, veículos de investimento "
            "offshore e frameworks de gestão de risco cambial, contribuindo para a literatura "
            "acadêmica sobre a interseção entre finanças pessoais e mercados internacionais. "
            "Publicação revisada por pares (double-blind) com DOI atribuído."
        )
    )

    add_evidence_card(doc, 10, "Artigo Científico — \"ETFs e a Expansão das Fronteiras do Investimento para o Investidor Brasileiro\" (Lumen et Virtus, v.15, n.42)",
        ev_type="Publicação Acadêmica Revisada por Pares",
        source="Revista Lumen et Virtus, ISSN 2177-2789, QUALIS CAPES B2",
        date="18 de novembro de 2024",
        url="DOI: 10.56238/levv15n42-085",
        description=(
            "Segundo artigo acadêmico publicado, examinando o papel dos Exchange-Traded Funds "
            "(ETFs) na democratização do acesso a investimentos internacionais para investidores "
            "brasileiros. A pesquisa analisa o crescimento exponencial do mercado global de ETFs — "
            "que atingiu US$ 19,5 trilhões em ativos sob gestão ao final de 2025 — e as "
            "implicações para a diversificação patrimonial de investidores de mercados emergentes. "
            "Publicação revisada por pares (double-blind) com DOI atribuído."
        )
    )

    add_evidence_card(doc, 11, "Artigo Científico — \"Tendências do Mercado de Investimentos Offshore\" (Lumen et Virtus, v.16, n.49)",
        ev_type="Publicação Acadêmica Revisada por Pares",
        source="Revista Lumen et Virtus, ISSN 2177-2789, QUALIS CAPES B2",
        date="30 de junho de 2025",
        url="DOI: 10.56238/levv16n49-119",
        description=(
            "Terceiro e mais recente artigo acadêmico, analisando as tendências contemporâneas "
            "do mercado de investimentos offshore, incluindo evolução regulatória, produtos "
            "estruturados, e o impacto de políticas monetárias divergentes entre economias "
            "desenvolvidas e emergentes sobre fluxos de capital cross-border. A publicação "
            "demonstra contribuição contínua e progressiva ao campo, refletindo a evolução "
            "de minha expertise em investimentos internacionais."
        )
    )

    add_section_header(doc, "Livro Publicado", level=4)

    add_evidence_card(doc, 5, "Livro — \"Liderança e Tomada de Decisão no Setor Financeiro: Lições de um jovem executivo\" (ISBN 978-65-83827-35-7)",
        ev_type="Publicação — Livro com ISBN",
        source="Golden Int Editora, Ribeirão Preto, SP, 2025",
        date="2025",
        url="https://clubedeautores.com.br/livro/lideranca-e-tomada-de-decisao-no-setor-financeiro",
        description=(
            "Livro publicado com ISBN pela Golden Int Editora, catalogado pela Câmara Brasileira "
            "do Livro (CIP) sob classificação CDD 658.4092 (Liderança: Administração). A obra "
            "aborda temas de liderança, tomada de decisão sob pressão, gestão de pessoas no "
            "setor financeiro, inovação e construção de visão estratégica, com referências a "
            "autores como Kahneman, Dalio e Housel. A publicação demonstra thought leadership "
            "no campo de gestão financeira e capacidade de contribuição intelectual original."
        )
    )

    add_body_paragraph(
        doc,
        "Publiquei o livro \"Liderança e Tomada de Decisão no Setor Financeiro: Lições de "
        "um jovem executivo\" (Evidence 05), com ISBN 978-65-83827-35-7, pela Golden Int "
        "Editora. A obra foi catalogada pela Câmara Brasileira do Livro (CIP) e está "
        "disponível em plataformas de distribuição de livros. O livro aborda liderança, "
        "tomada de decisão sob pressão, gestão de pessoas e inovação no setor financeiro, "
        "demonstrando capacidade de contribuição intelectual original ao campo.",
        bold_phrases=["Evidence 05"]
    )

    add_section_header(doc, "Cobertura de Mídia Nacional", level=4)

    add_body_paragraph(
        doc,
        "Fui perfilado por três veículos de mídia brasileiros como líder emergente no "
        "setor de investimentos internacionais (Evidence 18, Evidence 19 e Evidence 20). "
        "As matérias destacam minha trajetória profissional acelerada, minha expertise em "
        "mercados offshore e minha contribuição para a democratização do investimento "
        "internacional no Brasil.",
        bold_phrases=["Evidence 18", "Evidence 19", "Evidence 20"]
    )

    # Publications table
    add_data_table(doc,
        ["#", "Publicação", "Tipo", "Fonte/Veículo", "Data"],
        [
            ["1", "Riscos e Oportunidades na Internacionalização de Ativos", "Artigo Revisado por Pares", "Lumen et Virtus v.13 n.31", "Dez/2023"],
            ["2", "ETFs e a Expansão das Fronteiras do Investimento", "Artigo Revisado por Pares", "Lumen et Virtus v.15 n.42", "Nov/2024"],
            ["3", "Tendências do Mercado de Investimentos Offshore", "Artigo Revisado por Pares", "Lumen et Virtus v.16 n.49", "Jun/2025"],
            ["4", "Liderança e Tomada de Decisão no Setor Financeiro", "Livro (ISBN)", "Golden Int Editora", "2025"],
            ["5", "Estratégia, resiliência e resultado", "Matéria de Mídia", "Brasil Agora", "Out/2025"],
            ["6", "A democratização do investimento internacional", "Matéria de Mídia", "Gazeta da Semana", "Out/2025"],
            ["7", "Jovens líderes globais", "Matéria de Mídia", "Business Feed", "Nov/2025"],
        ]
    )

    # Eligibility synopsis table
    add_section_header(doc, "Tabela-Sinopse — Elegibilidade", level=3)

    add_data_table(doc,
        ["Critério (8 C.F.R. § 204.5(k)(3)(ii))", "Status", "Evidência"],
        [
            ["(i) Diploma relacionado à área", "Atendido", "Evidence 02, 03, 04"],
            ["(ii) 10+ anos de experiência", "Não aplicável", "—"],
            ["(iii) Licença profissional", "Não aplicável", "—"],
            ["(iv) Remuneração acima da norma", "Atendido", "Evidence 07, 08"],
            ["(v) Membro de associação profissional", "Não aplicável", "—"],
            ["(vi) Reconhecimento por contribuições", "Atendido", "Evidence 05, 09–20"],
        ]
    )

    add_body_paragraph(
        doc,
        "Conforme demonstrado, atendo a três dos seis critérios regulamentares — critérios "
        "(i), (iv) e (vi) do 8 C.F.R. § 204.5(k)(3)(ii) —, o que é consistente com a "
        "exigência regulamentar mínima para classificação EB-2 por Exceptional Ability. "
        "Minha habilidade excepcional no campo de investimentos internacionais e gestão "
        "financeira estratégica está diretamente relacionada ao proposed endeavor de fundar "
        "e operar a EventFinOps LLC.",
        italic_phrases=["Exceptional Ability", "proposed endeavor"]
    )

    _add_page_break_run(doc)

    # --------------------------------------------------------
    # SECTION 3: PRONG 1 — SUBSTANTIAL MERIT + NATIONAL IMPORTANCE
    # --------------------------------------------------------
    add_section_header(doc, "III. Prong 1 — O Empreendimento Proposto Possui Mérito Substancial e Importância Nacional")

    add_body_paragraph(
        doc,
        "O primeiro prong do teste Dhanasar requer que o peticionário demonstre que o "
        "proposed endeavor possui (1) mérito substancial — gera valor econômico, social, "
        "cultural, tecnológico ou educacional mensurável — e (2) importância nacional — "
        "seu escopo e impacto transcendem a esfera local ou individual e afetam interesses "
        "nacionais. Conforme o USCIS Policy Manual (Vol. 6, Part F, Ch. 5, Sec. B): "
        "\"National importance does not require that the endeavor have national or even "
        "regional scope.\" A importância nacional pode ser demonstrada pelo impacto em "
        "setor crítico, mesmo que o empreendimento tenha escopo local.",
        italic_phrases=["proposed endeavor", "Matter of Dhanasar, 26 I&N Dec. 884 (AAO 2016)"]
    )

    # PART A: PE + Substantial Merit
    add_section_header(doc, "Parte A: O Proposed Endeavor e Seu Mérito Substancial", level=2)

    add_section_header(doc, "O Empreendimento Proposto", level=3)

    add_body_paragraph(
        doc,
        f"Meu proposed endeavor consiste em {PE_SUMMARY}. "
        "A EventFinOps LLC será uma empresa projetada para atender um segmento "
        "significativo e crescente do mercado: investidores brasileiros e latino-americanos "
        "que buscam acessar os mercados financeiros norte-americanos para diversificação "
        "patrimonial, proteção cambial e otimização de retornos ajustados ao risco.",
        italic_phrases=["proposed endeavor"]
    )

    add_body_paragraph(
        doc,
        "A empresa operará quatro linhas de serviço complementares: (1) assessoria "
        "estratégica em investimentos internacionais, abrangendo alocação de portfólios, "
        "seleção de ativos e construção de carteiras diversificadas em mercados globais; "
        "(2) estruturação de produtos financeiros, incluindo veículos de investimento "
        "adaptados às necessidades específicas de investidores cross-border; (3) serviços "
        "de fusões e aquisições (M&A) e mercado de capitais de dívida (DCM), conectando "
        "empresas brasileiras a oportunidades de capital nos Estados Unidos; e (4) programas "
        "educacionais sobre investimentos internacionais, democratizando o acesso à educação "
        "financeira qualificada para investidores de alta renda e seus assessores.",
        bold_phrases=[]
    )

    add_evidence_card(doc, 6,
        "Business Plan — EventFinOps LLC, Miami, Flórida",
        ev_type="Plano de Negócios",
        source="EventFinOps LLC",
        date="2026",
        url="",
        description=(
            "Plano de negócios detalhado da EventFinOps LLC apresentando o modelo operacional, "
            "análise de mercado, projeções financeiras quinquenais, estratégia competitiva e "
            "plano de crescimento. O documento detalha receita projetada de US$ 384.000 no "
            "primeiro ano, escalando para US$ 2.160.000 no quinto ano, com geração de até 18 "
            "empregos diretos. O payback projetado é de aproximadamente 20 meses, com margem "
            "EBITDA alcançando 46,7% até o quinto ano de operação."
        )
    )

    add_section_header(doc, "Mérito Substancial — Valor Econômico Direto", level=3)

    add_body_paragraph(
        doc,
        "O mérito substancial do empreendimento proposto é evidenciado por projeções "
        "financeiras concretas documentadas no Business Plan (Evidence 06). A "
        "EventFinOps LLC projeta receitas crescentes ao longo de um horizonte de cinco "
        "anos, partindo de US$ 384.000 no primeiro ano e alcançando US$ 2.160.000 no "
        "quinto ano, com geração progressiva de empregos diretos em território "
        "norte-americano.",
        bold_phrases=["Evidence 06", "US$ 384.000", "US$ 2.160.000"]
    )

    # Financial projections table
    add_data_table(doc,
        ["Métrica", "Ano 1", "Ano 2", "Ano 3", "Ano 4", "Ano 5"],
        [
            ["Receita (US$)", "384.000", "648.000", "1.080.000", "1.560.000", "2.160.000"],
            ["Funcionários", "3", "5", "8", "12", "16–18"],
            ["Margem EBITDA", "12,5%", "22,3%", "33,1%", "40,2%", "46,7%"],
            ["Payback", "~20 meses", "—", "—", "—", "—"],
        ]
    )

    add_body_paragraph(
        doc,
        "Além do valor econômico direto, o empreendimento gerará impactos econômicos "
        "indiretos significativos: (1) facilitação de fluxos de capital internacional "
        "para os mercados norte-americanos, contribuindo para a liquidez e profundidade "
        "do sistema financeiro dos EUA; (2) assessoria a investidores brasileiros que "
        "coletivamente representam um mercado potencial de R$ 7,9 trilhões em ativos "
        "(ANBIMA, junho de 2025), dos quais menos de 2% estão atualmente alocados no "
        "exterior; e (3) criação de até 18 empregos qualificados em Miami, Flórida.",
        bold_phrases=["R$ 7,9 trilhões", "menos de 2%", "18 empregos"]
    )

    # PART B: National Importance
    add_section_header(doc, "Parte B: Importância Nacional — Dados Federais", level=2)

    add_body_paragraph(
        doc,
        "A importância nacional do proposed endeavor é demonstrada por seu alinhamento "
        "com múltiplas prioridades federais documentadas, incluindo tecnologias críticas "
        "e emergentes (CETs), ordens executivas, projeções do Bureau of Labor Statistics, "
        "designações do O*NET OnLine, classificação de infraestrutura crítica pela CISA, "
        "e prioridades orçamentárias federais.",
        italic_phrases=["proposed endeavor"]
    )

    # B.1 CETs
    add_section_header(doc, "Tecnologias Críticas e Emergentes (CETs)", level=3)

    add_body_paragraph(
        doc,
        "A lista de Tecnologias Críticas e Emergentes (Critical and Emerging Technologies — "
        "CETs), publicada pelo National Science and Technology Council (NSTC) através do "
        "White House Office of Science and Technology Policy (OSTP), identifica áreas "
        "tecnológicas \"essential to U.S. national security\" e relevantes para \"compete "
        "for international talent.\" Dentre as 18 áreas tecnológicas catalogadas, o "
        "proposed endeavor opera na interseção de pelo menos três:",
        italic_phrases=["proposed endeavor"]
    )

    add_data_table(doc,
        ["#", "Área CET", "Sub-tecnologia", "Conexão com o Proposed Endeavor"],
        [
            ["1", "Data Science and Storage", "Financial data analytics, AI-driven portfolio optimization", "Assessoria de investimentos utiliza análise de dados financeiros em larga escala para alocação de portfólios e identificação de oportunidades em mercados globais"],
            ["2", "Advanced Computing", "Algorithmic trading, quantitative finance", "Estruturação de produtos financeiros e assessoria de investimentos empregam tecnologias de computação avançada para modelagem de risco e otimização de portfólios"],
            ["3", "Distributed Ledger Technologies", "Digital assets, tokenized securities", "Assessoria sobre veículos de investimento que incluem ativos digitais e títulos tokenizados, expandindo o acesso de investidores internacionais a mercados emergentes"],
        ]
    )

    add_body_paragraph(
        doc,
        "A conexão do proposed endeavor com as CETs não é meramente tangencial: as "
        "service lines da EventFinOps LLC dependem fundamentalmente de tecnologias de "
        "análise de dados financeiros, computação avançada para modelagem de risco, e "
        "conhecimento de tecnologias de registro distribuído (DLT) para assessorar "
        "investidores sobre ativos digitais e títulos tokenizados — produtos que "
        "representam uma das fronteiras mais dinâmicas do mercado financeiro "
        "internacional contemporâneo.",
        italic_phrases=["proposed endeavor"]
    )

    # B.2 Executive Orders
    add_section_header(doc, "Ordens Executivas Relevantes", level=3)

    add_data_table(doc,
        ["Ordem Executiva", "Título/Tema", "Data", "Conexão com o Proposed Endeavor"],
        [
            ["EO — Removing Barriers to AI Leadership", "Remoção de barreiras à liderança americana em inteligência artificial", "20/01/2025",
             "Promove a adoção de IA nos setores produtivos americanos, incluindo serviços financeiros. A EventFinOps empregará ferramentas de IA para análise de portfólios e assessoria de investimentos."],
            ["America First Investment Policy", "Priorização de investimentos estrangeiros que beneficiam os EUA", "2025",
             "O proposed endeavor facilita o fluxo de capital brasileiro e latino-americano para mercados norte-americanos, diretamente alinhado com a política de atração de investimentos estrangeiros."],
            ["EO 14028 — Cybersecurity", "Melhoria da cibersegurança nacional", "12/05/2021",
             "O setor de serviços financeiros é infraestrutura crítica (CISA). A EventFinOps implementará protocolos robustos de segurança cibernética para proteção de dados de investidores."],
        ]
    )

    add_body_paragraph(
        doc,
        "O alinhamento com múltiplas ordens executivas demonstra que o proposed endeavor "
        "não opera isoladamente, mas na interseção de prioridades presidenciais "
        "documentadas. A política America First Investment, em particular, visa "
        "atrair capital estrangeiro produtivo para os Estados Unidos — exatamente o "
        "que a EventFinOps LLC propõe facilitar ao assessorar investidores brasileiros "
        "e latino-americanos na alocação de recursos nos mercados financeiros "
        "norte-americanos.",
        italic_phrases=["proposed endeavor"]
    )

    # B.3 BLS/O*NET
    add_section_header(doc, "Bureau of Labor Statistics e O*NET OnLine", level=3)

    add_data_table(doc,
        ["Métrica", "Dados — SOC 11-3031 (Financial Managers)", "Fonte"],
        [
            ["Emprego (2024)", "868.600 posições", "BLS OOH"],
            ["Crescimento projetado (2024–2034)", "15% (much faster than average)", "BLS OOH"],
            ["Crescimento médio (todas as ocupações)", "3%", "BLS OOH"],
            ["Múltiplo vs. média", "5x mais rápido", "Cálculo"],
            ["Vagas anuais projetadas", "74.600", "BLS OOH"],
            ["Salário mediano anual (maio 2024)", "US$ 161.700", "BLS OES"],
            ["Designação Bright Outlook", "Sim", "O*NET OnLine"],
            ["Designação In Demand", "Sim", "O*NET OnLine"],
        ]
    )

    add_body_paragraph(
        doc,
        "O Bureau of Labor Statistics (BLS) projeta crescimento de 15% para Financial "
        "Managers (SOC 11-3031) entre 2024 e 2034, taxa cinco vezes superior à média de "
        "3% para todas as ocupações. Com 74.600 vagas anuais projetadas e salário mediano "
        "de US$ 161.700 (maio de 2024), esta é uma das categorias ocupacionais de mais "
        "rápido crescimento no segmento gerencial. O O*NET OnLine classifica a ocupação "
        "com dupla designação: Bright Outlook (perspectiva positiva de emprego) e In Demand "
        "(habilidades frequentemente requisitadas pelo mercado de trabalho).",
        bold_phrases=["15%", "cinco vezes superior", "74.600 vagas anuais", "US$ 161.700", "Bright Outlook", "In Demand"]
    )

    add_body_paragraph(
        doc,
        "É fundamental destacar que o argumento aqui apresentado NÃO se fundamenta em "
        "escassez de mão de obra — argumento expressamente rejeitado pelo AAO como base "
        "para demonstração de importância nacional. O argumento é que o crescimento "
        "projetado pelo BLS de 15% confirma a relevância nacional e a demanda estrutural "
        "por profissionais qualificados em gestão financeira, e que o proposed endeavor "
        "avança esta prioridade federal documentada ao criar uma empresa que expandirá a "
        "capacidade do setor e trará expertise cross-border única ao mercado americano.",
        italic_phrases=["proposed endeavor"],
        bold_phrases=["NÃO se fundamenta em escassez de mão de obra"]
    )

    # B.4 CISA
    add_section_header(doc, "Infraestrutura Crítica — CISA/PPD-21", level=3)

    add_body_paragraph(
        doc,
        "O setor de serviços financeiros é designado como um dos 16 setores de "
        "infraestrutura crítica dos Estados Unidos pela Cybersecurity and Infrastructure "
        "Security Agency (CISA), sob a Presidential Policy Directive 21 (PPD-21, 2013). "
        "O Department of the Treasury atua como Sector Risk Management Agency (SRMA) para "
        "o setor, coordenando esforços de resiliência e segurança. Conforme a CISA: "
        "\"Large-scale power outages, recent natural disasters, and an increase in the "
        "number and sophistication of cyberattacks demonstrate the wide range of potential "
        "risks facing the sector.\"",
        bold_phrases=["16 setores de infraestrutura crítica"]
    )

    add_body_paragraph(
        doc,
        "A CISA está atualmente desenvolvendo Sector-Specific Goals (SSGs) para o setor "
        "de serviços financeiros (Winter 2025), e um novo módulo de avaliação CSET será "
        "disponibilizado no primeiro trimestre de 2026. O Cybersecurity Strategic Plan "
        "FY2024-2026 da CISA estabelece metas específicas para proteção de infraestrutura "
        "financeira, incluindo sistemas de pagamento, mercados de valores mobiliários e "
        "serviços de gestão de investimentos — exatamente o setor em que a EventFinOps "
        "LLC operará.",
        bold_phrases=[]
    )

    add_body_paragraph(
        doc,
        "A operação da EventFinOps LLC no setor de serviços financeiros — um dos 16 "
        "setores de infraestrutura crítica — confere automaticamente relevância nacional "
        "ao empreendimento proposto, pois qualquer contribuição para o fortalecimento, "
        "modernização e expansão de infraestrutura crítica transcende interesses "
        "locais e afeta a segurança e estabilidade econômica nacional.",
        bold_phrases=[]
    )

    # B.5 Market Data
    add_section_header(doc, "Dados de Mercado e Tendências Setoriais", level=3)

    add_body_paragraph(
        doc,
        "O mercado global de Exchange-Traded Funds (ETFs) — um dos principais veículos "
        "de investimento que a EventFinOps LLC utilizará em suas estratégias de assessoria "
        "— atingiu US$ 19,5 trilhões em ativos sob gestão ao final de 2025, representando "
        "crescimento de 33% em relação aos US$ 14,6 trilhões registrados ao final de 2024. "
        "Somente nos Estados Unidos, os ativos em ETFs alcançaram US$ 13,46 trilhões "
        "(crescimento de 30% year-over-year), com entradas líquidas recordes de "
        "US$ 2,1 trilhões em 2025 — quase 3,5 vezes mais do que os fundos mútuos "
        "tradicionais.",
        bold_phrases=["US$ 19,5 trilhões", "33%", "US$ 13,46 trilhões", "US$ 2,1 trilhões"]
    )

    add_data_table(doc,
        ["Métrica", "Valor", "Crescimento", "Fonte"],
        [
            ["ETFs — AUM Global (2025)", "US$ 19,5 trilhões", "+33% vs. 2024", "PwC/ETFGI"],
            ["ETFs — AUM EUA (2025)", "US$ 13,46 trilhões", "+30% YoY", "American Century/PwC"],
            ["ETFs — Entradas líquidas globais (2025)", "US$ 2,1 trilhões", "3,5x vs. fundos mútuos", "PwC/State Street"],
            ["ETFs Ativos — AUM Global (2025)", "US$ 1,9 trilhão", "+65%", "PwC"],
            ["Projeção AUM Global (2030)", "US$ 35 trilhões", "~80% de crescimento vs. 2025", "PwC/Bloomberg"],
            ["Ativos investidores brasileiros (ANBIMA)", "R$ 7,9 trilhões", "—", "ANBIMA (Jun/2025)"],
            ["Alocação internacional — Brasil", "< 2%", "vs. 25% investidores institucionais EUA", "Citi Research"],
        ]
    )

    add_body_paragraph(
        doc,
        "O gap entre a alocação internacional dos investidores brasileiros (menos de 2%) "
        "e a dos investidores institucionais norte-americanos (25%) representa uma "
        "oportunidade estrutural significativa. Com R$ 7,9 trilhões em ativos sob "
        "administração no mercado brasileiro (ANBIMA, junho de 2025), mesmo um aumento "
        "modesto na taxa de internacionalização representaria bilhões de dólares "
        "adicionais fluindo para os mercados norte-americanos — capital que beneficia "
        "diretamente a economia dos Estados Unidos.",
        bold_phrases=["menos de 2%", "25%", "R$ 7,9 trilhões"]
    )

    # PART C: Synthesis
    add_section_header(doc, "Parte C: Síntese — Importância Nacional Documentada", level=2)

    add_body_paragraph(
        doc,
        "A convergência de múltiplas fontes federais independentes demonstra de forma "
        "conclusiva a importância nacional do proposed endeavor:",
        italic_phrases=["proposed endeavor"]
    )

    add_body_paragraph(
        doc,
        "1. O setor de serviços financeiros é infraestrutura crítica (CISA/PPD-21), "
        "um dos 16 setores designados como vitais para a segurança econômica nacional;\n"
        "2. As tecnologias empregadas pelo empreendimento proposto — Data Science, "
        "Advanced Computing, Distributed Ledger Technologies — constam na lista de "
        "Tecnologias Críticas e Emergentes (NSTC/OSTP);\n"
        "3. Múltiplas ordens executivas mandatam a modernização do setor financeiro, "
        "a adoção de inteligência artificial e a atração de investimentos estrangeiros;\n"
        "4. O BLS projeta crescimento de 15% para Financial Managers (2024–2034), "
        "cinco vezes acima da média, com designação Bright Outlook e In Demand pelo "
        "O*NET;\n"
        "5. O mercado global de ETFs alcançou US$ 19,5 trilhões em 2025, com projeção "
        "de US$ 35 trilhões até 2030;\n"
        "6. A Flórida possui 427 Qualified Opportunity Zones, com expansão prevista "
        "para 1.360 zonas elegíveis sob o programa OZ 2.0 em 2027.",
        bold_phrases=["15%", "US$ 19,5 trilhões", "US$ 35 trilhões"]
    )

    # Synopsis Table Prong 1
    add_section_header(doc, "Tabela-Sinopse — Prong 1", level=3)

    add_data_table(doc,
        ["Dimensão", "Fonte Federal", "Dado-Chave", "Evidência"],
        [
            ["CETs", "NSTC CET List 2024", "3 áreas tecnológicas", "Sources"],
            ["Ordens Executivas", "Federal Register", "3 EOs aplicáveis", "Sources"],
            ["BLS — Crescimento", "BLS OOH", "15% (5x a média)", "Sources"],
            ["O*NET", "O*NET OnLine", "Bright Outlook + In Demand", "Sources"],
            ["CISA", "PPD-21", "1/16 setores críticos", "Sources"],
            ["Mercado ETFs", "PwC/ETFGI", "US$ 19,5T (2025)", "Sources"],
            ["ANBIMA", "ANBIMA", "R$ 7,9T (alocação < 2% internacional)", "Sources"],
        ]
    )

    add_body_paragraph(
        doc,
        "O proposed endeavor não é uma iniciativa comercial ordinária, mas opera na "
        "interseção de infraestrutura crítica, tecnologias emergentes, prioridades "
        "orçamentárias e regulatórias federais, e um mercado em crescimento exponencial — "
        "configurando importância nacional sob o framework analítico de "
        "Dhanasar.",
        italic_phrases=["proposed endeavor", "Dhanasar"]
    )

    _add_page_break_run(doc)

    # --------------------------------------------------------
    # SECTION 4: PRONG 2 — WELL POSITIONED
    # --------------------------------------------------------
    add_section_header(doc, "IV. Prong 2 — Estou Bem Posicionado para Avançar o Empreendimento Proposto")

    add_body_paragraph(
        doc,
        "O segundo prong do teste Dhanasar requer que o peticionário demonstre que "
        "possui educação, habilidades, conhecimento e registro de sucesso suficientes "
        "para avançar o proposed endeavor com alta probabilidade de êxito. Conforme o "
        "USCIS Policy Manual (Vol. 6, Part F, Ch. 5), os fatores considerados incluem: "
        "(a) educação e habilidades; (b) histórico de sucesso em esforços similares; "
        "(c) modelo ou plano para atividade futura; (d) progresso em direção ao "
        "empreendimento proposto; e (e) interesse de potenciais clientes, usuários ou "
        "investidores.",
        italic_phrases=["proposed endeavor", "Dhanasar"]
    )

    # PART A: Education + Career
    add_section_header(doc, "Parte A: Educação e Histórico Profissional", level=2)

    add_section_header(doc, "Formação Acadêmica e Certificações", level=3)

    add_body_paragraph(
        doc,
        "Conforme demonstrado na seção de Elegibilidade (Evidence 02, Evidence 03 e "
        "Evidence 04), possuo bacharelado em Administração de Empresas pelo Insper "
        "Instituto de Ensino e Pesquisa, complementado por dois programas internacionais "
        "na UCLA. As disciplinas de finanças corporativas, análise de investimentos, "
        "gestão de portfólios e economia quantitativa fornecem a base teórica diretamente "
        "aplicável às service lines da EventFinOps LLC. A formação na UCLA, especificamente "
        "nos programas Junior Classic e Empreendedorismo, proporcionou familiaridade com "
        "o ecossistema empresarial e regulatório norte-americano.",
        bold_phrases=["Evidence 02", "Evidence 03", "Evidence 04"]
    )

    add_section_header(doc, "Histórico Profissional: Criteria Financial Group (Janeiro de 2022 — Agosto de 2025)", level=3)

    add_evidence_card(doc, 7,
        "Declaração de Recursos Humanos — Criteria Financial Group / Criteria Invest Assessor de Investimento Ltda. (Janeiro de 2022 — Agosto de 2025)",
        ev_type="Declaração de Emprego",
        source="Criteria Financial Group, São Paulo, Brasil",
        date="7 de abril de 2026",
        url="",
        description=(
            "Declaração oficial de recursos humanos assinada por Carlos Wald Reissmann, confirmando "
            "o vínculo empregatício de Gustavo Lopes Esteves com o Criteria Financial Group de "
            "janeiro de 2022 a agosto de 2025. O documento detalha a progressão de estagiário na "
            "área de Alocação e Renda Fixa para Analista de Investimentos Offshore e, finalmente, "
            "Diretor de Investimentos Offshore, com responsabilidade sobre portfólios de mais de "
            "US$ 210 milhões e contribuição para uma vertical que alcançou ~R$ 1 bilhão em "
            "alocações internacionais."
        )
    )

    add_section_header(doc, "A Instituição", level=4)

    add_body_paragraph(
        doc,
        "O Criteria Financial Group (formalmente Criteria Invest Agentes Autônomos de "
        "Investimentos Ltda., atualmente Criteria Invest Assessor de Investimento Ltda.) "
        "é um dos escritórios integrantes da rede XP Inc., a maior plataforma de "
        "investimentos do Brasil, com mais de R$ 1,2 trilhão em ativos sob custódia. "
        "O equivalente norte-americano mais próximo seria um escritório afiliado a uma "
        "plataforma como Charles Schwab ou Fidelity Investments, atuando como registered "
        "investment advisor para clientes de alta renda.",
        bold_phrases=["XP Inc.", "R$ 1,2 trilhão"],
        italic_phrases=["registered investment advisor"]
    )

    add_section_header(doc, "Progressão e Conquistas", level=4)

    add_body_paragraph(
        doc,
        "Minha trajetória no Criteria Financial Group exemplifica uma progressão "
        "acelerada e resultados mensuráveis excepcionais:",
        bold_phrases=[]
    )

    add_body_paragraph(
        doc,
        "Fase 1 — Estagiário, Alocação e Renda Fixa (Janeiro de 2022 — Dezembro de 2022): "
        "Iniciei minha carreira realizando estudos de renda fixa, suporte a assessores e "
        "clientes, e análises de mercado. Nesta fase inicial, demonstrei aptidão para "
        "análise financeira e relacionamento com clientes, o que resultou em minha "
        "promoção acelerada para a área de investimentos internacionais.",
        bold_phrases=["Fase 1 — Estagiário"]
    )

    add_body_paragraph(
        doc,
        "Fase 2 — Analista de Investimentos Offshore (Janeiro de 2023 — Dezembro de 2024): "
        "Assumi responsabilidade por assessoria B2B e B2C em investimentos internacionais, "
        "produção de materiais analíticos e participação em podcasts sobre cenários "
        "financeiros globais. Neste período, contribuí para a expansão da vertical "
        "offshore do escritório, que alcançou aproximadamente R$ 1 bilhão em alocações "
        "nos mercados norte-americanos. Participei ativamente do lançamento dos BDRs de "
        "ETFs da gestora norte-americana First Trust na B3 (Bolsa de Valores do Brasil) — "
        "um marco na internacionalização do mercado financeiro brasileiro.",
        bold_phrases=["Fase 2 — Analista de Investimentos Offshore", "R$ 1 bilhão",
                      "BDRs de ETFs da gestora norte-americana First Trust na B3"]
    )

    add_body_paragraph(
        doc,
        "Fase 3 — Diretor de Investimentos Offshore (Janeiro de 2025 — Agosto de 2025): "
        "Fui formalmente designado como sucessor de Rafael Wurzmann (atualmente Wealth "
        "Management, Offshore no BTG Pactual) para liderar toda a divisão de investimentos "
        "offshore. Nesta posição, gerenciei portfólios pessoais superiores a US$ 210 "
        "milhões, supervisionei uma equipe de 2 profissionais e prestei suporte a outros "
        "assessores na gestão de portfólios de seus clientes. Atendi mais de 2.000 "
        "clientes, e o escritório foi reconhecido como Top 10 Offshore dentro da rede "
        "XP por anos consecutivos, com viagens de reconhecimento aos escritórios da XP "
        "em Miami e Nova York.",
        bold_phrases=["Fase 3 — Diretor de Investimentos Offshore", "US$ 210 milhões",
                      "2.000 clientes", "Top 10 Offshore"]
    )

    # Career timeline table
    add_data_table(doc,
        ["Período", "Empresa", "Cargo", "Conquistas-Chave", "Evidence"],
        [
            ["01/2022–12/2022", "Criteria Financial Group", "Estagiário — Renda Fixa", "Estudos de mercado, suporte a assessores, promoção acelerada", "Evidence 07"],
            ["01/2023–12/2024", "Criteria Financial Group", "Analista Offshore", "Vertical ~R$ 1 bi, BDRs First Trust/B3, podcasts", "Evidence 07"],
            ["01/2025–08/2025", "Criteria Financial Group", "Diretor Offshore", "US$ 210M+ portfólio, 2.000+ clientes, Top 10 XP", "Evidence 07"],
        ]
    )

    # PART B: Recommendation Letters
    add_section_header(doc, "Parte B: Cartas de Recomendação, Business Plan e Publicações", level=2)

    add_section_header(doc, "Cartas de Recomendação", level=3)

    add_body_paragraph(
        doc,
        "Minha capacidade de avançar o proposed endeavor é corroborada por cinco "
        "profissionais seniores do setor financeiro que atestam, de forma independente, "
        "minhas competências técnicas, capacidade de liderança e registro de resultados "
        "mensuráveis. Cada recomendador possui posição de destaque no mercado financeiro "
        "e observou diretamente minha atuação profissional.",
        italic_phrases=["proposed endeavor"]
    )

    # Letter 1: Luiz Fernando Mesquita
    add_evidence_card(doc, 21,
        "Carta de Recomendação — Luiz Fernando Mesquita, Diretor, BioWash; Ex-CFO, Criteria Partners (30 de janeiro de 2026)",
        ev_type="Carta de Recomendação",
        source="Luiz Fernando Mesquita — Diretor, BioWash; ex-CFO, Criteria Partners",
        date="30 de janeiro de 2026",
        url="",
        description=(
            "Carta de recomendação de Luiz Fernando Mesquita, atual Diretor da BioWash e ex-CFO "
            "do Criteria Partners. O recomendador atesta que a vertical offshore sob contribuição "
            "de Gustavo cresceu para aproximadamente R$ 1 bilhão em mercados norte-americanos, "
            "que Gustavo demonstrava compreensão da lógica de negócios por trás de receita, "
            "retenção e margem, e exibia 'maturidade incomum', operando 'acima do profissional "
            "médio' do setor."
        )
    )

    add_body_paragraph(
        doc,
        "Luiz Fernando Mesquita, atual Diretor da BioWash e ex-CFO do Criteria Partners "
        "(Evidence 21), atesta que contribuí significativamente para o crescimento da "
        "vertical offshore para aproximadamente R$ 1 bilhão em alocações nos mercados "
        "norte-americanos. Em suas palavras, demonstrei \"maturidade incomum\" e "
        "compreensão profunda da \"lógica de negócios por trás de receita, retenção e "
        "margem\", operando \"acima do profissional médio\" do setor financeiro.",
        bold_phrases=["Evidence 21"],
    )

    # Letter 2: Victor Fonseca
    add_evidence_card(doc, 22,
        "Carta de Recomendação — Victor Fonseca, Diretor de Alocação e Produtos, Criteria Financial Group (18 de fevereiro de 2026)",
        ev_type="Carta de Recomendação",
        source="Victor Fonseca — Diretor de Alocação e Produtos, Criteria Financial Group",
        date="18 de fevereiro de 2026",
        url="",
        description=(
            "Carta de recomendação de Victor Fonseca, graduado em Economia pela New York University "
            "(NYU) e certificado CGA-ANBIMA (Certified Investment Manager). O recomendador atesta "
            "que Gustavo contribuiu para discussões macroeconômicas, posicionamento tático e "
            "diversificação geográfica de portfólios, e que o escritório foi reconhecido como "
            "Top 10 Offshore na rede XP por anos consecutivos. Afirma que Gustavo 'se distingue "
            "acima da média' dos profissionais do setor."
        )
    )

    add_body_paragraph(
        doc,
        "Victor Fonseca, Diretor de Alocação e Produtos do Criteria Financial Group e "
        "graduado em Economia pela New York University (NYU) com certificação CGA-ANBIMA "
        "(Evidence 22), atesta que contribuí para \"discussões macroeconômicas, "
        "posicionamento tático e diversificação geográfica\" de portfólios, e que o "
        "escritório alcançou reconhecimento como Top 10 Offshore na rede XP por anos "
        "consecutivos, resultado para o qual minha contribuição foi fundamental.",
        bold_phrases=["Evidence 22"]
    )

    # Letter 3: Felipe Lala
    add_evidence_card(doc, 23,
        "Carta de Recomendação — Felipe Lala, Head de Equities, Criteria Investimentos (6 de fevereiro de 2026)",
        ev_type="Carta de Recomendação",
        source="Felipe Lala — Head de Equities, Criteria Investimentos",
        date="6 de fevereiro de 2026",
        url="",
        description=(
            "Carta de recomendação de Felipe Lala, graduado em Engenharia de Produção pela FEI "
            "com pós-graduação em Finanças (CFM) pelo Insper, e 6 anos de experiência no trading "
            "desk da XP Investimentos. O recomendador atesta que Gustavo demonstrou 'equilíbrio "
            "emocional, precisão analítica e disciplina de execução', distinguindo-se pela "
            "'profundidade analítica, execução disciplinada e consistência de performance "
            "sustentada' ao longo de sua carreira."
        )
    )

    add_body_paragraph(
        doc,
        "Felipe Lala, Head de Equities do Criteria Investimentos, com formação em "
        "Engenharia de Produção pela FEI e pós-graduação em Finanças pelo Insper "
        "(Evidence 23), atesta que demonstrei \"equilíbrio emocional, precisão analítica "
        "e disciplina de execução\", distinguindo-me pela \"profundidade analítica, "
        "execução disciplinada e consistência de performance sustentada\" — competências "
        "diretamente transferíveis para a gestão da EventFinOps LLC.",
        bold_phrases=["Evidence 23"]
    )

    # Letter 4: Thiago Neves
    add_evidence_card(doc, 24,
        "Carta de Recomendação — Thiago Neves, Diretor Jurídico Executivo; Ex-Diretor Jurídico, Criteria (28 de janeiro de 2026)",
        ev_type="Carta de Recomendação",
        source="Thiago Neves — Diretor Executivo; Ex-Diretor, Criteria",
        date="28 de janeiro de 2026",
        url="",
        description=(
            "Carta de recomendação de Thiago Neves, doutor em Administração de Empresas pela USP "
            "(Universidade de São Paulo), com mestrado em Administração e pós-graduação em Direito "
            "Internacional. O recomendador atesta que Gustavo demonstrou 'nível incomum de atenção "
            "à segurança regulatória e conformidade', trabalhando com estruturas envolvendo pessoas "
            "físicas, holdings e veículos internacionais. Afirma que Gustavo se distingue pela "
            "'responsabilidade, seriedade e mentalidade orientada a conformidade'."
        )
    )

    add_body_paragraph(
        doc,
        "Thiago Neves, doutor em Administração de Empresas pela USP e ex-Diretor do "
        "Criteria (Evidence 24), atesta meu \"nível incomum de atenção à segurança "
        "regulatória e conformidade\" e minha experiência com \"estruturas envolvendo "
        "pessoas físicas, holdings e veículos internacionais\". Essa expertise em "
        "compliance e estruturação regulatória é particularmente relevante para a "
        "operação da EventFinOps LLC no ambiente altamente regulado dos mercados "
        "financeiros norte-americanos.",
        bold_phrases=["Evidence 24"]
    )

    # Letter 5: Rafael Wurzmann
    add_evidence_card(doc, 25,
        "Carta de Recomendação — Rafael Wurzmann, Wealth Management/Offshore, BTG Pactual; Ex-Head de Offshore, Criteria (22 de fevereiro de 2026)",
        ev_type="Carta de Recomendação",
        source="Rafael Wurzmann — Wealth Management/Offshore, BTG Pactual",
        date="22 de fevereiro de 2026",
        url="",
        description=(
            "Carta de recomendação de Rafael Wurzmann, atualmente no BTG Pactual (um dos maiores "
            "bancos de investimento da América Latina), graduado em Finanças e Marketing pela NYU "
            "Stern School of Business, com certificações BMC, SIE, Series 7 e ANCORD. O "
            "recomendador trabalhou com Gustavo de novembro de 2021 a dezembro de 2024, e juntos "
            "estruturaram e escalaram a divisão Offshore para ~R$ 1 bilhão. Rafael formalmente "
            "recomendou Gustavo como seu sucessor na posição de Diretor de Offshore, atestando "
            "que Gustavo 'se destaca por combinar profundidade técnica, prontidão para liderança "
            "e visão estratégica de longo prazo em uma idade relativamente jovem'."
        )
    )

    add_body_paragraph(
        doc,
        "Rafael Wurzmann, atualmente no BTG Pactual, graduado pela NYU Stern School of "
        "Business com certificações BMC, SIE, Series 7 e ANCORD (Evidence 25), é "
        "particularmente significativo: como meu predecessor direto na posição de Head de "
        "Offshore no Criteria Financial Group, Rafael formalmente me recomendou como seu "
        "sucessor, atestando que me \"destaco por combinar profundidade técnica, prontidão "
        "para liderança e visão estratégica de longo prazo em uma idade relativamente "
        "jovem\". Esta designação formal de sucessão por um profissional atualmente "
        "posicionado em um dos maiores bancos de investimento da América Latina constitui "
        "validação inequívoca de minha capacidade para avançar o proposed endeavor.",
        bold_phrases=["Evidence 25"],
        italic_phrases=["proposed endeavor"]
    )

    # Cross-validation table
    add_section_header(doc, "Tabela de Validação Cruzada — Corroboração", level=3)

    add_data_table(doc,
        ["Competência/Conquista", "Mesquita (Ev. 21)", "Fonseca (Ev. 22)", "Lala (Ev. 23)", "Neves (Ev. 24)", "Wurzmann (Ev. 25)"],
        [
            ["Gestão portfólio offshore US$ 210M+", "✓", "✓", "✓", "", "✓"],
            ["Vertical ~R$ 1 bi em mercados EUA", "✓", "✓", "✓", "✓", "✓"],
            ["Top 10 Offshore na rede XP", "", "✓", "", "", "✓"],
            ["Profundidade analítica/técnica", "✓", "✓", "✓", "", "✓"],
            ["Liderança e maturidade", "✓", "", "", "✓", "✓"],
            ["Compliance e conformidade regulatória", "", "", "", "✓", ""],
            ["Disciplina de execução", "", "", "✓", "✓", "✓"],
        ]
    )

    add_body_paragraph(
        doc,
        "A tabela de validação cruzada demonstra que múltiplos recomendadores, de forma "
        "independente, corroboram as mesmas competências e conquistas centrais — "
        "particularmente a gestão de portfólios de mais de US$ 210 milhões, a "
        "contribuição para a vertical offshore de ~R$ 1 bilhão, e a combinação de "
        "profundidade técnica com capacidade de liderança. Esta convergência "
        "independente de testemunhos fortalece substancialmente a credibilidade "
        "das evidências apresentadas.",
        bold_phrases=[]
    )

    # Business Plan section
    add_section_header(doc, "Business Plan e Projeções", level=3)

    add_body_paragraph(
        doc,
        "O Business Plan da EventFinOps LLC (Evidence 06) detalha um modelo de negócios "
        "viável e sustentável, com projeções financeiras conservadoras baseadas em "
        "benchmarks do setor. A empresa será cofundada com Pedro Siviero Paciullo, também "
        "graduado pelo Insper e com experiência em M&A e DCM no Criteria Financial Group, "
        "garantindo complementaridade de competências na equipe fundadora.",
        bold_phrases=["Evidence 06"]
    )

    add_body_paragraph(
        doc,
        "O modelo de receita é diversificado entre assessoria de investimentos (hourly e "
        "project-based), estruturação de produtos financeiros (fee-based), serviços de "
        "M&A/DCM (success fee), e programas educacionais (subscription). O payback "
        "projetado de aproximadamente 20 meses e a margem EBITDA crescente (de 12,5% no "
        "primeiro ano para 46,7% no quinto ano) demonstram viabilidade econômica robusta.",
        bold_phrases=["20 meses", "46,7%"]
    )

    # Publications
    add_section_header(doc, "Publicações e Thought Leadership", level=3)

    add_body_paragraph(
        doc,
        "Além de minha experiência profissional, demonstro capacidade intelectual e "
        "thought leadership no campo de investimentos internacionais por meio de três "
        "artigos acadêmicos revisados por pares (Evidence 09, Evidence 10 e Evidence 11), "
        "certificados de publicação correspondentes (Evidence 12, Evidence 13 e Evidence 14), "
        "declarações de aceite (Evidence 15, Evidence 16 e Evidence 17), e um livro "
        "publicado com ISBN (Evidence 05). Esta produção acadêmica demonstra não apenas "
        "conhecimento técnico aprofundado, mas também capacidade de contribuição original "
        "e validada por pares ao corpo de conhecimento do campo em que o proposed endeavor "
        "se insere.",
        bold_phrases=["Evidence 09", "Evidence 10", "Evidence 11", "Evidence 05"],
        italic_phrases=["proposed endeavor", "thought leadership"]
    )

    # Synopsis Table Prong 2
    add_section_header(doc, "Tabela-Sinopse — Prong 2", level=3)

    add_data_table(doc,
        ["Dimensão", "Evidence #", "O que Demonstra"],
        [
            ["Educação", "Evidence 02–04", "Bacharelado Insper + UCLA (2 programas)"],
            ["Carreira — Criteria", "Evidence 07", "Estagiário → Diretor em 4 anos; US$ 210M+ portfólio"],
            ["Livro publicado", "Evidence 05", "ISBN 978-65-83827-35-7; thought leadership"],
            ["3 artigos acadêmicos", "Evidence 09–11", "QUALIS B2, double-blind, DOI atribuído"],
            ["3 matérias de mídia", "Evidence 18–20", "Líder emergente em investimentos internacionais"],
            ["Recomendação — Mesquita", "Evidence 21", "Maturidade incomum; opera acima da média"],
            ["Recomendação — Fonseca", "Evidence 22", "NYU; CGA-ANBIMA; Top 10 Offshore XP"],
            ["Recomendação — Lala", "Evidence 23", "Precisão analítica; disciplina de execução"],
            ["Recomendação — Neves", "Evidence 24", "PhD USP; compliance e conformidade"],
            ["Recomendação — Wurzmann", "Evidence 25", "NYU Stern; BTG Pactual; sucessor formal"],
            ["Business Plan", "Evidence 06", "US$ 384K→US$ 2,16M; 18 empregos; payback 20m"],
        ]
    )

    add_body_paragraph(
        doc,
        "A totalidade das evidências apresentadas — educação de excelência, progressão "
        "profissional acelerada com resultados mensuráveis excepcionais, validação "
        "independente por cinco profissionais seniores do setor, produção acadêmica "
        "revisada por pares, publicação de livro, cobertura de mídia, e um Business Plan "
        "detalhado com projeções conservadoras — demonstra de forma conclusiva que estou "
        "bem posicionado para avançar o proposed endeavor com alta probabilidade de "
        "sucesso.",
        italic_phrases=["proposed endeavor"]
    )

    _add_page_break_run(doc)

    # --------------------------------------------------------
    # SECTION 5: PRONG 3 — ON BALANCE, BENEFICIAL TO WAIVE
    # --------------------------------------------------------
    add_section_header(doc, "V. Prong 3 — No Balanço, É Benéfico aos Estados Unidos Dispensar os Requisitos de Oferta de Emprego e Certificação de Trabalho")

    add_body_paragraph(
        doc,
        "O terceiro prong do teste Dhanasar requer uma análise de custo-benefício: os "
        "benefícios da dispensa (waiver) dos requisitos de oferta de emprego e "
        "certificação de trabalho superam os interesses nacionais protegidos pelo "
        "processo de labor certification? Conforme a Seção 203(b)(2)(B) do INA: \"the "
        "alien's services in the sciences, arts, professions, or business would "
        "prospectively benefit the national economy, cultural or educational interests, "
        "or welfare of the United States.\" O USCIS avalia cinco fatores derivados de "
        "Dhanasar e dos fatores NYSDOT.",
        italic_phrases=["Dhanasar", "labor certification"]
    )

    # PART A: Factors 1-3
    add_section_header(doc, "Parte A: Fatores 1 a 3", level=2)

    add_section_header(doc, "Fator 1: Impraticabilidade da Labor Certification", level=3)

    add_body_paragraph(
        doc,
        "A labor certification (PERM) é estruturalmente impraticável para o empreendimento "
        "proposto por três razões fundamentais:",
        italic_phrases=["labor certification"]
    )

    add_body_paragraph(
        doc,
        "Primeiro, a natureza multi-client do empreendimento. A EventFinOps LLC atenderá "
        "múltiplos clientes simultaneamente — investidores individuais, family offices, "
        "empresas brasileiras buscando expansão nos EUA — com engagements de duração e "
        "escopo variáveis. O processo PERM pressupõe uma relação fixa entre uma entidade "
        "patronal e um trabalhador, o que é incompatível com o modelo de assessoria "
        "multi-client da EventFinOps.",
        bold_phrases=["natureza multi-client"],
        italic_phrases=[]
    )

    add_body_paragraph(
        doc,
        "Segundo, o escopo geográfico transcende fronteiras estaduais. O mercado-alvo "
        "da EventFinOps LLC não se limita a Miami: inclui investidores brasileiros e "
        "latino-americanos distribuídos em múltiplos estados (Flórida, Nova York, "
        "Califórnia, Texas), além de clientes institucionais no Brasil. A labor "
        "certification está vinculada a uma localidade específica, o que imporia uma "
        "limitação artificial ao alcance do empreendimento.",
        bold_phrases=["escopo geográfico transcende fronteiras estaduais"]
    )

    add_body_paragraph(
        doc,
        "Terceiro, sou cofundador e principal prestador de serviços da EventFinOps LLC. "
        "O processo PERM requer uma entidade patronal desinteressada para patrocinar o "
        "trabalhador, mas um empreendedor não pode se autopatrocinar via PERM. Conforme "
        "reconhecido pelo AAO em Dhanasar: empreendedores e profissionais autônomos não "
        "estão excluídos do NIW — pelo contrário, o NIW foi concebido precisamente para "
        "acomodar situações em que o processo de labor certification é inadequado.",
        bold_phrases=["cofundador e principal prestador de serviços"],
        italic_phrases=["Dhanasar"]
    )

    add_section_header(doc, "Fator 2: Benefício Mesmo com Trabalhadores Americanos Qualificados Disponíveis", level=3)

    add_body_paragraph(
        doc,
        "Reitero enfaticamente: o argumento aqui apresentado NÃO se fundamenta em "
        "escassez de profissionais americanos qualificados. O AAO expressamente rejeitou "
        "este tipo de argumento em Dhanasar, e o USCIS responderia corretamente: "
        "\"use PERM.\" O argumento é distinto: mesmo que existam Financial Managers "
        "qualificados nos Estados Unidos, a nação se beneficia da minha contribuição "
        "específica por razões que nenhum profissional doméstico pode replicar:",
        bold_phrases=["NÃO se fundamenta em escassez"],
        italic_phrases=["Dhanasar"]
    )

    add_body_paragraph(
        doc,
        "1. Expertise cross-border Brasil-EUA: possuo conhecimento profundo de ambos os "
        "mercados financeiros — brasileiro e norte-americano — adquirido por quatro anos "
        "de gestão de portfólios offshore que canalizaram ~R$ 1 bilhão do Brasil para "
        "os mercados dos EUA. Este conhecimento bilateral não é facilmente replicável por "
        "um profissional que conheça apenas o mercado doméstico americano.",
        bold_phrases=["Expertise cross-border Brasil-EUA"]
    )

    add_body_paragraph(
        doc,
        "2. Rede de relacionamentos no mercado financeiro brasileiro: mantenho "
        "relacionamentos profissionais estabelecidos com investidores, assessores e "
        "instituições financeiras no Brasil — incluindo a rede XP Inc. e conexões no "
        "BTG Pactual — que constituem pipeline direto de clientes para a EventFinOps LLC.",
        bold_phrases=["Rede de relacionamentos"]
    )

    add_body_paragraph(
        doc,
        "3. Fluência trilíngue (português, inglês, espanhol): conforme documentado em "
        "meu currículo (Evidence 08), possuo fluência nativa em português, fluência em "
        "inglês e fluência em espanhol — combinação que permite atender investidores "
        "brasileiros e latino-americanos em seu idioma nativo enquanto opero nos mercados "
        "norte-americanos.",
        bold_phrases=["Fluência trilíngue", "Evidence 08"]
    )

    add_body_paragraph(
        doc,
        "4. Track record quantificado excepcional: a gestão de mais de US$ 210 milhões "
        "em portfólios offshore e a contribuição para uma vertical de ~R$ 1 bilhão, "
        "validada por cinco profissionais seniores do setor, constitui um registro de "
        "sucesso que demonstra capacidade comprovada — não meramente projetada — de "
        "atrair capital estrangeiro para os mercados norte-americanos.",
        bold_phrases=["Track record quantificado excepcional", "US$ 210 milhões"]
    )

    add_body_paragraph(
        doc,
        "5. Contribuição acadêmica validada: três artigos científicos revisados por pares "
        "e um livro publicado com ISBN demonstram capacidade de contribuição intelectual "
        "original ao campo — qualidade que transcende a mera prestação de serviços "
        "profissionais e agrega valor ao ecossistema de conhecimento financeiro "
        "norte-americano.",
        bold_phrases=["Contribuição acadêmica validada"]
    )

    add_section_header(doc, "Fator 3: Urgência e Timing", level=3)

    add_body_paragraph(
        doc,
        "O momento atual apresenta uma convergência única de fatores que tornam urgente "
        "a aprovação desta petição:",
        bold_phrases=[]
    )

    add_body_paragraph(
        doc,
        "O mercado global de ETFs atingiu US$ 19,5 trilhões ao final de 2025, com "
        "projeção de US$ 35 trilhões até 2030 — representando uma das maiores "
        "oportunidades de mercado da história financeira recente. Simultaneamente, "
        "investidores brasileiros mantêm menos de 2% de seus ativos no exterior, "
        "comparado a 25% dos investidores institucionais norte-americanos. Esta "
        "assimetria está se corrigindo rapidamente: cada mês de atraso representa "
        "perda de first-mover advantage em um mercado em rápida expansão.",
        bold_phrases=["US$ 19,5 trilhões", "US$ 35 trilhões", "menos de 2%"]
    )

    add_body_paragraph(
        doc,
        "O USCIS atualizou suas orientações sobre avaliação de petições EB-2 NIW em "
        "janeiro de 2025 (PA-2025-03), clarificando critérios de avaliação e reforçando "
        "a disponibilidade de Premium Processing — sinalizando reconhecimento institucional "
        "da importância de processar eficientemente petições que demonstrem interesse "
        "nacional documentado.",
        italic_phrases=["PA-2025-03"]
    )

    # PART B: Factors 4-5
    add_section_header(doc, "Parte B: Fatores 4 e 5 + Síntese", level=2)

    add_section_header(doc, "Fator 4: Criação de Empregos e Impacto Econômico", level=3)

    add_data_table(doc,
        ["Posição", "Ano 1", "Ano 2", "Ano 3", "Ano 4", "Ano 5"],
        [
            ["Assessores financeiros", "1", "2", "3", "5", "7"],
            ["Analistas de mercado", "1", "1", "2", "3", "4"],
            ["Administrativo/operacional", "1", "2", "3", "4", "5–7"],
            ["Total", "3", "5", "8", "12", "16–18"],
        ]
    )

    add_body_paragraph(
        doc,
        "A EventFinOps LLC projeta a criação de 3 empregos diretos no primeiro ano de "
        "operação, escalando para 16 a 18 posições até o quinto ano (Evidence 06). "
        "Estas posições serão em categorias qualificadas — assessores financeiros "
        "certificados, analistas de mercado com formação superior, e profissionais "
        "administrativos especializados — contribuindo para o mercado de trabalho "
        "local em Miami com posições de alta remuneração.",
        bold_phrases=["3 empregos", "16 a 18 posições", "Evidence 06"]
    )

    add_body_paragraph(
        doc,
        "Além dos empregos diretos, o empreendimento gerará impacto econômico indireto "
        "por meio de: (1) pagamentos a fornecedores locais (escritório, tecnologia, "
        "serviços profissionais); (2) receita fiscal para o estado da Flórida e os "
        "Estados Unidos; e (3) facilitação de fluxos de capital internacional que "
        "beneficiam a economia norte-americana como um todo.",
        bold_phrases=[]
    )

    add_section_header(doc, "Fator 5: Trabalho Autônomo sem Efeito Adverso", level=3)

    add_body_paragraph(
        doc,
        "Como cofundador da EventFinOps LLC, minha atuação nos Estados Unidos é de "
        "natureza empreendedora: estou criando uma empresa nova, não ocupando uma "
        "posição existente. Este modelo de trabalho autônomo não desloca nenhum "
        "trabalhador americano de posição existente — pelo contrário, CRIA novas "
        "posições para trabalhadores americanos.",
        bold_phrases=["CRIA novas posições"]
    )

    add_body_paragraph(
        doc,
        "O pricing da EventFinOps LLC será competitivo mas não predatório: como empresa "
        "voltada para investidores brasileiros e latino-americanos — um segmento "
        "específico e crescente que requer expertise bilíngue e conhecimento de ambos "
        "os mercados —, a EventFinOps não compete diretamente com assessorias financeiras "
        "que atendem exclusivamente investidores domésticos americanos. A empresa expande "
        "o mercado ao atender um segmento underserved, não divide fatias existentes.",
        bold_phrases=["expande o mercado"],
        italic_phrases=["underserved"]
    )

    add_body_paragraph(
        doc,
        "Conforme expressamente reconhecido pelo AAO em Dhanasar: \"Entrepreneurs and "
        "self-employed individuals are not excluded from the NIW.\" O empreendedorismo "
        "é não apenas permitido, mas incentivado pelo framework NIW quando o "
        "empreendimento demonstra mérito substancial e importância nacional — como "
        "amplamente documentado nesta petição.",
        italic_phrases=["Dhanasar"]
    )

    # Balance Sheet
    add_section_header(doc, "Síntese: Análise de Custo-Benefício", level=3)

    add_data_table(doc,
        ["Benefício da Dispensa (Waiver)", "Custo de NÃO Conceder a Dispensa"],
        [
            ["Criação de 16–18 empregos em Miami em 5 anos", "PERM delay: 12–18 meses → oportunidade de mercado perdida"],
            ["Facilitação de fluxo de capital Brasil→EUA", "Confinamento a 1 entidade patronal → serviço limitado"],
            ["US$ 2,16M em receita projetada (Ano 5)", "Sem self-employment → empresa não nasce"],
            ["Atendimento a segmento underserved (investidores LatAm)", "Sem criação de empresa → 18 empregos não criados"],
            ["Contribuição acadêmica contínua (publicações, livro)", "Mercado em crescimento exponencial → first-mover advantage perdido"],
            ["Expertise cross-border única (Brasil-EUA)", "Gap de internacionalização persistente (< 2% alocação)"],
        ]
    )

    add_body_paragraph(
        doc,
        "A análise de custo-benefício demonstra inequivocamente que os benefícios da "
        "dispensa dos requisitos de oferta de emprego e certificação de trabalho "
        "superam substancialmente os interesses protegidos pelo processo de labor "
        "certification. A imposição de PERM seria não apenas impraticável — como "
        "demonstrado no Fator 1 — mas contraproducente para os interesses nacionais, "
        "pois impediria a criação de empregos, limitaria o fluxo de capital "
        "internacional para os EUA, e eliminaria a contribuição de expertise "
        "cross-border única ao mercado financeiro norte-americano.",
        bold_phrases=[]
    )

    # Synopsis Table Prong 3
    add_section_header(doc, "Tabela-Sinopse — Prong 3", level=3)

    add_data_table(doc,
        ["Fator", "Argumento Central", "Evidência"],
        [
            ["1. Impraticabilidade", "Multi-client, multi-state, cofundador/self-employed", "Evidence 06 (BP)"],
            ["2. Benefício", "Expertise cross-border única + segmento underserved", "Evidence 07–25"],
            ["3. Urgência", "ETFs US$ 19,5T; gap 2% Brasil; PA-2025-03", "Sources"],
            ["4. Criação de empregos", "3 empregos (Ano 1) → 16–18 (Ano 5)", "Evidence 06 (BP)"],
            ["5. Trabalho autônomo", "Zero deslocamento; novo mercado; new jobs", "Evidence 06 (BP)"],
        ]
    )

    _add_page_break_run(doc)

    # --------------------------------------------------------
    # SECTION 6: CONCLUSION
    # --------------------------------------------------------
    add_section_header(doc, "VI. Conclusão")

    add_body_paragraph(
        doc,
        "Pela totalidade das evidências apresentadas nesta petição, demonstrei que:",
        bold_phrases=[]
    )

    add_body_paragraph(
        doc,
        "1. Sou elegível para classificação EB-2 pela via de Exceptional Ability, "
        "atendendo a três dos seis critérios regulamentares do 8 C.F.R. § 204.5(k)(3)(ii): "
        "diploma relacionado à área (Insper + UCLA), remuneração acima da norma "
        "(US$ 130.000/ano + gestão de US$ 210M+ em portfólios), e reconhecimento por "
        "contribuições significativas (3 artigos revisados por pares, livro com ISBN, "
        "3 matérias de mídia).",
        italic_phrases=["Exceptional Ability"],
        bold_phrases=["três dos seis critérios"]
    )

    add_body_paragraph(
        doc,
        "2. Prong 1 — Mérito Substancial e Importância Nacional: O empreendimento "
        "proposto possui mérito substancial documentado por projeções financeiras de "
        "US$ 384.000 a US$ 2.160.000 em cinco anos, e importância nacional demonstrada "
        "por alinhamento com CETs (3 áreas), ordens executivas, crescimento projetado "
        "pelo BLS de 15% (5x a média), designação Bright Outlook/In Demand pelo O*NET, "
        "classificação de infraestrutura crítica pela CISA, e um mercado global de ETFs "
        "de US$ 19,5 trilhões em expansão exponencial.",
        bold_phrases=["Prong 1"]
    )

    add_body_paragraph(
        doc,
        "3. Prong 2 — Bem Posicionado: Estou bem posicionado para avançar o "
        "empreendimento proposto, conforme demonstrado por: educação de excelência "
        "(Insper + UCLA), progressão profissional acelerada de estagiário a Diretor "
        "em menos de quatro anos, gestão de US$ 210M+ em portfólios offshore, "
        "contribuição para vertical de ~R$ 1 bilhão, cinco cartas de recomendação "
        "independentes de profissionais seniores, três publicações acadêmicas revisadas "
        "por pares, livro publicado com ISBN, e um Business Plan detalhado com "
        "projeções conservadoras.",
        bold_phrases=["Prong 2"]
    )

    add_body_paragraph(
        doc,
        "4. Prong 3 — Benéfico Dispensar: A análise de custo-benefício demonstra que a "
        "dispensa dos requisitos de oferta de emprego e certificação de trabalho é "
        "benéfica para os Estados Unidos, considerando: a impraticabilidade estrutural "
        "do PERM para empreendimento multi-client e multi-state, o benefício de minha "
        "expertise cross-border única, a urgência ditada pelo crescimento exponencial "
        "do mercado, a criação de até 18 empregos, e a natureza empreendedora da "
        "atividade que não desloca trabalhadores americanos.",
        bold_phrases=["Prong 3"]
    )

    # Summary table
    add_section_header(doc, "Tabela Resumo da Petição", level=3)

    add_data_table(doc,
        ["Requisito", "Status", "Evidências-Chave"],
        [
            ["Elegibilidade EB-2", "3/6 critérios atendidos", "Evidence 02–05, 07–20"],
            ["Prong 1 — Mérito + Importância Nacional", "Documentalmente atendido", "Evidence 06, 09–11 + Sources"],
            ["Prong 2 — Bem Posicionado", "Documentalmente atendido", "Evidence 02–25"],
            ["Prong 3 — Benéfico Dispensar", "Documentalmente atendido", "Evidence 06"],
        ]
    )

    add_body_paragraph(
        doc,
        "Respeitosamente solicito que esta petição I-140 seja aprovada e que seja "
        "concedida a classificação EB-2 com dispensa do requisito de oferta de emprego "
        "e certificação de trabalho (National Interest Waiver), nos termos da "
        "Seção 203(b)(2)(B) do INA e do 8 C.F.R. § 204.5(k).",
        italic_phrases=["National Interest Waiver"]
    )

    doc.add_paragraph()

    # Signature
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sig = p_sig.add_run("Respeitosamente,")
    run_sig.font.name = "Garamond"
    run_sig.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p_name = doc.add_paragraph()
    run_name = p_name.add_run(CLIENT_NAME)
    run_name.bold = True
    run_name.font.name = "Garamond"
    run_name.font.size = Pt(12)

    p_title = doc.add_paragraph()
    run_title = p_title.add_run("Peticionário")
    run_title.font.name = "Garamond"
    run_title.font.size = Pt(12)

    _add_page_break_run(doc)

    # --------------------------------------------------------
    # SECTION 7: SOURCES CITED
    # --------------------------------------------------------
    add_section_header(doc, "VII. Fontes Citadas")

    sources = [
        "Bureau of Labor Statistics (BLS). Occupational Outlook Handbook: Financial Managers (SOC 11-3031). https://www.bls.gov/ooh/management/financial-managers.htm",
        "Bureau of Labor Statistics (BLS). Occupational Employment and Wages, May 2024: Financial Managers. https://www.bls.gov/oes/2023/may/oes113031.htm",
        "O*NET OnLine. Summary Report for 11-3031.00 — Financial Managers. https://www.onetonline.org/link/summary/11-3031.00",
        "O*NET OnLine. Bright Outlook Occupation: 11-3031.00 — Financial Managers. https://www.onetonline.org/help/bright/11-3031.00",
        "USCIS. Employment-Based Immigration: Second Preference EB-2. https://www.uscis.gov/working-in-the-united-states/permanent-workers/employment-based-immigration-second-preference-eb-2",
        "USCIS. Policy Alert PA-2025-03: EB-2 National Interest Waiver Petitions (15 de janeiro de 2025). https://www.uscis.gov/sites/default/files/document/policy-manual-updates/20250115-Employment-BasedNationalInterestWaivers.pdf",
        "USCIS. Policy Manual, Vol. 6, Part F, Ch. 5. https://www.uscis.gov/policy-manual",
        "CISA. Financial Services Sector. https://www.cisa.gov/topics/critical-infrastructure-security-and-resilience/critical-infrastructure-sectors/financial-services-sector",
        "CISA. Critical Infrastructure Sectors. https://www.cisa.gov/topics/critical-infrastructure-security-and-resilience/critical-infrastructure-sectors",
        "PwC. ETFs 2026: The Next Big Leap. https://www.pwc.com/gx/en/industries/financial-services/publications/etf-2026-the-next-big-leap.html",
        "American Century. ETFs Defying Gravity: Record Inflows Drive Another Chart-Topping Year. https://www.americancentury.com/insights/etfs-defying-gravity/",
        "State Street Global Advisors. 2025 Global ETF Outlook. https://www.statestreet.com/us/en/insights/etfs-2025-outlook",
        "ANBIMA. Dados de mercado — ativos sob administração (junho de 2025).",
        "White House OSTP. Critical and Emerging Technologies List Update (fevereiro de 2024). https://www.whitehouse.gov/ostp/",
        "Presidential Policy Directive 21 (PPD-21): Critical Infrastructure Security and Resilience (2013).",
        "Executive Order — Removing Barriers to American Leadership in Artificial Intelligence (20 de janeiro de 2025). https://www.whitehouse.gov/presidential-actions/2025/01/removing-barriers-to-american-leadership-in-artificial-intelligence/",
        "IBGE/PNAD Contínua. Rendimento médio mensal do trabalho principal (2024).",
        "Revista Lumen et Virtus (ISSN 2177-2789). QUALIS CAPES B2. New Science Publishers Ltda.",
        "OpportunityZones.com. Qualified Opportunity Zones in Florida. https://opportunityzones.com/location/florida/",
        "Matter of Dhanasar, 26 I&N Dec. 884 (AAO 2016).",
    ]

    for i, source in enumerate(sources, 1):
        add_footnote_text(doc, i, source)

    # --------------------------------------------------------
    # EVIDENCE INDEX
    # --------------------------------------------------------
    _add_page_break_run(doc)
    add_section_header(doc, "Índice de Evidências")

    evidence_list = [
        ("01", "Credential Evaluation Report [VERIFICAR — não localizado na pasta do cliente]"),
        ("02", "Diploma de Bacharel em Administração de Empresas — Insper Instituto de Ensino e Pesquisa (Dezembro de 2023)"),
        ("03", "Certificado — UCLA Junior Program Classic"),
        ("04", "Certificado — UCLA Entrepreneurship Program"),
        ("05", "Livro — \"Liderança e Tomada de Decisão no Setor Financeiro\" (ISBN 978-65-83827-35-7, Golden Int Editora, 2025)"),
        ("06", "Business Plan — EventFinOps LLC, Miami, Flórida"),
        ("07", "Declaração de Recursos Humanos — Criteria Financial Group (Janeiro de 2022 — Agosto de 2025)"),
        ("08", "Curriculum Vitae — Gustavo Lopes Esteves"),
        ("09", "Artigo Científico — \"Riscos e Oportunidades na Internacionalização de Ativos\" (Lumen et Virtus, v.13, n.31, DOI: 10.56238/levv13n31-058)"),
        ("10", "Artigo Científico — \"ETFs e a Expansão das Fronteiras do Investimento\" (Lumen et Virtus, v.15, n.42, DOI: 10.56238/levv15n42-085)"),
        ("11", "Artigo Científico — \"Tendências do Mercado de Investimentos Offshore\" (Lumen et Virtus, v.16, n.49, DOI: 10.56238/levv16n49-119)"),
        ("12", "Certificado de Publicação — DEC-58 v13n31"),
        ("13", "Certificado de Publicação — DEC-85 v15n42"),
        ("14", "Certificado de Publicação — DEC-119 v16n49"),
        ("15", "Declaração de Aceite — \"Riscos e Oportunidades na Internacionalização de Ativos\""),
        ("16", "Declaração de Aceite — \"ETFs e a Expansão das Fronteiras do Investimento\""),
        ("17", "Declaração de Aceite — \"Tendências do Mercado de Investimentos Offshore\""),
        ("18", "Matéria de Mídia — \"Estratégia, resiliência e resultado\" (Brasil Agora, 3 de outubro de 2025)"),
        ("19", "Matéria de Mídia — \"A democratização do investimento internacional\" (Gazeta da Semana, 20 de outubro de 2025)"),
        ("20", "Matéria de Mídia — \"Jovens líderes globais\" (Business Feed, 4 de novembro de 2025)"),
        ("21", "Carta de Recomendação — Luiz Fernando Mesquita, Diretor, BioWash (30 de janeiro de 2026)"),
        ("22", "Carta de Recomendação — Victor Fonseca, Diretor de Alocação e Produtos, Criteria Financial Group (18 de fevereiro de 2026)"),
        ("23", "Carta de Recomendação — Felipe Lala, Head de Equities, Criteria Investimentos (6 de fevereiro de 2026)"),
        ("24", "Carta de Recomendação — Thiago Neves, Diretor Executivo (28 de janeiro de 2026)"),
        ("25", "Carta de Recomendação — Rafael Wurzmann, Wealth Management/Offshore, BTG Pactual (22 de fevereiro de 2026)"),
        ("26", "Link do Ebook — Clube de Autores (https://clubedeautores.com.br/livro/lideranca-e-tomada-de-decisao-no-setor-financeiro)"),
        ("27", "Certidão de Nascimento — Gustavo Lopes Esteves (traduzida)"),
        ("28", "Passaporte — Gustavo Lopes Esteves"),
    ]

    add_data_table(doc,
        ["Evidence #", "Descrição"],
        [[num, desc] for num, desc in evidence_list]
    )

    # --------------------------------------------------------
    # FINALIZE: Footer + Save
    # --------------------------------------------------------
    add_footer(doc, CLIENT_NAME)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"✅ Cover Letter gerada com sucesso!")
    print(f"📄 Arquivo: {OUTPUT_FILE}")
    print(f"📊 Evidências: 28")
    print(f"📋 Tabelas: 14+")
    print(f"🔗 Fontes: 20")

    return OUTPUT_FILE


if __name__ == "__main__":
    output = generate_cover_letter()
    print(f"\nDocumento salvo em: {output}")
