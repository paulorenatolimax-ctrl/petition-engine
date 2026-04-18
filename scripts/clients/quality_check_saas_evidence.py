#!/usr/bin/env python3
"""
Quality Gate Scanner for SaaS Evidence Document
Checks all error rules from the generation prompt.
"""

import re
from docx import Document

DOC_PATH = "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2025/CAMILLA/_Forjado por Petition Engine/saas_evidence_Camilla_Santana_Pereira_Paes_de_Barros.docx"

def extract_text(path):
    doc = Document(path)
    text = []
    for p in doc.paragraphs:
        text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    return "\n".join(text)

def check_rules(text):
    results = []

    # CRITICAL/BLOCK rules
    critical_rules = [
        (r'\b(I|we)\s+believe\b', 'CRITICAL/BLOCK: "I/we believe" found'),
        (r'\b(I|we)\s+think\b', 'HIGH/BLOCK: "I/we think" found'),
        (r'\b(in conclusion|to summarize)\b', 'HIGH/BLOCK: "in conclusion/to summarize" found'),
        (r'(23-1011|29-1069|17-201[1-9]|13-2011)', 'CRITICAL/BLOCK: Prohibited SOC code found'),
        (r'\bprompt\b', 'CRITICAL/BLOCK: Word "prompt" found in output'),
        (r'(PROEX|Kortix|Carlos Avelino)', 'CRITICAL/BLOCK: Internal reference found'),
        (r'\b(introducao|peticao|informacao|certificacao|formacao|avaliacao|ocupacao|operacao|integracao|migracao|capacitacao|micropigmentacao)\b',
         'CRITICAL/BLOCK: Missing Portuguese accents found'),
        (r'\b(Version \d|Generated:|SaaS Evidence Architect|Petition Engine)\b',
         'CRITICAL/BLOCK: Internal system term found in output'),
        (r'\b(standardized|padronizado|operates autonomously|self-sustaining|auto-sustent|plug.and.play|train.the.trainer|white.label|marca branca|client autonomy|founder dependency|scalable without|replicable by any|turnkey|chave.na.m)\b',
         'CRITICAL/BLOCK: Anti-Cristine V2 prohibited term found'),
        (r'\b(consultoria|consulting)\b', 'HIGH/WARN: "consultoria/consulting" found'),
        (r'\b(denial|negativa anterior|RFE anterior|previous petition|prior filing|refile|segunda tentativa|nova submiss|peti..o anterior|corrected approach|lessons learned from denial)\b',
         'CRITICAL/BLOCK: Case history reference found'),
        (r'\b(petition|petitioner|EB-2|EB-1|NIW|USCIS|immigration|visa|I-140|green.card|adjudicator|Dhanasar|Kazarian|extraordinary.ability|national.interest.waiver)\b',
         'CRITICAL/BLOCK: Immigration term found in SaaS Evidence'),
    ]

    violations = 0
    warnings = 0

    for pattern, message in critical_rules:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            severity = "BLOCK" if "BLOCK" in message else "WARN"
            if severity == "BLOCK":
                violations += 1
            else:
                warnings += 1
            # Show context for each match
            for match in matches[:3]:
                match_str = match if isinstance(match, str) else match[0] if isinstance(match, tuple) else str(match)
                # Find context
                idx = text.lower().find(match_str.lower())
                if idx >= 0:
                    start = max(0, idx - 40)
                    end = min(len(text), idx + len(match_str) + 40)
                    context = text[start:end].replace("\n", " ")
                    results.append(f"  [{message}]")
                    results.append(f"    Match: '{match_str}' in: ...{context}...")
                else:
                    results.append(f"  [{message}] Match: '{match_str}'")
            if len(matches) > 3:
                results.append(f"    ... and {len(matches) - 3} more occurrences")

    return violations, warnings, results

def main():
    print("=" * 60)
    print("QUALITY GATE — SaaS Evidence Camilla Santana")
    print("=" * 60)

    text = extract_text(DOC_PATH)
    word_count = len(text.split())
    print(f"\nDocument word count: {word_count}")
    print(f"Document char count: {len(text)}")

    violations, warnings, results = check_rules(text)

    print(f"\n--- SCAN RESULTS ---")
    print(f"BLOCK violations: {violations}")
    print(f"WARN findings: {warnings}")

    if results:
        print("\nDetails:")
        for r in results:
            print(r)

    if violations == 0:
        print("\n[PASS] All CRITICAL/BLOCK rules passed.")
    else:
        print(f"\n[FAIL] {violations} BLOCK violation(s) found. Document needs fixes.")

    if warnings > 0:
        print(f"[WARN] {warnings} warning(s) found. Review recommended.")

    print("\n" + "=" * 60)
    return violations

if __name__ == "__main__":
    exit(main())
