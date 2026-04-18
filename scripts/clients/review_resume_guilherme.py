#!/usr/bin/env python3
"""
REVIEWED DOCX Generator — Guilherme Caldeira De Lello
Applies corrections from Separation of Concerns review.
Reads V1, outputs _REVIEWED version.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy
import os
import re

# ─── CONFIG ────────────────────────────────────────────────────
BASE_DIR = "/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2026/Guilherme Caldeira De Lello OS5195 15DEZ25 Kayenne AJUSTE/_Forjado por Petition Engine"
INPUT_FILE = os.path.join(BASE_DIR, "resume_eb1a_Guilherme_Caldeira_De_Lello_V1.docx")
OUTPUT_FILE = os.path.join(BASE_DIR, "resume_eb1a_Guilherme_Caldeira_De_Lello_REVIEWED.docx")

# Colors
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
BORDER_COLOR = "CCCCCC"
HEADER_BG = "D5E8F0"

# Exhibit numbering map — matches the order evidence blocks appear in the document
EXHIBIT_MAP = [
    ("Exhibit 1", "Tech & Learning Awards of Excellence — Best of 2024 (EduxGen.AI)", "Criterion 1"),
    ("Exhibit 2", "IDB Gobernarte Award 2020 — Innovation Category (CMSP)", "Criterion 1"),
    ("Exhibit 3", "WISE Award 2009 — Qatar Foundation (Distance Learning Amazon)", "Criterion 1"),
    ("Exhibit 4", "Assespro-RJ Award 2013 — Telecom & IT Infrastructure (SynCast)", "Criterion 1"),
    ("Exhibit 5", "REALMS — Software Registration (INPI BR512023001278-5)", "Criterion 5"),
    ("Exhibit 6", "Mano (Android) — Software Registration (INPI BR512018051989-0)", "Criterion 5"),
    ("Exhibit 7", "SynCast IP.TV Sync — Software Registration (INPI)", "Criterion 5"),
    ("Exhibit 8", "sEMG Hand Gesture Classification Research — SBIC Journal, 2024", "Criterion 5"),
    ("Exhibit 9", "WiSARD NILM Research — IEEE IPDPSW, 2020", "Criterion 5"),
    ("Exhibit 10", "Journal Article: Comparison of sEMG Classifiers — SBIC, 2024", "Criterion 6"),
    ("Exhibit 11", "Conference Paper: WiSARD NILM — IEEE IPDPSW, 2020", "Criterion 6"),
    ("Exhibit 12", "Master's Dissertation — COPPE/UFRJ, 2025", "Criterion 6"),
    ("Exhibit 13", "IP.TV LTDA / IP.TV LAB — Corporate Records & Project Docs", "Criterion 8"),
    ("Exhibit 14", "Centro de Midias SP (CMSP) — Government Project Documentation", "Criterion 8"),
    ("Exhibit 15", "VAT Tecnologia — Corporate Records & Award Documentation", "Criterion 8"),
    ("Exhibit 16", "M.Sc. Electrical Engineering — COPPE/UFRJ Diploma", "Academic"),
    ("Exhibit 17", "B.Eng. Electronic & Computing Engineering — UFRJ Diploma", "Academic"),
]


def add_page_numbers(doc):
    """Add page numbering to footer: Page X of Y Name – EB-1"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Clear existing
        for p in footer.paragraphs:
            p.clear()

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # "Page "
        run1 = p.add_run("Page ")
        run1.font.name = "Arial"
        run1.font.size = Pt(9)
        run1.font.color.rgb = MED_GRAY

        # PAGE field
        fld_char1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_page = p.add_run()
        run_page._r.append(fld_char1)
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_instr = p.add_run()
        run_instr._r.append(instr)
        fld_char2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_end = p.add_run()
        run_end._r.append(fld_char2)

        # " of "
        run_of = p.add_run(" of ")
        run_of.font.name = "Arial"
        run_of.font.size = Pt(9)
        run_of.font.color.rgb = MED_GRAY

        # NUMPAGES field
        fld_char3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_np = p.add_run()
        run_np._r.append(fld_char3)
        instr2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
        run_instr2 = p.add_run()
        run_instr2._r.append(instr2)
        fld_char4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_end2 = p.add_run()
        run_end2._r.append(fld_char4)

        # " Résumé Guilherme Caldeira De Lello – EB-1"
        run_name = p.add_run("  Résumé Guilherme Caldeira De Lello – EB-1")
        run_name.font.name = "Arial"
        run_name.font.size = Pt(9)
        run_name.font.color.rgb = MED_GRAY


def add_exhibit_numbers(doc):
    """Add exhibit numbers to evidence block tables."""
    exhibit_idx = 0
    for table in doc.tables:
        # Identify evidence block tables: have [THUMBNAIL] in first cell
        first_cell = table.cell(0, 0)
        if "[THUMBNAIL]" in first_cell.text:
            if exhibit_idx < len(EXHIBIT_MAP):
                exhibit_num = EXHIBIT_MAP[exhibit_idx][0]
                # Add exhibit number below [THUMBNAIL]
                p = first_cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(exhibit_num)
                run.font.name = "Arial"
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = DARK_GRAY
                exhibit_idx += 1


def add_exhibit_index(doc):
    """Add Exhibit Index table at the end of the document."""
    # Page break before index
    p_break = doc.add_paragraph()
    p_break.paragraph_format.page_break_before = True
    p_break.paragraph_format.space_before = Pt(24)
    p_break.paragraph_format.space_after = Pt(12)
    run_title = p_break.add_run("EXHIBIT INDEX")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0, 0, 0)

    # Separator line
    pPr = p_break._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{BORDER_COLOR}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True

    # Header row
    hdr = table.rows[0]
    headers = ["Exhibit #", "Description", "Related Criterion"]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = True
        # Header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG}" w:val="clear"/>')
        tcPr.append(shd)

    # Data rows
    for exhibit_num, description, criterion in EXHIBIT_MAP:
        row = table.add_row()
        values = [exhibit_num, description, criterion]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9.5)

    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


# ─── TEXT CORRECTIONS ──────────────────────────────────────────

TEXT_REPLACEMENTS = {
    # C2: Soften conclusory language
    "is a distinguished Architectural and Engineering Manager":
        "is an Architectural and Engineering Manager",

    "established himself as a preeminent figure at the intersection of engineering leadership and technological innovation in Brazil":
        "built a sustained record of achievement at the intersection of engineering leadership and technological innovation in Brazil",

    "at extraordinary scale":
        "at national scale",

    # C1: Reduce CMSP repetition (later mentions)
    # Only reduce in the leadership and addendum sections
    "Under Mr. De Lello's engineering leadership, IP.TV has developed and deployed multiple proprietary software platforms (REALMS, Mano, EduxGen.AI) serving educational institutions and government clients. The company's most significant achievement is the Centro de Midias SP (CMSP), a platform that delivered simultaneous educational access to over one million students and educators for the State of Sao Paulo during the COVID-19 pandemic. This deployment represents one of the largest emergency educational technology initiatives in Latin American history and was recognized by the Inter-American Development Bank with the Gobernarte 2020 Award.":
        "Under Mr. De Lello's engineering leadership, IP.TV has developed and deployed multiple proprietary software platforms (REALMS, Mano, EduxGen.AI) serving educational institutions and government clients. The company's most significant achievement, the Centro de Midias SP (CMSP) platform, is documented in detail under Criterion 1 above.",

    # Soften more conclusory language
    "a rare convergence of deep technical expertise":
        "a productive convergence of deep technical expertise",

    "extraordinary achievement spanning sixteen years":
        "sustained achievement spanning sixteen years",

    "extraordinary pace of deployment":
        "rapid pace of deployment",

    # C5: Vary repetitive "under Mr. De Lello's engineering leadership"
    "developed under Mr. De Lello's engineering leadership at IP.TV LAB EDUCATION LTDA":
        "developed by the IP.TV LAB EDUCATION LTDA team directed by Mr. De Lello",

    "developed under Mr. De Lello's technical leadership at IP.TV LTDA":
        "created under Mr. De Lello's direction at IP.TV LTDA",

    "developed under Mr. De Lello's technical direction at VAT Tecnologia":
        "built by the VAT Tecnologia team led by Mr. De Lello",

    "distinguished organizations":
        "established organizations",

    "distinguished institutions":
        "recognized institutions",

    # Addendum fixes — reduce padding
    "In the domain of technology enterprise leadership, Mr. De Lello founded and directed organizations that have delivered infrastructure-critical technology platforms at national scale. The Centro de Midias SP platform, serving over one million simultaneous users for the State of Sao Paulo's education system, stands as one of the most significant emergency educational technology deployments in Latin American history. This achievement was independently validated by the Inter-American Development Bank through the Gobernarte 2020 Award, selected from 127 initiatives across 12 nations.":
        "In the domain of technology enterprise leadership, Mr. De Lello founded and directed organizations that delivered infrastructure-critical platforms at national scale, as documented under Criteria 1 and 8 above.",

    "In the domain of applied research and original contributions, Mr. De Lello has produced peer-reviewed publications in IEEE proceedings and the Brazilian Society on Computational Intelligence journal, advancing the fields of sEMG-based hand gesture recognition for prosthetic devices and weightless neural networks for energy monitoring. His four INPI-registered software programs translate these research innovations into commercially deployed products serving educational institutions, assistive technology users, and enterprise clients.":
        "In the domain of applied research, his peer-reviewed publications and four INPI-registered software programs translate research innovations into commercially deployed products, as detailed under Criteria 5 and 6.",

    "In the domain of international recognition, Mr. De Lello has received four independent awards from organizations spanning three continents: the Inter-American Development Bank (Washington, D.C.), Tech & Learning Magazine (United States), the Qatar Foundation (Doha), and Assespro-RJ (Brazil). These recognitions span a fifteen-year period (2009-2024), confirming the sustained and evolving nature of his contributions to the field.":
        "His four independent international awards, spanning three continents and a fifteen-year period (2009-2024), confirm the sustained nature of his contributions, as documented under Criterion 1.",
}


def apply_text_corrections(doc):
    """Apply text replacements to paragraphs."""
    for para in doc.paragraphs:
        full_text = para.text
        for old, new in TEXT_REPLACEMENTS.items():
            if old in full_text:
                # Replace in runs
                # Strategy: find the replacement in the concatenated run text
                # and rebuild runs
                new_text = full_text.replace(old, new)
                if new_text != full_text:
                    # Clear all runs and set new text, preserving first run's formatting
                    if para.runs:
                        first_run = para.runs[0]
                        # Store formatting
                        font_name = first_run.font.name
                        font_size = first_run.font.size
                        font_bold = first_run.font.bold
                        font_italic = first_run.font.italic
                        font_color = first_run.font.color.rgb if first_run.font.color.type else RGBColor(0, 0, 0)

                        # Clear paragraph
                        for run in para.runs:
                            run._element.getparent().remove(run._element)

                        # Add new run with text
                        new_run = para.add_run(new_text)
                        new_run.font.name = font_name or "Arial"
                        new_run.font.size = font_size or Pt(10.5)
                        new_run.font.bold = font_bold
                        new_run.font.italic = font_italic
                        new_run.font.color.rgb = font_color

                    full_text = new_text  # update for chained replacements


# ─── MAIN ──────────────────────────────────────────────────────

print(f"Reading: {INPUT_FILE}")
doc = Document(INPUT_FILE)

print("Applying text corrections...")
apply_text_corrections(doc)

print("Adding page numbers to footer...")
add_page_numbers(doc)

print("Adding exhibit numbers to evidence blocks...")
add_exhibit_numbers(doc)

print("Adding Exhibit Index...")
add_exhibit_index(doc)

print(f"Saving: {OUTPUT_FILE}")
doc.save(OUTPUT_FILE)

print("Done! REVIEWED version generated.")
print(f"\nCorrections applied:")
print("  ✅ Page numbering in footer (Page X of Y Résumé Name – EB-1)")
print("  ✅ Exhibit numbers (1-17) added to evidence blocks")
print("  ✅ Exhibit Index table added at end of document")
print("  ✅ Conclusory language softened (distinguished, extraordinary, preeminent)")
print("  ✅ CMSP repetition reduced (later mentions reference Criterion 1)")
print("  ✅ Addendum condensed (cross-references criteria instead of repeating)")
print("  ✅ 'engineering leadership' formula varied")
print("\nStill requires manual action:")
print("  ⚠️  Fill [VERIFICAR] placeholders with actual client data")
print("  ⚠️  Verify Executive Orders 14277/14278 exist")
print("  ⚠️  Verify BLS data ($165,370 / $172,290) and add footnote")
print("  ⚠️  Verify Gobernarte data (127 initiatives, 12 nations)")
print("  ⚠️  Resolve EB-2-NIW vs EB-1A in clients.json")
print("  ⚠️  Run insert_thumbnails.py after all fixes")
