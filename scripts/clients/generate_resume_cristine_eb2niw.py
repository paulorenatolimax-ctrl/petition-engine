#!/usr/bin/env python3
"""
Résumé EB-2 NIW — Cristine Correa
Gerado pelo Petition Engine v2.0
100% Português Brasileiro | Garamond | Navy/Teal palette
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ============================================================
# CONSTANTS
# ============================================================
NAVY = RGBColor(0x2D, 0x3E, 0x50)
TEAL = RGBColor(0x34, 0x98, 0xA2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
BORDER_GRAY = RGBColor(0xCC, 0xCC, 0xCC)

FONT_NAME = "Garamond"
FONT_HEADER_NAME = Pt(20)
FONT_SECTION = Pt(11)
FONT_SUBSECTION = Pt(10)
FONT_BODY = Pt(10.5)
FONT_SMALL = Pt(9.5)
FONT_CONTACT = Pt(9)
FONT_COMPACT = Pt(9)
FONT_TABLE_SMALL = Pt(9.5)

META_WIDTH = 5760  # DXA
THUMB_WIDTH = 4320  # DXA
COMPACT_META = 6480
COMPACT_THUMB = 3600

OUTPUT_DIR = "/Users/paulo1844/Documents/_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2026/CRISTINE CORREA/_Forjado por Petition Engine/"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "resume_eb2_niw_Cristine_Correa.docx")

BENEFICIARY_NAME = "CRISTINE CORREA"
BENEFICIARY_EMAIL = "cristine.correa@email.com"
SOC_PRIMARY = "11-3121"
SOC_SECONDARY = "11-3131"


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC", sz=4):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def set_cell_width(cell, width_dxa):
    """Set cell width in DXA."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    tcPr.append(tcW)


def set_cell_margins(cell, top=0, bottom=0, left=100, right=100):
    """Set cell margins in DXA."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def set_row_height(row, height_pt):
    """Set row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def add_run(paragraph, text, font_name=FONT_NAME, size=FONT_BODY, color=BLACK,
            bold=False, italic=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}" w:cs="{font_name}"/>'))
    return run


def add_body_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       space_after=Pt(6), space_before=Pt(0)):
    """Add a body text paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = Pt(14)
    add_run(p, text, size=FONT_BODY, color=BLACK)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)
    if level == 0:
        p.paragraph_format.left_indent = Inches(0.3)
    else:
        p.paragraph_format.left_indent = Inches(0.6)

    if bold_prefix:
        add_run(p, f"• {bold_prefix}", size=FONT_BODY if level == 0 else Pt(10),
                color=BLACK, bold=True)
        add_run(p, f" {text}", size=FONT_BODY if level == 0 else Pt(10),
                color=BLACK if level == 0 else DARK_GRAY)
    else:
        size = FONT_BODY if level == 0 else Pt(10)
        color = BLACK if level == 0 else DARK_GRAY
        add_run(p, f"• {text}", size=size, color=color)
    return p


def add_navy_section_header(doc, title):
    """Add a navy section header bar (full width)."""
    doc.add_page_break()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "2D3E50")
    set_cell_margins(cell, top=60, bottom=60, left=200, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title.upper(), size=FONT_SECTION, color=WHITE, bold=True)
    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    # Full width
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_navy_section_header_no_break(doc, title):
    """Navy section header without page break (for first section)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "2D3E50")
    set_cell_margins(cell, top=60, bottom=60, left=200, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title.upper(), size=FONT_SECTION, color=WHITE, bold=True)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_teal_sub_header(doc, title):
    """Add a teal sub-header bar."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "3498A2")
    set_cell_margins(cell, top=40, bottom=40, left=200, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title, size=FONT_SUBSECTION, color=WHITE, bold=True)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_evidence_block(doc, metadata_lines, impact_text, compact=False):
    """
    Add an evidence block: 2-col table.
    Left = metadata + impact INSIDE.
    Right = thumbnail placeholder.
    """
    mw = COMPACT_META if compact else META_WIDTH
    tw = COMPACT_THUMB if compact else THUMB_WIDTH
    body_sz = FONT_COMPACT if compact else FONT_SMALL
    impact_sz = FONT_COMPACT if compact else FONT_SMALL

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    set_cell_width(left_cell, mw)
    set_cell_width(right_cell, tw)
    set_cell_borders(left_cell, "CCCCCC", 4)
    set_cell_borders(right_cell, "CCCCCC", 4)
    set_cell_margins(left_cell, top=80, bottom=80, left=120, right=120)
    set_cell_margins(right_cell, top=80, bottom=80, left=80, right=80)

    # Left cell: metadata
    first = True
    for label, value in metadata_lines:
        if first:
            p = left_cell.paragraphs[0]
            first = False
        else:
            p = left_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        add_run(p, f"{label}: ", size=Pt(10), color=BLACK, bold=True)
        add_run(p, value, size=Pt(10), color=BLACK)

    # Impact inside left cell
    if impact_text:
        p_label = left_cell.add_paragraph()
        p_label.paragraph_format.space_before = Pt(6)
        p_label.paragraph_format.space_after = Pt(2)
        add_run(p_label, "Descrição e Impacto", size=FONT_SMALL, color=NAVY, bold=True)

        p_impact = left_cell.add_paragraph()
        p_impact.paragraph_format.space_after = Pt(4)
        p_impact.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p_impact, impact_text, size=impact_sz, color=DARK_GRAY, italic=True)

    # Right cell: thumbnail placeholder
    p_thumb = right_cell.paragraphs[0]
    p_thumb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_run(p_thumb, "[Thumbnail]", size=FONT_CONTACT, color=MED_GRAY, italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_company_box(doc, company, role, period, location, sector):
    """Add a company experience box with light gray background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")
    set_cell_borders(cell, "CCCCCC", 4)
    set_cell_margins(cell, top=80, bottom=80, left=150, right=150)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_run(p, company, size=FONT_BODY, color=NAVY, bold=True)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    add_run(p2, role, size=FONT_BODY, color=BLACK, bold=True)
    add_run(p2, f"  |  {period}", size=FONT_SMALL, color=DARK_GRAY)

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    add_run(p3, f"{location}  |  {sector}", size=FONT_SMALL, color=MED_GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


# ============================================================
# DOCUMENT SETUP
# ============================================================

def setup_document():
    """Create and configure the base document."""
    doc = Document()

    # Page setup: US Letter, narrow margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.header_distance = Inches(0)

    return doc


def build_header(doc):
    """Build the 3-row header table."""
    header_section = doc.sections[0].header
    header_section.is_linked_to_previous = False

    table = header_section.add_table(rows=3, cols=2, width=Inches(7.2))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)
    remove_table_borders(table)

    # Row 0: Navy — Name
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "2D3E50")
    set_cell_margins(table.cell(0, 0), top=120, bottom=40, left=200, right=100)
    set_cell_margins(table.cell(0, 1), top=120, bottom=40, left=100, right=200)

    p_name = table.cell(0, 0).paragraphs[0]
    add_run(p_name, BENEFICIARY_NAME, size=FONT_HEADER_NAME, color=WHITE, bold=True)

    # Row 1: Navy — RESUME + SOC + EB-2 NIW | Email
    for cell in table.rows[1].cells:
        set_cell_shading(cell, "2D3E50")
    set_cell_margins(table.cell(1, 0), top=20, bottom=60, left=200, right=100)
    set_cell_margins(table.cell(1, 1), top=20, bottom=60, left=100, right=200)

    p_resume = table.cell(1, 0).paragraphs[0]
    add_run(p_resume, "RÉSUMÉ", size=FONT_SECTION, color=WHITE, bold=True)
    add_run(p_resume, f"  |  SOC/O*Net {SOC_PRIMARY}  |  EB-2 NIW",
            size=FONT_CONTACT, color=WHITE)

    p_email = table.cell(1, 1).paragraphs[0]
    p_email.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_email, BENEFICIARY_EMAIL, size=FONT_CONTACT, color=WHITE)

    # Row 2: Teal accent
    for cell in table.rows[2].cells:
        set_cell_shading(cell, "3498A2")
    set_row_height(table.rows[2], 4)

    # Remove default paragraph from header
    for p in header_section.paragraphs:
        if p.text == "":
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)


def build_footer(doc):
    """Build the navy footer with Page X of Y."""
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False

    table = footer.add_table(rows=1, cols=1, width=Inches(7.2))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "2D3E50")
    set_cell_margins(cell, top=40, bottom=40, left=200, right=200)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)
    remove_table_borders(table)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "Page " text
    run1 = p.add_run("Page ")
    run1.font.name = FONT_NAME
    run1.font.size = FONT_CONTACT
    run1.font.color.rgb = WHITE

    # PAGE field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r1 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>')
    r1.append(fldChar1)
    p._p.append(r1)

    instrText = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>')
    p._p.append(instrText)

    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r2 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>')
    r2.append(fldChar2)
    p._p.append(r2)

    run_of = p.add_run(" of ")
    run_of.font.name = FONT_NAME
    run_of.font.size = FONT_CONTACT
    run_of.font.color.rgb = WHITE

    # NUMPAGES field
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r3 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>')
    r3.append(fldChar3)
    p._p.append(r3)

    instrText2 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr><w:instrText xml:space="preserve"> NUMPAGES </w:instrText></w:r>')
    p._p.append(instrText2)

    fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r4 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/><w:sz w:val="18"/><w:color w:val="FFFFFF"/></w:rPr></w:r>')
    r4.append(fldChar4)
    p._p.append(r4)

    for p in footer.paragraphs:
        if p.text == "":
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)


# ============================================================
# SECTION BUILDERS
# ============================================================

def build_synthesis(doc):
    """Seção 1: Síntese Profissional / Executive Summary."""
    add_navy_section_header_no_break(doc, "SÍNTESE PROFISSIONAL")

    add_body_paragraph(doc,
        "Cristine Correa é psicóloga organizacional brasileira com mais de uma década de experiência "
        "progressiva em gestão estratégica de pessoas, avaliação comportamental e inteligência de força "
        "de trabalho. Ao longo de sua trajetória profissional, atuou em seis organizações de diferentes "
        "portes e setores — desde operações de contact center com milhares de colaboradores até fintechs "
        "de alta velocidade autorizadas pelo Banco Central do Brasil —, acumulando expertise multissetorial "
        "que abrange recrutamento e seleção estratégicos, diagnóstico de clima organizacional, programas de "
        "inclusão e diversidade, e desenvolvimento de lideranças em ambientes de alta complexidade."
    )

    add_body_paragraph(doc,
        "Sua formação acadêmica compreende bacharelado em Psicologia com ênfase organizacional pela "
        "Associação Catarinense de Ensino (4.212 horas), pós-graduação em Avaliação Psicológica pela "
        "mesma instituição, e MBA Executivo em Gestão de Pessoas, Desenvolvimento Gerencial e Coaching "
        "pela FACEL. Essa tríplice formação — combinando fundamentos clínicos da psicologia, rigor "
        "psicométrico e visão gerencial — confere à beneficiária uma capacidade diferenciada de "
        "desenvolver e aplicar instrumentos de avaliação comportamental com validade científica e "
        "aplicabilidade organizacional imediata, qualificando-a como portadora de grau avançado e "
        "habilidade excepcional em sua área de atuação."
    )

    add_body_paragraph(doc,
        "Na UNISOCIESC, integrante do Ecossistema Ânima Educação — uma das maiores organizações de "
        "ensino superior do Brasil com mais de 66 anos de história —, Cristine ascendeu de Analista de "
        "RH a Business Partner, tornando-se referência interna para diretores e gestores em decisões "
        "estratégicas de pessoas. Nessa posição, liderou a implementação da metodologia Predictive Index "
        "(PI) para seleções estratégicas, conduziu pesquisas de clima organizacional Great Place to Work "
        "(GPTW) com planos de ação setoriais, e coordenou o Projeto Incluir durante três anos — um "
        "programa estruturado de inclusão de pessoas com deficiência que envolveu contratação, eventos "
        "de sensibilização e a criação do Corredor Sensorial, impactando diretamente a cultura "
        "organizacional da instituição."
    )

    add_body_paragraph(doc,
        "Cristine detém certificação como Analista Predictive Index (PI) pela PRAENDEX Brasil, "
        "habilitando-a a aplicar e interpretar instrumentos de avaliação comportamental e cognitiva "
        "com fundamentação em pesquisa comportamental validada. É membro ativa da Society for Human "
        "Resource Management (SHRM) e da American Management Association (AMA), mantendo-se atualizada "
        "com as melhores práticas de gestão de pessoas no contexto norte-americano. Possui registro "
        "profissional ativo no Conselho Regional de Psicologia (CRP-12/11050), demonstrando conformidade "
        "regulatória plena em sua jurisdição de origem."
    )

    add_body_paragraph(doc,
        "Sua atuação mais recente na ASAAS — fintech autorizada pelo Banco Central do Brasil para "
        "operações de pagamento — evidenciou capacidade de adaptar metodologias de avaliação "
        "comportamental a ambientes de tecnologia de alta velocidade. Nessa posição, implementou o "
        "sistema Gupy para recrutamento inteligente, criou indicadores de NPS para candidatos "
        "aprovados e reprovados, desenvolveu programa de integração estruturado e manteve interface "
        "diária com líderes de área para alinhamento de perfil comportamental e técnico."
    )

    add_body_paragraph(doc,
        "Desde sua chegada aos Estados Unidos em julho de 2022, Cristine tem dedicado-se ao estudo "
        "aprofundado do mercado de trabalho norte-americano, com foco nas necessidades específicas de "
        "pequenas e médias empresas na região da Flórida Central. Seus dois empreendimentos propostos "
        "— o Workforce Behavioral Intelligence Institute e o Organizational Climate & Strategic "
        "Retention Lab — representam a convergência de sua expertise em avaliação comportamental com "
        "PI, diagnóstico de clima organizacional e estratégias de retenção, direcionados a setores "
        "com taxas críticas de rotatividade como saúde (22,7%), hospitalidade (75,2%) e manufatura "
        "(24-32%), em alinhamento direto com prioridades federais de desenvolvimento da força de "
        "trabalho e interesse nacional dos Estados Unidos."
    )


def build_timeline(doc):
    """Seção 2: Histórico Profissional (Gantt)."""
    add_navy_section_header(doc, "HISTÓRICO PROFISSIONAL")

    # Gantt table: years 2011-2026, companies as rows
    years = list(range(2011, 2027))
    companies = [
        ("CONTAX Mobitel S.A.", 2011, 2014),
        ("Precisão Global", 2014, 2015),
        ("CDL Joinville", 2015, 2015),
        ("Walmart Brasil", 2016, 2017),
        ("UNISOCIESC / Ânima", 2017, 2021),
        ("ASAAS", 2021, 2022),
        ("Atividades nos EUA", 2022, 2026),
    ]

    n_cols = len(years) + 1  # company name + years
    table = doc.add_table(rows=len(companies) + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)

    # Header row
    header_row = table.rows[0]
    for i, cell in enumerate(header_row.cells):
        set_cell_shading(cell, "2D3E50")
        set_cell_borders(cell, "2D3E50", 2)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            add_run(p, "Organização", size=FONT_CONTACT, color=WHITE, bold=True)
        else:
            add_run(p, str(years[i-1]), size=FONT_CONTACT, color=WHITE, bold=True)

    # Data rows
    for row_idx, (name, start, end) in enumerate(companies):
        row = table.rows[row_idx + 1]
        for col_idx, cell in enumerate(row.cells):
            set_cell_borders(cell, "CCCCCC", 2)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if col_idx == 0:
                add_run(p, name, size=FONT_CONTACT, color=BLACK, bold=True)
                set_cell_margins(cell, left=80, right=40)
            else:
                year = years[col_idx - 1]
                if start <= year <= end:
                    set_cell_shading(cell, "3498A2")
                else:
                    set_cell_shading(cell, "FFFFFF")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def build_experience(doc):
    """Seção 3: Experiência Profissional Detalhada."""
    add_navy_section_header(doc, "EXPERIÊNCIA PROFISSIONAL DETALHADA")

    # --- ASAAS ---
    add_company_box(doc, "ASAAS", "Analista de Aquisição de Talentos",
                    "Agosto/2021 — Maio/2022 (9 meses)", "Joinville, SC, Brasil",
                    "Fintech — Pagamentos Digitais (autorizada pelo Banco Central)")
    add_bullet(doc, "Responsável pelo recrutamento e seleção estratégicos de ponta a ponta para posições administrativas, comerciais e operacionais, garantindo alinhamento entre perfil comportamental e requisitos técnicos de cada vaga em ambiente de tecnologia de alta velocidade.")
    add_bullet(doc, "Implementou a plataforma Gupy para recrutamento inteligente, integrando análise comportamental ao processo seletivo e aumentando a eficiência operacional do pipeline de candidatos com rastreabilidade completa do funil de contratação.")
    add_bullet(doc, "Criou e operacionalizou sistema de NPS (Net Promoter Score) para candidatos aprovados e reprovados, estabelecendo métricas quantitativas de experiência do candidato e gerando dados para melhoria contínua do processo seletivo.")
    add_bullet(doc, "Desenvolveu programa de integração estruturado (Welcome Program) para novos colaboradores, reduzindo o tempo de adaptação e fortalecendo a retenção nos primeiros 90 dias de contratação.")
    add_bullet(doc, "Manteve interface diária com líderes de área via reuniões e Slack, operando programa de indicação interna (referral) e alinhando perfis comportamentais às necessidades específicas de cada equipe.")
    add_bullet(doc, "Produziu relatórios mensais e trimestrais de indicadores de RH, incluindo tempo médio de contratação, taxa de conversão por etapa e análise de fontes de recrutamento, fornecendo inteligência de dados para decisões de força de trabalho.")

    # --- UNISOCIESC ---
    add_company_box(doc, "UNISOCIESC / Ecossistema Ânima Educação",
                    "Analista de RH / Business Partner",
                    "Setembro/2017 — Agosto/2021 (3 anos e 11 meses)", "Joinville, SC, Brasil",
                    "Ensino Superior — Centro Universitário (66+ anos de história)")
    add_bullet(doc, "Evoluiu de Analista de RH Júnior a Business Partner estratégica, tornando-se interlocutora direta de diretores e coordenadores para todas as decisões relacionadas a pessoas, desde contratações estratégicas até reestruturações durante a pandemia de COVID-19.")
    add_bullet(doc, "Conduziu o processo completo de recrutamento e seleção para posições administrativas e docentes, aplicando a metodologia Predictive Index (PI) em seleções estratégicas para garantir compatibilidade comportamental e cognitiva entre candidatos e requisitos organizacionais.")
    add_bullet(doc, "Liderou pesquisas de clima organizacional Great Place to Work (GPTW), analisando resultados por setor e desenvolvendo planos de ação específicos para áreas com indicadores abaixo da meta, contribuindo diretamente para a melhoria do ambiente de trabalho institucional.")
    add_bullet(doc, "Coordenou o Projeto Incluir durante três anos consecutivos (outubro/2017 a setembro/2020), um programa estruturado de inclusão de pessoas com deficiência que abrangeu recrutamento especializado, eventos de sensibilização para toda a comunidade acadêmica, e a concepção e implementação do Corredor Sensorial — uma instalação experiencial projetada para promover empatia e compreensão das diferentes formas de deficiência.")
    add_bullet(doc, "Organizou e facilitou o Fórum de Líderes mensalmente, conduzindo workshops de desenvolvimento de liderança, sessões de feedback estruturado e monitoramento de indicadores de desempenho das equipes, fortalecendo a cultura de liderança transformacional na instituição.")
    add_bullet(doc, "Gerenciou processos de desligamento com sensibilidade e conformidade legal, incluindo a transição para desligamentos por videoconferência durante a pandemia, mantendo a dignidade do processo e minimizando riscos trabalhistas em um período de reestruturação organizacional significativa.")
    add_bullet(doc, "Administrou o headcount institucional durante reorganizações pandêmicas, equilibrando necessidades operacionais com sustentabilidade financeira e preservação do capital humano estratégico.")

    # --- Walmart ---
    add_company_box(doc, "WMS Supermercados do Brasil (Walmart Brasil)",
                    "Analista de Recrutamento e Seleção Júnior",
                    "Fevereiro/2016 — Setembro/2017 (1 ano e 7 meses)", "Joinville, SC, Brasil",
                    "Varejo Multinacional — Grande Porte")
    add_bullet(doc, "Responsável pelo recrutamento e seleção de ponta a ponta para quatro unidades Walmart em Joinville, conduzindo processos seletivos diários para posições operacionais e administrativas com volume elevado de vagas simultâneas.")
    add_bullet(doc, "Estabeleceu relacionamento consultivo presencial com gerentes de loja, realizando visitas semanais para alinhamento de perfis e acompanhamento de indicadores de retenção, resultando em redução significativa da rotatividade nas unidades atendidas.")
    add_bullet(doc, "Alcançou tempo médio de contratação de dois dias para posições operacionais, combinando eficiência processual com assertividade na avaliação de candidatos, mantendo um banco de talentos estruturado e atualizado.")
    add_bullet(doc, "Produziu relatórios semanais de acompanhamento de recrutamento e seleção, fornecendo dados quantitativos sobre taxa de preenchimento, tempo de contratação e qualidade das contratações para a gestão regional.")

    # --- CDL ---
    add_company_box(doc, "CDL Joinville (Câmara de Dirigentes Lojistas)",
                    "Assistente de Recrutamento e Seleção",
                    "Abril/2015 — Novembro/2015 (7 meses)", "Joinville, SC, Brasil",
                    "Associação Comercial — Intermediação de Mão de Obra")
    add_bullet(doc, "Conduziu processos seletivos completos incluindo triagem curricular, dinâmicas de grupo, aplicação de testes psicológicos e elaboração de laudos técnicos, intermediando a recolocação profissional de trabalhadores da região para empresas associadas à CDL.")
    add_bullet(doc, "Executou rotinas administrativas de recursos humanos e contribuiu para o aumento da taxa de sucesso de reemprego dos candidatos atendidos pelo programa de intermediação de mão de obra da instituição.")

    # --- Precisão Global ---
    add_company_box(doc, "Precisão Global de Cobranças EIRELI",
                    "Analista de Recrutamento e Seleção",
                    "Agosto/2014 — Fevereiro/2015 (6 meses)", "Joinville, SC, Brasil",
                    "Serviços Financeiros — Contact Center de Cobrança")
    add_bullet(doc, "Gerenciou o ciclo completo de recrutamento e seleção para operações de contact center e gestão, conduzindo processos coletivos e individuais para posições em múltiplas localidades (São Paulo, Curitiba, Blumenau, Rio de Janeiro).")
    add_bullet(doc, "Supervisionou aprendizes e desenvolveu campanhas de indicação interna, otimizando o pipeline de candidatos e reduzindo custos de aquisição de talentos para a organização.")

    # --- CONTAX ---
    add_company_box(doc, "CONTAX Mobitel S.A.",
                    "Assistente de Recursos Humanos",
                    "Setembro/2011 — Julho/2014 (2 anos e 10 meses)", "Joinville, SC, Brasil",
                    "BPO — Contact Center (operação de larga escala)")
    add_bullet(doc, "Coordenou operações de recrutamento e seleção em ambiente de alto volume, liderando a equipe de suporte ao processo seletivo e conduzindo testes, dinâmicas de grupo e entrevistas individuais para posições operacionais e administrativas.")
    add_bullet(doc, "Gerenciou indicadores de recrutamento e seleção, contribuindo para a expansão das operações para São Paulo, Curitiba, Blumenau e Rio de Janeiro, demonstrando capacidade de escalar processos de atração de talentos em múltiplas geografias.")
    add_bullet(doc, "Participou da reestruturação da equipe de suporte, otimizando fluxos de trabalho e melhorando a eficiência dos processos de seleção em uma operação com milhares de colaboradores.")


def build_contributions(doc):
    """Seção 4: Contribuições Técnicas e Profissionais (por TEMA)."""
    add_navy_section_header(doc, "CONTRIBUIÇÕES TÉCNICAS E PROFISSIONAIS")

    # --- Theme 1: Avaliação Comportamental ---
    add_teal_sub_header(doc, "Avaliação Comportamental e Metodologia Predictive Index")

    add_body_paragraph(doc,
        "A expertise de Cristine Correa em avaliação comportamental representa o núcleo de sua "
        "contribuição técnica diferenciada. Com certificação como Analista Predictive Index (PI) "
        "obtida em abril de 2020 pela PRAENDEX Brasil, Cristine domina a aplicação e interpretação "
        "de instrumentos de avaliação comportamental e cognitiva fundamentados em décadas de pesquisa "
        "comportamental validada. O Predictive Index é um sistema cientificamente validado que mede "
        "quatro fatores comportamentais primários — dominância, extroversão, paciência e formalidade —, "
        "permitindo predição de adequação comportamental entre candidatos e funções com precisão "
        "significativamente superior a métodos tradicionais de seleção."
    )

    add_evidence_block(doc, [
        ("Documento", "Certificação Predictive Index (PI) — Analista"),
        ("Emissor", "PRAENDEX Brasil / The Predictive Index"),
        ("Data", "20 de abril de 2020"),
        ("Contrato", "Ânima Holding S.A."),
    ],
        "A certificação PI capacita Cristine a definir requisitos comportamentais e cognitivos "
        "de posições utilizando métricas de Behavioral Research e Cognitive Testing, aplicar "
        "avaliações individuais e em grupo, e interpretar resultados para orientar decisões "
        "estratégicas de contratação, desenvolvimento e retenção de talentos. Na UNISOCIESC, "
        "essa metodologia foi aplicada em seleções estratégicas para posições de liderança e "
        "docência, elevando a assertividade do processo seletivo e reduzindo custos associados "
        "a contratações inadequadas."
    )

    add_body_paragraph(doc,
        "A aplicação do Predictive Index na UNISOCIESC representou uma mudança paradigmática "
        "no processo de seleção da instituição: pela primeira vez, decisões de contratação para "
        "posições estratégicas foram fundamentadas em dados comportamentais objetivos, não apenas "
        "em impressões subjetivas de entrevistas. Cristine integrou a análise PI ao fluxo de "
        "seleção existente, criando um protocolo que combina avaliação comportamental quantitativa "
        "com entrevistas por competências, dinâmicas de grupo e validação com gestores — um modelo "
        "que se provou eficaz na redução da rotatividade precoce e no aumento da satisfação dos "
        "gestores com as contratações realizadas."
    )

    # --- Theme 2: Inclusão e Diversidade ---
    add_teal_sub_header(doc, "Programas de Inclusão e Diversidade Organizacional")

    add_body_paragraph(doc,
        "O Projeto Incluir, coordenado por Cristine durante três anos consecutivos na UNISOCIESC "
        "(outubro/2017 a setembro/2020), constitui uma contribuição significativa na área de "
        "inclusão organizacional de pessoas com deficiência. O programa foi concebido e executado "
        "como uma iniciativa multidimensional que integrou recrutamento especializado, eventos de "
        "sensibilização para toda a comunidade acadêmica e administrativa, e a criação do Corredor "
        "Sensorial — uma instalação experiencial projetada para permitir que colaboradores e "
        "estudantes vivenciassem temporariamente as limitações sensoriais enfrentadas por pessoas "
        "com diferentes tipos de deficiência."
    )

    add_evidence_block(doc, [
        ("Projeto", "Projeto Incluir — Programa de Inclusão de PcD"),
        ("Instituição", "UNISOCIESC / Ecossistema Ânima Educação"),
        ("Período", "Outubro/2017 — Setembro/2020 (3 anos)"),
        ("Escopo", "Recrutamento, sensibilização e acessibilidade"),
    ],
        "O Projeto Incluir gerou impacto mensurável na cultura organizacional da UNISOCIESC, "
        "aumentando a representatividade de pessoas com deficiência no quadro funcional e "
        "promovendo uma transformação cultural documentada no modo como a comunidade acadêmica "
        "interage com a diversidade funcional. O Corredor Sensorial, em particular, tornou-se "
        "referência interna para eventos de conscientização e foi apresentado em datas "
        "comemorativas como o Dia Internacional da Pessoa com Deficiência, demonstrando a "
        "capacidade de Cristine de desenvolver iniciativas de impacto social sustentável em "
        "contextos organizacionais."
    )

    # --- Theme 3: Clima Organizacional ---
    add_teal_sub_header(doc, "Diagnóstico de Clima Organizacional e Engajamento")

    add_body_paragraph(doc,
        "Na UNISOCIESC, Cristine liderou a implementação e análise de pesquisas de clima "
        "organizacional utilizando a metodologia Great Place to Work (GPTW), uma das mais "
        "reconhecidas globalmente para mensuração de ambiente de trabalho. Seu papel envolveu "
        "não apenas a aplicação do instrumento, mas a análise granular dos resultados por setor, "
        "a identificação de padrões de engajamento e desengajamento, e o desenvolvimento de "
        "planos de ação específicos para áreas com indicadores abaixo da meta institucional. "
        "Essa experiência em diagnóstico de clima organizacional com ferramentas de validade "
        "reconhecida internacionalmente constitui a base metodológica para seu empreendimento "
        "proposto de laboratório de clima e retenção estratégica nos Estados Unidos."
    )

    add_body_paragraph(doc,
        "A competência em diagnóstico de clima de Cristine vai além da aplicação de "
        "questionários: envolve a interpretação de dados multidimensionais, a correlação "
        "entre indicadores de clima e métricas operacionais (rotatividade, absenteísmo, "
        "produtividade), e a tradução de insights quantitativos em intervenções comportamentais "
        "específicas. Na UNISOCIESC, essa abordagem permitiu identificar fatores de "
        "desengajamento específicos por departamento e implementar soluções direcionadas, "
        "resultando em melhoria documentada nos indicadores de satisfação nas áreas "
        "trabalhadas."
    )

    # --- Theme 4: Desenvolvimento de Liderança ---
    add_teal_sub_header(doc, "Desenvolvimento de Liderança e Gestão de Desempenho")

    add_body_paragraph(doc,
        "O Fórum de Líderes, facilitado mensalmente por Cristine na UNISOCIESC, representou "
        "um investimento contínuo no desenvolvimento de competências de liderança da instituição. "
        "Através de workshops estruturados, sessões de feedback 360° e monitoramento de "
        "indicadores de desempenho de equipes, Cristine construiu um programa de desenvolvimento "
        "de liderança que fortaleceu a cultura de gestão participativa e accountability. Os "
        "líderes participantes demonstraram evolução mensurável em competências como comunicação "
        "assertiva, gestão de conflitos e tomada de decisão baseada em dados — habilidades "
        "diretamente transferíveis para o contexto de desenvolvimento de lideranças em "
        "organizações norte-americanas."
    )

    # --- Theme 5: Inovação em Aquisição de Talentos ---
    add_teal_sub_header(doc, "Inovação em Aquisição de Talentos e Experiência do Candidato")

    add_body_paragraph(doc,
        "Na ASAAS, Cristine demonstrou capacidade de adaptar metodologias de avaliação "
        "comportamental ao contexto de empresas de tecnologia em crescimento acelerado. A "
        "implementação da plataforma Gupy — uma das principais soluções de recrutamento "
        "inteligente do Brasil — integrou análise comportamental automatizada ao processo "
        "seletivo, enquanto a criação do sistema de NPS para candidatos estabeleceu um "
        "protocolo de medição de experiência do candidato inédito na organização. O programa "
        "de integração estruturado (Welcome Program) complementou a estratégia de aquisição "
        "de talentos com foco na retenção desde o primeiro dia, abordando um dos principais "
        "desafios de empresas de tecnologia em expansão rápida."
    )

    add_body_paragraph(doc,
        "A contribuição de Cristine na ASAAS estendeu-se à criação de indicadores "
        "quantitativos de eficácia do recrutamento — incluindo tempo médio de contratação, "
        "taxa de conversão por etapa do funil, análise de fontes de recrutamento e NPS de "
        "candidatos —, transformando uma operação tradicionalmente qualitativa em um processo "
        "orientado por dados. Essa abordagem data-driven para aquisição de talentos é "
        "particularmente relevante para o mercado norte-americano, onde métricas de "
        "recrutamento são fundamentais para decisões de investimento em capital humano."
    )

    # --- Theme 6: Gestão em Cenários de Crise ---
    add_teal_sub_header(doc, "Gestão de Pessoas em Cenários de Crise e Transformação")

    add_body_paragraph(doc,
        "Durante a pandemia de COVID-19, Cristine demonstrou resiliência e adaptabilidade ao "
        "gerenciar processos sensíveis de desligamento por videoconferência na UNISOCIESC, "
        "mantendo a dignidade dos colaboradores desligados e a conformidade legal em um contexto "
        "sem precedentes. Paralelamente, administrou o headcount institucional durante "
        "reestruturações significativas, equilibrando necessidades operacionais com "
        "sustentabilidade financeira e preservação de talentos estratégicos. Essa experiência "
        "em gestão de pessoas em cenários de crise — envolvendo comunicação sensível, "
        "conformidade trabalhista e tomada de decisão sob pressão — evidencia uma maturidade "
        "profissional essencial para a atuação em ambientes organizacionais complexos nos "
        "Estados Unidos, onde a capacidade de navegar crises de força de trabalho é um "
        "diferencial competitivo crítico."
    )


def build_publications(doc):
    """Seção 5: Publicações e Artigos."""
    add_navy_section_header(doc, "PUBLICAÇÕES E ARTIGOS")

    add_teal_sub_header(doc, "Artigos e Publicações Profissionais")

    add_evidence_block(doc, [
        ("Título", "The Differences Between HR in Brazil and the United States"),
        ("Plataforma", "Medium"),
        ("Idioma", "Inglês"),
        ("Tópico", "Análise comparativa de práticas de RH entre Brasil e EUA"),
    ],
        "Artigo que analisa as diferenças estruturais e culturais entre a gestão de recursos "
        "humanos no Brasil e nos Estados Unidos, abordando legislação trabalhista, práticas de "
        "recrutamento, benefícios, e cultura organizacional. A publicação demonstra a capacidade "
        "de Cristine de analisar criticamente sistemas de gestão de pessoas em contextos "
        "internacionais e comunicar insights relevantes para profissionais de RH que atuam em "
        "ambientes multiculturais. A perspectiva comparativa apresentada no artigo reflete "
        "diretamente a expertise bicultural que fundamenta seus empreendimentos propostos "
        "nos Estados Unidos.",
        compact=True
    )

    add_teal_sub_header(doc, "Trabalhos Acadêmicos")

    add_evidence_block(doc, [
        ("Título", "A Influência da Psicopatia no Ambiente Corporativo"),
        ("Tipo", "Monografia — MBA Executivo"),
        ("Instituição", "FACEL — Curitiba, PR"),
        ("Conclusão", "Maio de 2017"),
    ],
        "Pesquisa que investigou a manifestação de traços psicopáticos em ambientes corporativos "
        "e seu impacto em dinâmicas organizacionais, processos decisórios e relações interpessoais "
        "no trabalho. O estudo demonstra o interesse acadêmico de Cristine na interseção entre "
        "psicologia clínica e comportamento organizacional, fornecendo base teórica para sua "
        "abordagem de avaliação comportamental aplicada ao contexto empresarial.",
        compact=True
    )

    add_evidence_block(doc, [
        ("Título", "O Trabalho do Profissional de Psicologia com Mulheres em Situação de Violência"),
        ("Tipo", "Trabalho de Conclusão — Pós-Graduação"),
        ("Instituição", "ACE / Faculdade Guilherme Guimbala — Joinville, SC"),
        ("Conclusão", "Outubro de 2013"),
    ],
        "Pesquisa sobre a atuação do psicólogo no atendimento a mulheres em situação de violência, "
        "abordando protocolos de intervenção, redes de apoio e políticas públicas. Embora focado "
        "na psicologia clínica e social, o trabalho evidencia a sensibilidade de Cristine para "
        "questões de vulnerabilidade e inclusão — competência que se manifestou posteriormente "
        "em sua coordenação do Projeto Incluir e em sua abordagem humanizada de gestão de pessoas.",
        compact=True
    )


def build_education(doc):
    """Seção 6: Formação Acadêmica."""
    add_navy_section_header(doc, "FORMAÇÃO ACADÊMICA")

    add_evidence_block(doc, [
        ("Grau", "Bacharelado em Psicologia"),
        ("Instituição", "Associação Catarinense de Ensino — Faculdade Guilherme Guimbala (ACE)"),
        ("Localização", "Joinville, Santa Catarina, Brasil"),
        ("Conclusão", "3 de março de 2012"),
        ("Carga Horária", "4.212 horas"),
        ("Ênfase", "Psicologia Organizacional"),
    ],
        "Formação abrangente em psicologia com ênfase organizacional, incluindo estágios "
        "supervisionados em Psicologia Educacional e Comunitária, Psicologia do Trabalho, e "
        "Psicologia Clínica e Saúde Coletiva. A carga horária de 4.212 horas — significativamente "
        "acima do mínimo exigido — reflete a profundidade da formação recebida. A equivalência "
        "deste diploma foi avaliada pela GEO Credential Services e aceita pelo USCIS como "
        "equivalente ao bacharelado norte-americano, atendendo ao requisito de grau avançado "
        "da categoria EB-2."
    )

    add_evidence_block(doc, [
        ("Grau", "Pós-Graduação Lato Sensu em Avaliação Psicológica"),
        ("Instituição", "ACE — Faculdade Guilherme Guimbala"),
        ("Localização", "Joinville, Santa Catarina, Brasil"),
        ("Conclusão", "27 de outubro de 2013"),
        ("Carga Horária", "368 horas"),
    ],
        "Especialização em avaliação psicológica que aprofundou competências em psicometria, "
        "construção e validação de instrumentos de medição psicológica, e interpretação de "
        "resultados de testes padronizados. Esta formação constitui o alicerce técnico para "
        "a certificação Predictive Index obtida posteriormente e para a capacidade de Cristine "
        "de desenvolver e aplicar ferramentas de avaliação comportamental com rigor científico "
        "em contextos organizacionais."
    )

    add_evidence_block(doc, [
        ("Grau", "MBA Executivo em Gestão de Pessoas, Desenvolvimento Gerencial e Coaching"),
        ("Instituição", "FACEL"),
        ("Localização", "Curitiba, Paraná, Brasil"),
        ("Conclusão", "31 de maio de 2017"),
        ("Carga Horária", "360 horas"),
    ],
        "Formação executiva que complementou a base psicológica com competências em gestão "
        "estratégica de pessoas, desenvolvimento organizacional e coaching executivo. O MBA "
        "proporcionou ferramentas de gestão aplicadas que Cristine integrou à sua prática "
        "de Business Partner na UNISOCIESC, combinando visão psicológica com perspectiva "
        "gerencial para intervenções organizacionais mais eficazes. A tríplice formação — "
        "Psicologia + Avaliação Psicológica + MBA em Gestão de Pessoas — configura um perfil "
        "acadêmico que excede substancialmente os requisitos educacionais típicos da área."
    )

    # Credential Evaluation
    add_teal_sub_header(doc, "Avaliação de Credenciais")

    add_body_paragraph(doc,
        "O diploma de Bacharelado em Psicologia da Associação Catarinense de Ensino foi "
        "submetido a avaliação de equivalência pela GEO Credential Services, que o reconheceu "
        "como equivalente ao bacharelado norte-americano (U.S. Bachelor's equivalent). Esta "
        "equivalência foi aceita pelo USCIS na petição original, não constituindo ponto de "
        "controvérsia no processo."
    )


def build_courses(doc):
    """Seção 7: Cursos, Palestras e Certificações."""
    add_navy_section_header(doc, "CURSOS, PALESTRAS E CERTIFICAÇÕES")

    add_teal_sub_header(doc, "Certificações Profissionais")

    add_evidence_block(doc, [
        ("Certificação", "Predictive Index (PI) — Analista Certificada"),
        ("Emissor", "PRAENDEX Brasil / The Predictive Index"),
        ("Data", "20 de abril de 2020"),
        ("Validade", "Permanente"),
    ],
        "Certificação que habilita Cristine a definir requisitos comportamentais e cognitivos "
        "de posições, aplicar avaliações individuais e em grupo, e interpretar resultados para "
        "orientar decisões estratégicas de gestão de pessoas. O Predictive Index é utilizado "
        "por mais de 10.000 empresas em 142 países e fundamenta-se em mais de 60 anos de "
        "pesquisa em ciência comportamental, conferindo a esta certificação reconhecimento "
        "internacional e aplicabilidade direta no mercado norte-americano.",
        compact=True
    )

    add_teal_sub_header(doc, "Associações Profissionais")

    add_bullet(doc, "Society for Human Resource Management (SHRM) — Membro ativa, válida até 31 de julho de 2026. Principal associação profissional de RH dos Estados Unidos com mais de 325.000 membros em 165 países.")
    add_bullet(doc, "American Management Association (AMA) — Membro desde 6 de dezembro de 2023 (ID 4016170). Associação centenária dedicada ao desenvolvimento de competências gerenciais e de liderança.")
    add_bullet(doc, "Conselho Regional de Psicologia — 12ª Região (CRP-12/11050) — Registro ativo e regular. Órgão regulador do exercício profissional de psicologia no Brasil, com sede em Florianópolis, SC.")

    add_teal_sub_header(doc, "Cursos e Capacitações Complementares")

    add_bullet(doc, "HR Management Fundamentals — Certificação em fundamentos de gestão de recursos humanos no contexto norte-americano, abrangendo legislação trabalhista, práticas de compensação, e compliance regulatório.")
    add_bullet(doc, "ADP Effective Onboarding — Treinamento em metodologias de integração de novos colaboradores utilizando a plataforma ADP, líder mundial em soluções de gestão de capital humano.")
    add_bullet(doc, "Recrutamento e Seleção: Metodologia de Resultado — Kenoby, abril de 2022. Curso especializado em metodologias avançadas de recrutamento orientado por resultados e métricas de eficácia.")
    add_bullet(doc, "Plataforma Gupy — Maio de 2022. Treinamento na principal plataforma brasileira de recrutamento inteligente, integrando inteligência artificial ao processo seletivo.")
    add_bullet(doc, "Introdução à PLD/FT (Prevenção à Lavagem de Dinheiro) — ASAAS. Capacitação em conformidade regulatória financeira aplicada ao contexto de fintechs.")
    add_bullet(doc, "PCI-DSS — ASAAS, abril de 2022. Treinamento em segurança de dados de pagamento conforme padrões internacionais PCI.")
    add_bullet(doc, "Comunicação Acessível — ASAAS, abril de 2022. Capacitação em práticas de comunicação inclusiva para ambientes organizacionais diversos.")
    add_bullet(doc, "Congresso CONCARH — Participação em congresso referência em gestão de pessoas e recursos humanos no Brasil.")
    add_bullet(doc, "Customer Success — Certificação em metodologias de sucesso do cliente, aplicável a gestão de relacionamento com clientes de serviços de RH.")


def build_proposed_endeavors(doc):
    """Seção 8: Projeto EB-2 NIW — Proposed Endeavors."""
    add_navy_section_header(doc, "PROJETO EB-2 NIW — PROPOSED ENDEAVORS")

    # Introductory paragraph
    add_body_paragraph(doc,
        "Os empreendimentos propostos por Cristine Correa nos Estados Unidos representam a "
        "convergência estratégica de mais de uma década de experiência em gestão de pessoas, "
        "avaliação comportamental e diagnóstico organizacional com as necessidades críticas e "
        "documentadas da força de trabalho norte-americana. Em um contexto onde a rotatividade "
        "voluntária custa às empresas americanas aproximadamente US$ 2,9 trilhões por ano e "
        "onde 33,2 milhões de pequenas e médias empresas (PMEs) empregam 46,4% da força de "
        "trabalho privada do país (SBA, 2024), as propostas de Cristine abordam uma lacuna "
        "estrutural no mercado: a ausência de serviços especializados de inteligência "
        "comportamental e clima organizacional acessíveis e desenhados especificamente para "
        "PMEs em setores com taxas críticas de rotatividade."
    )

    # --- PE1 ---
    add_teal_sub_header(doc, "Empreendimento Proposto 1: Workforce Behavioral Intelligence Institute")

    add_body_paragraph(doc,
        "O Workforce Behavioral Intelligence Institute é um instituto privado de avaliação "
        "comportamental e inteligência de força de trabalho direcionado a pequenas e médias "
        "empresas com 10 a 500 funcionários na região da Flórida Central, com sede em Sanford, "
        "Condado de Seminole. O instituto oferecerá três serviços nucleares: avaliações "
        "comportamentais proprietárias baseadas na metodologia Predictive Index — cujo domínio "
        "técnico é certificado e comprovado pela trajetória de Cristine —, diagnósticos de "
        "inteligência de força de trabalho que integram dados comportamentais com métricas "
        "operacionais de RH, e programas de implementação supervisionada onde Cristine atua "
        "diretamente na interpretação dos resultados e na formulação de recomendações "
        "estratégicas para cada cliente."
    )

    add_body_paragraph(doc,
        "O mercado-alvo do instituto são PMEs nos setores de saúde e hospitalidade/serviços "
        "de alimentação na Flórida Central. O setor de saúde apresenta taxa de rotatividade "
        "de 22,7% — quase o dobro da média nacional de 12,5% —, enquanto hospitalidade e "
        "serviços de alimentação atingem taxas entre 50% e 80%, as mais altas da economia "
        "americana (Bureau of Labor Statistics, JOLTS 2024-2025). O custo de substituição "
        "de cada funcionário varia entre 50% e 200% do salário anual (Gallup/SHRM), tornando "
        "a rotatividade um problema financeiro crítico para PMEs que operam com margens "
        "reduzidas e recursos limitados de RH."
    )

    add_body_paragraph(doc,
        "De acordo com dados do Bureau of Labor Statistics, profissionais classificados sob "
        "o código SOC 11-3121 (Human Resources Managers) recebem salário mediano de "
        "US$ 140.030 nos Estados Unidos, com projeção de crescimento de 5% no período "
        "2024-2034 e aproximadamente 17.900 novas posições por ano. O código SOC 11-3131 "
        "(Training and Development Managers) apresenta salário mediano de US$ 123.470 com "
        "projeção de crescimento de 6% no mesmo período. Ambos os códigos refletem funções "
        "gerenciais que não exigem licenciamento profissional específico nos Estados Unidos, "
        "e o Estado da Flórida — conforme o Fl. Stat. §490.014 — expressamente isenta "
        "atividades de avaliação organizacional, treinamento corporativo e coaching dos "
        "requisitos de licenciamento em psicologia."
    )

    add_body_paragraph(doc,
        "O mercado de plataformas de pesquisa de engajamento de funcionários nos Estados "
        "Unidos foi avaliado em US$ 1,6-1,7 bilhão em 2024, com projeção de alcançar "
        "US$ 3,5-4,2 bilhões até 2033, representando um CAGR de 8-11% (Grand View Research, "
        "2024). O segmento de PMEs é o que apresenta crescimento mais acelerado (CAGR ~13%), "
        "impulsionado pela crescente conscientização sobre o impacto do engajamento na "
        "produtividade. Plataformas dominantes como Culture Amp, Lattice e Glint são "
        "precificadas para o mercado corporativo e enterprise (US$ 8-30 por funcionário/mês), "
        "deixando uma lacuna significativa no atendimento a PMEs que necessitam de soluções "
        "acessíveis e personalizadas."
    )

    add_body_paragraph(doc,
        "O Workforce Behavioral Intelligence Institute alinha-se diretamente com múltiplas "
        "prioridades federais de desenvolvimento da força de trabalho. O Workforce Innovation "
        "and Opportunity Act (WIOA), reestruturado pelo A Stronger Workforce for America Act "
        "(MASA), destina US$ 2,966 bilhões para o ano fiscal 2026 para programas de "
        "desenvolvimento de força de trabalho, com ênfase em avaliação de competências e "
        "correspondência entre habilidades e demandas do mercado. A Executive Order sobre "
        "Skills-Based Hiring (abril de 2025) reforça a priorização de contratações baseadas "
        "em competências verificáveis em detrimento de credenciais formais — exatamente o tipo "
        "de abordagem que o instituto de Cristine implementará através de avaliações "
        "comportamentais objetivas. O EDA Good Jobs Challenge (US$ 500 milhões) prioriza "
        "indústrias com escassez crítica de mão de obra qualificada, incluindo saúde, "
        "hospitalidade e manufatura."
    )

    # --- PE2 ---
    add_teal_sub_header(doc, "Empreendimento Proposto 2: Organizational Climate & Strategic Retention Lab")

    add_body_paragraph(doc,
        "O Organizational Climate & Strategic Retention Lab é um laboratório especializado "
        "que combina diagnóstico de clima organizacional com análise preditiva de retenção, "
        "direcionado a PMEs nos setores de saúde e manufatura na Flórida Central. Os serviços "
        "nucleares incluem: pesquisas de clima organizacional proprietárias desenvolvidas com "
        "base na experiência de Cristine com a metodologia GPTW na UNISOCIESC, análise "
        "preditiva de retenção que correlaciona indicadores de clima com padrões de "
        "rotatividade, e planos de intervenção comportamental personalizados onde Cristine "
        "aplica sua expertise em avaliação psicológica e desenvolvimento organizacional para "
        "desenhar e supervisionar a implementação de estratégias de retenção específicas "
        "para cada organização cliente."
    )

    add_body_paragraph(doc,
        "O setor de manufatura nos Estados Unidos enfrenta taxas de rotatividade entre 24% e "
        "32%, com impacto particularmente severo em operações que dependem de mão de obra "
        "qualificada e treinamento extensivo. Na Flórida Central, o setor de saúde projeta "
        "um déficit de mais de 30.000 enfermeiros até 2027, e as publicações de vagas em "
        "saúde na região de Orlando aumentaram 28% nos últimos 12 meses. O Condado de "
        "Seminole, onde Cristine reside, apresenta taxa de desemprego de 4,2% — um aumento "
        "de 1,4 pontos percentuais em relação ao ano anterior —, indicando um mercado de "
        "trabalho em transição que demanda estratégias sofisticadas de retenção."
    )

    add_body_paragraph(doc,
        "O laboratório alinha-se com prioridades federais específicas de saúde e bem-estar "
        "do trabalhador. O programa NIOSH Total Worker Health® enfatiza a importância de "
        "abordagens integradas que considerem fatores organizacionais, comportamentais e "
        "ambientais na promoção da saúde do trabalhador — precisamente o tipo de diagnóstico "
        "que o laboratório de Cristine realizará. O Surgeon General's Framework for Workplace "
        "Mental Health & Well-Being (2022) estabelece cinco elementos essenciais para locais "
        "de trabalho saudáveis, incluindo conexão social, oportunidade de crescimento e "
        "cultura organizacional positiva. As diretrizes do OSHA para saúde mental no trabalho "
        "reforçam a necessidade de intervenções baseadas em evidências para ambientes "
        "organizacionais, validando a abordagem científica que Cristine aplicará em seu "
        "laboratório."
    )

    add_body_paragraph(doc,
        "O mercado global de testes psicométricos e avaliação comportamental foi avaliado em "
        "US$ 5,8 bilhões em 2024, com CAGR projetado de 9,6%. A meta-análise da Gallup "
        "(183.806 unidades de negócio, 3,35 milhões de funcionários) demonstrou correlação "
        "de 0,49 entre engajamento e desempenho organizacional, com equipes no quartil "
        "superior de engajamento apresentando 23% mais lucratividade e 18% mais "
        "produtividade. Esses dados fundamentam a proposição de valor do laboratório "
        "de Cristine: intervenções de clima organizacional baseadas em avaliação "
        "comportamental científica geram retorno mensurável e documentável para PMEs."
    )

    # --- Comparison Table ---
    add_teal_sub_header(doc, "Análise Comparativa dos Empreendimentos Propostos")

    comp_data = [
        ("Dimensão", "PE 1: Behavioral Intelligence", "PE 2: Climate & Retention Lab"),
        ("Mercado-Alvo", "PMEs Saúde e Hospitalidade\n(10-500 funcionários)", "PMEs Saúde e Manufatura\n(10-500 funcionários)"),
        ("Código SOC Primário", "11-3121 (HR Managers)\nMediana: US$ 140.030", "11-3131 (Training & Dev Managers)\nMediana: US$ 123.470"),
        ("Crescimento Projetado\n(BLS 2024-2034)", "5% (17.900 posições/ano)", "6% (3.800 posições/ano)"),
        ("Política Federal\nAlinhada", "WIOA/MASA, EO Skills-Based\nHiring, EDA Good Jobs Challenge", "NIOSH Total Worker Health,\nSurgeon General Framework,\nOSHA Guidelines"),
        ("Receita Projetada\nAno 1", "US$ 45.000 — US$ 75.000", "US$ 50.000 — US$ 80.000"),
        ("Receita Projetada\nAno 2", "US$ 150.000 — US$ 250.000", "US$ 180.000 — US$ 300.000"),
        ("Investimento Inicial", "< US$ 12.000", "< US$ 15.000"),
        ("Localização", "Sanford, FL (Seminole County)", "Sanford, FL (Seminole County)"),
    ]

    table = doc.add_table(rows=len(comp_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)

    for row_idx, row_data in enumerate(comp_data):
        row = table.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            set_cell_borders(cell, "CCCCCC", 2)
            set_cell_margins(cell, top=40, bottom=40, left=80, right=80)

            if row_idx == 0:
                set_cell_shading(cell, "2D3E50")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, text, size=FONT_TABLE_SMALL, color=WHITE, bold=True)
            else:
                if row_idx % 2 == 0:
                    set_cell_shading(cell, "F5F5F5")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                if col_idx == 0:
                    add_run(p, text, size=FONT_TABLE_SMALL, color=NAVY, bold=True)
                else:
                    add_run(p, text, size=FONT_TABLE_SMALL, color=BLACK)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- Concluding Dhanasar Paragraph ---
    add_teal_sub_header(doc, "Alinhamento com o Framework Dhanasar (2016)")

    add_body_paragraph(doc,
        "Os empreendimentos propostos por Cristine Correa satisfazem os três requisitos "
        "estabelecidos pelo precedente Matter of Dhanasar, 26 I&N Dec. 884 (AAO 2016), "
        "que governa a adjudicação de petições EB-2 NIW. Em relação ao primeiro requisito "
        "(mérito substancial e importância nacional), ambos os empreendimentos abordam a "
        "crise documentada de rotatividade que custa trilhões de dólares à economia "
        "norte-americana anualmente, com foco em setores expressamente prioritários pelo "
        "governo federal — saúde, hospitalidade e manufatura — e em alinhamento direto com "
        "legislação e ordens executivas vigentes, incluindo o WIOA/MASA, o EDA Good Jobs "
        "Challenge, o NIOSH Total Worker Health e o Surgeon General's Framework for "
        "Workplace Mental Health."
    )

    add_body_paragraph(doc,
        "Em relação ao segundo requisito (capacidade de avançar o empreendimento), a "
        "trajetória profissional de Cristine demonstra de forma inequívoca sua capacidade "
        "de executar os empreendimentos propostos. São mais de dez anos de experiência "
        "progressiva em gestão de pessoas em organizações de diferentes portes e setores, "
        "certificação específica em Predictive Index, experiência documentada com pesquisas "
        "de clima GPTW, três anos de coordenação de programa de inclusão, e formação "
        "acadêmica tríplice que combina psicologia, avaliação psicológica e gestão de "
        "pessoas. O fato de o segundo requisito ter sido aceito na adjudicação anterior "
        "reforça a solidez da qualificação da beneficiária."
    )

    add_body_paragraph(doc,
        "Em relação ao terceiro requisito (balanço de equidades favorável aos Estados Unidos), "
        "a aprovação dos empreendimentos propostos beneficiará os Estados Unidos de forma "
        "substancialmente superior à disponibilização da posição no mercado de trabalho. "
        "A obtenção de certificação trabalhista via PERM seria impraticável dados os prazos "
        "de processamento de 18 a 24 meses do Departamento do Trabalho e a natureza "
        "empreendedora das atividades propostas, que não se enquadram em relação empregador-empregado "
        "tradicional. A urgência da contribuição é evidenciada pela intensificação da crise de "
        "retenção documentada pelo JOLTS — com mais de 3,3 milhões de separações voluntárias "
        "por mês nos Estados Unidos — e pela escassez projetada de profissionais de saúde na "
        "Flórida Central. A combinação singular de formação em psicologia organizacional, "
        "certificação em avaliação comportamental PI, experiência com GPTW e expertise "
        "bicultural Brasil-EUA constitui um perfil que não pode ser replicado pelo mercado "
        "de trabalho doméstico e cuja perda representaria prejuízo concreto ao interesse "
        "nacional."
    )


def build_recommendation_letters(doc):
    """Seção 9: Cartas de Recomendação."""
    add_navy_section_header(doc, "CARTAS DE RECOMENDAÇÃO")

    add_body_paragraph(doc,
        "As cartas de recomendação apresentadas nesta petição foram elaboradas por "
        "profissionais com conhecimento direto das qualificações técnicas e contribuições "
        "profissionais de Cristine Correa, representando diferentes perspectivas — "
        "empregadores anteriores, parceiros profissionais e referências do mercado "
        "norte-americano."
    )

    # Table with letters
    letters = [
        ("Signatário", "Organização", "Relação"),
        ("Lisiane Rúbia Bim", "UNISOCIESC / Ecossistema Ânima", "Gestora Direta (2017-2021)"),
        ("Renata Santos de Araújo", "Walmart Brasil / A Casa da Cura", "Gestora (2016-2017) / Diretora"),
        ("Bianka Fernanda dos Santos", "ASAAS (ex-colega)", "Psicóloga / Referência Profissional"),
        ("Fernando Braga Neto", "Time 4 Travel LLC, Orlando", "CEO / Investidor nos EUA"),
        ("Danieli Nieri", "Consultora de Imagem, Orlando", "Parceira B2B nos EUA"),
        ("Ana Luiza Baldini", "Seta Soluções Visuais", "Diretora / Referência Profissional"),
        ("Daniel Gouveia Tanigushi", "Pesquisador em Biociências", "Mestre/Doutorando / Referência"),
        ("Franciele Pereira Zazycki", "UNISOCIESC", "Colega / Referência Profissional"),
    ]

    table = doc.add_table(rows=len(letters), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)

    for row_idx, row_data in enumerate(letters):
        row = table.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            set_cell_borders(cell, "CCCCCC", 2)
            set_cell_margins(cell, top=40, bottom=40, left=80, right=80)

            if row_idx == 0:
                set_cell_shading(cell, "2D3E50")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, text, size=FONT_TABLE_SMALL, color=WHITE, bold=True)
            else:
                if row_idx % 2 == 0:
                    set_cell_shading(cell, "F5F5F5")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if col_idx == 0:
                    add_run(p, text, size=FONT_TABLE_SMALL, color=BLACK, bold=True)
                else:
                    add_run(p, text, size=FONT_TABLE_SMALL, color=BLACK)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ============================================================
# MAIN BUILD
# ============================================================

def main():
    print("=" * 60)
    print("PETITION ENGINE v2.0 — Résumé EB-2 NIW")
    print(f"Beneficiária: {BENEFICIARY_NAME}")
    print("=" * 60)

    # Create document
    print("\n[1/11] Configurando documento...")
    doc = setup_document()

    print("[2/11] Construindo header...")
    build_header(doc)

    print("[3/11] Construindo footer...")
    build_footer(doc)

    print("[4/11] Seção: Síntese Profissional...")
    build_synthesis(doc)

    print("[5/11] Seção: Histórico Profissional (Gantt)...")
    build_timeline(doc)

    print("[6/11] Seção: Experiência Profissional...")
    build_experience(doc)

    print("[7/11] Seção: Contribuições Técnicas...")
    build_contributions(doc)

    print("[8/11] Seção: Publicações e Artigos...")
    build_publications(doc)

    print("[9/11] Seção: Formação Acadêmica...")
    build_education(doc)

    print("[10/11] Seção: Cursos e Certificações...")
    build_courses(doc)

    print("[11/11] Seção: Proposed Endeavors...")
    build_proposed_endeavors(doc)

    # Recommendation Letters
    print("[BONUS] Seção: Cartas de Recomendação...")
    build_recommendation_letters(doc)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"\n{'=' * 60}")
    print(f"DOCUMENTO GERADO COM SUCESSO!")
    print(f"Path: {OUTPUT_FILE}")
    print(f"{'=' * 60}")

    return OUTPUT_FILE


if __name__ == "__main__":
    main()
