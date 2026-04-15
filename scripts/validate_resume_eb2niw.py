#!/usr/bin/env python3
"""
COMPREHENSIVE RESUME VALIDATOR — EB-2 NIW
Combines QUALITY_REVIEWER.md checks + ALL error_rules.json patterns.

Usage:
    python3 validate_resume_eb2niw.py /path/to/resume.docx
"""

import sys
import os
import re
from collections import OrderedDict

try:
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree
except ImportError:
    print("ERROR: Missing dependencies. Install with:")
    print("  pip3 install python-docx lxml")
    sys.exit(1)

# ═══════════════════════════════════════════════════════════════
# CONFIG
# ═══════════════════════════════════════════════════════════════

ALLOWED_FONTS = {"garamond"}

ALLOWED_COLORS = {
    "2D3E50", "3498A2", "FFFFFF", "000000",
    "333333", "666666", "F5F5F5", "CCCCCC", "FAFAFA",
}

# ═══════════════════════════════════════════════════════════════
# ERROR RULES — REGEX PATTERNS
# ═══════════════════════════════════════════════════════════════

BLOCK_PATTERNS = [
    # Opinion language
    (r'\b(I|we)\s+believe\b', "BLOCK", "Opinion language: 'I/we believe'"),
    (r'\b(I|we)\s+think\b', "BLOCK", "Opinion language: 'I/we think'"),
    # Terminology
    (r'proposed\s+(venture|business)', "BLOCK", "Wrong terminology: 'proposed venture/business'"),
    (r'\b(in conclusion|to summarize)\b', "BLOCK", "Forbidden transition: 'in conclusion/to summarize'"),
    # Forbidden SOC codes in body
    (r'\b(23-1011|29-1069|17-201[1-9]|13-2011)\b', "BLOCK", "Forbidden SOC code found"),
    # Meta/internal leaks
    (r'\bprompt\b', "BLOCK", "Internal term leak: 'prompt'"),
    (r'\b(PROEX|Kortix|Carlos Avelino)\b', "BLOCK", "Internal reference leak"),
    # Missing accents (Portuguese)
    (r'\b(introducao|peticao|informacao|certificacao|formacao|avaliacao|ocupacao|operacao|integracao|migracao|capacitacao)\b', "BLOCK", "Missing accent on Portuguese word"),
    # Version/system metadata
    (r'\b(Version \d|Generated:|SaaS Evidence Architect|Petition Engine)\b', "BLOCK", "System metadata leak"),
    # Scalability forbidden terms
    (r'\b(standardized|padronizado|operates autonomously|self-sustaining|auto-sustent|plug.and.play|train.the.trainer|white.label|marca branca|client autonomy|founder dependency|scalable without|replicable by any|turnkey|chave.na.m)\b', "BLOCK", "Forbidden scalability term (USCIS red flag)"),
    # Denial/prior filing references
    (r'\b(denial|negativa anterior|RFE anterior|previous petition|prior filing|refile|segunda tentativa|nova submiss[ãa]o|peti[çc][ãa]o anterior)\b', "BLOCK", "Reference to prior filing/denial"),
    # Legal team references
    (r'\b(equipe jur[ií]dica|advogado|escrit[oó]rio de advocacia|representa[çc][aã]o legal)\b', "BLOCK", "Legal team reference (should not appear in resume)"),
    # Judicial references
    (r'\b(tribunal|ju[ií]z|senten[çc]a|julgamento|vara|processo judicial)\b', "BLOCK", "Judicial reference"),
    # Sworn translation references
    (r'\b(tradu[çc][aã]o juramentada|tradutor juramentado|tradutor p[uú]blico)\b', "BLOCK", "Sworn translation reference"),
    # RAG system references
    (r'\b(RAG I|RAG II|RAG III|RAGs|reposit[oó]rio de argumenta[çc][aã]o)\b', "BLOCK", "RAG system reference leak"),
    # Engine/generation references
    (r'\b(Petition Engine|Forjado por|gerado automaticamente|gerado por)\b', "BLOCK", "Generation engine reference leak"),
    # Internal tools
    (r'\b(Obsidian|formato \.md|markdown)\b', "BLOCK", "Internal tool reference"),
    # Draft/version markers
    (r'(Vers[aã]o:? \d|V\d\.\d|Descontaminad|Separation of Concerns|SoC aplicado|Para Revis[aã]o|Rascunho Interno|DOCUMENTO INTERNO|Quality Gate|auto.learning)', "BLOCK", "Draft/version marker"),
    # TODO/placeholder markers
    (r'\[(VERIFICAR|TODO|INSERIR|PENDENTE|TBD|XXX|COMPLETAR|PREENCHER)\]', "BLOCK", "Placeholder marker found"),
    # Cross-references to letters
    (r'\b(conforme carta|como atestado|segundo recomenda|carta de recomenda|recommendation letter|support letter|as attested by)\b', "BLOCK", "Cross-reference to recommendation letter"),
    # Exhibit numbering (should use Evidência)
    (r'\bExhibit\s+\d', "BLOCK", "Exhibit numbering (should use 'Evidência')"),
    # Proposed endeavor references (should NOT be in resume)
    (r'\b(proposed endeavor|endeavor\s*(1|2|3)|three endeavor|tres endeavor|opcao de endeavor)\b', "BLOCK", "Proposed endeavor reference in resume"),
    # Financial projections
    (r'\b(revenue|receita bruta|faturamento|projecao financeira|lucro liquido|net income|ROI|EBITDA|fluxo de caixa|margem de lucro|profit margin|financial projection)\b', "BLOCK", "Financial projection term"),
    # R$ currency
    (r'R\$', "BLOCK", "Brazilian currency R$ found"),
]

WARN_PATTERNS = [
    (r'\b(consultoria|consulting)\b', "WARN", "Potentially problematic term: 'consultoria/consulting'"),
]

ALL_PATTERNS = BLOCK_PATTERNS + WARN_PATTERNS

# ═══════════════════════════════════════════════════════════════
# FUNCTIONS
# ═══════════════════════════════════════════════════════════════

def load_document(path):
    """Load DOCX and extract all text."""
    doc = Document(path)
    para_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + "\n"
    merged = para_text + "\n" + table_text
    return doc, merged


def check_fonts(doc):
    """[S0] Check all fonts are Garamond."""
    issues = []
    font_counts = {}
    total_runs = 0

    def scan_runs(paragraphs):
        nonlocal total_runs
        for para in paragraphs:
            for run in para.runs:
                total_runs += 1
                fname = run.font.name
                if fname:
                    key = fname.strip()
                    font_counts[key] = font_counts.get(key, 0) + 1

    # Paragraphs
    scan_runs(doc.paragraphs)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan_runs(cell.paragraphs)

    # Header/footer
    for section in doc.sections:
        if section.header:
            scan_runs(section.header.paragraphs)
            for t in section.header.tables:
                for r in t.rows:
                    for c in r.cells:
                        scan_runs(c.paragraphs)
        if section.footer:
            scan_runs(section.footer.paragraphs)
            for t in section.footer.tables:
                for r in t.rows:
                    for c in r.cells:
                        scan_runs(c.paragraphs)

    # Analyze
    bad_fonts = {}
    for fname, count in font_counts.items():
        if fname.lower() not in ALLOWED_FONTS:
            bad_fonts[fname] = count

    if bad_fonts:
        for fname, count in bad_fonts.items():
            severity = "S0" if "arial" in fname.lower() or "calibri" in fname.lower() else "S1"
            issues.append((severity, f"Non-Garamond font '{fname}' found in {count} runs"))

    return issues, total_runs, font_counts


def check_page_setup(doc):
    """[S1] Check page dimensions and margins."""
    issues = []
    section = doc.sections[0]

    w = round(section.page_width.inches, 1)
    h = round(section.page_height.inches, 1)
    if w != 8.5 or h != 11.0:
        issues.append(("S1", f"Paper size WRONG: {w}\"x{h}\" (expected 8.5\"x11\")"))

    lm = round(section.left_margin.inches, 2)
    rm = round(section.right_margin.inches, 2)
    if abs(lm - 0.65) > 0.05:
        issues.append(("S1", f"Left margin: {lm}\" (expected ~0.65\")"))
    if abs(rm - 0.65) > 0.05:
        issues.append(("S1", f"Right margin: {rm}\" (expected ~0.65\")"))

    return issues


def check_regex_patterns(merged_text):
    """Run ALL regex patterns against extracted text."""
    issues = []
    for pattern, action, desc in ALL_PATTERNS:
        flags = re.IGNORECASE
        matches = list(re.finditer(pattern, merged_text, flags))
        if matches:
            # Collect unique match texts
            unique_matches = set()
            for m in matches:
                unique_matches.add(m.group())
            severity = "S0" if action == "BLOCK" else "S2"
            match_preview = ", ".join(list(unique_matches)[:5])
            issues.append((
                severity,
                f"[{action}] {desc} — {len(matches)} occurrence(s): [{match_preview}]"
            ))
    return issues


def check_word_count(merged_text):
    """Check minimum word count >= 3500."""
    issues = []
    words = merged_text.split()
    word_count = len(words)
    if word_count < 3500:
        issues.append(("S1", f"Word count TOO LOW: {word_count} words (minimum 3,500)"))
    return issues, word_count


def check_structure(doc, merged_text):
    """Check document structure requirements."""
    issues = []
    upper = merged_text.upper()

    # Required sections for EB-2 NIW resume
    required_sections = [
        ("Executive Summary / Síntese", ["EXECUTIVE SUMMARY", "SÍNTESE PROFISSIONAL", "PROFESSIONAL SUMMARY"]),
        ("Career Timeline / Histórico", ["CAREER TIMELINE", "HISTÓRICO PROFISSIONAL", "PROFESSIONAL CAREER", "PROFESSIONAL EXPERIENCE", "EXPERIÊNCIA PROFISSIONAL"]),
    ]

    for label, needles in required_sections:
        found = any(n in upper for n in needles)
        if not found:
            issues.append(("S1", f"Missing section: {label}"))

    # Check for "Responsabilidades" and "Resultados" sub-sections
    has_responsabilidades = "RESPONSABILIDADES" in upper or "RESPONSIBILITIES" in upper
    has_resultados = "RESULTADOS" in upper or "KEY RESULTS" in upper or "RESULTS" in upper
    if not has_responsabilidades:
        issues.append(("S2", "No 'Responsabilidades' sub-sections found in experience entries"))
    if not has_resultados:
        issues.append(("S2", "No 'Resultados' sub-sections found in experience entries"))

    # Check NO "Proposed Endeavors" section
    if "PROPOSED ENDEAVOR" in upper:
        issues.append(("S0", "Resume contains 'Proposed Endeavors' section (should NOT exist in resume)"))

    # Check for evidence placeholders
    thumbnail_pattern = re.compile(r'\[THUMBNAIL.*?Evid[eê]ncia\s*\d+\]', re.IGNORECASE)
    thumbnail_matches = thumbnail_pattern.findall(merged_text)
    if thumbnail_matches:
        issues.append(("S2", f"{len(thumbnail_matches)} [THUMBNAIL] placeholders found (images not yet inserted)"))

    # Check for Exhibit numbering (should use Evidência)
    exhibit_pattern = re.compile(r'\bExhibit\s+\d', re.IGNORECASE)
    if exhibit_pattern.search(merged_text):
        issues.append(("S0", "Uses 'Exhibit' numbering instead of 'Evidência'"))

    return issues


def check_header_footer(doc):
    """Check header and footer presence."""
    issues = []
    section = doc.sections[0]

    # Header
    header = section.header
    if not header.paragraphs and not header.tables:
        issues.append(("S1", "HEADER MISSING"))
    else:
        header_text = " ".join(p.text for p in header.paragraphs)
        for t in header.tables:
            for r in t.rows:
                for c in r.cells:
                    header_text += " " + c.text
        if not header_text.strip():
            issues.append(("S2", "Header exists but appears empty"))

    # Footer
    footer = section.footer
    footer_text = ""
    if footer.paragraphs:
        footer_text = " ".join(p.text for p in footer.paragraphs)
    for t in getattr(footer, 'tables', []):
        for r in t.rows:
            for c in r.cells:
                footer_text += " " + c.text

    if not footer_text.strip():
        issues.append(("S1", "FOOTER MISSING or empty"))
    else:
        # Check Page X of Y
        page_pattern = re.compile(r'page\s+\d+\s+of\s+\d+|p[aá]gina\s+\d+\s+de\s+\d+', re.IGNORECASE)
        if not page_pattern.search(footer_text):
            # python-docx may not see dynamic fields, so just note it
            issues.append(("S3", "Footer text doesn't show 'Page X of Y' (may use dynamic fields)"))

    return issues


def check_images(doc):
    """Check for images and thumbnail placeholders."""
    issues = []
    img_count = len(doc.inline_shapes)

    if img_count == 0:
        issues.append(("S1", "ZERO images in document (evidence blocks without thumbnails)"))

    placeholder_count = 0
    for para in doc.paragraphs:
        if "[THUMBNAIL" in para.text.upper():
            placeholder_count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "[THUMBNAIL" in cell.text.upper():
                    placeholder_count += 1

    if placeholder_count > 0:
        issues.append(("S2", f"{placeholder_count} [THUMBNAIL] placeholders still present (images not inserted)"))

    return issues, img_count, placeholder_count


def check_evidence_blocks(doc):
    """Check evidence block structure."""
    issues = []
    tables_2col = 0
    tables_with_impact = 0

    for table in doc.tables:
        if len(table.columns) == 2 and len(table.rows) >= 1:
            tables_2col += 1
            cell_text = ""
            for row in table.rows:
                for cell in row.cells:
                    cell_text += cell.text + " "
            if any(kw in cell_text for kw in ["Impact", "Impacto", "Description", "Descrição"]):
                tables_with_impact += 1

    if tables_2col > 0 and tables_with_impact == 0:
        issues.append(("S1", f"No evidence blocks have 'Impact/Impacto' inside ({tables_2col} 2-col tables found)"))

    return issues, tables_2col, tables_with_impact


def check_colors(doc):
    """Check cell shading colors."""
    issues = []
    unexpected = set()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill and fill.upper() not in ALLOWED_COLORS and fill != "auto":
                            unexpected.add(fill.upper())

    if unexpected:
        issues.append(("S2", f"Unexpected cell shading colors: {unexpected}"))

    return issues


def check_soc_codes_in_body(merged_text, doc):
    """SOC codes (XX-XXXX) should only be in header, not body text."""
    issues = []
    # Get header text
    header_text = ""
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                header_text += p.text + " "
            for t in section.header.tables:
                for r in t.rows:
                    for c in r.cells:
                        header_text += c.text + " "

    # Get body-only text (paragraphs)
    body_text = "\n".join(p.text for p in doc.paragraphs)

    soc_pattern = re.compile(r'\b\d{2}-\d{4}(\.\d{2})?\b')
    body_matches = list(soc_pattern.finditer(body_text))

    # Filter: some SOC codes are legitimate in body (like BLS references)
    # Only flag the specifically forbidden ones
    forbidden_soc = {"23-1011", "29-1069", "17-2011", "17-2012", "17-2013",
                     "17-2014", "17-2015", "17-2016", "17-2017", "17-2018",
                     "17-2019", "13-2011"}
    for m in body_matches:
        code = m.group().split('.')[0]  # strip .XX suffix
        if code in forbidden_soc:
            issues.append(("S0", f"Forbidden SOC code {m.group()} found in body text"))

    return issues


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════

def run_full_validation(docx_path):
    print("=" * 72)
    print(f"  COMPREHENSIVE EB-2 NIW RESUME VALIDATOR")
    print(f"  File: {os.path.basename(docx_path)}")
    print(f"  Size: {os.path.getsize(docx_path):,} bytes")
    print("=" * 72)

    doc, merged = load_document(docx_path)

    all_issues = []

    # 1. FONTS
    print("\n[1/11] Checking fonts (must be 100% Garamond)...")
    font_issues, total_runs, font_counts = check_fonts(doc)
    all_issues.extend(font_issues)
    print(f"        {total_runs} runs scanned")
    print(f"        Font distribution: {dict(font_counts)}")

    # 2. PAGE SETUP
    print("[2/11] Checking page setup...")
    all_issues.extend(check_page_setup(doc))

    # 3. REGEX PATTERNS (ALL error rules)
    print(f"[3/11] Running {len(ALL_PATTERNS)} regex patterns...")
    regex_issues = check_regex_patterns(merged)
    all_issues.extend(regex_issues)
    print(f"        {len(regex_issues)} pattern violations found")

    # 4. WORD COUNT
    print("[4/11] Checking word count (min 3,500)...")
    wc_issues, word_count = check_word_count(merged)
    all_issues.extend(wc_issues)
    print(f"        Total words: {word_count:,}")

    # 5. STRUCTURE
    print("[5/11] Checking document structure...")
    struct_issues = check_structure(doc, merged)
    all_issues.extend(struct_issues)

    # 6. HEADER/FOOTER
    print("[6/11] Checking header and footer...")
    all_issues.extend(check_header_footer(doc))

    # 7. IMAGES
    print("[7/11] Checking images and thumbnails...")
    img_issues, img_count, placeholder_count = check_images(doc)
    all_issues.extend(img_issues)
    print(f"        {img_count} images, {placeholder_count} placeholders")

    # 8. EVIDENCE BLOCKS
    print("[8/11] Checking evidence blocks...")
    eb_issues, tables_2col, tables_impact = check_evidence_blocks(doc)
    all_issues.extend(eb_issues)
    print(f"        {tables_2col} evidence blocks, {tables_impact} with impact")

    # 9. COLORS
    print("[9/11] Checking cell shading colors...")
    all_issues.extend(check_colors(doc))

    # 10. SOC CODES IN BODY
    print("[10/11] Checking for forbidden SOC codes in body...")
    all_issues.extend(check_soc_codes_in_body(merged, doc))

    # 11. EB-2 NIW SPECIFIC
    print("[11/11] Running EB-2 NIW specific checks...")
    upper = merged.upper()
    # Should NOT have EB-1A criteria sections
    if "CRITERION 1" in upper or "CRITERION 2" in upper:
        all_issues.append(("S0", "EB-1A Criterion sections found — this is EB-2 NIW"))

    # ═══════════════════════════════════════════════════════════
    # REPORT
    # ═══════════════════════════════════════════════════════════
    print("\n" + "=" * 72)
    print("  VALIDATION REPORT")
    print("=" * 72)

    # Stats
    print(f"\n  Document Stats:")
    print(f"    Paragraphs:      {len(doc.paragraphs)}")
    print(f"    Tables:          {len(doc.tables)}")
    print(f"    Images:          {img_count}")
    print(f"    Evidence Blocks: {tables_2col} ({tables_impact} with impact)")
    print(f"    Word Count:      {word_count:,}")
    print(f"    Runs Scanned:    {total_runs}")

    # Categorize
    s0 = [(s, m) for s, m in all_issues if s == "S0"]
    s1 = [(s, m) for s, m in all_issues if s == "S1"]
    s2 = [(s, m) for s, m in all_issues if s == "S2"]
    s3 = [(s, m) for s, m in all_issues if s == "S3"]

    print(f"\n  Issue Summary:")
    print(f"    S0 CRITICAL:  {len(s0)}")
    print(f"    S1 SEVERE:    {len(s1)}")
    print(f"    S2 MODERATE:  {len(s2)}")
    print(f"    S3 MINOR:     {len(s3)}")
    print(f"    TOTAL:        {len(all_issues)}")

    if s0:
        print(f"\n  {'!'*60}")
        print(f"  S0 — CRITICAL (BLOCKS DELIVERY)")
        print(f"  {'!'*60}")
        for _, msg in s0:
            print(f"    [S0] {msg}")

    if s1:
        print(f"\n  {'!'*50}")
        print(f"  S1 — SEVERE (BLOCKS DELIVERY)")
        print(f"  {'!'*50}")
        for _, msg in s1:
            print(f"    [S1] {msg}")

    if s2:
        print(f"\n  {'='*40}")
        print(f"  S2 — MODERATE (FIX BEFORE DELIVERY)")
        print(f"  {'='*40}")
        for _, msg in s2:
            print(f"    [S2] {msg}")

    if s3:
        print(f"\n  {'-'*30}")
        print(f"  S3 — MINOR (DOCUMENT)")
        print(f"  {'-'*30}")
        for _, msg in s3:
            print(f"    [S3] {msg}")

    # Verdict
    print(f"\n{'=' * 72}")
    if s0:
        print(f"  VERDICT: FAILED — {len(s0)} CRITICAL issues block delivery")
        print(f"  ACTION:  REBUILD required. Fix S0 issues immediately.")
    elif s1:
        print(f"  VERDICT: FAILED — {len(s1)} SEVERE issues block delivery")
        print(f"  ACTION:  Fix all S1 issues before delivery.")
    elif s2:
        print(f"  VERDICT: CONDITIONAL PASS — {len(s2)} moderate issues")
        print(f"  ACTION:  Fix S2 issues if possible before delivery.")
    else:
        print(f"  VERDICT: PASSED — Document ready for delivery")
    print(f"{'=' * 72}")

    return len(s0), len(s1), len(s2), len(s3)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 validate_resume_eb2niw.py <resume.docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    if not os.path.exists(docx_path):
        print(f"ERROR: File not found: {docx_path}")
        sys.exit(1)

    s0, s1, s2, s3 = run_full_validation(docx_path)
    sys.exit(1 if (s0 + s1) > 0 else 0)
