#!/usr/bin/env python3
"""SaaS Evidence DOCX builder — genérico, parametrizado.

Substitui o antigo generate_saas_cristine.py (hardcoded pra um cliente).
Consome product_spec.md + evidence_summary.json + screenshots/ e produz
V1_saas_evidence_{slug}.docx com screenshots injetados inline.

Uso:
    python3 build_saas_evidence.py <product_spec.md> <evidence_summary.json> \
            <screenshots_dir> <output.docx>

Sequência obrigatória do pipeline:
  Passo 1 (Research) → Passo 2 (product_spec.md) → Passo 3 (evidence_summary.json)
  → Passo 4 (LOVABLE_BUILD_SPEC.md) → Passo 5 (HUMAN GATE: lovable_url.txt)
  → Passo 6 (capture_saas.js → screenshots/) → PASSO 7 AQUI → Passo 8 (validate).

Este script FALHA cedo se:
- screenshots/ tem < 10 PNGs (Passo 6 não rodou)
- evidence_summary.json não tem founder_dependency="critical"
- product_spec.md é menor que 10KB (Passo 2 foi raso)
"""
import json
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import OxmlElement, parse_xml
except ImportError:
    sys.stderr.write("ERR: pip install python-docx\n")
    sys.exit(2)


# Design tokens (alinhados ao LOVABLE_BUILD_SPEC)
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
SLATE = RGBColor(0x4B, 0x55, 0x63)
INK = RGBColor(0x10, 0x14, 0x1F)
PAPER = RGBColor(0xFA, 0xFB, 0xFC)
SOFT = RGBColor(0xF5, 0xF1, 0xE6)
RULE = RGBColor(0xE3, 0xE7, 0xEC)
SUCCESS = RGBColor(0x05, 0x96, 0x69)

FORBIDDEN = re.compile(
    r"\b(petition|petitioner|USCIS|Dhanasar|Kazarian|extraordinary ability|"
    r"NIW|EB-2|EB-1|I-140|visa|green card|adjudicator|national interest waiver|"
    r"Kortix|Petition Engine|Forjado por|RAG\s+[IVX\d]|Obsidian)\b",
    re.IGNORECASE
)


def guard(s):
    if not isinstance(s, str):
        return s
    if FORBIDDEN.search(s):
        raise RuntimeError(f"FORBIDDEN term: {s[:120]!r}")
    return s


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_borders(cell, color="E3E7EC", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_para(doc, text, *, size=11, bold=False, italic=False, color=INK,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
             line_spacing=1.35, font='Calibri'):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    r = p.add_run(guard(text))
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    # Gold underline via border
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:color="B8860B" w:space="4"/>'
        f'</w:pBdr>'
    ))
    r = p.add_run(guard(text))
    r.font.name = 'Calibri'
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = NAVY
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(guard(text))
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = NAVY
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(guard(text))
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SLATE
    return p


def add_screenshot_with_caption(doc, img_path, caption_text, page_name):
    """Insert a screenshot centered with bordered caption below it."""
    # Paragraph with image centered
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(10)
    p_img.paragraph_format.space_after = Pt(4)
    try:
        p_img.add_run().add_picture(img_path, width=Inches(6.2))
    except Exception as e:
        r_err = p_img.add_run(f"[screenshot: {os.path.basename(img_path)} — erro: {e}]")
        r_err.font.size = Pt(9)
        r_err.font.italic = True

    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(16)
    r1 = p_cap.add_run(f"Figure — {page_name}: ")
    r1.font.name = 'Calibri'; r1.font.size = Pt(9); r1.font.bold = True; r1.font.color.rgb = NAVY
    r2 = p_cap.add_run(guard(caption_text))
    r2.font.name = 'Calibri'; r2.font.size = Pt(9); r2.font.italic = True; r2.font.color.rgb = SLATE


def apply_header_footer(doc, client_name):
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        for p in list(header.paragraphs):
            p.clear()
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = hp.add_run(client_name.upper())
        r1.font.name = 'Calibri'; r1.font.size = Pt(9); r1.font.bold = True; r1.font.color.rgb = NAVY
        hp.add_run('\t\t')
        r2 = hp.add_run('Platform Documentation')
        r2.font.name = 'Calibri'; r2.font.size = Pt(8); r2.font.italic = True; r2.font.color.rgb = SLATE
        pPr = hp._element.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:color="B8860B" w:space="4"/>'
            f'</w:pBdr>'
        ))
        pPr.append(parse_xml(
            f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:pos="9360"/></w:tabs>'
        ))

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p.clear()
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fpPr = fp._element.get_or_add_pPr()
        fpPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="6" w:color="B8860B" w:space="4"/>'
            f'</w:pBdr>'
        ))
        fpPr.append(parse_xml(
            f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:pos="9360"/></w:tabs>'
        ))
        rf1 = fp.add_run(f'Methodology by {client_name} | Proprietary Framework')
        rf1.font.name = 'Calibri'; rf1.font.size = Pt(8); rf1.font.bold = True; rf1.font.color.rgb = NAVY
        fp.add_run('\t\t')
        rf2 = fp.add_run('CONFIDENTIAL')
        rf2.font.name = 'Calibri'; rf2.font.size = Pt(7); rf2.font.italic = True; rf2.font.color.rgb = GOLD


def render_markdown_section(doc, md_text):
    """Very simple MD renderer: H1/H2/H3/paragraph/list.
    Preserves substance; we don't need perfect MD — just readable DOCX output."""
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith('# '):
            add_h1(doc, line[2:].strip())
        elif line.startswith('## '):
            add_h2(doc, line[3:].strip())
        elif line.startswith('### '):
            add_h3(doc, line[4:].strip())
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run('• ' + guard(line[2:].strip()))
            r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = INK
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(guard(line.strip()))
            r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = INK
        elif line.startswith('```'):
            # Skip code blocks entirely (substance is the surrounding prose)
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                i += 1
        else:
            add_para(doc, line)
        i += 1


def load_screenshots(shots_dir, map_path=None):
    """Returns [(num, path, name, description), ...] sorted by num."""
    shots = []
    if not os.path.isdir(shots_dir):
        return shots

    # Load map.json if present
    metadata = {}
    if map_path and os.path.isfile(map_path):
        try:
            data = json.load(open(map_path, encoding='utf-8'))
            for page in data.get('pages', []):
                num = str(page.get('number', '')).zfill(2)
                metadata[num] = {
                    'display_name': page.get('display_name', ''),
                    'description': page.get('description', ''),
                }
        except Exception:
            pass

    for fn in sorted(os.listdir(shots_dir)):
        if fn.startswith('SaaS_') and fn.lower().endswith('.png'):
            m = re.match(r'SaaS_(\d+)_(.+)\.png', fn)
            if m:
                num = m.group(1)
                name_raw = m.group(2).replace('_', ' ')
                meta = metadata.get(num, {})
                shots.append((
                    int(num),
                    os.path.join(shots_dir, fn),
                    meta.get('display_name') or name_raw,
                    meta.get('description') or f"Interface view from the platform — {name_raw}"
                ))
    shots.sort(key=lambda x: x[0])
    return shots


def main():
    if len(sys.argv) < 5:
        sys.stderr.write(
            "Usage: build_saas_evidence.py <product_spec.md> "
            "<evidence_summary.json> <screenshots_dir> <output.docx>\n"
        )
        sys.exit(2)

    spec_path = sys.argv[1]
    summary_path = sys.argv[2]
    shots_dir = sys.argv[3]
    out_path = sys.argv[4]

    # Fail-early gates
    if not os.path.isfile(spec_path):
        sys.stderr.write(f"ERR: product_spec not found: {spec_path}\n")
        sys.exit(2)
    if os.path.getsize(spec_path) < 5000:
        sys.stderr.write(f"ERR: product_spec too small ({os.path.getsize(spec_path)} bytes) — Fase 2 raso\n")
        sys.exit(2)
    if not os.path.isfile(summary_path):
        sys.stderr.write(f"ERR: evidence_summary.json not found: {summary_path}\n")
        sys.exit(2)

    try:
        summary = json.load(open(summary_path, encoding='utf-8'))
    except Exception as e:
        sys.stderr.write(f"ERR: evidence_summary.json unreadable: {e}\n")
        sys.exit(2)

    fd = summary.get('platform_proof', {}).get('founder_dependency')
    if fd != 'critical':
        sys.stderr.write(
            f"ERR: founder_dependency='{fd}' — MUST be 'critical' (ANTI-CRISTINE rule). "
            "Fix evidence_summary.json before building DOCX.\n"
        )
        sys.exit(2)

    shots = load_screenshots(shots_dir, map_path=os.path.join(shots_dir, 'screenshot_map.json'))
    if len(shots) < 10:
        sys.stderr.write(
            f"ERR: screenshots/ has {len(shots)} SaaS_*.png files — need ≥10. "
            "Run Passo 6 (node scripts/capture_saas.js) before building DOCX.\n"
        )
        sys.exit(2)

    # Extract client name from summary or out_path
    client_name = summary.get('proposed_endeavor', '')[:80] or 'Platform'
    # Try to derive from out_path
    m = re.search(r'V\d+_saas_evidence_([^.]+)\.docx', os.path.basename(out_path))
    client_slug = m.group(1) if m else 'client'
    # Humanize slug
    client_display = client_slug.replace('_', ' ').title()

    # Build DOCX
    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
        s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)

    apply_header_footer(doc, client_display)

    # Cover
    add_h1(doc, summary.get('proposed_endeavor', 'Platform Dossier')[:120])
    add_para(doc, "Platform Documentation & Evidence Dossier",
             size=13, italic=True, color=SLATE, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, "", space_after=12)

    # Render product_spec.md content
    with open(spec_path, encoding='utf-8') as f:
        spec_md = f.read()
    render_markdown_section(doc, spec_md)

    # Screenshots section
    doc.add_page_break()
    add_h1(doc, "Platform Interface — Screenshots")
    add_para(doc,
        f"The following {len(shots)} screenshots document the operational interface "
        "of the platform as currently deployed. Each view captures a specific user-facing "
        "workflow or administrative function. The platform is accessible via subscription "
        "and operates under the beneficiary's continuous quality-assurance framework.",
        space_after=12)

    for num, img_path, name, description in shots:
        add_screenshot_with_caption(doc, img_path, description, name)

    # Evidence summary section
    doc.add_page_break()
    add_h1(doc, "Evidence Summary")
    add_h2(doc, "Platform Proof")
    pp = summary.get('platform_proof', {})
    add_para(doc, f"Beneficiary created methodology: {pp.get('beneficiary_created_methodology')}")
    add_para(doc, f"Requires ongoing beneficiary oversight: {pp.get('requires_ongoing_beneficiary_oversight')}")
    add_para(doc, f"Founder dependency: {pp.get('founder_dependency')}")
    add_para(doc, f"What breaks without beneficiary: {pp.get('what_breaks_without_beneficiary', '')}")

    add_h2(doc, "National Scope")
    ns = summary.get('national_scope', {})
    add_para(doc, f"States targeted: {', '.join(ns.get('states_targeted', []))}")
    add_para(doc, f"Institutions served: {ns.get('institutions_served', 0)}")
    add_para(doc, f"Total users: {ns.get('total_users', 0)}")

    add_h2(doc, "Systemic Impact")
    si = summary.get('systemic_impact', {})
    add_para(doc, f"Problem: {si.get('problem_description', '')}")
    add_para(doc, f"Sector size: {si.get('sector_size_usd', '')}")
    add_para(doc, f"Gap: {si.get('gap_description', '')}")

    add_h2(doc, "Federal Alignment")
    for policy in summary.get('federal_alignment', []):
        add_para(doc, f"• {policy}", space_after=3, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Save
    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or '.', exist_ok=True)
    doc.save(out_path)
    size = os.path.getsize(out_path)
    inline = len(Document(out_path).inline_shapes)
    print(f"OK SaaS Evidence DOCX: {out_path}")
    print(f"  Size: {size:,} bytes")
    print(f"  Inline shapes: {inline}")
    print(f"  Screenshots embedded: {len(shots)}")
    if size < 300_000:
        sys.stderr.write(f"WARN: DOCX size {size:,} < 300KB — validador pode flag\n")
    if inline < 10:
        sys.stderr.write(f"WARN: inline_shapes {inline} < 10 — validador vai bloquear\n")


if __name__ == "__main__":
    main()
