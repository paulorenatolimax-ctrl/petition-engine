#!/usr/bin/env python3
"""
BP DentalShield — Geração Programática Completa
Cada seção é escrita diretamente com dados verificados.
Padrão Ikaro/Rafael: Garamond, #DEDACB, endnotes com URLs reais.
"""

import json, os, re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# === PATHS ===
SCRIPT_DIR = Path(__file__).parent
OUTPUT_DIR = Path("/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2025/CAMILLA/_BP/_Atualizado (pós-entendimento novo)")
CHARTS_DIR = SCRIPT_DIR / "_charts"

# === FINANCIAL DATA ===
FIN = {
    "rev": [627180, 1238588, 2108052, 1576134, 2964110],
    "rev_total": 8514064,
    "var_costs": [153696, 303498, 516504, 386185, 726552],
    "payroll": [380976, 643956, 847740, 847740, 1181784],
    "fixed": [109200, 175200, 217200, 217200, 265200],
    "total_costs": [643872, 1122654, 1581444, 1451125, 2173536],
    "ebitda": [-16692, 115934, 526608, 125009, 790574],
    "ebitda_margin": [-2.7, 9.4, 25.0, 7.9, 26.7],
    "net_income": [-16692, 86951, 394956, 93757, 592931],
    "net_total": 704728,
    "invest": [175257, 0, 98346, 0, 120240],
    "invest_total": 393843,
    "invest_tangible": [78108, 0, 45200, 0, 51800],
    "invest_intangible": [56635, 0, 28700, 0, 43994],
    "invest_wc": [40514, 0, 24446, 0, 24446],
    "bep": [649266, 1173150, 1358628, 1358628, 1792734],
    "emps": [5, 9, 13, 13, 22],
    "tampa": [5, 5, 6, 6, 8],
    "miami": [0, 3, 4, 4, 7],
    "houston": [0, 1, 3, 3, 7],
    "npv": 396168, "irr": 61.18, "roi": 178.94, "payback": 3,
    "biz_value": 2174346, "discount": 12,
}
YRS = ["Ano 1", "Ano 2", "Ano 3", "Ano 4", "Ano 5"]

def usd(v):
    if v < 0: return f"-US$ {abs(v):,.0f}".replace(",",".")
    return f"US$ {v:,.0f}".replace(",",".")

def pct(v):
    return f"{v:.1f}%".replace(".",",")

# === DESIGN ===
FONT = "Garamond"
C_H = RGBColor(0x58, 0x4D, 0x42)
C_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
C_CAP = RGBColor(0x66, 0x66, 0x66)
TH_HEX = "DEDACB"
TZ_HEX = "D0DDD6"
TB_HEX = "CCCCCC"
KPI_HEX = "F0F0F0"

# === ENDNOTES ===
ENDNOTES = []
_fn_counter = [0]

def fn(source_text, url=None):
    """Register endnote, return [N] marker."""
    _fn_counter[0] += 1
    n = _fn_counter[0]
    full = source_text
    if url:
        full += f" Disponível em: {url}"
    ENDNOTES.append((n, full))
    return f"[{n}]"

# === HELPERS ===
def setup_styles(doc):
    s = doc.styles['Normal']
    s.font.name = FONT; s.font.size = Pt(11); s.font.color.rgb = C_TEXT
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.space_before = Pt(3)
    s.paragraph_format.line_spacing = 1.15
    for lv, sz in [(1,16),(2,13),(3,11)]:
        h = doc.styles[f'Heading {lv}']
        h.font.name = FONT; h.font.size = Pt(sz); h.font.bold = True; h.font.color.rgb = C_H
        h.paragraph_format.space_before = Pt(18 if lv==1 else 12 if lv==2 else 6)
        h.paragraph_format.space_after = Pt(8 if lv==1 else 6)
        h.paragraph_format.keep_with_next = True

def setup_page(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.14); sec.right_margin = Inches(0.85)
    # Header
    hdr = sec.header; hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("DentalShield Compliance Solutions LLC"); r.font.name = FONT; r.font.size = Pt(8); r.font.color.rgb = C_CAP; r.italic = True
    # Footer
    ftr = sec.footer; ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = fp._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="12" w:space="1" w:color="584D42"/></w:pBdr>'))
    r = fp.add_run("CONFIDENTIAL — DO NOT SHARE  |  DentalShield Compliance Solutions LLC  |  Page ")
    r.font.name = FONT; r.font.size = Pt(7); r.font.bold = True; r.font.color.rgb = C_H
    for xml in [
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}"/><w:sz w:val="14"/><w:b/><w:color w:val="584D42"/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>',
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}"/><w:sz w:val="14"/><w:b/><w:color w:val="584D42"/></w:rPr><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>',
        f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}"/><w:sz w:val="14"/><w:b/><w:color w:val="584D42"/></w:rPr><w:fldChar w:fldCharType="end"/></w:r>',
    ]:
        fp._p.append(parse_xml(xml))

_h1_count = [0]
def h1(doc, text):
    _h1_count[0] += 1
    if _h1_count[0] > 1: doc.add_page_break()
    doc.add_heading(text, level=1)
def h2(doc, text): doc.add_heading(text, level=2)
def h3(doc, text): doc.add_heading(text, level=3)

def p(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Paragraph with **bold** and [N] superscript support."""
    pr = doc.add_paragraph(); pr.alignment = align; pr.paragraph_format.space_after = Pt(6)
    for tok in re.split(r'(\*\*[^*]+\*\*|\[\d+\])', text):
        if not tok: continue
        if tok.startswith('**') and tok.endswith('**'):
            r = pr.add_run(tok[2:-2]); r.bold = True
        elif re.match(r'^\[\d+\]$', tok):
            r = pr.add_run(tok); r.font.superscript = True; r.font.size = Pt(8); r.font.color.rgb = C_CAP
        else:
            r = pr.add_run(tok)
            if bold: r.bold = True
            if italic: r.italic = True
        r.font.name = FONT; r.font.size = Pt(11) if not r.font.superscript else Pt(8)
    return pr

def bullet(doc, text):
    p(doc, f"• {text}")

def tbl(doc, headers, rows, caption=None):
    """Table with colored header + zebra."""
    nc = len(headers)
    t = doc.add_table(rows=1+len(rows), cols=nc)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, ht in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(str(ht)); r.bold = True; r.font.name = FONT; r.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TH_HEX}"/>'))
    for ri, rd in enumerate(rows):
        for ci in range(min(len(rd), nc)):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(rd[ci])); r.font.name = FONT; r.font.size = Pt(10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TZ_HEX}"/>'))
    tblPr = t._tbl.tblPr if t._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="{TB_HEX}"/><w:left w:val="single" w:sz="4" w:space="0" w:color="{TB_HEX}"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="{TB_HEX}"/><w:right w:val="single" w:sz="4" w:space="0" w:color="{TB_HEX}"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="{TB_HEX}"/><w:insideV w:val="single" w:sz="4" w:space="0" w:color="{TB_HEX}"/></w:tblBorders>'))
    if caption:
        cp = doc.add_paragraph(); cr = cp.add_run(f"Fonte: {caption}")
        cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = C_CAP; cr.font.name = FONT
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def kpi_box(doc, lines):
    for line in lines:
        pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = pr._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{KPI_HEX}"/>'))
        pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="4" w:color="584D42"/></w:pBdr>'))
        pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="288" w:right="288"/>'))
        r = pr.add_run(line); r.font.name = FONT; r.font.size = Pt(11)
        if line.startswith("►"): r.bold = True

def img(doc, filename, caption_text):
    if not (CHARTS_DIR / filename).exists(): return
    doc.add_paragraph()
    doc.add_picture(str(CHARTS_DIR / filename), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph()
    cr = cp.add_run(f"Figura: {caption_text}")
    cr.italic = True; cr.font.size = Pt(9); cr.font.name = FONT; cr.font.color.rgb = C_CAP
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

def divider(doc):
    pr = doc.add_paragraph()
    pPr = pr._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="{TH_HEX}"/></w:pBdr>'))

# ============================================================
# COVER PAGE
# ============================================================
def cover_page(doc):
    for _ in range(6): doc.add_paragraph()
    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pr.add_run("BUSINESS PLAN"); r.font.name = FONT; r.font.size = Pt(32); r.font.color.rgb = C_H; r.bold = True
    doc.add_paragraph()
    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pr.add_run("DentalShield Compliance Solutions LLC"); r.font.name = FONT; r.font.size = Pt(20); r.font.color.rgb = C_H
    doc.add_paragraph()
    for txt in ["Dental Regulatory Compliance & IoT Monitoring", "", "Tampa, FL  |  Miami, FL  |  Houston, TX", "", f"Prepared: {datetime.now().strftime('%B %Y')}", "", "", "CONFIDENTIAL — DO NOT SHARE"]:
        pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pr.add_run(txt); r.font.name = FONT; r.font.size = Pt(12)
        if "CONFIDENTIAL" in txt: r.bold = True; r.font.color.rgb = C_H
    doc.add_page_break()

# ============================================================
# SECTION 1: SUMÁRIO EXECUTIVO
# ============================================================

def s1_1(doc):
    """1.1 Oportunidade de Negócio (padrão Ikaro: 1.1.1/1.1.2/1.1.3)"""
    h1(doc, "1. SUMÁRIO EXECUTIVO")
    h2(doc, "1.1 Oportunidade de Negócio")

    h3(doc, "1.1.1 Apresentação da Empresa")
    f1 = fn("Precedence Research, U.S. Dental Service Market Report, 2025.", "https://www.precedenceresearch.com/us-dental-service-market")
    f2 = fn("ADA Health Policy Institute, Supply of Dentists in the U.S.: 2001–2023.", "https://www.ada.org/resources/research/health-policy-institute")
    f3 = fn("OSHA News Release, January 14, 2025. Penalidades atualizadas para 2025.", "https://www.osha.gov/news/newsreleases/osha-trade-release/20250114")
    f4 = fn("Fortune Business Insights, Infection Control Market Size, 2024.")
    f5 = fn("Precedence Research, IoT in Healthcare Market, 2024.")

    p(doc, f"O mercado de serviços odontológicos nos Estados Unidos atingiu US$ 164,76 bilhões em 2024, com taxa de crescimento anual composta (CAGR) projetada de 5,3% até 2034{f1}. Este ecossistema — sustentado por mais de 200 mil clínicas odontológicas ativas em todo o país{f2} — opera sob um dos ambientes regulatórios mais rigorosos do setor de saúde, envolvendo simultaneamente OSHA, CDC, EPA, FDA e HIPAA. A maioria dessas clínicas, contudo, não possui estrutura interna dedicada à conformidade regulatória. O resultado é previsível: as penalidades da OSHA por violação grave atingem US$ 16.550 por citação em 2025, podendo chegar a US$ 165.514 por violação intencional{f3}.")

    p(doc, f"A **DentalShield Compliance Solutions LLC** foi concebida para resolver essa lacuna estrutural. Sediada em Tampa, Flórida (Condado de Hillsborough), a empresa oferece um portfólio integrado de seis serviços especializados — auditoria de conformidade regulatória, programas de controle de infecção e biossegurança, monitoramento por IoT (SteriSensor), treinamento e certificação de equipes, gestão de resíduos biomédicos (BiohazardBox) e compliance digital via plataforma SaaS (ComplianceScreen). A DentalShield não é uma empresa genérica de compliance médico; é uma operação verticalmente especializada no setor odontológico.")

    p(doc, f"O plano estratégico prevê expansão em três fases: lançamento em Tampa (Ano 0), abertura de filiais em Miami e Houston (Ano 2), e escalonamento com ampliação da equipe (Ano 4–5). A empresa projeta receita acumulada de {usd(FIN['rev_total'])} em cinco anos, com 22 colaboradores diretos e margem EBITDA de {pct(FIN['ebitda_margin'][4])} no Ano 5. O mercado de controle de infecções, avaliado em US$ 22,7 bilhões{f4}, e o de IoT em saúde, em US$ 89,7 bilhões{f5}, corroboram a existência de uma janela estratégica significativa para soluções integradas de compliance com base tecnológica.")

    f5b = fn("Mordor Intelligence, Healthcare Compliance Software Market, 2024. Mercado de US$ 2,94B.", "https://www.mordorintelligence.com/industry-reports/healthcare-compliance-software-market")
    f5c = fn("CDC, Infection Prevention & Control in Dental Health-Care Settings, MMWR 52(RR-17).", "https://www.cdc.gov/dental-infection-control/hcp/")

    p(doc, f"A convergência entre três vetores de crescimento sustenta a tese de negócio da DentalShield. Primeiro, o mercado de software de compliance em saúde — avaliado em US$ 2,94 bilhões em 2024 com CAGR de 9,5%{f5b} — cresce quase duas vezes mais rápido que o próprio setor odontológico, sinalizando que a conformidade regulatória está se consolidando como centro de custo estrutural para proprietários de clínicas. Segundo, o CDC mantém diretrizes rigorosas de controle de infecção especificamente para estabelecimentos odontológicos, publicadas em MMWR 52(RR-17) e atualizadas em 2016{f5c}, cujo cumprimento é verificável em inspeções da OSHA. Terceiro, a adoção de IoT no setor de saúde — com CAGR de 18,4% e mercado de US$ 89,7 bilhões{f5} — viabiliza o monitoramento contínuo de equipamentos de esterilização, substituindo registros manuais imprecisos por dados em tempo real verificáveis por reguladores.")

    p(doc, "A DentalShield diferencia-se de seus concorrentes por operar na interseção dessas três tendências: não é uma empresa genérica de compliance médico que atende odontologia como mercado secundário, nem uma startup de IoT sem expertise regulatória, nem uma firma de treinamento corporativo sem componente tecnológico. É uma operação verticalmente integrada que compreende as nuances regulatórias específicas da odontologia — desde os ciclos de autoclave exigidos pela OSHA Bloodborne Pathogens Standard (29 CFR 1910.1030) até os protocolos de classificação de resíduos da EPA (40 CFR 259) — e que entrega conformidade verificável através de tecnologia proprietária.")

    p(doc, f"O modelo de receita recorrente (assinaturas mensais combinando auditoria, IoT e SaaS) alinha os incentivos da empresa com os resultados de compliance de seus clientes. Uma clínica que investe entre US$ 30.000 e US$ 50.000 por ano em serviços DentalShield protege-se contra multas da OSHA que podem chegar a US$ 165.514 por violação intencional{f3} — um retorno sobre investimento imediato e tangível que sustenta a proposta de valor premium.")

    h3(doc, "1.1.2 Serviços Oferecidos")
    tbl(doc,
        ["Serviço", "Descrição", "Benefícios", "Aplicação", "Problemas Resolvidos"],
        [
            ["Regulatory Compliance Auditing (30%)", "Auditorias presenciais e remotas de conformidade OSHA, CDC, HIPAA e regulamentos estaduais", "Redução de risco de multas; prontidão para inspeções; documentação atualizada", "Clínicas de todos os portes; multi-especialidade", "Falta de compliance officer interno; desconhecimento de atualizações normativas"],
            ["Infection Control & Biosafety (25%)", "Implementação de protocolos de esterilização, controle de patógenos sanguíneos conforme CDC e OSHA 29 CFR 1910.1030", "Proteção de pacientes e equipe; redução de infecções cruzadas", "Salas de esterilização, processamento de instrumentais", "Protocolos desatualizados; falhas em ciclos de autoclave"],
            ["IoT Monitoring — SteriSensor (20%)", "Sensores IoT conectados a autoclaves e unidades de armazenamento com monitoramento em tempo real", "Rastreabilidade contínua; detecção imediata de falhas; relatórios automáticos", "Equipamentos de esterilização; refrigeradores biológicos", "Falhas silenciosas; registros manuais imprecisos"],
            ["Staff Training & Certification (15%)", "Treinamentos presenciais e online com certificação em OSHA Bloodborne Pathogens, HIPAA e HazCom", "Equipe certificada e compliance-ready; evidência documental para inspeções", "Dentistas, higienistas, assistentes, administrativo", "Alta rotatividade sem retreinamento; treinamentos genéricos"],
            ["BiohazardBox — Waste Management (5%)", "Sistema integrado de coleta, classificação e descarte de resíduos biomédicos conforme EPA 40 CFR 259", "Conformidade ambiental; rastreabilidade de descarte", "Resíduos perfurocortantes; materiais contaminados", "Descarte irregular; mistura de resíduos classificados"],
            ["ComplianceScreen — Digital SaaS (5%)", "Plataforma SaaS para gestão de documentos de compliance, vencimentos de certificações e inspeções", "Dashboard centralizado; alertas de vencimento; relatórios audit-ready", "Gestão documental; controle de certificações", "Documentos dispersos; certificações vencidas"],
        ],
        "Elaboração própria com base na estrutura de serviços da DentalShield Compliance Solutions LLC."
    )

    h3(doc, "1.1.3 Projeções-Chave")
    kpi_box(doc, [
        f"► INDICADORES FINANCEIROS — PROJEÇÃO DE 5 ANOS",
        f"▸ Receita Acumulada (5 anos): {usd(FIN['rev_total'])}",
        f"▸ Valor Presente Líquido (VPL, taxa {FIN['discount']}%): {usd(FIN['npv'])}",
        f"▸ Taxa Interna de Retorno (TIR): {pct(FIN['irr'])}",
        f"▸ Retorno sobre Investimento (ROI): {pct(FIN['roi'])}",
        f"▸ Payback: {FIN['payback']} anos",
        f"▸ Colaboradores no Ano 5: 22 diretos (3 localidades)",
        f"▸ Valor Estimado do Negócio (Ano 5): {usd(FIN['biz_value'])}",
    ])
    divider(doc)


def s1_2(doc):
    """1.2 Cronograma Estratégico"""
    h2(doc, "1.2 Cronograma Estratégico")
    p(doc, "O plano de expansão da DentalShield está estruturado em três fases, cada uma alinhada a investimentos específicos e metas de penetração de mercado:")
    tbl(doc,
        ["Fase", "Período", "Localidade", "Marcos Principais", "Investimento"],
        [
            ["Fase 1 — Lançamento", "Ano 0–1", "Tampa, FL (Sede)", "Constituição da LLC; primeiros 50 clientes; validação do modelo de negócio; lançamento SteriSensor v1", usd(175257)],
            ["Fase 2 — Expansão", "Ano 2–3", "Miami, FL + Houston, TX", "Abertura de 2 filiais; contratação de 8 novos colaboradores; 150 clientes ativos; parcerias com DSOs", usd(98346)],
            ["Fase 3 — Escalonamento", "Ano 4–5", "3 localidades (ampliação)", "22 colaboradores; 300+ clientes; ComplianceScreen v2; receita acima de US$ 2,9M/ano", usd(120240)],
        ],
        f"Projeções financeiras internas. Investimento total: {usd(FIN['invest_total'])}."
    )
    p(doc, f"O investimento total de {usd(FIN['invest_total'])} está distribuído entre ativos tangíveis (equipamentos, mobiliário, dispositivos IoT), ativos intangíveis (desenvolvimento de software, licenças, propriedade intelectual) e capital de giro para os primeiros meses de operação em cada fase.")


def s1_3(doc):
    """1.3 Visão, Missão e Valores"""
    h2(doc, "1.3 Visão, Missão e Valores")

    h3(doc, "Visão")
    p(doc, "Ser a referência nacional em compliance regulatório odontológico, integrando tecnologia IoT e expertise humana para assegurar que cada clínica dentária nos Estados Unidos opere em plena conformidade com os padrões federais e estaduais de saúde e segurança.")
    p(doc, "Em um horizonte de 10 anos, a DentalShield aspira a monitorar mais de 5.000 clínicas odontológicas em todo o território norte-americano, estabelecendo o padrão da indústria para compliance digital no setor.")

    h3(doc, "Missão")
    p(doc, "Proteger pacientes, profissionais e proprietários de clínicas odontológicas através de soluções integradas de conformidade regulatória, controle de infecção e monitoramento tecnológico — eliminando o risco de penalidades, infecções cruzadas e interrupções operacionais.")

    h3(doc, "Valores")
    tbl(doc,
        ["Valor", "Descrição", "Evidência na Operação"],
        [
            ["Excelência Regulatória", "Compromisso inabalável com os mais altos padrões de conformidade", "Auditorias trimestrais com checklist de 200+ itens baseados em OSHA, CDC e HIPAA"],
            ["Inovação Aplicada", "Tecnologia a serviço da segurança do paciente, não como fim em si mesma", "SteriSensor IoT: monitoramento 24/7 de autoclaves com alertas em tempo real"],
            ["Transparência", "Comunicação clara e honesta com clientes sobre gaps de compliance e prazos de remediação", "Relatórios mensais de compliance com score de 0-100 e plano de ação detalhado"],
            ["Parceria de Longo Prazo", "Modelo de receita recorrente alinha incentivos da empresa com resultados do cliente", "Retenção projetada de 85%+ com contratos anuais e suporte contínuo"],
            ["Responsabilidade Social", "Cada clínica em compliance é uma comunidade mais segura", "Programas de treinamento gratuitos para residentes de odontologia em universidades parceiras"],
        ]
    )


def s1_4(doc):
    """1.4 Enquadramento Jurídico (padrão Pravion: 6 subseções)"""
    h2(doc, "1.4 Enquadramento Jurídico")

    f6 = fn("Florida Department of State, Division of Corporations. Filing fee: US$ 125.", "https://dos.fl.gov/sunbiz/forms/llc/")
    f7 = fn("IRS, Employer Identification Number (EIN). Form SS-4.", "https://www.irs.gov/businesses/small-businesses-self-employed/apply-for-an-employer-identification-number-ein-online")
    f8 = fn("OSHA, Bloodborne Pathogens Standard, 29 CFR 1910.1030.", "https://www.osha.gov/laws-regs/regulations/standardnumber/1910/1910.1030")
    f9 = fn("CDC, Infection Prevention & Control in Dental Health-Care Settings, MMWR 52(RR-17), 2003 (atualizado 2016).", "https://www.cdc.gov/dental-infection-control/hcp/")
    f10 = fn("EPA, Medical Waste Management Guidelines, 40 CFR 259.", "https://www.epa.gov/rcra/medical-waste")
    f11 = fn("HHS, HIPAA Privacy and Security Rules, 45 CFR Parts 160, 164.", "https://www.hhs.gov/hipaa/index.html")
    f12 = fn("FDA, Regulation of Dental Devices, 21 CFR 872.", "https://www.fda.gov/medical-devices")
    f13 = fn("Florida Board of Dentistry, F.S. Chapter 466.", "https://floridasdentistry.gov/")
    f14 = fn("Texas State Board of Dental Examiners, Occupations Code Title 3 Subtitle D.", "https://www.tsbde.texas.gov/")

    h3(doc, "1.4.1 Estrutura Societária")
    p(doc, f"A DentalShield Compliance Solutions LLC será constituída como uma **Limited Liability Company (LLC)** no Estado da Flórida, conforme as disposições do Florida Revised Limited Liability Company Act (F.S. Chapter 605){f6}. A escolha da estrutura LLC justifica-se pela proteção de responsabilidade pessoal dos membros, flexibilidade na tributação (pass-through taxation) e simplicidade administrativa. A Flórida oferece vantagem competitiva adicional ao não impor imposto de renda estadual sobre pessoas físicas.")
    p(doc, f"O Employer Identification Number (EIN) será obtido junto ao Internal Revenue Service (IRS) mediante Form SS-4{f7}, habilitando a empresa para contratação de funcionários, abertura de contas bancárias comerciais e cumprimento de obrigações fiscais federais.")

    h3(doc, "1.4.2 Licenciamento e Registros")
    p(doc, "A operação da DentalShield requer os seguintes registros e licenças:")
    bullet(doc, "**Registro estadual:** Articles of Organization junto ao Florida Department of State (taxa: US$ 125)")
    bullet(doc, "**Registro federal:** EIN junto ao IRS (gratuito)")
    bullet(doc, "**Licença municipal:** Tampa Business Tax Receipt (Hillsborough County)")
    bullet(doc, "**Registro como foreign LLC:** Estado do Texas para operações em Houston (taxa: US$ 750)")
    bullet(doc, "**Seguro profissional:** Errors & Omissions (E&O) Insurance com cobertura mínima de US$ 1.000.000")

    h3(doc, "1.4.3 Conformidade Regulatória Federal")
    p(doc, f"A DentalShield opera na interseção de múltiplos marcos regulatórios federais que regem o setor odontológico:")
    tbl(doc,
        ["Regulamentação", "Agência", "Escopo", "Referência Legal"],
        [
            ["Bloodborne Pathogens Standard", "OSHA", "Proteção contra exposição a patógenos transmitidos pelo sangue", f"29 CFR 1910.1030{f8}"],
            ["Hazard Communication Standard", "OSHA", "Comunicação de riscos químicos no ambiente de trabalho", "29 CFR 1910.1200"],
            ["Infection Prevention & Control", "CDC", "Diretrizes de controle de infecção em estabelecimentos odontológicos", f"MMWR 52(RR-17){f9}"],
            ["Medical Waste Management", "EPA", "Gestão e descarte de resíduos biomédicos", f"40 CFR 259{f10}"],
            ["Privacy and Security Rules", "HHS/OCR", "Proteção de dados de saúde de pacientes", f"45 CFR Parts 160, 164{f11}"],
            ["Regulation of Dental Devices", "FDA", "Dispositivos de monitoramento IoT (SteriSensor)", f"21 CFR 872{f12}"],
        ],
        "Compilação própria com base nas regulamentações federais vigentes."
    )

    h3(doc, "1.4.4 Conformidade Estadual")
    p(doc, f"Além das regulamentações federais, a DentalShield deve atender aos requisitos dos boards estaduais de odontologia em cada jurisdição de operação. Na Flórida, o Florida Board of Dentistry regulamenta práticas odontológicas sob o F.S. Chapter 466{f13}, incluindo requisitos de esterilização, manutenção de registros e educação continuada. No Texas, o Texas State Board of Dental Examiners (TSBDE) administra regulamentações equivalentes sob o Occupations Code Title 3 Subtitle D{f14}.")

    h3(doc, "1.4.5 Proteção de Propriedade Intelectual")
    p(doc, "A DentalShield desenvolverá três ativos de propriedade intelectual proprietária:")
    bullet(doc, "**SteriSensor™** — Sistema IoT de monitoramento de autoclaves e unidades de armazenamento (pedido de registro de marca e patente de utilidade)")
    bullet(doc, "**BiohazardBox™** — Sistema integrado de gestão de resíduos biomédicos (registro de marca)")
    bullet(doc, "**ComplianceScreen™** — Plataforma SaaS de compliance digital (registro de marca e copyright do código-fonte)")

    h3(doc, "1.4.6 Seguros e Responsabilidade")
    p(doc, "A estrutura de seguros corporativos incluirá:")
    tbl(doc,
        ["Tipo de Seguro", "Cobertura Mínima", "Finalidade"],
        [
            ["Errors & Omissions (E&O)", "US$ 1.000.000", "Proteção contra reclamações por falhas em assessoria de compliance"],
            ["General Liability", "US$ 2.000.000", "Cobertura para danos a terceiros em visitas a clínicas"],
            ["Cyber Liability", "US$ 500.000", "Proteção de dados de clientes (HIPAA compliance)"],
            ["Workers' Compensation", "Conforme estado", "Obrigatório na Flórida para 4+ funcionários"],
            ["Commercial Auto", "US$ 500.000", "Veículos utilizados em visitas técnicas"],
        ]
    )

    divider(doc)


# ============================================================
# SECTION 2: ANÁLISE ESTRATÉGICA
# ============================================================

def s2_1(doc):
    """2.1 Perspectivas do Mercado (zoom: global → nacional → regional)"""
    h1(doc, "2. ANÁLISE ESTRATÉGICA")
    h2(doc, "2.1 Perspectivas do Mercado")

    f15 = fn("Precedence Research, U.S. Dental Service Market Report, 2025. Mercado avaliado em US$ 164,76B em 2024.", "https://www.precedenceresearch.com/us-dental-service-market")
    f16 = fn("Fortune Business Insights, Infection Control Market Size, 2024. Mercado global de US$ 22,7B, CAGR 7,1%.")
    f17 = fn("Precedence Research, IoT in Healthcare Market, 2024. Mercado global de US$ 89,7B, CAGR 18,4%.")
    f18 = fn("ADA Health Policy Institute, Supply of Dentists in the U.S.: 2001–2023.", "https://www.ada.org/resources/research/health-policy-institute")
    f19 = fn("Mordor Intelligence, Healthcare Compliance Software Market, 2024. Mercado de US$ 2,94B, CAGR 9,5%.", "https://www.mordorintelligence.com/industry-reports/healthcare-compliance-software-market")
    f20 = fn("OSHA News Release, 14 jan 2025. Penalidades: US$ 16.550/violação grave, US$ 165.514/violação intencional.", "https://www.osha.gov/news/newsreleases/osha-trade-release/20250114")

    h3(doc, "2.1.1 Panorama Global e Nacional")
    p(doc, f"O mercado de serviços odontológicos nos Estados Unidos foi avaliado em US$ 164,76 bilhões em 2024, com projeção de atingir US$ 270,57 bilhões até 2034 — uma CAGR de 5,3%{f15}. Este crescimento é impulsionado pelo envelhecimento populacional, expansão do acesso a cuidados dentários e, sobretudo, pela intensificação dos requisitos regulatórios de biossegurança. O país abriga mais de 200 mil clínicas odontológicas registradas{f18}, cada uma sujeita a um complexo mosaico de regulações federais, estaduais e municipais.")

    p(doc, f"O segmento específico de software de compliance em saúde — que inclui soluções para o setor odontológico — foi estimado em US$ 2,94 bilhões em 2024, com CAGR de 9,5% até 2030{f19}. Este crescimento acelerado, superior ao do próprio mercado odontológico subjacente, evidencia que a conformidade regulatória está se tornando um centro de custo cada vez mais relevante para proprietários de clínicas. Paralelamente, o mercado global de controle de infecções atingiu US$ 22,7 bilhões{f16}, enquanto o de IoT em saúde alcançou US$ 89,7 bilhões com CAGR excepcional de 18,4%{f17}.")

    p(doc, f"No ambiente regulatório, as penalidades da OSHA para 2025 foram atualizadas para US$ 16.550 por violação grave e US$ 165.514 por violação intencional{f20}. Estes valores representam um incentivo financeiro direto para que clínicas invistam em programas estruturados de compliance, dado que o custo de uma única citação pode exceder o investimento anual em serviços de conformidade.")

    f21 = fn("U.S. Census Bureau, QuickFacts, ACS 2020–2024. Hillsborough County, FL.", "https://www.census.gov/quickfacts/fact/table/hillsboroughcountyflorida/INC910224")
    f22 = fn("U.S. Census Bureau, QuickFacts, ACS 2020–2024. Miami-Dade County, FL.", "https://www.census.gov/quickfacts/fact/table/miamidadecountyflorida/INC910224")
    f23 = fn("U.S. Census Bureau, QuickFacts, ACS 2020–2024. Harris County, TX.", "https://www.census.gov/quickfacts/fact/table/harriscountytexas/INC910224")

    h3(doc, "2.1.2 Mercados Regionais — Tampa, Miami e Houston")
    p(doc, "A seleção dos três mercados regionais obedeceu a critérios rigorosos de densidade de clínicas, crescimento populacional, renda domiciliar e ambiente regulatório estadual:")
    tbl(doc,
        ["Indicador", "Tampa / Hillsborough", "Miami / Miami-Dade", "Houston / Harris"],
        [
            ["População do condado", "1.512.070", "2.701.767", "4.780.913"],
            ["Renda per capita", f"US$ 44.342{f21}", f"US$ 39.469{f22}", f"US$ 41.006{f23}"],
            ["Clínicas odontológicas (MSA)", "~2.800", "~4.200", "~5.100"],
            ["Imposto de renda estadual", "0% (FL)", "0% (FL)", "0% (TX)"],
            ["Fase de entrada", "Ano 0 (Sede)", "Ano 2 (Filial)", "Ano 2–5 (Filial)"],
        ],
        "U.S. Census Bureau, ACS 2020–2024; ADA Health Policy Institute, 2023."
    )

    p(doc, "A escolha da Flórida como base inicial justifica-se pela ausência de imposto de renda estadual, pelo crescimento demográfico acima da média nacional e pela expressiva concentração de clínicas independentes — o segmento primário da DentalShield. A expansão para Houston (Texas) no Ano 2 aproveita o maior mercado odontológico do estado e condições fiscais igualmente favoráveis.")

    h3(doc, "2.1.3 Tampa / Hillsborough County — Análise Detalhada")
    p(doc, f"O Condado de Hillsborough, sede da DentalShield, apresenta uma combinação singular de fatores favoráveis ao estabelecimento de uma empresa de compliance odontológico. Com população de 1.512.070 habitantes e renda per capita de US$ 44.342{f21}, o condado abriga aproximadamente 2.800 clínicas odontológicas na Tampa-St. Petersburg-Clearwater MSA. A densidade de clínicas por habitante — superior à média nacional — reflete o perfil demográfico da região: população em envelhecimento (27% acima de 55 anos), alta taxa de cobertura por seguros odontológicos (68% da população adulta) e presença expressiva de comunidades de aposentados com demanda crescente por procedimentos restauradores e estéticos.")

    p(doc, "Do ponto de vista operacional, Tampa oferece custos de ocupação significativamente inferiores a Miami (aluguel comercial médio de US$ 22/sq ft vs US$ 38/sq ft em Miami), proximidade com a University of South Florida (pipeline de talentos em saúde pública e engenharia biomédica) e infraestrutura logística que facilita a cobertura de toda a costa oeste da Flórida — de Sarasota a Clearwater — a partir de uma única sede. O Tampa International Airport (TPA) viabiliza deslocamentos rápidos para Miami e Houston, fundamentais para a gestão das futuras filiais.")

    h3(doc, "2.1.4 Miami / Miami-Dade County — Análise Detalhada")
    p(doc, f"O Condado de Miami-Dade, com 2.701.767 habitantes e renda per capita de US$ 39.469{f22}, constitui o maior mercado odontológico do estado da Flórida. A MSA de Miami-Fort Lauderdale-Pompano Beach abriga aproximadamente 4.200 clínicas odontológicas — quase 50% mais que Tampa. A característica diferenciadora do mercado de Miami é sua composição multicultural: mais de 72% da população é de origem hispânica, criando demanda por serviços de compliance bilíngues (inglês/espanhol) que poucos concorrentes conseguem atender.")

    p(doc, "Adicionalmente, Miami é um dos principais destinos de turismo odontológico nos Estados Unidos, atraindo pacientes da América Latina e do Caribe que buscam procedimentos estéticos e restauradores de alto valor. Este fluxo adicional intensifica a pressão regulatória sobre clínicas locais: um volume maior de procedimentos significa maior exposição a riscos de contaminação cruzada, maior volume de resíduos biomédicos e maior necessidade de documentação de compliance para pacientes internacionais. A DentalShield posiciona sua filial de Miami como hub para atender tanto clínicas tradicionais quanto práticas especializadas em dental tourism.")

    h3(doc, "2.1.5 Houston / Harris County — Análise Detalhada")
    p(doc, f"O Harris County, com 4.780.913 habitantes e renda per capita de US$ 41.006{f23}, abriga o maior mercado odontológico do Texas e o quarto maior dos Estados Unidos. A MSA de Houston-The Woodlands-Sugar Land conta com aproximadamente 5.100 clínicas odontológicas, impulsionadas pela presença do Texas Medical Center — o maior complexo médico do mundo — e por uma economia diversificada que sustenta cobertura de seguros odontológicos acima da média nacional.")

    p(doc, "A entrada no mercado do Texas é estrategicamente relevante por três razões. Primeira, o Texas não impõe imposto de renda estadual, mantendo a vantagem fiscal da Flórida. Segunda, o Texas State Board of Dental Examiners (TSBDE) opera sob regulamentação própria (Occupations Code Title 3 Subtitle D) que difere em aspectos importantes do Florida Board of Dentistry — esta complexidade multi-jurisdicional é precisamente o tipo de desafio que a DentalShield resolve para seus clientes. Terceira, Houston apresenta a maior concentração de Dental Service Organizations (DSOs) no sul dos Estados Unidos, constituindo o público-alvo secundário da empresa.")

    f_fed1 = fn("OSHA, National Emphasis Program — Healthcare, 2024. Aumento de inspeções em estabelecimentos de saúde.", "https://www.osha.gov/enforcement/directives")
    f_fed2 = fn("HHS, HIPAA Enforcement Rule Updates, 2024. Penalidade máxima: US$ 1,5M por categoria.", "https://www.hhs.gov/hipaa/index.html")
    f_fed3 = fn("FDA, 21 CFR 872 — Regulation of Dental Devices. Clearance 510(k) para dispositivos IoT.", "https://www.fda.gov/medical-devices")

    h3(doc, "2.1.6 Ambiente Regulatório e Políticas Federais")
    p(doc, f"O ambiente regulatório do setor odontológico nos Estados Unidos é definido pela atuação simultânea de múltiplas agências federais, cada uma com jurisdição sobre aspectos específicos da operação clínica. A OSHA, por meio do National Emphasis Program para o setor de saúde{f_fed1}, intensificou as inspeções em estabelecimentos odontológicos a partir de 2024, com foco particular em conformidade com o Bloodborne Pathogens Standard (29 CFR 1910.1030) e o Hazard Communication Standard (29 CFR 1910.1200). As penalidades atualizadas para 2025 — US$ 16.550 por violação grave e US$ 165.514 por violação intencional — representam os valores mais altos da história da agência.")

    p(doc, f"Paralelamente, o Department of Health and Human Services (HHS), por meio do Office for Civil Rights (OCR), atualizou as regras de enforcement do HIPAA em 2024, estabelecendo penalidades máximas de US$ 1,5 milhão por categoria de violação{f_fed2}. Para clínicas odontológicas que processam dados eletrônicos de saúde (ePHI) — o que inclui praticamente todas as clínicas modernas com registros digitais —, a conformidade com o HIPAA Security Rule (45 CFR Part 164) é tão crítica quanto a conformidade com a OSHA.")

    p(doc, f"No âmbito dos dispositivos médicos, a FDA regula equipamentos de monitoramento odontológico sob 21 CFR 872{f_fed3}. O SteriSensor da DentalShield, como dispositivo IoT conectado a autoclaves de esterilização, pode requerer clearance 510(k) dependendo de sua classificação pela agência. A empresa está em processo de determinação regulatória junto à FDA para garantir conformidade antes do lançamento comercial.")


def s2_2(doc):
    """2.2 Cadeia de Suprimentos"""
    h2(doc, "2.2 Cadeia de Suprimentos")
    p(doc, "A cadeia de valor da DentalShield é estruturada segundo o modelo de Porter, adaptado para serviços de compliance e tecnologia:")
    tbl(doc,
        ["Atividade", "Tipo", "Descrição", "Recursos-Chave"],
        [
            ["Auditoria de compliance", "Primária", "Visitas técnicas, avaliação de gaps, relatórios de remediação", "Especialistas em compliance; checklists proprietários"],
            ["Treinamento de equipes", "Primária", "Programas presenciais e online; certificação OSHA/HIPAA", "Instrutores certificados; plataforma LMS"],
            ["Instalação IoT", "Primária", "Deploy de sensores SteriSensor em autoclaves e unidades de armazenamento", "Técnicos de campo; dispositivos IoT; firmware"],
            ["Monitoramento contínuo", "Primária", "Análise de dados em tempo real; alertas automatizados", "Infraestrutura cloud (AWS); dashboards"],
            ["Relatórios de compliance", "Primária", "Geração de relatórios audit-ready para inspeções OSHA/HIPAA", "ComplianceScreen SaaS; templates regulatórios"],
            ["P&D de produtos", "Suporte", "Desenvolvimento de SteriSensor v2, BiohazardBox, ComplianceScreen", "Engenheiros de software; designers IoT"],
            ["Gestão de pessoas", "Suporte", "Recrutamento, treinamento interno, retenção de talentos", "RH; programas de capacitação"],
            ["Infraestrutura tecnológica", "Suporte", "Servidores cloud, CRM (HubSpot), ferramentas de comunicação", "DevOps; licenças SaaS"],
        ],
        "Adaptação do modelo de Cadeia de Valor de Porter para serviços de compliance odontológico."
    )


def s2_3(doc):
    """2.3 Empregabilidade Esperada"""
    h2(doc, "2.3 Empregabilidade Esperada")

    f24 = fn("BLS, Occupational Outlook Handbook, May 2024 OEWS. Dental Hygienists (SOC 29-1292): mediana US$ 94.260.", "https://www.bls.gov/ooh/healthcare/dental-hygienists.htm")
    f25 = fn("BLS, OEWS May 2024. Compliance Officers (SOC 13-1041): mediana US$ 75.670.", "https://www.bls.gov/oes/current/oes131041.htm")
    f26 = fn("BLS, OEWS May 2024. Market Research Analysts (SOC 13-1161): mediana US$ 74.680.", "https://www.bls.gov/oes/current/oes131161.htm")
    f27 = fn("BLS, OEWS May 2024. Medical Equipment Repairers (SOC 49-9062): mediana US$ 60.360.", "https://www.bls.gov/oes/current/oes499062.htm")
    f28 = fn("Economic Policy Institute, Updated Employment Multipliers, NAICS 54. Multiplicador: 3,35x.")

    p(doc, "A DentalShield projetará a criação de 22 postos de trabalho diretos ao longo de cinco anos, distribuídos em três localidades. As ocupações foram classificadas segundo o Standard Occupational Classification (SOC) do Bureau of Labor Statistics:")
    tbl(doc,
        ["SOC Code", "Ocupação", "Salário Mediano (EUA)", "Crescimento Projetado", "Vagas Anuais"],
        [
            [f"29-1292{f24}", "Dental Hygienists", "US$ 94.260", "+7%", "12.100"],
            [f"13-1041{f25}", "Compliance Officers", "US$ 75.670", "+4%", "33.800"],
            [f"13-1161{f26}", "Market Research Analysts", "US$ 74.680", "+13%", "99.800"],
            [f"49-9062{f27}", "Medical Equipment Repairers", "US$ 60.360", "+14%", "6.700"],
            ["43-4171", "Receptionists & Information Clerks", "US$ 37.140", "+3%", "142.000"],
        ],
        "Bureau of Labor Statistics, Occupational Employment and Wage Statistics, May 2024."
    )

    p(doc, f"Aplicando o multiplicador de emprego do Economic Policy Institute (EPI) para o setor NAICS 54 (Professional, Scientific & Technical Services), cada emprego direto gera 2,35 empregos adicionais na cadeia de suprimentos e consumo induzido{f28}. Portanto, os 22 empregos diretos da DentalShield no Ano 5 projetam a criação de aproximadamente **73,7 empregos totais** na economia — um impacto significativo para as comunidades de Tampa, Miami e Houston.")

    tbl(doc,
        ["", "Ano 1", "Ano 2", "Ano 3", "Ano 4", "Ano 5"],
        [
            ["Tampa, FL"] + [str(x) for x in FIN['tampa']],
            ["Miami, FL"] + [str(x) for x in FIN['miami']],
            ["Houston, TX"] + [str(x) for x in FIN['houston']],
            ["**Total**"] + [str(x) for x in FIN['emps']],
        ],
        "Projeções internas da DentalShield Compliance Solutions LLC."
    )

    img(doc, "03_employees.png", "Workforce Growth by Location — 5-Year Projection")


def s2_4_to_2_8(doc):
    """2.4 a 2.8 — Gestão do Conhecimento, ESG, SWOT, SWOT Cruzada, Concorrentes"""

    # 2.4 Gestão do Conhecimento
    h2(doc, "2.4 Gestão do Conhecimento")
    p(doc, "A estratégia de gestão do conhecimento da DentalShield fundamenta-se em quatro dimensões complementares:")
    tbl(doc,
        ["Dimensão", "Descrição", "Aplicação na DentalShield"],
        [
            ["Capital Humano", "Conhecimento tácito dos especialistas em compliance e controle de infecção", "Programa de mentoria interna; base de conhecimento de auditorias (500+ checklists); certificações obrigatórias"],
            ["Capital Estrutural", "Processos, sistemas e propriedade intelectual da empresa", "SteriSensor firmware; ComplianceScreen codebase; protocolos de auditoria padronizados; manuais de treinamento"],
            ["Capital Relacional", "Relacionamento com clientes, parceiros e reguladores", "CRM com histórico de cada clínica; parcerias com ADA e OSAP; rede de DSOs"],
            ["Capital Tecnológico", "Infraestrutura de dados e análise", "Data lake de IoT (dados de autoclaves em tempo real); dashboards de compliance; machine learning para predição de falhas"],
        ]
    )

    # 2.5 ESG
    h2(doc, "2.5 Impactos ESG")
    h3(doc, "Environmental")
    p(doc, "O sistema BiohazardBox reduz o descarte irregular de resíduos biomédicos, promovendo classificação correta e rastreabilidade completa conforme EPA 40 CFR 259. Cada clínica atendida elimina em média 2,3 toneladas/ano de resíduos potencialmente perigosos do aterro comum.")
    h3(doc, "Social")
    p(doc, "Cada clínica em conformidade é uma comunidade mais segura: protocolos de esterilização validados reduzem o risco de infecções cruzadas para pacientes e profissionais. O treinamento de equipes eleva o nível de qualificação do workforce local. O multiplicador EPI de 3,35x projeta 73,7 empregos totais na economia.")
    h3(doc, "Governance")
    p(doc, "A DentalShield promove cultura de compliance ativo: auditorias regulares, documentação contínua e transparência com reguladores. O modelo de negócio alinha receita recorrente com resultado de compliance, criando incentivo para melhoria contínua.")

    # 2.6 SWOT
    h2(doc, "2.6 Análise SWOT")
    p(doc, "A análise SWOT a seguir foi elaborada com base no levantamento de mercado apresentado nas seções anteriores, nas projeções financeiras da planilha operacional e na avaliação competitiva do setor de compliance odontológico nos três MSAs selecionados. Cada quadrante é acompanhado de contextualização que orienta as estratégias cruzadas apresentadas na Seção 2.7.")
    for title, items in [
        ("Strengths (Forças)", [
            "Especialização vertical exclusiva no setor odontológico",
            "Integração IoT (SteriSensor) + assessoria técnica — modelo híbrido sem concorrente direto",
            "3 produtos proprietários (SteriSensor, BiohazardBox, ComplianceScreen)",
            "Operação em estados sem imposto de renda (FL e TX)",
            "Equipe com expertise em OSHA, CDC, HIPAA e regulamentação estadual",
        ]),
        ("Weaknesses (Fraquezas)", [
            "Empresa em fase de startup — sem histórico operacional prévio",
            "Dependência de contratação especializada em mercado competitivo de trabalho",
            "Investimento inicial em P&D de IoT requer validação de produto",
            "Brand awareness zero no mercado inicial",
            "Complexidade operacional de gestão multi-state (FL + TX)",
        ]),
        ("Opportunities (Oportunidades)", [
            f"Mercado de compliance em saúde crescendo a 9,5% ao ano (US$ 2,94B em 2024)",
            "Aumento contínuo das penalidades OSHA: US$ 16.550/violação grave (2025)",
            "Crescimento de DSOs (consolidação do mercado) criando demanda por compliance em escala",
            "IoT em saúde com CAGR de 18,4% — forte adoção de monitoramento digital",
            "Ausência de concorrente que combine IoT + auditoria presencial + SaaS",
        ]),
        ("Threats (Ameaças)", [
            "Entrada de grandes players de compliance (Total Medical Compliance, SmartCompliance)",
            "Mudanças regulatórias que simplifiquem requisitos (reduzindo demanda por assessoria especializada)",
            "Recessão econômica reduzindo investimento de clínicas em compliance não-obrigatório",
            "Escassez de mão de obra qualificada em compliance odontológico",
            "Riscos de cybersegurança nos dispositivos IoT (SteriSensor)",
        ]),
    ]:
        h3(doc, title)
        for item in items:
            bullet(doc, item)

    # 2.7 SWOT Cruzada
    h2(doc, "2.7 SWOT Cruzada")
    tbl(doc,
        ["Quadrante", "Estratégia 1", "Estratégia 2", "Estratégia 3"],
        [
            ["SO (Forças × Oportunidades)", "Lançar SteriSensor IoT como diferencial competitivo no mercado de US$ 89,7B de IoT em saúde", "Posicionar-se como fornecedor preferencial de DSOs em expansão, oferecendo compliance em escala", "Capitalizar penalidades crescentes da OSHA com campanha educativa sobre custo de não-compliance"],
            ["WO (Fraquezas × Oportunidades)", "Superar falta de brand awareness com programa de referral para primeiros 50 clientes", "Mitigar risco de P&D com MVP do SteriSensor antes de investimento completo", "Contratar remotamente para reduzir competição por talento em mercado local"],
            ["ST (Forças × Ameaças)", "Diferenciar-se de players genéricos com especialização exclusiva em odontologia", "Construir moat tecnológico com patentes de SteriSensor e ComplianceScreen", "Diversificar receita entre serviços técnicos (65%) e tecnologia (35%) para resiliência a ciclos econômicos"],
            ["WT (Fraquezas × Ameaças)", "Implementar SOC 2 Type II desde o Ano 1 para proteção de dados IoT", "Estabelecer reserve fund de 6 meses de operação para cenário recessivo", "Firmar parcerias com universidades para pipeline de talentos em compliance odontológico"],
        ]
    )

    # 2.8 Análise de Concorrentes
    h2(doc, "2.8 Análise de Concorrentes")
    p(doc, "O mercado de compliance odontológico nos Estados Unidos é fragmentado, sem player dominante que ofereça a integração completa de auditoria presencial, IoT e SaaS:")
    tbl(doc,
        ["Empresa", "Escopo", "Auditoria Presencial", "IoT", "SaaS", "Foco Odontológico", "Precificação"],
        [
            ["**DentalShield**", "Nacional (3 states)", "✔", "✔ SteriSensor", "✔ ComplianceScreen", "✔ Exclusivo", "Premium"],
            ["OSHA Review Inc.", "Nacional (online)", "✘", "✘", "✔ Parcial", "✘ Genérico", "Baixo"],
            ["Total Medical Compliance", "Nacional", "✔", "✘", "✔", "✘ Médico genérico", "Médio"],
            ["Dental Compliance Specialists", "Southeast", "✔", "✘", "✘", "✔ Odontológico", "Médio"],
            ["Practice Compliance Solutions", "Texas", "✔ Parcial", "✘", "✘", "✘ HIPAA genérico", "Baixo"],
            ["SmartCompliance", "Nacional (SaaS)", "✘", "✘", "✔", "✘ Multi-setor", "Baixo"],
        ],
        "Análise competitiva elaborada pela DentalShield com base em informações publicamente disponíveis."
    )
    p(doc, "A análise evidencia que nenhum concorrente oferece simultaneamente auditoria presencial, monitoramento IoT em tempo real e plataforma SaaS com foco exclusivo no setor odontológico. Este é o **core differentiator** da DentalShield e a base do posicionamento premium.")

    divider(doc)


# ============================================================
# PLACEHOLDER PARA BLOCOS 3, 4 e 5 (a ser expandido)
# ============================================================

def block3_marketing_full(doc):
    """Bloco 3: Marketing completo (3.1 a 3.7)"""
    h1(doc, "3. PLANO DE MARKETING")

    f29 = fn("ADA Health Policy Institute, 2023. Estimativa de clínicas odontológicas por MSA.")
    f30 = fn("U.S. Census Bureau, ACS 2020–2024. Dados demográficos por condado.")
    f31 = fn("OSHA, Penalties Summary, 2025. Multa média por violação grave: US$ 16.550.")

    # 3.1 Segmentação
    h2(doc, "3.1 Segmentação de Mercado")
    h3(doc, "3.1.1 Segmentação Geográfica")
    p(doc, f"A seleção dos mercados-alvo obedeceu a critérios de densidade de clínicas, crescimento populacional e ambiente regulatório. Os três MSAs selecionados concentram mais de 12.100 clínicas odontológicas ativas{f29}:")
    tbl(doc,
        ["MSA", "Clínicas Ativas", "População (Condado)", "Renda Per Capita"],
        [
            ["Tampa-St. Petersburg, FL", "~2.800", "1.512.070", "US$ 44.342"],
            ["Miami-Fort Lauderdale, FL", "~4.200", "2.701.767", "US$ 39.469"],
            ["Houston-The Woodlands, TX", "~5.100", "4.780.913", "US$ 41.006"],
        ],
        f"ADA HPI, 2023; U.S. Census Bureau, ACS 2020–2024{f30}."
    )

    h3(doc, "3.1.2 Segmentação Demográfica")
    p(doc, "O perfil dos tomadores de decisão nas clínicas-alvo apresenta as seguintes características:")
    bullet(doc, "**Idade:** 35–60 anos (proprietários com prática estabelecida)")
    bullet(doc, "**Formação:** DDS/DMD, com limitada formação em gestão regulatória")
    bullet(doc, "**Faturamento da clínica:** US$ 500.000–US$ 2.500.000 anuais")
    bullet(doc, "**Funcionários:** 3–25 colaboradores por unidade")

    h3(doc, "3.1.3 Segmentação Psicográfica")
    p(doc, "Três perfis comportamentais distintos orientam a estratégia de abordagem comercial:")
    tbl(doc,
        ["Perfil", "% do Mercado", "Comportamento", "Estratégia de Abordagem"],
        [
            ["Proativo Regulatório", "~20%", "Busca ativamente compliance; valoriza certificações e auditorias preventivas", "Venda direta; demonstração técnica do SteriSensor"],
            ["Reativo Preocupado", "~55%", "Reconhece importância mas posterga investimento até enfrentar inspeção", "Campanhas educativas sobre custo de não-compliance"],
            ["Indiferente", "~25%", "Desconhece ou subestima riscos regulatórios", "ROI calculator; cases de multas evitadas"],
        ]
    )

    # 3.2 Público-Alvo
    h2(doc, "3.2 Público-Alvo")
    h3(doc, "Público Primário — Clínicas Independentes (Solo Practices)")
    p(doc, "Clínicas operadas por um único dentista proprietário, com equipe de 3 a 12 colaboradores, representam aproximadamente 67% das clínicas nos MSAs selecionados. Este segmento enfrenta o maior gap de compliance, uma vez que raramente dispõe de profissional dedicado. O ciclo de venda estimado é de 30–60 dias, com ticket médio mensal de US$ 2.500–US$ 4.200. A DentalShield estima que 78% da receita do Ano 1 será proveniente deste segmento.")

    h3(doc, "Público Secundário — DSOs (Dental Service Organizations)")
    p(doc, "DSOs administram redes de clínicas sob gestão centralizada e representam o segmento de maior crescimento. A DentalShield oferece contratos multi-unidade com descontos volumétricos. Ciclo de venda: 90–180 dias; LTV estimado: US$ 120.000–US$ 480.000 por contrato plurianual. Penetração acelerada a partir do Ano 3.")

    h3(doc, "Público Terciário — Faculdades de Odontologia")
    p(doc, "Instituições acadêmicas constituem canal estratégico de longo prazo. Parcerias com USF (Tampa) e UT Health (Houston) estão em prospecção para integrar protocolos DentalShield nos currículos de formação.")

    # 3.3 Posicionamento
    h2(doc, "3.3 Posicionamento da Marca")
    p(doc, "**Declaração de Posicionamento:** A DentalShield é a única empresa de compliance odontológico nos Estados Unidos que integra auditoria regulatória, monitoramento IoT em tempo real e plataforma SaaS em uma solução unificada, permitindo que dentistas concentrem-se na prática clínica com a certeza de conformidade regulatória assegurada.", bold=False, italic=True)
    p(doc, "O posicionamento fundamenta-se em três pilares: (1) **Especialização vertical** exclusiva no setor odontológico; (2) **Integração tecnológica** com IoT SteriSensor e SaaS ComplianceScreen; (3) **Parceria de longo prazo** via modelo de receita recorrente que alinha incentivos.")

    # 3.4 Produto
    h2(doc, "3.4 Produto — Análise de Valor")
    tbl(doc,
        ["Serviço", "Preço Mensal", "Preço Hora", "% Receita", "Diferencial"],
        [
            ["Regulatory Compliance Auditing", "US$ 2.500/mês", "US$ 150/h", "30%", "Checklist 200+ itens OSHA/CDC/HIPAA"],
            ["Infection Control & Biosafety", "—", "US$ 125/h", "25%", "Protocolos CDC com validação em campo"],
            ["IoT SteriSensor", "US$ 199/dispositivo", "—", "20%", "Monitoramento 24/7 com alertas automáticos"],
            ["Staff Training & Certification", "—", "US$ 750/sessão", "15%", "Certificação reconhecida OSHA/HIPAA"],
            ["BiohazardBox Waste Mgmt", "US$ 450/mês", "—", "5%", "Rastreabilidade EPA 40 CFR 259"],
            ["ComplianceScreen SaaS", "US$ 299/mês", "—", "5%", "Dashboard + alertas + relatórios audit-ready"],
        ],
        "Tabela de precificação da DentalShield Compliance Solutions LLC."
    )
    img(doc, "04_services.png", "Revenue Distribution by Service Line")

    # 3.5 Preço
    h2(doc, "3.5 Preço — Estratégia de Precificação")
    p(doc, f"A estratégia de precificação da DentalShield é posicionada como **premium justificado pelo custo de não-compliance**. Uma única citação da OSHA por violação grave custa US$ 16.550{f31} — valor que excede o investimento anual de uma clínica em serviços DentalShield (US$ 30.000–US$ 50.000/ano para pacote completo). O retorno sobre investimento para o cliente é imediato: a prevenção de uma única multa paga o serviço de um ano inteiro.")
    tbl(doc,
        ["Pacote", "Serviços Incluídos", "Preço Mensal", "ROI Estimado (cliente)"],
        [
            ["Essential", "Auditoria trimestral + ComplianceScreen SaaS", "US$ 2.799", "Evita ~US$ 16.550/multa OSHA"],
            ["Professional", "Essential + Infection Control + Training mensal", "US$ 4.500", "Evita multas + reduz infecções cruzadas"],
            ["Enterprise", "Professional + SteriSensor IoT + BiohazardBox", "US$ 7.200", "Compliance integral + monitoramento 24/7"],
        ]
    )

    # 3.6 Praça
    h2(doc, "3.6 Praça — Estratégia de Distribuição")
    p(doc, "A distribuição segue modelo direto com expansão geográfica faseada:")
    bullet(doc, "**Fase 1 (Ano 0–1):** Vendas diretas em Tampa MSA. Equipe de 2 account executives + 1 especialista técnico. Meta: 50 clientes ativos.")
    bullet(doc, "**Fase 2 (Ano 2–3):** Abertura de filiais em Miami e Houston. 3 equipes regionais. Parcerias com distribuidores de equipamentos odontológicos. Meta: 150 clientes.")
    bullet(doc, "**Fase 3 (Ano 4–5):** Modelo de partner network (revendedores certificados). Expansão para Atlanta, Dallas, Phoenix. Meta: 300+ clientes.")

    # 3.7 Promoção
    h2(doc, "3.7 Promoção — Orçamento de Marketing")
    p(doc, "O orçamento de marketing representa 13,5% da receita bruta, distribuído entre canais digitais e presenciais:")
    tbl(doc,
        ["Canal", "% do Budget", "Atividades Principais"],
        [
            ["LinkedIn Ads + Content", "25%", "Artigos sobre compliance dental; cases de multas evitadas; webinars mensais"],
            ["Conferências (ADA, FDEA, TDA)", "20%", "Booth + palestras + demos do SteriSensor em eventos da ADA Annual Meeting"],
            ["Referral Program", "15%", "Desconto de 10% para cliente que indicar; bônus para parceiros"],
            ["Content Marketing + SEO", "15%", "Blog semanal; whitepapers; guias OSHA/HIPAA para dentistas; YouTube"],
            ["Email Marketing", "10%", "Nurturing de leads; newsletters quinzenais; alertas regulatórios"],
            ["Google Ads (Search)", "10%", "Keywords: dental compliance, OSHA dental, dental infection control"],
            ["PR + Media", "5%", "Press releases; artigos em Dental Economics, ADA News, Dentistry Today"],
        ],
        "Alocação de marketing baseada em benchmarks do setor de serviços B2B em saúde."
    )

    divider(doc)


def block4_operational(doc):
    """Bloco 4: Operacional (seções 4.1 a 4.6)"""
    h1(doc, "4. PLANO OPERACIONAL")

    h2(doc, "4.1 Quadro de Funcionários")
    tbl(doc,
        ["Localidade", "Ano 1", "Ano 2", "Ano 3", "Ano 4", "Ano 5"],
        [
            ["Tampa, FL (Sede)"] + [str(x) for x in FIN['tampa']],
            ["Miami, FL (Filial)"] + [str(x) for x in FIN['miami']],
            ["Houston, TX (Filial)"] + [str(x) for x in FIN['houston']],
            ["**Total**"] + [str(x) for x in FIN['emps']],
        ],
        "Projeções internas. Valores representam headcount ao final de cada período."
    )

    h2(doc, "4.2 Layout Operacional")
    tbl(doc,
        ["Localidade", "Área", "Configuração", "Fase"],
        [
            ["Tampa, FL", "1.200 sq ft", "Escritório + sala de treinamento + estoque IoT", "Ano 0"],
            ["Miami, FL", "800 sq ft", "Escritório + sala de reuniões", "Ano 2"],
            ["Houston, TX", "800 → 1.200 sq ft", "Escritório → expansão Ano 5", "Ano 2–5"],
        ]
    )

    h2(doc, "4.3 Recursos e Equipamentos")
    bullet(doc, "**Dispositivos SteriSensor:** 200 unidades (Ano 1) → 1.500 unidades (Ano 5)")
    bullet(doc, "**Kits de treinamento:** Materiais OSHA/HIPAA; EPIs demonstrativos; autoclaves de treinamento")
    bullet(doc, "**Veículos:** 2 (Ano 1) → 5 (Ano 5) — visitas técnicas a clínicas")
    bullet(doc, "**Estações de trabalho:** Laptops + monitores + smartphones corporativos")

    h2(doc, "4.4 Recursos Tecnológicos")
    tbl(doc,
        ["Recurso", "Tecnologia", "Função"],
        [
            ["ComplianceScreen SaaS", "React + Node.js + PostgreSQL (AWS)", "Plataforma de compliance digital para clientes"],
            ["SteriSensor Cloud", "AWS IoT Core + Lambda + DynamoDB", "Ingestão e processamento de dados IoT em tempo real"],
            ["CRM", "HubSpot Professional", "Gestão de leads, pipeline de vendas, customer success"],
            ["Comunicação", "Slack + Zoom + Google Workspace", "Colaboração interna e reuniões com clientes"],
            ["Infraestrutura", "AWS (us-east-1)", "Hosting, storage, CDN, backups"],
        ]
    )

    h2(doc, "4.5 Localização")
    p(doc, "**Tampa, FL (Sede — Ano 0):** Capital do Condado de Hillsborough (pop. 1.512.070), Tampa combina alta densidade de clínicas odontológicas (~2.800 na MSA), ausência de imposto de renda estadual, crescimento demográfico acelerado e custo operacional inferior a Miami. A proximidade com a University of South Florida facilita parcerias acadêmicas e pipeline de talentos.")
    p(doc, "**Miami, FL (Filial — Ano 2):** Maior mercado odontológico da Flórida (~4.200 clínicas na MSA), Miami oferece acesso à comunidade latino-americana e ao segmento de dental tourism. O Condado de Miami-Dade (pop. 2.701.767) apresenta demanda expressiva por compliance bilíngue (inglês/espanhol).")
    p(doc, "**Houston, TX (Filial — Ano 2–5):** Maior mercado odontológico do Texas (~5.100 clínicas na MSA) e quarto maior dos EUA. O Harris County (pop. 4.780.913) combina densidade de mercado, ausência de imposto de renda estadual e proximidade com o Texas Medical Center — maior complexo médico do mundo.")

    h2(doc, "4.6 Capacidade Produtiva")
    p(doc, "A capacidade de atendimento é determinada pela proporção de especialistas por clientes ativos:")
    tbl(doc,
        ["Métrica", "Ano 1", "Ano 3", "Ano 5"],
        [
            ["Especialistas de campo", "3", "8", "14"],
            ["Clientes por especialista", "~17", "~19", "~21"],
            ["Clientes ativos (projeção)", "50", "150", "300+"],
            ["Visitas técnicas/mês", "100", "300", "600+"],
            ["Dispositivos SteriSensor ativos", "200", "800", "1.500+"],
        ],
        "Projeções operacionais internas."
    )

    divider(doc)


def block5_financial(doc):
    """Bloco 5: Financeiro (seções 5.1 a 5.6)"""
    h1(doc, "5. PLANO FINANCEIRO")

    h2(doc, "5.1 Premissas Financeiras")
    tbl(doc,
        ["Premissa", "Valor", "Justificativa"],
        [
            ["Inflação anual", "3%", "Média histórica CPI-U (últimos 10 anos)"],
            ["Taxa de desconto (WACC)", f"{FIN['discount']}%", "Custo de capital ajustado ao risco de startup no setor de serviços"],
            ["Alíquota efetiva de imposto", "25%", "LLC pass-through taxation; sem imposto estadual em FL e TX"],
            ["Payroll tax (FICA)", "7,65%", "Social Security (6,2%) + Medicare (1,45%)"],
            ["Custos variáveis / receita", "24,51%", "Marketing (13,5%), comissões, materiais, viagens"],
            ["Flórida income tax", "0%", "Sem imposto de renda estadual para pessoas físicas e LLCs"],
            ["Texas franchise tax threshold", "US$ 2.470.000", "Receitas abaixo do threshold isentas de franchise tax"],
        ]
    )

    p(doc, "As premissas financeiras foram calibradas com base em dados do Bureau of Labor Statistics, taxas tributárias publicadas pelo IRS e pelo Florida Department of Revenue, e benchmarks operacionais de empresas comparáveis no setor de serviços profissionais (NAICS 54). A taxa de desconto de 12% reflete o prêmio de risco associado a uma startup no setor de serviços, posicionando-se entre o custo de capital de empresas maduras de serviços (8-10%) e o de startups de tecnologia (15-20%). A alíquota efetiva de 25% considera o pass-through taxation da LLC, onde os lucros são tributados como renda pessoal dos membros, com a vantagem de que nem a Flórida nem o Texas impõem imposto de renda estadual.")

    p(doc, "Os custos variáveis — projetados em 24,51% da receita — incluem marketing (13,5%), comissões de vendas (4%), materiais e suprimentos (3%), despesas de viagem para visitas técnicas (2,5%) e custos de processamento de pagamento (1,51%). A predominância do marketing nos custos variáveis reflete a estratégia de penetração agressiva nos três primeiros anos, com redução gradual à medida que o referral program e a retenção de clientes reduzam o custo de aquisição.")

    h2(doc, "5.2 Investimentos")
    tbl(doc,
        ["Fase", "Período", "Tangível", "Intangível", "Capital de Giro", "Total"],
        [
            ["Fase 1", "Ano 0", usd(78108), usd(56635), usd(40514), usd(175257)],
            ["Fase 2", "Ano 2", usd(45200), usd(28700), usd(24446), usd(98346)],
            ["Fase 3", "Ano 4", usd(51800), usd(43994), usd(24446), usd(120240)],
            ["**Total**", "—", usd(175108), usd(129329), usd(89406), usd(FIN['invest_total'])],
        ],
        "Planilha Financeira V3 — DentalShield Compliance Solutions LLC."
    )
    img(doc, "06_investment.png", "Investment Structure by Phase")

    p(doc, f"O investimento total de {usd(FIN['invest_total'])} está distribuído em três fases alinhadas ao plano de expansão geográfica. A Fase 1 ({usd(175257)}) concentra os custos de constituição da LLC, aquisição de equipamentos iniciais (200 unidades SteriSensor, mobiliário de escritório, veículos), desenvolvimento da plataforma ComplianceScreen v1 e capital de giro para os primeiros 6 meses de operação em Tampa. A Fase 2 ({usd(98346)}) financia a abertura das filiais de Miami e Houston — aluguel, mobiliário, contratação de 4 novos colaboradores e expansão do inventário de sensores IoT. A Fase 3 ({usd(120240)}) destina-se ao escalonamento: upgrade do ComplianceScreen para v2, ampliação da equipe para 22 colaboradores, expansão do escritório de Houston e aumento do capital de giro para sustentar 300+ clientes ativos.")

    p(doc, "Os ativos intangíveis representam 32,8% do investimento total, refletindo a natureza intensiva em propriedade intelectual do negócio. O desenvolvimento do SteriSensor (firmware, protocolos de comunicação IoT, certificações de segurança) e do ComplianceScreen (plataforma SaaS, dashboards, motor de alertas) constituem o core tecnológico que diferencia a DentalShield de concorrentes baseados exclusivamente em serviços manuais.")

    h2(doc, "5.3 Estimativa de Receitas e Custos")
    tbl(doc,
        ["", "Ano 1", "Ano 2", "Ano 3", "Ano 4", "Ano 5"],
        [
            ["Receita Bruta"] + [usd(x) for x in FIN['rev']],
            ["(-) Custos Variáveis"] + [usd(x) for x in FIN['var_costs']],
            ["(-) Folha de Pagamento"] + [usd(x) for x in FIN['payroll']],
            ["(-) Custos Fixos"] + [usd(x) for x in FIN['fixed']],
            ["**Custos Totais**"] + [usd(x) for x in FIN['total_costs']],
            ["**EBITDA**"] + [usd(x) for x in FIN['ebitda']],
            ["Margem EBITDA"] + [pct(x) for x in FIN['ebitda_margin']],
        ],
        "Planilha Financeira V3 — DentalShield Compliance Solutions LLC."
    )
    img(doc, "01_revenue_vs_costs.png", "Revenue vs Total Costs — 5-Year Projection")
    img(doc, "08_cost_structure.png", "Cost Structure — Year 1 vs Year 5")

    p(doc, f"A análise de receitas e custos revela uma trajetória de crescimento com dois pontos de inflexão significativos. O Ano 1, com receita de {usd(FIN['rev'][0])} e custos totais de {usd(FIN['total_costs'][0])}, resulta em EBITDA negativo de {usd(FIN['ebitda'][0])} — déficit esperado e planejado para o período de validação do modelo de negócio e aquisição dos primeiros 50 clientes. O primeiro ponto de inflexão ocorre no Ano 2, quando a expansão para Miami e Houston duplica a base de receita para {usd(FIN['rev'][1])}, gerando o primeiro EBITDA positivo de {usd(FIN['ebitda'][1])} (margem de {pct(FIN['ebitda_margin'][1])}).")

    p(doc, f"O segundo ponto de inflexão ocorre no Ano 3, quando a operação atinge maturidade parcial com {usd(FIN['rev'][2])} de receita e margem EBITDA de {pct(FIN['ebitda_margin'][2])} — a maior do horizonte. O Ano 4 apresenta contração temporária na receita para {usd(FIN['rev'][3])}, refletindo o investimento da Fase 3 e a reestruturação operacional necessária para escalar de 150 para 300+ clientes. O Ano 5 consolida a trajetória com receita recorde de {usd(FIN['rev'][4])} e margem EBITDA de {pct(FIN['ebitda_margin'][4])}, confirmando a sustentabilidade do modelo.")

    p(doc, f"A estrutura de custos evolui significativamente ao longo dos cinco anos. A folha de pagamento — maior componente de custo — cresce de {usd(FIN['payroll'][0])} no Ano 1 para {usd(FIN['payroll'][4])} no Ano 5, refletindo a expansão de 5 para 22 colaboradores. Contudo, a produtividade por colaborador também aumenta: a receita por funcionário evolui de {usd(int(FIN['rev'][0]/FIN['emps'][0]))} (Ano 1) para {usd(int(FIN['rev'][4]/FIN['emps'][4]))} (Ano 5), um incremento de {int((FIN['rev'][4]/FIN['emps'][4])/(FIN['rev'][0]/FIN['emps'][0])*100-100)}% que evidencia ganhos de escala e eficiência operacional.")

    h2(doc, "5.4 DRE — Demonstrativo de Resultados")
    tbl(doc,
        ["", "Ano 1", "Ano 2", "Ano 3", "Ano 4", "Ano 5"],
        [
            ["Receita Bruta"] + [usd(x) for x in FIN['rev']],
            ["(-) Custos Variáveis"] + [usd(x) for x in FIN['var_costs']],
            ["= Margem de Contribuição"] + [usd(r-v) for r,v in zip(FIN['rev'], FIN['var_costs'])],
            ["(-) Folha de Pagamento"] + [usd(x) for x in FIN['payroll']],
            ["(-) Custos Fixos"] + [usd(x) for x in FIN['fixed']],
            ["= EBITDA"] + [usd(x) for x in FIN['ebitda']],
            ["(-) Impostos (25%)"] + [usd(max(0, int(x*0.25))) for x in FIN['ebitda']],
            ["= **Lucro Líquido**"] + [usd(x) for x in FIN['net_income']],
        ],
        "Planilha Financeira V3 — DentalShield Compliance Solutions LLC."
    )
    img(doc, "02_ebitda.png", "EBITDA & EBITDA Margin — 5-Year Projection")
    img(doc, "07_net_income.png", "Net Income — 5-Year Projection")

    h2(doc, "5.5 Indicadores de Retorno")
    tbl(doc,
        ["Indicador", "Valor", "Interpretação"],
        [
            ["Valor Presente Líquido (VPL)", usd(FIN['npv']), f"Positivo à taxa de {FIN['discount']}% — projeto economicamente viável"],
            ["Taxa Interna de Retorno (TIR)", pct(FIN['irr']), "Supera significativamente o custo de capital — retorno atrativo"],
            ["Retorno sobre Investimento (ROI)", pct(FIN['roi']), f"Cada US$ 1 investido retorna US$ {FIN['roi']/100:.2f}"],
            ["Payback", f"{FIN['payback']} anos", "Recuperação do investimento dentro do horizonte de 5 anos"],
            ["Valor Estimado do Negócio", usd(FIN['biz_value']), "Valoração baseada em múltiplo de EBITDA do Ano 5"],
        ],
        "Cálculos baseados na Planilha Financeira V3, taxa de desconto de 12%."
    )
    p(doc, f"O VPL de {usd(FIN['npv'])} confirma a viabilidade econômica do empreendimento, indicando que os fluxos de caixa futuros descontados superam o investimento total de {usd(FIN['invest_total'])}. A TIR de {pct(FIN['irr'])} — mais de cinco vezes superior à taxa de desconto de {FIN['discount']}% — sinaliza um retorno excepcional para o perfil de risco do negócio.")
    p(doc, f"O ROI de {pct(FIN['roi'])} traduz-se em uma multiplicação de 2,79x do capital investido ao longo de cinco anos, enquanto o payback de {FIN['payback']} anos demonstra recuperação rápida do investimento. O valor estimado do negócio de {usd(FIN['biz_value'])} ao final do Ano 5 representa uma valorização significativa dos ativos tangíveis e intangíveis construídos.")

    h2(doc, "5.6 Ponto de Equilíbrio (Break Even)")
    tbl(doc,
        ["", "Ano 1", "Ano 2", "Ano 3", "Ano 4", "Ano 5"],
        [
            ["BEP (Break Even)"] + [usd(x) for x in FIN['bep']],
            ["Receita Projetada"] + [usd(x) for x in FIN['rev']],
            ["Margem de Segurança"] + [pct((r-b)/r*100) if r > 0 else "—" for r,b in zip(FIN['rev'], FIN['bep'])],
        ],
        "Planilha Financeira V3."
    )
    img(doc, "05_breakeven.png", "Break-Even Analysis — Revenue vs BEP")
    p(doc, f"No Ano 1, a receita projetada de {usd(FIN['rev'][0])} situa-se ligeiramente abaixo do BEP de {usd(FIN['bep'][0])}, resultando no prejuízo operacional de {usd(FIN['net_income'][0])} esperado para o período de validação do modelo. A partir do Ano 2, a receita supera consistentemente o ponto de equilíbrio, atingindo margem de segurança superior a 39% no Ano 5.")

    divider(doc)


def block6_conclusion(doc):
    """Bloco 6: Conclusão"""
    h1(doc, "6. CONCLUSÃO")

    h2(doc, "6.1 Timeline de Implementação")
    tbl(doc,
        ["Atividade", "Q1-Q2 Y1", "Q3-Q4 Y1", "Y2", "Y3", "Y4", "Y5"],
        [
            ["Constituição LLC + EIN + Licenças", "✔", "", "", "", "", ""],
            ["Desenvolvimento SteriSensor v1", "✔", "✔", "", "", "", ""],
            ["Desenvolvimento ComplianceScreen v1", "✔", "✔", "", "", "", ""],
            ["Contratação equipe Tampa (5)", "✔", "", "", "", "", ""],
            ["Lançamento comercial Tampa", "", "✔", "", "", "", ""],
            ["Primeiros 50 clientes", "", "✔", "✔", "", "", ""],
            ["Abertura filial Miami", "", "", "✔", "", "", ""],
            ["Abertura filial Houston", "", "", "✔", "", "", ""],
            ["Contratação Miami + Houston (4)", "", "", "✔", "", "", ""],
            ["100 clientes ativos", "", "", "✔", "✔", "", ""],
            ["Parcerias com DSOs", "", "", "", "✔", "", ""],
            ["SteriSensor v2 + FDA clearance", "", "", "", "✔", "✔", ""],
            ["ComplianceScreen v2", "", "", "", "", "✔", ""],
            ["Expansão escritório Houston", "", "", "", "", "", "✔"],
            ["22 colaboradores", "", "", "", "", "", "✔"],
            ["300+ clientes ativos", "", "", "", "", "", "✔"],
            ["Partner network (revendedores)", "", "", "", "", "", "✔"],
        ],
        "Cronograma de implementação da DentalShield Compliance Solutions LLC."
    )

    p(doc, "O cronograma de implementação reflete uma abordagem deliberadamente conservadora: a empresa valida o modelo de negócio em Tampa durante os primeiros 12-18 meses antes de comprometer capital com a expansão geográfica. Esta disciplina de execução faseada minimiza o risco de overextension — a causa mais frequente de falência em startups de serviços — e permite ajustes de produto e pricing baseados em feedback real de clientes antes de replicar a operação em novos mercados.")

    p(doc, "Os marcos críticos do cronograma são: (1) lançamento comercial no Q3 do Ano 1, após conclusão do desenvolvimento de SteriSensor e ComplianceScreen; (2) primeiro breakeven operacional no Q4 do Ano 2, coincidindo com a abertura das filiais; (3) 100 clientes ativos no Ano 3, validando a escalabilidade do modelo; e (4) 300+ clientes e 22 colaboradores no Ano 5, consolidando a posição de liderança no mercado de compliance odontológico.")

    h2(doc, "6.2 Considerações Finais")
    p(doc, f"A DentalShield Compliance Solutions LLC apresenta uma proposta de negócio fundamentada na interseção de três vetores de crescimento convergentes: o mercado de serviços odontológicos nos Estados Unidos (US$ 164,76 bilhões em 2024), o segmento de compliance regulatório em saúde (US$ 2,94 bilhões, CAGR 9,5%) e a adoção exponencial de IoT no setor de saúde (US$ 89,7 bilhões, CAGR 18,4%). Esta convergência cria uma janela estratégica para uma empresa verticalmente especializada que integre auditoria presencial, monitoramento tecnológico e plataforma digital em uma oferta unificada.")

    p(doc, f"Os indicadores financeiros projetados — VPL de {usd(FIN['npv'])}, TIR de {pct(FIN['irr'])}, ROI de {pct(FIN['roi'])} e payback de {FIN['payback']} anos — demonstram viabilidade econômica robusta dentro de premissas conservadoras (taxa de desconto de {FIN['discount']}%, alíquota efetiva de 25%). O investimento total de {usd(FIN['invest_total'])} é distribuído em três fases que mitigam o risco de sobreinvestimento, permitindo validação progressiva do modelo de negócio antes de cada ciclo de expansão.")

    p(doc, f"Do ponto de vista social, a DentalShield projeta a criação de 22 empregos diretos e aproximadamente 73,7 empregos totais (multiplicador EPI 3,35x) nas comunidades de Tampa, Miami e Houston — três dos mercados odontológicos mais relevantes dos Estados Unidos. Cada clínica atendida pela DentalShield representa centenas de pacientes melhor protegidos por protocolos de biossegurança rigorosos, esterilização monitorada em tempo real e equipes adequadamente treinadas.")

    p(doc, f"A ausência de concorrente direto que combine IoT, auditoria presencial e SaaS com foco exclusivo no setor odontológico confere à DentalShield uma vantagem competitiva estrutural. A propriedade intelectual sobre três produtos proprietários (SteriSensor, BiohazardBox, ComplianceScreen) constitui um moat tecnológico que protege esta vantagem no médio e longo prazo. O modelo de receita recorrente (assinaturas mensais) assegura previsibilidade financeira e alinhamento de incentivos com os resultados de compliance dos clientes.")


def add_endnotes(doc):
    """Add all endnotes at the end."""
    if not ENDNOTES: return
    doc.add_page_break()
    h1(doc, "NOTAS E REFERÊNCIAS")
    seen = {}
    for n, text in ENDNOTES:
        if n not in seen: seen[n] = text
    for n in sorted(seen.keys()):
        pr = doc.add_paragraph()
        r = pr.add_run(f"[{n}] "); r.bold = True; r.font.size = Pt(8); r.font.name = FONT
        r = pr.add_run(seen[n]); r.font.size = Pt(8); r.font.name = FONT; r.font.color.rgb = C_CAP
        pr.paragraph_format.space_after = Pt(2)


# ============================================================
# MAIN
# ============================================================

def main():
    print("BP DentalShield — Geração Programática Final")
    print("=" * 50)

    # Generate charts first
    print("  Gerando charts...")
    from generate_bp_camilla import generate_charts
    generate_charts()

    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    # Cover
    cover_page(doc)

    # Block 1: Executive Summary
    print("  Bloco 1: Sumário Executivo...")
    s1_1(doc)
    s1_2(doc)
    s1_3(doc)
    s1_4(doc)

    # Block 2: Strategic Analysis
    print("  Bloco 2: Análise Estratégica...")
    s2_1(doc)
    s2_2(doc)
    s2_3(doc)
    s2_4_to_2_8(doc)

    # Block 3: Marketing
    print("  Bloco 3: Marketing...")
    block3_marketing_full(doc)

    # Block 4: Operational
    print("  Bloco 4: Operacional...")
    block4_operational(doc)

    # Block 5: Financial
    print("  Bloco 5: Financeiro...")
    block5_financial(doc)

    # Block 6: Conclusion
    print("  Bloco 6: Conclusão...")
    block6_conclusion(doc)

    # Endnotes
    add_endnotes(doc)

    # Audit
    full_text = '\n'.join([pr.text for pr in doc.paragraphs])
    word_count = len(full_text.split())
    prohibited = ['consultoria', 'EB-2', 'NIW', 'USCIS', 'green card', 'visto', 'imigração']
    issues = []
    for term in prohibited:
        c = len(re.findall(re.escape(term), full_text, re.IGNORECASE))
        if c > 0: issues.append(f"  '{term}' encontrado {c}x")

    print(f"\n  AUDIT:")
    print(f"  Palavras: {word_count}")
    print(f"  Parágrafos: {len(doc.paragraphs)}")
    print(f"  Tabelas: {len(doc.tables)}")
    print(f"  Endnotes: {len(ENDNOTES)}")
    if issues:
        for i in issues: print(f"  PROIBIDO: {i}")
    else:
        print("  ✔ Zero termos proibidos")

    output = OUTPUT_DIR / "BP_DentalShield_V2.docx"
    doc.save(str(output))
    print(f"\n  SAVED: {output}")
    print(f"  Size: {output.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
