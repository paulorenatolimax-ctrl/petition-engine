#!/usr/bin/env python3
"""
Post-processor: Convert [N] markers to REAL Word footnotes (bottom of page).
Opens the generated DOCX, finds [N] markers, creates footnote definitions,
inserts footnote references inline.
"""

import re, copy
from pathlib import Path
from lxml import etree
from docx import Document
from docx.opc.part import Part
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
FONT = "Garamond"

def create_footnotes_part(doc):
    """Create the footnotes.xml part with proper separators."""
    # Build footnotes XML
    footnotes_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
    <w:footnote w:type="separator" w:id="0">
        <w:p>
            <w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
            <w:r><w:separator/></w:r>
        </w:p>
    </w:footnote>
    <w:footnote w:type="continuationSeparator" w:id="-1">
        <w:p>
            <w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
            <w:r><w:continuationSeparator/></w:r>
        </w:p>
    </w:footnote>
</w:footnotes>'''

    return footnotes_xml


def add_footnote_to_xml(footnotes_root, fn_id, fn_text):
    """Add a footnote definition to the footnotes XML."""
    footnote = etree.SubElement(footnotes_root, f'{W}footnote')
    footnote.set(f'{W}id', str(fn_id))

    p_elem = etree.SubElement(footnote, f'{W}p')

    # Paragraph properties with FootnoteText style
    pPr = etree.SubElement(p_elem, f'{W}pPr')
    pStyle = etree.SubElement(pPr, f'{W}pStyle')
    pStyle.set(f'{W}val', 'FootnoteText')
    spacing = etree.SubElement(pPr, f'{W}spacing')
    spacing.set(f'{W}after', '0')
    spacing.set(f'{W}line', '240')
    spacing.set(f'{W}lineRule', 'auto')

    # Footnote reference run (the superscript number in the footnote itself)
    r1 = etree.SubElement(p_elem, f'{W}r')
    rPr1 = etree.SubElement(r1, f'{W}rPr')
    rStyle1 = etree.SubElement(rPr1, f'{W}rStyle')
    rStyle1.set(f'{W}val', 'FootnoteReference')
    etree.SubElement(r1, f'{W}footnoteRef')

    # Space after number
    r2 = etree.SubElement(p_elem, f'{W}r')
    rPr2 = etree.SubElement(r2, f'{W}rPr')
    sz2 = etree.SubElement(rPr2, f'{W}sz')
    sz2.set(f'{W}val', '16')  # 8pt
    szCs2 = etree.SubElement(rPr2, f'{W}szCs')
    szCs2.set(f'{W}val', '16')
    rFonts2 = etree.SubElement(rPr2, f'{W}rFonts')
    rFonts2.set(f'{W}ascii', FONT)
    rFonts2.set(f'{W}hAnsi', FONT)
    t2 = etree.SubElement(r2, f'{W}t')
    t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t2.text = f' {fn_text}'


def insert_footnote_reference(run_element, fn_id):
    """Replace a text run containing [N] with a footnote reference."""
    # Create footnote reference run
    r = etree.Element(f'{W}r')
    rPr = etree.SubElement(r, f'{W}rPr')
    rStyle = etree.SubElement(rPr, f'{W}rStyle')
    rStyle.set(f'{W}val', 'FootnoteReference')
    fnRef = etree.SubElement(r, f'{W}footnoteReference')
    fnRef.set(f'{W}id', str(fn_id))
    return r


def process_docx(input_path, output_path, endnotes_map):
    """
    Open DOCX, find [N] markers, replace with real footnotes.
    endnotes_map: dict of {N: "source text"}
    """
    doc = Document(input_path)

    # Step 1: Create footnotes XML
    footnotes_xml = create_footnotes_part(doc)
    footnotes_root = etree.fromstring(footnotes_xml.encode('utf-8'))

    # Step 2: Add all footnote definitions
    for fn_id in sorted(endnotes_map.keys()):
        add_footnote_to_xml(footnotes_root, fn_id, endnotes_map[fn_id])

    # Step 3: Find and replace [N] markers in the document body
    replacements_made = 0
    for paragraph in doc.paragraphs:
        p_elem = paragraph._p
        runs = list(p_elem.findall(f'{W}r'))

        for run in runs:
            t_elem = run.find(f'{W}t')
            if t_elem is None or t_elem.text is None:
                continue

            text = t_elem.text
            # Find all [N] patterns
            matches = list(re.finditer(r'\[(\d+)\]', text))
            if not matches:
                continue

            # Need to split the run into parts: text before, footnote ref, text after, etc.
            parent = run.getparent()
            run_index = list(parent).index(run)

            # Get the run properties to copy to new runs
            rPr = run.find(f'{W}rPr')

            # Build new elements
            new_elements = []
            last_end = 0

            for match in matches:
                fn_num = int(match.group(1))
                start, end = match.start(), match.end()

                # Text before the [N]
                if start > last_end:
                    before_text = text[last_end:start]
                    if before_text:
                        new_r = etree.Element(f'{W}r')
                        if rPr is not None:
                            new_r.append(copy.deepcopy(rPr))
                        new_t = etree.SubElement(new_r, f'{W}t')
                        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        new_t.text = before_text
                        new_elements.append(new_r)

                # Footnote reference (only if we have the footnote defined)
                if fn_num in endnotes_map:
                    fn_ref = insert_footnote_reference(run, fn_num)
                    new_elements.append(fn_ref)
                    replacements_made += 1
                else:
                    # Keep the [N] as text if not in our map
                    new_r = etree.Element(f'{W}r')
                    if rPr is not None:
                        new_r.append(copy.deepcopy(rPr))
                    new_t = etree.SubElement(new_r, f'{W}t')
                    new_t.text = f'[{fn_num}]'
                    new_elements.append(new_r)

                last_end = end

            # Text after the last match
            if last_end < len(text):
                after_text = text[last_end:]
                if after_text:
                    new_r = etree.Element(f'{W}r')
                    if rPr is not None:
                        new_r.append(copy.deepcopy(rPr))
                    new_t = etree.SubElement(new_r, f'{W}t')
                    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    new_t.text = after_text
                    new_elements.append(new_r)

            # Replace the original run with new elements
            for new_elem in reversed(new_elements):
                run.addnext(new_elem)
            parent.remove(run)

    # Step 4: Remove the endnotes section from the document body
    # Find and remove the last H1 "NOTAS E REFERÊNCIAS" and everything after
    found_endnotes_heading = False
    paragraphs_to_remove = []
    for para in doc.paragraphs:
        try:
            txt = para.text or ""
        except:
            txt = ""
        if "NOTAS E REFERÊNCIAS" in txt or "ENDNOTES" in txt:
            found_endnotes_heading = True
        if found_endnotes_heading:
            paragraphs_to_remove.append(para)

    for para in paragraphs_to_remove:
        parent = para._p.getparent()
        if parent is not None:
            parent.remove(para._p)

    # Step 5: Save footnotes part into the DOCX package
    # We need to add the footnotes.xml as a part
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    footnotes_bytes = etree.tostring(footnotes_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Get the main document part
    main_part = doc.part

    # Check if footnotes relationship already exists
    existing_rels = [rel for rel in main_part.rels.values()
                     if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes']

    if existing_rels:
        # Update existing footnotes part
        fn_part = existing_rels[0].target_part
        fn_part._blob = footnotes_bytes
    else:
        # Create new footnotes part
        fn_partname = PackURI('/word/footnotes.xml')
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
        fn_part = Part(fn_partname, content_type, footnotes_bytes, main_part.package)
        main_part.relate_to(fn_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')

    # Step 6: Ensure FootnoteReference style exists
    # This is handled by Word itself when it opens the file

    doc.save(output_path)
    print(f"  Footnotes: {replacements_made} references inserted")
    print(f"  Endnotes section removed from body")
    return replacements_made


def main():
    """Post-process the BP to add real footnotes."""
    SCRIPT_DIR = Path(__file__).parent
    OUTPUT_DIR = Path("/Users/paulo1844/Documents/_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2025/CAMILLA/_BP/_Atualizado (pós-entendimento novo)")

    input_path = OUTPUT_DIR / "BP_DentalShield_V2.docx"
    output_path = OUTPUT_DIR / "BP_DentalShield_V2_FOOTNOTES.docx"

    if not input_path.exists():
        print(f"ERROR: {input_path} not found")
        return

    # Import the endnotes from the generation script
    import sys
    sys.path.insert(0, str(SCRIPT_DIR))

    # Re-run endnote collection by importing
    from generate_bp_final import ENDNOTES, fn, _fn_counter

    # Reset and re-collect
    ENDNOTES.clear()
    _fn_counter[0] = 0

    # We need to re-generate to collect endnotes, or parse them from the existing file
    # Easier: parse the endnotes section from the existing DOCX
    doc_temp = Document(str(input_path))
    endnotes_map = {}
    collecting = False
    for para in doc_temp.paragraphs:
        if "NOTAS E REFERÊNCIAS" in para.text or "ENDNOTES" in para.text:
            collecting = True
            continue
        if collecting and para.text.strip():
            match = re.match(r'\[(\d+)\]\s*(.*)', para.text.strip())
            if match:
                endnotes_map[int(match.group(1))] = match.group(2)

    print(f"  Found {len(endnotes_map)} endnotes to convert to footnotes")

    if not endnotes_map:
        print("  ERROR: No endnotes found in document")
        return

    # Process
    count = process_docx(str(input_path), str(output_path), endnotes_map)

    print(f"\n  SAVED: {output_path}")
    print(f"  Size: {output_path.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
