#!/usr/bin/env python3
"""
BP DOCX Assembler V2 — Programmatic generation (no markdown parsing)
Reads raw block .md files and builds DOCX directly using python-docx helpers.
Based on generate_bp_deni.py patterns.

Usage: python3 assemble_docx.py
"""

import re, os, json
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# === PATHS ===
SCRIPT_DIR = Path(__file__).parent
OUTPUT_DIR = Path("/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2025/CAMILLA/_BP/_Atualizado (pós-entendimento novo)")

# === DESIGN CONSTANTS ===
FONT = "Garamond"
COLOR_H1 = RGBColor(0x58, 0x4D, 0x42)       # Brown dark
COLOR_H2 = RGBColor(0x58, 0x4D, 0x42)
COLOR_H3 = RGBColor(0x2A, 0x2A, 0x2A)       # Dark gray
COLOR_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_CAPTION = RGBColor(0x66, 0x66, 0x66)
COLOR_FOOTER = RGBColor(0x58, 0x4D, 0x42)
TABLE_HEADER_HEX = "DEDACB"
TABLE_ZEBRA_HEX = "D0DDD6"
TABLE_BORDER_HEX = "CCCCCC"
KPI_BG_HEX = "F0F0F0"

# ============================================================
# HELPER FUNCTIONS (following generate_bp_deni.py pattern)
# ============================================================

def setup_styles(doc):
    """Configure document-wide styles."""
    # Normal
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_TEXT
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(3)
    pf.line_spacing = 1.15

    # Headings
    for level, size, color in [(1, 16, COLOR_H1), (2, 13, COLOR_H2), (3, 11, COLOR_H3)]:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = FONT
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = color
        hs.paragraph_format.space_before = Pt(18 if level == 1 else 12 if level == 2 else 6)
        hs.paragraph_format.space_after = Pt(8 if level == 1 else 6)
        hs.paragraph_format.keep_with_next = True


def setup_page(doc):
    """Configure page layout, header, footer."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.14)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("DentalShield Compliance Solutions LLC")
    run.font.name = FONT
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_CAPTION
    run.italic = True

    # Footer with brown bar
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Top border on footer
    pPr = fp._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="12" w:space="1" w:color="584D42"/></w:pBdr>')
    pPr.append(pBdr)

    run = fp.add_run("CONFIDENTIAL — DO NOT SHARE  |  DentalShield Compliance Solutions LLC  |  Page ")
    run.font.name = FONT
    run.font.size = Pt(7)
    run.font.bold = True
    run.font.color.rgb = COLOR_FOOTER

    # Page number field
    fld = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}"/><w:sz w:val="14"/><w:b/><w:color w:val="584D42"/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>')
    fp._p.append(fld)
    fld2 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}"/><w:sz w:val="14"/><w:b/><w:color w:val="584D42"/></w:rPr><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>')
    fp._p.append(fld2)
    fld3 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}"/><w:sz w:val="14"/><w:b/><w:color w:val="584D42"/></w:rPr><w:fldChar w:fldCharType="end"/></w:r>')
    fp._p.append(fld3)


def h1(doc, text):
    """Add H1 heading with page break before (except first)."""
    existing_h1 = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
    if len(existing_h1) > 0:
        doc.add_page_break()
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def para(doc, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add paragraph with inline formatting: **bold**, [N] superscript footnotes."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)

    # Split by bold markers and footnote refs
    tokens = re.split(r'(\*\*[^*]+\*\*|\[\d+\])', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = p.add_run(tok[2:-2])
            run.bold = True
            run.font.name = FONT
            run.font.size = Pt(11)
        elif re.match(r'^\[\d+\]$', tok):
            run = p.add_run(tok)
            run.font.superscript = True
            run.font.size = Pt(8)
            run.font.name = FONT
            run.font.color.rgb = COLOR_CAPTION
        else:
            run = p.add_run(tok)
            run.font.name = FONT
            run.font.size = Pt(11)
            if bold:
                run.bold = True
            if italic:
                run.italic = True

    return p


def add_table(doc, headers, rows, caption=None):
    """Add formatted table with colored header and zebra striping."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        cp = cell.paragraphs[0]
        run = cp.add_run(str(header_text))
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(10)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_HEADER_HEX}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx in range(min(len(row_data), num_cols)):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            cp = cell.paragraphs[0]
            run = cp.add_run(str(row_data[c_idx]))
            run.font.name = FONT
            run.font.size = Pt(10)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_ZEBRA_HEX}"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_HEX}"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_HEX}"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_HEX}"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_HEX}"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_HEX}"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_HEX}"/>
    </w:tblBorders>''')
    tblPr.append(borders)

    # Caption
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(f'Fonte: {caption}')
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_CAPTION
        run.font.name = FONT
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Spacing after
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(3)
    sp.paragraph_format.space_before = Pt(0)


def add_kpi_box(doc, lines):
    """Add a KPI highlight box with gray background."""
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Gray background via shading
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{KPI_BG_HEX}"/>')
        pPr.append(shading)
        # Left border
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="4" w:color="584D42"/></w:pBdr>')
        pPr.append(pBdr)
        # Indent
        ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="288" w:right="288"/>')
        pPr.append(ind)

        run = p.add_run(line)
        run.font.name = FONT
        run.font.size = Pt(11)
        if line.startswith('►'):
            run.bold = True


def add_image(doc, image_path, caption_text, width=Inches(6.3)):
    """Insert image with caption."""
    if not os.path.exists(image_path):
        return
    doc.add_paragraph()  # spacing
    doc.add_picture(image_path, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(f'Figura: {caption_text}')
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = FONT
    run.font.color.rgb = COLOR_CAPTION
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_divider(doc):
    """Add a thin horizontal divider line."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="{TABLE_HEADER_HEX}"/></w:pBdr>')
    pPr.append(pBdr)


# ============================================================
# SMART MARKDOWN → DOCX CONVERTER
# ============================================================

def convert_md_to_docx(doc, md_content, endnotes_list):
    """
    Convert markdown content to DOCX elements using helper functions.
    Much more robust than the previous regex-based parser.
    """
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_headers = []
    table_rows = []
    table_caption = None
    collecting_footnotes = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # === EMPTY LINE ===
        if not stripped:
            # Flush table if we were in one
            if in_table and table_headers:
                add_table(doc, table_headers, table_rows, table_caption)
                table_headers, table_rows, table_caption = [], [], None
                in_table = False
            i += 1
            continue

        # === FOOTNOTE COLLECTION (at end of sections) ===
        fn_match = re.match(r'^\[(\d+)\]\s+(.+)', stripped)
        if fn_match:
            endnotes_list.append((int(fn_match.group(1)), fn_match.group(2)))
            i += 1
            continue

        # Also catch "Fontes" / "Fonte" section markers
        if re.match(r'^\*?\*?Fontes?\b', stripped) or stripped.startswith('**Fontes'):
            collecting_footnotes = True
            i += 1
            continue

        if collecting_footnotes:
            fn2 = re.match(r'^\[?(\d+)\]?\s*(.+)', stripped)
            if fn2:
                endnotes_list.append((int(fn2.group(1)), fn2.group(2)))
                i += 1
                continue
            else:
                collecting_footnotes = False
                # Fall through to process this line normally

        # === HORIZONTAL RULE ===
        if stripped == '---' or stripped == '***':
            if in_table and table_headers:
                add_table(doc, table_headers, table_rows, table_caption)
                table_headers, table_rows, table_caption = [], [], None
                in_table = False
            add_divider(doc)
            i += 1
            continue

        # === HEADINGS ===
        if stripped.startswith('# ') and not stripped.startswith('## '):
            if in_table and table_headers:
                add_table(doc, table_headers, table_rows, table_caption)
                table_headers, table_rows, table_caption = [], [], None
                in_table = False
            h1(doc, stripped[2:].strip())
            i += 1
            continue

        if stripped.startswith('## '):
            if in_table and table_headers:
                add_table(doc, table_headers, table_rows, table_caption)
                table_headers, table_rows, table_caption = [], [], None
                in_table = False
            h1(doc, stripped[3:].strip())
            i += 1
            continue

        if stripped.startswith('### '):
            if in_table and table_headers:
                add_table(doc, table_headers, table_rows, table_caption)
                table_headers, table_rows, table_caption = [], [], None
                in_table = False
            h2(doc, stripped[4:].strip())
            i += 1
            continue

        if stripped.startswith('#### '):
            if in_table and table_headers:
                add_table(doc, table_headers, table_rows, table_caption)
                table_headers, table_rows, table_caption = [], [], None
                in_table = False
            h3(doc, stripped[5:].strip())
            i += 1
            continue

        # === BLOCKQUOTE (KPI box) ===
        if stripped.startswith('> '):
            kpi_lines = []
            while i < len(lines) and lines[i].strip().startswith('> '):
                kpi_lines.append(lines[i].strip()[2:].strip())
                i += 1
            add_kpi_box(doc, kpi_lines)
            continue

        # === TABLE ===
        if '|' in stripped and not stripped.startswith('```'):
            cells = [c.strip() for c in stripped.split('|')]
            # Remove empty first/last from leading/trailing pipes
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]

            if not cells:
                i += 1
                continue

            # Separator row (---|---|---)
            if all(re.match(r'^[-:]+$', c.strip()) for c in cells if c.strip()):
                i += 1
                continue

            if not in_table:
                # This is the header row
                table_headers = cells
                in_table = True
            else:
                # Data row — pad to match header count
                while len(cells) < len(table_headers):
                    cells.append('')
                table_rows.append(cells[:len(table_headers)])

            i += 1
            continue

        # === BULLET POINTS ===
        if stripped.startswith('- ') or stripped.startswith('* ') or re.match(r'^\d+\.\s', stripped):
            if in_table and table_headers:
                add_table(doc, table_headers, table_rows, table_caption)
                table_headers, table_rows, table_caption = [], [], None
                in_table = False

            bullet_text = re.sub(r'^[-*]\s+|^\d+\.\s+', '', stripped)
            para(doc, f'• {bullet_text}')
            i += 1
            continue

        # === CODE BLOCKS (skip) ===
        if stripped.startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1  # skip closing ```
            continue

        # === REGULAR PARAGRAPH ===
        if in_table and table_headers:
            add_table(doc, table_headers, table_rows, table_caption)
            table_headers, table_rows, table_caption = [], [], None
            in_table = False

        para(doc, stripped)
        i += 1

    # Flush remaining table
    if in_table and table_headers:
        add_table(doc, table_headers, table_rows, table_caption)


# ============================================================
# COVER PAGE
# ============================================================

def add_cover_page(doc):
    """Add professional cover page."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BUSINESS PLAN")
    run.font.name = FONT
    run.font.size = Pt(32)
    run.font.color.rgb = COLOR_H1
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DentalShield Compliance Solutions LLC")
    run.font.name = FONT
    run.font.size = Pt(20)
    run.font.color.rgb = COLOR_H1

    doc.add_paragraph()

    for text in [
        "Dental Regulatory Compliance & IoT Monitoring",
        "",
        "Tampa, FL  |  Miami, FL  |  Houston, TX",
        "",
        f"Prepared: {datetime.now().strftime('%B %Y')}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_TEXT

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL — DO NOT SHARE")
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_FOOTER

    doc.add_page_break()


# ============================================================
# ENDNOTES SECTION
# ============================================================

def add_endnotes_section(doc, endnotes):
    """Add endnotes/references at the end."""
    if not endnotes:
        return

    doc.add_page_break()
    h1(doc, "NOTAS E REFERÊNCIAS / ENDNOTES")

    # Deduplicate and sort
    seen = {}
    for num, text in endnotes:
        if num not in seen:
            seen[num] = text

    for num in sorted(seen.keys()):
        p = doc.add_paragraph()
        run = p.add_run(f'[{num}] ')
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = FONT
        run = p.add_run(seen[num])
        run.font.size = Pt(8)
        run.font.name = FONT
        run.font.color.rgb = COLOR_CAPTION
        p.paragraph_format.space_after = Pt(2)


# ============================================================
# MAIN
# ============================================================

def main():
    print("BP DOCX Assembler V2 — DentalShield")
    print("=" * 50)

    # Create document
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    # Cover page
    add_cover_page(doc)

    # Endnotes collection
    all_endnotes = []

    # Process 4 blocks
    charts_dir = SCRIPT_DIR / "_charts"
    chart_map = {
        "3.4": ("04_services.png", "Revenue Distribution by Service Line"),
        "4.1": ("03_employees.png", "Workforce Growth by Location — 5-Year Projection"),
        "5.2": ("06_investment.png", "Investment Structure by Phase"),
        "5.3a": ("01_revenue_vs_costs.png", "Revenue vs Total Costs — 5-Year Projection"),
        "5.3b": ("08_cost_structure.png", "Cost Structure — Year 1 vs Year 5"),
        "5.4a": ("02_ebitda.png", "EBITDA & EBITDA Margin — 5-Year Projection"),
        "5.4b": ("07_net_income.png", "Net Income — 5-Year Projection"),
        "5.6": ("05_breakeven.png", "Break-Even Analysis — Revenue vs BEP"),
    }

    for block_num in range(1, 5):
        block_file = SCRIPT_DIR / f"_output_block{block_num}.md"
        if not block_file.exists():
            print(f"  WARNING: Block {block_num} not found, skipping")
            continue

        print(f"  Processing Block {block_num}...")
        with open(block_file) as f:
            content = f.read()

        # Fix prohibited terms
        content = content.replace('consultórios', 'clínicas odontológicas')
        content = content.replace('consultório', 'clínica odontológica')
        content = content.replace('Consultórios', 'Clínicas Odontológicas')
        content = content.replace('Consultório', 'Clínica Odontológica')
        content = content.replace('consultorios', 'clínicas odontológicas')
        content = content.replace('consultorio', 'clínica odontológica')
        content = content.replace('consultoria', 'assessoria especializada')
        content = content.replace('consultor ', 'especialista ')
        content = content.replace('Consultor', 'Especialista')

        convert_md_to_docx(doc, content, all_endnotes)

        # Insert charts after relevant sections
        if block_num == 3:
            for key in ["3.4", "4.1"]:
                if key in chart_map:
                    fname, caption = chart_map[key]
                    add_image(doc, str(charts_dir / fname), caption)

        if block_num == 4:
            for key in ["5.2", "5.3a", "5.3b", "5.4a", "5.4b", "5.6"]:
                if key in chart_map:
                    fname, caption = chart_map[key]
                    add_image(doc, str(charts_dir / fname), caption)

    # Add endnotes
    add_endnotes_section(doc, all_endnotes)

    # Audit
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    word_count = len(full_text.split())
    table_count = len(doc.tables)
    para_count = len(doc.paragraphs)

    prohibited = ['consultoria', 'consultor ', 'EB-2', 'NIW', 'USCIS', 'green card',
                   'visto', 'imigração', 'immigration', 'petition', 'sponsor', 'waiver']
    issues = []
    for term in prohibited:
        count = len(re.findall(re.escape(term), full_text, re.IGNORECASE))
        if count > 0:
            issues.append(f"  CRITICAL: '{term}' found {count}x")

    print(f"\n{'='*50}")
    print(f"  AUDIT:")
    print(f"  Words: {word_count}")
    print(f"  Paragraphs: {para_count}")
    print(f"  Tables: {table_count}")
    print(f"  Endnotes: {len(all_endnotes)}")
    if issues:
        for issue in issues:
            print(issue)
    else:
        print("  Zero prohibited terms!")
    print(f"{'='*50}")

    # Save
    output_path = OUTPUT_DIR / "BP_DentalShield_V2.docx"
    doc.save(str(output_path))
    print(f"\n  SAVED: {output_path}")
    print(f"  Size: {output_path.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
