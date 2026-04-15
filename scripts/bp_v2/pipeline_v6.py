#!/usr/bin/env python3
"""
BP DentalShield V6 — Pipeline Definitivo
Fixes: cover page, TOC, page breaks, footnote size 9pt single spacing,
       white tables, justified text, charts inline, dedup footnote refs
"""
import subprocess, re, os
from pathlib import Path

BP_DIR = Path("/Users/paulo1844/Documents/_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2025/CAMILLA/_BP/_Atualizado (pós-entendimento novo)")
CHARTS_DIR = Path("/Users/paulo1844/Documents/_OMNI/_IMIGRAÇÃO/Sistema Automatizado/petition-engine/scripts/bp_v2/_charts")
MD_INPUT  = BP_DIR / "BP_DentalShield_FULL.md"
REF_DOC   = BP_DIR / "reference.docx"
MD_FIXED  = BP_DIR / "_BP_fixed_v6.md"
OUTPUT    = BP_DIR / "BP_DentalShield_V6_FINAL.docx"


def fix_markdown(md_text):
    # 1. Chart markers → images
    def replace_chart(m):
        fn, cap = m.group(1).strip(), m.group(2).strip()
        p = CHARTS_DIR / fn
        return f'\n![{cap}]({p}){{ width=85% }}\n' if p.exists() else f'\n*[Gráfico: {cap}]*\n'
    md_text = re.sub(r'<!--\s*CHART:\s*([^|]+)\|\s*([^>]+)-->', replace_chart, md_text)

    # 2. IDEM footnotes
    for fn, txt in {27:'*Idem*, nota 20.',28:'*Idem*, nota 21.',35:'*Idem*, nota 2.',
                     39:'*Idem*, nota 1.',40:'*Idem*, nota 12.',41:'*Idem*, nota 10.',42:'*Idem*, nota 11.'}.items():
        md_text = re.sub(rf'(\[\^{fn}\]:)\s*.+$', rf'\1 {txt}', md_text, flags=re.MULTILINE)

    # 3. Clean
    md_text = re.sub(r'(?m)^.*Fim do bloco.*$\n?', '', md_text)
    md_text = re.sub(r'/\s*English Translation', '', md_text)

    # 4. Deduplicate body footnote references: if same [^N] appears multiple times,
    #    keep only the FIRST occurrence, remove subsequent ones
    fn_seen = set()
    def dedup_fn(m):
        fn_num = m.group(0)
        if fn_num in fn_seen:
            return ''  # Remove duplicate reference
        fn_seen.add(fn_num)
        return fn_num

    # Process line by line to only dedup in body text (not in definitions)
    lines = md_text.split('\n')
    new_lines = []
    for line in lines:
        if re.match(r'^\[\^\d+\]:', line):
            new_lines.append(line)  # Keep definitions as-is
        else:
            fn_seen_line = set()
            # Remove duplicate refs within same line AND across lines
            def dedup_in_line(m):
                key = m.group(0)
                if key in fn_seen:
                    return ''
                fn_seen.add(key)
                return key
            line = re.sub(r'\[\^\d+\]', dedup_in_line, line)
            new_lines.append(line)
    md_text = '\n'.join(new_lines)

    # 5. Page breaks before # and ## headings (correct single backslash)
    lines = md_text.split('\n')
    out = []
    for i, line in enumerate(lines):
        if i > 10 and (re.match(r'^#\s+\d+\.', line) or re.match(r'^##\s+\d+\.\d+', line)):
            out.extend(['', '\\newpage', ''])
        out.append(line)
    md_text = '\n'.join(out)

    return md_text


def run_pandoc(md_path, output_path):
    cmd = [
        'pandoc', str(md_path),
        '--from', 'markdown+footnotes+pipe_tables+yaml_metadata_block+implicit_figures',
        '--to', 'docx',
        '--reference-doc', str(REF_DOC),
        '--toc', '--toc-depth=3',
        '-o', str(output_path),
    ]
    r = subprocess.run(cmd, capture_output=True, text=True)
    if r.returncode != 0:
        print(f"PANDOC ERROR: {r.stderr}")
        return False
    w = [l for l in (r.stderr or '').strip().split('\n') if l.strip()]
    print(f"  Pandoc OK ({len(w)} warnings)")
    return True


def post_process(docx_path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(str(docx_path))

    # ── COVER PAGE ──
    # Insert cover page BEFORE everything else
    # Add a page break at the very beginning, then insert cover content before it
    cover_paras = []

    # Add empty paragraphs for vertical centering
    for _ in range(6):
        p = doc.paragraphs[0].insert_paragraph_before('')
        cover_paras.append(p)

    # Company name
    p = doc.paragraphs[0].insert_paragraph_before('')
    run = p.add_run('BUSINESS PLAN')
    run.font.name = 'Garamond'
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x58, 0x4D, 0x42)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_paras.append(p)

    p = doc.paragraphs[0].insert_paragraph_before('')
    cover_paras.append(p)

    p = doc.paragraphs[0].insert_paragraph_before('')
    run = p.add_run('DentalShield Compliance Solutions LLC')
    run.font.name = 'Garamond'
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x58, 0x4D, 0x42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_paras.append(p)

    p = doc.paragraphs[0].insert_paragraph_before('')
    cover_paras.append(p)

    p = doc.paragraphs[0].insert_paragraph_before('')
    run = p.add_run('Dental Regulatory Compliance & IoT Monitoring')
    run.font.name = 'Garamond'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_paras.append(p)

    for _ in range(4):
        p = doc.paragraphs[0].insert_paragraph_before('')
        cover_paras.append(p)

    p = doc.paragraphs[0].insert_paragraph_before('')
    run = p.add_run('PROPOSED BY:')
    run.font.name = 'Garamond'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x58, 0x4D, 0x42)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_paras.append(p)

    p = doc.paragraphs[0].insert_paragraph_before('')
    run = p.add_run('Camilla de Oliveira Lopes')
    run.font.name = 'Garamond'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x58, 0x4D, 0x42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_paras.append(p)

    p = doc.paragraphs[0].insert_paragraph_before('')
    run = p.add_run('Tampa, Florida — April 2026')
    run.font.name = 'Garamond'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_paras.append(p)

    # Page break after cover
    p = doc.paragraphs[0].insert_paragraph_before('')
    run = p.add_run()
    run.add_break()  # page break
    from docx.enum.text import WD_BREAK
    # Actually use proper page break
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

    # ── JUSTIFY ALL PARAGRAPHS ──
    for para in doc.paragraphs:
        sn = (para.style.name or '').lower()
        if 'toc' in sn or 'title' in sn or 'subtitle' in sn:
            continue
        if 'heading' in sn:
            for r in para.runs:
                r.font.name = 'Garamond'
                r.font.color.rgb = RGBColor(0x58, 0x4D, 0x42)
            continue
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── FIX TABLES ──
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                for shd in tcPr.findall(qn('w:shd')):
                    tcPr.remove(shd)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DEDACB' if i == 0 else 'FFFFFF')
                tcPr.append(shd)
                if i == 0:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.bold = True
                            r.font.color.rgb = RGBColor(0x58, 0x4D, 0x42)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Garamond'
                        r.font.size = Pt(9)

        # Borders
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        for b in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(b)
        borders = OxmlElement('w:tblBorders')
        for bn in ['top','left','bottom','right','insideH','insideV']:
            b = OxmlElement(f'w:{bn}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'B0A898')
            borders.append(b)
        tblPr.append(borders)

    # ── GARAMOND ON BODY ──
    for para in doc.paragraphs:
        sn = (para.style.name or '').lower()
        if 'heading' in sn or 'title' in sn:
            continue
        for r in para.runs:
            r.font.name = 'Garamond'
            if r.font.size is None or ('footnote' not in sn):
                if 'footnote' not in sn:
                    r.font.size = Pt(11)

    # ── FIX FOOTNOTES: size 9pt, single spacing ──
    # Access footnotes via the footnotes part
    try:
        fn_part = doc.part.package.part_related_by(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
        )
        if fn_part is not None:
            from lxml import etree
            W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            fn_root = etree.fromstring(fn_part.blob)
            fn_count = 0

            for footnote in fn_root.findall(f'{{{W}}}footnote'):
                fn_type = footnote.get(f'{{{W}}}type')
                if fn_type in ('separator', 'continuationSeparator'):
                    continue

                for para in footnote.findall(f'{{{W}}}p'):
                    # Set paragraph spacing to single (240 twips = single)
                    pPr = para.find(f'{{{W}}}pPr')
                    if pPr is None:
                        pPr = etree.SubElement(para, f'{{{W}}}pPr')
                        para.insert(0, pPr)

                    # Remove existing spacing
                    for sp in pPr.findall(f'{{{W}}}spacing'):
                        pPr.remove(sp)
                    sp = etree.SubElement(pPr, f'{{{W}}}spacing')
                    sp.set(f'{{{W}}}after', '0')
                    sp.set(f'{{{W}}}before', '0')
                    sp.set(f'{{{W}}}line', '240')
                    sp.set(f'{{{W}}}lineRule', 'auto')

                    # Set font on all runs
                    for run in para.findall(f'{{{W}}}r'):
                        rPr = run.find(f'{{{W}}}rPr')
                        if rPr is None:
                            rPr = etree.SubElement(run, f'{{{W}}}rPr')
                            run.insert(0, rPr)

                        # Remove existing size
                        for sz in rPr.findall(f'{{{W}}}sz'):
                            rPr.remove(sz)
                        for sz in rPr.findall(f'{{{W}}}szCs'):
                            rPr.remove(sz)

                        # Set size 9pt = 18 half-points
                        sz = etree.SubElement(rPr, f'{{{W}}}sz')
                        sz.set(f'{{{W}}}val', '18')
                        szCs = etree.SubElement(rPr, f'{{{W}}}szCs')
                        szCs.set(f'{{{W}}}val', '18')

                        # Set Garamond
                        for rf in rPr.findall(f'{{{W}}}rFonts'):
                            rPr.remove(rf)
                        rFonts = etree.SubElement(rPr, f'{{{W}}}rFonts')
                        rFonts.set(f'{{{W}}}ascii', 'Garamond')
                        rFonts.set(f'{{{W}}}hAnsi', 'Garamond')

                fn_count += 1

            # Save modified footnotes back
            fn_part._blob = etree.tostring(fn_root, xml_declaration=True, encoding='UTF-8', standalone=True)
            print(f"  Fixed {fn_count} footnotes: 9pt Garamond, single spacing")
    except Exception as e:
        print(f"  Footnote fix error: {e}")

    # ── RESIZE IMAGES ──
    for para in doc.paragraphs:
        for run in para.runs:
            for d in run._element.findall(qn('w:drawing')):
                for extent in d.iter(qn('wp:extent')):
                    cx = int(extent.get('cx', 0))
                    cy = int(extent.get('cy', 0))
                    if cx > 0 and cx != 5400000:
                        ratio = 5400000 / cx
                        extent.set('cx', str(5400000))
                        extent.set('cy', str(int(cy * ratio)))

    doc.save(str(docx_path))
    print(f"  Saved: {docx_path.name} ({docx_path.stat().st_size/1024:.0f} KB)")


def main():
    print("=" * 60)
    print("BP DentalShield V6 — Pipeline Definitivo")
    print("=" * 60)

    md = MD_INPUT.read_text(encoding='utf-8')
    print(f"\n[1/3] Fixing markdown ({len(md.split())} words)...")
    md = fix_markdown(md)

    fn_defs = len(re.findall(r'^\[\^\d+\]:', md, re.MULTILINE))
    fn_refs = len(re.findall(r'\[\^\d+\](?!:)', md))
    charts = len(re.findall(r'!\[.*?\]\(.+?\.png\)', md))
    print(f"  Footnotes: {fn_defs} defs, {fn_refs} refs (deduped)")
    print(f"  Charts: {charts}")

    MD_FIXED.write_text(md, encoding='utf-8')

    print(f"\n[2/3] Pandoc + TOC...")
    if not run_pandoc(MD_FIXED, OUTPUT):
        return

    print(f"\n[3/3] Post-processing (cover, tables, footnotes 9pt)...")
    post_process(OUTPUT)

    MD_FIXED.unlink()
    print(f"\n{'='*60}")
    print(f"DONE: {OUTPUT.name}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
