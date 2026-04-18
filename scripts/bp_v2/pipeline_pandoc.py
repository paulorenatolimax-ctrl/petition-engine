#!/usr/bin/env python3
"""
BP DentalShield — Complete Pipeline V4
1. Fix markdown (footnote dedup, chart images, clean residues)
2. Pandoc → DOCX (real footnotes + reference-doc branding)
3. python-docx post-processing (WHITE tables, JUSTIFIED text, chart sizing)
"""

import subprocess, re, os, sys, shutil
from pathlib import Path

# ── Paths ─────────────────────────────────────────────────────────────
BP_DIR = Path("/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2025/CAMILLA/_BP/_Atualizado (pós-entendimento novo)")
CHARTS_DIR = Path("/Users/paulo1844/Documents/_OMNI/_IMIGRAÇÃO/Sistema Automatizado/petition-engine/scripts/bp_v2/_charts")
MD_INPUT  = BP_DIR / "BP_DentalShield_FULL.md"
REF_DOC   = BP_DIR / "reference.docx"
MD_FIXED  = BP_DIR / "_BP_fixed_for_pandoc.md"
OUTPUT    = BP_DIR / "BP_DentalShield_V4_FINAL.docx"


# ── Step 1: Fix Markdown ──────────────────────────────────────────────
def fix_markdown(md_text: str) -> str:
    """Fix footnotes, chart markers, residues, and add page breaks."""

    # 1a. Replace chart markers with pandoc-compatible image syntax
    def replace_chart(match):
        filename = match.group(1).strip()
        caption  = match.group(2).strip()
        img_path = CHARTS_DIR / filename
        if img_path.exists():
            # Use pandoc image syntax with width attribute
            return f'\n![{caption}]({img_path}){{ width=85% }}\n'
        else:
            print(f"  WARNING: chart not found: {filename}")
            return f'\n*[Gráfico: {caption}]*\n'

    md_text = re.sub(
        r'<!--\s*CHART:\s*([^|]+)\|\s*([^>]+)-->',
        replace_chart,
        md_text
    )

    # 1b. Fix duplicate footnote definitions → IDEM
    idem_replacements = {
        27: '*Idem*, nota 20.',
        28: '*Idem*, nota 21.',
        35: '*Idem*, nota 2.',
        39: '*Idem*, nota 1.',
        40: '*Idem*, nota 12.',
        41: '*Idem*, nota 10.',
        42: '*Idem*, nota 11.',
    }
    for fn_num, idem_text in idem_replacements.items():
        pattern = rf'(\[\^{fn_num}\]:)\s*.+$'
        md_text = re.sub(pattern, rf'\1 {idem_text}', md_text, flags=re.MULTILINE)

    # 1c. Remove processing residues
    md_text = re.sub(r'(?m)^.*Fim do bloco.*$\n?', '', md_text)
    md_text = re.sub(r'(?m)^.*BLOCO \d+\s*[—–-].*$\n?', '', md_text)
    md_text = re.sub(r'/\s*English Translation', '', md_text)

    # 1d. Fix "clínica odontológica odontológico" → "consultório odontológico"
    md_text = md_text.replace('clínica odontológica odontológico', 'consultório odontológico')
    md_text = md_text.replace('clínicas odontológicas odontológicos', 'consultórios odontológicos')

    # 1e. Add page breaks before major sections (## headings)
    # Pandoc recognizes \newpage
    lines = md_text.split('\n')
    new_lines = []
    first_h2 = True
    for i, line in enumerate(lines):
        if re.match(r'^## \d', line):
            if first_h2:
                first_h2 = False
            else:
                new_lines.append('')
                new_lines.append('\\newpage')
                new_lines.append('')
        new_lines.append(line)
    md_text = '\n'.join(new_lines)

    # 1f. Fix YAML frontmatter for pandoc
    # Ensure proper title block
    if md_text.startswith('---'):
        # Already has YAML frontmatter - good
        pass

    return md_text


# ── Step 2: Pandoc Conversion ─────────────────────────────────────────
def run_pandoc(md_path: Path, output_path: Path) -> bool:
    """Convert fixed markdown to DOCX with real footnotes."""
    cmd = [
        'pandoc',
        str(md_path),
        '--from', 'markdown+footnotes+pipe_tables+yaml_metadata_block+implicit_figures',
        '--to', 'docx',
        '-o', str(output_path),
    ]

    # Use reference-doc if available
    if REF_DOC.exists():
        cmd.extend(['--reference-doc', str(REF_DOC)])
        print(f"  Using reference-doc: {REF_DOC.name}")
    else:
        print("  WARNING: no reference.docx found, using pandoc defaults")

    print(f"  Running: {' '.join(cmd[:6])}...")
    result = subprocess.run(cmd, capture_output=True, text=True)

    if result.returncode != 0:
        print(f"  PANDOC ERROR: {result.stderr}")
        return False

    if result.stderr:
        # Pandoc warnings (not errors)
        warnings = [l for l in result.stderr.strip().split('\n') if l.strip()]
        if warnings:
            print(f"  Pandoc warnings: {len(warnings)}")
            for w in warnings[:5]:
                print(f"    {w}")

    print(f"  Pandoc conversion OK → {output_path.name}")
    return True


# ── Step 3: Post-Process DOCX ─────────────────────────────────────────
def post_process(docx_path: Path):
    """Fix tables (WHITE), justify text, size images, clean up."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    print("  Post-processing DOCX...")
    doc = Document(str(docx_path))

    # ── 3a. JUSTIFY all paragraphs ──
    justified = 0
    for para in doc.paragraphs:
        style_name = para.style.name.lower() if para.style.name else ''
        # Skip TOC entries and headings
        if 'toc' in style_name or 'heading' in style_name or 'title' in style_name or 'subtitle' in style_name:
            continue
        # Skip footnote text (already small)
        if 'footnote' in style_name:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            justified += 1
            continue
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        justified += 1
    print(f"    Justified {justified} paragraphs")

    # ── 3b. Fix ALL tables: WHITE body, #DEDACB header ──
    header_color = 'DEDACB'
    tables_fixed = 0

    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)

                # Remove ALL existing shading
                for shd in tcPr.findall(qn('w:shd')):
                    tcPr.remove(shd)

                # Apply new shading
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')

                if i == 0:
                    # HEADER ROW: #DEDACB
                    shd.set(qn('w:fill'), header_color)
                    # Bold and color header text
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0x58, 0x4D, 0x42)
                else:
                    # DATA ROWS: WHITE (no zebra, no shading)
                    shd.set(qn('w:fill'), 'FFFFFF')

                tcPr.append(shd)

                # Set font in all cells
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Garamond'
                        run.font.size = Pt(9.5)

        # Add table borders
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Remove existing borders
        for borders in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(borders)

        # Add clean borders
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'B0A898')  # Light brown border
            borders.append(border)
        tblPr.append(borders)

        tables_fixed += 1

    print(f"    Fixed {tables_fixed} tables (WHITE body, #{header_color} header)")

    # ── 3c. Set Garamond font on body text ──
    fonts_set = 0
    for para in doc.paragraphs:
        style_name = para.style.name.lower() if para.style.name else ''
        if 'heading' in style_name or 'title' in style_name:
            for run in para.runs:
                run.font.name = 'Garamond'
                run.font.color.rgb = RGBColor(0x58, 0x4D, 0x42)
            continue
        if 'footnote' in style_name:
            for run in para.runs:
                run.font.name = 'Garamond'
                run.font.size = Pt(8)
            continue
        for run in para.runs:
            run.font.name = 'Garamond'
            if run.font.size is None:
                run.font.size = Pt(11)
            fonts_set += 1
    print(f"    Set Garamond on {fonts_set} runs")

    # ── 3d. Resize inline images ──
    images_resized = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            images_resized += 1

    # Resize images in paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            drawing = run._element.findall(qn('w:drawing'))
            for d in drawing:
                # Find the extent element and set width
                for extent in d.iter(qn('wp:extent')):
                    # Set width to ~15cm (5.9 inches) = 5400000 EMU
                    current_cx = int(extent.get('cx', 0))
                    current_cy = int(extent.get('cy', 0))
                    if current_cx > 0:
                        target_cx = 5400000  # ~15cm
                        ratio = target_cx / current_cx
                        extent.set('cx', str(target_cx))
                        extent.set('cy', str(int(current_cy * ratio)))

                # Also fix inline extent
                for inline in d.iter(qn('wp:inline')):
                    for extent in inline.iter(qn('wp:extent')):
                        current_cx = int(extent.get('cx', 0))
                        current_cy = int(extent.get('cy', 0))
                        if current_cx > 0 and current_cx != 5400000:
                            target_cx = 5400000
                            ratio = target_cx / current_cx
                            extent.set('cx', str(target_cx))
                            extent.set('cy', str(int(current_cy * ratio)))

    print(f"    Found {images_resized} embedded images")

    # ── 3e. Save ──
    doc.save(str(docx_path))
    print(f"    Saved: {docx_path.name}")


# ── Main ──────────────────────────────────────────────────────────────
def main():
    print("=" * 60)
    print("BP DentalShield — Pipeline V4 (Pandoc + Post-Processing)")
    print("=" * 60)

    # Read markdown
    print("\n[1/3] Fixing markdown...")
    md_text = MD_INPUT.read_text(encoding='utf-8')
    print(f"  Input: {len(md_text.split())} words")

    fixed = fix_markdown(md_text)

    # Count footnotes
    fn_defs = re.findall(r'^\[\^\d+\]:', fixed, re.MULTILINE)
    fn_refs = re.findall(r'\[\^\d+\](?!:)', fixed)
    idem_count = fixed.count('*Idem*')
    print(f"  Footnote definitions: {len(fn_defs)} ({idem_count} IDEM)")
    print(f"  Footnote references in body: {len(fn_refs)}")

    # Count charts
    chart_count = len(re.findall(r'!\[.*?\]\(.+?\.png\)', fixed))
    print(f"  Inline chart images: {chart_count}")

    # Write fixed markdown
    MD_FIXED.write_text(fixed, encoding='utf-8')
    print(f"  Written: {MD_FIXED.name}")

    # Run pandoc
    print("\n[2/3] Running Pandoc...")
    success = run_pandoc(MD_FIXED, OUTPUT)
    if not success:
        print("FATAL: Pandoc conversion failed!")
        sys.exit(1)

    size_kb = OUTPUT.stat().st_size / 1024
    print(f"  Output size: {size_kb:.0f} KB")

    # Post-process
    print("\n[3/3] Post-processing DOCX...")
    post_process(OUTPUT)

    final_size = OUTPUT.stat().st_size / 1024
    print(f"\n{'=' * 60}")
    print(f"DONE: {OUTPUT.name}")
    print(f"Size: {final_size:.0f} KB")
    print(f"Open in Word to verify footnotes, tables, charts.")
    print(f"{'=' * 60}")

    # Cleanup temp file
    if MD_FIXED.exists():
        MD_FIXED.unlink()
        print(f"  Cleaned up: {MD_FIXED.name}")


if __name__ == "__main__":
    main()
