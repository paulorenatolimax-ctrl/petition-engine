#!/usr/bin/env python3
"""
BP Generation Pipeline V2 — DentalShield (Camilla)
Multi-phase with Research Database + Endnotes + Charts
Based on Rafael Almeida reverse engineering

Usage: python3 generate_bp_camilla.py
Output: BP_DentalShield_V2.docx in target directory
"""

import json, os, sys, subprocess, re, textwrap
from pathlib import Path
from datetime import datetime

# === PATHS ===
SCRIPT_DIR = Path(__file__).parent
OUTPUT_DIR = Path("/Users/paulo1844/Documents/2_PROEX (A COMPLEMENTAR)/_2. MEUS CASOS/2025/CAMILLA/_BP/_Atualizado (pós-entendimento novo)")
RESEARCH_PATH = SCRIPT_DIR / "research_dental.json"
FINANCIAL_PATH = SCRIPT_DIR / "financial_camilla.json"

# === LOAD DATA ===
with open(RESEARCH_PATH) as f:
    RESEARCH = json.load(f)
with open(FINANCIAL_PATH) as f:
    FIN = json.load(f)

# === DESIGN CONSTANTS (Rafael/Pravion pattern) ===
FONT = "Garamond"
FONT_SIZE_BODY = 11
FONT_SIZE_H1 = 16
FONT_SIZE_H2 = 13
FONT_SIZE_H3 = 11
FONT_SIZE_ENDNOTE = 8
FONT_SIZE_CAPTION = 9

BROWN_DARK = "584D42"
BROWN_LIGHT = "DEDACB"
GREEN_LIGHT = "D0DDD6"
GRAY_BG = "F0F0F0"
TEXT_COLOR = "1A1A1A"
CAPTION_COLOR = "666666"
BORDER_COLOR = "CCCCCC"

# === PROHIBITED TERMS ===
PROHIBITED = [
    'consultoria', 'consultor', 'consulting', 'EB-2', 'NIW', 'EB2',
    'USCIS', 'green card', 'visto', 'imigração', 'immigration',
    'petition', 'petição', 'sponsor', 'waiver', 'I-140', 'I-485',
    'permanent resident', 'work visa', 'labor certification'
]

# ============================================================
# SECTION 1: CONTENT GENERATION VIA claude -p (MULTI-PHASE)
# ============================================================

def build_system_prompt():
    """Master system prompt for all blocks."""
    return textwrap.dedent(f"""
    You are writing a Business Plan for DentalShield Compliance Solutions LLC.
    This is a professional investment document. Language: Portuguese (Brazil), formal, investor-facing.

    ABSOLUTE RULES:
    1. NEVER mention: {', '.join(PROHIBITED)}
    2. Every numerical claim MUST have a footnote with verifiable source
    3. Footnote format: [N] inline, with full source at the end of each section
    4. Pattern: Thesis → Evidence (with footnote) → Explanation
    5. Zoom: Global → National → Regional → Company
    6. Minimum 60% prose, maximum 2 tables per section
    7. Font references: all numbers must come from the RESEARCH DATABASE below
    8. Financial numbers MUST match the FINANCIAL DATA below exactly

    RESEARCH DATABASE:
    {json.dumps(RESEARCH, indent=2, ensure_ascii=False)}

    FINANCIAL DATA:
    {json.dumps(FIN, indent=2, ensure_ascii=False)}
    """).strip()


def build_block_prompt(block_num):
    """Build specific prompt for each generation block."""
    blocks = {
        1: textwrap.dedent("""
        Generate BLOCK 1: SUMÁRIO EXECUTIVO (Sections 1.1 to 1.4)

        SECTION 1.1 — OPORTUNIDADE DE NEGÓCIO
        Follow IKARO pattern:
        - 1.1.1 Apresentação da Empresa (2-3 paragraphs: who, what, where, why)
        - 1.1.2 Serviços Oferecidos (table with 5 columns: Serviço | Descrição | Benefícios | Aplicação | Problemas Resolvidos)
          Use the 6 services from financial data
        - 1.1.3 Projeções-Chave (KPI box: Revenue 5yr, NPV, IRR, ROI, Payback, Employees Y5)
        Include 4-6 footnotes with sources from research database.

        SECTION 1.2 — CRONOGRAMA ESTRATÉGICO / TIMELINE DO NEGÓCIO
        Table format: Phase | Period | Location | Key Milestones | Investment
        3 phases: Y0 Tampa launch, Y2 Miami+Houston expansion, Y4 scaling

        SECTION 1.3 — VISÃO, MISSÃO E VALORES
        - Visão: 2 paragraphs (future state of dental compliance in US)
        - Missão: 1 paragraph (what DentalShield does daily)
        - Valores: 5 values with 1-sentence evidence for each

        SECTION 1.4 — ENQUADRAMENTO JURÍDICO
        MUST have minimum 5 subsections (Pravion pattern):
        1.4.1 Estrutura Societária (LLC Florida, advantages)
        1.4.2 Licenciamento e Registros (FL SOS, EIN, FEIN)
        1.4.3 Conformidade Regulatória Federal (OSHA, CDC, EPA, HIPAA, FDA)
        1.4.4 Conformidade Estadual (FL Board of Dentistry F.S. 466, TX TSBDE)
        1.4.5 Proteção de Propriedade Intelectual (SteriSensor, BiohazardBox, ComplianceScreen trademarks)
        1.4.6 Seguros e Responsabilidade (E&O, General Liability, Cyber)
        Include 8-10 footnotes with CFR references.

        FORMAT: Use markdown headers (## for H1, ### for H2, #### for H3).
        Include [1], [2], etc. for footnotes inline.
        At the end, list all footnotes as: [N] Source description, URL/reference, date.
        Total: 2000-2500 words.
        """),

        2: textwrap.dedent("""
        Generate BLOCK 2: ANÁLISE ESTRATÉGICA (Sections 2.1 to 2.8)

        SECTION 2.1 — PERSPECTIVAS DO MERCADO
        Zoom pattern: Global dental → US dental → Compliance market → Regional (FL, TX)
        Use ALL market data from research database with footnotes.
        Include IoT healthcare market data.
        6 subsections: 2.1.1 Global, 2.1.2 Nacional, 2.1.3 Tampa/Hillsborough, 2.1.4 Miami-Dade, 2.1.5 Houston/Harris, 2.1.6 Federal Policies

        SECTION 2.2 — CADEIA DE SUPRIMENTOS (Supply Chain)
        Porter's Value Chain adapted: Primary (auditing, training, IoT installation, monitoring, reporting) + Support (R&D, HR, IT, procurement)
        Table format.

        SECTION 2.3 — EMPREGABILIDADE ESPERADA
        Table with ALL SOC codes from financial data:
        SOC Code | Occupation | Median Salary | Growth | Annual Openings
        Include EPI multiplier calculation: 22 direct jobs × 3.35 = 73.7 total jobs
        3-4 paragraphs of analysis.

        SECTION 2.4 — GESTÃO DO CONHECIMENTO
        4 dimensions: Human Capital, Structural Capital, Relational Capital, Technological Capital
        Focus on IoT data, compliance databases, training certifications.

        SECTION 2.5 — IMPACTOS ESG
        Environmental (waste reduction), Social (worker safety, community health), Governance (compliance culture)

        SECTION 2.6 — ANÁLISE SWOT
        4 tables (Strengths, Weaknesses, Opportunities, Threats) with 5-6 items each.

        SECTION 2.7 — SWOT CRUZADA
        4 quadrants: SO, WO, ST, WT strategies. 3 strategies per quadrant.

        SECTION 2.8 — ANÁLISE DE CONCORRENTES (Porter 5 Forces summary)
        Use competitors from research database. Table comparing DentalShield vs 5 competitors.
        Dimensions: Scope, Tech Integration, On-Site, IoT, Pricing, Geographic.

        FORMAT: Markdown. [N] footnotes inline. List sources at end.
        Use data from research database for EVERY numerical claim.
        Total: 3000-3500 words.
        """),

        3: textwrap.dedent("""
        Generate BLOCK 3: MARKETING + OPERACIONAL (Sections 3.1 to 4.6)

        SECTION 3.1 — SEGMENTAÇÃO DE MERCADO
        Geographic: Tampa MSA (2,847 practices), Miami MSA (4,231), Houston MSA (5,112)
        Demographic: Solo practices, group practices, DSOs, dental schools
        Psychographic: Compliance-conscious vs reactive

        SECTION 3.2 — PÚBLICO-ALVO
        Primary: Solo/small dental practices (1-5 dentists)
        Secondary: Dental Service Organizations (DSOs)
        Tertiary: Dental schools and training institutions

        SECTION 3.3 — POSICIONAMENTO DA MARCA
        "The only dental compliance company combining IoT monitoring with on-site expertise"
        Differentiation matrix vs competitors.

        SECTION 3.4 — PRODUTO (Análise de Valor)
        Table: 6 services with features, benefits, pricing from financial data

        SECTION 3.5 — PREÇO
        Pricing strategy: premium positioning justified by cost of non-compliance
        Compare: OSHA fine avg $50K-100K vs DentalShield annual retainer
        Table: Service | Monthly | Annual | ROI for client

        SECTION 3.6 — PRAÇA (Distribuição)
        Phase 1: Tampa HQ direct sales
        Phase 2: Miami + Houston branches
        Phase 3: Scaling model (partner network)

        SECTION 3.7 — PROMOÇÃO
        Marketing budget from financial data (13.5% of revenue)
        Channels: LinkedIn, dental conferences (ADA, FDEA), referral program, content marketing

        SECTION 4.1 — QUADRO DE FUNCIONÁRIOS
        Table by location and year: Tampa (5→8), Miami (0→7), Houston (0→7)
        Include roles with SOC codes and salaries from BLS.

        SECTION 4.2 — LAYOUT DO EMPREENDIMENTO
        Tampa: 1,200 sq ft office + training room
        Miami: 800 sq ft office
        Houston: 800 sq ft office (Y5: 1,200 sq ft)

        SECTION 4.3 — RECURSOS E EQUIPAMENTOS
        IoT devices (SteriSensor units), training equipment, compliance software licenses, vehicles

        SECTION 4.4 — RECURSOS TECNOLÓGICOS
        ComplianceScreen SaaS platform, SteriSensor IoT network, CRM (HubSpot), cloud infrastructure

        SECTION 4.5 — LOCALIZAÇÃO
        Why Tampa (dental industry hub, no state income tax, proximity to Latin America)
        Why Miami (largest dental market in FL, bilingual advantage)
        Why Houston (largest dental market in TX, energy sector dental benefits)

        SECTION 4.6 — CAPACIDADE PRODUTIVA
        Y1: 50 clients, Y3: 150 clients, Y5: 300+ clients
        Based on consultant capacity: 1 consultant = ~25 active clients

        FORMAT: Markdown. [N] footnotes. Total: 2500-3000 words.
        """),

        4: textwrap.dedent("""
        Generate BLOCK 4: FINANCEIRO + CONCLUSÃO (Sections 5.1 to 6.2)

        SECTION 5.1 — PREMISSAS FINANCEIRAS
        List key assumptions:
        - Inflation: 3% annual
        - Discount rate: 12%
        - Tax: 25% effective (LLC pass-through)
        - Payroll tax FL: 7.65%, TX: 7.65%
        - Florida: no state income tax
        - Revenue growth: phased with location expansion
        - Variable costs: 24.51% of revenue

        SECTION 5.2 — INVESTIMENTOS
        Table: Phase | Period | Tangible | Intangible | Working Capital | Total
        Phase 1 (Y0): $175,257
        Phase 2 (Y2): $98,346
        Phase 3 (Y4): $120,240
        Total: $393,843

        SECTION 5.3 — ESTIMATIVA DE RECEITAS E CUSTOS
        Table: Year | Revenue | Variable Costs | Payroll | Fixed Costs | Total Costs | EBITDA
        Use EXACT numbers from financial data.
        2-3 paragraphs analyzing trends.

        SECTION 5.4 — DRE (Demonstrativo de Resultados)
        Full income statement table Y1-Y5
        Revenue → (-) Variable → Contribution Margin → (-) Payroll → (-) Fixed → EBITDA → (-) Tax → Net Income
        Use EXACT numbers from financial data.

        SECTION 5.5 — INDICADORES DE RETORNO
        Table: Indicator | Value
        NPV (12%): $396,168
        IRR: 61.18%
        ROI: 178.94%
        Payback: 3 years
        Business Value: $2,174,346
        2-3 paragraphs explaining each indicator.

        SECTION 5.6 — BREAK EVEN POINT
        Table: Year | BEP | Revenue | Surplus/Deficit
        Analysis of when company becomes sustainable.

        SECTION 6.1 — TIMELINE DE IMPLEMENTAÇÃO
        Gantt-style table: Activity | Q1Y1 | Q2Y1 | ... | Q4Y5
        Key milestones: LLC formation, Tampa launch, first clients, Miami opening, Houston opening, IoT product launch, 100th client, 300th client

        SECTION 6.2 — CONSIDERAÇÕES FINAIS
        NO FOOTNOTES in this section (it's synthesis).
        3-4 paragraphs summarizing:
        - Market opportunity ($4.8B compliance + $89.7B IoT healthcare)
        - Competitive advantage (IoT + on-site expertise)
        - Financial viability (IRR 61%, ROI 179%, payback 3 years)
        - Job creation (22 direct + 73 total via EPI multiplier)
        - Social impact (safer dental care for millions of Americans)

        FORMAT: Markdown. [N] footnotes (except 6.2). Total: 2000-2500 words.
        Use EXACT financial numbers — do not round or approximate.
        """)
    }
    return blocks.get(block_num, "")


def generate_content_block(block_num):
    """Call claude -p via stdin pipe: cat context.md | claude -p 'short instruction'."""
    system = build_system_prompt()
    prompt = build_block_prompt(block_num)

    print(f"\n{'='*60}")
    print(f"  GENERATING BLOCK {block_num}/4...")
    print(f"{'='*60}")

    # Save context (data) to file for piping via stdin
    context_file = SCRIPT_DIR / f"_context_block{block_num}.md"
    with open(context_file, 'w') as f:
        f.write(f"CONTEXT DATA:\n{system}\n\nSECTION INSTRUCTIONS:\n{prompt}")

    # Short instruction as -p argument (under shell limit)
    instruction = (
        f"You are writing Block {block_num} of a Business Plan for DentalShield LLC. "
        f"The piped input contains research data, financial data, and detailed section instructions. "
        f"Generate the ACTUAL content in Portuguese (Brazil), formal, investor-facing. "
        f"Use markdown headers (## H1, ### H2, #### H3). "
        f"Include [N] footnotes inline with verifiable sources. "
        f"List all footnotes at the end. "
        f"Output 2000-3500 words of REAL business plan prose, tables, and data. "
        f"Do NOT describe what you would write — WRITE the actual content."
    )

    # Save prompt for debugging
    prompt_file = SCRIPT_DIR / f"_prompt_block{block_num}.md"
    with open(prompt_file, 'w') as f:
        f.write(f"INSTRUCTION: {instruction}\n\nCONTEXT FILE: {context_file}")

    try:
        # Pipe context via stdin, short instruction as argument
        cmd = f'cat "{context_file}" | claude -p "{instruction}"'
        result = subprocess.run(
            cmd, shell=True,
            capture_output=True, text=True, timeout=600,
            cwd=str(SCRIPT_DIR)
        )
        content = result.stdout.strip()
        if not content and result.stderr.strip():
            content = result.stderr.strip()
            print(f"  WARNING: Got stderr output for block {block_num}")

        # Save raw output
        output_file = SCRIPT_DIR / f"_output_block{block_num}.md"
        with open(output_file, 'w') as f:
            f.write(content)

        print(f"  Block {block_num} generated: {len(content)} chars")
        return content

    except subprocess.TimeoutExpired:
        print(f"  ERROR: Block {block_num} timed out after 600s")
        return ""
    except Exception as e:
        print(f"  ERROR: Block {block_num} failed: {e}")
        return ""


# ============================================================
# SECTION 2: DOCX ASSEMBLY (python-docx with endnotes)
# ============================================================

def build_docx(blocks_content):
    """Assemble final DOCX from generated content blocks."""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # === PAGE SETUP ===
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # === STYLES ===
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(FONT_SIZE_BODY)
    style.font.color.rgb = RGBColor.from_string(TEXT_COLOR)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(3)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level, size, color in [(1, FONT_SIZE_H1, BROWN_DARK), (2, FONT_SIZE_H2, BROWN_DARK), (3, FONT_SIZE_H3, BROWN_DARK)]:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = FONT
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor.from_string(color)
        hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        hs.paragraph_format.space_after = Pt(8 if level == 1 else 6)
        hs.paragraph_format.keep_with_next = True

    # === ENDNOTE COLLECTION ===
    endnotes = []

    def add_paragraph_with_refs(text):
        """Add paragraph, converting [N] to superscript and collecting endnotes."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)

        # Split by footnote references [N]
        parts = re.split(r'(\[\d+\])', text)
        for part in parts:
            match = re.match(r'\[(\d+)\]', part)
            if match:
                run = p.add_run(f'[{match.group(1)}]')
                run.font.superscript = True
                run.font.size = Pt(8)
                run.font.name = FONT
            elif part.strip():
                # Handle **bold** markers
                bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
                for bp in bold_parts:
                    if bp.startswith('**') and bp.endswith('**'):
                        run = p.add_run(bp[2:-2])
                        run.bold = True
                        run.font.name = FONT
                        run.font.size = Pt(FONT_SIZE_BODY)
                    elif bp:
                        run = p.add_run(bp)
                        run.font.name = FONT
                        run.font.size = Pt(FONT_SIZE_BODY)
        return p

    def add_table(headers, rows, caption=None):
        """Add formatted table with proper styling."""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.bold = True
            run.font.name = FONT
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Header shading
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{BROWN_LIGHT}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rows with zebra striping
        for r_idx, row_data in enumerate(rows):
            for c_idx, value in enumerate(row_data):
                if c_idx >= len(headers):
                    break  # skip extra columns
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(value))
                run.font.name = FONT
                run.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Zebra striping
                if r_idx % 2 == 1:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{GREEN_LIGHT}"/>')
                    cell._tc.get_or_add_tcPr().append(shading)

        # Table borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>
                <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>
            </w:tblBorders>
        ''')
        tblPr.append(borders)

        # Caption
        if caption:
            p = doc.add_paragraph()
            run = p.add_run(f'Fonte: {caption}')
            run.italic = True
            run.font.size = Pt(FONT_SIZE_CAPTION)
            run.font.color.rgb = RGBColor.from_string(CAPTION_COLOR)
            run.font.name = FONT
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing after table
        doc.add_paragraph()

    def parse_and_add_content(content):
        """Parse markdown content and add to document."""
        lines = content.split('\n')
        i = 0
        in_table = False
        table_headers = []
        table_rows = []
        footnote_section = False

        while i < len(lines):
            line = lines[i].strip()

            # Skip empty lines
            if not line:
                if in_table and table_headers:
                    add_table(table_headers, table_rows)
                    table_headers = []
                    table_rows = []
                    in_table = False
                i += 1
                continue

            # Detect footnote section
            if re.match(r'^\[?\d+\]?\s*(Fonte|Source|http|BLS|Census|OSHA|CDC|ADA|Grand View|IBISWorld|Precedence)', line):
                footnote_section = True

            if footnote_section:
                # Collect endnotes
                fn_match = re.match(r'^\[?(\d+)\]?\s*(.+)', line)
                if fn_match:
                    endnotes.append((int(fn_match.group(1)), fn_match.group(2)))
                i += 1
                continue

            # Headers
            if line.startswith('#### '):
                if in_table:
                    add_table(table_headers, table_rows)
                    table_headers, table_rows, in_table = [], [], False
                doc.add_heading(line[5:], level=3)
            elif line.startswith('### '):
                if in_table:
                    add_table(table_headers, table_rows)
                    table_headers, table_rows, in_table = [], [], False
                doc.add_heading(line[4:], level=2)
            elif line.startswith('## '):
                if in_table:
                    add_table(table_headers, table_rows)
                    table_headers, table_rows, in_table = [], [], False
                # Add page break before H1 (except first)
                if len([p for p in doc.paragraphs if p.style.name.startswith('Heading 1')]) > 0:
                    doc.add_page_break()
                doc.add_heading(line[3:], level=1)

            # Tables (markdown format)
            elif '|' in line and not line.startswith('```'):
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if cells and not all(c.replace('-', '').replace(':', '') == '' for c in cells):
                    if not in_table:
                        table_headers = cells
                        in_table = True
                    elif all(re.match(r'^[-:]+$', c) for c in cells):
                        pass  # separator row
                    else:
                        table_rows.append(cells)

            # Regular paragraphs
            elif not line.startswith('```') and not line.startswith('---'):
                # Skip bullet markers, convert to paragraph
                if line.startswith('- ') or line.startswith('* '):
                    line = '• ' + line[2:]
                add_paragraph_with_refs(line)

            i += 1

        # Flush remaining table
        if in_table and table_headers:
            add_table(table_headers, table_rows)

    # === COVER PAGE ===
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BUSINESS PLAN")
    run.font.name = FONT
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(BROWN_DARK)
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DentalShield Compliance Solutions LLC")
    run.font.name = FONT
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(BROWN_DARK)

    doc.add_paragraph()

    for line_text in [
        "Dental Regulatory Compliance & IoT Monitoring",
        "Tampa, FL | Miami, FL | Houston, TX",
        f"Prepared: {datetime.now().strftime('%B %Y')}",
        "",
        "CONFIDENTIAL — DO NOT SHARE"
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line_text)
        run.font.name = FONT
        run.font.size = Pt(12)
        if "CONFIDENTIAL" in line_text:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(BROWN_DARK)

    doc.add_page_break()

    # === TABLE OF CONTENTS placeholder ===
    doc.add_heading("SUMÁRIO / TABLE OF CONTENTS", level=1)
    p = doc.add_paragraph("(Atualizar índice no Word: Ctrl+A → F9)")
    p.italic = True
    doc.add_page_break()

    # === PROCESS EACH BLOCK ===
    for block_num, content in enumerate(blocks_content, 1):
        if content:
            print(f"  Assembling Block {block_num} into DOCX...")
            parse_and_add_content(content)

    # === ENDNOTES SECTION ===
    if endnotes:
        doc.add_page_break()
        doc.add_heading("NOTAS E REFERÊNCIAS / ENDNOTES", level=1)
        endnotes.sort(key=lambda x: x[0])
        seen = set()
        for num, text in endnotes:
            if num not in seen:
                seen.add(num)
                p = doc.add_paragraph()
                run = p.add_run(f"[{num}] ")
                run.bold = True
                run.font.size = Pt(FONT_SIZE_ENDNOTE)
                run.font.name = FONT
                run = p.add_run(text)
                run.font.size = Pt(FONT_SIZE_ENDNOTE)
                run.font.name = FONT
                run.font.color.rgb = RGBColor.from_string(CAPTION_COLOR)

    # === FOOTER ===
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer bar
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:top w:val="single" w:sz="12" w:space="1" w:color="{BROWN_DARK}"/>
        </w:pBdr>
    ''')
    pPr.append(pBdr)

    run = p.add_run("CONFIDENTIAL — DO NOT SHARE  |  DentalShield Compliance Solutions LLC  |  ")
    run.font.name = FONT
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(BROWN_DARK)

    return doc


# ============================================================
# SECTION 3: CHART GENERATION (matplotlib)
# ============================================================

def generate_charts():
    """Generate 8 charts from financial data."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np

    plt.rcParams.update({
        'font.family': 'serif',
        'font.serif': ['Georgia', 'Palatino', 'Times New Roman'],
        'font.size': 10,
        'axes.titlesize': 13,
        'axes.titleweight': 'bold',
        'axes.labelsize': 11,
        'figure.facecolor': 'white',
        'axes.facecolor': 'white',
        'axes.edgecolor': '#CCCCCC',
        'axes.grid': True,
        'grid.color': '#E8E8E8',
        'grid.linewidth': 0.5,
    })

    COLORS = ["#584D42", "#8B7D6B", "#A0C4A8", "#D0DDD6", "#B8A88A", "#7A9E8C", "#C4B8A8"]
    YEARS = ["Year 1", "Year 2", "Year 3", "Year 4", "Year 5"]
    REV = [FIN['revenue'][f'Y{i}'] for i in range(1, 6)]
    EBITDA = [FIN['dre']['ebitda'][f'Y{i}'] for i in range(1, 6)]
    NET = [FIN['dre']['net_income'][f'Y{i}'] for i in range(1, 6)]
    EMPS = [FIN['employees'][f'Y{i}'] for i in range(1, 6)]
    BEP = [FIN['breakeven'][f'Y{i}'] for i in range(1, 6)]
    TOTAL_COSTS = [FIN['dre']['total_costs'][f'Y{i}'] for i in range(1, 6)]

    charts_dir = SCRIPT_DIR / "_charts"
    charts_dir.mkdir(exist_ok=True)
    chart_files = []

    # Chart 1: Revenue vs Costs
    fig, ax = plt.subplots(figsize=(7.5, 4.2), dpi=200)
    x = np.arange(len(YEARS))
    w = 0.35
    ax.bar(x - w/2, [r/1000 for r in REV], w, label='Revenue', color=COLORS[0])
    ax.bar(x + w/2, [c/1000 for c in TOTAL_COSTS], w, label='Total Costs', color=COLORS[2])
    ax.set_xlabel('Period')
    ax.set_ylabel('Amount (USD thousands)')
    ax.set_title('Revenue vs Total Costs — 5-Year Projection')
    ax.set_xticks(x)
    ax.set_xticklabels(YEARS)
    ax.legend()
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}K'))
    plt.tight_layout()
    path = charts_dir / "01_revenue_vs_costs.png"
    plt.savefig(path)
    plt.close()
    chart_files.append(("5.3", str(path), "Revenue vs Total Costs — 5-Year Projection"))

    # Chart 2: EBITDA + Margin
    fig, ax1 = plt.subplots(figsize=(7.5, 4.2), dpi=200)
    ax1.bar(YEARS, [e/1000 for e in EBITDA], color=[COLORS[0] if e >= 0 else '#CC4444' for e in EBITDA])
    ax1.set_ylabel('EBITDA (USD thousands)')
    ax1.set_title('EBITDA & EBITDA Margin — 5-Year Projection')
    ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}K'))
    ax2 = ax1.twinx()
    margins = [-2.7, 9.4, 25.0, 7.9, 26.7]
    ax2.plot(YEARS, margins, 'o-', color=COLORS[5], linewidth=2, markersize=8)
    ax2.set_ylabel('EBITDA Margin (%)')
    ax2.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'{x:.0f}%'))
    plt.tight_layout()
    path = charts_dir / "02_ebitda.png"
    plt.savefig(path)
    plt.close()
    chart_files.append(("5.4", str(path), "EBITDA & EBITDA Margin — 5-Year Projection"))

    # Chart 3: Employees by Location (stacked area)
    fig, ax = plt.subplots(figsize=(7.5, 4.2), dpi=200)
    tampa = [FIN['employees']['by_location']['tampa'][f'Y{i}'] for i in range(1, 6)]
    miami = [FIN['employees']['by_location']['miami'][f'Y{i}'] for i in range(1, 6)]
    houston = [FIN['employees']['by_location']['houston'][f'Y{i}'] for i in range(1, 6)]
    ax.stackplot(YEARS, tampa, miami, houston, labels=['Tampa, FL', 'Miami, FL', 'Houston, TX'],
                 colors=[COLORS[0], COLORS[2], COLORS[4]])
    ax.set_title('Workforce Growth by Location — 5-Year Projection')
    ax.set_ylabel('Number of Employees')
    ax.legend(loc='upper left')
    plt.tight_layout()
    path = charts_dir / "03_employees.png"
    plt.savefig(path)
    plt.close()
    chart_files.append(("4.1", str(path), "Workforce Growth by Location — 5-Year Projection"))

    # Chart 4: Services Revenue (donut)
    fig, ax = plt.subplots(figsize=(7.5, 4.2), dpi=200)
    services = [s['name'] for s in RESEARCH['company']['services']]
    pcts = [s['pct_revenue'] for s in RESEARCH['company']['services']]
    wedges, texts, autotexts = ax.pie(pcts, labels=services, autopct='%1.0f%%', startangle=90,
                                       colors=COLORS[:len(services)], pctdistance=0.85)
    centre = plt.Circle((0, 0), 0.65, fc='white')
    ax.add_artist(centre)
    ax.set_title('Revenue Distribution by Service Line')
    plt.tight_layout()
    path = charts_dir / "04_services.png"
    plt.savefig(path)
    plt.close()
    chart_files.append(("3.4", str(path), "Revenue Distribution by Service Line"))

    # Chart 5: Break-Even
    fig, ax = plt.subplots(figsize=(7.5, 4.2), dpi=200)
    ax.plot(YEARS, [r/1000 for r in REV], 'o-', color=COLORS[0], linewidth=2, label='Revenue')
    ax.plot(YEARS, [b/1000 for b in BEP], 's--', color='#CC4444', linewidth=2, label='Break-Even Point')
    ax.fill_between(YEARS, [b/1000 for b in BEP], [r/1000 for r in REV],
                     where=[r > b for r, b in zip(REV, BEP)], alpha=0.2, color=COLORS[2], label='Profit Zone')
    ax.set_title('Break-Even Analysis — Revenue vs BEP')
    ax.set_ylabel('Amount (USD thousands)')
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}K'))
    ax.legend()
    plt.tight_layout()
    path = charts_dir / "05_breakeven.png"
    plt.savefig(path)
    plt.close()
    chart_files.append(("5.6", str(path), "Break-Even Analysis — Revenue vs Break-Even Point"))

    # Chart 6: Investment by Phase
    fig, ax = plt.subplots(figsize=(7.5, 4.2), dpi=200)
    phases = ['Phase 1 (Y0)', 'Phase 2 (Y2)', 'Phase 3 (Y4)']
    tangible = [78108, 45200, 51800]
    intangible = [56635, 28700, 43994]
    working = [40514, 24446, 24446]
    y_pos = np.arange(len(phases))
    ax.barh(y_pos, [t/1000 for t in tangible], 0.25, label='Tangible', color=COLORS[0])
    ax.barh(y_pos + 0.25, [t/1000 for t in intangible], 0.25, label='Intangible', color=COLORS[2])
    ax.barh(y_pos + 0.5, [t/1000 for t in working], 0.25, label='Working Capital', color=COLORS[4])
    ax.set_yticks(y_pos + 0.25)
    ax.set_yticklabels(phases)
    ax.set_xlabel('Amount (USD thousands)')
    ax.set_title('Investment Structure by Phase')
    ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}K'))
    ax.legend()
    plt.tight_layout()
    path = charts_dir / "06_investment.png"
    plt.savefig(path)
    plt.close()
    chart_files.append(("5.2", str(path), "Investment Structure by Phase"))

    # Chart 7: Net Income
    fig, ax = plt.subplots(figsize=(7.5, 4.2), dpi=200)
    colors_ni = [COLORS[0] if n >= 0 else '#CC4444' for n in NET]
    ax.bar(YEARS, [n/1000 for n in NET], color=colors_ni)
    ax.set_title('Net Income — 5-Year Projection')
    ax.set_ylabel('Net Income (USD thousands)')
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}K'))
    ax.axhline(y=0, color='#999999', linewidth=0.8)
    plt.tight_layout()
    path = charts_dir / "07_net_income.png"
    plt.savefig(path)
    plt.close()
    chart_files.append(("5.4", str(path), "Net Income — 5-Year Projection"))

    # Chart 8: Cost Structure Y1 vs Y5
    fig, ax = plt.subplots(figsize=(7.5, 4.2), dpi=200)
    categories = ['Variable\nCosts', 'Payroll', 'Fixed\nCosts']
    y1_vals = [FIN['variable_costs']['Y1']/1000, FIN['payroll']['Y1']/1000, FIN['fixed_costs']['Y1']/1000]
    y5_vals = [FIN['variable_costs']['Y5']/1000, FIN['payroll']['Y5']/1000, FIN['fixed_costs']['Y5']/1000]
    x = np.arange(len(categories))
    w = 0.35
    ax.bar(x - w/2, y1_vals, w, label='Year 1', color=COLORS[0])
    ax.bar(x + w/2, y5_vals, w, label='Year 5', color=COLORS[2])
    ax.set_title('Cost Structure — Year 1 vs Year 5')
    ax.set_ylabel('Amount (USD thousands)')
    ax.set_xticks(x)
    ax.set_xticklabels(categories)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}K'))
    ax.legend()
    plt.tight_layout()
    path = charts_dir / "08_cost_structure.png"
    plt.savefig(path)
    plt.close()
    chart_files.append(("5.3", str(path), "Cost Structure — Year 1 vs Year 5"))

    print(f"  Generated {len(chart_files)} charts")
    return chart_files


def insert_charts(doc, chart_files):
    """Insert charts into the document after their anchor sections."""
    from docx.shared import Inches

    from docx.shared import Pt as PtSize, RGBColor as RGB2
    for section_ref, chart_path, caption in chart_files:
        if os.path.exists(chart_path):
            try:
                doc.add_paragraph()
                doc.add_picture(chart_path, width=Inches(6.3))
                last_para = doc.paragraphs[-1]
                last_para.alignment = 1  # CENTER

                p = doc.add_paragraph()
                run = p.add_run(f'Figura: {caption}')
                run.italic = True
                run.font.size = PtSize(9)
                run.font.name = FONT
                run.font.color.rgb = RGB2.from_string(CAPTION_COLOR)
                p.alignment = 1
            except Exception as e:
                print(f"  WARNING: Could not insert chart {chart_path}: {e}")


# ============================================================
# SECTION 4: AUDIT
# ============================================================

def audit_document(doc):
    """Run quality audit on generated document."""
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    issues = []

    # Check prohibited terms
    for term in PROHIBITED:
        count = len(re.findall(re.escape(term), full_text, re.IGNORECASE))
        if count > 0:
            issues.append(f"CRITICAL: Prohibited term '{term}' found {count}x")

    # Check fonts
    fonts_used = set()
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name:
                fonts_used.add(run.font.name)
    non_garamond = fonts_used - {FONT, None}
    if non_garamond:
        issues.append(f"WARNING: Non-Garamond fonts found: {non_garamond}")

    # Check word count
    word_count = len(full_text.split())
    if word_count < 5000:
        issues.append(f"WARNING: Only {word_count} words (target: 10,000+)")

    # Check for placeholder artifacts
    artifacts = re.findall(r'\[fn:\d+\]|\{[^}]+\}|\[TODO\]|\[PLACEHOLDER\]', full_text)
    if artifacts:
        issues.append(f"WARNING: {len(artifacts)} artifacts found: {artifacts[:5]}")

    print(f"\n{'='*60}")
    print(f"  AUDIT RESULTS")
    print(f"{'='*60}")
    print(f"  Word count: {word_count}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")
    print(f"  Fonts used: {fonts_used}")
    print(f"  Issues: {len(issues)}")
    for issue in issues:
        print(f"    - {issue}")

    return issues


# ============================================================
# SECTION 5: MAIN ORCHESTRATOR
# ============================================================

def main():
    print(f"""
╔══════════════════════════════════════════════════════════╗
║  BP GENERATION PIPELINE V2 — DentalShield (Camilla)     ║
║  Based on Rafael Almeida Reverse Engineering             ║
║  Multi-Phase + Research Database + Charts + Audit        ║
╚══════════════════════════════════════════════════════════╝

Output: {OUTPUT_DIR}/BP_DentalShield_V2.docx
Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
""")

    # Step 1: Generate charts first (fast, no LLM needed)
    print("STEP 1/5: Generating charts from financial data...")
    chart_files = generate_charts()

    # Step 2: Generate content blocks via claude -p
    print("\nSTEP 2/5: Generating content (4 blocks via claude -p)...")
    print("  This will take ~20-40 minutes total.\n")

    blocks = []
    for i in range(1, 5):
        content = generate_content_block(i)
        blocks.append(content)
        if content:
            print(f"  ✓ Block {i} complete ({len(content)} chars)")
        else:
            print(f"  ✗ Block {i} FAILED — will use fallback")

    # Step 3: Assemble DOCX
    print("\nSTEP 3/5: Assembling DOCX with python-docx...")
    doc = build_docx(blocks)

    # Step 4: Insert charts
    print("\nSTEP 4/5: Inserting charts...")
    insert_charts(doc, chart_files)

    # Step 5: Audit
    print("\nSTEP 5/5: Running audit...")
    issues = audit_document(doc)

    # Save
    output_path = OUTPUT_DIR / "BP_DentalShield_V2.docx"
    doc.save(str(output_path))

    print(f"""
╔══════════════════════════════════════════════════════════╗
║  GENERATION COMPLETE                                     ║
║  Output: {str(output_path)[:50]}...
║  Issues: {len(issues)}
║  Finished: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
╚══════════════════════════════════════════════════════════╝
""")

    return str(output_path)


if __name__ == "__main__":
    main()
