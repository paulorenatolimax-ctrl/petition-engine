#!/usr/bin/env python3
"""
Petition Engine — DOCX Inserter
Inserts thumbnails into DOCX documents with formatted captions.
"""

import sys
import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_photographic_report(thumbnails: list, output_path: str, client_name: str, visa_type: str, title: str = "PHOTOGRAPHIC EVIDENCE REPORT") -> dict:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Petitioner: {client_name}\nVisa Category: {visa_type}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    inserted = 0
    errors = []

    for i, thumb in enumerate(thumbnails):
        img_path = thumb.get('path', '')
        caption = thumb.get('caption', f'Evidence {i + 1}')
        exhibit = thumb.get('exhibit_number', f'Exhibit {i + 1}')

        if not os.path.exists(img_path):
            errors.append(f"Imagem não encontrada: {img_path}")
            continue

        exhibit_para = doc.add_paragraph()
        exhibit_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = exhibit_para.add_run(exhibit)
        run.bold = True
        run.font.size = Pt(12)

        try:
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            inserted += 1
        except Exception as e:
            errors.append(f"Erro ao inserir {img_path}: {str(e)}")
            continue

        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_para.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

        if i < len(thumbnails) - 1:
            doc.add_paragraph()

    doc.save(output_path)

    return {
        "status": "ok",
        "output": output_path,
        "total_evidence": len(thumbnails),
        "inserted": inserted,
        "errors": errors,
    }


def insert_thumbnails_into_existing(docx_path: str, thumbnails: list, output_path: str, section_title: str = "EVIDENCE APPENDIX") -> dict:
    doc = Document(docx_path)
    doc.add_page_break()

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(section_title)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    inserted = 0
    errors = []

    for i, thumb in enumerate(thumbnails):
        img_path = thumb.get('path', '')
        caption = thumb.get('caption', f'Evidence {i + 1}')
        exhibit = thumb.get('exhibit_number', f'Exhibit {i + 1}')

        if not os.path.exists(img_path):
            errors.append(f"Não encontrado: {img_path}")
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(exhibit)
        run.bold = True
        run.font.size = Pt(11)

        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5))
            inserted += 1
        except Exception as e:
            errors.append(f"Erro: {img_path} — {str(e)}")
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)

        doc.add_paragraph()

    doc.save(output_path)

    return {
        "status": "ok",
        "output": output_path,
        "inserted": inserted,
        "errors": errors,
    }


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(json.dumps({
            "error": "Uso: docx_inserter.py <mode> [args...]",
            "modes": {
                "report": "docx_inserter.py report <thumbnails_json> <output.docx> <client_name> <visa_type>",
                "insert": "docx_inserter.py insert <existing.docx> <thumbnails_json> <output.docx>",
            }
        }))
        sys.exit(1)

    mode = sys.argv[1]

    if mode == 'report':
        thumbnails = json.loads(sys.argv[2])
        output = sys.argv[3]
        client_name = sys.argv[4]
        visa_type = sys.argv[5]
        result = create_photographic_report(thumbnails, output, client_name, visa_type)
        print(json.dumps(result, indent=2))
    elif mode == 'insert':
        docx_path = sys.argv[2]
        thumbnails = json.loads(sys.argv[3])
        output = sys.argv[4]
        result = insert_thumbnails_into_existing(docx_path, thumbnails, output)
        print(json.dumps(result, indent=2))
    else:
        print(json.dumps({"error": f"Modo desconhecido: {mode}"}))
        sys.exit(1)
