#!/usr/bin/env python3
"""Validador determinístico bloqueante da Resposta RFE EB-2 NIW.

Uso:
    python3 validate_rfe_response.py <rfe_response.docx> <officer_quotes.json> \\
            <research_sources.json> <cover_letter_last_exhibit_num> [report.md]

Exit 0 se OK, exit 1 se errors. Gera report markdown.

Checa (inspirado em impacto v3.0 + anteprojeto v2.0):
- Zero termos afrontosos (lista expandida)
- Zero terminologia EB-1A (Kazarian, extraordinary ability, top of field)
- Zero exposição de infra (PROEX internal refs, Kortix, Petition Engine, Obsidian)
- Numeração Exhibit continua da CL (primeira nova evidência = last_exhibit + 1)
- ≥10 sources inline [Source: url, accessed date] OU footnotes com URLs
- ≥5 domínios autoritativos (uscis.gov, bls.gov, census.gov, whitehouse.gov, bea.gov, etc)
- ≥1 Matter of Dhanasar reference
- ≥1 USCIS Policy Manual PA-2025-03 OR equivalent reference
- ≥N officer quotes verbatim presentes no corpo (N = len(officer_quotes.json))
- Thumbnails presentes (inline_shapes ≥ estimated new evidence count)
- Paras + tables + images em ranges plausíveis vs Deni benchmark
"""
import json
import os
import re
import sys

try:
    from docx import Document
except ImportError:
    sys.stderr.write("ERR: python-docx not installed\n")
    sys.exit(2)

# ─── Lists ───────────────────────────────────────────────────────────────────
FORBIDDEN_CONFRONTATIONAL = [
    r"\bThe officer erred\b", r"\bThe officer was (wrong|incorrect|mistaken)\b",
    r"\bo oficial (errou|equivocou|est[aá] equivocado|est[aá] incorreto)\b",
    r"\boficial\s+equivocad[oa]\b",
    r"\b(foi |é )?refutad[oa]\b", r"\bREFUTA(ÇÃO|DO|MOS)\b",
    r"\bcompletely contradic(ts|ted)\b", r"\boficial\s+errado\b",
    r"\bthe adjudicator\s+(erred|failed to understand|overlooked the obvious)\b",
    r"\bthe record (clearly )?refutes\b",
    r"\bconfrontativo\b", r"\bonfronto a decisão\b",
]

FORBIDDEN_EB1A = [
    r"\bKazarian\b", r"\bextraordinary ability\b", r"\bextraordinary-ability\b",
    r"\btop of (the|my) field\b", r"\bcomparable evidence\b"
]

FORBIDDEN_INFRA_INTERNAL = [
    r"\bKortix\b", r"\bPetition Engine\b", r"\bForjado por\b",
    r"\bObsidian\b", r"\bRAG\s+(I{1,3}|\d)\b",
    r"\bCowork(?!\s+space)\b",
]

REQUIRED_LEGAL_FRAMEWORK = [
    r"Matter of Dhanasar",
    r"(USCIS Policy Manual|Volume\s+6.*Part\s+F.*Chapter\s+5|PA-2025-03|8\s*CFR\s*204\.5)"
]

AUTHORITATIVE_DOMAINS = [
    "uscis.gov", "bls.gov", "bea.gov", "census.gov", "dol.gov",
    "whitehouse.gov", "commerce.gov", "eda.gov", "sba.gov",
    "federalregister.gov", "ustr.gov", "state.gov",
    "apprenticeship.gov", "e-verify.gov", "osha.gov",
    "cbo.gov", "gao.gov", "doc.gov"
]

SOURCE_MARKER = re.compile(
    r"\[Source:\s*(https?://[^\s,\]]+)[^,\]]*,?\s*accessed\s+(\d{4}-\d{2}-\d{2})\s*\]",
    re.IGNORECASE
)

EXHIBIT_REF = re.compile(r"\b(?:Exhibit|Evidence|Anexo|Evidência|Attachment)\s+(\d+)\b", re.IGNORECASE)

# Prong 3 must always be addressed (Barbara syllogism — conclusion of P1+P2)
# Even if officer did not explicitly object to P3, response must close the logical chain
PRONG3_MARKERS = [
    r"\bProng\s+3\b", r"\bon\s+balance\b",
    r"\bbeneficial\s+to\s+(?:the\s+United\s+States\s+to\s+)?waive\b",
    r"\blabor\s+certification\b", r"\bconventional\s+PERM\b"
]


# ─── Helpers ─────────────────────────────────────────────────────────────────
def _err(lst, msg):
    lst.append(msg)


def extract_all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        for container in (section.header, section.footer):
            if container:
                for p in container.paragraphs:
                    parts.append(p.text)
    return "\n".join(p for p in parts if p)


# ─── Checks ──────────────────────────────────────────────────────────────────
def check_forbidden(text, errors, passed):
    found_confront = [p for p in FORBIDDEN_CONFRONTATIONAL if re.search(p, text, re.I)]
    if found_confront:
        for p in found_confront:
            m = re.search(p, text, re.I)
            ctx = text[max(0, m.start() - 60):m.end() + 60]
            _err(errors, f"confrontational term: /{p}/ near '{ctx[:120]}'")
    else:
        _err(passed, "zero confrontational terminology")

    found_eb1a = [p for p in FORBIDDEN_EB1A if re.search(p, text, re.I)]
    if found_eb1a:
        for p in found_eb1a:
            _err(errors, f"EB-1A terminology leak: {p}")
    else:
        _err(passed, "zero EB-1A terminology")

    found_infra = [p for p in FORBIDDEN_INFRA_INTERNAL if re.search(p, text)]
    if found_infra:
        for p in found_infra:
            _err(errors, f"internal infrastructure exposure: {p}")
    else:
        _err(passed, "zero internal infrastructure exposure")


def check_legal_framework(text, errors, passed):
    missing = []
    for pattern in REQUIRED_LEGAL_FRAMEWORK:
        if not re.search(pattern, text, re.IGNORECASE):
            missing.append(pattern)
    if missing:
        for m in missing:
            _err(errors, f"missing required legal framework reference: {m}")
    else:
        _err(passed, "Matter of Dhanasar + USCIS Policy Manual / 8 CFR 204.5 referenced")


def check_sources(text, errors, warnings, passed):
    sources = SOURCE_MARKER.findall(text)
    footnote_urls = re.findall(r"https?://[^\s\)]+", text)
    total_source_refs = len(sources) + len(footnote_urls)

    if len(sources) + len(footnote_urls) < 10:
        _err(errors, f"only {len(sources)} inline [Source:] markers + {len(footnote_urls)} URL mentions = {total_source_refs}; minimum 10 required")
    else:
        _err(passed, f"{len(sources)} inline markers + {len(footnote_urls)} URL mentions = {total_source_refs} total refs")

    domains_cited = set()
    for url in (s[0] for s in sources):
        for d in AUTHORITATIVE_DOMAINS:
            if d in url:
                domains_cited.add(d)
                break
    for url in footnote_urls:
        for d in AUTHORITATIVE_DOMAINS:
            if d in url:
                domains_cited.add(d)
                break

    if len(domains_cited) < 5:
        _err(errors, f"only {len(domains_cited)} authoritative domains cited; minimum 5")
    else:
        _err(passed, f"{len(domains_cited)} authoritative domains: {', '.join(sorted(domains_cited))}")


def check_exhibit_numbering(text, cl_last_exhibit, errors, warnings, passed):
    """New exhibits should start at cl_last_exhibit + 1, continuing sequence."""
    exhibits = sorted(set(int(m.group(1)) for m in EXHIBIT_REF.finditer(text)))
    if not exhibits:
        _err(warnings, "no Exhibit N references found in body — verify evidence integration")
        return
    min_exhibit = min(exhibits)
    max_exhibit = max(exhibits)
    expected_start = cl_last_exhibit + 1

    # All exhibits should be >= cl_last_exhibit - 5 (allows some back-refs to CL old exhibits)
    # New exhibits should start at expected_start or higher
    new_exhibits = [e for e in exhibits if e >= expected_start]
    if not new_exhibits:
        _err(errors, f"no new exhibits found >= Exhibit {expected_start} (CL ended at {cl_last_exhibit}); new evidence numbering broken")
    elif new_exhibits[0] != expected_start:
        _err(warnings, f"first new exhibit = {new_exhibits[0]}; expected {expected_start} (CL last was {cl_last_exhibit})")
    else:
        _err(passed, f"exhibit numbering continuous: CL ended Exhibit {cl_last_exhibit}; new evidence starts at Exhibit {expected_start}")


def check_prong3_present(text, errors, warnings, passed):
    """Silogismo Bárbara: P1+P2 válidos → P3 válido. Mesmo sem objeção explícita do
    oficial, a resposta DEVE fechar o silogismo com seção/argumentação de P3."""
    hits = sum(1 for p in PRONG3_MARKERS if re.search(p, text, re.IGNORECASE))
    if hits < 2:
        _err(errors,
             f"Prong 3 absent or insufficiently addressed ({hits}/5 markers found). "
             "Per silogismo Bárbara (Aristóteles), P3 is the conclusion of P1+P2 and must always close the logical chain, "
             "even if the officer did not explicitly object to it.")
    else:
        _err(passed, f"Prong 3 addressed ({hits}/5 syllogism markers present)")


def check_officer_quotes_embedded(text, quotes_json_path, errors, warnings, passed):
    if not os.path.isfile(quotes_json_path):
        _err(warnings, f"officer_quotes.json not found at {quotes_json_path}; skipping quote-embedding check")
        return
    try:
        data = json.load(open(quotes_json_path))
        quotes = data.get("quotes", [])
    except Exception as e:
        _err(warnings, f"officer_quotes.json unreadable: {e}")
        return

    if not quotes:
        _err(warnings, "officer_quotes.json has 0 quotes")
        return

    found = 0
    missing = []
    for q in quotes:
        verbatim = q.get("verbatim", "")
        # Match first 50 chars to allow for light formatting differences
        probe = re.escape(verbatim[:50].strip())
        if probe and re.search(probe, text):
            found += 1
        else:
            missing.append(q.get("id", "?"))

    total = len(quotes)
    if found < total * 0.5:
        _err(errors, f"only {found}/{total} officer quotes embedded in response (<50%); missing IDs: {missing[:5]}")
    elif found < total:
        _err(warnings, f"{found}/{total} officer quotes embedded; missing: {missing[:5]}")
    else:
        _err(passed, f"all {total} officer quotes embedded verbatim")


def check_metrics(doc, file_size, errors, warnings, passed):
    paras = len(doc.paragraphs)
    tables = len(doc.tables)
    images = len(doc.inline_shapes)
    text = extract_all_text(doc)
    words = len(text.split())

    # Benchmarks: Deni Rubens (EB-2 NIW) 811p/19t/53i/27852w/92pg — Marcelo Góes (EB-1A) 90t/69i/39008w/327pg
    # Paulo's guidance: RFE response must be IGUAL OU MELHOR than Deni (same visa type)
    BENCH = {"paragraphs": 811, "tables": 19, "images": 53, "words": 27852, "bytes": 3_900_000}

    if paras < BENCH["paragraphs"] * 0.5:  # below 50% = block
        _err(errors, f"paragraphs={paras} < {int(BENCH['paragraphs']*0.5)} (50% of Deni benchmark {BENCH['paragraphs']}); response below quality floor")
    elif paras < BENCH["paragraphs"] * 0.75:
        _err(warnings, f"paragraphs={paras} < {int(BENCH['paragraphs']*0.75)} (75% of Deni benchmark) — expand further to match benchmark")
    elif paras < BENCH["paragraphs"]:
        _err(warnings, f"paragraphs={paras} below Deni benchmark {BENCH['paragraphs']}")
    else:
        _err(passed, f"paragraphs={paras} meets/exceeds Deni benchmark ({BENCH['paragraphs']})")

    if tables < BENCH["tables"] * 0.5:
        _err(errors, f"tables={tables} < {int(BENCH['tables']*0.5)} (50% of benchmark); missing structural tables")
    elif tables < BENCH["tables"]:
        _err(warnings, f"tables={tables} below Deni benchmark {BENCH['tables']}")
    else:
        _err(passed, f"tables={tables} meets Deni benchmark")

    if images < BENCH["images"] * 0.3:
        _err(errors, f"inline_images={images} < {int(BENCH['images']*0.3)} (30% of benchmark); thumbnails critically missing")
    elif images < BENCH["images"]:
        _err(warnings, f"inline_images={images} below Deni benchmark {BENCH['images']} — thumbnails to add")
    else:
        _err(passed, f"inline_images={images} meets Deni benchmark")

    if words < BENCH["words"] * 0.5:
        _err(errors, f"words={words} < {int(BENCH['words']*0.5)} (50% of Deni benchmark {BENCH['words']}); response not substantive enough")
    elif words < BENCH["words"]:
        _err(warnings, f"words={words} below Deni benchmark {BENCH['words']} — expand argumentation")
    else:
        _err(passed, f"words={words} meets/exceeds Deni benchmark")

    if file_size < 500_000:
        _err(warnings, f"file_size={file_size:,} bytes < 500KB (benchmark ~3.9MB); thumbnails will increase this")

    return {"paragraphs": paras, "tables": tables, "images": images, "words": words, "bytes": file_size}


# ─── Main ────────────────────────────────────────────────────────────────────
def main():
    if len(sys.argv) < 5:
        sys.stderr.write(
            "Usage: validate_rfe_response.py <rfe.docx> <officer_quotes.json> "
            "<research_sources.json> <cl_last_exhibit_num> [report.md]\n"
        )
        sys.exit(2)

    docx_path = sys.argv[1]
    quotes_path = sys.argv[2]
    sources_path = sys.argv[3]
    try:
        cl_last = int(sys.argv[4])
    except ValueError:
        sys.stderr.write(f"ERR: cl_last_exhibit_num must be integer, got {sys.argv[4]!r}\n")
        sys.exit(2)
    report_path = sys.argv[5] if len(sys.argv) > 5 else None

    if not os.path.isfile(docx_path):
        sys.stderr.write(f"ERR: DOCX not found: {docx_path}\n")
        sys.exit(2)

    file_size = os.path.getsize(docx_path)
    try:
        doc = Document(docx_path)
    except Exception as e:
        sys.stderr.write(f"ERR: cannot open DOCX: {e}\n")
        sys.exit(2)

    text = extract_all_text(doc)
    errors, warnings, passed = [], [], []

    check_forbidden(text, errors, passed)
    check_legal_framework(text, errors, passed)
    check_sources(text, errors, warnings, passed)
    check_exhibit_numbering(text, cl_last, errors, warnings, passed)
    check_prong3_present(text, errors, warnings, passed)
    check_officer_quotes_embedded(text, quotes_path, errors, warnings, passed)
    metrics = check_metrics(doc, file_size, errors, warnings, passed)

    status = "PASSED" if not errors else "FAILED"
    lines = [
        f"# RFE Response Validator — {status}",
        "",
        f"**File:** `{docx_path}`",
        "",
        "## Metrics",
        ""
    ]
    for k, v in metrics.items():
        lines.append(f"- {k}: {v:,}" if isinstance(v, int) else f"- {k}: {v}")
    lines.append("")
    lines.append(f"## Checks — {len(passed)} passed · {len(warnings)} warnings · {len(errors)} errors")
    lines.append("")
    if passed:
        lines.append("### ✓ Passed")
        for p in passed:
            lines.append(f"- {p}")
        lines.append("")
    if warnings:
        lines.append("### ⚠ Warnings (non-blocking)")
        for w in warnings:
            lines.append(f"- {w}")
        lines.append("")
    if errors:
        lines.append("### ✗ Errors (BLOCKING)")
        for e in errors:
            lines.append(f"- {e}")
        lines.append("")

    report = "\n".join(lines)
    print(report)
    if report_path:
        os.makedirs(os.path.dirname(os.path.abspath(report_path)) or ".", exist_ok=True)
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(report + "\n")

    sys.exit(1 if errors else 0)


if __name__ == "__main__":
    main()
