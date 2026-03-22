#!/usr/bin/env python3
"""Extract text from DOCX files."""
import sys
try:
    from docx import Document
except ImportError:
    print("[ERRO: python-docx não instalado. Rode: pip3 install python-docx]")
    sys.exit(1)

def extract(path):
    try:
        doc = Document(path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text.append(' | '.join(cells))
        return '\n'.join(text)
    except Exception as e:
        return f"[ERRO: {str(e)}]"

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Uso: python3 extract_docx.py <caminho.docx>")
        sys.exit(1)
    print(extract(sys.argv[1]))
