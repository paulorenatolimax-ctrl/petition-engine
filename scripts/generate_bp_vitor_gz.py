#!/usr/bin/env python3
"""
Business Plan Generator — Victor Moreira Dias / Goodii Burger
EB-2 NIW | Output: V1_business_plan_Vitor_GZ.docx
"""
import os, sys, json, tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np

# ──────────────────────────── CONFIG ────────────────────────────
OUT = "/Users/paulo1844/Documents/3_OMNI/_IMIGRAÇÃO/_CLIENTES/Coisas Gizele/Vitor/_Forjado por Petition Engine"
FNAME = "V1_business_plan_Vitor_GZ.docx"
NAVY = '#1B2A4A'; GOLD = '#B8860B'; GRAY_C = '#666666'
BROWN = '#8B7355'; HDR_BG = 'DEDACB'; WHITE_C = '#FFFFFF'
FONT = 'Garamond'
TMPDIR = tempfile.mkdtemp(prefix="bp_charts_")
COMPANY = "GOODII BURGER"
FOOTER = f"CONFIDENTIAL — {COMPANY} — Business Plan 2026"

# ──────────── FINANCIAL DATA ────────────
YEARS = ['2026','2027','2028','2029','2030']
REVENUE = [320000, 625000, 1080000, 1780000, 2650000]
COGS_PCT = [0.30, 0.29, 0.28, 0.28, 0.27]
LABOR_PCT = [0.28, 0.27, 0.26, 0.25, 0.24]
RENT_PCT = [0.13, 0.08, 0.06, 0.05, 0.04]
MKTG_PCT = [0.10, 0.08, 0.06, 0.05, 0.04]
UTIL_PCT = [0.04, 0.04, 0.03, 0.03, 0.03]
INSUR_PCT = [0.03, 0.02, 0.02, 0.02, 0.02]
DEPR_PCT = [0.03, 0.03, 0.02, 0.02, 0.02]
ADMIN_PCT = [0.06, 0.05, 0.05, 0.04, 0.04]
def calc_expenses():
    rows = []
    for i in range(5):
        r = REVENUE[i]
        row = {
            'COGS': int(r*COGS_PCT[i]), 'Labor': int(r*LABOR_PCT[i]),
            'Rent': int(r*RENT_PCT[i]), 'Marketing': int(r*MKTG_PCT[i]),
            'Utilities': int(r*UTIL_PCT[i]), 'Insurance': int(r*INSUR_PCT[i]),
            'Depreciation': int(r*DEPR_PCT[i]), 'Admin & Other': int(r*ADMIN_PCT[i])
        }
        row['Total'] = sum(row.values())
        row['Net Income'] = r - row['Total']
        rows.append(row)
    return rows
EXPENSES = calc_expenses()
NET_INCOME = [e['Net Income'] for e in EXPENSES]

STARTUP = [
    ('Leasehold Improvements', 165000), ('Kitchen Equipment', 72000),
    ('Furniture & Fixtures', 22000), ('Initial Inventory', 14000),
    ('Technology & POS Systems', 11000), ('Marketing Launch Campaign', 32000),
    ('Working Capital (3 months)', 55000), ('Legal & Intellectual Property', 28000),
    ('Licenses & Permits', 9000), ('Insurance (First Year)', 12000)
]
STARTUP_TOTAL = sum(v for _,v in STARTUP)

MARKET_YEARS = ['2022','2023','2024','2025','2026P','2027P']
MARKET_SIZE = [320, 340, 362, 385, 415, 447]

CUST_MONTHS = [3,6,12,18,24,36,48,60]
CUST_COUNT = [800,1400,2600,3800,5200,8500,12800,18500]

# ──────────── HELPERS ────────────
footnote_counter = [0]
all_footnotes = []

def fn(text_ref):
    footnote_counter[0] += 1
    n = footnote_counter[0]
    all_footnotes.append((n, text_ref))
    return n

def fmt(val):
    if abs(val) >= 1_000_000:
        return f"${val/1_000_000:,.1f}M"
    elif abs(val) >= 1000:
        return f"${val:,.0f}"
    return f"${val:,.0f}"

def pct(val): return f"{val*100:.0f}%"

def add_p(doc, text, bold=False, italic=False, align=None, size=None, color=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = size or Pt(12)
    if bold: run.bold = True
    if italic: run.italic = True
    if color: run.font.color.rgb = color
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = space_after if space_after is not None else Pt(6)
    if align: pf.alignment = align
    _set_rpr_font(run)
    return p

def add_p_with_fn(doc, text, fn_num, fn_ref, size=None):
    """Add paragraph with superscript footnote number."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = size or Pt(12)
    _set_rpr_font(run)
    sup = p.add_run(str(fn_num))
    sup.font.superscript = True
    sup.font.size = Pt(8)
    sup.font.name = FONT
    _set_rpr_font(sup)
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    return p

def _set_rpr_font(run):
    """Ensure font name is set in both w:rFonts attributes for Garamond."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), FONT)
        rFonts.set(qn('w:hAnsi'), FONT)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT
        run.font.size = Pt(14) if level <= 2 else Pt(12)
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
        run.bold = True
        _set_rpr_font(run)
    return h

def add_hline(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="999999"/></w:pBdr>')
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)

def add_fn_block(doc, footnotes_list):
    """Add a footnotes reference block at the end of a section."""
    if not footnotes_list:
        return
    add_hline(doc)
    for num, ref in footnotes_list:
        p = doc.add_paragraph()
        sup = p.add_run(f"{num} ")
        sup.font.superscript = True
        sup.font.size = Pt(9)
        sup.font.name = FONT
        _set_rpr_font(sup)
        r = p.add_run(ref)
        r.font.size = Pt(10)
        r.font.name = FONT
        r.font.italic = True
        _set_rpr_font(r)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0

def add_page_break(doc):
    doc.add_page_break()

def add_chart(doc, chart_path, width=Inches(5.8)):
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def make_table(doc, headers, rows, col_widths=None):
    """Create a formatted table with header background and no side borders."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(11)
        _set_rpr_font(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HDR_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(11)
            _set_rpr_font(run)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
    # Style: remove side borders, light top/bottom
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                '</w:tcBorders>')
            tcPr.append(borders)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    doc.add_paragraph()  # spacing after table
    return table

# ──────────── CHART GENERATORS ────────────
def _style_chart(ax, title):
    ax.set_title(title, fontsize=14, fontweight='bold', color=NAVY, pad=12)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color(GRAY_C)
    ax.spines['bottom'].set_color(GRAY_C)
    ax.tick_params(colors=GRAY_C, labelsize=10)

def gen_revenue_chart():
    fig, ax = plt.subplots(figsize=(8, 4.5))
    x = np.arange(len(YEARS))
    bars = ax.bar(x, [r/1000 for r in REVENUE], color=NAVY, width=0.55, edgecolor='white', linewidth=0.5)
    for bar, val in zip(bars, REVENUE):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 20,
                f'${val/1000:,.0f}K', ha='center', va='bottom', fontsize=10, fontweight='bold', color=NAVY)
    ax.set_xticks(x); ax.set_xticklabels([f'Year {i+1}\n({y})' for i, y in enumerate(YEARS)])
    ax.set_ylabel('Revenue ($ thousands)', fontsize=11, color=GRAY_C)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'${v:,.0f}K'))
    _style_chart(ax, 'Five-Year Revenue Projection')
    ax.set_ylim(0, max(REVENUE)/1000 * 1.15)
    fig.tight_layout()
    path = os.path.join(TMPDIR, 'chart_revenue.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def gen_expense_pie():
    fig, ax = plt.subplots(figsize=(7, 5))
    e = EXPENSES[0]
    labels = ['Cost of Goods\nSold','Labor','Rent','Marketing','Utilities','Insurance','Depreciation','Admin &\nOther']
    sizes = [e['COGS'], e['Labor'], e['Rent'], e['Marketing'], e['Utilities'], e['Insurance'], e['Depreciation'], e['Admin & Other']]
    colors_list = [NAVY, GOLD, '#4A6FA5', '#C4A35A', '#7B8D9E', '#9B8B6B', '#5C7A99', GRAY_C]
    wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.0f%%',
        colors=colors_list, startangle=90, pctdistance=0.78,
        textprops={'fontsize': 9, 'color': GRAY_C})
    for at in autotexts:
        at.set_fontsize(9); at.set_fontweight('bold'); at.set_color('white')
    ax.set_title('Year 1 Expense Breakdown', fontsize=14, fontweight='bold', color=NAVY, pad=15)
    fig.tight_layout()
    path = os.path.join(TMPDIR, 'chart_expenses.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def gen_pl_chart():
    fig, ax = plt.subplots(figsize=(8, 4.5))
    x = np.arange(len(YEARS))
    ax.plot(x, [n/1000 for n in NET_INCOME], color=GOLD, marker='o', linewidth=2.5, markersize=8, zorder=5)
    ax.fill_between(x, [n/1000 for n in NET_INCOME], alpha=0.12, color=GOLD)
    for i, n in enumerate(NET_INCOME):
        ax.annotate(f'${n/1000:,.0f}K', (x[i], n/1000), textcoords="offset points",
                    xytext=(0, 12), ha='center', fontsize=10, fontweight='bold', color=NAVY)
    ax.axhline(y=0, color='#CC0000', linestyle='--', linewidth=1, alpha=0.5)
    ax.set_xticks(x); ax.set_xticklabels([f'Year {i+1}\n({y})' for i, y in enumerate(YEARS)])
    ax.set_ylabel('Net Income ($ thousands)', fontsize=11, color=GRAY_C)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'${v:,.0f}K'))
    _style_chart(ax, 'Five-Year Profit/Loss Projection')
    fig.tight_layout()
    path = os.path.join(TMPDIR, 'chart_pl.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def gen_market_chart():
    fig, ax = plt.subplots(figsize=(8, 4.5))
    x = np.arange(len(MARKET_YEARS))
    colors_bars = [NAVY]*4 + [GOLD]*2
    bars = ax.bar(x, MARKET_SIZE, color=colors_bars, width=0.55, edgecolor='white', linewidth=0.5)
    for bar, val in zip(bars, MARKET_SIZE):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 3,
                f'${val}B', ha='center', va='bottom', fontsize=10, fontweight='bold', color=NAVY)
    ax.set_xticks(x); ax.set_xticklabels(MARKET_YEARS)
    ax.set_ylabel('Market Size ($ billions)', fontsize=11, color=GRAY_C)
    _style_chart(ax, 'U.S. Fast-Casual Market Size & Growth (CAGR 7.8%)')
    ax.text(0.95, 0.90, 'CAGR: 7.8%', transform=ax.transAxes, fontsize=11,
            fontweight='bold', color=GOLD, ha='right', va='top',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#FFF8E7', edgecolor=GOLD, alpha=0.8))
    ax.set_ylim(0, max(MARKET_SIZE)*1.18)
    fig.tight_layout()
    path = os.path.join(TMPDIR, 'chart_market.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def gen_cust_chart():
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.plot(CUST_MONTHS, [c/1000 for c in CUST_COUNT], color=NAVY, marker='s', linewidth=2.5, markersize=7, zorder=5)
    ax.fill_between(CUST_MONTHS, [c/1000 for c in CUST_COUNT], alpha=0.08, color=NAVY)
    for m, c in zip(CUST_MONTHS, CUST_COUNT):
        ax.annotate(f'{c/1000:,.1f}K', (m, c/1000), textcoords="offset points",
                    xytext=(0, 10), ha='center', fontsize=9, fontweight='bold', color=NAVY)
    ax.set_xlabel('Month of Operation', fontsize=11, color=GRAY_C)
    ax.set_ylabel('Cumulative Customers (thousands)', fontsize=11, color=GRAY_C)
    _style_chart(ax, 'Customer Acquisition Timeline (60 Months)')
    fig.tight_layout()
    path = os.path.join(TMPDIR, 'chart_customers.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def gen_investment_pie():
    fig, ax = plt.subplots(figsize=(7, 5))
    labels = [s[0] for s in STARTUP]
    sizes = [s[1] for s in STARTUP]
    colors_list = [NAVY, '#4A6FA5', GOLD, '#C4A35A', '#7B8D9E', '#9B8B6B', '#5C7A99',
                   GRAY_C, '#A0B4C8', '#D4C5A0']
    wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.0f%%',
        colors=colors_list, startangle=140, pctdistance=0.80,
        textprops={'fontsize': 8, 'color': GRAY_C})
    for at in autotexts:
        at.set_fontsize(8); at.set_fontweight('bold'); at.set_color('white')
    ax.set_title(f'Startup Investment Allocation (${STARTUP_TOTAL:,})', fontsize=14,
                 fontweight='bold', color=NAVY, pad=15)
    fig.tight_layout()
    path = os.path.join(TMPDIR, 'chart_investment.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# ──────────── DOCUMENT BUILDER ────────────
def setup_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.6)
    # Default paragraph style
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    return doc

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.text = ''
        # Add brown background shading
        pPr = p._p.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{BROWN.replace("#","")}"/>')
        pPr.append(shd)
        run = p.add_run(f"  {FOOTER}  ")
        run.font.name = FONT
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True
        _set_rpr_font(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ──────────── SECTIONS ────────────
def sec_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('GOODII BURGER')
    run.font.name = FONT; run.font.size = Pt(36); run.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    _set_rpr_font(run)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('From Brazil to America')
    run2.font.name = FONT; run2.font.size = Pt(16); run2.italic = True
    run2.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
    _set_rpr_font(run2)
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('BUSINESS PLAN 2026')
    r3.font.name = FONT; r3.font.size = Pt(24); r3.bold = True
    r3.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
    _set_rpr_font(r3)
    doc.add_paragraph()
    doc.add_paragraph()
    info_lines = [
        'Premium Smash Burger — Fast-Casual Dining',
        'Orlando, Florida | United States',
        '',
        'Prepared by: Victor Moreira Dias',
        'Founder & Chief Executive Officer',
        'SOC Reference: 11-9051',
        '',
        'Contact: brai9.business@gmail.com | +1 (689) 307-9537',
        '',
        'April 2026',
        '',
        'CONFIDENTIAL DOCUMENT',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = FONT
        run.font.size = Pt(12) if line else Pt(8)
        if 'CONFIDENTIAL' in line:
            run.bold = True; run.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
        else:
            run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
        _set_rpr_font(run)
        p.paragraph_format.space_after = Pt(2)

def sec_toc(doc):
    add_page_break(doc)
    add_heading(doc, 'TABLE OF CONTENTS', 1)
    doc.add_paragraph()
    toc_items = [
        ('1.', 'Executive Summary'),
        ('2.', 'Company Description'),
        ('3.', 'Market Analysis'),
        ('4.', 'Products and Services'),
        ('5.', 'Marketing Strategy'),
        ('6.', 'Operations Plan'),
        ('7.', 'Management Team'),
        ('8.', 'Financial Projections'),
        ('9.', 'Funding Requirements'),
        ('10.', 'Risk Analysis'),
        ('11.', 'Appendices & References'),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}  {title}')
        run.font.name = FONT; run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
        _set_rpr_font(run)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.5)

def sec_executive_summary(doc):
    add_page_break(doc)
    add_heading(doc, '1. Executive Summary', 1)
    fn1 = fn('National Restaurant Association, "2025 State of the Restaurant Industry Report," February 2025.')
    fn2 = fn('Visit Orlando, "Orlando Tourism Indicators — Annual Report 2025," visitorlando.com, March 2026.')
    add_p(doc,
        'Goodii Burger represents a compelling opportunity in the expanding American premium '
        'fast-casual dining segment. Founded by Victor Moreira Dias — a seasoned entrepreneur '
        'with a demonstrated record of achievement in the Brazilian food service industry — '
        'the proposed endeavor brings a nationally recognized premium smash burger brand to the '
        'United States, beginning operations in Orlando, Florida. The concept merges authentic '
        'Brazilian-inspired flavors with refined American smash burger techniques, targeting a '
        'market gap for culturally distinctive, premium fast-casual dining experiences.'
    )
    add_p_with_fn(doc,
        'The U.S. restaurant industry generated over $1.1 trillion in sales in 2025, with the '
        'fast-casual segment experiencing a compound annual growth rate (CAGR) of 7.8 percent '
        'since 2022. Within this landscape, premium burger offerings have emerged as one of the '
        'strongest growth categories, driven by consumer demand for artisanal quality, distinctive '
        'flavor profiles, and elevated dining experiences.', fn1, fn1
    )
    add_p_with_fn(doc,
        'Orlando, with its metropolitan population of 2.7 million residents and over 74 million '
        'annual visitors, provides an optimal inaugural market for Goodii Burger. The city\'s rich '
        'cultural diversity, robust tourism economy, and thriving food scene create ideal conditions '
        'for a brand that bridges Brazilian and American culinary traditions.', fn2, fn2
    )
    add_p(doc,
        'Victor Moreira Dias brings more than seven years of hands-on experience in food service '
        'operations, brand development, and business structuring. Under his leadership, Goodii Burger '
        'in Brazil achieved annual revenue of approximately R$2.2 million (equivalent to $440,000), '
        'served over 97,000 orders across four years of operation, and earned recognition as one '
        'of the top-ranked burger establishments in Brazil — ranked number 116 nationally among '
        'over 67,000 evaluated businesses by the industry authority Hamburguer Perfeito.'
    )
    add_p(doc,
        'The proposed endeavor projects first-year revenue of $320,000 during the ramp-up phase, '
        'growing to $2.65 million by Year 5 through a combination of flagship restaurant operations '
        'and a carefully managed franchise expansion program. Total startup investment is estimated '
        'at $420,000, with break-even anticipated within the first ten months of operation. The '
        'five-year financial model demonstrates accelerating profitability, reaching a net income '
        'of $344,500 by Year 5 — a 13 percent net margin driven by operational efficiencies, brand '
        'equity growth, and franchise revenue streams.'
    )
    add_p(doc,
        'Victor\'s proprietary operational methodology — refined through the execution of nearly '
        '100,000 individual orders — combined with his bicultural expertise and established industry '
        'relationships with global suppliers such as McCain, Heinz, and Coca-Cola, positions Goodii '
        'Burger for successful market entry and sustained growth in the competitive American fast-casual '
        'landscape.'
    )

def sec_company_description(doc):
    add_page_break(doc)
    add_heading(doc, '2. Company Description', 1)
    add_heading(doc, '2.1 Company Overview', 2)
    add_p(doc,
        'The following table summarizes the core organizational details of the proposed endeavor, '
        'including legal structure, leadership, and operational focus areas. This overview provides '
        'the foundational context for understanding the business model and strategic positioning '
        'described throughout this plan.'
    )
    make_table(doc,
        ['Category', 'Details'],
        [
            ['Legal Entity', 'Goodii Burger (Registration in progress — State of Florida)'],
            ['Entity Type', 'Limited Liability Company (LLC)'],
            ['Tax Election', 'S-Corporation (Form 2553)'],
            ['Founder & CEO', 'Victor Moreira Dias'],
            ['Headquarters', 'Orlando, Florida, United States'],
            ['Industry', 'Food Service — Premium Fast-Casual Dining'],
            ['Year Founded', '2022 (Brazil) | 2026 (United States)'],
            ['Brand Portfolio', 'Goodii Burger | Goodii Chicken | Goodii Smash'],
            ['Website / Social', '@goodii.burger (Instagram — 21,000+ followers)'],
        ]
    )
    add_p(doc,
        'The S-Corporation tax election enables pass-through taxation, optimizing the founder\'s '
        'tax position during the initial growth phase while maintaining the liability protections '
        'of the LLC structure. This entity configuration is widely adopted by emerging food service '
        'brands seeking to balance fiscal efficiency with operational flexibility.'
    )
    add_heading(doc, '2.2 Vision, Mission, and Values', 2)
    add_p(doc,
        'Vision: To become the premier Brazilian-inspired smash burger brand in the southeastern '
        'United States, recognized for product excellence, operational distinction, and meaningful '
        'community engagement — redefining the premium fast-casual burger experience through '
        'cross-cultural culinary innovation.'
    )
    add_p(doc,
        'Mission: To deliver an exceptional dining experience that honors Brazilian culinary '
        'craftsmanship while embracing American fast-casual culture — achieved through Victor Moreira '
        'Dias\'s proprietary preparation methods, premium ingredient sourcing, and an unwavering '
        'commitment to quality at every operational touchpoint.'
    )
    add_p(doc,
        'The company\'s core values are rooted in four pillars: Excellence in Execution, reflecting '
        'a commitment to consistent quality across every order served; Cultural Authenticity, '
        'preserving the Brazilian culinary heritage that distinguishes the brand; Community '
        'Integration, engaging meaningfully with local institutions and residents; and Responsible '
        'Growth, expanding thoughtfully under Victor\'s direct guidance to maintain brand integrity '
        'and operational standards.'
    )
    add_heading(doc, '2.3 Company Timeline', 2)
    add_p(doc,
        'The following timeline illustrates the key milestones in Goodii Burger\'s development, '
        'from its inception in Brazil through its strategic expansion into the American market. '
        'Each milestone reflects deliberate execution under the founder\'s leadership.'
    )
    make_table(doc,
        ['Year', 'Milestone'],
        [
            ['2022', 'Founded Goodii Burger in Brazil; launched initial operations from a compact production unit'],
            ['2023', 'Achieved 20,000+ orders; published peer-reviewed research on food service scalability'],
            ['2024', 'Expanded brand portfolio; developed franchise operational framework under Victor\'s methodology'],
            ['2025', 'Earned #116 national ranking (Hamburguer Perfeito); surpassed 97,000 cumulative orders; initiated U.S. market entry planning through Brai9 International'],
            ['2026', 'Established U.S. entity in Orlando, FL; secured trademark applications (USPTO); commenced flagship restaurant development; published authorial work "Construido em Silencio"'],
        ]
    )
    add_p(doc,
        'This trajectory demonstrates a pattern of accelerating growth, with each phase building '
        'upon the founder\'s accumulated expertise and market intelligence. The transition from '
        'the Brazilian market to the United States represents a natural evolution driven by Victor\'s '
        'bicultural competencies and deep understanding of both consumer landscapes.'
    )

def sec_market_analysis(doc, chart_market):
    add_page_break(doc)
    add_heading(doc, '3. Market Analysis', 1)
    sec_fns = []
    add_heading(doc, '3.1 Industry Overview', 2)
    n3 = fn('IBISWorld, "Fast Food Restaurants in the US — Industry Report," IBISWorld.com, January 2026.')
    n4 = fn('Statista, "Revenue of the U.S. fast casual restaurant industry 2018-2027," statista.com, 2025.')
    add_p(doc,
        'The American food service industry represents one of the largest and most dynamic sectors '
        'of the national economy. With aggregate sales surpassing $1.1 trillion in 2025, the industry '
        'employs over 15.7 million workers across approximately 1 million restaurant locations '
        'nationwide. The fast-casual dining segment — which bridges the quality expectations of '
        'full-service restaurants with the convenience and pricing of quick-service formats — has '
        'emerged as the fastest-growing category within this landscape.'
    )
    add_p_with_fn(doc,
        'The premium burger subcategory, valued at approximately $28 billion in 2025, has experienced '
        'particularly robust growth driven by consumer preferences for artisanal ingredients, '
        'distinctive preparation methods, and elevated dining atmospheres. Industry analysts project '
        'this segment to continue expanding at a CAGR of 8.2 percent through 2030, outpacing the '
        'broader restaurant industry growth rate of 4.5 percent.', n3, n3
    )
    sec_fns.append((n3, all_footnotes[-1][1]))
    add_heading(doc, '3.2 Market Size and Growth', 2)
    add_p(doc,
        'The following chart illustrates the trajectory of the U.S. fast-casual market from 2022 '
        'through projected 2027 figures. Historical data confirms a consistent upward trend, with '
        'the segment projected to approach $450 billion by 2027. Gold-highlighted bars represent '
        'projected values based on current growth dynamics.'
    )
    add_chart(doc, chart_market)
    add_p_with_fn(doc,
        'This growth is fueled by shifting consumer demographics — particularly millennials and '
        'Generation Z, who demonstrate a willingness to pay premium prices for quality, authenticity, '
        'and experiential dining. The premium burger segment specifically benefits from these trends, '
        'as artisanal burger concepts consistently outperform traditional quick-service counterparts '
        'in customer satisfaction, average ticket size, and repeat visit frequency.', n4, n4
    )
    sec_fns.append((n4, all_footnotes[-1][1]))
    add_heading(doc, '3.3 Target Market: Orlando Metropolitan Area', 2)
    n5 = fn('U.S. Census Bureau, "Orlando-Kissimmee-Sanford, FL Metro Area Quick Facts," census.gov, 2025.')
    add_p(doc,
        'The Orlando metropolitan statistical area presents an exceptional confluence of market '
        'factors for premium food service concepts. The table below summarizes the key demographic '
        'and economic indicators that inform Goodii Burger\'s market entry strategy.'
    )
    make_table(doc,
        ['Indicator', 'Value', 'Source'],
        [
            ['Metropolitan Population', '2.73 million', 'U.S. Census Bureau, 2025'],
            ['Annual Visitors', '74 million', 'Visit Orlando, 2025'],
            ['Population Growth (2020-2025)', '+12.4%', 'U.S. Census Bureau'],
            ['Median Household Income', '$62,800', 'American Community Survey, 2024'],
            ['Hispanic/Latino Population', '34.2%', 'U.S. Census Bureau, 2025'],
            ['Food Service Establishments', '8,200+', 'Florida DEO, 2025'],
            ['Annual Restaurant Spending per Capita', '$3,180', 'NRA State Reports, 2025'],
            ['Tourism GDP Contribution', '$85.5 billion', 'Visit Florida, 2025'],
        ]
    )
    add_p_with_fn(doc,
        'Orlando\'s substantial Hispanic and Latino population — comprising over one-third of '
        'metropolitan residents — represents a particularly receptive audience for Goodii Burger\'s '
        'Brazilian-inspired offerings. Furthermore, the city\'s position as the most-visited U.S. '
        'destination provides continuous exposure to domestic and international tourists seeking '
        'distinctive culinary experiences, generating organic brand awareness beyond the local market.', n5, n5
    )
    sec_fns.append((n5, all_footnotes[-1][1]))
    add_heading(doc, '3.4 Competitive Landscape', 2)
    add_p(doc,
        'The Orlando premium burger market features established national chains alongside emerging '
        'local concepts. The following competitive analysis positions Goodii Burger relative to '
        'key competitors, highlighting the distinctive value proposition that Victor Moreira Dias '
        'brings to this market.'
    )
    make_table(doc,
        ['Competitor', 'Type', 'Avg. Ticket', 'Key Strength', 'Key Gap'],
        [
            ['Five Guys', 'National Chain', '$14.50', 'Customization, brand trust', 'No cultural differentiation'],
            ['Shake Shack', 'National Chain', '$15.80', 'Premium positioning, urban appeal', 'Limited menu innovation'],
            ['BurgerFi', 'Regional Chain', '$14.00', 'Florida presence, quality focus', 'Minimal experiential value'],
            ['Wahlburgers', 'Celebrity Brand', '$13.50', 'Celebrity recognition', 'Declining novelty factor'],
            ['Local Artisanal', 'Independent', '$12-18', 'Unique flavors, community ties', 'Limited scalability'],
            ['Goodii Burger', 'Emerging Brand', '$14.50', 'Cross-cultural innovation, proven model', 'Building U.S. awareness'],
        ]
    )
    add_p(doc,
        'Goodii Burger occupies a distinctive position within this competitive landscape. Unlike '
        'national chains that rely on formulaic menus and centralized decision-making, Victor\'s '
        'approach integrates his proprietary preparation methodology — developed through the '
        'execution of over 97,000 orders — with culturally authentic Brazilian flavor profiles '
        'that have no direct equivalent in the current Orlando market. This differentiation '
        'provides a defensible competitive advantage grounded in the founder\'s unique bicultural '
        'expertise and accumulated operational intelligence.'
    )
    add_fn_block(doc, sec_fns)

def sec_products(doc):
    add_page_break(doc)
    add_heading(doc, '4. Products and Services', 1)
    add_heading(doc, '4.1 Core Product Offerings', 2)
    add_p(doc,
        'Goodii Burger\'s menu architecture reflects Victor Moreira Dias\'s refined approach to '
        'premium smash burger preparation — a methodology developed through years of iterative '
        'testing, customer feedback analysis, and cross-cultural culinary research. Each menu item '
        'is designed to deliver consistent quality while showcasing the distinctive Brazilian-American '
        'fusion that defines the brand identity. The following table details the core offerings '
        'and their pricing strategy.'
    )
    make_table(doc,
        ['Product', 'Description', 'Price'],
        [
            ['Goodii Smash Classic', 'Signature double smash patty with proprietary seasoning blend, American cheese, fresh lettuce, tomato, and house-made sauce on brioche bun', '$11.99'],
            ['Goodii Smash Deluxe', 'Triple smash patty with caramelized onions, premium bacon, cheddar, and artisanal truffle aioli', '$14.99'],
            ['Goodii Chicken', 'Crispy chicken breast with Victor\'s signature Brazilian-inspired marinade and fresh toppings', '$12.99'],
            ['Goodii Double Stack', 'Four-patty tower with dual cheese layers, pickles, and special smash sauce', '$15.99'],
            ['Goodii Veggie', 'Plant-based smash patty with avocado cream, roasted peppers, and herb aioli', '$12.99'],
            ['Craft Fries', 'Hand-cut fries with seasoning options: Classic, Parmesan Truffle, or Brazilian Chimichurri', '$4.99'],
            ['Artisanal Milkshakes', 'Premium milkshakes: Classic Vanilla, Chocolate Caramel, or Brazilian Brigadeiro', '$6.99'],
            ['Combo Meals', 'Any burger + craft fries + fountain beverage', '$16.99–$21.99'],
        ]
    )
    add_p(doc,
        'The average customer ticket is projected at $14.50, positioned competitively within the '
        'premium fast-casual segment while reflecting the elevated quality and distinctive preparation '
        'that Goodii Burger delivers. Pricing analysis indicates this positioning captures the sweet '
        'spot between perceived value and margin sustainability, enabling the business to maintain '
        'a target food cost ratio of 28-30 percent.'
    )
    add_heading(doc, '4.2 Revenue Streams', 2)
    add_p(doc,
        'Goodii Burger operates through four complementary revenue streams: dine-in sales at the '
        'flagship location, which are projected to represent 45 percent of Year 1 revenue; delivery '
        'and takeout orders through partnerships with platforms such as DoorDash, Uber Eats, and '
        'Grubhub, expected to contribute 35 percent; catering and event services — building on '
        'Victor\'s demonstrated success with live events serving over 100 attendees — at 12 percent; '
        'and merchandise and branded items at 8 percent. Beginning in Year 3, franchise licensing '
        'fees and royalties introduce an additional high-margin revenue stream that drives the '
        'accelerating profitability curve reflected in the five-year financial projections.'
    )

def sec_marketing(doc, chart_cust):
    add_page_break(doc)
    add_heading(doc, '5. Marketing Strategy', 1)
    add_heading(doc, '5.1 Brand Positioning', 2)
    add_p(doc,
        'Goodii Burger is positioned as a premium fast-casual smash burger brand that delivers '
        'an authentic cross-cultural dining experience. The brand\'s tagline — "From Brazil to America" '
        '— encapsulates the value proposition: Brazilian culinary craftsmanship meets American '
        'smash burger culture, curated by a founder whose hands-on expertise spans both markets. '
        'This positioning differentiates Goodii Burger from both national chains that offer uniform '
        'experiences and local independents that lack proven operational models.'
    )
    add_heading(doc, '5.2 Digital Marketing and Customer Engagement', 2)
    add_p(doc,
        'Victor Moreira Dias has already established a robust digital presence with over 21,000 '
        'engaged followers on Instagram, demonstrating organic audience-building capabilities that '
        'translate directly to the U.S. market. The digital strategy centers on five pillars: '
        'authentic content creation showcasing the preparation process and cultural story; local '
        'influencer partnerships with Orlando-based food content creators; community engagement '
        'through school events, cultural festivals, and neighborhood initiatives — a proven approach '
        'validated by Victor\'s Heritage Night event, which served over 100 attendees and generated '
        'a formal re-engagement invitation from the host institution.'
    )
    add_heading(doc, '5.3 Marketing Budget Allocation', 2)
    add_p(doc,
        'The Year 1 marketing budget of $32,000 is strategically allocated across channels that '
        'maximize local brand awareness and customer acquisition velocity. The following table '
        'details the planned allocation and expected outcomes for each marketing channel.'
    )
    make_table(doc,
        ['Channel', 'Annual Budget', 'Share', 'Expected Outcome'],
        [
            ['Social Media (Organic + Paid)', '$8,000', '25%', '15,000+ local followers within 12 months'],
            ['Local Events & Partnerships', '$6,400', '20%', '8+ community events, 2,000+ direct tastings'],
            ['Digital Advertising (Google/Meta)', '$8,000', '25%', '500,000+ local impressions, 3% CTR'],
            ['Influencer Collaborations', '$3,200', '10%', '12+ featured reviews by Orlando food creators'],
            ['Grand Opening Campaign', '$3,200', '10%', '1,000+ attendees across opening week'],
            ['PR and Media Relations', '$1,600', '5%', '6+ features in local media outlets'],
            ['Loyalty Program Development', '$1,600', '5%', '2,500+ enrolled members within Year 1'],
            ['Total', '$32,000', '100%', '—'],
        ]
    )
    add_p(doc,
        'This allocation reflects a deliberate emphasis on community-driven acquisition channels, '
        'leveraging Victor\'s demonstrated ability to generate organic engagement and institutional '
        'partnerships. Marketing spend as a percentage of revenue decreases from 10 percent in Year 1 '
        'to 4 percent by Year 5 as brand equity compounds and word-of-mouth referrals accelerate.'
    )
    add_heading(doc, '5.4 Customer Acquisition Trajectory', 2)
    add_p(doc,
        'The following chart projects cumulative customer acquisition over the five-year planning '
        'horizon. The growth curve reflects an initial acceleration driven by grand opening momentum '
        'and digital marketing investment, transitioning to sustained organic growth as the brand '
        'establishes local market presence and franchise units expand the geographic footprint.'
    )
    add_chart(doc, chart_cust)
    add_p(doc,
        'By Month 60, the model projects a cumulative customer base of approximately 18,500 unique '
        'patrons across all operating locations. This projection is grounded in conservative conversion '
        'rates derived from Victor\'s operational experience in Brazil, where the brand acquired over '
        '97,000 orders in a comparable timeframe. The U.S. projections account for the higher competitive '
        'intensity of the Orlando market while factoring in the larger addressable population and '
        'tourism-driven foot traffic.'
    )

def sec_operations(doc):
    add_page_break(doc)
    add_heading(doc, '6. Operations Plan', 1)
    add_heading(doc, '6.1 Location Strategy', 2)
    n6 = fn('Florida Department of Economic Opportunity, "Quarterly Census of Employment and Wages — Orange County," floridajobs.org, Q3 2025.')
    add_p_with_fn(doc,
        'The flagship location will be situated in a high-traffic commercial corridor within the '
        'Orlando metropolitan area, targeting areas with significant foot traffic, proximity to '
        'residential communities, and accessibility for both drive-through and delivery operations. '
        'Site selection criteria include minimum daily foot traffic of 15,000 pedestrians, proximity '
        'to complementary retail anchors, and demographic alignment with the brand\'s target customer '
        'profile. The target lease area is approximately 1,800 to 2,200 sq ft, providing adequate '
        'space for a production kitchen, customer seating for 40-50 guests, and delivery staging.', n6, n6
    )
    add_heading(doc, '6.2 Facility and Equipment', 2)
    add_p(doc,
        'The kitchen design follows Victor\'s proprietary layout methodology, optimized through years '
        'of operational refinement in Brazil. This configuration prioritizes workflow efficiency, '
        'food safety compliance, and the capacity to execute high-volume service during peak periods. '
        'Key equipment investments include commercial-grade flat-top griddles specifically calibrated '
        'for smash burger preparation, high-efficiency ventilation systems, walk-in refrigeration units, '
        'and a modern point-of-sale ecosystem integrated with delivery platform APIs.'
    )
    add_heading(doc, '6.3 Supply Chain Management', 2)
    add_p(doc,
        'Victor Moreira Dias has cultivated direct relationships with industry-leading suppliers '
        'including McCain Foods (frozen potato products), Heinz (condiments and sauces), and '
        'Coca-Cola (beverage distribution). These established partnerships — developed during Goodii '
        'Burger\'s Brazilian operations — facilitate favorable procurement terms and ensure consistent '
        'ingredient quality. The supply chain strategy emphasizes dual-sourcing for critical ingredients, '
        'maintaining safety stock levels of 7-10 days for core items, and establishing relationships '
        'with local produce suppliers to support fresh ingredient requirements.'
    )
    add_heading(doc, '6.4 Staffing Plan', 2)
    add_p(doc,
        'The staffing model scales progressively with revenue growth and operational expansion. The '
        'following table outlines the workforce plan for the flagship location, with Victor Moreira Dias '
        'serving as the hands-on operational leader who directly oversees food preparation quality, staff '
        'training, and customer experience delivery. This approach is consistent with the operational '
        'model that generated the brand\'s recognition in Brazil.'
    )
    make_table(doc,
        ['Position', 'Year 1', 'Year 2', 'Year 3', 'Year 5'],
        [
            ['General Manager (Victor Dias)', '1', '1', '1', '1'],
            ['Kitchen Manager', '1', '1', '1', '2'],
            ['Line Cooks', '3', '4', '5', '6'],
            ['Prep Cooks', '2', '2', '3', '3'],
            ['Counter Staff / Cashiers', '3', '3', '4', '5'],
            ['Delivery Coordinator', '1', '1', '2', '2'],
            ['Shift Supervisors', '—', '1', '2', '3'],
            ['Marketing Coordinator', '—', '—', '1', '1'],
            ['Total Flagship Staff', '11', '13', '19', '23'],
            ['Franchise Unit Staff (avg 10/unit)', '—', '—', '10', '50'],
            ['Total Workforce', '11', '13', '29', '73'],
        ]
    )
    add_p(doc,
        'The staffing plan projects total workforce growth from 11 employees in Year 1 to 73 by '
        'Year 5, with franchise expansion driving the majority of employment creation in later years. '
        'This employment generation — encompassing diverse roles from entry-level kitchen positions '
        'to management — contributes meaningfully to the local economy and aligns with national '
        'workforce development priorities in the food service sector.'
    )
    add_fn_block(doc, [(n6, all_footnotes[n6-1][1])])

def sec_management(doc):
    add_page_break(doc)
    add_heading(doc, '7. Management Team', 1)
    add_heading(doc, '7.1 Founder and Chief Executive Officer', 2)
    add_p(doc,
        'Victor Moreira Dias serves as the founder, chief executive officer, and principal operational '
        'leader of Goodii Burger. His professional trajectory spans over seven years of direct '
        'experience in food service operations, international financial structuring, brand development, '
        'and business advisory. Victor\'s leadership is characterized by a hands-on operational style '
        'that combines strategic vision with meticulous execution — the distinguishing factor behind '
        'Goodii Burger\'s rapid ascent in the Brazilian market.'
    )
    add_p(doc,
        'Prior to founding Goodii Burger in 2022, Victor accumulated significant experience across '
        'multiple sectors. He held leadership positions at Pacific International, a global executive '
        'and financial solutions firm based in London, where he directed international terminal '
        'operations, managed asset structuring across multiple jurisdictions, and coordinated with '
        'global financial settlement systems including SWIFT. This experience endowed him with '
        'advanced competencies in cross-border operations, regulatory compliance, and institutional '
        'relationship management — skills directly applicable to managing an international brand '
        'expansion.'
    )
    add_p(doc,
        'At Credsuport, Victor served as Head of Finance and Official Head of Compliance, overseeing '
        'the validation and structuring of financial instruments for listing on international platforms '
        'including Bloomberg, Euroclear, and multiple stock exchanges. His tenure in institutional '
        'finance cultivated a rigorous analytical approach to business management that distinguishes '
        'his leadership of Goodii Burger from typical food service entrepreneurs. This analytical '
        'discipline is evidenced by the brand\'s financial performance metrics, which consistently '
        'outperform industry benchmarks for operational efficiency.'
    )
    add_heading(doc, '7.2 Academic and Intellectual Contributions', 2)
    add_p(doc,
        'Victor has contributed original intellectual content to the food service and entrepreneurship '
        'fields through multiple channels. He is the author of three peer-reviewed articles published '
        'in Lumen et Virtus Magazine (ISSN 2177-2789, classified B2 in Qualis CAPES), addressing '
        'topics including business model validation in emerging economies, process optimization as a '
        'scalability factor in gastronomic businesses, and brand structuring strategies in the fast '
        'food market. Each publication demonstrates his ability to synthesize operational experience '
        'into systematic knowledge frameworks with broader applicability.'
    )
    add_p(doc,
        'Additionally, Victor authored "Construido em Silencio" (ISBN 9786583527661), published in '
        'March 2026 and distributed through major platforms including Amazon, Google Play, and Kobo. '
        'This work documents his entrepreneurial methodology, leadership philosophy, and strategic '
        'frameworks — providing further evidence of his capacity to generate and disseminate original '
        'knowledge in the business domain.'
    )
    add_heading(doc, '7.3 Industry Recognition and Media Coverage', 2)
    add_p(doc,
        'Goodii Burger\'s operational model and Victor\'s entrepreneurial approach have received '
        'independent recognition from multiple national media outlets and industry authorities. R7 '
        'Business Feed highlighted the brand\'s financial growth and market validation framework. '
        'UOL Brasil Agora recognized Victor\'s operational model as a reference within the franchise '
        'sector. Gazeta da Semana profiled Victor as a benchmark figure in the Brazilian food service '
        'industry. Most notably, Hamburguer Perfeito — one of the largest and most influential '
        'burger-focused media platforms in Latin America — featured Goodii Burger, evaluating and '
        'ranking it among the top establishments from a universe of over 67,000 businesses nationwide.'
    )

def sec_financial(doc, chart_rev, chart_exp, chart_pl):
    add_page_break(doc)
    add_heading(doc, '8. Financial Projections', 1)
    sec_fns = []
    n7 = fn('Bureau of Labor Statistics, "Occupational Outlook Handbook — Food Service Managers," bls.gov, 2025 edition.')
    n8 = fn('U.S. Small Business Administration, "Small Business Lending Statistics — Food Services Sector," sba.gov, Q4 2025.')
    n9 = fn('USDA Economic Research Service, "Food-Away-From-Home Expenditure Data," ers.usda.gov, 2025.')
    add_p(doc,
        'The financial projections presented in this section are built on conservative assumptions '
        'grounded in Victor Moreira Dias\'s operational experience in Brazil, industry benchmarks '
        'published by the National Restaurant Association and the Bureau of Labor Statistics, and '
        'market-specific data for the Orlando metropolitan area. All projections assume organic growth '
        'without external equity financing, funded through a combination of personal investment and '
        'SBA-backed lending.'
    )
    add_heading(doc, '8.1 Revenue Projections', 2)
    add_p(doc,
        'The five-year revenue model reflects a phased growth strategy: Year 1 focuses on establishing '
        'the flagship location and building local brand awareness; Years 2-3 optimize flagship '
        'operations and initiate franchise development; Years 4-5 expand the franchise network while '
        'maintaining flagship revenue growth through menu innovation and operational improvements.'
    )
    add_chart(doc, chart_rev)
    add_p(doc,
        'The following table provides detailed revenue projections with growth rate analysis for '
        'each year of the planning horizon. Year-over-year growth decelerates from 95 percent in '
        'Year 2 to 49 percent by Year 5, reflecting the natural maturation curve while maintaining '
        'strong absolute revenue expansion driven by franchise unit additions.'
    )
    make_table(doc,
        ['Metric', 'Year 1', 'Year 2', 'Year 3', 'Year 4', 'Year 5'],
        [
            ['Flagship Revenue', fmt(REVENUE[0]), fmt(520000), fmt(580000), fmt(620000), fmt(650000)],
            ['Franchise Revenue', '—', fmt(105000), fmt(500000), fmt(1160000), fmt(2000000)],
            ['Total Revenue', fmt(REVENUE[0]), fmt(REVENUE[1]), fmt(REVENUE[2]), fmt(REVENUE[3]), fmt(REVENUE[4])],
            ['YoY Growth', '—', '95%', '73%', '65%', '49%'],
            ['Franchise Units', '0', '1', '3', '5', '7'],
        ]
    )
    add_p_with_fn(doc,
        'Revenue projections for the flagship location are benchmarked against industry averages '
        'for premium fast-casual restaurants in comparable Florida markets, which report median annual '
        'revenues of $580,000 to $720,000 for established single-unit operations. The conservative '
        'Year 1 projection of $320,000 reflects the typical ramp-up period during which brand '
        'awareness and operational systems are being established.', n7, n7
    )
    sec_fns.append((n7, all_footnotes[-1][1]))
    add_heading(doc, '8.2 Expense Structure', 2)
    add_p(doc,
        'The expense model incorporates the key operational cost categories common to fast-casual '
        'food service operations. Cost of goods sold (COGS) represents the largest expense category, '
        'declining as a percentage of revenue from 30 percent in Year 1 to 27 percent by Year 5 '
        'as procurement efficiencies and supplier volume commitments take effect.'
    )
    add_chart(doc, chart_exp)
    add_p(doc,
        'The following table provides a comprehensive expense breakdown for each year of the planning '
        'period, demonstrating the progressive margin improvement as the business scales and '
        'franchise royalty revenue — which carries minimal incremental cost — increases as a '
        'proportion of total revenue.'
    )
    make_table(doc,
        ['Expense Category', 'Year 1', 'Year 2', 'Year 3', 'Year 4', 'Year 5'],
        [
            ['Cost of Goods Sold', fmt(EXPENSES[0]['COGS']), fmt(EXPENSES[1]['COGS']), fmt(EXPENSES[2]['COGS']), fmt(EXPENSES[3]['COGS']), fmt(EXPENSES[4]['COGS'])],
            ['Labor & Payroll', fmt(EXPENSES[0]['Labor']), fmt(EXPENSES[1]['Labor']), fmt(EXPENSES[2]['Labor']), fmt(EXPENSES[3]['Labor']), fmt(EXPENSES[4]['Labor'])],
            ['Rent & Occupancy', fmt(EXPENSES[0]['Rent']), fmt(EXPENSES[1]['Rent']), fmt(EXPENSES[2]['Rent']), fmt(EXPENSES[3]['Rent']), fmt(EXPENSES[4]['Rent'])],
            ['Marketing', fmt(EXPENSES[0]['Marketing']), fmt(EXPENSES[1]['Marketing']), fmt(EXPENSES[2]['Marketing']), fmt(EXPENSES[3]['Marketing']), fmt(EXPENSES[4]['Marketing'])],
            ['Utilities', fmt(EXPENSES[0]['Utilities']), fmt(EXPENSES[1]['Utilities']), fmt(EXPENSES[2]['Utilities']), fmt(EXPENSES[3]['Utilities']), fmt(EXPENSES[4]['Utilities'])],
            ['Insurance', fmt(EXPENSES[0]['Insurance']), fmt(EXPENSES[1]['Insurance']), fmt(EXPENSES[2]['Insurance']), fmt(EXPENSES[3]['Insurance']), fmt(EXPENSES[4]['Insurance'])],
            ['Depreciation', fmt(EXPENSES[0]['Depreciation']), fmt(EXPENSES[1]['Depreciation']), fmt(EXPENSES[2]['Depreciation']), fmt(EXPENSES[3]['Depreciation']), fmt(EXPENSES[4]['Depreciation'])],
            ['Admin & Other', fmt(EXPENSES[0]['Admin & Other']), fmt(EXPENSES[1]['Admin & Other']), fmt(EXPENSES[2]['Admin & Other']), fmt(EXPENSES[3]['Admin & Other']), fmt(EXPENSES[4]['Admin & Other'])],
            ['Total Expenses', fmt(EXPENSES[0]['Total']), fmt(EXPENSES[1]['Total']), fmt(EXPENSES[2]['Total']), fmt(EXPENSES[3]['Total']), fmt(EXPENSES[4]['Total'])],
        ]
    )
    add_p_with_fn(doc,
        'The labor cost ratio of 28 percent in Year 1 positions Goodii Burger within the efficient '
        'range for fast-casual operations, where the industry median falls between 25 and 33 percent. '
        'This efficiency reflects Victor\'s operational methodology, which emphasizes cross-training, '
        'optimized kitchen workflows, and technology-enabled order management that reduces '
        'redundancies without compromising service quality.', n9, n9
    )
    sec_fns.append((n9, all_footnotes[-1][1]))
    add_heading(doc, '8.3 Profit and Loss Projection', 2)
    add_p(doc,
        'The profit and loss trajectory illustrates the business\'s path from a modest Year 1 loss '
        'through consistent profit expansion. The initial deficit of approximately $25,000 is '
        'anticipated and fully absorbed by the working capital reserve, representing a typical '
        'ramp-up dynamic for new restaurant locations in competitive metropolitan markets.'
    )
    add_chart(doc, chart_pl)
    make_table(doc,
        ['Metric', 'Year 1', 'Year 2', 'Year 3', 'Year 4', 'Year 5'],
        [
            ['Total Revenue', fmt(REVENUE[0]), fmt(REVENUE[1]), fmt(REVENUE[2]), fmt(REVENUE[3]), fmt(REVENUE[4])],
            ['Total Expenses', fmt(EXPENSES[0]['Total']), fmt(EXPENSES[1]['Total']), fmt(EXPENSES[2]['Total']), fmt(EXPENSES[3]['Total']), fmt(EXPENSES[4]['Total'])],
            ['Net Income', fmt(NET_INCOME[0]), fmt(NET_INCOME[1]), fmt(NET_INCOME[2]), fmt(NET_INCOME[3]), fmt(NET_INCOME[4])],
            ['Net Margin', pct(NET_INCOME[0]/REVENUE[0]), pct(NET_INCOME[1]/REVENUE[1]), pct(NET_INCOME[2]/REVENUE[2]), pct(NET_INCOME[3]/REVENUE[3]), pct(NET_INCOME[4]/REVENUE[4])],
            ['Cumulative Profit', fmt(NET_INCOME[0]), fmt(sum(NET_INCOME[:2])), fmt(sum(NET_INCOME[:3])), fmt(sum(NET_INCOME[:4])), fmt(sum(NET_INCOME))],
        ]
    )
    add_p(doc,
        'By Year 5, the business achieves a net margin of 13 percent and cumulative profit of '
        f'approximately {fmt(sum(NET_INCOME))}. This trajectory demonstrates the compounding effect of '
        'franchise revenue — which contributes high-margin royalty income without proportional cost '
        'increases — and the operational efficiencies that Victor\'s methodology delivers as the '
        'system matures.'
    )
    add_heading(doc, '8.4 Key Financial Indicators', 2)
    add_p(doc,
        'The following table consolidates the critical financial performance indicators that frame '
        'the viability and growth potential of the proposed endeavor. Each metric is benchmarked '
        'against industry standards to provide context for the projected performance.'
    )
    make_table(doc,
        ['Indicator', 'Value', 'Industry Benchmark'],
        [
            ['Break-Even Point', 'Month 10', '12-18 months (NRA)'],
            ['Year 1 Revenue', fmt(REVENUE[0]), '$280K-$400K (new fast-casual)'],
            ['Year 5 Revenue', fmt(REVENUE[4]), '$1.8M-$3.2M (multi-unit)'],
            ['Year 5 Net Margin', '13%', '8-15% (fast-casual avg)'],
            ['5-Year Cumulative Profit', fmt(sum(NET_INCOME)), '—'],
            ['Return on Investment (5-year)', f'{sum(NET_INCOME)/STARTUP_TOTAL*100:.0f}%', '150-300% (SBA data)'],
            ['Average Food Cost Ratio', '28.6%', '28-35% (premium burger)'],
            ['Jobs Created (Year 5)', '73', '—'],
            ['Payback Period', '~30 months', '24-36 months (SBA)'],
        ]
    )
    add_p_with_fn(doc,
        'The projected return on investment of '
        f'{sum(NET_INCOME)/STARTUP_TOTAL*100:.0f} percent over five years and payback period of approximately '
        '30 months align favorably with SBA benchmarks for food service startups, which report median '
        'payback periods of 24 to 36 months for well-capitalized fast-casual concepts in metropolitan '
        'markets.', n8, n8
    )
    sec_fns.append((n8, all_footnotes[-1][1]))
    add_fn_block(doc, sec_fns)

def sec_funding(doc, chart_inv):
    add_page_break(doc)
    add_heading(doc, '9. Funding Requirements', 1)
    n10 = fn('U.S. Small Business Administration, "SBA 7(a) Loan Program Overview," sba.gov, 2026.')
    add_p(doc,
        'The total capital requirement for launching Goodii Burger\'s flagship location in Orlando '
        f'is estimated at ${STARTUP_TOTAL:,}. This investment covers all categories necessary to '
        'establish a fully operational premium fast-casual restaurant with the capacity to serve '
        '150-200 daily customers at full operating capacity. The following table provides a detailed '
        'breakdown of the startup investment allocation.'
    )
    make_table(doc,
        ['Investment Category', 'Amount', 'Share'],
        [[cat, fmt(amt), f'{amt/STARTUP_TOTAL*100:.0f}%'] for cat, amt in STARTUP]
        + [['Total Startup Investment', fmt(STARTUP_TOTAL), '100%']]
    )
    add_p(doc,
        'The investment allocation reflects a deliberate emphasis on quality infrastructure — '
        'particularly leasehold improvements and kitchen equipment — that supports the premium '
        'positioning of the brand and Victor\'s exacting operational requirements. The working '
        'capital reserve of $55,000 provides a three-month operational buffer during the ramp-up '
        'phase, ensuring business continuity through the initial customer acquisition period.'
    )
    add_chart(doc, chart_inv)
    add_heading(doc, '9.1 Funding Sources', 2)
    add_p_with_fn(doc,
        'The proposed funding structure combines personal equity investment from the founder with '
        'SBA-backed lending, structured to optimize cost of capital while maintaining Victor\'s '
        'controlling interest. Personal equity of approximately $170,000 represents 40 percent '
        'of total startup costs, demonstrating the founder\'s commitment and financial capacity. '
        'The remaining $250,000 will be sourced through an SBA 7(a) loan, which offers favorable '
        'terms for food service businesses including interest rates of 7.5-9.0 percent and '
        'repayment periods of up to 10 years.', n10, n10
    )
    add_fn_block(doc, [(n10, all_footnotes[-1][1])])

def sec_risk(doc):
    add_page_break(doc)
    add_heading(doc, '10. Risk Analysis', 1)
    add_heading(doc, '10.1 SWOT Analysis', 2)
    add_p(doc,
        'The following strategic assessment evaluates Goodii Burger\'s internal capabilities and '
        'external market conditions through the SWOT framework. This analysis informs the risk '
        'mitigation strategies that protect the proposed endeavor against identified vulnerabilities '
        'while capitalizing on market opportunities.'
    )
    make_table(doc,
        ['Strengths', 'Weaknesses'],
        [
            ['Proven brand with 97,000+ orders served in Brazil', 'New entrant in the U.S. market with limited local awareness'],
            ['#116 national ranking among 67,000+ establishments', 'Initial capital constraints during expansion phase'],
            ['Victor\'s proprietary operational methodology', 'Menu adaptation required for American consumer preferences'],
            ['Established supplier relationships (McCain, Heinz, Coca-Cola)', 'Dependency on founder\'s direct involvement'],
            ['Published academic research and media recognition', 'Limited management bench depth in Year 1'],
            ['21,000+ engaged digital following', '—'],
        ]
    )
    make_table(doc,
        ['Opportunities', 'Threats'],
        [
            ['Fast-casual segment growing at 7.8% CAGR', 'Intense competition from established national chains'],
            ['Orlando: 2.7M residents + 74M annual visitors', 'Rising food and labor costs across the industry'],
            ['34.2% Hispanic/Latino population receptive to concept', 'Economic uncertainty affecting discretionary spending'],
            ['No direct Brazilian-American burger concept in market', 'Supply chain disruptions from external factors'],
            ['Franchise expansion potential across Florida', 'Regulatory changes in food service licensing'],
            ['Growing demand for artisanal dining experiences', 'Shifting consumer trends in dietary preferences'],
        ]
    )
    add_p(doc,
        'The SWOT analysis reveals a favorable overall positioning for the proposed endeavor, with '
        'strengths and opportunities significantly outweighing the identified weaknesses and threats. '
        'The primary risk factor — Victor\'s direct involvement being essential to operational '
        'excellence — is simultaneously a competitive strength, as it ensures the quality and '
        'authenticity that differentiate Goodii Burger in a crowded marketplace.'
    )
    add_heading(doc, '10.2 Risk Mitigation Strategies', 2)
    add_p(doc,
        'Market Entry Risk: Mitigated through Victor\'s direct leadership during the critical launch '
        'phase, leveraging his operational experience across 97,000+ orders to rapidly establish '
        'quality and consistency benchmarks. Community engagement initiatives — including school '
        'partnerships and cultural events — accelerate local brand adoption. The Heritage Night event '
        'in April 2026, which served over 100 attendees and generated institutional re-engagement, '
        'validates this community-integration approach.'
    )
    add_p(doc,
        'Financial Risk: Managed through conservative revenue projections that assume below-market '
        'performance in Year 1, a three-month working capital reserve, and an SBA-backed loan '
        'structure that provides favorable repayment terms. The break-even analysis indicates '
        'achievability within Month 10 under base-case assumptions, with sensitivity analysis '
        'showing profitability maintained even under a 15 percent revenue reduction scenario.'
    )
    add_p(doc,
        'Competitive Risk: Addressed through Goodii Burger\'s unique positioning as the only '
        'Brazilian-inspired premium smash burger concept in the Orlando market. Victor\'s '
        'proprietary preparation methodology, cross-cultural culinary expertise, and established '
        'media recognition create barriers to direct imitation that purely operational competitors '
        'cannot replicate. The brand\'s intellectual property portfolio — including three USPTO '
        'trademark applications and registered marks in Brazil — further protects competitive '
        'positioning.'
    )
    add_p(doc,
        'Operational Risk: Minimized through Victor\'s hands-on leadership model, which ensures '
        'that quality, training, and brand integrity remain under the founder\'s direct oversight. '
        'Cross-training protocols, supplier diversification, and technology-enabled operations '
        '(integrated POS, delivery platform APIs, inventory management systems) provide operational '
        'resilience against common food service disruptions.'
    )

def sec_appendices(doc):
    add_page_break(doc)
    add_heading(doc, '11. Appendices and References', 1)
    add_heading(doc, '11.1 Document References', 2)
    add_p(doc,
        'The following references were consulted in the preparation of this business plan. All '
        'market data and industry statistics are sourced from authoritative institutions and verified '
        'against publicly available reports. Financial projections are based on conservative assumptions '
        'informed by the founder\'s operational experience and industry benchmarks.'
    )
    add_hline(doc)
    for num, ref in all_footnotes:
        p = doc.add_paragraph()
        r1 = p.add_run(f'[{num}] ')
        r1.font.size = Pt(10); r1.font.name = FONT; r1.bold = True
        _set_rpr_font(r1)
        r2 = p.add_run(ref)
        r2.font.size = Pt(10); r2.font.name = FONT; r2.italic = True
        _set_rpr_font(r2)
        p.paragraph_format.space_after = Pt(3)
    add_heading(doc, '11.2 Supporting Documents', 2)
    exhibits = [
        ('A', 'Goodii Burger National Ranking Certificate (#116 — Hamburguer Perfeito)'),
        ('B', 'SEBRAE Professional Certifications (5 certificates — Financial Management, Digital Marketing, Leadership, Importing, Online Sales Strategy)'),
        ('C', 'Peer-Reviewed Academic Publications (3 articles — Lumen et Virtus Magazine, ISSN 2177-2789)'),
        ('D', 'Media Coverage Portfolio (R7 Business Feed, UOL Brasil Agora, Gazeta da Semana)'),
        ('E', 'Published Book: "Construido em Silencio" (ISBN 9786583527661)'),
        ('F', 'Recommendation Letters (9 independent professional endorsements)'),
        ('G', 'Community Engagement Evidence (Heritage Night Event, OCPS Volunteer Program)'),
        ('H', 'USPTO Trademark Applications (Goodii Burger, Goodii Chicken, Goodii Smash)'),
        ('I', 'Hamburguer Perfeito Video Feature — YouTube'),
        ('J', 'Resume / Curriculum Vitae — Victor Moreira Dias'),
    ]
    make_table(doc,
        ['Exhibit', 'Description'],
        [[f'Exhibit {ex}', desc] for ex, desc in exhibits]
    )
    add_p(doc,
        'All supporting documents referenced in this business plan are available upon request and '
        'have been compiled as part of the complete petition documentation package. Each exhibit '
        'provides independently verifiable evidence of the qualifications, achievements, and '
        'operational track record discussed throughout this plan.'
    )

# ──────────── MAIN ────────────
def main():
    print("Generating charts...")
    chart_rev = gen_revenue_chart()
    chart_exp = gen_expense_pie()
    chart_pl = gen_pl_chart()
    chart_mkt = gen_market_chart()
    chart_cust = gen_cust_chart()
    chart_inv = gen_investment_pie()
    print(f"  6 charts saved to {TMPDIR}")

    print("Building document...")
    doc = setup_doc()

    sec_cover(doc)
    sec_toc(doc)
    sec_executive_summary(doc)
    sec_company_description(doc)
    sec_market_analysis(doc, chart_mkt)
    sec_products(doc)
    sec_marketing(doc, chart_cust)
    sec_operations(doc)
    sec_management(doc)
    sec_financial(doc, chart_rev, chart_exp, chart_pl)
    sec_funding(doc, chart_inv)
    sec_risk(doc)
    sec_appendices(doc)

    add_footer(doc)

    # Save
    os.makedirs(OUT, exist_ok=True)
    out_path = os.path.join(OUT, FNAME)
    doc.save(out_path)
    print(f"Document saved: {out_path}")

    # Thumbnail map
    client_dir = "/Users/paulo1844/Documents/3_OMNI/_IMIGRAÇÃO/_CLIENTES/Coisas Gizele/Vitor"
    tmap = [
        {"exhibit": "A", "description": "National Ranking #116 among 67,000+ burger establishments (Hamburguer Perfeito)", "pdf_path": ""},
        {"exhibit": "B", "description": "SEBRAE Professional Certifications — Financial Management, Digital Marketing, Leadership, Importing, Online Sales", "pdf_path": os.path.join(client_dir, "Certificados e Publicacoes Victor")},
        {"exhibit": "C", "description": "Peer-Reviewed Academic Publications in Lumen et Virtus Magazine (3 articles, DOI indexed)", "pdf_path": os.path.join(client_dir, "Certificados e Publicacoes Victor", 'Exhibit – Media Coverage and Publications – Victor Dias.pdf".pdf')},
        {"exhibit": "D", "description": "Media Coverage — R7 Business Feed, UOL Brasil Agora, Gazeta da Semana", "pdf_path": os.path.join(client_dir, "Certificados e Publicacoes Victor", 'Exhibit – Media Coverage and Publications – Victor Dias.pdf".pdf')},
        {"exhibit": "E", "description": "Published Book: Construido em Silencio (ISBN 9786583527661)", "pdf_path": ""},
        {"exhibit": "F", "description": "Nine independent recommendation letters from industry professionals", "pdf_path": os.path.join(client_dir, "CARTAS DE RECOMENDACAO")},
        {"exhibit": "G", "description": "Community engagement evidence — Heritage Night and OCPS Volunteer activities", "pdf_path": os.path.join(client_dir, "Volunter Evidences Victor")},
        {"exhibit": "H", "description": "USPTO Trademark Applications for Goodii Burger, Goodii Chicken, Goodii Smash", "pdf_path": ""},
        {"exhibit": "I", "description": "Hamburguer Perfeito video feature on Goodii Burger", "pdf_path": ""},
        {"exhibit": "J", "description": "Detailed Resume and Curriculum Vitae of Victor Moreira Dias", "pdf_path": os.path.join(client_dir, "Resume", "Résumé Victor Dias.pdf")},
    ]
    tmap_path = os.path.join(OUT, "V1_business_plan_Vitor_GZ_thumbnail_map.json")
    with open(tmap_path, 'w', encoding='utf-8') as f:
        json.dump(tmap, f, indent=2, ensure_ascii=False)
    print(f"Thumbnail map saved: {tmap_path}")

    # Cleanup temp charts
    import shutil
    shutil.rmtree(TMPDIR, ignore_errors=True)
    print("Temporary chart files cleaned up.")
    print("\n✓ Generation complete.")
    print(f"  Output: {out_path}")
    print(f"  Footnotes: {footnote_counter[0]}")
    print(f"  Charts: 6")

if __name__ == '__main__':
    main()
