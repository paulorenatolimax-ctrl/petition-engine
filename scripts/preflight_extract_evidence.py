#!/usr/bin/env python3
"""
preflight_extract_evidence.py — Pre-flight evidence text extraction.

Executes BEFORE any prompt build in cover-letter pipelines. For each evidence
file (.pdf/.docx/.png/.jpg/.jpeg) found under the client's evidence directory:

  1. Compute SHA256 of file bytes (idempotency + cache key)
  2. Extract text:
     - PDF: pdfplumber per-page; fallback to OCR via pdftoppm + pytesseract
       when extracted text is < 50 chars (scanned PDFs, image-only PDFs)
     - DOCX: python-docx (paragraphs + tables) + OCR on embedded images
     - PNG/JPG/JPEG: pytesseract direct
  3. Persist {sha256}.txt under outputDir
  4. Append manifest entry

Exits 0 iff every file extracted >= 50 chars. Otherwise exits 1 and the
manifest still records the failures (the pipeline reads it to abort with a
specific error citing which files starved).

Background: this exists because Cowork (and earlier sessions) sometimes wrote
about evidences by reading the FILENAME instead of the CONTENT — producing
factually wrong narrative tied to whatever the file was named, not what it
actually contained. r220 is the rule that this script enforces.

Usage:
  python3 scripts/preflight_extract_evidence.py \\
      --client-id gustavo_nelson_sales_chaves \\
      --evidence-dir "/path/to/client/Evidence" \\
      --output-dir   "data/evidence_extracted/gustavo_nelson_sales_chaves"

If --output-dir is omitted, defaults to data/evidence_extracted/{client-id}/
relative to the petition-engine repo root (cwd of caller).
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

CHAR_THRESHOLD = 50
SUPPORTED_EXTS = {".pdf", ".docx", ".png", ".jpg", ".jpeg"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg"}
SKIP_DIR_NAMES = {
    "_Forjado por Petition Engine",
    "phases",
    ".git",
    "node_modules",
    ".next",
    "__pycache__",
    "_LIXO",
    "lixo",
}


def sha256_of_file(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(64 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def has_tesseract() -> bool:
    return shutil.which("tesseract") is not None


def has_pdftoppm() -> bool:
    return shutil.which("pdftoppm") is not None


def ocr_image_file(path: Path) -> str:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except ImportError:
        return ""
    try:
        with Image.open(path) as img:
            return pytesseract.image_to_string(img, lang="por+eng") or ""
    except Exception:
        return ""


def extract_pdf(path: Path) -> tuple[str, bool]:
    """Return (text, ocr_used). pdfplumber first; OCR fallback if < threshold."""
    text_parts: list[str] = []
    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    text_parts.append(t)
    except Exception:
        pass

    text = "\n".join(text_parts)
    if len(text.strip()) >= CHAR_THRESHOLD:
        return text, False

    if not (has_tesseract() and has_pdftoppm()):
        return text, False

    with tempfile.TemporaryDirectory() as tmp:
        try:
            subprocess.run(
                ["pdftoppm", "-r", "200", "-png", str(path), str(Path(tmp) / "page")],
                check=True, capture_output=True, timeout=300,
            )
        except (subprocess.SubprocessError, FileNotFoundError):
            return text, False
        ocr_parts: list[str] = []
        for img_path in sorted(Path(tmp).glob("page*.png")):
            ocr_parts.append(ocr_image_file(img_path))
        ocr_text = "\n".join(ocr_parts).strip()
        if ocr_text:
            return ocr_text, True
    return text, False


def extract_docx(path: Path) -> tuple[str, bool]:
    """Return (text, ocr_used). python-docx for paragraphs+tables; OCR on embedded images if text is starved."""
    text_parts: list[str] = []
    ocr_used = False
    try:
        from docx import Document  # type: ignore
        doc = Document(str(path))
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
    except Exception:
        pass

    text = "\n".join(text_parts)
    if len(text.strip()) >= CHAR_THRESHOLD:
        return text, False

    if not has_tesseract():
        return text, False

    try:
        with zipfile.ZipFile(path) as zf, tempfile.TemporaryDirectory() as tmp:
            ocr_parts: list[str] = []
            for member in zf.namelist():
                if not member.startswith("word/media/"):
                    continue
                lower = member.lower()
                if not (lower.endswith(".png") or lower.endswith(".jpg") or lower.endswith(".jpeg")):
                    continue
                out = Path(tmp) / Path(member).name
                with zf.open(member) as src, out.open("wb") as dst:
                    shutil.copyfileobj(src, dst)
                ocr_parts.append(ocr_image_file(out))
            ocr_text = "\n".join(p for p in ocr_parts if p.strip())
            if ocr_text:
                ocr_used = True
                text = (text + "\n" + ocr_text).strip() if text.strip() else ocr_text
    except (zipfile.BadZipFile, OSError):
        pass

    return text, ocr_used


def extract_image(path: Path) -> tuple[str, bool]:
    if not has_tesseract():
        return "", False
    return ocr_image_file(path), True


def walk_evidences(root: Path) -> list[Path]:
    """Recursive walk filtering noise dirs and non-evidence file types."""
    out: list[Path] = []
    if not root.exists():
        return out
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d not in SKIP_DIR_NAMES and not d.startswith(".")]
        for fn in filenames:
            if fn.startswith(".") or fn.startswith("~$"):
                continue
            ext = Path(fn).suffix.lower()
            if ext not in SUPPORTED_EXTS:
                continue
            out.append(Path(dirpath) / fn)
    out.sort()
    return out


def classify_status(char_count: int, ext: str, photo_mode: bool, threshold: int) -> str:
    """Returns 'ok', 'image_only', or 'starved'.

    image_only: image file (.png/.jpg/.jpeg) with text below threshold AND --photo-evidence-mode active.
    These are visual evidences (photos of subject's work, e.g. dolls/products) where text-extraction
    isn't expected. Pipeline must describe them via filename + contextual inference, NOT pretend it
    read content. Distinguished from 'starved' so r220 can pass while still flagging textual files
    that genuinely failed extraction.
    """
    if char_count >= threshold:
        return "ok"
    if photo_mode and ext in IMAGE_EXTS:
        return "image_only"
    return "starved"


def build_manifest_entry(evidence_root: Path, file: Path, output_dir: Path, photo_mode: bool, threshold: int) -> dict[str, Any]:
    digest = sha256_of_file(file)
    txt_path = output_dir / f"{digest}.txt"
    ext = file.suffix.lower()

    if txt_path.exists() and txt_path.stat().st_size > 0:
        cached = txt_path.read_text(encoding="utf-8", errors="replace")
        return {
            "filename": file.name,
            "relative_path": str(file.relative_to(evidence_root)) if evidence_root in file.parents or evidence_root == file.parent else str(file),
            "absolute_path": str(file),
            "sha256": digest,
            "type": ext.lstrip("."),
            "char_count": len(cached.strip()),
            "ocr_used": None,
            "extracted_text_path": str(txt_path),
            "status": classify_status(len(cached.strip()), ext, photo_mode, threshold),
            "cached": True,
        }

    text = ""
    ocr_used = False
    error: str | None = None
    try:
        if ext == ".pdf":
            text, ocr_used = extract_pdf(file)
        elif ext == ".docx":
            text, ocr_used = extract_docx(file)
        elif ext in {".png", ".jpg", ".jpeg"}:
            text, ocr_used = extract_image(file)
    except Exception as e:
        error = f"{type(e).__name__}: {e}"

    text = (text or "").strip()
    txt_path.write_text(text, encoding="utf-8")

    try:
        rel = str(file.relative_to(evidence_root))
    except ValueError:
        rel = str(file)

    entry: dict[str, Any] = {
        "filename": file.name,
        "relative_path": rel,
        "absolute_path": str(file),
        "sha256": digest,
        "type": ext.lstrip("."),
        "char_count": len(text),
        "ocr_used": ocr_used,
        "extracted_text_path": str(txt_path),
        "status": classify_status(len(text), ext, photo_mode, threshold),
        "cached": False,
    }
    if error:
        entry["error"] = error
    return entry


def main() -> int:
    ap = argparse.ArgumentParser(description="Pre-flight evidence text extraction (rule r220).")
    ap.add_argument("--client-id", required=True)
    ap.add_argument("--evidence-dir", required=True, help="Root directory of client evidences")
    ap.add_argument("--output-dir", default=None, help="Where to write {sha256}.txt + manifest.json")
    ap.add_argument("--threshold", type=int, default=CHAR_THRESHOLD)
    ap.add_argument(
        "--photo-evidence-mode",
        action="store_true",
        help=("When set, .jpg/.png/.jpeg files that extract < threshold chars are tagged "
              "'image_only' (not 'starved'). Pipeline accepts these as visual-only evidences "
              "to be described via filename + contextual inference. Use ONLY for clients whose "
              "evidence is dominantly photographic (e.g. fashion design portfolios, art works). "
              "Default off — abuse re-introduces the Cowork 'reading filename instead of content' bug."),
    )
    args = ap.parse_args()

    evidence_root = Path(args.evidence_dir).resolve()
    if args.output_dir:
        output_dir = Path(args.output_dir).resolve()
    else:
        output_dir = Path.cwd() / "data" / "evidence_extracted" / args.client_id
    output_dir.mkdir(parents=True, exist_ok=True)

    files = walk_evidences(evidence_root)
    entries = [
        build_manifest_entry(evidence_root, f, output_dir, args.photo_evidence_mode, args.threshold)
        for f in files
    ]

    ok_count = sum(1 for e in entries if e["status"] == "ok")
    image_only_count = sum(1 for e in entries if e["status"] == "image_only")
    starved_count = sum(1 for e in entries if e["status"] == "starved")

    manifest = {
        "client_id": args.client_id,
        "evidence_root": str(evidence_root),
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "threshold": args.threshold,
        "photo_evidence_mode": args.photo_evidence_mode,
        "evidence_count": len(entries),
        "ok_count": ok_count,
        "image_only_count": image_only_count,
        "starved_count": starved_count,
        "entries": entries,
    }
    manifest_path = output_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    starved = [e for e in entries if e["status"] == "starved"]
    print(f"[preflight] client={args.client_id} files={len(entries)} ok={ok_count} image_only={image_only_count} starved={len(starved)} photo_mode={args.photo_evidence_mode}")
    if starved:
        print(f"[preflight] STARVED (< {args.threshold} chars; non-image OR photo_mode=off):")
        for e in starved[:20]:
            print(f"  - {e['filename']} ({e['char_count']} chars)")
        if len(starved) > 20:
            print(f"  ... and {len(starved) - 20} more")
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
