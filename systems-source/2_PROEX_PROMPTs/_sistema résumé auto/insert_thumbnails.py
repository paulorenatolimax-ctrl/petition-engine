#!/usr/bin/env python3
"""
insert_thumbnails.py — Universal thumbnail insertion for DOCX documents.

Supports TWO placeholder patterns:
1. Résumé pattern:      [THUMBNAIL — Exhibit X] or [THUMBNAIL — Exhibits X–Y]
                        → Maps by exhibit number from DOCX exhibit index
2. Cover letter pattern: [THUMBNAIL] (simple, no number)
                        → Uses adjacent cell metadata to fuzzy-match PDFs

Accepts optional JSON override: --map thumbnail_map.json
If thumbnail_map.json exists in the same folder as the DOCX, loads it automatically.

Dependencies: pip install PyMuPDF python-docx Pillow
"""

import sys
import os
import re
import json
import argparse
import tempfile
from difflib import SequenceMatcher
from typing import Optional
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
THUMBNAIL_WIDTH_INCHES = 2.6       # Standard width
THUMBNAIL_WIDTH_COMPACT = 2.0      # Compact width (for multi-exhibit cells)
THUMBNAIL_DPI = 150                # Render resolution
BORDER_COLOR = "#CCCCCC"
BORDER_WIDTH_PT = 0.5

# Patterns that indicate a translation certificate page (page to SKIP)
CERT_PATTERNS = [
    r"i\s+hereby\s+certif",
    r"certif(?:y|ied)\s+(?:that|translation)",
    r"certified\s+translation",
    r"sworn\s+translat",
    r"tradutor(?:a)?\s+juramentad",
    r"notari(?:y|zed)",
    r"tabelião|tabeliã",
    r"translation\s+company",
    r"empresa\s+de\s+tradução",
    r"certificate\s+of\s+accuracy",
    r"certificado\s+de\s+fidelidade",
    r"translat(?:ion|or)\s+certificate",
    r"i,?\s+the\s+undersigned",
    r"accuracy\s+of\s+(?:the|this)\s+translation",
]
CERT_RE = re.compile("|".join(CERT_PATTERNS), re.IGNORECASE)

# --- Placeholder patterns ---
# Pattern 1 (résumé): [THUMBNAIL — Exhibit 11] or [THUMBNAIL — Exhibits 5–7]
RESUME_PLACEHOLDER_RE = re.compile(
    r"\[THUMBNAIL\s*[-—–]\s*Exhibits?\s+([\d]+)(?:\s*[-—–]\s*(\d+))?\]"
)
# Pattern 2 (cover letter): [THUMBNAIL] alone (no exhibit number)
COVER_PLACEHOLDER_RE = re.compile(r"^\s*\[THUMBNAIL\]\s*$")

# Evidence number extraction from metadata cell text
EVIDENCE_NUM_RE = re.compile(r"Evidence\s+(\d+)\.", re.IGNORECASE)

# Stopwords for fuzzy matching (low information tokens)
STOPWORDS = {
    "the", "a", "an", "of", "in", "on", "at", "to", "for", "and", "or",
    "is", "are", "was", "were", "be", "been", "de", "da", "do", "dos",
    "das", "em", "no", "na", "nos", "nas", "com", "por", "para", "um",
    "uma", "que", "se", "e", "o", "evidence", "type", "source", "date",
    "url", "n/a", "descrição", "relevância", "description", "relevance",
    "tipo", "fonte", "data",
}


# ---------------------------------------------------------------------------
# Helper: detect translation certificate page
# ---------------------------------------------------------------------------
def is_translation_certificate(page_text: str) -> bool:
    """Return True if the page text matches translation certificate patterns."""
    matches = len(CERT_RE.findall(page_text))
    return matches >= 2


# ---------------------------------------------------------------------------
# Helper: extract best page from PDF as PNG
# ---------------------------------------------------------------------------
def extract_thumbnail(pdf_path: str, temp_dir: str, label: str) -> Optional[str]:
    """
    Extract the best thumbnail page from a PDF.
    Returns path to PNG file, or None if PDF can't be processed.
    label: display label (e.g. "Exhibit 5" or "Evidence 12")
    """
    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        print(f"  [WARN] Cannot open PDF {pdf_path}: {e}")
        return None

    if doc.page_count == 0:
        print(f"  [WARN] PDF has 0 pages: {pdf_path}")
        doc.close()
        return None

    # Check page 1 text for certificate patterns
    page_idx = 0
    page1 = doc[0]
    page1_text = page1.get_text("text")

    if doc.page_count > 1 and is_translation_certificate(page1_text):
        print(f"  [INFO] {label}: Page 1 is translation certificate → using page 2")
        page_idx = 1
    else:
        print(f"  [INFO] {label}: Using page 1")

    # Render page to PNG
    page = doc[page_idx]
    mat = fitz.Matrix(THUMBNAIL_DPI / 72, THUMBNAIL_DPI / 72)
    pix = page.get_pixmap(matrix=mat)

    # Use a safe filename
    safe_label = re.sub(r"[^\w]", "_", label)
    png_path = os.path.join(temp_dir, f"{safe_label}.png")
    pix.save(png_path)
    doc.close()

    return png_path


# ---------------------------------------------------------------------------
# Helper: tokenize text for fuzzy matching
# ---------------------------------------------------------------------------
def tokenize(text: str) -> set:
    """Extract meaningful tokens from text, lowercased, without stopwords."""
    tokens = re.findall(r"[a-zA-ZÀ-ÿ0-9]+", text.lower())
    return {t for t in tokens if t not in STOPWORDS and len(t) > 1}


# ---------------------------------------------------------------------------
# Helper: fuzzy match a title against PDF filenames
# ---------------------------------------------------------------------------
def fuzzy_match_pdf(title: str, pdf_files: list, used_pdfs: set,
                    evidence_num: Optional[int] = None) -> Optional[str]:
    """
    Find the best PDF match for a given evidence title.
    Returns the filename or None.
    """
    title_tokens = tokenize(title)
    if not title_tokens:
        return None

    best_match = None
    best_score = 0.0

    for pdf_file in pdf_files:
        if pdf_file in used_pdfs:
            continue

        pdf_name = os.path.splitext(pdf_file)[0]
        pdf_tokens = tokenize(pdf_name)

        if not pdf_tokens:
            continue

        # Token overlap score (Jaccard-like)
        overlap = title_tokens & pdf_tokens
        if not overlap:
            continue

        # Score = overlap size / min(title tokens, pdf tokens)
        # This favors PDFs that match many of the title's keywords
        score = len(overlap) / min(len(title_tokens), len(pdf_tokens))

        # Bonus: if evidence number appears in filename (e.g. "Evidence_12")
        if evidence_num is not None:
            ev_patterns = [f"evidence_{evidence_num}", f"evidence {evidence_num}",
                           f"ev_{evidence_num}", f"ev {evidence_num}",
                           f"exhibit_{evidence_num}", f"exhibit {evidence_num}"]
            pdf_lower = pdf_file.lower()
            for pat in ev_patterns:
                if pat in pdf_lower:
                    score += 0.5
                    break

        # Bonus: SequenceMatcher ratio for overall string similarity
        seq_ratio = SequenceMatcher(None, title.lower(), pdf_name.lower()).ratio()
        score += seq_ratio * 0.3

        if score > best_score:
            best_score = score
            best_match = pdf_file

    # Require minimum score threshold to avoid garbage matches
    if best_score >= 0.25:
        return best_match
    return None


# ---------------------------------------------------------------------------
# Helper: build exhibit→PDF map using keyword rules (résumé pattern)
# ---------------------------------------------------------------------------
def build_exhibit_pdf_map(evidence_dir: str, exhibit_index: dict) -> tuple:
    """
    Try to map exhibit numbers to PDF files in the evidence directory.
    Returns ({exhibit_num: pdf_path}, [unmapped_exhibit_nums]).

    Uses keyword-based rules derived from exhibit descriptions.
    Falls back to generic fuzzy matching.
    """
    pdf_files = [os.path.relpath(os.path.join(dp, f), evidence_dir) for dp, _, fns in os.walk(evidence_dir) for f in fns if f.lower().endswith(".pdf") and '_Forjado' not in dp]
    exhibit_map = {}
    unmapped = []
    used_pdfs = set()

    # First pass: generic fuzzy matching based on exhibit descriptions
    # Sort by exhibit number for deterministic ordering
    for num in sorted(exhibit_index.keys()):
        desc = exhibit_index[num]
        match = fuzzy_match_pdf(desc, pdf_files, used_pdfs, evidence_num=num)
        if match:
            exhibit_map[num] = os.path.join(evidence_dir, match)
            used_pdfs.add(match)
        else:
            unmapped.append(num)

    return exhibit_map, unmapped


# ---------------------------------------------------------------------------
# Helper: extract metadata from adjacent cell (cover letter pattern)
# ---------------------------------------------------------------------------
def extract_metadata_from_row(table, row_idx: int, thumbnail_cell_idx: int) -> dict:
    """
    Given a table row and the cell index containing [THUMBNAIL],
    look at the adjacent cell (metadata column) and extract:
    - evidence_num: int or None
    - title: str (the evidence title)
    - full_text: str (all text in the metadata cell)
    """
    row = table.rows[row_idx]
    cells = row.cells

    # The metadata cell is the OTHER cell in the row (not the thumbnail cell)
    meta_cell_idx = 1 if thumbnail_cell_idx == 0 else 0
    if meta_cell_idx >= len(cells):
        return {"evidence_num": None, "title": "", "full_text": ""}

    meta_cell = cells[meta_cell_idx]
    full_text = meta_cell.text.strip()

    # Extract evidence number from first line
    evidence_num = None
    title = ""
    lines = [l.strip() for l in full_text.split("\n") if l.strip()]

    if lines:
        first_line = lines[0]
        ev_match = EVIDENCE_NUM_RE.search(first_line)
        if ev_match:
            evidence_num = int(ev_match.group(1))
            # Title is everything after "Evidence XX. "
            title = EVIDENCE_NUM_RE.sub("", first_line).strip()
        else:
            title = first_line

    return {
        "evidence_num": evidence_num,
        "title": title,
        "full_text": full_text,
    }


# ---------------------------------------------------------------------------
# Helper: extract exhibit index from DOCX
# ---------------------------------------------------------------------------
def extract_exhibit_index(doc: Document) -> dict:
    """
    Find the EXHIBIT INDEX table in the DOCX and return {exhibit_num: description}.
    """
    exhibit_index = {}
    for table in doc.tables:
        if len(table.rows) < 3:
            continue
        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
        if "exhibit" in header_cells and "description" in header_cells:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    match = re.match(r"Exhibit\s+(\d+)", cells[0])
                    if match:
                        exhibit_index[int(match.group(1))] = cells[1]
            break
    return exhibit_index


# ---------------------------------------------------------------------------
# Helper: add border to inline image
# ---------------------------------------------------------------------------
def add_image_border(run):
    """Add a gray border around the image in a run using DrawingML properties."""
    drawing = run._element.find(qn("w:drawing"))
    if drawing is None:
        return

    inline = drawing.find(qn("wp:inline"))
    if inline is None:
        return

    graphic = inline.find(qn("a:graphic"))
    if graphic is None:
        return
    graphic_data = graphic.find(qn("a:graphicData"))
    if graphic_data is None:
        return
    pic = graphic_data.find(qn("pic:pic"))
    if pic is None:
        return
    sp_pr = pic.find(qn("pic:spPr"))
    if sp_pr is None:
        return

    r_val = int(BORDER_COLOR[1:3], 16)
    g_val = int(BORDER_COLOR[3:5], 16)
    b_val = int(BORDER_COLOR[5:7], 16)
    border_emu = int(BORDER_WIDTH_PT * 12700)

    ln_xml = (
        f'<a:ln xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" w="{border_emu}">'
        f'  <a:solidFill>'
        f'    <a:srgbClr val="{r_val:02X}{g_val:02X}{b_val:02X}"/>'
        f'  </a:solidFill>'
        f'</a:ln>'
    )
    ln_element = parse_xml(ln_xml)
    sp_pr.append(ln_element)


# ---------------------------------------------------------------------------
# Helper: load manual overrides from thumbnail_map.json
# ---------------------------------------------------------------------------
def load_override_map(map_path: str) -> dict:
    """
    Load thumbnail_map.json and return {evidence_num: pdf_path} for filled entries.
    """
    if not os.path.exists(map_path):
        return {}

    try:
        with open(map_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        print(f"  [WARN] Cannot load override map {map_path}: {e}")
        return {}

    overrides = {}
    # Support both formats: {"unmapped_exhibits": {"Exhibit 5": {"pdf_path": "..."}}}
    # and flat: {"5": "/path/to/file.pdf", "Evidence 12": "/path/to/file.pdf"}
    if "unmapped_exhibits" in data:
        for key, val in data["unmapped_exhibits"].items():
            num_match = re.search(r"(\d+)", key)
            if num_match and isinstance(val, dict) and val.get("pdf_path"):
                overrides[int(num_match.group(1))] = val["pdf_path"]
    else:
        for key, val in data.items():
            if key == "note":
                continue
            num_match = re.search(r"(\d+)", str(key))
            if num_match and val and isinstance(val, str) and os.path.exists(val):
                overrides[int(num_match.group(1))] = val

    return overrides


# ---------------------------------------------------------------------------
# Main: process DOCX
# ---------------------------------------------------------------------------
def process_docx(docx_path: str, evidence_dir: str, map_path: Optional[str] = None):
    """Main processing function."""
    print("=" * 70)
    print("INSERT THUMBNAILS (Universal)")
    print("=" * 70)
    print(f"DOCX: {docx_path}")
    print(f"Evidence: {evidence_dir}")
    if map_path:
        print(f"Override map: {map_path}")
    print()

    if not os.path.exists(docx_path):
        print(f"[ERROR] DOCX not found: {docx_path}")
        sys.exit(1)
    if not os.path.isdir(evidence_dir):
        print(f"[ERROR] Evidence directory not found: {evidence_dir}")
        sys.exit(1)

    doc = Document(docx_path)

    # Load PDF file list from evidence directory
    pdf_files = [os.path.relpath(os.path.join(dp, f), evidence_dir) for dp, _, fns in os.walk(evidence_dir) for f in fns if f.lower().endswith(".pdf") and '_Forjado' not in dp]
    print(f"[0/5] Found {len(pdf_files)} PDFs in evidence directory")

    # Step 1: Load manual override map
    print("[1/5] Loading override map...")
    override_map = {}
    # Auto-detect thumbnail_map.json next to DOCX
    auto_map = os.path.join(os.path.dirname(docx_path), "thumbnail_map.json")
    if map_path:
        override_map = load_override_map(map_path)
    elif os.path.exists(auto_map):
        print(f"  Auto-detected: {auto_map}")
        override_map = load_override_map(auto_map)

    if override_map:
        print(f"  Loaded {len(override_map)} manual override(s)")
    else:
        print("  No overrides loaded")

    # Step 2: Extract exhibit index from DOCX (for résumé pattern)
    print("[2/5] Extracting exhibit index from DOCX...")
    exhibit_index = extract_exhibit_index(doc)
    if exhibit_index:
        print(f"  Found {len(exhibit_index)} exhibits in index")
    else:
        print("  No exhibit index found (normal for cover letters)")

    # Step 3: Find ALL placeholders (both patterns)
    print("[3/5] Scanning for THUMBNAIL placeholders...")

    # Résumé placeholders: (table_idx, row_idx, cell_idx, para_idx, exhibit_nums)
    resume_placeholders = []
    # Cover letter placeholders: (table_idx, row_idx, cell_idx, para_idx, metadata)
    cover_placeholders = []

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    if not text:
                        continue

                    # Try résumé pattern first (more specific)
                    resume_match = RESUME_PLACEHOLDER_RE.search(text)
                    if resume_match:
                        start = int(resume_match.group(1))
                        end = int(resume_match.group(2)) if resume_match.group(2) else start
                        exhibit_nums = list(range(start, end + 1))
                        resume_placeholders.append((t_idx, r_idx, c_idx, p_idx, exhibit_nums))
                        print(f"  [RÉSUMÉ]  {text} (Table {t_idx})")
                        continue

                    # Try cover letter pattern
                    if COVER_PLACEHOLDER_RE.match(text):
                        metadata = extract_metadata_from_row(table, r_idx, c_idx)
                        cover_placeholders.append((t_idx, r_idx, c_idx, p_idx, metadata))
                        ev_label = f"Evidence {metadata['evidence_num']}" if metadata['evidence_num'] else "Unknown"
                        print(f"  [COVER]   {text} → {ev_label}. {metadata['title'][:50]}... (Table {t_idx})")
                        continue

    total_placeholders = len(resume_placeholders) + len(cover_placeholders)
    if total_placeholders == 0:
        print("  No placeholders found. Nothing to do.")
        return

    print(f"  Total: {total_placeholders} ({len(resume_placeholders)} résumé + {len(cover_placeholders)} cover)")

    # Step 4: Map evidence to PDFs
    print("[4/5] Mapping evidence to PDF files...")

    # === Résumé mapping (exhibit-based) ===
    exhibit_map = {}
    if resume_placeholders:
        all_exhibit_nums = set()
        for _, _, _, _, nums in resume_placeholders:
            all_exhibit_nums.update(nums)

        exhibit_map, unmapped_resume = build_exhibit_pdf_map(evidence_dir, exhibit_index)

        # Fallback: direct number matching (EVIDENCIA X - *.pdf → Exhibit X)
        pdf_files = [os.path.relpath(os.path.join(dp, f), evidence_dir) for dp, _, fns in os.walk(evidence_dir) for f in fns if f.lower().endswith(".pdf") and '_Forjado' not in dp]
        for num in all_exhibit_nums:
            if num not in exhibit_map:
                # Try patterns: "EVIDENCIA X", "EVIDÊNCIA X", "Evidence X", "Exhibit X", "Exhibit_X"
                for pdf in pdf_files:
                    patterns = [
                        rf"(?:EVIDENCIA|EVIDÊNCIA|Evidence|Exhibit)[_ -]*{num}\b",
                        rf"^{num}[_ -]",  # files starting with the number
                    ]
                    for pat in patterns:
                        if re.search(pat, pdf, re.IGNORECASE):
                            exhibit_map[num] = os.path.join(evidence_dir, pdf)
                            break
                    if num in exhibit_map:
                        break

        # Apply overrides
        for num, path in override_map.items():
            if os.path.exists(path):
                exhibit_map[num] = path

        for num in sorted(all_exhibit_nums):
            if num in exhibit_map:
                print(f"  Exhibit {num} → {os.path.basename(exhibit_map[num])}")
            else:
                print(f"  Exhibit {num} → [NOT MAPPED]")

    # === Cover letter mapping (metadata-based fuzzy match) ===
    cover_map = {}  # (t_idx, r_idx) → pdf_path
    used_pdfs = set(os.path.basename(p) for p in exhibit_map.values())
    unmapped_cover = []

    if cover_placeholders:
        for t_idx, r_idx, c_idx, p_idx, metadata in cover_placeholders:
            ev_num = metadata["evidence_num"]
            title = metadata["title"]
            full_text = metadata["full_text"]

            # Check override first
            if ev_num is not None and ev_num in override_map:
                path = override_map[ev_num]
                if os.path.exists(path):
                    cover_map[(t_idx, r_idx)] = path
                    used_pdfs.add(os.path.basename(path))
                    print(f"  Evidence {ev_num} → {os.path.basename(path)} [OVERRIDE]")
                    continue

            # Try fuzzy matching with title + full metadata text
            search_text = f"{title} {full_text}"
            match = fuzzy_match_pdf(search_text, pdf_files, used_pdfs, evidence_num=ev_num)
            if match:
                full_path = os.path.join(evidence_dir, match)
                cover_map[(t_idx, r_idx)] = full_path
                used_pdfs.add(match)
                ev_label = f"Evidence {ev_num}" if ev_num else "?"
                print(f"  {ev_label} → {match}")
            else:
                ev_label = f"Evidence {ev_num}" if ev_num else "?"
                print(f"  {ev_label} → [NOT MAPPED] ({title[:40]}...)")
                unmapped_cover.append({
                    "evidence_num": ev_num,
                    "title": title,
                    "table_idx": t_idx,
                    "row_idx": r_idx,
                })

    # Step 5: Extract thumbnails and insert
    print("[5/5] Extracting thumbnails and inserting into DOCX...")

    with tempfile.TemporaryDirectory() as temp_dir:
        thumbnail_cache = {}  # key → png_path

        # Pre-extract résumé thumbnails
        if resume_placeholders:
            all_exhibit_nums = set()
            for _, _, _, _, nums in resume_placeholders:
                all_exhibit_nums.update(nums)
            for num in sorted(all_exhibit_nums):
                if num in exhibit_map:
                    png = extract_thumbnail(exhibit_map[num], temp_dir, f"Exhibit {num}")
                    if png:
                        thumbnail_cache[f"exhibit_{num}"] = png

        # Pre-extract cover thumbnails
        for key, pdf_path in cover_map.items():
            label = f"Cover_T{key[0]}_R{key[1]}"
            png = extract_thumbnail(pdf_path, temp_dir, label)
            if png:
                thumbnail_cache[key] = png

        replaced = 0
        skipped = 0

        # --- Process résumé placeholders ---
        for t_idx, r_idx, c_idx, p_idx, exhibit_nums in resume_placeholders:
            table = doc.tables[t_idx]
            cell = table.rows[r_idx].cells[c_idx]
            para = cell.paragraphs[p_idx]

            available = [n for n in exhibit_nums if f"exhibit_{n}" in thumbnail_cache]

            if not available:
                print(f"  [SKIP] Table {t_idx}: No PDFs for Exhibits {exhibit_nums}")
                skipped += 1
                continue

            is_multi = len(available) > 1
            width = Inches(THUMBNAIL_WIDTH_COMPACT if is_multi else THUMBNAIL_WIDTH_INCHES)

            for run in para.runs:
                run.text = ""
            para.alignment = 1  # CENTER

            for i, num in enumerate(available):
                png_path = thumbnail_cache[f"exhibit_{num}"]
                run = para.add_run()
                run.add_picture(png_path, width=width)
                add_image_border(run)
                if i < len(available) - 1:
                    spacing_run = para.add_run()
                    spacing_run.add_break()

            replaced += 1
            label = f"Exhibit{'s' if len(exhibit_nums) > 1 else ''} {exhibit_nums[0]}"
            if len(exhibit_nums) > 1:
                label += f"–{exhibit_nums[-1]}"
            print(f"  [OK] Table {t_idx}: Inserted {len(available)} thumbnail(s) for {label}")

        # --- Process cover letter placeholders ---
        for t_idx, r_idx, c_idx, p_idx, metadata in cover_placeholders:
            key = (t_idx, r_idx)
            if key not in thumbnail_cache:
                ev_label = f"Evidence {metadata['evidence_num']}" if metadata['evidence_num'] else "Unknown"
                print(f"  [SKIP] Table {t_idx}: No PDF for {ev_label}")
                skipped += 1
                continue

            table = doc.tables[t_idx]
            cell = table.rows[r_idx].cells[c_idx]
            para = cell.paragraphs[p_idx]

            width = Inches(THUMBNAIL_WIDTH_INCHES)

            for run in para.runs:
                run.text = ""
            para.alignment = 1  # CENTER

            png_path = thumbnail_cache[key]
            run = para.add_run()
            run.add_picture(png_path, width=width)
            add_image_border(run)

            replaced += 1
            ev_label = f"Evidence {metadata['evidence_num']}" if metadata['evidence_num'] else "?"
            print(f"  [OK] Table {t_idx}: Inserted thumbnail for {ev_label}")

    # Save output
    base, ext = os.path.splitext(docx_path)
    output_path = f"{base}_with_thumbnails{ext}"
    doc.save(output_path)

    print()
    print("=" * 70)
    print("DONE!")
    print(f"  Replaced: {replaced}/{total_placeholders} placeholders")
    print(f"  Skipped:  {skipped}/{total_placeholders} (no PDF mapped)")
    print(f"  Output:   {output_path}")
    print("=" * 70)

    # Save unmapped evidence for manual resolution
    all_unmapped = []

    # Résumé unmapped
    if resume_placeholders:
        all_exhibit_nums = set()
        for _, _, _, _, nums in resume_placeholders:
            all_exhibit_nums.update(nums)
        for n in sorted(all_exhibit_nums):
            if n not in exhibit_map:
                all_unmapped.append({
                    "key": f"Exhibit {n}",
                    "num": n,
                    "description": exhibit_index.get(n, "Unknown"),
                })

    # Cover letter unmapped
    for item in unmapped_cover:
        all_unmapped.append({
            "key": f"Evidence {item['evidence_num']}" if item['evidence_num'] else f"Table {item['table_idx']} Row {item['row_idx']}",
            "num": item["evidence_num"],
            "description": item["title"],
        })

    if all_unmapped:
        out_map_path = os.path.join(os.path.dirname(docx_path), "thumbnail_map.json")
        map_data = {
            "note": "Fill in 'pdf_path' for each unmapped evidence, then re-run the script with --map.",
            "unmapped_exhibits": {}
        }
        for item in all_unmapped:
            map_data["unmapped_exhibits"][item["key"]] = {
                "description": item["description"],
                "pdf_path": ""
            }

        with open(out_map_path, "w", encoding="utf-8") as f:
            json.dump(map_data, f, indent=2, ensure_ascii=False)

        print(f"\n  [INFO] {len(all_unmapped)} unmapped evidence(s) saved to:")
        print(f"         {out_map_path}")
        print(f"         Fill in the paths and re-run.")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Insert PDF thumbnails into DOCX files. "
                    "Supports both [THUMBNAIL — Exhibit X] (résumé) and [THUMBNAIL] (cover letter) patterns."
    )
    parser.add_argument(
        "--docx", required=True,
        help="Path to DOCX file with thumbnail placeholders"
    )
    parser.add_argument(
        "--evidence", required=True,
        help="Path to folder containing evidence PDFs"
    )
    parser.add_argument(
        "--map", dest="map_file", default=None,
        help="Path to thumbnail_map.json with manual overrides (optional)"
    )

    # Also support positional args for backward compatibility
    # python insert_thumbnails.py <docx> <evidence>
    if len(sys.argv) >= 2 and not sys.argv[1].startswith("--"):
        # Legacy positional mode
        if len(sys.argv) < 3:
            parser.print_help()
            sys.exit(1)
        docx_path = sys.argv[1]
        evidence_dir = sys.argv[2]
        map_file = sys.argv[3] if len(sys.argv) > 3 else None
        process_docx(docx_path, evidence_dir, map_file)
    else:
        args = parser.parse_args()
        process_docx(args.docx, args.evidence, args.map_file)


if __name__ == "__main__":
    main()
