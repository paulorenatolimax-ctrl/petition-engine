#!/usr/bin/env python3
"""
QUALITY GATE — Résumé EB-1A / EB-2 NIW  (mecânico, rápido, definitivo)
=======================================================================
Roda APÓS cada build. Retorna exit code 1 se reprovado.

USO:
    python3 quality_gate_resume.py <resume.docx> [eb1a|eb2niw]

CHECKS (em ordem de severidade):
  S0  GRAVÍSSIMO — bloqueia entrega imediatamente
  S1  GRAVE      — bloqueia entrega
  S2  MODERADO   — corrigir antes de entregar
  S3  MENOR      — documentar

Verificações cobertas:
  ✔ fontes        — 100% Garamond (ou Arial nos benchmarks antigos)
  ✔ palavras      — contagem mínima de palavras no corpo
  ✔ Exhibit       — palavra "Exhibit" proibida no corpo
  ✔ placeholders  — texto do tipo [THUMBNAIL], [TODO], [VERIFICAR], etc.
  ✔ termos proibidos — R$, jargão jurídico, linguagem de marketing
  ✔ SOC no corpo  — código SOC (XX-XXXX) NÃO deve aparecer fora do header
  ✔ anti-Cristine — nenhum nome de outro cliente no documento
  ✔ cartas citadas — seção de cartas de recomendação existe e é não-vazia
  ✔ Principais Responsabilidades — sub-header obrigatório por empresa
  ✔ Principais Resultados        — sub-header obrigatório por empresa
  ✔ footer        — footer com "Page" presente
  ✔ thumbnails    — imagens presentes; zero placeholders [THUMBNAIL]
  ✔ inferências   — números grandes sem fonte parentética (heurística)
"""

import sys
import os
import re
from docx import Document
from docx.oxml.ns import qn

# ═══════════════════════════════════════════════════════════════════════
# CONFIGURAÇÃO
# ═══════════════════════════════════════════════════════════════════════

# Fontes permitidas (case-insensitive)
ALLOWED_FONTS = {"garamond", "arial"}

# Paleta de cores permitidas (hex sem #)
ALLOWED_COLORS = {
    "2D3E50", "2D4F5F",        # Navy (variações)
    "3498A2",                   # Teal
    "FFFFFF",                   # White
    "000000",                   # Black
    "333333",                   # Dark gray
    "666666",                   # Med gray
    "F5F5F5",                   # Light gray
    "CCCCCC",                   # Border gray
    "FAFAFA",                   # Alt row
    "AUTO", "auto",
}

# Mínimo de palavras no corpo do documento
MIN_WORD_COUNT = 600

# Termos proibidos no corpo (case-insensitive)
FORBIDDEN_TERMS = [
    "R$",
    "pursuant to",
    "8 CFR",
    "Kazarian",
    "preponderance of evidence",
    "the beneficiary meets",
    "o beneficiário atende",
    "this criterion",
    "incrível",
    "revolucionário",
    "transformador",
    "100% do crescimento",
    "100% atribuível",
    "único no Brasil",
    "o mais importante",
    "nenhum outro profissional",
]

# Termos proibidos APENAS em EB-2 NIW
# (são a linguagem errada para este tipo de visto)
FORBIDDEN_TERMS_EB2NIW_ONLY = [
    "extraordinary ability",  # standard do EB-1A, não EB-2 NIW
]

# Termos proibidos APENAS em EB-1A
FORBIDDEN_TERMS_EB1A_ONLY = [
    "Dhanasar",  # framework do EB-2 NIW — não pertence ao EB-1A
]

# Nomes de clientes que NUNCA devem aparecer em outro documento
# Adicione novos clientes aqui conforme necessário
FORBIDDEN_CLIENT_NAMES = [
    "Carlos Avelino",
    "Bruno Cipriano",
    "Cristine",
    "PROEX",
    "Renato Silveira",   # só proibido em documentos de outros clientes
    "Julio Caleiro",
    "Júlio Caleiro",
]

# Placeholders proibidos
FORBIDDEN_PLACEHOLDERS = [
    "[THUMBNAIL]",
    "[TODO]",
    "[A PREENCHER]",
    "[VERIFICAR]",
    "[A CRUZAR COM CL]",
    "[PLACEHOLDER]",
    "[IMAGEM]",
    "[EVIDÊNCIA A COLETAR]",
    "[INSERT",
    "TBD",
    "PLACEHOLDER",
]

# Padrão SOC/O*Net  (XX-XXXX ou XX-XXXX.XX)
SOC_PATTERN = re.compile(r'\b\d{2}-\d{4}(\.\d{2})?\b')

# Padrão de inferência: número grande (>999) sem fonte parentética nas próximas 300 chars
# Detecta casos como "2.5 million downloads" sem "(Source: ...)" próximo
INFERENCE_PATTERN = re.compile(
    r'\b(\d{1,3}(?:[.,]\d{3})*(?:\.\d+)?)\s*(?:million|billion|thousand|milhões?|bilhões?|mil)\b',
    re.IGNORECASE
)
SOURCE_PATTERN = re.compile(
    r'\((?:Source|Fonte|BLS|Gartner|Statista|McKinsey|IDC|Forrester|IBGE|Valor|Forbes|'
    r'acesso em|disponível em|available at|cited in|apud|cf\.|see )',
    re.IGNORECASE
)

# Sub-headers obrigatórios na seção de Experiência Profissional
REQUIRED_EXPERIENCE_HEADERS = [
    "Principais Responsabilidades",
    "Principais Resultados",
]

# Nomes de seções de cartas de recomendação
LETTERS_SECTION_KEYWORDS = [
    "CARTAS DE RECOMENDAÇÃO",
    "LETTERS OF RECOMMENDATION",
    "RECOMMENDATION LETTERS",
    "CARTAS DE APOIO",
]


# ═══════════════════════════════════════════════════════════════════════
# UTILITÁRIOS
# ═══════════════════════════════════════════════════════════════════════

def load_doc(path):
    """Carrega .docx e extrai texto completo (parágrafos + tabelas + section headers)."""
    doc = Document(path)

    para_texts = [p.text for p in doc.paragraphs]
    table_texts = []
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                table_texts.append(cell.text)

    # Incluir section headers (onde fica o nome/SOC/visa no DNA Visual V4)
    header_texts = []
    for section in doc.sections:
        hdr = section.header
        for p in hdr.paragraphs:
            header_texts.append(p.text)
        for tbl in hdr.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    header_texts.append(cell.text)

    body_text = "\n".join(para_texts + table_texts + header_texts)
    return doc, body_text


def get_section_headers(doc):
    """Retorna lista de textos de células de tabela que são headers de seção
    (células curtas, todo-maiúsculas, que representam seções do documento como
    'EXPERIÊNCIA PROFISSIONAL', 'PROPOSED ENDEAVORS', etc.)."""
    headers = []
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                # Header de seção: curto (< 60 chars), todo caps ou quase todo caps
                if txt and len(txt) < 60 and txt == txt.upper() and len(txt) > 4:
                    headers.append(txt)
    return headers


def get_header_text(doc):
    """Extrai texto dos headers de todas as seções."""
    texts = []
    for section in doc.sections:
        hdr = section.header
        for p in hdr.paragraphs:
            texts.append(p.text)
        for tbl in hdr.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    texts.append(cell.text)
    return "\n".join(texts)


def get_footer_text(doc):
    """Extrai texto dos footers de todas as seções."""
    texts = []
    for section in doc.sections:
        ftr = section.footer
        for p in ftr.paragraphs:
            texts.append(p.text)
        for tbl in ftr.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    texts.append(cell.text)
    return "\n".join(texts)


def count_images(doc):
    """Conta imagens inline no documento."""
    return len(doc.inline_shapes)


def count_runs_with_font(doc):
    """Retorna set de fontes encontradas e contagem de runs por fonte."""
    fonts_found = {}
    total_runs = 0

    def scan_paragraphs(paragraphs):
        nonlocal total_runs
        for para in paragraphs:
            for run in para.runs:
                total_runs += 1
                fname = run.font.name or "None"
                fname_key = fname.lower()
                fonts_found[fname_key] = fonts_found.get(fname_key, 0) + 1

    scan_paragraphs(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                scan_paragraphs(cell.paragraphs)

    return fonts_found, total_runs


# ═══════════════════════════════════════════════════════════════════════
# VERIFICAÇÕES
# ═══════════════════════════════════════════════════════════════════════

def check_fonts(doc):
    """[S0/S1] Verifica que todas as fontes são Garamond (ou Arial para benchmarks antigos)."""
    issues = []
    fonts_found, total_runs = count_runs_with_font(doc)

    bad_fonts = {f: c for f, c in fonts_found.items()
                 if f not in ALLOWED_FONTS and f != "none"}

    if bad_fonts:
        for fname, count in sorted(bad_fonts.items(), key=lambda x: -x[1]):
            severity = "S0" if "calibri" in fname or "times" in fname else "S1"
            issues.append((severity,
                f"Fonte proibida '{fname}': {count} runs "
                f"({count/max(total_runs,1)*100:.0f}% do documento)"))

    return issues, total_runs, fonts_found


def check_word_count(body_text):
    """[S1] Verifica contagem mínima de palavras."""
    issues = []
    words = len(body_text.split())
    if words < MIN_WORD_COUNT:
        issues.append(("S1",
            f"Documento muito curto: {words} palavras (mínimo: {MIN_WORD_COUNT}). "
            f"Résumé incompleto."))
    return issues, words


def check_exhibit(body_text):
    """[S1] Verifica presença da palavra 'Exhibit' (pertence à Cover Letter, não ao Résumé)."""
    issues = []
    # Procura "Exhibit" como palavra standalone (exceto dentro de palavras como 'exhibition')
    matches = re.findall(r'\bExhibit\b', body_text, re.IGNORECASE)
    if matches:
        issues.append(("S1",
            f"Palavra 'Exhibit' encontrada {len(matches)}x — "
            f"use 'Evidence' ou 'Evidência'. 'Exhibit' é terminologia da Cover Letter."))
    return issues


def check_placeholders(body_text):
    """[S0/S1] Verifica placeholders não preenchidos."""
    issues = []
    found = []

    for ph in FORBIDDEN_PLACEHOLDERS:
        count = body_text.count(ph)
        if count > 0:
            found.append((ph, count))

    # Detectar padrão genérico [UPPERCASE WORDS]
    generic_phs = re.findall(r'\[[A-ZÁÉÍÓÚÂÊÎÔÛÃÕ][A-ZÁÉÍÓÚÂÊÎÔÛÃÕ\s]{3,}\]', body_text)
    for gph in set(generic_phs):
        if gph not in [f[0] for f in found]:
            found.append((gph, body_text.count(gph)))

    if found:
        severity = "S0" if any("[THUMBNAIL]" in f[0] for f in found) else "S1"
        for ph, count in found:
            sev = "S0" if ph == "[THUMBNAIL]" else "S1"
            issues.append((sev, f"Placeholder não preenchido: {ph} ({count}x)"))

    return issues


def check_forbidden_terms(body_text, doc_type="eb2niw"):
    """[S1/S2] Verifica termos proibidos (globais + específicos por tipo)."""
    issues = []

    # Termos globais
    for term in FORBIDDEN_TERMS:
        if term.lower() in body_text.lower():
            severity = "S1" if term in ("R$", "8 CFR", "Kazarian") else "S2"
            issues.append((severity, f"Termo proibido encontrado: '{term}'"))

    # Termos específicos por tipo
    extra = FORBIDDEN_TERMS_EB2NIW_ONLY if doc_type == "eb2niw" else FORBIDDEN_TERMS_EB1A_ONLY
    for term in extra:
        if term.lower() in body_text.lower():
            issues.append(("S2", f"Termo proibido para {doc_type.upper()}: '{term}'"))

    return issues


def check_soc_in_body(doc, body_text, header_text):
    """[S1] Verifica se código SOC (XX-XXXX) aparece no CORPO (deveria ser só no header)."""
    issues = []

    # SOC no header é esperado — remover do body_text para verificar
    # Pegar só texto dos parágrafos do body (não header, não footer)
    para_text = "\n".join(p.text for p in doc.paragraphs)

    matches = SOC_PATTERN.findall(para_text)
    if matches:
        # Verificar se não é dentro de uma seção de Proposed Endeavors (onde BLS é permitido)
        # Uma heurística: se há muitas ocorrências, provavelmente estão corretas
        # Se há apenas 1-2 e não no contexto de BLS/O*Net, pode ser erro
        unique = list(set(matches))
        # Contar contextos de BLS — nesses contextos é esperado
        bls_context = len(re.findall(
            r'(?:BLS|O\*Net|SOC|Occupation Code)[^\n]{0,50}' + SOC_PATTERN.pattern,
            para_text, re.IGNORECASE
        ))
        suspicious = len(unique) - bls_context
        if suspicious > 0:
            issues.append(("S2",
                f"Código SOC/O*Net ({', '.join(unique)}) aparece no corpo do documento. "
                f"SOC vai no header e nas Proposed Endeavors (com contexto BLS). "
                f"Verificar se não está solto em texto narrativo."))
    return issues


def check_anti_client_names(body_text, current_client=None):
    """[S1] Verifica que nenhum nome de outro cliente aparece no documento."""
    issues = []
    for name in FORBIDDEN_CLIENT_NAMES:
        if current_client and name.lower() in current_client.lower():
            continue  # é o próprio cliente
        if name.lower() in body_text.lower():
            issues.append(("S1",
                f"Nome de outro cliente encontrado: '{name}'. "
                f"Benchmarks são para estrutura, nunca para conteúdo."))
    return issues


def check_letters_section(body_text):
    """[S2] Verifica que a seção de cartas de recomendação existe e tem conteúdo."""
    issues = []
    upper = body_text.upper()
    found = any(kw in upper for kw in LETTERS_SECTION_KEYWORDS)

    if not found:
        issues.append(("S2",
            "Seção de Cartas de Recomendação NÃO encontrada. "
            "Verifique se foi incluída (pode estar em outro arquivo do merge)."))
        return issues

    # Verificar se a seção tem pelo menos um nome/pessoa citada
    # Heurística: checar se há pelo menos 2 linhas após o header da seção
    for kw in LETTERS_SECTION_KEYWORDS:
        idx = upper.find(kw)
        if idx != -1:
            after = body_text[idx:idx+500]
            # Contar linhas não-vazias após o header
            lines = [l.strip() for l in after.split('\n') if l.strip() and kw not in l.upper()]
            if len(lines) < 2:
                issues.append(("S2",
                    f"Seção '{kw}' encontrada mas parece vazia. "
                    f"Verificar se as cartas estão listadas na tabela."))
            break

    return issues


def check_experience_headers(body_text, doc_type):
    """[S1] Verifica presença de 'Principais Responsabilidades' e 'Principais Resultados'
    na seção de Experiência Profissional (EB-2 NIW)."""
    issues = []
    if doc_type != "eb2niw":
        return issues

    # Verificar se há pelo menos 1 seção de Experiência Profissional
    exp_keywords = ["EXPERIÊNCIA PROFISSIONAL", "PROFESSIONAL EXPERIENCE", "CAREER EXPERIENCE"]
    upper = body_text.upper()
    has_experience = any(kw in upper for kw in exp_keywords)

    if not has_experience:
        # Sem seção de experiência detectada — reportar mas não duplo-penalizar
        return issues

    for header in REQUIRED_EXPERIENCE_HEADERS:
        count = len(re.findall(re.escape(header), body_text, re.IGNORECASE))
        if count == 0:
            issues.append(("S1",
                f"Sub-header obrigatório ausente: '{header}' "
                f"— deve aparecer em CADA bloco de empresa na Experiência Profissional."))
        elif count == 1:
            issues.append(("S2",
                f"'{header}' aparece apenas 1x. "
                f"Esperado em múltiplos blocos de empresa. Verificar completude."))

    return issues, body_text.count("Principais Responsabilidades"), body_text.count("Principais Resultados")


def check_footer(doc):
    """[S1] Verifica que o footer existe e contém 'Page'."""
    issues = []
    footer_text = get_footer_text(doc)

    if not footer_text.strip():
        issues.append(("S1", "FOOTER AUSENTE — nenhum conteúdo nos footers do documento."))
    elif "page" not in footer_text.lower() and "página" not in footer_text.lower():
        issues.append(("S2",
            f"Footer existe mas sem 'Page X of Y'. "
            f"Conteúdo encontrado: '{footer_text[:80].strip()}'"))
    return issues, footer_text


def check_thumbnails(doc, body_text):
    """[S0/S1/S2] Verifica imagens e placeholders de thumbnail."""
    issues = []
    img_count = count_images(doc)

    # Placeholders exatos [THUMBNAIL] não substituídos → S0 obrigatório
    thumb_phs = body_text.count("[THUMBNAIL]")
    if thumb_phs > 0:
        issues.append(("S0",
            f"{thumb_phs} placeholder(s) [THUMBNAIL] não substituídos. "
            f"REBUILD OBRIGATÓRIO."))

    # Sem imagens: distinguir entre "build falhou" e "thumbnails pendentes"
    if img_count == 0:
        # Verificar se existem placeholders descritivos tipo "[THUMBNAIL — Evidência X]"
        has_thumb_placeholders = "[THUMBNAIL" in body_text
        if has_thumb_placeholders:
            # Estrutura de evidência existe, só faltam os arquivos de imagem
            issues.append(("S2",
                "ZERO imagens no documento. Estrutura de Evidence blocks presente "
                "(placeholders [THUMBNAIL — Evidência X] encontrados), mas as imagens "
                "não foram inseridas ainda. Rodar script de inserção de thumbnails."))
        else:
            # Nenhuma imagem E nenhum placeholder → build falhou ou não tem evidence blocks
            issues.append(("S1",
                "ZERO imagens no documento e sem placeholders de thumbnail. "
                "Evidence blocks devem ter thumbnails — ou o build falhou, "
                "ou as imagens não foram inseridas."))

    return issues, img_count


def check_inferences(doc):
    """[S2] Detecta números grandes sem fonte parentética próxima (inferências sem suporte)."""
    issues = []
    suspicious = []

    all_paras = list(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    for para in all_paras:
        text = para.text
        if not text.strip():
            continue

        matches = list(INFERENCE_PATTERN.finditer(text))
        for m in matches:
            # Verificar se há uma fonte nos próximos 300 chars do parágrafo
            context = text[m.start():min(m.end()+300, len(text))]
            has_source = bool(SOURCE_PATTERN.search(context))

            if not has_source:
                suspicious.append(f"{m.group()} — \"{text[:80].strip()}...\"")

    if len(suspicious) > 3:
        issues.append(("S2",
            f"{len(suspicious)} estatísticas grandes sem fonte parentética detectadas. "
            f"Verifique se cada número tem '(Source: ...)' ou '(Fonte: ...)' próximo."))
        for s in suspicious[:5]:
            issues.append(("S3", f"  Sem fonte: {s}"))

    return issues, len(suspicious)


def check_eb2niw_specific(body_text, doc=None):
    """[S0/S1] Verificações específicas para EB-2 NIW."""
    issues = []
    upper = body_text.upper()

    # REGRA r122: résumé NUNCA tem Proposed Endeavors como SEÇÃO FORMAL
    # (verificar presença de header de seção, não apenas menção no texto)
    if doc is not None:
        section_headers = get_section_headers(doc)
        has_proposed_section = any(
            "PROPOSED ENDEAVOR" in h.upper() or "PROJETO EB-2 NIW" in h.upper()
            for h in section_headers
        )
        if has_proposed_section:
            issues.append(("S1",
                "Seção PROPOSED ENDEAVORS encontrada no résumé. "
                "REGRA r122: Proposed Endeavors pertence à Cover Letter / Anteprojeto, NUNCA ao résumé. "
                "Remover esta seção do résumé."))

    # Dhanasar pertence à Cover Letter, não ao résumé
    if "DHANASAR" in upper:
        issues.append(("S1",
            "Referência ao framework Dhanasar encontrada no résumé. "
            "Dhanasar pertence à Cover Letter / argumentação legal, nunca ao résumé. "
            "Remover esta referência."))

    # Código SOC deve estar no header (ex: 11-9041) — obrigatório no EB-2 NIW
    if not SOC_PATTERN.search(body_text):
        issues.append(("S1",
            "Código SOC/BLS (formato XX-XXXX) ausente. "
            "O header do résumé EB-2 NIW deve incluir o código SOC da ocupação "
            "(ex: 'RÉSUMÉ  |  EB-2 NIW  |  11-9041')."))

    # NÃO deve ter seções C1-C10 (isso é EB-1A)
    criterion_pattern = re.compile(r'\bCRITERION\s+[1-9]\b', re.IGNORECASE)
    if criterion_pattern.search(body_text):
        issues.append(("S1",
            "Seções 'CRITERION X' encontradas — isso é estrutura EB-1A, não EB-2 NIW. "
            "EB-2 NIW organiza por TEMA, não por critério."))

    # Header deve mencionar EB-2 NIW
    if "EB-1A" in upper and "EB-2 NIW" not in upper:
        issues.append(("S0",
            "Header menciona 'EB-1A' mas documento é EB-2 NIW. Corrigir o header."))

    return issues


def check_eb1a_specific(body_text):
    """[S1] Verificações específicas para EB-1A."""
    issues = []
    upper = body_text.upper()

    # Deve ter pelo menos um Criterion
    if "CRITERION" not in upper:
        issues.append(("S1",
            "Nenhuma seção por Critério (CRITERION) encontrada. "
            "EB-1A deve ter pelo menos 1 seção de critério batido."))

    # Deve ter Executive Summary / Síntese
    if "EXECUTIVE SUMMARY" not in upper and "SÍNTESE PROFISSIONAL" not in upper:
        issues.append(("S1", "Seção EXECUTIVE SUMMARY / SÍNTESE PROFISSIONAL ausente."))

    # NÃO deve ter seções de Proposed Endeavors em nenhum résumé (r122)
    if "PROPOSED ENDEAVOR" in upper:
        issues.append(("S1",
            "Seção 'PROPOSED ENDEAVORS' encontrada em documento EB-1A. "
            "REGRA r122: Proposed Endeavors pertence à Cover Letter, nunca ao résumé."))

    return issues


# ═══════════════════════════════════════════════════════════════════════
# EXECUTOR PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════

def run_quality_gate(docx_path, doc_type="eb2niw", current_client=None):
    """
    Executa todas as verificações e retorna (s0, s1, s2, s3) contagens.
    Imprime relatório completo no stdout.
    """
    print("=" * 72)
    print(f"  QUALITY GATE — {os.path.basename(docx_path)}")
    print(f"  Tipo: {doc_type.upper()}")
    print(f"  Tamanho: {os.path.getsize(docx_path):,} bytes")
    print("=" * 72)

    doc, body_text = load_doc(docx_path)
    header_text = get_header_text(doc)

    all_issues = []

    step = 0
    total_steps = 13

    def step_header(name):
        nonlocal step
        step += 1
        print(f"\n[{step:02d}/{total_steps}] {name}...")

    # ── 1. FONTES ──
    step_header("Fontes")
    font_issues, total_runs, fonts_found = check_fonts(doc)
    all_issues.extend(font_issues)
    dominant = max(fonts_found, key=fonts_found.get) if fonts_found else "n/a"
    print(f"       {total_runs} runs | Fonte dominante: '{dominant}' "
          f"({fonts_found.get(dominant,0)/max(total_runs,1)*100:.0f}%)")

    # ── 2. CONTAGEM DE PALAVRAS ──
    step_header("Contagem de palavras")
    wc_issues, word_count = check_word_count(body_text)
    all_issues.extend(wc_issues)
    print(f"       {word_count:,} palavras")

    # ── 3. EXHIBIT ──
    step_header("Palavra 'Exhibit'")
    all_issues.extend(check_exhibit(body_text))

    # ── 4. PLACEHOLDERS ──
    step_header("Placeholders")
    all_issues.extend(check_placeholders(body_text))

    # ── 5. TERMOS PROIBIDOS ──
    step_header("Termos proibidos")
    all_issues.extend(check_forbidden_terms(body_text, doc_type))

    # ── 6. SOC NO CORPO ──
    step_header("SOC no corpo")
    all_issues.extend(check_soc_in_body(doc, body_text, header_text))

    # ── 7. ANTI-CRISTINE (outros clientes) ──
    step_header("Anti-Cristine (nomes proibidos)")
    all_issues.extend(check_anti_client_names(body_text, current_client))

    # ── 8. CARTAS CITADAS ──
    step_header("Cartas de recomendação")
    all_issues.extend(check_letters_section(body_text))

    # ── 9. PRINCIPAIS RESPONSABILIDADES / RESULTADOS ──
    step_header("Principais Responsabilidades + Resultados")
    exp_result = check_experience_headers(body_text, doc_type)
    if isinstance(exp_result, tuple):
        exp_issues, resp_count, res_count = exp_result
        all_issues.extend(exp_issues)
        print(f"       Responsabilidades: {resp_count}x | Resultados: {res_count}x")
    else:
        all_issues.extend(exp_result)

    # ── 10. FOOTER ──
    step_header("Footer")
    footer_issues, footer_text = check_footer(doc)
    all_issues.extend(footer_issues)
    if footer_text.strip():
        print(f"       Footer: '{footer_text.strip()[:60]}'")

    # ── 11. THUMBNAILS ──
    step_header("Thumbnails")
    thumb_issues, img_count = check_thumbnails(doc, body_text)
    all_issues.extend(thumb_issues)
    print(f"       {img_count} imagem(ns) no documento")

    # ── 12. INFERÊNCIAS SEM FONTE ──
    step_header("Inferências sem fonte")
    inf_issues, inf_count = check_inferences(doc)
    all_issues.extend(inf_issues)
    print(f"       {inf_count} estatísticas grandes detectadas")

    # ── 13. ESPECÍFICO DO TIPO ──
    step_header(f"Regras específicas {doc_type.upper()}")
    if doc_type == "eb2niw":
        all_issues.extend(check_eb2niw_specific(body_text, doc=doc))
    elif doc_type == "eb1a":
        all_issues.extend(check_eb1a_specific(body_text))

    # ═══════════════════════════════════════════════════════════════════
    # RELATÓRIO FINAL
    # ═══════════════════════════════════════════════════════════════════

    s0 = [(s, m) for s, m in all_issues if s == "S0"]
    s1 = [(s, m) for s, m in all_issues if s == "S1"]
    s2 = [(s, m) for s, m in all_issues if s == "S2"]
    s3 = [(s, m) for s, m in all_issues if s == "S3"]

    print("\n" + "=" * 72)
    print("  RELATÓRIO")
    print("=" * 72)
    print(f"  Palavras:  {word_count:,}")
    print(f"  Imagens:   {img_count}")
    print(f"  Tabelas:   {len(doc.tables)}")
    print(f"  Parágrafos:{len(doc.paragraphs)}")
    print()
    print(f"  S0 GRAVÍSSIMO: {len(s0)}")
    print(f"  S1 GRAVE:      {len(s1)}")
    print(f"  S2 MODERADO:   {len(s2)}")
    print(f"  S3 MENOR:      {len(s3)}")

    if s0:
        print("\n  !!!  S0 — GRAVÍSSIMO — REBUILD OBRIGATÓRIO  !!!")
        for _, msg in s0:
            print(f"  [S0] {msg}")

    if s1:
        print("\n  !!  S1 — GRAVE — BLOQUEIA ENTREGA  !!")
        for _, msg in s1:
            print(f"  [S1] {msg}")

    if s2:
        print("\n  !  S2 — MODERADO — CORRIGIR ANTES DE ENTREGAR  !")
        for _, msg in s2:
            print(f"  [S2] {msg}")

    if s3:
        print("\n  S3 — MENOR — DOCUMENTAR")
        for _, msg in s3:
            print(f"  [S3] {msg}")

    bloqueantes = len(s0) + len(s1)
    print("\n" + "=" * 72)
    if bloqueantes > 0:
        print(f"  VEREDICTO: ✗ REPROVADO — {bloqueantes} erro(s) S0/S1 bloqueiam entrega")
        print(f"  Corrija os erros acima e re-execute este script antes de entregar.")
    elif s2:
        print(f"  VEREDICTO: ⚠ APROVADO COM RESSALVAS — {len(s2)} moderado(s)")
        print(f"  Corrija S2 se possível antes de entregar ao Paulo.")
    else:
        print("  VEREDICTO: ✓ APROVADO — Documento pronto para entrega")
    print("=" * 72)

    return len(s0), len(s1), len(s2), len(s3)


# ═══════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        print("\nUso: python3 quality_gate_resume.py <resume.docx> [eb1a|eb2niw] [nome_cliente]")
        sys.exit(1)

    docx_path = sys.argv[1]
    doc_type = sys.argv[2].lower() if len(sys.argv) > 2 else "eb2niw"
    current_client = sys.argv[3] if len(sys.argv) > 3 else None

    if not os.path.exists(docx_path):
        print(f"ERRO: Arquivo não encontrado: {docx_path}")
        sys.exit(2)

    if doc_type not in ("eb1a", "eb2niw"):
        print(f"ERRO: Tipo inválido '{doc_type}'. Use 'eb1a' ou 'eb2niw'.")
        sys.exit(2)

    s0, s1, s2, s3 = run_quality_gate(docx_path, doc_type, current_client)

    # Exit code 1 = REPROVADO (tem S0 ou S1)
    # Exit code 0 = APROVADO (só S2/S3 ou limpo)
    sys.exit(1 if (s0 + s1) > 0 else 0)
