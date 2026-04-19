#!/usr/bin/env python3
"""
Generate VALUATION_MODEL_CONSOLIDADO.docx from markdown content
using the Cristine Style V2 design system.
"""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─── COLORS ───────────────────────────────────────────────────────────
NAVY_DARK   = RGBColor(0x1B, 0x2A, 0x3D)
WARM_BEIGE  = RGBColor(0xED, 0xE8, 0xDF)
BRONZE      = RGBColor(0x8B, 0x73, 0x55)
DEEP_BROWN  = RGBColor(0x5A, 0x55, 0x50)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)

NAVY_HEX    = "1B2A3D"
BEIGE_HEX   = "EDE8DF"
BRONZE_HEX  = "8B7355"
GRAY_BORDER = "CCCCCC"

# ─── INPUT/OUTPUT ─────────────────────────────────────────────────────
BASE_DIR = "/Users/paulo1844/Documents/AIOS_Petition Engine/Audit/_Valuation"
INPUT_MD = os.path.join(BASE_DIR, "VALUATION_MODEL_CONSOLIDADO.md")
OUTPUT_DOCX = os.path.join(BASE_DIR, "VALUATION_MODEL_CONSOLIDADO.docx")


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC", size="4"):
    """Set thin borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def add_bronze_divider(doc, thickness_pt=6):
    """Add a Bronze horizontal line as a paragraph with bottom border."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{thickness_pt * 2}" w:space="1" w:color="{BRONZE_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_paragraph_with_format(doc, text, size=11, color=BLACK, bold=False, italic=False,
                               alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6),
                               space_before=Pt(0), font_name="Calibri"):
    """Add a paragraph with specific formatting."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    # Set Calibri for East Asian font too
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rPr.append(rFonts)
    return p


def format_run(run, size=11, color=BLACK, bold=False, italic=False, font_name="Calibri"):
    """Apply formatting to a run."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_section_heading(doc, text, level=1):
    """Add a section heading with Bronze bottom border."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)

    if level == 1:
        display_text = text.upper()
        run = p.add_run(display_text)
        format_run(run, size=14, color=NAVY_DARK, bold=True)
        # Bronze bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="{BRONZE_HEX}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    elif level == 2:
        run = p.add_run(text)
        format_run(run, size=12, color=NAVY_DARK, bold=True)
        pf.space_before = Pt(12)
    elif level == 3:
        run = p.add_run(text)
        format_run(run, size=11, color=NAVY_DARK, bold=True)
        pf.space_before = Pt(10)
    return p


def apply_rich_text(paragraph, text, base_size=11, base_color=BLACK):
    """Parse markdown-like bold/italic in text and add runs accordingly."""
    # Pattern to find **bold** and *italic* and combined ***bolditalic***
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            inner = part[3:-3]
            run = paragraph.add_run(inner)
            format_run(run, size=base_size, color=base_color, bold=True, italic=True)
        elif part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            format_run(run, size=base_size, color=base_color, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            format_run(run, size=base_size, color=base_color, italic=True)
        else:
            if part:
                run = paragraph.add_run(part)
                format_run(run, size=base_size, color=base_color)


def apply_rich_text_to_cell(cell, text, size=9, color=BLACK, bold=False):
    """Apply rich text formatting inside a table cell."""
    # Clear default paragraph
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)

    # Parse bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = p.add_run(inner)
            format_run(run, size=size, color=color, bold=True)
        elif part.startswith('~') and part.endswith('~'):
            inner = part[1:-1]
            run = p.add_run(inner)
            format_run(run, size=size, color=color)
        else:
            if part:
                run = p.add_run(part)
                format_run(run, size=size, color=color, bold=bold)


def add_styled_table(doc, headers, rows):
    """Add a table with Cristine V2 styling."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, NAVY_HEX)
        set_cell_borders(cell, color=GRAY_BORDER)
        # Set text
        for p in cell.paragraphs:
            p.clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(3)
        pf.space_after = Pt(3)
        run = p.add_run(header_text.strip())
        format_run(run, size=9, color=WHITE, bold=True)

    # Style data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            set_cell_borders(cell, color=GRAY_BORDER)

            if col_idx == 0:
                # First column: WarmBeige bg, bold
                set_cell_shading(cell, BEIGE_HEX)
                apply_rich_text_to_cell(cell, cell_text.strip(), size=9, color=BLACK, bold=True)
            else:
                # Data cells: White bg
                set_cell_shading(cell, "FFFFFF")
                apply_rich_text_to_cell(cell, cell_text.strip(), size=9, color=BLACK, bold=False)

    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def create_cover_page(doc):
    """Create the cover page with the Cristine V2 design."""
    # Top spacing
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    # Top Bronze divider
    add_bronze_divider(doc, thickness_pt=6)

    # Spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Title
    add_paragraph_with_format(doc, "MODELO DE VALUATION CONSOLIDADO",
                               size=20, color=NAVY_DARK, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               space_after=Pt(4))

    # Subtitle
    add_paragraph_with_format(doc, "OMNI \u2014 AIOS / Petition Engine",
                               size=18, color=NAVY_DARK, bold=False,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               space_after=Pt(16))

    # Bronze divider (small)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500" * 20)
    format_run(run, size=12, color=BRONZE)
    p.paragraph_format.space_after = Pt(16)

    # Based on line
    add_paragraph_with_format(doc, "Baseado em 4 Pesquisas de Deep Research",
                               size=11, color=DEEP_BROWN,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               space_after=Pt(4))

    # Date
    add_paragraph_with_format(doc, "Abril 2026",
                               size=9, color=DEEP_BROWN,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               space_after=Pt(20))

    # Confidential notice
    add_paragraph_with_format(doc, "Documento Confidencial \u2014 Uso Interno / Data Room",
                               size=9, color=DEEP_BROWN, italic=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               space_after=Pt(12))

    # Bottom Bronze divider
    add_bronze_divider(doc, thickness_pt=6)

    # Spacing before metadata
    doc.add_paragraph().paragraph_format.space_after = Pt(30)

    # Metadata block
    meta_lines = [
        ("Documento preparado por:", " Equipe de An\u00e1lise de Investimentos \u2014 OMNI"),
        ("Data:", " 02 de abril de 2026"),
        ("Classifica\u00e7\u00e3o:", " Confidencial \u2014 Uso Interno / Data Room"),
        ("Vers\u00e3o:", " 1.0 \u2014 Consolida\u00e7\u00e3o Final"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_after = Pt(2)
        run_label = p.add_run(label)
        format_run(run_label, size=10, color=NAVY_DARK, bold=True)
        run_value = p.add_run(value)
        format_run(run_value, size=10, color=DEEP_BROWN)

    # Methodological note
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.left_indent = Inches(0.5)
    pf.right_indent = Inches(0.5)
    # Add border for blockquote effect
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="{BRONZE_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run("Nota metodol\u00f3gica: ")
    format_run(run, size=9, color=NAVY_DARK, bold=True)
    run2 = p.add_run(
        "Todo n\u00famero apresentado neste documento \u00e9 rastre\u00e1vel a uma fonte "
        "prim\u00e1ria identificada nas 4 pesquisas de refer\u00eancia. Quando estimativas "
        "s\u00e3o utilizadas, s\u00e3o explicitamente marcadas como tal. Nenhum dado foi fabricado."
    )
    format_run(run2, size=9, color=DEEP_BROWN, italic=True)

    # Page break
    doc.add_page_break()


def setup_headers_footers(doc):
    """Set up running headers and footers."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run("MODELO DE VALUATION   \u2022   OMNI / AIOS   \u2022   Abril 2026")
        format_run(run, size=8, color=DEEP_BROWN)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = fp.add_run("Documento Confidencial   |   Abril 2026   |   p. ")
        format_run(run1, size=8, color=DEEP_BROWN)
        # Add page number field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_pg = fp.add_run()
        run_pg._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_pg2 = fp.add_run()
        run_pg2._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_pg3 = fp.add_run()
        run_pg3._r.append(fldChar2)
        format_run(run_pg, size=8, color=DEEP_BROWN)


def parse_markdown_table(lines):
    """Parse markdown table lines into headers and rows."""
    headers = []
    rows = []
    for i, line in enumerate(lines):
        line = line.strip()
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if i == 0:
            headers = cells
        elif all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
            continue  # separator line
        else:
            # Pad or trim to match header count
            while len(cells) < len(headers):
                cells.append('')
            rows.append(cells[:len(headers)])
    return headers, rows


def parse_markdown(filepath):
    """Parse the markdown file into structured blocks."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    blocks = []
    i = 0
    total = len(lines)

    while i < total:
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip the very first lines (title/subtitle) — we handle them in cover page
        if i < 15:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            i += 1
            continue

        # Empty line
        if stripped == '':
            i += 1
            continue

        # Code block
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < total and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # skip closing ```
            blocks.append(('code', '\n'.join(code_lines)))
            continue

        # Table
        if stripped.startswith('|') and i + 1 < total and lines[i + 1].strip().startswith('|'):
            table_lines = []
            while i < total and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            blocks.append(('table', table_lines))
            continue

        # Headings
        if stripped.startswith('## ') and not stripped.startswith('### '):
            text = stripped[3:].strip()
            blocks.append(('h1', text))
            i += 1
            continue

        if stripped.startswith('### '):
            text = stripped[4:].strip()
            blocks.append(('h2', text))
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            text = stripped[2:].strip()
            blocks.append(('blockquote', text))
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- '):
            text = stripped[2:].strip()
            blocks.append(('bullet', text))
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            text = num_match.group(0)
            blocks.append(('numbered', text))
            i += 1
            continue

        # Bold-prefixed line (like **Source:**)
        if stripped.startswith('**') or stripped.startswith('*'):
            blocks.append(('body', stripped))
            i += 1
            continue

        # Regular body text
        blocks.append(('body', stripped))
        i += 1

    return blocks


def add_code_block(doc, text):
    """Add a styled code/calculation block."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.left_indent = Inches(0.3)
    # Light background via shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F3EF" w:val="clear"/>')
    pPr.append(shading)
    # Left border
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="8" w:space="8" w:color="{BRONZE_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    format_run(run, size=9, color=NAVY_DARK, font_name="Consolas")


def add_blockquote(doc, text):
    """Add a blockquote with Bronze left border."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Inches(0.4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="{BRONZE_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    apply_rich_text(p, text, base_size=10, base_color=DEEP_BROWN)


def add_bullet(doc, text):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.left_indent = Inches(0.4)
    pf.first_line_indent = Inches(-0.2)
    # Bullet character
    run_bullet = p.add_run("\u2022  ")
    format_run(run_bullet, size=11, color=BRONZE)
    apply_rich_text(p, text, base_size=11, base_color=BLACK)


def add_numbered_item(doc, text):
    """Add a numbered list item."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.left_indent = Inches(0.4)
    apply_rich_text(p, text, base_size=11, base_color=BLACK)


def add_body_text(doc, text):
    """Add justified body text with rich formatting."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(4)
    apply_rich_text(p, text, base_size=11, base_color=BLACK)


def build_document():
    """Main function to build the DOCX."""
    doc = Document()

    # ─── PAGE SETUP ───────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.10)
        section.right_margin = Inches(1.00)

    # ─── DEFAULT FONT ─────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = BLACK

    # ─── COVER PAGE ───────────────────────────────────────────────────
    create_cover_page(doc)

    # ─── PARSE CONTENT ────────────────────────────────────────────────
    blocks = parse_markdown(INPUT_MD)

    # ─── RENDER BLOCKS ────────────────────────────────────────────────
    for block_type, content in blocks:
        if block_type == 'h1':
            add_section_heading(doc, content, level=1)

        elif block_type == 'h2':
            add_section_heading(doc, content, level=2)

        elif block_type == 'h3':
            add_section_heading(doc, content, level=3)

        elif block_type == 'table':
            headers, rows = parse_markdown_table(content)
            if headers and rows:
                add_styled_table(doc, headers, rows)

        elif block_type == 'code':
            # Split long code blocks into per-line paragraphs for better formatting
            code_text = content.strip()
            if len(code_text) > 0:
                add_code_block(doc, code_text)

        elif block_type == 'blockquote':
            add_blockquote(doc, content)

        elif block_type == 'bullet':
            add_bullet(doc, content)

        elif block_type == 'numbered':
            add_numbered_item(doc, content)

        elif block_type == 'body':
            if content.strip():
                add_body_text(doc, content)

    # ─── FINAL PAGE: DOCUMENT FOOTER ──────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    add_bronze_divider(doc, thickness_pt=4)

    final_lines = [
        ("Documento gerado em:", " 02 de abril de 2026"),
        ("Baseado em:", " 4 pesquisas de deep research consolidadas"),
        ("Classifica\u00e7\u00e3o:", " Confidencial \u2014 Uso Interno / Data Room"),
        ("Pr\u00f3ximos passos:", " Valida\u00e7\u00e3o com advisors legais e financeiros antes de uso em pitch deck"),
    ]
    for label, value in final_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_after = Pt(2)
        run_l = p.add_run(label)
        format_run(run_l, size=10, color=NAVY_DARK, bold=True)
        run_v = p.add_run(value)
        format_run(run_v, size=10, color=DEEP_BROWN)

    add_bronze_divider(doc, thickness_pt=4)

    # ─── HEADERS & FOOTERS ────────────────────────────────────────────
    setup_headers_footers(doc)

    # ─── SAVE ─────────────────────────────────────────────────────────
    doc.save(OUTPUT_DOCX)
    print(f"[OK] Document saved to: {OUTPUT_DOCX}")
    print(f"[OK] Total blocks processed: {len(blocks)}")


if __name__ == "__main__":
    build_document()
