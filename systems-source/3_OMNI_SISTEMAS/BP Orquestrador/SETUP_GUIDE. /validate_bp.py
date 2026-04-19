#!/usr/bin/env python3
"""
Validação automática do Business Plan — Vieira Operations LLC
Rodar após cada geração. Loop até 0 erros.
"""
from docx import Document
import sys
import os

def validate(docx_path):
    doc = Document(docx_path)
    errors = []
    warnings = []
    info = []
    
    total_paras = len(doc.paragraphs)
    info.append(f"Total de parágrafos: {total_paras}")
    
    # ========================================
    # 1. PARÁGRAFOS LONGOS
    # ========================================
    long_500 = 0
    long_400 = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if len(text) > 500:
            long_500 += 1
            errors.append(f"Para {i} [{p.style.name}] tem {len(text)}ch (MAX 400): «{text[:60]}...»")
        elif len(text) > 400:
            long_400 += 1
            warnings.append(f"Para {i} [{p.style.name}] tem {len(text)}ch (recomendado <400): «{text[:60]}...»")
    info.append(f"Parágrafos >500ch: {long_500} | >400ch: {long_400}")
    
    # ========================================
    # 2. SUBTÍTULOS ÓRFÃOS (Normal style que parece heading)
    # ========================================
    orphan_count = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name
        if style == 'Normal' and 20 < len(text) < 100:
            if not text.startswith('•') and not text.startswith('✔') and not text.startswith('━'):
                if not text.startswith('---') and not text[0:2].isdigit():
                    if text[0].isupper() and ':' not in text[:5]:
                        orphan_count += 1
                        errors.append(f"Subtítulo órfão para {i} [{style}]: «{text}»")
    info.append(f"Subtítulos órfãos: {orphan_count}")
    
    # ========================================
    # 3. CONTEÚDO DUPLICADO
    # ========================================
    seen = {}
    dupes = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if len(text) > 200:
            key = text[:100]
            if key in seen:
                dupes += 1
                errors.append(f"DUPLICAÇÃO: para {i} ≈ para {seen[key]}: «{text[:60]}...»")
            seen[key] = i
    info.append(f"Blocos duplicados: {dupes}")
    
    # ========================================
    # 4. SEÇÕES OBRIGATÓRIAS
    # ========================================
    headings = []
    for i, p in enumerate(doc.paragraphs):
        if 'Heading' in p.style.name:
            headings.append((i, p.text.strip(), p.style.name))
    
    heading_texts_lower = [h[1].lower() for h in headings]
    
    required_sections = {
        'Serviços Oferecidos': ['serviços oferecidos', 'serviços da empresa', 'portfólio de serviços'],
        'Marketing Mix': ['marketing mix'],
        'Perfil Comportamental': ['perfil comportamental', 'behavioral profile'],
        'Perfil Psicográfico': ['perfil psicográfico', 'psychographic'],
        'Perfil Geográfico': ['perfil geográfico', 'geographic profile'],
        'Recursos Físicos': ['recursos físicos', 'recursos fisicos'],
        'Quadro de Funcionários': ['quadro de funcionários', 'quadro de funcionarios'],
        'Executive Summary': ['executive summary', 'sumário executivo'],
        'Análise SWOT': ['swot'],
        'Porter / Cinco Forças': ['concorrentes', 'novos entrantes', 'fornecedores', 'substitutos'],
        'Marketing 4.0': ['marketing 4.0'],
        'Financial Plan': ['financial plan', 'plano financeiro'],
        'Break Even': ['break even'],
        'Considerações Finais': ['considerações finais', 'final considerations'],
    }
    
    missing = 0
    for section_name, keywords in required_sections.items():
        found = any(
            any(kw in ht for kw in keywords)
            for ht in heading_texts_lower
        )
        if not found:
            missing += 1
            errors.append(f"Seção AUSENTE: {section_name}")
    info.append(f"Seções obrigatórias ausentes: {missing}")
    
    # ========================================
    # 5. PERFIL DEMOGRÁFICO REPETIDO
    # ========================================
    demo_count = sum(1 for h in headings if 'demográfico' in h[1].lower() and 'Heading3' in h[2])
    if demo_count > 1:
        errors.append(f"'Perfil Demográfico' aparece {demo_count}x como H3 — deve aparecer apenas 1x")
    
    # ========================================
    # 6. HEADINGS SEM CONTEÚDO
    # ========================================
    empty_headings = 0
    for idx, (para_idx, text, style) in enumerate(headings):
        if idx + 1 < len(headings):
            next_para_idx = headings[idx + 1][0]
            content = [p for p in doc.paragraphs[para_idx+1:next_para_idx] if p.text.strip()]
            if not content:
                empty_headings += 1
                warnings.append(f"Heading sem conteúdo: «{text}»")
    info.append(f"Headings sem conteúdo: {empty_headings}")
    
    # ========================================
    # 7. NUMERAÇÃO
    # ========================================
    h2_numbers = []
    for _, text, style in headings:
        if 'Heading2' in style and text[0].isdigit():
            parts = text.split('.')
            try:
                nums = [int(p) for p in parts[:3] if p.strip().isdigit()]
                h2_numbers.append(('.'.join(str(n) for n in nums), text))
            except:
                pass
    
    # Check for known bad patterns
    for num, text in h2_numbers:
        if num in ['3.3.3', '3.3.4']:
            errors.append(f"Numeração incorreta: {num} — deveria ser 3.2.3/3.2.4: «{text}»")
    
    # ========================================
    # 8. SEPARADORES MARKDOWN
    # ========================================
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '---':
            errors.append(f"Separador markdown '---' encontrado em para {i}")
    
    # ========================================
    # 9. TABELAS
    # ========================================
    table_count = len(doc.tables)
    info.append(f"Total de tabelas: {table_count}")
    if table_count < 10:
        warnings.append(f"Apenas {table_count} tabelas — benchmark tem 15-20+")
    
    # ========================================
    # RELATÓRIO
    # ========================================
    print(f"\n{'='*70}")
    print(f"  VALIDAÇÃO: {os.path.basename(docx_path)}")
    print(f"{'='*70}")
    
    print(f"\n📊 INFO:")
    for i in info:
        print(f"  ℹ️  {i}")
    
    if errors:
        print(f"\n🔴 ERROS ({len(errors)}):")
        for e in errors:
            print(f"  ❌ {e}")
    
    if warnings:
        print(f"\n🟡 AVISOS ({len(warnings)}):")
        for w in warnings:
            print(f"  ⚠️  {w}")
    
    if not errors and not warnings:
        print(f"\n✅ DOCUMENTO PERFEITO — 0 erros, 0 avisos!")
    elif not errors:
        print(f"\n✅ DOCUMENTO APROVADO — 0 erros, {len(warnings)} avisos")
    else:
        print(f"\n❌ DOCUMENTO REPROVADO — {len(errors)} erros, {len(warnings)} avisos")
        print(f"   Corrigir todos os erros e rodar novamente.")
    
    print(f"{'='*70}\n")
    return len(errors) == 0


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Uso: python validate_bp.py <caminho_do_docx>")
        sys.exit(1)
    
    success = validate(sys.argv[1])
    sys.exit(0 if success else 1)
