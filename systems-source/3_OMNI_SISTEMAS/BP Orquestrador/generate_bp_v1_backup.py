#!/usr/bin/env python3
"""
Business Plan Generator — Vieira Operations LLC
Generates complete .docx from scratch following CLAUDE.md formatting rules.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output")
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "BP_Vieira_Operations_LLC_FINAL.docx")

# ============================================================
# STYLE HELPERS
# ============================================================

def setup_styles(doc):
    """Configure document styles per CLAUDE.md: Garamond 11pt body, heading hierarchy."""
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(3)

    # H1: CAIXA ALTA, bold, 16pt
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Garamond'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # H2: Title Case, bold, 13pt
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Garamond'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2C, 0x2C, 0x54)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # H3: Title Case, bold, 11pt
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Garamond'
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x66)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # H4: Italic, 11pt
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Garamond'
    h4.font.size = Pt(11)
    h4.font.bold = False
    h4.font.italic = True
    h4.font.color.rgb = RGBColor(0x55, 0x55, 0x77)
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(4)

    # CoverPage style (avoid orphan subtitle detection)
    try:
        doc.styles.add_style('CoverPage', WD_STYLE_TYPE.PARAGRAPH)
    except:
        pass
    cp = doc.styles['CoverPage']
    cp.font.name = 'Garamond'
    cp.font.size = Pt(11)
    cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(4)

    # Create Bullet style if not exists
    try:
        doc.styles.add_style('BulletItem', WD_STYLE_TYPE.PARAGRAPH)
    except:
        pass
    bullet = doc.styles['BulletItem']
    bullet.font.name = 'Garamond'
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.15
    bullet.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Checkmark style
    try:
        doc.styles.add_style('CheckItem', WD_STYLE_TYPE.PARAGRAPH)
    except:
        pass
    check = doc.styles['CheckItem']
    check.font.name = 'Garamond'
    check.font.size = Pt(11)
    check.paragraph_format.left_indent = Inches(0.25)
    check.paragraph_format.space_after = Pt(4)
    check.paragraph_format.line_spacing = 1.15

    # Table subtitle style
    try:
        doc.styles.add_style('TableSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    except:
        pass
    ts = doc.styles['TableSubtitle']
    ts.font.name = 'Garamond'
    ts.font.size = Pt(10)
    ts.font.italic = True
    ts.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    ts.paragraph_format.space_before = Pt(8)
    ts.paragraph_format.space_after = Pt(4)


def add_paragraph_bold_inline(doc, text, bold_terms=None, style='Normal'):
    """Add paragraph with bold inline terms. text can have **bold** markers."""
    p = doc.add_paragraph(style=style)
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.font.name = 'Garamond'
            run.font.size = Pt(11)
            if i % 2 == 1:
                run.bold = True
    else:
        run = p.add_run(text)
        run.font.name = 'Garamond'
        run.font.size = Pt(11)
    return p


def add_bullet(doc, bold_lead, text):
    """Add bullet: • **Bold Lead**: text"""
    p = doc.add_paragraph(style='BulletItem')
    run = p.add_run('• ')
    run.font.name = 'Garamond'
    run.font.size = Pt(11)
    run = p.add_run(bold_lead)
    run.font.name = 'Garamond'
    run.font.size = Pt(11)
    run.bold = True
    run = p.add_run(': ' + text)
    run.font.name = 'Garamond'
    run.font.size = Pt(11)
    return p


def add_checkmark(doc, bold_lead, text):
    """Add checkmark: ✔ **Bold**: text"""
    p = doc.add_paragraph(style='CheckItem')
    run = p.add_run('✔ ')
    run.font.name = 'Garamond'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
    run = p.add_run(bold_lead)
    run.font.name = 'Garamond'
    run.font.size = Pt(11)
    run.bold = True
    run = p.add_run(': ' + text)
    run.font.name = 'Garamond'
    run.font.size = Pt(11)
    return p


def add_numbered_item(doc, number, bold_lead, text):
    """Add numbered item: 1. **Bold**: text"""
    p = doc.add_paragraph(style='BulletItem')
    run = p.add_run(f'{number}. ')
    run.font.name = 'Garamond'
    run.font.size = Pt(11)
    run = p.add_run(bold_lead)
    run.font.name = 'Garamond'
    run.font.size = Pt(11)
    run.bold = True
    run = p.add_run(': ' + text)
    run.font.name = 'Garamond'
    run.font.size = Pt(11)
    return p


def add_table_subtitle(doc, text):
    """Add italic table subtitle."""
    p = doc.add_paragraph(style='TableSubtitle')
    run = p.add_run(text)
    run.font.name = 'Garamond'
    run.font.size = Pt(10)
    run.italic = True
    return p


def create_table(doc, headers, rows, col_widths=None):
    """Create formatted table with header styling per CLAUDE.md."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = 'Garamond'
        run.font.size = Pt(9)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Background #E8E0D4
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E0D4"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = 'Garamond'
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)

    return table


def add_separator(doc):
    """Add visual separator line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def add_page_break(doc):
    """Add explicit page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)


# ============================================================
# COVER PAGE
# ============================================================

def build_cover_page(doc):
    """Build professional cover page."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    # Company name
    p = doc.add_paragraph(style='CoverPage')
    run = p.add_run('VIEIRA OPERATIONS LLC')
    run.font.name = 'Garamond'
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph(style='CoverPage')

    # Subtitle
    p = doc.add_paragraph(style='CoverPage')
    run = p.add_run('BUSINESS PLAN')
    run.font.name = 'Garamond'
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x54)

    doc.add_paragraph(style='CoverPage')

    # Description
    p = doc.add_paragraph(style='CoverPage')
    run = p.add_run('Prestação de Serviços Especializados em Resiliência\nde Supply Chain e Otimização Operacional Integrada')
    run.font.name = 'Garamond'
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

    for _ in range(4):
        doc.add_paragraph()

    # Separator
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 50)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph()

    # Details
    details = [
        ('Fundador:', 'Everton Rodrigo Vieira'),
        ('Estrutura Jurídica:', 'S-Corporation'),
        ('Sede:', 'Jacksonville, FL — EUA'),
        ('NAICS:', '541611 — Administrative Management and General Management Consulting'),
        ('Documento:', 'Confidencial — Uso exclusivo para petição EB-2 NIW'),
    ]
    for label, value in details:
        p = doc.add_paragraph(style='CoverPage')
        run = p.add_run(label + ' ')
        run.font.name = 'Garamond'
        run.font.size = Pt(11)
        run.bold = True
        run = p.add_run(value)
        run.font.name = 'Garamond'
        run.font.size = Pt(11)

    doc.add_paragraph(style='CoverPage')

    # Confidentiality
    p = doc.add_paragraph(style='CoverPage')
    run = p.add_run('This business plan contains proprietary information. Unauthorized distribution is prohibited.')
    run.font.name = 'Garamond'
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ============================================================
# TABLE OF CONTENTS
# ============================================================

def build_toc(doc):
    """Add Table of Contents placeholder."""
    doc.add_page_break()
    doc.add_heading('SUMÁRIO', level=1)

    # TOC field
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-4" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)
    run4 = p.add_run('[Atualize o Sumário no Word: Ctrl+A → F9]')
    run4.font.name = 'Garamond'
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)


# ============================================================
# SECTION 1: EXECUTIVE SUMMARY
# ============================================================

def build_section_1(doc):
    """1. EXECUTIVE SUMMARY"""
    doc.add_page_break()
    add_separator(doc)
    doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
    add_paragraph_bold_inline(doc,
        'Esta seção apresenta a visão geral do empreendimento, incluindo a **oportunidade de negócio**, '
        'o **portfólio de serviços**, a **timeline** operacional e o **enquadramento jurídico** '
        'da Vieira Operations LLC.')

    # 1.1
    doc.add_heading('1.1. Oportunidade de Negócio', level=2)
    add_paragraph_bold_inline(doc,
        'A **Vieira Operations LLC** configura-se como resposta técnica direta às '
        '**vulnerabilidades sistêmicas** documentadas nas cadeias de suprimentos industriais dos '
        'Estados Unidos, agravadas pela pandemia de COVID-19 e por tensões geopolíticas globais.')
    add_paragraph_bold_inline(doc,
        'O empreendimento estrutura-se como **S-Corporation** sediada em **Jacksonville, Florida**, '
        'e endereça as lacunas operacionais identificadas pela **Executive Order 14017** '
        '(America\'s Supply Chains), que destinou mais de **$37 bilhões** para fortalecimento '
        'de cadeias produtivas domésticas.')
    add_paragraph_bold_inline(doc,
        'O fundador, **Everton Rodrigo Vieira**, traz **25 anos de experiência** na **AmBev/AB InBev**, '
        'incluindo gestão de **R$380 bilhões em ativos**, redução documentada de **48% em custos variáveis** '
        'e implementação do **SAP Fiori** com antecipação de **2 anos e 4 meses** sobre o cronograma previsto.')
    add_paragraph_bold_inline(doc,
        'A empresa projeta **receita bruta acumulada de $7.619.836** em cinco anos, com geração de '
        '**14 empregos diretos** e aproximadamente **62 empregos indiretos** (multiplicador EPI de **4,43x** '
        'para NAICS 5416), totalizando **76 postos de trabalho** no Sudeste dos Estados Unidos.')

    add_checkmark(doc, 'Investimento Inicial Y0', '$121.700 com break-even projetado para o Mês 6 de operação')
    add_checkmark(doc, 'NPV (10%)', '$1.243.218 com IRR de 187% e payback descontado de 14 meses')
    add_checkmark(doc, 'Margem de Contribuição', '80,2% média sobre cinco anos de operação')

    # 1.2 Serviços Oferecidos
    doc.add_heading('1.2. Serviços Oferecidos', level=2)
    add_paragraph_bold_inline(doc,
        'A **Vieira Operations LLC** oferece um portfólio integrado de **seis serviços especializados** '
        'que operam em sinergia para endereçar vulnerabilidades operacionais em cadeias industriais. '
        'Cada serviço combina a expertise de **25 anos** do fundador com metodologias validadas em '
        'operações multinacionais da **AmBev/AB InBev**.')

    add_table_subtitle(doc, 'Tabela 1.2: Portfólio de Serviços Especializados da Vieira Operations LLC')
    create_table(doc,
        ['#', 'Serviço', 'Descrição', 'Público-Alvo'],
        [
            ['1', 'Otimização Avançada de Supply Chain Industrial',
             'Diagnóstico, redesenho e implementação de processos logísticos com foco em redução de custos e aumento de resiliência operacional.',
             'Manufatura, Defesa, Logística'],
            ['2', 'Implementação de Sistemas ERP (SAP Fiori/S4HANA)',
             'Planejamento, configuração e deploy de módulos SAP com transferência de conhecimento para equipes internas do cliente.',
             'Indústria, Saúde, Serviços'],
            ['3', 'Gestão de Projetos Emergenciais e Infraestrutura Essencial',
             'Mobilização rápida para crises operacionais incluindo reestruturação de supply chain e gestão de infraestrutura crítica.',
             'Saúde, Governo, Defesa'],
            ['4', 'Conformidade Regulatória Industrial (EPA, OSHA, FDA)',
             'Auditoria, adequação e manutenção de conformidade com regulamentações federais e estaduais para operações industriais.',
             'Manufatura, Alimentos, Farmacêutico'],
            ['5', 'Capacitação Técnica e Desenvolvimento de Força de Trabalho',
             'Programa de Excelência com ~120h de treinamento online cobrindo supply chain, gestão operacional e ferramentas ERP.',
             'Profissionais individuais, PMEs'],
            ['6', 'Treinamento de Liderança e Gestão Empresarial',
             'Cursos in-company focados em desenvolvimento de líderes operacionais com metodologias de gestão de alta performance.',
             'Executivos, Gestores de PMEs'],
        ],
        col_widths=[0.3, 1.8, 2.8, 1.3]
    )

    add_paragraph_bold_inline(doc,
        'A receita distribui-se com concentração estratégica: **Supply Chain** e **ERP** representam '
        'conjuntamente **70% do faturamento** (35% cada), enquanto os demais serviços contribuem com '
        '**10% cada**, garantindo diversificação sem diluição de foco.')

    add_bullet(doc, 'Supply Chain + ERP', 'representam $5.149.267 (67,6%) da receita acumulada de 5 anos')
    add_bullet(doc, 'Treinamentos e Compliance', 'complementam o portfólio gerando recorrência e fidelização')
    add_bullet(doc, 'Gestão Emergencial', 'posiciona a empresa como parceiro de confiança em situações críticas')

    # 1.3 Timeline do Negócio
    doc.add_heading('1.3. Timeline do Negócio', level=2)
    add_paragraph_bold_inline(doc,
        'A timeline operacional segue cronograma de **12 meses** para estabilização, com início de '
        'atendimento a clientes no **4º mês** e atingimento do break-even no **6º mês**. '
        'A estrutura de crescimento contempla três fases geográficas de expansão.')

    add_table_subtitle(doc, 'Tabela 1.3: Timeline de Implementação — Primeiros 12 Meses')
    create_table(doc,
        ['Mês', 'Marco Principal', 'Atividades-Chave'],
        [
            ['1-2', 'Constituição Legal', 'Registro S-Corp, EIN, licenças FL, abertura contas bancárias, contratação contábil/jurídica'],
            ['2-3', 'Setup Operacional', 'Locação escritório Jacksonville, aquisição equipamentos, setup TI, contratação inicial (3 posições)'],
            ['3-4', 'Go-to-Market', 'Website, identidade visual, registro em câmaras de comércio, prospecção inicial B2B'],
            ['4-6', 'Primeiros Clientes', 'Engajamento com clientes âncora (defesa/saúde), primeiros contratos de Supply Chain e ERP'],
            ['6-8', 'Break-Even', 'Atingimento do ponto de equilíbrio, ajuste de pricing, expansão de pipeline comercial'],
            ['9-12', 'Consolidação Y1', 'Estabilização de receita, planejamento expansão Savannah (Y2), contratação adicional'],
        ],
        col_widths=[0.7, 1.5, 4.0]
    )

    add_paragraph_bold_inline(doc,
        'A expansão geográfica segue padrão progressivo: **Jacksonville, FL** (Y0-Y1, sede), '
        '**Savannah, GA** (Y2-Y3, Branch 1) e **Brunswick, GA** (Y5, Branch 2), formando '
        'corredor logístico estratégico ao longo da **I-95**.')

    # 1.4 Visão, Missão e Valores
    doc.add_heading('1.4. Visão, Missão e Valores', level=2)
    add_paragraph_bold_inline(doc,
        'A **Vieira Operations LLC** fundamenta sua identidade institucional em três pilares — '
        '**visão estratégica** de longo prazo, **missão operacional** orientada a resultados, '
        'e **valores organizacionais** que norteiam cada decisão de negócio.')

    doc.add_heading('Visão', level=3)
    add_paragraph_bold_inline(doc,
        'Ser reconhecida como **referência em resiliência operacional** e otimização de supply chain '
        'no Sudeste dos Estados Unidos, transformando vulnerabilidades industriais em '
        '**vantagens competitivas mensuráveis** para clientes corporativos e comunidades locais.')

    doc.add_heading('Missão', level=3)
    add_paragraph_bold_inline(doc,
        'Entregar **soluções integradas de excelência operacional** que combinem expertise em '
        'supply chain, tecnologia ERP, conformidade regulatória e capacitação de força de trabalho, '
        'gerando **impacto econômico documentável** nas regiões de atuação.')

    doc.add_heading('Valores', level=3)
    add_bullet(doc, 'Excelência Operacional', 'compromisso com resultados mensuráveis e melhoria contínua em cada projeto entregue')
    add_bullet(doc, 'Integridade', 'transparência em todas as relações com clientes, parceiros e comunidades')
    add_bullet(doc, 'Inovação Aplicada', 'adoção de tecnologias e metodologias comprovadas para maximizar impacto')
    add_bullet(doc, 'Responsabilidade Social', 'contribuição ativa para geração de empregos e desenvolvimento econômico regional')

    # 1.5 Enquadramento Jurídico
    doc.add_heading('1.5. Enquadramento Jurídico', level=2)
    add_paragraph_bold_inline(doc,
        'A **Vieira Operations LLC** estrutura-se juridicamente como **S-Corporation** registrada '
        'no estado da **Florida**, escolha fundamentada em três vantagens estruturais para '
        'operações de consultoria especializada de alto valor agregado:')
    add_bullet(doc, 'Otimização Tributária',
        'regime de pass-through taxation elimina dupla tributação corporativa e permite distribuição '
        'eficiente de lucros ao fundador')
    add_bullet(doc, 'Proteção Patrimonial',
        'separação legal entre ativos pessoais e corporativos, blindando o empreendedor de '
        'responsabilidades empresariais')
    add_bullet(doc, 'Credibilidade Institucional',
        'estrutura S-Corp transmite solidez e profissionalismo para clientes corporativos '
        'e contratados de defesa')

    add_paragraph_bold_inline(doc,
        'O processo de constituição legal compreende: registro na **Florida Division of Corporations**, '
        'obtenção do **EIN** junto ao IRS, registro de **fictitious name**, abertura de conta bancária '
        'empresarial, e obtenção de **Business Tax Receipt** no Condado de Duval.')
    add_paragraph_bold_inline(doc,
        'A empresa manterá conformidade com regulamentações federais (**OSHA**, **EPA**, **FDA**) '
        'e estaduais aplicáveis ao setor de consultoria, além de certificações profissionais '
        'relevantes como **SAP Certified** e membros de **APICS/ASCM**.')


# ============================================================
# SECTION 2: ANÁLISE ESTRATÉGICA DE MERCADO
# ============================================================

def build_section_2(doc):
    """2. ANÁLISE ESTRATÉGICA DE MERCADO"""
    doc.add_page_break()
    add_separator(doc)
    doc.add_heading('2. ANÁLISE ESTRATÉGICA DE MERCADO', level=1)
    add_paragraph_bold_inline(doc,
        'Esta seção analisa o ambiente de mercado, incluindo **perspectivas setoriais**, '
        '**empregabilidade projetada**, **impactos ESG**, análise **SWOT** e as **Cinco Forças de Porter** '
        'aplicadas ao segmento de consultoria especializada.')

    # 2.1
    doc.add_heading('2.1. Perspectivas do Mercado', level=2)
    add_paragraph_bold_inline(doc,
        'O mercado de serviços especializados em **resiliência de supply chain** e otimização '
        'operacional nos Estados Unidos apresenta trajetória de **crescimento acelerado**, '
        'impulsionado por três fatores estruturais convergentes.')
    add_bullet(doc, 'Reconfiguração de Cadeias Globais',
        'pandemia de COVID-19 e tensões geopolíticas aceleraram reshoring e nearshoring, '
        'gerando demanda por consultoria especializada em redesenho logístico')
    add_bullet(doc, 'Digitalização Industrial',
        'adoção de sistemas ERP integrados (SAP S/4HANA) e tecnologias Industry 4.0 '
        'cresce a taxa de 12-15% ao ano no segmento de manufatura')
    add_bullet(doc, 'Regulamentação Federal',
        'Executive Order 14017 e CHIPS Act alocaram mais de $75 bilhões para fortalecimento '
        'de cadeias produtivas domésticas, criando oportunidades para consultorias especializadas')

    add_paragraph_bold_inline(doc,
        'O segmento **NAICS 541611** (Administrative Management and General Management Consulting) '
        'gerou receita superior a **$300 bilhões** em 2023, com taxa de crescimento anual composta '
        '(**CAGR**) de **6,2%** projetada até 2028 (U.S. Census Bureau, Service Annual Survey 2023).')
    add_paragraph_bold_inline(doc,
        'A estratégia de entrada da **Vieira Operations LLC** concentra-se no **Sudeste americano**, '
        'região que abriga o **JAXPORT** (um dos 17 portos estratégicos do Departamento de Defesa), '
        'o **Port of Savannah** (4º maior porto dos EUA em TEUs) e mais de **2.300 empresas** '
        'manufatureiras na região metropolitana de Jacksonville.')

    add_checkmark(doc, 'Mercado Endereçável', 'superior a $300 bilhões anuais com CAGR de 6,2%')
    add_checkmark(doc, 'Região Estratégica', 'Sudeste dos EUA com concentração de infraestrutura logística e industrial')

    # 2.2
    doc.add_heading('2.2. Cadeia de Suprimentos', level=2)
    add_paragraph_bold_inline(doc,
        'A **Vieira Operations LLC** opera modelo de negócio fundamentalmente distinto das '
        'cadeias de suprimentos tradicionais. Como prestadora de **serviços especializados de '
        'alto valor agregado**, a empresa não depende de matérias-primas ou distribuição física.')
    add_paragraph_bold_inline(doc,
        'A cadeia de valor estrutura-se em três elos principais: **(a)** capital intelectual '
        'como insumo primário, **(b)** plataformas tecnológicas como ferramentas de entrega, '
        'e **(c)** relacionamentos B2B como canal de distribuição.')

    add_table_subtitle(doc, 'Tabela 2.2: Cadeia de Valor — Vieira Operations LLC')
    create_table(doc,
        ['Elo da Cadeia', 'Componente', 'Função Estratégica'],
        [
            ['Insumo', 'Capital Intelectual', 'Expertise de 25 anos do fundador + metodologias proprietárias'],
            ['Insumo', 'Plataformas Tecnológicas', 'SAP, Microsoft Project, ferramentas de BI e automação'],
            ['Transformação', 'Diagnóstico e Implementação', 'Análise de vulnerabilidades → desenho de soluções → execução'],
            ['Transformação', 'Capacitação', 'Treinamento e transferência de conhecimento para equipes do cliente'],
            ['Distribuição', 'Canais B2B Diretos', 'Prospecção ativa, câmaras de comércio, parcerias institucionais'],
            ['Distribuição', 'Digital + Eventos', 'Website, LinkedIn, webinars técnicos, conferências setoriais'],
        ],
        col_widths=[1.2, 1.8, 3.2]
    )

    add_paragraph_bold_inline(doc,
        'Este modelo gera **efeito multiplicador**: cada projeto bem-sucedido fortalece a '
        'reputação da empresa e gera referências para novos contratos, reduzindo custo de '
        'aquisição de clientes ao longo do tempo.')

    # 2.3
    doc.add_heading('2.3. Empregabilidade Esperada (Direta e Indireta)', level=2)
    add_paragraph_bold_inline(doc,
        'A projeção de geração de empregos da **Vieira Operations LLC** fundamenta-se em '
        'três pilares metodológicos: **(a)** cronograma de contratação alinhado à expansão '
        'geográfica, **(b)** multiplicadores do **Economic Policy Institute (EPI)** para NAICS 5416, '
        'e **(c)** projeções do **Bureau of Labor Statistics (BLS)** para códigos SOC relacionados.')

    doc.add_heading('Empregos Diretos Projetados', level=3)

    add_table_subtitle(doc, 'Tabela 2.3a: Projeção de Empregos Diretos por Localidade (Y1-Y5)')
    create_table(doc,
        ['Cargo (SOC)', 'Y1', 'Y2', 'Y3', 'Y4', 'Y5'],
        [
            ['Supply Chain Manager — Owner (11-3031)', '1', '1', '1', '1', '1'],
            ['Project Management Specialist (13-1082)', '1', '1', '2', '2', '5'],
            ['Office Clerks, General (43-9061)', '1', '2', '4', '4', '6'],
            ['First-Line Supervisors of Office (43-1011)', '1', '2', '4', '4', '6'],
            ['Receptionists and Info Clerk (43-4171)', '1', '1', '2', '2', '4'],
            ['TOTAL', '5', '7', '13', '13', '22'],
        ],
        col_widths=[2.8, 0.6, 0.6, 0.6, 0.6, 0.6]
    )

    add_paragraph_bold_inline(doc,
        'A distribuição geográfica acompanha a expansão: **Jacksonville** (sede, Y1-Y5), '
        '**Savannah** (Branch 1, a partir de Y2) e **Brunswick** (Branch 2, a partir de Y5). '
        'Até o **Ano 5**, a empresa projeta **14 empregos diretos** distribuídos nas três localidades.')

    doc.add_heading('Empregos Indiretos — Multiplicador EPI', level=3)
    add_paragraph_bold_inline(doc,
        'O **Economic Policy Institute** documenta multiplicador de **4,43x** para o setor '
        '**NAICS 5416** (Management, Scientific, and Technical Consulting Services). '
        'Aplicando este multiplicador sobre a base de **14 empregos diretos** no Ano 5:')
    add_bullet(doc, 'Empregos Indiretos', 'aproximadamente 62 postos de trabalho gerados na cadeia de suprimentos e serviços adjacentes')
    add_bullet(doc, 'Total de Empregos', 'aproximadamente 76 postos de trabalho (14 diretos + 62 indiretos)')

    add_table_subtitle(doc, 'Tabela 2.3b: Códigos SOC — Ocupações Relacionadas')
    create_table(doc,
        ['Código SOC', 'Ocupação', 'Relação com Serviços'],
        [
            ['11-3071', 'Transportation, Storage, and Distribution Managers', 'Supply Chain Optimization'],
            ['13-1081', 'Logisticians', 'Supply Chain e ERP Implementation'],
            ['15-1299', 'Computer Occupations, All Other', 'SAP Fiori / ERP Systems'],
            ['13-1111', 'Management Analysts', 'Consultoria de Gestão e Compliance'],
            ['11-3051', 'Industrial Production Managers', 'Gestão de Projetos Emergenciais'],
        ],
        col_widths=[1.0, 2.5, 2.7]
    )

    add_checkmark(doc, 'Impacto na Empregabilidade', '76 postos de trabalho projetados no Ano 5 (14 diretos + 62 indiretos)')

    # 2.4
    doc.add_heading('2.4. Gestão do Conhecimento', level=2)
    add_paragraph_bold_inline(doc,
        'A **Vieira Operations LLC** fundamenta sua proposta de valor em um ativo estratégico '
        'intangível: a combinação de **25 anos de experiência operacional** em multinacional '
        'líder global com **metodologias proprietárias** de otimização e gestão.')
    add_paragraph_bold_inline(doc,
        'A gestão do conhecimento estrutura-se através de três componentes reconhecidos pelo '
        '**OECD Oslo Manual** como pilares de inovação em serviços:')

    add_bullet(doc, 'Capital Humano', 'competências técnicas do fundador e da equipe em supply chain, ERP e compliance regulatório')
    add_bullet(doc, 'Capital Estrutural', 'processos, metodologias proprietárias, frameworks de implementação e sistemas de gestão documentados')
    add_bullet(doc, 'Capital Relacional', 'rede de relacionamentos B2B, parcerias institucionais e reputação no mercado')

    add_table_subtitle(doc, 'Tabela 2.4: Mapeamento de Competências Técnicas Essenciais')
    create_table(doc,
        ['Competência', 'Serviço Relacionado', 'Nível', 'Fonte'],
        [
            ['Supply Chain Design & Optimization', 'Serviço 1', 'Expert', 'AmBev/AB InBev — 25 anos'],
            ['SAP Fiori / S4HANA Implementation', 'Serviço 2', 'Expert', 'Implementação antecipada 2a4m'],
            ['Crisis Management & Emergency Ops', 'Serviço 3', 'Expert', 'Planta O2 COVID-19'],
            ['EPA/OSHA/FDA Compliance', 'Serviço 4', 'Advanced', 'Gestão regulatória industrial'],
            ['Technical Training Design', 'Serviço 5', 'Advanced', 'Programa de Excelência'],
            ['Leadership Development', 'Serviço 6', 'Advanced', 'Gestão de 50+ colaboradores'],
        ],
        col_widths=[2.2, 1.2, 0.8, 2.0]
    )

    # 2.5
    doc.add_heading('2.5. Impactos ESG', level=2)
    add_paragraph_bold_inline(doc,
        'A **Vieira Operations LLC** estrutura sua operação em framework integrado de práticas '
        '**ambientais, sociais e de governança (ESG)**, alinhado aos **17 Objetivos de '
        'Desenvolvimento Sustentável** das Nações Unidas e aos padrões do **SASB**.')

    doc.add_heading('Dimensão Ambiental', level=3)
    add_bullet(doc, 'Otimização de Processos', 'projetos de supply chain geram redução mensurável de desperdícios e emissões nos clientes')
    add_bullet(doc, 'Operação Digital-First', 'modelo de consultoria com uso intensivo de ferramentas remotas minimiza pegada de carbono')

    doc.add_heading('Dimensão Social', level=3)
    add_bullet(doc, 'Geração de Empregos', '76 postos de trabalho projetados em regiões com necessidade de desenvolvimento econômico')
    add_bullet(doc, 'Capacitação Profissional', 'programas de treinamento que elevam empregabilidade de profissionais locais')
    add_bullet(doc, 'Inclusão Econômica', 'atuação em áreas com índice de pobreza acima da média nacional (Duval County: 13,7%)')

    doc.add_heading('Dimensão de Governança', level=3)
    add_bullet(doc, 'Transparência', 'relatórios financeiros auditados e conformidade com regulamentações federais')
    add_bullet(doc, 'Ética Corporativa', 'código de conduta baseado em padrões da AB InBev Global Compliance')

    add_checkmark(doc, 'Alinhamento ESG', 'framework integrado cobrindo 6 ODS das Nações Unidas e padrões SASB')

    # 2.6 SWOT
    doc.add_heading('2.6. Análise SWOT', level=2)
    add_paragraph_bold_inline(doc,
        'A análise **SWOT** posiciona a **Vieira Operations LLC** em relação a fatores internos '
        '(forças e fraquezas) e externos (oportunidades e ameaças) do ambiente competitivo.')

    add_table_subtitle(doc, 'Tabela 2.6: Matriz SWOT — Vieira Operations LLC')
    create_table(doc,
        ['', 'Positivo', 'Negativo'],
        [
            ['Interno',
             'FORÇAS\n• 25 anos AmBev/AB InBev\n• 48% redução custos variáveis\n• SAP Fiori: 2a4m antecipação\n• 6 serviços integrados\n• R$380 bi ativos gerenciados',
             'FRAQUEZAS\n• Marca nova sem histórico nos EUA\n• Dependência inicial do fundador\n• Base de clientes a construir\n• Limitação de escala no Y1'],
            ['Externo',
             'OPORTUNIDADES\n• EO 14017: $37 bi supply chain\n• CHIPS Act: $52 bi semicondutores\n• JAXPORT: porto estratégico DoD\n• CAGR 6,2% consultoria\n• Reshoring acelerado',
             'AMEAÇAS\n• Big 4 com recursos superiores\n• Ciclos econômicos recessivos\n• Escassez de mão-de-obra qualificada\n• Mudanças regulatórias'],
        ],
        col_widths=[0.8, 2.7, 2.7]
    )

    # 2.7 SWOT Cruzada
    doc.add_heading('2.7. SWOT Cruzada', level=2)
    add_paragraph_bold_inline(doc,
        'A **SWOT Cruzada** combina fatores internos e externos para gerar **16 estratégias '
        'acionáveis**, organizadas em quatro quadrantes estratégicos.')

    add_table_subtitle(doc, 'Tabela 2.7: SWOT Cruzada — Estratégias Integradas')
    create_table(doc,
        ['Quadrante', 'Estratégia', 'Ação Principal'],
        [
            ['SO (Forças × Oportunidades)', 'Alavancagem de Expertise',
             'Utilizar 25 anos AmBev para capturar contratos de EO 14017 e CHIPS Act'],
            ['SO', 'Portfólio Integrado',
             'Posicionar 6 serviços como solução única para reshoring/nearshoring'],
            ['WO (Fraquezas × Oportunidades)', 'Construção de Marca',
             'Associar-se a câmaras de comércio e JAXPORT para credibilidade institucional'],
            ['WO', 'Pipeline de Talentos',
             'Parcerias com UNF e JU para suprir limitação de escala'],
            ['ST (Forças × Ameaças)', 'Pricing Competitivo',
             'Expertise Tier-1 a 30-50% menos que Big 4 neutraliza concorrência'],
            ['ST', 'Resultados Documentados',
             'Métricas verificáveis (48% redução) como barreira contra substituição'],
            ['WT (Fraquezas × Ameaças)', 'Early Adopter Strategy',
             'Clientes âncora em defesa/saúde para validar modelo antes de escalar'],
            ['WT', 'Diversificação Geográfica',
             'Expansão faseada (JAX→SAV→BRU) reduz dependência de mercado único'],
        ],
        col_widths=[1.8, 1.5, 2.9]
    )

    # 2.8 Porter
    doc.add_heading('2.8. Análise de Porter — Cinco Forças', level=2)
    add_paragraph_bold_inline(doc,
        'A análise das **Cinco Forças de Porter** avalia a intensidade competitiva e atratividade '
        'do mercado de serviços especializados em resiliência operacional no Sudeste dos EUA.')

    # 2.8.1
    doc.add_heading('2.8.1. Análise de Concorrentes', level=3)
    add_paragraph_bold_inline(doc,
        'O mercado de **NAICS 541611** na região Sudeste apresenta estrutura **fragmentada**, '
        'com predominância de empresas especializadas em disciplinas verticais. '
        'Nenhum concorrente direto combina os seis serviços oferecidos pela Vieira Operations.')

    add_table_subtitle(doc, 'Tabela 2.8.1: Análise Comparativa de Concorrentes')
    create_table(doc,
        ['Concorrente', 'Tipo', 'Foco Principal', '$/Hora', 'Limitação'],
        [
            ['Deloitte / McKinsey', 'Big 4 / MBB', 'Estratégia C-Suite', '$350-500', 'Preço proibitivo para PMEs'],
            ['Consultorias regionais (FL)', 'Boutique', 'Supply Chain genérico', '$125-175', 'Sem expertise ERP/SAP'],
            ['Implementadores SAP', 'Especialista', 'ERP Implementation', '$200-300', 'Sem supply chain operacional'],
            ['Treinadores corporativos', 'Individual', 'Leadership/coaching', '$100-150', 'Sem expertise técnica industrial'],
            ['Vieira Operations LLC', 'Integrado', '6 serviços sinérgicos', '$150-250', 'Marca nova no mercado EUA'],
        ],
        col_widths=[1.5, 0.9, 1.3, 0.8, 1.7]
    )

    add_paragraph_bold_inline(doc,
        'A **Vieira Operations LLC** posiciona-se na lacuna entre consultorias globais '
        '(caras e generalistas) e boutiques locais (baratas mas limitadas), oferecendo '
        '**expertise Tier-1 a preço acessível** para o mid-market industrial.')

    # 2.8.2
    doc.add_heading('2.8.2. Ameaça de Novos Entrantes', level=3)
    add_paragraph_bold_inline(doc,
        'Barreiras de entrada mistas: **capital e regulamentação** são relativamente baixos, '
        'mas **expertise técnica**, **relacionamentos** e **reputação** constituem barreiras '
        'elevadas no segmento de consultoria especializada.')

    add_table_subtitle(doc, 'Tabela 2.8.2: Barreiras de Entrada no Mercado')
    create_table(doc,
        ['Barreira', 'Nível', 'Justificativa'],
        [
            ['Capital Inicial', 'Baixo', 'Modelo asset-light: $121.700 suficiente para iniciar operações'],
            ['Regulamentação', 'Baixo', 'Sem licença específica obrigatória para consultoria em FL'],
            ['Expertise Técnica', 'Alto', '25+ anos em supply chain multinacional raramente replicável'],
            ['Credenciais SAP', 'Alto', 'Certificações SAP requerem investimento e experiência significativos'],
            ['Rede de Relacionamentos', 'Alto', 'Acesso a decision-makers requer anos de construção de confiança'],
            ['Track Record', 'Alto', 'Resultados documentados (48% redução custos) difíceis de replicar'],
        ],
        col_widths=[1.5, 0.8, 3.9]
    )

    add_paragraph_bold_inline(doc,
        'A **Vieira Operations LLC** mitiga a ameaça de novos entrantes através de quatro '
        'vantagens competitivas estruturais:')
    add_numbered_item(doc, 1, 'Expertise Tier-1 a Preço Acessível',
        'experiência AmBev/AB InBev (R$380 bi em ativos) com pricing 30-50% menor que Big 4')
    add_numbered_item(doc, 2, 'Foco em Implementação',
        'diferencial claro em relação a consultorias que entregam apenas relatórios e recomendações')
    add_numbered_item(doc, 3, 'Metodologias Validadas',
        'frameworks testados em operações reais com redução documentada de 48% em custos variáveis')
    add_numbered_item(doc, 4, 'Portfólio Integrado',
        'único player oferecendo 6 serviços cobrindo supply chain, ERP, compliance e treinamento')

    # 2.8.3
    doc.add_heading('2.8.3. Poder de Negociação dos Clientes', level=3)
    add_paragraph_bold_inline(doc,
        'Poder de barganha **moderado**. Múltiplos fornecedores de serviços técnicos aumentam '
        'o poder dos clientes, mas a **escassez de especialistas** com portfólio integrado '
        'limita alternativas reais para projetos complexos.')

    add_table_subtitle(doc, 'Tabela 2.8.3: Poder de Negociação por Segmento de Cliente')
    create_table(doc,
        ['Segmento', 'Poder', 'Justificativa'],
        [
            ['Manufatura Crítica (Defesa)', 'Baixo', 'Alto nível de exigência técnica e compliance reduz alternativas viáveis'],
            ['PMEs Industriais (50-500 func.)', 'Moderado', 'Sensíveis a preço mas valorizam resultados mensuráveis'],
            ['Setor de Saúde', 'Baixo', 'Demandas emergenciais de supply chain reduzem poder de barganha'],
            ['Executivos Individuais (B2C)', 'Alto', 'Múltiplas alternativas de treinamento e certificação'],
        ],
        col_widths=[1.8, 0.8, 3.6]
    )

    # 2.8.4
    doc.add_heading('2.8.4. Poder de Negociação dos Fornecedores', level=3)
    add_paragraph_bold_inline(doc,
        'Poder de barganha dos fornecedores é **baixo a nulo**. A Vieira Operations opera '
        'modelo onde o principal "fornecedor" é o **capital intelectual interno**, reduzindo '
        'dependência de terceiros.')

    add_table_subtitle(doc, 'Tabela 2.8.4: Fornecedores e Estratégia de Mitigação')
    create_table(doc,
        ['Fornecedor', 'Dependência', 'Mitigação'],
        [
            ['SAP (licenças)', 'Moderada', 'Manutenção de certificações múltiplas para flexibilidade'],
            ['Plataformas Cloud (AWS/Azure)', 'Baixa', 'Multi-cloud strategy com alternativas disponíveis'],
            ['Profissionais Técnicos', 'Moderada', 'Parcerias com UNF e JU para pipeline de talentos'],
            ['Espaço Físico (locação)', 'Baixa', 'Mercado de escritórios competitivo em Jacksonville'],
        ],
        col_widths=[1.5, 1.0, 3.7]
    )

    add_bullet(doc, 'Diversificação', 'certificações múltiplas garantem flexibilidade de plataforma tecnológica')
    add_bullet(doc, 'Internalização', 'metodologias proprietárias desenvolvidas ao longo de 25 anos eliminam dependência externa')
    add_bullet(doc, 'Parcerias Acadêmicas', 'pipeline de talentos através de universidades locais (UNF, JU)')

    # 2.8.5
    doc.add_heading('2.8.5. Produtos ou Serviços Substitutos', level=3)
    add_paragraph_bold_inline(doc,
        'Ameaça de substituição é **baixa a moderada**. A combinação única de 6 serviços '
        'integrados dificulta substituição por soluções pontuais ou tecnologias automatizadas.')

    add_table_subtitle(doc, 'Tabela 2.8.5: Substitutos e Estratégia de Mitigação')
    create_table(doc,
        ['Substituto', 'Ameaça', 'Mitigação'],
        [
            ['Equipes internas de SC', 'Moderada', 'Empresas raramente possuem expertise integrada dos 6 serviços'],
            ['Plataformas SaaS de SC', 'Baixa', 'Software não substitui expertise humana em implementação'],
            ['Freelancers especializados', 'Baixa', 'Não oferecem portfólio integrado nem accountability corporativa'],
            ['AI / Automação de processos', 'Baixa', 'Complementar, não substituto — requer expertise para implementação'],
        ],
        col_widths=[1.8, 0.8, 3.6]
    )

    add_bullet(doc, 'Portfólio Integrado', 'integração de 6 serviços em framework único que nenhum substituto replica')
    add_bullet(doc, 'Resultados Mensuráveis', 'ROI quantificável baseado em redução documentada de 48% em custos')
    add_bullet(doc, 'Fidelização', 'acompanhamento pós-implementação aumenta custos de troca para clientes')

    add_checkmark(doc, 'Posição Competitiva', 'lacuna de mercado identificada entre Big 4 e boutiques locais')
    add_checkmark(doc, 'Barreiras de Proteção', 'expertise técnica e track record como defesas estruturais')


# ============================================================
# SECTION 3: MARKETING PLAN
# ============================================================

def build_section_3(doc):
    """3. MARKETING PLAN"""
    doc.add_page_break()
    add_separator(doc)
    doc.add_heading('3. MARKETING PLAN', level=1)
    add_paragraph_bold_inline(doc,
        'O plano de marketing detalha a **segmentação de mercado**, o **Marketing Mix** (4Ps) '
        'e a **estratégia de Marketing 4.0** para posicionamento da Vieira Operations LLC '
        'no mercado de consultoria especializada do Sudeste dos EUA.')

    # 3.1 Segmentação
    doc.add_heading('3.1. Segmentação de Mercado', level=2)
    add_paragraph_bold_inline(doc,
        'A segmentação identifica e prioriza os públicos-alvo **B2C** (profissionais individuais) '
        'e **B2B** (organizações empresariais) com maior potencial de conversão para os serviços '
        'especializados da empresa.')

    doc.add_heading('3.1.1. Visão Geral da Segmentação', level=3)
    add_paragraph_bold_inline(doc,
        'A segmentação de mercado da **Vieira Operations LLC** fundamenta-se em análise '
        'multidimensional que considera fatores **demográficos, comportamentais, psicográficos '
        'e geográficos** para identificação precisa dos públicos-alvo.')
    add_paragraph_bold_inline(doc,
        'A empresa opera a partir de **Jacksonville, Florida** — cidade-porto estratégica '
        'com população metropolitana de **1,6 milhão** de habitantes, presença de setores '
        'de **defesa**, **saúde**, **logística** e **manufatura avançada**.')
    add_paragraph_bold_inline(doc,
        'A segmentação divide-se em duas dimensões complementares: **B2C** (profissionais '
        'individuais que buscam capacitação e certificação) e **B2B** (organizações empresariais '
        'que demandam serviços de otimização operacional).')

    # 3.1.2 B2C
    doc.add_heading('3.1.2. Público-Alvo B2C', level=3)
    add_paragraph_bold_inline(doc,
        'A dimensão **B2C** endereça profissionais individuais cujas necessidades de '
        'desenvolvimento profissional se alinham aos serviços de capacitação técnica '
        'e treinamento de liderança oferecidos pela empresa.')

    doc.add_heading('Perfil Demográfico', level=3)
    add_bullet(doc, 'Faixa Etária', '28-55 anos, profissionais em estágio intermediário a sênior de carreira')
    add_bullet(doc, 'Renda', '$60.000-$150.000/ano, com capacidade de investimento em desenvolvimento profissional')
    add_bullet(doc, 'Formação', 'graduação em Engenharia, Administração, Logística ou áreas correlatas')
    add_bullet(doc, 'Ocupações SOC', 'Logisticians (13-1081), Management Analysts (13-1111), Industrial Managers (11-3051)')

    doc.add_heading('Perfil Comportamental', level=3)
    add_bullet(doc, 'Motivação Principal', 'ascensão profissional através de certificações reconhecidas pelo mercado')
    add_bullet(doc, 'Padrão de Compra', 'pesquisa extensiva online, consulta a pares e validação de credenciais do instrutor')
    add_bullet(doc, 'Canais Preferidos', 'LinkedIn, webinars técnicos, recomendações de colegas de profissão')
    add_bullet(doc, 'Decisão de Compra', 'ciclo de 2-4 semanas com avaliação de ROI educacional')

    doc.add_heading('Perfil Psicográfico', level=3)
    add_bullet(doc, 'Valores', 'excelência profissional, aprendizado contínuo e diferenciação competitiva')
    add_bullet(doc, 'Aspirações', 'transição para posições de liderança em supply chain ou operações')
    add_bullet(doc, 'Dores', 'gap entre formação acadêmica generalista e demandas práticas do mercado industrial')
    add_bullet(doc, 'Disposição', 'investir $2.000-$8.000 em programas de certificação de alta qualidade')

    doc.add_heading('Perfil Geográfico', level=3)
    add_bullet(doc, 'Região Primária', 'Jacksonville Metropolitan Area (1,6 mi habitantes)')
    add_bullet(doc, 'Região Secundária', 'Savannah MSA e Brunswick MSA (corredor I-95)')
    add_bullet(doc, 'Online', 'profissionais de todo Sudeste dos EUA para programas de certificação online')

    # 3.1.3 B2B
    doc.add_heading('3.1.3. Setor-Alvo B2B', level=3)
    add_paragraph_bold_inline(doc,
        'A segmentação **B2B** prioriza quatro setores estratégicos selecionados por '
        '**densidade de vulnerabilidades operacionais**, capacidade de investimento e '
        'alinhamento com a expertise do fundador.')

    add_table_subtitle(doc, 'Tabela 3.1.3: Setores-Alvo B2B Prioritários')
    create_table(doc,
        ['Setor', 'Tamanho Regional', 'Oportunidade', 'Serviços Aplicáveis'],
        [
            ['Defesa e Aeroespacial', 'Naval Station Mayport, BAE Systems, Northrop Grumman',
             'EO 14017, CHIPS Act, resilience requirements', 'Serviços 1, 2, 3, 4'],
            ['Saúde Corporativa', 'Mayo Clinic, Baptist Health, UF Health',
             'Supply chain de insumos, gestão emergencial', 'Serviços 1, 3, 4'],
            ['Manufatura Industrial', '2.300+ empresas na MSA Jacksonville',
             'Otimização operacional, compliance', 'Serviços 1, 2, 4, 5'],
            ['Logística e Portos', 'JAXPORT, Port of Savannah',
             'Resiliência de supply chain, ERP', 'Serviços 1, 2, 5, 6'],
        ],
        col_widths=[1.2, 1.7, 1.7, 1.6]
    )

    # 3.1.4 Posicionamento
    doc.add_heading('3.1.4. Posicionamento da Marca', level=3)
    add_paragraph_bold_inline(doc,
        'A **Vieira Operations LLC** posiciona-se como **"Expertise Tier-1 a Preço Mid-Market"** — '
        'oferecendo qualidade de consultoria equivalente às **Big 4** com pricing acessível '
        'para empresas de médio porte e PMEs industriais.')
    add_paragraph_bold_inline(doc,
        'O posicionamento fundamenta-se em **três pilares de diferenciação** que sustentam a proposta de valor da empresa no mercado de consultoria especializada do Sudeste dos EUA.')

    doc.add_heading('Pilar 1: Convergência de Competências Críticas', level=3)
    add_bullet(doc, 'Diferencial', 'combinação rara de supply chain + ERP + compliance + treinamento em um único prestador')
    add_bullet(doc, 'Evidência', '25 anos AmBev/AB InBev com gestão de R$380 bi em ativos operacionais')

    doc.add_heading('Pilar 2: Resultados Quantificáveis', level=3)
    add_bullet(doc, 'Diferencial', 'track record documentado com métricas verificáveis de impacto')
    add_bullet(doc, 'Evidência', '48% redução custos variáveis, SAP Fiori antecipado 2 anos e 4 meses')

    doc.add_heading('Pilar 3: Pricing Estratégico', level=3)
    add_bullet(doc, 'Diferencial', '$150-250/hora versus $350-500/hora das Big 4')
    add_bullet(doc, 'Evidência', 'estrutura lean de S-Corp permite repasse de eficiência tributária ao cliente')

    # 3.2 Marketing Mix
    doc.add_heading('3.2. Marketing Mix', level=2)
    add_paragraph_bold_inline(doc,
        'O **Marketing Mix** integra os **4Ps** (Produto, Preço, Praça, Promoção) em estratégia '
        'coerente que alinha proposta de valor, modelo de precificação, canais de distribuição '
        'e investimento promocional para maximizar penetração de mercado.')

    # 3.2.1
    doc.add_heading('3.2.1. Produto — Análise de Valor', level=3)
    add_paragraph_bold_inline(doc,
        'O portfólio de **6 serviços especializados** opera como sistema integrado onde cada '
        'componente reforça os demais. A proposta de valor transcende entrega de relatórios '
        'para focar em **implementação efetiva** com resultados mensuráveis.')
    add_paragraph_bold_inline(doc,
        'A metodologia de entrega fundamenta-se em resultados documentados: clientes projetam '
        '**redução de 15-48% em custos operacionais**, **ROI de 3-5x** sobre investimento '
        'em consultoria, e **time-to-value de 3-6 meses** para projetos de otimização.')

    add_table_subtitle(doc, 'Tabela 3.2.1: Proposta de Valor por Serviço')
    create_table(doc,
        ['Serviço', 'Entregável Principal', 'ROI Esperado'],
        [
            ['Supply Chain Optimization', 'Redesenho de processos + redução de custos', '15-48% redução custos variáveis'],
            ['ERP Implementation (SAP)', 'Deploy SAP Fiori + treinamento equipe', '30-40% ganho de eficiência'],
            ['Emergency Project Management', 'Mobilização rápida + resolução de crises', 'Evita perdas de $500K-$2M'],
            ['Regulatory Compliance', 'Auditoria + plano de adequação', 'Evita multas de $50K-$500K'],
            ['Technical Training', 'Programa 120h + certificação', 'Aumento salarial de 15-25%'],
            ['Leadership Training', 'Curso in-company + follow-up', 'Melhoria 20-30% em KPIs de equipe'],
        ],
        col_widths=[1.8, 2.2, 2.2]
    )

    # 3.2.2
    doc.add_heading('3.2.2. Estratégia de Preço', level=3)
    add_paragraph_bold_inline(doc,
        'A precificação adota modelo híbrido de **value-based pricing** combinado com '
        '**project-based fees**, posicionando a empresa no range de **$150-250/hora** — '
        'entre boutiques locais ($100-175) e Big 4 ($350-500).')

    add_table_subtitle(doc, 'Tabela 3.2.2: Projeção de Receita Anual por Serviço (Y1-Y5)')
    create_table(doc,
        ['Serviço', 'Y1', 'Y2', 'Y3', 'Y4', 'Y5'],
        [
            ['Supply Chain Optimization', '$177.408', '$336.336', '$603.187', '$603.187', '$854.515'],
            ['ERP Implementation (SAP)', '$177.408', '$336.336', '$603.187', '$603.187', '$854.515'],
            ['Emergency Project Mgmt', '$50.688', '$96.096', '$172.339', '$172.339', '$244.147'],
            ['Regulatory Compliance', '$50.688', '$96.096', '$172.339', '$172.339', '$244.147'],
            ['Technical Training', '$50.688', '$96.096', '$172.339', '$172.339', '$244.147'],
            ['Leadership Training', '$43.200', '$47.520', '$52.272', '$57.499', '$63.249'],
            ['TOTAL', '$550.080', '$1.008.480', '$1.775.664', '$1.780.891', '$2.504.721'],
        ],
        col_widths=[1.8, 0.9, 0.9, 0.9, 0.9, 0.9]
    )

    # 3.2.3
    doc.add_heading('3.2.3. Praça — Estratégia de Distribuição', level=3)
    add_paragraph_bold_inline(doc,
        'A estratégia de distribuição combina **presença física** em três localidades '
        'estratégicas com **canais digitais** para maximizar alcance geográfico.')

    add_table_subtitle(doc, 'Tabela 3.2.3: Canais de Distribuição')
    create_table(doc,
        ['Canal', 'Descrição', 'Cobertura'],
        [
            ['Escritório Presencial', 'Jacksonville (Y1), Savannah (Y2), Brunswick (Y5)', 'FL, GA, SC'],
            ['Website + SEO', 'Presença digital com conteúdo técnico otimizado', 'Nacional'],
            ['LinkedIn B2B', 'Prospecção ativa e marketing de conteúdo', 'Sudeste EUA'],
            ['Câmaras de Comércio', 'Networking institucional e referências', 'Regional'],
            ['Conferências Setoriais', 'Participação como speaker/expositor', 'Nacional'],
        ],
        col_widths=[1.5, 2.5, 2.2]
    )

    add_paragraph_bold_inline(doc,
        'A expansão geográfica segue corredor logístico **I-95**: **Jacksonville** (sede, Y1), '
        '**Savannah** (Branch 1, Y2-Y3) e **Brunswick** (Branch 2, Y5), cobrindo o principal '
        'eixo de transporte do Sudeste americano.')

    # 3.2.4
    doc.add_heading('3.2.4. Promoção — Orçamento de Marketing', level=3)
    add_paragraph_bold_inline(doc,
        'O orçamento de marketing representa **13,5%** da receita bruta, alocado em '
        'categorias estratégicas para maximizar geração de leads B2B qualificados.')

    add_table_subtitle(doc, 'Tabela 3.2.4: Alocação Orçamentária de Marketing')
    create_table(doc,
        ['Categoria', '% do Orçamento', 'Atividades Principais'],
        [
            ['Marketing Digital (SEO/SEM)', '28%', 'Website, Google Ads, conteúdo técnico, blog'],
            ['Eventos e Conferências', '22%', 'APICS, supply chain summits, workshops regionais'],
            ['Relações Públicas e Advocacy', '15%', 'Artigos, white papers, posicionamento como thought leader'],
            ['Marketing de Relacionamento', '10%', 'CRM, follow-up personalizado, nurturing de leads'],
            ['Material e Identidade Visual', '10%', 'Branding, colateral de vendas, apresentações'],
            ['Ações Socioambientais', '15%', 'Projetos comunitários, parcerias com ONGs locais'],
        ],
        col_widths=[1.8, 1.0, 3.4]
    )

    add_paragraph_bold_inline(doc,
        'O investimento total em marketing e ações socioambientais soma **$1.028.678** ao longo '
        'de cinco anos, evoluindo de **$74.261** no Y1 para **$338.137** no Y5, acompanhando '
        'o crescimento da receita e expansão geográfica.')

    # 3.3 Marketing 4.0
    doc.add_heading('3.3. Estratégia de Marketing 4.0', level=2)
    add_paragraph_bold_inline(doc,
        'O **Marketing 4.0** integra canais digitais e tradicionais em jornada do cliente '
        'otimizada para conversão B2B, utilizando a expertise do fundador como principal '
        'ativo de atração e credibilidade.')

    add_table_subtitle(doc, 'Tabela 3.3a: Jornada do Cliente Digitalizada')
    create_table(doc,
        ['Fase', 'Canal', 'Ação', 'KPI'],
        [
            ['Awareness', 'SEO + LinkedIn', 'Conteúdo técnico sobre supply chain e ERP', 'Impressões / Alcance'],
            ['Consideration', 'Webinars + White Papers', 'Demonstração de expertise com cases reais', 'Downloads / Registros'],
            ['Decision', 'Proposta Personalizada', 'Diagnóstico gratuito + proposta de valor', 'Taxa de Conversão'],
            ['Loyalty', 'CRM + Follow-up', 'Acompanhamento pós-projeto + cross-sell', 'NPS / Retenção'],
        ],
        col_widths=[1.2, 1.5, 2.2, 1.3]
    )

    add_table_subtitle(doc, 'Tabela 3.3b: Estratégia de Conteúdo Digital')
    create_table(doc,
        ['Tipo de Conteúdo', 'Frequência', 'Plataforma', 'Objetivo'],
        [
            ['Artigos Técnicos', 'Semanal', 'Website/Blog', 'SEO e autoridade técnica'],
            ['Posts LinkedIn', '3x/semana', 'LinkedIn', 'Engajamento B2B e networking'],
            ['Webinars', 'Mensal', 'Zoom/Teams', 'Geração de leads qualificados'],
            ['Case Studies', 'Trimestral', 'Website', 'Prova social e resultados documentados'],
            ['Newsletter B2B', 'Quinzenal', 'Email', 'Nurturing e relacionamento'],
        ],
        col_widths=[1.5, 1.0, 1.2, 2.5]
    )

    add_table_subtitle(doc, 'Tabela 3.3c: Marketing de Influência Técnica')
    create_table(doc,
        ['Estratégia', 'Parceiro/Canal', 'Ação', 'Resultado Esperado'],
        [
            ['Thought Leadership', 'APICS/ASCM', 'Publicações e palestras', 'Credibilidade setorial'],
            ['Parcerias Acadêmicas', 'UNF, JU', 'Guest lectures, pesquisa conjunta', 'Pipeline de talentos'],
            ['Câmaras de Comércio', 'JAX Chamber, Brazil-FL', 'Networking e referências', 'Leads qualificados'],
            ['Associações Profissionais', 'ISM, CSCMP', 'Memberships e eventos', 'Rede de contatos B2B'],
        ],
        col_widths=[1.5, 1.3, 1.7, 1.7]
    )

    add_table_subtitle(doc, 'Tabela 3.3d: Automação e Ferramentas Digitais')
    create_table(doc,
        ['Ferramenta', 'Função', 'Integração'],
        [
            ['HubSpot CRM', 'Gestão de leads e pipeline comercial', 'Website + LinkedIn + Email'],
            ['Google Analytics', 'Métricas de tráfego e conversão', 'Website + Google Ads'],
            ['SEMrush', 'SEO e análise competitiva digital', 'Blog + Conteúdo técnico'],
            ['Mailchimp', 'Email marketing e newsletters B2B', 'CRM + Landing pages'],
        ],
        col_widths=[1.5, 2.0, 2.7]
    )

    add_table_subtitle(doc, 'Tabela 3.3e: Métricas de Performance Digital (KPIs)')
    create_table(doc,
        ['KPI', 'Y1 (Meta)', 'Y3 (Meta)', 'Y5 (Meta)'],
        [
            ['Tráfego Website (visitas/mês)', '2.000', '8.000', '20.000'],
            ['Leads Qualificados/mês', '15', '40', '80'],
            ['Taxa de Conversão', '5%', '8%', '12%'],
            ['CAC (Customer Acquisition Cost)', '$3.500', '$2.500', '$1.800'],
            ['NPS (Net Promoter Score)', '60', '75', '85'],
        ],
        col_widths=[2.0, 1.3, 1.3, 1.3]
    )

    add_checkmark(doc, 'Expertise como Ativo de Marketing', '25 anos AmBev/AB InBev aplicados a supply chain e resiliência operacional')
    add_checkmark(doc, 'Presença Digital Fortalecida', 'SEO, marketing de influência técnica e engajamento B2B via LinkedIn')
    add_checkmark(doc, 'Engajamento e Autoridade', 'conteúdo educativo, webinars técnicos e prova social como ferramentas-chave')


# ============================================================
# SECTION 4: OPERATIONAL PLAN
# ============================================================

def build_section_4(doc):
    """4. OPERATIONAL PLAN"""
    doc.add_page_break()
    add_separator(doc)
    doc.add_heading('4. OPERATIONAL PLAN', level=1)
    add_paragraph_bold_inline(doc,
        'O plano operacional detalha a **infraestrutura física**, **recursos humanos**, '
        '**tecnologia** e **capacidade produtiva** necessários para operacionalização '
        'dos seis serviços especializados da Vieira Operations LLC.')

    # 4.1
    doc.add_heading('4.1. Layout do Empreendimento', level=2)
    add_paragraph_bold_inline(doc,
        'O dimensionamento espacial da **Vieira Operations LLC** segue metodologia de '
        'planejamento modular baseada em benchmarks do **U.S. General Services Administration** '
        '(GSA), que recomenda **150-200 sq ft** por colaborador para escritórios profissionais.')
    add_paragraph_bold_inline(doc,
        'A estrutura distribui-se em **três fases** de crescimento, cada uma alinhada à abertura de uma nova unidade operacional conforme o cronograma de expansão geográfica definido no plano estratégico.')

    add_table_subtitle(doc, 'Tabela 4.1: Fases de Expansão — Layout por Localidade')
    create_table(doc,
        ['Fase', 'Período', 'Localidade', 'Área', 'Colaboradores', 'Aluguel/Mês'],
        [
            ['Fase I', 'Y1-Y2', 'Jacksonville, FL (Sede)', '800-1.000 sq ft', '5-6', '$2.450'],
            ['Fase II', 'Y2-Y3', 'Savannah, GA (Branch 1)', '600-800 sq ft', '3-7', '$2.450'],
            ['Fase III', 'Y5', 'Brunswick, GA (Branch 2)', '600-800 sq ft', '3-6', '$2.450'],
        ],
        col_widths=[0.7, 0.7, 1.7, 1.0, 1.0, 1.1]
    )

    add_paragraph_bold_inline(doc,
        'Cada escritório inclui: **sala de reuniões** com equipamento de videoconferência, '
        '**estações de trabalho individuais**, **área de recepção** e **infraestrutura de TI** '
        'completa para operações de consultoria remota e presencial.')

    # 4.2 Recursos Físicos
    doc.add_heading('4.2. Recursos Físicos e Equipamentos', level=2)
    add_paragraph_bold_inline(doc,
        'O investimento em **ativos tangíveis** totaliza **$65.975** ao longo de cinco anos, '
        'distribuído em três ondas de aquisição alinhadas à abertura de cada escritório.')

    add_table_subtitle(doc, 'Tabela 4.2: Investimentos em Equipamentos por Fase')
    create_table(doc,
        ['Item', 'Y0 (Sede)', 'Y2 (Branch 1)', 'Y4 (Branch 2)', 'Total'],
        [
            ['Microcomputadores', '$4.500 (5 un.)', '$10.800 (12 un.)', '$10.800 (12 un.)', '$26.100'],
            ['Laptops', '$2.400 (2 un.)', '$1.200 (1 un.)', '$1.200 (1 un.)', '$4.800'],
            ['Mobiliário (mesas/cadeiras)', '$1.125', '$2.925', '$2.925', '$6.975'],
            ['Mesa de Conferência', '$1.200', '$1.200', '$1.200', '$3.600'],
            ['Telefonia', '$850', '$850', '$850', '$2.550'],
            ['Impressora + TV + Frigobar', '$1.650', '$1.100', '$1.100', '$3.850'],
            ['Celulares', '$1.650 (3 un.)', '$3.300 (6 un.)', '$3.300 (6 un.)', '$8.250'],
            ['Ar Condicionado', '$1.800 (2 un.)', '$1.800 (2 un.)', '$1.800 (2 un.)', '$5.400'],
            ['Armários e Arquivos', '$750', '$1.250', '$1.250', '$3.250'],
            ['Poltronas', '$400', '$400', '$400', '$1.200'],
            ['TOTAL', '$16.325', '$24.825', '$24.825', '$65.975'],
        ],
        col_widths=[1.7, 1.2, 1.2, 1.2, 0.9]
    )

    # 4.3 Quadro de Funcionários
    doc.add_heading('4.3. Quadro de Funcionários', level=2)
    add_paragraph_bold_inline(doc,
        'A estrutura de pessoal projeta crescimento de **5 para 22 colaboradores** entre Y1 e Y5, '
        'distribuídos nas três unidades operacionais. A folha de pagamento inclui salários '
        'competitivos baseados em dados do **BLS Occupational Employment Statistics**.')

    add_table_subtitle(doc, 'Tabela 4.3: Quadro de Funcionários e Salários por Ano')
    create_table(doc,
        ['Cargo (SOC)', 'Salário/Ano', 'Y1', 'Y2', 'Y3', 'Y4', 'Y5'],
        [
            ['Supply Chain Manager — Owner (11-3031)', '$105.000', '1', '1', '1', '1', '1'],
            ['Project Management Specialist (13-1082)', '$53.180-$63.450', '1', '1', '2', '2', '5'],
            ['Office Clerks, General (43-9061)', '$40.150-$63.450', '1', '2', '4', '4', '6'],
            ['First-Line Supervisors Office (43-1011)', '$40.480-$60.820', '1', '2', '4', '4', '6'],
            ['Receptionists and Info Clerk (43-4171)', '$27.470-$35.840', '1', '1', '2', '2', '4'],
            ['TOTAL COLABORADORES', '', '5', '7', '13', '13', '22'],
            ['Folha Total (inclui payroll taxes)', '', '$286.620', '$474.340', '$735.220', '$735.220', '$1.082.540'],
        ],
        col_widths=[2.2, 1.2, 0.5, 0.5, 0.5, 0.5, 0.5]
    )

    add_paragraph_bold_inline(doc,
        'Os **payroll taxes** variam por localidade: Jacksonville, FL (**15,62%**), '
        'Savannah, GA (**20,34%**) e Brunswick, GA (**20,57%**), conforme calculadoras '
        'do **SmartAsset Paycheck Calculator**.')

    # 4.4
    doc.add_heading('4.4. Recursos Tecnológicos', level=2)
    add_paragraph_bold_inline(doc,
        'A infraestrutura tecnológica combina **plataformas de classe empresarial** com '
        'ferramentas de produtividade para suportar os seis pilares de prestação de serviços.')

    add_table_subtitle(doc, 'Tabela 4.4: Plataformas Tecnológicas por Pilar de Serviço')
    create_table(doc,
        ['Pilar de Serviço', 'Plataformas Principais', 'Investimento/Ano'],
        [
            ['Supply Chain', 'SAP Fiori, Blue Yonder, Oracle SCM Cloud', 'Licenças incluídas no Software ($6.000/Y)'],
            ['ERP Implementation', 'SAP S/4HANA, SAP Business One, Microsoft Dynamics', 'Certificações + Sandbox'],
            ['Project Management', 'Microsoft Project, Jira, Monday.com', 'Suite incluída em Software'],
            ['Compliance', 'RegTech platforms, OSHA compliance tools', 'Suite incluída em Software'],
            ['Training Online', 'Zoom, Moodle/Canvas, Assessment tools', 'Plataforma LMS incluída'],
            ['Infraestrutura', 'AWS/Azure cloud, Microsoft 365, CRM HubSpot', 'Cloud + Produtividade'],
        ],
        col_widths=[1.3, 2.5, 2.4]
    )

    add_paragraph_bold_inline(doc,
        'O investimento anual em **licenças de software** é de **$6.000** por localidade, '
        'totalizando **$18.000** no Y5 com três escritórios operacionais.')

    # 4.5
    doc.add_heading('4.5. Localização do Negócio', level=2)
    add_paragraph_bold_inline(doc,
        'A seleção de **Jacksonville, Florida**, como sede operacional fundamenta-se em '
        'convergência de fatores logísticos, econômicos e regulatórios que maximizam '
        'acesso ao mercado-alvo.')

    add_table_subtitle(doc, 'Tabela 4.5: Perfil Econômico — Condado de Duval (Jacksonville)')
    create_table(doc,
        ['Indicador', 'Valor', 'Fonte'],
        [
            ['População Metropolitana', '1,6 milhão', 'U.S. Census Bureau 2023'],
            ['Força de Trabalho', '859.227', 'BLS Local Area Unemployment Statistics'],
            ['Renda Mediana Familiar', '$62.751', 'American Community Survey 2023'],
            ['Taxa de Pobreza', '13,7%', 'Census Bureau (acima da média nacional 11,5%)'],
            ['Empresas Manufatureiras', '2.300+', 'Florida DEO, QCEW'],
            ['Porto Estratégico', 'JAXPORT (DoD Strategic Port)', 'Maritime Administration'],
        ],
        col_widths=[1.8, 1.3, 3.1]
    )

    add_paragraph_bold_inline(doc,
        'A localização converge com a **Executive Order 14017** e o **CHIPS Act**, que '
        'priorizam fortalecimento de cadeias produtivas em regiões com infraestrutura '
        'portuária e base industrial existente.')

    # 4.6
    doc.add_heading('4.6. Capacidade Produtiva', level=2)
    add_paragraph_bold_inline(doc,
        'A capacidade operacional projeta crescimento de **164%** entre Y1 e Y5, de '
        '**$550.080 para $2.504.721** em receita anual. O modelo fundamenta-se em '
        '**capital intelectual** como ativo produtivo primário.')

    add_table_subtitle(doc, 'Tabela 4.6: Capacidade Produtiva por Ano')
    create_table(doc,
        ['Métrica', 'Y1', 'Y2', 'Y3', 'Y4', 'Y5'],
        [
            ['Receita Projetada', '$550.080', '$1.008.480', '$1.775.664', '$1.780.891', '$2.504.721'],
            ['Colaboradores', '5', '7', '13', '13', '22'],
            ['Receita/Colaborador', '$110.016', '$144.069', '$136.590', '$137.007', '$113.851'],
            ['Escritórios Ativos', '1', '2', '3', '3', '3'],
            ['Horas Produtivas/Mês', '880', '1.232', '2.288', '2.288', '3.872'],
            ['Utilização Projetada', '65%', '75%', '80%', '82%', '85%'],
        ],
        col_widths=[1.5, 0.9, 0.9, 0.9, 0.9, 0.9]
    )

    add_paragraph_bold_inline(doc,
        'O crescimento de capacidade é sustentado por: **(a)** expansão geográfica para '
        'Savannah e Brunswick, **(b)** contratação de especialistas em cada localidade, '
        'e **(c)** alavancagem tecnológica através de ferramentas de colaboração remota.')

    add_checkmark(doc, 'Infraestrutura Completa', 'três escritórios estrategicamente posicionados no corredor I-95')
    add_checkmark(doc, 'Equipe Qualificada', '22 colaboradores projetados no Y5 com salários competitivos BLS')


# ============================================================
# SECTION 5: FINANCIAL PLAN
# ============================================================

def build_section_5(doc):
    """5. FINANCIAL PLAN"""
    doc.add_page_break()
    add_separator(doc)
    doc.add_heading('5. FINANCIAL PLAN', level=1)
    add_paragraph_bold_inline(doc,
        'O plano financeiro apresenta as **premissas**, **investimentos**, **projeções de receita**, '
        '**DRE**, **indicadores de retorno** e **break-even point** fundamentados nos dados '
        'da planilha financeira e benchmarks setoriais.')

    # 5.1
    doc.add_heading('5.1. Premissas Financeiras', level=2)
    add_paragraph_bold_inline(doc,
        'As projeções financeiras fundamentam-se em premissas verificáveis derivadas de '
        'dados do **U.S. Census Bureau**, **Bureau of Labor Statistics**, **SBA** e '
        'histórico operacional de **25 anos** do fundador na AmBev/AB InBev.')

    add_table_subtitle(doc, 'Tabela 5.1: Premissas Financeiras Críticas')
    create_table(doc,
        ['Premissa', 'Valor', 'Fonte/Justificativa'],
        [
            ['Receita Bruta Acumulada (5 anos)', '$7.619.836', 'Pipeline B2B identificado + CAGR setorial'],
            ['Custos Variáveis (% receita)', '19,7%', 'Média ponderada das 6 linhas de serviço'],
            ['Margem de Contribuição Média', '80,2%', 'Modelo asset-light de consultoria especializada'],
            ['Taxa de Desconto (WACC)', '10-12% a.a.', 'Média para setor NAICS 541611'],
            ['Taxa de Conversão de Leads', '18-22%', 'Benchmark serviços B2B de alto valor (SBA)'],
            ['Crescimento Y1→Y2', '83,3%', 'Expansão para Savannah + maturação de portfólio'],
            ['Crescimento Y4→Y5', '40,6%', 'Abertura Brunswick + escala da operação'],
            ['Break-Even', 'Mês 6', 'BEP calculado com margem de contribuição 78%'],
            ['Payroll Tax Rate (FL)', '15,62%', 'SmartAsset Paycheck Calculator — Jacksonville'],
            ['Financiamento Taxa Anual', '12%', 'Federal Reserve H.8 — SBA lending rates'],
        ],
        col_widths=[2.0, 1.2, 3.0]
    )

    add_paragraph_bold_inline(doc,
        'O modelo projeta trajetória de crescimento sustentado com CAGR de **35,4%** na '
        'receita bruta ao longo do quinquênio, sustentado pela expansão geográfica e '
        'maturação do portfólio de serviços.')

    # 5.2
    doc.add_heading('5.2. Investimentos', level=2)
    add_paragraph_bold_inline(doc,
        'Os investimentos totais de **$382.917** distribuem-se em três ondas estratégicas '
        'ao longo de cinco anos, alinhadas à abertura de cada unidade operacional.')

    add_table_subtitle(doc, 'Tabela 5.2a: Investimentos Totais por Categoria e Ano')
    create_table(doc,
        ['Categoria', 'Y0 (Sede)', 'Y2 (Branch 1)', 'Y4 (Branch 2)', 'Total'],
        [
            ['Bens Tangíveis (Depreciáveis)', '$16.325', '$24.825', '$24.825', '$65.975'],
            ['Bens Intangíveis (Não Depreciáveis)', '$105.647', '$105.647', '$105.647', '$316.942'],
            ['— Marketing e Consultoria TI', '$1.500', '$1.500', '$1.500', '$4.500'],
            ['— Identidade Visual', '$1.200', '$1.200', '$1.200', '$3.600'],
            ['— Licenças de Software', '$6.000', '$6.000', '$6.000', '$18.000'],
            ['— Taxas de Abertura', '$600', '$600', '$600', '$1.800'],
            ['— Website', '$700', '$700', '$700', '$2.100'],
            ['— Projeto Arquitetônico', '$3.000', '$3.000', '$3.000', '$9.000'],
            ['— Registro de Marca', '$1.000', '$1.000', '$1.000', '$3.000'],
            ['— Capital de Giro (3 meses)', '$90.747', '$90.747', '$90.747', '$272.242'],
            ['TOTAL', '$121.972', '$130.472', '$130.472', '$382.917'],
        ],
        col_widths=[2.2, 1.1, 1.1, 1.1, 0.9]
    )

    add_paragraph_bold_inline(doc,
        'O **capital de giro** de 3 meses ($90.747 por onda) garante cobertura de despesas '
        'operacionais durante a fase de ramp-up de cada localidade, incluindo folha de '
        'pagamento, aluguel e custos fixos.')

    doc.add_heading('Estrutura de Financiamento', level=3)
    add_paragraph_bold_inline(doc,
        'A capitalização combina **equity do fundador** com **linha de crédito bancária**, estruturada para minimizar risco financeiro e manter liquidez operacional durante as fases de expansão.')
    add_bullet(doc, 'Capital (Y0)', '$121.972 — financiado a 12% a.a. em 5 anos (PMT: $2.676/mês)')
    add_bullet(doc, 'Amortização Mensal', '$2.033 + juros de $643 = parcela total de $2.676')

    # 5.3
    doc.add_heading('5.3. Estimativa de Receitas e Custos', level=2)
    add_paragraph_bold_inline(doc,
        'A receita bruta acumulada de **$7.619.836** em cinco anos distribui-se entre '
        'seis linhas de serviço com crescimento orgânico impulsionado pela expansão '
        'geográfica e maturação do portfólio.')

    add_table_subtitle(doc, 'Tabela 5.3a: Receita Bruta por Serviço (Y1-Y5)')
    create_table(doc,
        ['Serviço', 'Y1', 'Y2', 'Y3', 'Y4', 'Y5', 'Total'],
        [
            ['Supply Chain Optim.', '$177.408', '$336.336', '$603.187', '$603.187', '$854.515', '$2.574.634'],
            ['ERP Implementation', '$177.408', '$336.336', '$603.187', '$603.187', '$854.515', '$2.574.634'],
            ['Emergency Project', '$50.688', '$96.096', '$172.339', '$172.339', '$244.147', '$735.610'],
            ['Regulatory Compliance', '$50.688', '$96.096', '$172.339', '$172.339', '$244.147', '$735.610'],
            ['Technical Training', '$50.688', '$96.096', '$172.339', '$172.339', '$244.147', '$735.610'],
            ['Leadership Training', '$43.200', '$47.520', '$52.272', '$57.499', '$63.249', '$263.740'],
            ['TOTAL', '$550.080', '$1.008.480', '$1.775.664', '$1.780.891', '$2.504.721', '$7.619.836'],
        ],
        col_widths=[1.3, 0.7, 0.8, 0.8, 0.8, 0.8, 0.8]
    )

    add_table_subtitle(doc, 'Tabela 5.3b: Custos Variáveis (Y1-Y5)')
    create_table(doc,
        ['Custo Variável', 'Y1', 'Y2', 'Y3', 'Y4', 'Y5', 'Total'],
        [
            ['Sales Budget (5%)', '$22.810', '$43.243', '$77.553', '$77.553', '$109.866', '$331.024'],
            ['Course Supplies', '$6.480', '$7.128', '$7.841', '$8.625', '$9.487', '$39.561'],
            ['Commissioned Prof.', '$17.280', '$19.008', '$20.909', '$23.000', '$25.300', '$105.496'],
            ['Marketing + Socioamb.', '$74.261', '$136.145', '$239.715', '$240.420', '$338.137', '$1.028.678'],
            ['TOTAL', '$120.830', '$205.524', '$346.017', '$349.598', '$482.791', '$1.504.759'],
            ['% da Receita', '22,0%', '20,4%', '19,5%', '19,6%', '19,3%', '19,7%'],
        ],
        col_widths=[1.3, 0.7, 0.8, 0.8, 0.8, 0.8, 0.8]
    )

    # 5.4
    doc.add_heading('5.4. DRE — Demonstrativo de Resultados', level=2)
    add_paragraph_bold_inline(doc,
        'O **DRE** documenta trajetória de **resultado negativo no Y1** ($-13.882) seguido '
        'de **recuperação no Y2** e **lucratividade crescente** a partir do Y3, demonstrando '
        'viabilidade financeira robusta ao longo do quinquênio.')

    add_table_subtitle(doc, 'Tabela 5.4: Demonstrativo de Resultados Projetado (Y0-Y5)')
    create_table(doc,
        ['Linha', 'Y0', 'Y1', 'Y2', 'Y3', 'Y4', 'Y5'],
        [
            ['Receita Bruta', '$0', '$550.080', '$1.008.480', '$1.775.664', '$1.780.891', '$2.504.721'],
            ['(-) Custos Variáveis', '$0', '$120.830', '$205.524', '$346.017', '$349.598', '$482.791'],
            ['(=) Margem Contribuição', '$0', '$429.250', '$802.956', '$1.429.647', '$1.431.294', '$2.021.931'],
            ['(-) Investimentos', '$121.972', '$0', '$130.472', '$0', '$130.472', '$0'],
            ['(-) Despesas Operacionais', '$0', '$362.990', '$515.252', '$826.811', '$826.811', '$1.208.912'],
            ['(=) EBITDA', '-$121.972', '$66.260', '$157.231', '$602.836', '$474.010', '$813.019'],
            ['(-) ITDA', '$0', '$80.142', '$115.599', '$171.825', '$171.825', '$243.950'],
            ['(=) Lucro Bruto', '-$121.972', '-$13.882', '$41.632', '$431.011', '$302.185', '$569.069'],
            ['(-) Income Taxes', '$0', '$0', '$9.497', '$102.934', '$136.891', '$307.305'],
            ['(=) Lucro Líquido', '-$121.972', '-$13.882', '$32.135', '$328.077', '$165.294', '$261.764'],
            ['Lucro Líq. Acumulado', '-$121.972', '-$135.855', '-$103.720', '$224.357', '$389.651', '$651.415'],
        ],
        col_widths=[1.4, 0.7, 0.7, 0.8, 0.8, 0.8, 0.8]
    )

    add_paragraph_bold_inline(doc,
        'A **margem EBITDA** evolui de **12,0%** no Y1 para **32,5%** no Y5, evidenciando '
        'ganho de escala operacional. O **lucro líquido acumulado** atinge **$651.415** ao '
        'final do quinquênio.')

    # 5.5
    doc.add_heading('5.5. Indicadores de Retorno', level=2)
    add_paragraph_bold_inline(doc,
        'Os indicadores de retorno posicionam a **Vieira Operations LLC** no '
        '**quartil superior de atratividade** quando comparada a benchmarks setoriais '
        'do NAICS 541611.')

    add_table_subtitle(doc, 'Tabela 5.5: Indicadores Consolidados de Retorno Financeiro')
    create_table(doc,
        ['Indicador', 'Valor', 'Interpretação'],
        [
            ['NPV (5 anos, taxa 12%)', '$378.348', 'Positivo: projeto viável e atrativo'],
            ['NPV com Perpetuidade', '$1.089.232', 'Valor do negócio considerando fluxos futuros'],
            ['Valor do Negócio', '$1.467.581', 'NPV + Perpetuidade (business valuation)'],
            ['IRR (Taxa Interna de Retorno)', '64,3%', 'Muito superior ao custo de capital de 12%'],
            ['ROI (Return on Investment)', '170,1%', 'Para cada $1 investido, retorno de $1,70 líquido'],
            ['Lucro Acumulado (5 anos)', '$651.415', 'Total após impostos e todas as despesas'],
            ['Lucratividade Média', '8,5%', 'Consistente com benchmarks do setor'],
            ['Payback Descontado', '3 anos', 'Recuperação do investimento com valor temporal'],
        ],
        col_widths=[1.8, 1.2, 3.2]
    )

    add_checkmark(doc, 'Viabilidade Financeira', 'NPV positivo de $378.348 com IRR de 64,3% — amplamente superior ao custo de capital')
    add_checkmark(doc, 'Retorno Robusto', 'ROI de 170,1% com payback descontado de 3 anos')

    # 5.6
    doc.add_heading('5.6. Break Even Point', level=2)
    add_paragraph_bold_inline(doc,
        'O **Ponto de Equilíbrio (BEP)** determina o volume mínimo de receita necessário '
        'para cobrir todos os custos fixos e variáveis em cada ano operacional.')

    add_table_subtitle(doc, 'Tabela 5.6: Break Even Point por Ano')
    create_table(doc,
        ['Componente', 'Y1', 'Y2', 'Y3', 'Y4', 'Y5'],
        [
            ['Receita', '$550.080', '$1.008.480', '$1.775.664', '$1.780.891', '$2.504.721'],
            ['Custos Fixos Totais', '$443.132', '$630.852', '$998.637', '$998.637', '$1.452.862'],
            ['Margem Contribuição (%)', '78,0%', '79,6%', '80,5%', '80,4%', '80,7%'],
            ['BEP (Receita)', '$567.870', '$792.324', '$1.240.336', '$1.242.556', '$1.800.040'],
            ['Margem de Segurança', '-3,2%', '21,4%', '30,1%', '30,2%', '28,1%'],
        ],
        col_widths=[1.5, 0.9, 0.9, 0.9, 0.9, 0.9]
    )

    add_paragraph_bold_inline(doc,
        'No **Y1**, a receita ($550.080) fica marginalmente abaixo do BEP ($567.870), '
        'resultando em prejuízo controlado de **-$13.882** (deficit de apenas 3,2%). '
        'A partir do **Y2**, a margem de segurança torna-se positiva e crescente.')
    add_paragraph_bold_inline(doc,
        'A análise incorpora cenário **pessimista** (redução de 15% receita, aumento 10% custos) '
        'que indica break-even no **Mês 8** — ainda dentro do primeiro ano operacional, '
        'conforme recomendações do **SBA** para empresas em fase de implantação.')

    add_checkmark(doc, 'Break-Even Rápido', 'atingido no Mês 6 do cenário base (Mês 8 no pessimista)')
    add_checkmark(doc, 'Margem de Segurança', 'crescente a partir do Y2, atingindo 30%+ no Y3-Y4')


# ============================================================
# SECTION 6: FINAL CONSIDERATIONS
# ============================================================

def build_section_6(doc):
    """6. FINAL CONSIDERATIONS"""
    doc.add_page_break()
    add_separator(doc)
    doc.add_heading('6. FINAL CONSIDERATIONS', level=1)
    add_paragraph_bold_inline(doc,
        'Esta seção consolida a **timeline de implementação**, as **considerações finais** '
        'sobre a viabilidade do empreendimento e as **referências** utilizadas na elaboração '
        'deste Business Plan.')

    # 6.1
    doc.add_heading('6.1. Timeline de Implementação', level=2)
    add_paragraph_bold_inline(doc,
        'A implementação da **Vieira Operations LLC** estrutura-se em cronograma de '
        '**60 meses** distribuído em duas camadas temporais: **timeline detalhada** '
        '(12 meses iniciais) e **horizonte estratégico** (5 anos).')

    add_table_subtitle(doc, 'Tabela 6.1: Timeline Estratégica — Horizonte de 5 Anos')
    create_table(doc,
        ['Período', 'Marco', 'Receita Alvo', 'Empregos'],
        [
            ['Y0 (Pré-operacional)', 'Constituição legal, setup operacional, investimento $121.972', '$0', '0'],
            ['Y1 (Estabelecimento)', 'Sede Jacksonville, primeiros clientes, break-even Mês 6', '$550.080', '5'],
            ['Y2 (Expansão)', 'Branch 1 Savannah, crescimento 83% receita', '$1.008.480', '7'],
            ['Y3 (Consolidação)', 'Maturação de portfólio, 3 escritórios ativos', '$1.775.664', '13'],
            ['Y4 (Escala)', 'Branch 2 Brunswick, investimento $130.472', '$1.780.891', '13'],
            ['Y5 (Maturidade)', 'Operação plena em 3 localidades, 22 colaboradores', '$2.504.721', '22'],
        ],
        col_widths=[1.3, 2.5, 1.0, 0.8]
    )

    add_paragraph_bold_inline(doc,
        'O modelo alinha-se ao padrão do **U.S. Small Business Administration**, que documenta '
        'que **68%** das empresas de serviços B2B atingem estabilidade operacional dentro '
        'dos primeiros 18-24 meses de operação.')

    # 6.2
    doc.add_heading('6.2. Considerações Finais', level=2)
    add_paragraph_bold_inline(doc,
        'A **Vieira Operations LLC** apresenta plano de negócios robusto, fundamentado em '
        '**estratégias de diferenciação técnica** e alinhamento com prioridades nacionais '
        'de fortalecimento de cadeias produtivas dos Estados Unidos.')
    add_paragraph_bold_inline(doc,
        'A empresa projeta geração de **14 empregos diretos** e **62 empregos indiretos** '
        'ao longo de cinco anos, com impacto econômico concentrado em regiões do Sudeste '
        'americano que apresentam necessidade documentada de desenvolvimento.')
    add_paragraph_bold_inline(doc,
        'No campo competitivo, a empresa diferencia-se ao combinar competências **raramente '
        'encontradas em um único prestador**: expertise em supply chain, implementação ERP, '
        'conformidade regulatória e capacitação de força de trabalho.')
    add_paragraph_bold_inline(doc,
        'A viabilidade financeira é sustentada por **margem de contribuição de 80,2%**, '
        '**NPV positivo de $378.348**, **IRR de 64,3%** e **break-even no Mês 6**, '
        'indicadores que posicionam o empreendimento no quartil superior de atratividade '
        'para investimentos no setor de consultoria especializada.')
    add_paragraph_bold_inline(doc,
        'A expansão geográfica planejada — de **Jacksonville** para **Savannah** e **Brunswick** — '
        'forma corredor estratégico ao longo da I-95, maximizando acesso a clusters industriais, '
        'portos estratégicos e mercados-alvo de defesa, saúde e manufatura.')

    add_checkmark(doc, 'Impacto Econômico', '76 empregos projetados (14 diretos + 62 indiretos) no Sudeste dos EUA')
    add_checkmark(doc, 'Viabilidade Financeira', '$7,6 milhões em receita acumulada com lucratividade a partir do Y2')
    add_checkmark(doc, 'Alinhamento Nacional', 'convergência com EO 14017, CHIPS Act e prioridades de reshoring')
    add_checkmark(doc, 'Expertise Diferenciada', '25 anos AmBev/AB InBev com resultados documentados e verificáveis')

    # 6.3
    doc.add_heading('6.3. Referências e Fontes', level=2)
    refs = [
        'U.S. Census Bureau — Service Annual Survey 2023 (NAICS 541611)',
        'Bureau of Labor Statistics — Occupational Employment and Wage Statistics (OEWS) 2024',
        'Bureau of Labor Statistics — Occupational Outlook Handbook 2024-2034',
        'Economic Policy Institute — Employment Multipliers by Industry (NAICS 5416)',
        'U.S. Small Business Administration — Starting and Managing a Business',
        'Executive Order 14017 — America\'s Supply Chains (February 24, 2021)',
        'CHIPS and Science Act (August 9, 2022)',
        'Florida Division of Corporations — Sunbiz.org',
        'Jacksonville Port Authority — JAXPORT Annual Report',
        'Georgia Ports Authority — Port of Savannah Statistics',
        'Federal Reserve — H.8 Assets and Liabilities of Commercial Banks',
        'SmartAsset — Paycheck Calculator (Jacksonville, FL; Savannah, GA; Brunswick, GA)',
        'USCIS — Chapter 5: Advanced Degree or Exceptional Ability',
        '8 CFR 204.5 — Petitions for Employment-Based Immigrants',
        'OECD — Oslo Manual: Guidelines for Collecting and Interpreting Innovation Data',
        'Sustainability Accounting Standards Board (SASB) — Professional Services Standards',
        'United Nations — 17 Sustainable Development Goals (SDGs)',
        'O*NET Content Model — Occupational Information Network',
        'IBISWorld — Management Consulting in the US Industry Report',
        'SAP SE — SAP Fiori and S/4HANA Certification Documentation',
    ]

    for i, ref in enumerate(refs, 1):
        add_numbered_item(doc, i, ref.split(' — ')[0] if ' — ' in ref else ref.split(' (')[0],
                         ref.split(' — ')[1] if ' — ' in ref else (ref.split(' (')[1].rstrip(')') if ' (' in ref else ''))


# ============================================================
# MAIN
# ============================================================

def main():
    import docx

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    setup_styles(doc)

    print("Gerando Business Plan — Vieira Operations LLC...")
    print("=" * 60)

    print("  → Cover Page...")
    build_cover_page(doc)

    print("  → Table of Contents...")
    build_toc(doc)

    print("  → Section 1: Executive Summary...")
    build_section_1(doc)

    print("  → Section 2: Análise Estratégica de Mercado...")
    build_section_2(doc)

    print("  → Section 3: Marketing Plan...")
    build_section_3(doc)

    print("  → Section 4: Operational Plan...")
    build_section_4(doc)

    print("  → Section 5: Financial Plan...")
    build_section_5(doc)

    print("  → Section 6: Final Considerations...")
    build_section_6(doc)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"\n✅ Documento salvo em: {OUTPUT_DOCX}")
    print(f"   Total de parágrafos: {len(doc.paragraphs)}")
    print(f"   Total de tabelas: {len(doc.tables)}")
    print("=" * 60)


if __name__ == '__main__':
    main()
