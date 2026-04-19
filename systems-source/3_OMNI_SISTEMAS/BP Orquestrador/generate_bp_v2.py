#!/usr/bin/env python3
"""
Business Plan Generator V2 — Vieira Operations LLC
Complete rewrite with professional visual design matching Ikaro benchmark.
- Garamond throughout
- Table headers #E8E0D4
- Footer bar with CONFIDENTIAL
- Header with company name
- 55-65 pages target
- Balanced text/table/bullet mix
"""

import os
import subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
from copy import deepcopy

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "BP_Vieira_Operations_LLC_FINAL.docx")

# ============================================================
# COLORS
# ============================================================
COLOR_HEADER_BG = "E8E0D4"       # Beige for table headers
COLOR_FOOTER_BAR = "3B4A3A"      # Dark olive green for footer bar
COLOR_BORDER = "CCCCCC"          # Light gray for table borders
COLOR_H1 = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_H2 = RGBColor(0x2C, 0x2C, 0x54)
COLOR_H3 = RGBColor(0x33, 0x33, 0x66)
COLOR_H4 = RGBColor(0x55, 0x55, 0x77)
COLOR_CHECK = RGBColor(0x22, 0x8B, 0x22)

# ============================================================
# STYLE SETUP
# ============================================================

def setup_styles(doc):
    """Configure all document styles."""
    # Normal
    style = doc.styles['Normal']
    style.font.name = 'Garamond'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(3)
    # Set East Asian font fallback
    rpr = style.element.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond" w:cs="Garamond"/></w:rPr>')
        style.element.append(rpr)

    # H1: CAIXA ALTA, bold, 16pt
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Garamond'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_H1
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # H2: Title Case, bold, 13pt
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Garamond'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_H2
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # H3: Title Case, bold, 11pt
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Garamond'
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_H3
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # H4: Italic, 11pt
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Garamond'
    h4.font.size = Pt(11)
    h4.font.bold = False
    h4.font.italic = True
    h4.font.color.rgb = COLOR_H4
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(4)

    # Custom styles
    for name in ['CoverPage', 'BulletItem', 'CheckItem', 'TableSubtitle', 'NumberedItem', 'HighlightBox']:
        try:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except:
            pass

    # Footnote styles
    try:
        fnt = doc.styles.add_style('FootnoteText', WD_STYLE_TYPE.PARAGRAPH)
        fnt.font.name = 'Garamond'
        fnt.font.size = Pt(8)
        fnt.paragraph_format.space_after = Pt(2)
        fnt.paragraph_format.line_spacing = 1.0
    except:
        pass
    try:
        from docx.enum.style import WD_STYLE_TYPE as WST
        fnr = doc.styles.add_style('FootnoteReference', WST.CHARACTER)
        fnr.font.superscript = True
        fnr.font.size = Pt(8)
    except:
        pass

    # CoverPage
    cp = doc.styles['CoverPage']
    cp.font.name = 'Garamond'
    cp.font.size = Pt(11)
    cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(4)

    # BulletItem
    bullet = doc.styles['BulletItem']
    bullet.font.name = 'Garamond'
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.15
    bullet.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # CheckItem
    check = doc.styles['CheckItem']
    check.font.name = 'Garamond'
    check.font.size = Pt(11)
    check.paragraph_format.left_indent = Inches(0.25)
    check.paragraph_format.space_after = Pt(4)
    check.paragraph_format.line_spacing = 1.15

    # TableSubtitle
    ts = doc.styles['TableSubtitle']
    ts.font.name = 'Garamond'
    ts.font.size = Pt(10)
    ts.font.italic = True
    ts.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    ts.paragraph_format.space_before = Pt(8)
    ts.paragraph_format.space_after = Pt(4)
    ts.paragraph_format.left_indent = Inches(0)
    ts.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # NumberedItem
    ni = doc.styles['NumberedItem']
    ni.font.name = 'Garamond'
    ni.font.size = Pt(11)
    ni.paragraph_format.left_indent = Inches(0.5)
    ni.paragraph_format.first_line_indent = Inches(-0.25)
    ni.paragraph_format.space_after = Pt(4)
    ni.paragraph_format.line_spacing = 1.15
    ni.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # HighlightBox (for key conclusions with shaded background)
    hb = doc.styles['HighlightBox']
    hb.font.name = 'Garamond'
    hb.font.size = Pt(11)
    hb.font.bold = True
    hb.paragraph_format.space_before = Pt(6)
    hb.paragraph_format.space_after = Pt(6)
    hb.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ============================================================
# PAGE SETUP (margins, header, footer)
# ============================================================

def setup_page(doc):
    """Configure page layout with header and footer."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.14)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    _setup_header(section)
    _setup_footer(section)


def _setup_header(section):
    """Add company name to header right-aligned."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("VIEIRA OPERATIONS LLC")
    run.font.name = 'Garamond'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.small_caps = True
    # Add thin bottom border to header paragraph
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{COLOR_BORDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _setup_footer(section):
    """Add professional footer with colored bar and page numbers."""
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear default
    for p in footer.paragraphs:
        p.clear()

    # Line 1: CONFIDENTIAL with dark background
    p1 = footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr1 = p1._element.get_or_add_pPr()
    # Add shading (colored bar background)
    shd1 = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{COLOR_FOOTER_BAR}"/>')
    pPr1.append(shd1)
    # Set spacing
    spacing1 = parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:before="60"/>')
    pPr1.append(spacing1)
    # Add indent for padding effect
    ind1 = parse_xml(f'<w:ind {nsdecls("w")} w:left="72" w:right="72"/>')
    pPr1.append(ind1)

    run1 = p1.add_run("CONFIDENTIAL \u2014 DO NOT SHARE.")
    run1.font.name = 'Garamond'
    run1.font.size = Pt(7)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Line 2: Disclaimer + Page number
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr2 = p2._element.get_or_add_pPr()
    spacing2 = parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:before="40"/>')
    pPr2.append(spacing2)

    run2 = p2.add_run("This business plan contains proprietary information. "
                       "Unauthorized distribution or disclosure is prohibited without written consent.")
    run2.font.name = 'Garamond'
    run2.font.size = Pt(7)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Add tab and page number
    run_tab = p2.add_run("\t")
    run_tab.font.size = Pt(7)

    # Page X of Y using field codes
    run_page_label = p2.add_run("Page ")
    run_page_label.font.name = 'Garamond'
    run_page_label.font.size = Pt(7)
    run_page_label.font.bold = True
    run_page_label.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # PAGE field
    fld_begin = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="14"/><w:b/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>')
    fld_code = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="14"/><w:b/></w:rPr><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>')
    fld_sep = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="separate"/></w:r>')
    fld_end = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>')
    p2._element.append(fld_begin)
    p2._element.append(fld_code)
    p2._element.append(fld_sep)
    p2._element.append(fld_end)

    run_of = p2.add_run(" of ")
    run_of.font.name = 'Garamond'
    run_of.font.size = Pt(7)
    run_of.font.bold = True
    run_of.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # NUMPAGES field
    fld_begin2 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="14"/><w:b/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>')
    fld_code2 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/><w:sz w:val="14"/><w:b/></w:rPr><w:instrText xml:space="preserve"> NUMPAGES </w:instrText></w:r>')
    fld_sep2 = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="separate"/></w:r>')
    fld_end2 = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>')
    p2._element.append(fld_begin2)
    p2._element.append(fld_code2)
    p2._element.append(fld_sep2)
    p2._element.append(fld_end2)

    # Set right-aligned tab stop for page number
    tabs = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:pos="9360"/></w:tabs>')
    pPr2.append(tabs)


# ============================================================
# CONTENT HELPERS
# ============================================================

def p_text(doc, text, style='Normal'):
    """Add paragraph with **bold** markers parsed inline."""
    p = doc.add_paragraph(style=style)
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            if part:
                run = p.add_run(part)
                run.font.name = 'Garamond'
                run.font.size = Pt(11) if style != 'TableSubtitle' else Pt(10)
                if i % 2 == 1:
                    run.bold = True
    else:
        run = p.add_run(text)
        run.font.name = 'Garamond'
        run.font.size = Pt(11)
    return p


def bullet(doc, bold_lead, text):
    """Add bullet: bullet **Bold Lead**: text"""
    p = doc.add_paragraph(style='BulletItem')
    r1 = p.add_run('\u2022 ')
    r1.font.name = 'Garamond'; r1.font.size = Pt(11)
    r2 = p.add_run(bold_lead)
    r2.font.name = 'Garamond'; r2.font.size = Pt(11); r2.bold = True
    r3 = p.add_run(': ' + text)
    r3.font.name = 'Garamond'; r3.font.size = Pt(11)
    return p


def check(doc, bold_lead, text):
    """Add checkmark: check **Bold**: text"""
    p = doc.add_paragraph(style='CheckItem')
    r1 = p.add_run('\u2714 ')
    r1.font.name = 'Garamond'; r1.font.size = Pt(11); r1.font.color.rgb = COLOR_CHECK
    r2 = p.add_run(bold_lead)
    r2.font.name = 'Garamond'; r2.font.size = Pt(11); r2.bold = True
    r3 = p.add_run(': ' + text)
    r3.font.name = 'Garamond'; r3.font.size = Pt(11)
    return p


def num_item(doc, number, bold_lead, text):
    """Add numbered item: N. **Bold**: text"""
    p = doc.add_paragraph(style='NumberedItem')
    r1 = p.add_run(f'{number}. ')
    r1.font.name = 'Garamond'; r1.font.size = Pt(11)
    r2 = p.add_run(bold_lead)
    r2.font.name = 'Garamond'; r2.font.size = Pt(11); r2.bold = True
    r3 = p.add_run(': ' + text)
    r3.font.name = 'Garamond'; r3.font.size = Pt(11)
    return p


def table_sub(doc, text):
    """Add italic table subtitle."""
    p = doc.add_paragraph(style='TableSubtitle')
    run = p.add_run(text)
    run.font.name = 'Garamond'; run.font.size = Pt(10); run.italic = True
    return p


def highlight_box(doc, text):
    """Add a shaded highlight box paragraph for key conclusions."""
    p = doc.add_paragraph(style='HighlightBox')
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F5F0E8"/>')
    pPr.append(shd)
    # Add border
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="4" w:color="B8A88A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            if part:
                run = p.add_run(part)
                run.font.name = 'Garamond'; run.font.size = Pt(11)
                if i % 2 == 1:
                    run.bold = True
    else:
        run = p.add_run(text)
        run.font.name = 'Garamond'; run.font.size = Pt(11); run.bold = True
    return p


def make_table(doc, headers, rows, col_widths=None):
    """Create table with #E8E0D4 header, #CCCCCC borders, Garamond 9pt."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = 'Garamond'; run.font.size = Pt(9); run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{COLOR_HEADER_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        # Cell padding
        _set_cell_padding(cell, top=40, bottom=40, left=80, right=80)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            # Support bold markers in cell text
            text = str(cell_text)
            if '**' in text:
                parts = text.split('**')
                for pi, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.name = 'Garamond'; run.font.size = Pt(9)
                        if pi % 2 == 1:
                            run.bold = True
            else:
                run = p.add_run(text)
                run.font.name = 'Garamond'; run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_cell_padding(cell, top=40, bottom=40, left=80, right=80)

    # Full width — match page margins (8.5" - 1.14" - 0.85" = 6.51")
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tbl_width)

    # Borders
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)
    return table


def _set_cell_padding(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins/padding in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def separator(doc):
    """No-op — decorative separators removed per design review."""
    pass


# ============================================================
# FOOTNOTES
# ============================================================

_fn_counter = 0
_fn_part = None

def _ensure_footnotes_part(doc):
    """Get or create the footnotes XML part."""
    global _fn_part
    if _fn_part is not None:
        return _fn_part
    # Check existing relationships
    for rel in doc.part.rels.values():
        if 'footnotes' in rel.reltype:
            _fn_part = rel.target_part
            return _fn_part
    # Create new footnotes part
    from docx.opc.part import Part as OpcPart
    from docx.opc.packuri import PackURI
    xml_bytes = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<w:footnote w:type="separator" w:id="-1">'
        '<w:p><w:r><w:separator/></w:r></w:p>'
        '</w:footnote>'
        '<w:footnote w:type="continuationSeparator" w:id="0">'
        '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
        '</w:footnote>'
        '</w:footnotes>'
    ).encode('utf-8')
    _fn_part = OpcPart(
        PackURI('/word/footnotes.xml'),
        'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml',
        xml_bytes,
        doc.part.package,
    )
    doc.part.relate_to(
        _fn_part,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes',
    )
    return _fn_part


def add_footnote(doc, paragraph, text):
    """Add a Word footnote to a paragraph. Returns the footnote ID."""
    global _fn_counter
    _fn_counter += 1
    fn_id = _fn_counter

    fn_part = _ensure_footnotes_part(doc)
    fn_xml = etree.fromstring(fn_part.blob)
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Build footnote element
    footnote = etree.SubElement(fn_xml, f'{{{W}}}footnote')
    footnote.set(f'{{{W}}}id', str(fn_id))

    fn_p = etree.SubElement(footnote, f'{{{W}}}p')
    fn_pPr = etree.SubElement(fn_p, f'{{{W}}}pPr')
    fn_pStyle = etree.SubElement(fn_pPr, f'{{{W}}}pStyle')
    fn_pStyle.set(f'{{{W}}}val', 'FootnoteText')

    # Footnote self-reference
    r1 = etree.SubElement(fn_p, f'{{{W}}}r')
    r1_rPr = etree.SubElement(r1, f'{{{W}}}rPr')
    r1_rStyle = etree.SubElement(r1_rPr, f'{{{W}}}rStyle')
    r1_rStyle.set(f'{{{W}}}val', 'FootnoteReference')
    etree.SubElement(r1, f'{{{W}}}footnoteRef')

    # Space
    r2 = etree.SubElement(fn_p, f'{{{W}}}r')
    t2 = etree.SubElement(r2, f'{{{W}}}t')
    t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t2.text = ' '

    # Footnote text
    r3 = etree.SubElement(fn_p, f'{{{W}}}r')
    r3_rPr = etree.SubElement(r3, f'{{{W}}}rPr')
    r3_sz = etree.SubElement(r3_rPr, f'{{{W}}}sz')
    r3_sz.set(f'{{{W}}}val', '16')
    r3_fonts = etree.SubElement(r3_rPr, f'{{{W}}}rFonts')
    r3_fonts.set(f'{{{W}}}ascii', 'Garamond')
    r3_fonts.set(f'{{{W}}}hAnsi', 'Garamond')
    t3 = etree.SubElement(r3, f'{{{W}}}t')
    t3.text = text

    # Update part blob
    fn_part._blob = etree.tostring(fn_xml, xml_declaration=True,
                                    encoding='UTF-8', standalone=True)

    # Add footnote reference in body paragraph
    ref_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        f'  <w:footnoteReference w:id="{fn_id}"/>'
        f'</w:r>'
    )
    paragraph._element.append(ref_run)
    return fn_id


def page_break(doc):
    """Add page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ============================================================
# COVER PAGE
# ============================================================

def build_cover(doc):
    """Professional cover page."""
    for _ in range(6):
        p = doc.add_paragraph(style='CoverPage')
        p.add_run('')

    p = doc.add_paragraph(style='CoverPage')
    run = p.add_run('VIEIRA OPERATIONS LLC')
    run.font.name = 'Garamond'; run.font.size = Pt(28); run.bold = True
    run.font.color.rgb = COLOR_H1

    p = doc.add_paragraph(style='CoverPage')
    run = p.add_run('BUSINESS PLAN')
    run.font.name = 'Garamond'; run.font.size = Pt(22); run.bold = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph(style='CoverPage')
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('Prestacao de Servicos Especializados em Resiliencia\n'
                     'de Supply Chain e Otimizacao Operacional Integrada')
    run.font.name = 'Garamond'; run.font.size = Pt(12); run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Separator line
    p = doc.add_paragraph(style='CoverPage')
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run('\u2501' * 50)
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0xBB, 0xAA, 0x88)

    # Company details
    details = [
        ("Fundador:", "Everton Rodrigo Vieira"),
        ("Estrutura Juridica:", "S-Corporation"),
        ("Sede:", "Jacksonville, FL \u2014 EUA"),
        ("NAICS:", "541611 \u2014 Administrative Management and General Management Consulting"),
        ("Documento:", "Confidencial \u2014 Distribui\u00e7\u00e3o restrita"),
    ]
    for label, value in details:
        p = doc.add_paragraph(style='CoverPage')
        r1 = p.add_run(label + ' ')
        r1.font.name = 'Garamond'; r1.font.size = Pt(11); r1.bold = True
        r2 = p.add_run(value)
        r2.font.name = 'Garamond'; r2.font.size = Pt(11)

    # Confidentiality notice
    for _ in range(3):
        doc.add_paragraph(style='CoverPage')

    p = doc.add_paragraph(style='CoverPage')
    run = p.add_run('This business plan contains proprietary information. '
                     'Unauthorized distribution is prohibited.')
    run.font.name = 'Garamond'; run.font.size = Pt(9); run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    page_break(doc)


# ============================================================
# TABLE OF CONTENTS
# ============================================================

def build_toc(doc):
    """Add Table of Contents heading + marker for static TOC."""
    doc.add_heading('SUM\u00c1RIO', level=1)
    # Marker paragraph — will be replaced by populate_toc
    marker = doc.add_paragraph()
    marker.text = '__TOC_MARKER__'
    marker.style = doc.styles['Normal']
    page_break(doc)


def populate_toc(doc):
    """Build a TOC field with static entries as fallback content.

    The TOC field code tells Word to auto-generate the TOC when the document
    is opened. The static entries serve as visible fallback for viewers that
    don't update fields (e.g., python-docx output viewed without Word).

    When Word opens the document (e.g., during PDF conversion), it replaces
    the fallback entries with the real TOC including page numbers.
    """
    # Collect headings (skip SUMÁRIO)
    headings = []
    for para in doc.paragraphs:
        style_name = para.style.name
        if style_name in ('Heading 1', 'Heading 2', 'Heading 3'):
            text = para.text.strip()
            if not text or text == 'SUM\u00c1RIO' or text == 'SUMARIO':
                continue
            level = int(style_name[-1])
            headings.append((level, text))

    # Find marker paragraph
    marker_elem = None
    for para in doc.paragraphs:
        if para.text.strip() == '__TOC_MARKER__':
            marker_elem = para._element
            break
    if marker_elem is None:
        print("Warning: TOC marker not found.")
        return

    body = doc.element.body
    marker_idx = list(body).index(marker_elem)

    # Build a single paragraph with TOC field begin + instrText + separate
    toc_begin_xml = (
        f'<w:p {nsdecls("w")}>'
        '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:instrText xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u '
        '</w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        '</w:p>'
    )
    body.insert(marker_idx, parse_xml(toc_begin_xml))
    marker_idx += 1

    # Build static entries as fallback content (between separate and end)
    for level, text in headings:
        indent_left = (level - 1) * 360
        safe_text = (text.replace('&', '&amp;').replace('<', '&lt;')
                         .replace('>', '&gt;').replace('"', '&quot;'))

        font_props = '<w:rFonts w:ascii="Garamond" w:hAnsi="Garamond"/>'
        if level == 1:
            font_props += '<w:b/><w:sz w:val="22"/>'
        elif level == 2:
            font_props += '<w:sz w:val="21"/>'
        else:
            font_props += '<w:sz w:val="20"/>'

        p_xml = (
            f'<w:p {nsdecls("w")}>'
            '<w:pPr>'
            f'<w:ind w:left="{indent_left}"/>'
            '<w:tabs><w:tab w:val="right" w:leader="dot" w:pos="9360"/></w:tabs>'
            '<w:spacing w:after="40" w:before="20"/>'
            f'<w:rPr>{font_props}</w:rPr>'
            '</w:pPr>'
            f'<w:r><w:rPr>{font_props}</w:rPr>'
            f'<w:t xml:space="preserve">{safe_text}</w:t></w:r>'
            '</w:p>'
        )
        body.insert(marker_idx, parse_xml(p_xml))
        marker_idx += 1

    # TOC field end paragraph
    toc_end_xml = (
        f'<w:p {nsdecls("w")}>'
        '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
        '</w:p>'
    )
    body.insert(marker_idx, parse_xml(toc_end_xml))

    # Remove marker
    body.remove(marker_elem)


# ============================================================
# MAIN BUILDER
# ============================================================

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    build_cover(doc)
    build_toc(doc)

    # Import and build each section
    from bp_sections import (
        build_section_1,
        build_section_2,
        build_section_3,
        build_section_4,
        build_section_5,
        build_section_6,
    )

    build_section_1(doc)
    build_section_2(doc)
    build_section_3(doc)
    build_section_4(doc)
    build_section_5(doc)
    build_section_6(doc)

    # Populate static TOC entries from headings
    populate_toc(doc)

    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")

    # Quick stats
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")

    # Update TOC via Word and export PDF
    OUTPUT_PDF = OUTPUT_DOCX.replace('.docx', '.pdf')
    update_toc_and_export_pdf(OUTPUT_DOCX, OUTPUT_PDF)


def update_toc_and_export_pdf(docx_path, pdf_path):
    """Use Microsoft Word via AppleScript to update TOC fields, then docx2pdf for PDF."""
    abs_docx = os.path.abspath(docx_path)

    # Step 1: AppleScript to open in Word, update TOC, save, close
    script = f'''
    tell application "Microsoft Word"
        activate
        delay 3
        open "{abs_docx}"
        delay 15
        set d to active document
        set tc to table of contents 1 of d
        update tc
        delay 5
        save d
        delay 2
        close d saving no
    end tell
    '''
    try:
        result = subprocess.run(
            ['osascript', '-e', script],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0:
            print("TOC updated via Word AppleScript.")
        else:
            print(f"Word TOC update warning: {result.stderr}")
    except Exception as e:
        print(f"Word TOC update failed: {e}")
        print("Note: Open docx in Word, press Ctrl+A then F9 to update TOC manually.")

    # Step 2: Convert to PDF via docx2pdf
    _fallback_pdf(docx_path, pdf_path)


def _fallback_pdf(docx_path, pdf_path):
    """Fallback: convert via docx2pdf without TOC update."""
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"PDF (without TOC update): {pdf_path}")
        print("Note: Open docx in Word, press Ctrl+A then F9 to update TOC.")
    except ImportError:
        print("docx2pdf not installed. Install with: pip install docx2pdf")


if __name__ == '__main__':
    main()
