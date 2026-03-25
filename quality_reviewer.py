#!/usr/bin/env python3
"""
QUALITY REVIEWER — Verificação Automatizada de Résumé V4
Execute após cada build/merge. Lê o .docx final e reporta problemas.

USO:
    python3 quality_reviewer.py /path/to/resume.docx [eb1a|eb2niw]
"""
import sys
import os
import re
from collections import OrderedDict
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# Cores permitidas (hex sem #)
ALLOWED_COLORS = {
    "2D3E50",  # NAVY
    "3498A2",  # TEAL
    "FFFFFF",  # WHITE
    "000000",  # BLACK
    "333333",  # DARK_GRAY
    "666666",  # MED_GRAY
    "F5F5F5",  # LIGHT_GRAY
    "CCCCCC",  # BORDER_GRAY
    "FAFAFA",  # alternate row (close to F5F5F5)
}

ALLOWED_FONTS = {"Garamond", "garamond"}
FORBIDDEN_WORDS = ["R$"]


def load_document(path):
    doc = Document(path)
    para_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + "\n"
    merged = para_text + "\n" + table_text
    return doc, merged


def check_fonts(doc):
    issues = []
    arial_count = 0
    other_fonts = set()
    total_runs = 0

    def check_runs(paragraphs):
        nonlocal arial_count, total_runs
        for para in paragraphs:
            for run in para.runs:
                total_runs += 1
                fname = run.font.name
                if fname and fname.lower() not in ALLOWED_FONTS and fname.lower() != "garamond":
                    if "arial" in fname.lower():
                        arial_count += 1
                    other_fonts.add(fname)

    check_runs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                check_runs(cell.paragraphs)

    if arial_count > 0:
        issues.append(("S0", f"ARIAL ENCONTRADO: {arial_count} runs com Arial (ZERO TOLERÂNCIA)"))
    if other_fonts - {"Garamond", "garamond", None}:
        cleaned = {f for f in other_fonts if f and f.lower() != "garamond"}
        if cleaned:
            issues.append(("S1", f"Fontes não-Garamond encontradas: {cleaned}"))

    return issues, total_runs


def check_page_setup(doc):
    issues = []
    section = doc.sections[0]
    w = round(section.page_width.inches, 1)
    h = round(section.page_height.inches, 1)
    if w != 8.5 or h != 11.0:
        issues.append(("S1", f"Paper size ERRADO: {w}\"x{h}\" (esperado: 8.5\"x11\")"))

    lm = round(section.left_margin.inches, 2)
    rm = round(section.right_margin.inches, 2)
    if abs(lm - 0.65) > 0.03:
        issues.append(("S1", f"Margem esquerda ERRADA: {lm}\" (esperado: 0.65\")"))
    if abs(rm - 0.65) > 0.03:
        issues.append(("S1", f"Margem direita ERRADA: {rm}\" (esperado: 0.65\")"))

    return issues


def check_forbidden_words(merged_text, xml_str):
    issues = []
    for word in FORBIDDEN_WORDS:
        if word in merged_text or word in xml_str:
            issues.append(("S1", f"Palavra proibida encontrada: '{word}'"))
    return issues


def check_images(doc):
    issues = []
    img_count = len(doc.inline_shapes)
    placeholder_count = 0
    for para in doc.paragraphs:
        if "[THUMBNAIL]" in para.text or "[THUMBNAIL" in para.text:
            placeholder_count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "[THUMBNAIL]" in p.text or "[THUMBNAIL" in p.text:
                        placeholder_count += 1

    if img_count == 0 and placeholder_count > 0:
        issues.append(("S2", f"{placeholder_count} thumbnails com placeholder [THUMBNAIL] (imagem faltando)"))
    elif img_count == 0 and placeholder_count == 0:
        issues.append(("S2", "Zero imagens e zero placeholders (considerar adicionar thumbnails)"))

    return issues, img_count, placeholder_count


def check_evidence_blocks(doc):
    issues = []
    tables_with_impact = 0
    tables_2col = 0

    for table in doc.tables:
        if len(table.columns) == 2 and len(table.rows) == 1:
            tables_2col += 1
            left_text = table.rows[0].cells[0].text
            if "Impact" in left_text or "Impacto" in left_text or "Description" in left_text:
                tables_with_impact += 1

    if tables_2col > 0 and tables_with_impact == 0:
        issues.append(("S1", f"NENHUM evidence block tem 'Impact' dentro ({tables_2col} tables 2-col)"))

    return issues, tables_2col, tables_with_impact


def check_paragraph_length(doc, min_lines=3):
    issues = []
    short_paras = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or len(text) < 20:
            continue
        if text.startswith(("\u25cf", "\u25cb", "Page ", "Description", "\u2014")):
            continue
        est_lines = max(1, len(text) / 85)
        if est_lines < 2.5 and len(text) > 50:
            short_paras.append((i, len(text), text[:60]))

    if len(short_paras) > 5:
        issues.append(("S2", f"{len(short_paras)} parágrafos potencialmente curtos (< 3 linhas)"))

    return issues


def check_header_footer(doc):
    issues = []
    section = doc.sections[0]
    header = section.header
    if not header.paragraphs and not header.tables:
        issues.append(("S1", "HEADER AUSENTE"))
    elif len(header.tables) == 0:
        issues.append(("S2", "Header sem tabela (pode não ter navy bar)"))

    footer = section.footer
    if not footer.paragraphs and not footer.tables:
        issues.append(("S1", "FOOTER AUSENTE"))

    return issues


def check_sections_eb2niw(merged_text):
    issues = []
    required = [
        ("Summary / Síntese", ["SUMMARY", "SÍNTESE PROFISSIONAL"]),
        ("Career Timeline / Histórico", ["CAREER TIMELINE", "HISTÓRICO PROFISSIONAL", "PROFESSIONAL CAREER"]),
        ("Proposed Endeavors", ["PROPOSED ENDEAVOR", "PROJETO EB-2 NIW"]),
        ("Contribuições Técnicas", ["CONTRIBUI", "ORIGINAL CONTRIBUTION", "TECHNICAL CONTRIBUTION"]),
    ]
    upper = merged_text.upper()
    for label, needles in required:
        found = any(n in upper for n in needles)
        if not found:
            issues.append(("S1", f"Seção ausente: {label}"))

    if "DHANASAR" not in upper:
        # Check for Dhanasar-like references (the V2.0 system requires this)
        pass  # Removed per FORBIDDEN_CONTENT v1.0 guidance

    bls_pattern = r'\d{2}-\d{4}'
    if not re.search(bls_pattern, merged_text):
        issues.append(("S2", "Nenhum código BLS/O*Net encontrado (formato XX-XXXX)"))

    if "CRITERION 1" in upper or "CRITERION 2" in upper:
        issues.append(("S1", "Seções por Critério (C1-C10) encontradas — isso é EB-1A, não EB-2 NIW"))

    return issues


def check_colors_in_xml(doc):
    issues = []
    unexpected_colors = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill and fill.upper() not in ALLOWED_COLORS and fill != "auto":
                            unexpected_colors.add(fill)
    if unexpected_colors:
        issues.append(("S2", f"Cores inesperadas em cell shading: {unexpected_colors}"))
    return issues


def check_consistency(merged_text):
    issues = []
    upper = merged_text.upper()
    has_pt = "SÍNTESE" in upper or "CONTRIBUIÇÕES" in upper or "FORMAÇÃO" in upper
    has_en = "EXECUTIVE SUMMARY" in upper or "ACADEMIC BACKGROUND" in upper
    if has_pt and has_en:
        issues.append(("S3", "Headers misturando idiomas (PT e EN)"))
    return issues


def run_review(docx_path, doc_type="eb2niw"):
    print("=" * 70)
    print(f"QUALITY REVIEWER — {os.path.basename(docx_path)}")
    print(f"Tipo: {doc_type.upper()}")
    print(f"Tamanho: {os.path.getsize(docx_path):,} bytes")
    print("=" * 70)

    doc, merged = load_document(docx_path)
    xml_str = etree.tostring(doc.element.body, encoding='unicode')

    all_issues = []

    print("\n[1/9] Verificando fontes...")
    font_issues, total_runs = check_fonts(doc)
    all_issues.extend(font_issues)
    print(f"      {total_runs} runs verificados")

    print("[2/9] Verificando page setup...")
    all_issues.extend(check_page_setup(doc))

    print("[3/9] Verificando palavras proibidas...")
    all_issues.extend(check_forbidden_words(merged, xml_str))

    print("[4/9] Verificando imagens...")
    img_issues, img_count, placeholder_count = check_images(doc)
    all_issues.extend(img_issues)
    print(f"      {img_count} imagens, {placeholder_count} placeholders")

    print("[5/9] Verificando evidence blocks...")
    eb_issues, tables_2col, tables_impact = check_evidence_blocks(doc)
    all_issues.extend(eb_issues)
    print(f"      {tables_2col} evidence blocks, {tables_impact} com impact dentro")

    print("[6/9] Verificando comprimento de parágrafos...")
    all_issues.extend(check_paragraph_length(doc))

    print("[7/9] Verificando header e footer...")
    all_issues.extend(check_header_footer(doc))

    print("[8/9] Verificando seções obrigatórias...")
    all_issues.extend(check_sections_eb2niw(merged))

    print("[9/9] Verificando cores e consistência...")
    all_issues.extend(check_colors_in_xml(doc))
    all_issues.extend(check_consistency(merged))

    # Report
    print("\n" + "=" * 70)
    print("RELATÓRIO DE QUALIDADE")
    print("=" * 70)

    print(f"\nParagraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Images: {img_count}")
    print(f"Evidence Blocks (2-col): {tables_2col}")
    print(f"  com impact dentro: {tables_impact}")
    print(f"Runs verificados: {total_runs}")

    s0 = [(s, m) for s, m in all_issues if s == "S0"]
    s1 = [(s, m) for s, m in all_issues if s == "S1"]
    s2 = [(s, m) for s, m in all_issues if s == "S2"]
    s3 = [(s, m) for s, m in all_issues if s == "S3"]

    print(f"\n--- ISSUES ---")
    print(f"S0 GRAVÍSSIMO: {len(s0)}")
    print(f"S1 GRAVE:      {len(s1)}")
    print(f"S2 MODERADO:   {len(s2)}")
    print(f"S3 MENOR:      {len(s3)}")

    if s0:
        print(f"\n!!! S0 — GRAVÍSSIMO (BLOQUEIA ENTREGA) !!!")
        for _, msg in s0:
            print(f"  [S0] {msg}")
    if s1:
        print(f"\n!! S1 — GRAVE (BLOQUEIA ENTREGA) !!")
        for _, msg in s1:
            print(f"  [S1] {msg}")
    if s2:
        print(f"\n! S2 — MODERADO (CORRIGIR) !")
        for _, msg in s2:
            print(f"  [S2] {msg}")
    if s3:
        print(f"\nS3 — MENOR (DOCUMENTAR)")
        for _, msg in s3:
            print(f"  [S3] {msg}")

    print(f"\n{'=' * 70}")
    if s0 or s1:
        print("VEREDICTO: REPROVADO — Corrigir issues S0/S1 antes de entregar")
        print(f"  {len(s0)} gravíssimos + {len(s1)} graves = REBUILD NECESSÁRIO")
    elif s2:
        print("VEREDICTO: APROVADO COM RESSALVAS — Corrigir issues S2 se possível")
    else:
        print("VEREDICTO: APROVADO — Documento pronto para entrega")
    print("=" * 70)

    return len(s0), len(s1), len(s2), len(s3)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python3 quality_reviewer.py <resume.docx> [eb1a|eb2niw]")
        sys.exit(1)

    docx_path = sys.argv[1]
    doc_type = sys.argv[2] if len(sys.argv) > 2 else "eb2niw"

    s0, s1, s2, s3 = run_review(docx_path, doc_type)
    sys.exit(1 if (s0 + s1) > 0 else 0)
